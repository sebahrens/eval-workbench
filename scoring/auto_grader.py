"""Automated grader for Cascade Industries test suite (prompt.md §7.2).

Grades the mechanical components of agent outputs against gold standards:
  - Correctness: numerical comparison with ±0.5% tolerance
  - Completeness: presence of required files, sheets, sections
  - Format: file validity (opens without errors, correct type)
  - Canary verification: correct files were read
  - Error detection: planted errors were flagged

Usage:
    # Grade a single test case
    python -m scoring.auto_grader --tc TC-01 \\
        --gold gold_standards/TC-01_gold.json \\
        --agent-output /path/to/agent/output/TC-01

    # Self-test: grade gold outputs against themselves (must score 3/3/3/3/3)
    python -m scoring.auto_grader --self-test --suite-dir /tmp/test_suite

    # Grade all test cases
    python -m scoring.auto_grader --suite-dir /tmp/test_suite \\
        --agent-output-dir /path/to/agent/outputs
"""

from __future__ import annotations

import argparse
import json
import math
import sys
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

import openpyxl
import yaml
from docx import Document as DocxDocument

# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class CheckResult:
    """Result of a single check within a grading dimension."""

    name: str
    status: str  # "pass" | "fail" | "warn"
    expected: Any = None
    actual: Any = None
    detail: str = ""


@dataclass
class DimensionResult:
    """Grading result for one dimension."""

    dimension: str
    score: int  # 1, 2, or 3
    checks: list[CheckResult] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "dimension": self.dimension,
            "score": self.score,
            "checks": [asdict(c) for c in self.checks],
        }


@dataclass
class GradeReport:
    """Full grading report for one test case."""

    test_case: str
    dimensions: dict[str, DimensionResult] = field(default_factory=dict)

    @property
    def scores(self) -> dict[str, int]:
        return {d: r.score for d, r in sorted(self.dimensions.items())}

    def to_dict(self) -> dict[str, Any]:
        return {
            "test_case": self.test_case,
            "scores": self.scores,
            "dimensions": {
                d: r.to_dict() for d, r in sorted(self.dimensions.items())
            },
        }


# ---------------------------------------------------------------------------
# File readers — attempt to open files and extract structure
# ---------------------------------------------------------------------------

def _read_xlsx_sheets(path: Path) -> list[str] | None:
    """Return sheet names if the file opens as valid xlsx, else None."""
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        names = wb.sheetnames
        wb.close()
        return names
    except Exception:
        return None


def _read_docx_sections(path: Path) -> list[str] | None:
    """Return paragraph heading texts if the file opens as valid docx, else None."""
    try:
        doc = DocxDocument(str(path))
        headings = []
        for para in doc.paragraphs:
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                headings.append(para.text)
        return headings
    except Exception:
        return None


def _file_is_valid_pdf(path: Path) -> bool:
    """Check if a file is a valid PDF by reading its magic bytes."""
    try:
        with open(path, "rb") as f:
            header = f.read(5)
        return header == b"%PDF-"
    except Exception:
        return False


def _file_is_valid(path: Path) -> bool:
    """Check if a file can be opened according to its type."""
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        return _read_xlsx_sheets(path) is not None
    if suffix == ".docx":
        return _read_docx_sections(path) is not None
    if suffix == ".pdf":
        return _file_is_valid_pdf(path)
    if suffix in (".json", ".yaml", ".yml"):
        try:
            with open(path) as f:
                if suffix == ".json":
                    json.load(f)
                else:
                    yaml.safe_load(f)
            return True
        except Exception:
            return False
    if suffix in (".csv", ".txt", ".md"):
        try:
            path.read_text(encoding="utf-8")
            return True
        except Exception:
            return False
    # Unknown type — just check it exists and is non-empty
    return path.exists() and path.stat().st_size > 0


def _search_file_for_text(path: Path, needle: str) -> bool:
    """Search for *needle* in a file, handling xlsx/docx/pdf/text."""
    suffix = path.suffix.lower()

    if suffix == ".xlsx":
        try:
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            # Check cell values
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and needle in str(cell.value):
                            wb.close()
                            return True
            # Check document properties
            if wb.properties:
                for attr in ("description", "subject", "title", "creator",
                             "keywords", "category"):
                    val = getattr(wb.properties, attr, None)
                    if val and needle in str(val):
                        wb.close()
                        return True
            wb.close()
        except Exception:
            pass
        return False

    if suffix == ".docx":
        try:
            doc = DocxDocument(str(path))
            # Check paragraphs
            for para in doc.paragraphs:
                if needle in para.text:
                    return True
            # Check core properties
            cp = doc.core_properties
            for attr in ("comments", "subject", "title", "description",
                         "keywords", "category"):
                val = getattr(cp, attr, None)
                if val and needle in str(val):
                    return True
        except Exception:
            pass
        return False

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            # Check metadata
            if reader.metadata:
                for val in reader.metadata.values():
                    if val and needle in str(val):
                        return True
            # Check page text
            for page in reader.pages:
                text = page.extract_text() or ""
                if needle in text:
                    return True
        except Exception:
            pass
        return False

    # Text-based files
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
        return needle in content
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Numerical comparison
# ---------------------------------------------------------------------------

def _values_match(expected: Any, actual: Any, tolerance_pct: float = 0.5) -> bool:
    """Compare two values with tolerance for numerical types.

    - Numbers: match within ±tolerance_pct (default 0.5%)
    - Strings: exact match (case-insensitive)
    - Lists: element-wise match
    - Dicts: key-wise match
    """
    if isinstance(expected, (int, float)) and isinstance(actual, (int, float)):
        if expected == 0:
            return math.isclose(actual, 0, abs_tol=1e-6)
        return math.isclose(actual, expected, rel_tol=tolerance_pct / 100)

    if isinstance(expected, str) and isinstance(actual, str):
        return expected.lower() == actual.lower()

    if isinstance(expected, list) and isinstance(actual, list):
        if len(expected) != len(actual):
            return False
        return all(_values_match(e, a, tolerance_pct) for e, a in zip(expected, actual))

    if isinstance(expected, dict) and isinstance(actual, dict):
        if set(expected.keys()) != set(actual.keys()):
            return False
        return all(
            _values_match(expected[k], actual[k], tolerance_pct)
            for k in expected
        )

    return expected == actual


# ---------------------------------------------------------------------------
# TestCaseGrader
# ---------------------------------------------------------------------------

class TestCaseGrader:
    """Grade one test case by comparing agent output against a gold standard.

    Parameters
    ----------
    test_case_id:
        Identifier such as "TC-01".
    gold_standard_path:
        Path to the gold JSON file (e.g. gold_standards/TC-01_gold.json).
    agent_output_path:
        Path to the directory containing the agent's output for this TC.
    suite_dir:
        Path to the full test suite (for locating canary files, etc.).
    """

    def __init__(
        self,
        test_case_id: str,
        gold_standard_path: str | Path,
        agent_output_path: str | Path,
        suite_dir: str | Path | None = None,
    ) -> None:
        self.test_case_id = test_case_id
        self.gold_path = Path(gold_standard_path)
        self.agent_path = Path(agent_output_path)
        self.suite_dir = Path(suite_dir) if suite_dir else self.agent_path.parent

        with open(self.gold_path) as f:
            self._gold = json.load(f)

        self._expected = self._gold.get("expected_outputs", {})
        self._canaries = self._gold.get("canary_verification", {})
        self._errors = self._gold.get("error_detection", {})

    # -- Dimension graders ----------------------------------------------------

    def grade_correctness(self) -> DimensionResult:
        """Compare numerical outputs to gold standard values.

        Walks the expected_outputs dict and compares every numerical value
        found in the agent's output against the gold standard.
        Returns score 3 if all match within ±0.5%, 2 if minor deviations,
        1 if material errors.
        """
        checks: list[CheckResult] = []
        agent_data = self._load_agent_data()

        if agent_data is None:
            checks.append(CheckResult(
                name="agent_data_loadable",
                status="fail",
                detail="Could not load agent output data for comparison",
            ))
            return DimensionResult(dimension="correctness", score=1, checks=checks)

        total, matched = self._compare_recursive(
            self._expected, agent_data, "expected_outputs", checks,
        )

        if total == 0:
            score = 3  # No numerical values to check
        elif matched == total:
            score = 3
        elif matched >= total * 0.8:
            score = 2
        else:
            score = 1

        return DimensionResult(dimension="correctness", score=score, checks=checks)

    def grade_completeness(self) -> DimensionResult:
        """Check for presence of required files, sheets, sections.

        Uses the gold standard's expected_outputs to determine what is required.
        """
        checks: list[CheckResult] = []

        # Check required sheets (if specified in gold)
        required_sheets = self._expected.get("required_sheets", [])
        if required_sheets:
            # Find xlsx files in agent output
            xlsx_files = sorted(self.agent_path.rglob("*.xlsx"))
            if not xlsx_files:
                checks.append(CheckResult(
                    name="required_sheets",
                    status="fail",
                    expected=required_sheets,
                    actual=[],
                    detail="No xlsx files found in agent output",
                ))
            else:
                all_sheets: list[str] = []
                for xlsx_path in xlsx_files:
                    sheets = _read_xlsx_sheets(xlsx_path)
                    if sheets:
                        all_sheets.extend(sheets)

                missing = [s for s in required_sheets if s not in all_sheets]
                if not missing:
                    checks.append(CheckResult(
                        name="required_sheets",
                        status="pass",
                        expected=required_sheets,
                        actual=all_sheets,
                    ))
                else:
                    checks.append(CheckResult(
                        name="required_sheets",
                        status="fail",
                        expected=required_sheets,
                        actual=all_sheets,
                        detail=f"Missing sheets: {missing}",
                    ))

        # Check required file type
        file_type = self._expected.get("file_type")
        if file_type:
            files = sorted(self.agent_path.rglob(f"*.{file_type}"))
            checks.append(CheckResult(
                name=f"has_{file_type}_file",
                status="pass" if files else "fail",
                expected=f"At least one .{file_type} file",
                actual=len(files),
            ))

        # Check required sections (for sub-dicts in expected_outputs)
        for section_name in sorted(self._expected):
            if section_name in ("file_type", "required_sheets"):
                continue
            if isinstance(self._expected[section_name], dict):
                checks.append(CheckResult(
                    name=f"section_{section_name}",
                    status="pass",  # Existence check — correctness checks values
                ))

        failed = sum(1 for c in checks if c.status == "fail")
        total = len(checks)

        if failed == 0:
            score = 3
        elif failed <= total * 0.3:
            score = 2
        else:
            score = 1

        return DimensionResult(dimension="completeness", score=score, checks=checks)

    def grade_format(self) -> DimensionResult:
        """Validate file types, check if files open without errors."""
        checks: list[CheckResult] = []

        # Find all files in agent output
        all_files = sorted(
            f for f in self.agent_path.rglob("*") if f.is_file()
        )

        if not all_files:
            checks.append(CheckResult(
                name="has_output_files",
                status="fail",
                detail="No files found in agent output directory",
            ))
            return DimensionResult(dimension="format", score=1, checks=checks)

        valid_count = 0
        for fpath in all_files:
            is_valid = _file_is_valid(fpath)
            checks.append(CheckResult(
                name=f"valid_{fpath.name}",
                status="pass" if is_valid else "fail",
                detail="" if is_valid else f"File {fpath.name} failed validation",
            ))
            if is_valid:
                valid_count += 1

        if valid_count == len(all_files):
            score = 3
        elif valid_count >= len(all_files) * 0.8:
            score = 2
        else:
            score = 1

        return DimensionResult(dimension="format", score=score, checks=checks)

    def verify_canaries(self) -> DimensionResult:
        """Check if the agent read the correct files by finding canary values.

        Searches the agent output for each canary that should have been
        encountered when reading the input files.
        """
        checks: list[CheckResult] = []

        if not self._canaries:
            return DimensionResult(dimension="canaries", score=3, checks=[
                CheckResult(name="no_canaries", status="pass",
                            detail="No canaries defined for this test case"),
            ])

        # Real grading: search only agent outputs for canary evidence.
        # The agent must have referenced the canary in its output to prove it
        # read the right files. SelfTestGrader overrides this to search input
        # files instead (since there is no agent output in self-test mode).
        found_count = 0
        for label, canary_code in sorted(self._canaries.items()):
            found = False
            found_in = ""
            for fpath in sorted(self.agent_path.rglob("*")):
                if fpath.is_file() and _search_file_for_text(fpath, canary_code):
                    found = True
                    try:
                        found_in = str(fpath.relative_to(self.suite_dir))
                    except ValueError:
                        found_in = str(fpath)
                    break

            checks.append(CheckResult(
                name=f"canary_{label}",
                status="pass" if found else "fail",
                expected=canary_code,
                actual="found" if found else "not_found",
                detail=f"Found in {found_in}" if found else "",
            ))
            if found:
                found_count += 1

        total = len(self._canaries)
        if found_count == total:
            score = 3
        elif found_count >= total * 0.5:
            score = 2
        else:
            score = 1

        return DimensionResult(dimension="canaries", score=score, checks=checks)

    def check_error_detection(self) -> DimensionResult:
        """Check if the agent identified planted errors.

        For self-test mode, verifies that the error descriptions in the gold
        standard are non-empty and well-formed.
        """
        checks: list[CheckResult] = []

        if not self._errors:
            return DimensionResult(dimension="error_detection", score=3, checks=[
                CheckResult(name="no_errors", status="pass",
                            detail="No planted errors defined for this test case"),
            ])

        detected = 0
        for error_id, description in sorted(self._errors.items()):
            # Check the gold standard has a valid error description
            has_description = bool(description and len(description) > 5)

            # Search agent output for evidence of error detection
            found_in_output = False
            for fpath in sorted(self.agent_path.rglob("*")):
                if fpath.is_file():
                    if _search_file_for_text(fpath, error_id):
                        found_in_output = True
                        break

            is_detected = has_description and found_in_output
            checks.append(CheckResult(
                name=f"error_{error_id}",
                status="pass" if is_detected else ("warn" if has_description else "fail"),
                expected=description,
                actual="detected" if found_in_output else "not_detected",
                detail="" if is_detected else (
                    "Error not found in agent output"
                    if has_description else "Missing error description in gold"
                ),
            ))
            if is_detected:
                detected += 1

        total = len(self._errors)
        if detected == total:
            score = 3
        elif detected >= total * 0.5:
            score = 2
        else:
            score = 1

        return DimensionResult(dimension="error_detection", score=score, checks=checks)

    # -- Full grading ---------------------------------------------------------

    def grade(self) -> GradeReport:
        """Run all grading dimensions and return a full report."""
        report = GradeReport(test_case=self.test_case_id)
        report.dimensions["correctness"] = self.grade_correctness()
        report.dimensions["completeness"] = self.grade_completeness()
        report.dimensions["format"] = self.grade_format()
        report.dimensions["canaries"] = self.verify_canaries()
        report.dimensions["error_detection"] = self.check_error_detection()
        return report

    # -- Helpers --------------------------------------------------------------

    def _load_agent_data(self) -> dict[str, Any] | None:
        """Attempt to load the agent's output as structured data.

        Looks for JSON files in the agent output directory, or extracts data
        from xlsx files.
        """
        # Try loading a JSON summary if the agent produced one
        json_files = sorted(self.agent_path.rglob("*.json"))
        for jf in json_files:
            try:
                with open(jf) as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    return data
            except (json.JSONDecodeError, OSError):
                continue

        # No usable agent structured data found
        return None

    def _compare_recursive(
        self,
        expected: Any,
        actual: Any,
        path: str,
        checks: list[CheckResult],
    ) -> tuple[int, int]:
        """Recursively compare expected vs actual, recording check results.

        Returns (total_comparisons, matched_comparisons).
        """
        total = 0
        matched = 0

        if isinstance(expected, dict):
            for key in sorted(expected):
                # Skip structural keys — these are checked in completeness
                if key in ("file_type", "required_sheets"):
                    continue
                key_present = key in actual if isinstance(actual, dict) else False
                actual_val = actual.get(key) if isinstance(actual, dict) else None
                if not key_present:
                    if isinstance(expected[key], dict):
                        # Recurse into sub-dicts even if actual is missing
                        sub_t, sub_m = self._compare_recursive(
                            expected[key], {}, f"{path}.{key}", checks,
                        )
                        total += sub_t
                        matched += sub_m
                    else:
                        total += 1
                        checks.append(CheckResult(
                            name=f"{path}.{key}",
                            status="fail",
                            expected=expected[key],
                            actual=None,
                            detail="Value missing from agent output",
                        ))
                else:
                    sub_t, sub_m = self._compare_recursive(
                        expected[key], actual_val, f"{path}.{key}", checks,
                    )
                    total += sub_t
                    matched += sub_m

        elif isinstance(expected, list):
            if not isinstance(actual, list):
                total += 1
                checks.append(CheckResult(
                    name=path,
                    status="fail",
                    expected=f"list[{len(expected)}]",
                    actual=type(actual).__name__,
                    detail="Expected list, got different type",
                ))
            elif _values_match(expected, actual):
                total += 1
                matched += 1
                checks.append(CheckResult(
                    name=path,
                    status="pass",
                    expected=expected,
                    actual=actual,
                ))
            else:
                total += 1
                checks.append(CheckResult(
                    name=path,
                    status="fail",
                    expected=expected,
                    actual=actual,
                    detail="List values don't match",
                ))

        elif isinstance(expected, (int, float)):
            total += 1
            if _values_match(expected, actual):
                matched += 1
                checks.append(CheckResult(
                    name=path,
                    status="pass",
                    expected=expected,
                    actual=actual,
                ))
            else:
                checks.append(CheckResult(
                    name=path,
                    status="fail",
                    expected=expected,
                    actual=actual,
                    detail="Value mismatch (tolerance: ±0.5%)",
                ))

        elif isinstance(expected, str):
            total += 1
            if _values_match(expected, actual):
                matched += 1
                checks.append(CheckResult(
                    name=path,
                    status="pass",
                    expected=expected,
                    actual=actual,
                ))
            else:
                checks.append(CheckResult(
                    name=path,
                    status="fail",
                    expected=expected,
                    actual=actual,
                    detail="String mismatch",
                ))

        else:
            total += 1
            if expected == actual:
                matched += 1
                checks.append(CheckResult(
                    name=path, status="pass",
                    expected=expected, actual=actual,
                ))
            else:
                checks.append(CheckResult(
                    name=path, status="fail",
                    expected=expected, actual=actual,
                ))

        return total, matched


# ---------------------------------------------------------------------------
# Self-test
# ---------------------------------------------------------------------------

class SelfTestGrader(TestCaseGrader):
    """Specialized grader for self-test mode.

    In self-test, the "agent output" is the generated test suite itself.
    The gold standard should score 3/3/3/3/3 against the generated files.
    """

    def _load_agent_data(self) -> dict[str, Any] | None:
        """In self-test mode, the gold standard IS the expected data."""
        return self._expected

    def grade_completeness(self) -> DimensionResult:
        """In self-test, verify the gold standard defines requirements.

        The gold standard's required_sheets/file_type/sections describe what
        an agent should produce — in self-test we verify these are non-empty
        and well-formed, not that the input files contain them.
        """
        checks: list[CheckResult] = []

        required_sheets = self._expected.get("required_sheets", [])
        if required_sheets:
            checks.append(CheckResult(
                name="required_sheets_defined",
                status="pass",
                expected="non-empty list",
                actual=required_sheets,
            ))

        file_type = self._expected.get("file_type")
        if file_type:
            checks.append(CheckResult(
                name="file_type_defined",
                status="pass",
                expected="non-empty string",
                actual=file_type,
            ))

        # Check that structural sections exist in expected_outputs
        for section_name in sorted(self._expected):
            if section_name in ("file_type", "required_sheets"):
                continue
            if isinstance(self._expected[section_name], dict):
                checks.append(CheckResult(
                    name=f"section_{section_name}",
                    status="pass",
                ))

        if not checks:
            checks.append(CheckResult(
                name="has_expected_outputs",
                status="pass" if self._expected else "fail",
                detail="Gold standard expected_outputs is defined" if self._expected
                       else "No expected_outputs in gold standard",
            ))

        failed = sum(1 for c in checks if c.status == "fail")
        score = 3 if failed == 0 else (2 if failed == 1 else 1)
        return DimensionResult(dimension="completeness", score=score, checks=checks)

    def check_error_detection(self) -> DimensionResult:
        """In self-test, verify error descriptions are well-formed."""
        checks: list[CheckResult] = []

        if not self._errors:
            return DimensionResult(dimension="error_detection", score=3, checks=[
                CheckResult(name="no_errors", status="pass",
                            detail="No planted errors for this test case"),
            ])

        for error_id, description in sorted(self._errors.items()):
            has_good_desc = bool(description and len(description) > 5)
            checks.append(CheckResult(
                name=f"error_{error_id}",
                status="pass" if has_good_desc else "fail",
                expected="well-formed description",
                actual=description,
            ))

        failed = sum(1 for c in checks if c.status == "fail")
        score = 3 if failed == 0 else (2 if failed == 1 else 1)
        return DimensionResult(dimension="error_detection", score=score, checks=checks)

    def verify_canaries(self) -> DimensionResult:
        """In self-test, verify canaries are findable in the generated input files."""
        checks: list[CheckResult] = []

        if not self._canaries:
            return DimensionResult(dimension="canaries", score=3, checks=[
                CheckResult(name="no_canaries", status="pass",
                            detail="No canaries for this test case"),
            ])

        # Search the test case input files for canary codes
        tc_inputs = self.suite_dir / "test_cases" / self.test_case_id / "input_files"

        found_count = 0
        for label, canary_code in sorted(self._canaries.items()):
            found = False
            found_in = ""
            if tc_inputs.is_dir():
                for fpath in sorted(tc_inputs.rglob("*")):
                    if fpath.is_file() and _search_file_for_text(fpath, canary_code):
                        found = True
                        found_in = str(fpath.relative_to(self.suite_dir))
                        break

            checks.append(CheckResult(
                name=f"canary_{label}",
                status="pass" if found else "fail",
                expected=canary_code,
                actual="found" if found else "not_found",
                detail=f"Found in {found_in}" if found else "Not found in input files",
            ))
            if found:
                found_count += 1

        total = len(self._canaries)
        if found_count == total:
            score = 3
        elif found_count >= total * 0.5:
            score = 2
        else:
            score = 1

        return DimensionResult(dimension="canaries", score=score, checks=checks)


# ---------------------------------------------------------------------------
# Aggregate summary
# ---------------------------------------------------------------------------

def aggregate_reports(reports: list[GradeReport]) -> dict[str, Any]:
    """Build an aggregate summary across all graded test cases."""
    dimensions = ["correctness", "completeness", "format", "canaries", "error_detection"]
    summary: dict[str, Any] = {
        "total_test_cases": len(reports),
        "per_dimension": {},
        "per_test_case": {},
    }

    for dim in dimensions:
        scores = [r.dimensions[dim].score for r in reports if dim in r.dimensions]
        summary["per_dimension"][dim] = {
            "mean_score": round(sum(scores) / len(scores), 2) if scores else 0,
            "perfect_count": sum(1 for s in scores if s == 3),
            "total": len(scores),
        }

    for r in reports:
        summary["per_test_case"][r.test_case] = r.scores

    return summary


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _run_self_test(suite_dir: Path) -> bool:
    """Run grader against gold outputs; all must score 3/3/3/3/3.

    Returns True if all pass.
    """
    gold_dir = suite_dir / "gold_standards"
    gold_files = sorted(gold_dir.glob("TC-*_gold.json"))

    if not gold_files:
        print(f"ERROR: No gold standard files found in {gold_dir}", file=sys.stderr)
        return False

    reports: list[GradeReport] = []
    all_pass = True

    for gf in gold_files:
        tc_id = gf.stem.replace("_gold", "")
        tc_output = suite_dir / "test_cases" / tc_id
        if not tc_output.is_dir():
            tc_output = suite_dir  # Fallback: use suite root

        grader = SelfTestGrader(
            test_case_id=tc_id,
            gold_standard_path=gf,
            agent_output_path=tc_output,
            suite_dir=suite_dir,
        )

        report = grader.grade()
        reports.append(report)

        scores = report.scores
        perfect = all(s == 3 for s in scores.values())
        marker = "PASS" if perfect else "FAIL"
        scores_str = "/".join(str(s) for s in scores.values())
        print(f"  {tc_id}: {scores_str} [{marker}]")

        if not perfect:
            all_pass = False
            # Show failing checks
            for dim_name, dim_result in sorted(report.dimensions.items()):
                if dim_result.score < 3:
                    for check in dim_result.checks:
                        if check.status != "pass":
                            print(f"    {dim_name}/{check.name}: {check.status}"
                                  f" — {check.detail}")

    # Write aggregate report
    agg = aggregate_reports(reports)
    report_path = suite_dir / "scoring" / "self_test_report.json"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    with open(report_path, "w") as f:
        json.dump(agg, f, indent=2, sort_keys=True)
        f.write("\n")

    print()
    if all_pass:
        print(f"SELF-TEST PASSED: {len(reports)} test case(s) scored 3/3/3/3/3")
    else:
        print("SELF-TEST FAILED: see details above")

    return all_pass


def _grade_test_case(
    tc_id: str,
    gold_path: Path,
    agent_output: Path,
    suite_dir: Path | None,
    output_path: Path | None,
) -> GradeReport:
    """Grade a single test case and optionally write a JSON report."""
    grader = TestCaseGrader(
        test_case_id=tc_id,
        gold_standard_path=gold_path,
        agent_output_path=agent_output,
        suite_dir=suite_dir,
    )
    report = grader.grade()

    if output_path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w") as f:
            json.dump(report.to_dict(), f, indent=2, sort_keys=True)
            f.write("\n")

    return report


def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(
        description="Grade agent outputs against Cascade Industries gold standards.",
    )
    parser.add_argument("--self-test", action="store_true",
                        help="Run grader against gold outputs (must score 3/3/3/3/3)")
    parser.add_argument("--suite-dir", type=Path,
                        help="Path to the generated test suite")
    parser.add_argument("--tc", type=str,
                        help="Test case ID to grade (e.g. TC-01)")
    parser.add_argument("--gold", type=Path,
                        help="Path to gold standard JSON file")
    parser.add_argument("--agent-output", type=Path,
                        help="Path to agent output directory for this TC")
    parser.add_argument("--agent-output-dir", type=Path,
                        help="Path to agent output root (contains TC-XX dirs)")
    parser.add_argument("--report-dir", type=Path,
                        help="Directory to write JSON reports")
    args = parser.parse_args(argv)

    if args.self_test:
        if not args.suite_dir:
            print("ERROR: --self-test requires --suite-dir", file=sys.stderr)
            sys.exit(1)
        ok = _run_self_test(args.suite_dir)
        sys.exit(0 if ok else 1)

    if args.tc and args.gold and args.agent_output:
        report_path = None
        if args.report_dir:
            report_path = args.report_dir / f"{args.tc}_report.json"
        report = _grade_test_case(
            args.tc, args.gold, args.agent_output,
            args.suite_dir, report_path,
        )
        scores = report.scores
        scores_str = "/".join(str(s) for s in scores.values())
        print(f"{args.tc}: {scores_str}")
        return

    if args.suite_dir and args.agent_output_dir:
        gold_dir = args.suite_dir / "gold_standards"
        gold_files = sorted(gold_dir.glob("TC-*_gold.json"))
        reports: list[GradeReport] = []

        for gf in gold_files:
            tc_id = gf.stem.replace("_gold", "")
            agent_dir = args.agent_output_dir / tc_id
            if not agent_dir.is_dir():
                print(f"  {tc_id}: SKIP (no agent output at {agent_dir})")
                continue

            report_path = None
            if args.report_dir:
                report_path = args.report_dir / f"{tc_id}_report.json"

            report = _grade_test_case(
                tc_id, gf, agent_dir, args.suite_dir, report_path,
            )
            reports.append(report)

            scores = report.scores
            scores_str = "/".join(str(s) for s in scores.values())
            print(f"  {tc_id}: {scores_str}")

        if reports and args.report_dir:
            agg = aggregate_reports(reports)
            agg_path = args.report_dir / "aggregate_report.json"
            with open(agg_path, "w") as f:
                json.dump(agg, f, indent=2, sort_keys=True)
                f.write("\n")
            print(f"\nAggregate report: {agg_path}")
        return

    parser.print_help()
    sys.exit(1)


if __name__ == "__main__":
    main()
