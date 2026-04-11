"""Tests for TC-03 — Substantive Analytical Procedures — Revenue (Audit, Complex).

Verifies:
- revenue_by_product_monthly_fy2024_fy2025.xlsx: 24 months, 6 product lines,
  unit volumes and ASPs
- industry_benchmark_report.pdf: 12 pages, benchmarks on pages 4, 7, 11
- management_rep_letter.docx: asserts ~8% growth vs actual 9.2%
- ERR-009 planted error (stale industry benchmark on page 4)
- ERR-018 planted error (classification error in March 2025 row)
- Canary embedding in all files
- Gold standard structure, revenue analysis, and scoring hints
- Prompt and expected behavior markdown files
"""

from __future__ import annotations

import json
import tempfile
from decimal import Decimal
from pathlib import Path

import openpyxl
from docx import Document
from pypdf import PdfReader

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc03 import (
    _ASP_TABLE,
    _INDUSTRY_BENCHMARKS,
    emit_tc03,
)
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model
from generator.model.revenue import (
    PRODUCT_LINES,
    validate_consolidated_growth,
    validate_product_line_growth,
)

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc03 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None

# All canary keys used by TC-03
_CANARY_KEYS = sorted([
    "tc03_revenue_by_product",
    "tc03_industry_benchmark",
    "tc03_mgmt_rep_letter",
])


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc03_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc03(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard
        emit_gold("TC-03", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-03/input_files"


# ---------------------------------------------------------------------------
# Revenue XLSX — 24 months × 6 product lines
# ---------------------------------------------------------------------------


class TestRevenueXLSX:
    """Verify revenue_by_product_monthly_fy2024_fy2025.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx"
        assert path.exists()

    def test_has_monthly_revenue_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Monthly Revenue" in wb.sheetnames

    def test_24_months_covered(self) -> None:
        """Should have data spanning FY2024 and FY2025 (12 months each)."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Monthly Revenue"]
        year_month_pairs = set()
        for row in range(2, ws.max_row + 1):
            year = ws.cell(row=row, column=1).value
            month = ws.cell(row=row, column=2).value
            if isinstance(year, int) and isinstance(month, str):
                year_month_pairs.add((year, month))
        # 12 months × 2 years = 24 unique (year, month) combos
        assert len(year_month_pairs) == 24, (
            f"Expected 24 (year, month) pairs, got {len(year_month_pairs)}"
        )

    def test_6_product_lines(self) -> None:
        """All 6 product lines should appear."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Monthly Revenue"]
        product_lines = set()
        for row in range(2, ws.max_row + 1):
            val = ws.cell(row=row, column=3).value
            if isinstance(val, str) and val.strip():
                product_lines.add(val.strip())
        # ERR-018 changes one "Specialty Coatings" to "Advanced Composites",
        # so Advanced Composites will appear but Specialty Coatings may still
        # appear in other months.
        assert len(product_lines) >= 6, (
            f"Expected ≥6 product lines, got {len(product_lines)}: {product_lines}"
        )

    def test_has_unit_volume_column(self) -> None:
        """Should have a units/volume column with numeric data."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Monthly Revenue"]
        # Header row is row 4 (rows 1-3 are title/subtitle/blank)
        headers = []
        for col in range(1, ws.max_column + 1):
            val = ws.cell(row=4, column=col).value
            if val is not None:
                headers.append(str(val).lower())
        assert any("unit" in h or "volume" in h for h in headers), (
            f"No units/volume column found in headers: {headers}"
        )

    def test_has_asp_column(self) -> None:
        """Should have an ASP (average selling price) column."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Monthly Revenue"]
        headers = []
        for col in range(1, ws.max_column + 1):
            val = ws.cell(row=4, column=col).value
            if val is not None:
                headers.append(str(val).lower())
        assert any("asp" in h or "avg" in h or "price" in h for h in headers), (
            f"No ASP column found in headers: {headers}"
        )

    def test_revenue_amounts_are_positive(self) -> None:
        """All revenue values should be positive."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Monthly Revenue"]
        # Find revenue column from header row 4
        rev_col = None
        for col in range(1, ws.max_column + 1):
            val = ws.cell(row=4, column=col).value
            if val is not None and "revenue" in str(val).lower():
                rev_col = col
                break
        assert rev_col is not None, "No revenue column found"
        for row in range(5, ws.max_row + 1):
            val = ws.cell(row=row, column=rev_col).value
            if isinstance(val, (int, float)):
                assert val > 0, f"Row {row}: revenue should be positive, got {val}"


# ---------------------------------------------------------------------------
# Industry Benchmark PDF — 12 pages
# ---------------------------------------------------------------------------


class TestBenchmarkPDF:
    """Verify industry_benchmark_report.pdf structure."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/industry_benchmark_report.pdf"
        assert path.exists()

    def test_has_12_pages(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/industry_benchmark_report.pdf"
        reader = PdfReader(str(path))
        assert len(reader.pages) == 12, (
            f"Expected 12 pages, got {len(reader.pages)}"
        )

    def test_page4_has_manufacturing_benchmark(self) -> None:
        """Page 4 should contain Industrial Manufacturing benchmark data."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/industry_benchmark_report.pdf"
        reader = PdfReader(str(path))
        text = reader.pages[3].extract_text().lower()  # 0-indexed
        assert "industrial" in text or "manufacturing" in text, (
            "Page 4 should mention Industrial Manufacturing"
        )

    def test_page7_has_materials_benchmark(self) -> None:
        """Page 7 should contain Advanced Materials benchmark data."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/industry_benchmark_report.pdf"
        reader = PdfReader(str(path))
        text = reader.pages[6].extract_text().lower()
        assert "advanced" in text or "materials" in text or "composites" in text, (
            "Page 7 should mention Advanced Materials"
        )

    def test_page11_has_logistics_benchmark(self) -> None:
        """Page 11 should contain Logistics benchmark data."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/industry_benchmark_report.pdf"
        reader = PdfReader(str(path))
        text = reader.pages[10].extract_text().lower()
        assert "logistics" in text or "distribution" in text, (
            "Page 11 should mention Logistics & Distribution"
        )

    def test_benchmark_growth_rates_present(self) -> None:
        """All benchmark growth rates should appear somewhere in the PDF."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/industry_benchmark_report.pdf"
        reader = PdfReader(str(path))
        full_text = ""
        for page in reader.pages:
            full_text += page.extract_text() + "\n"
        # Note: ERR-009 replaces 4.2% with 3.8% for Industrial Manufacturing on page 4,
        # so we check for 3.8% (the stale value) or 4.2% (in case fix is applied)
        assert "3.8%" in full_text or "4.2%" in full_text, (
            "Manufacturing growth rate not found"
        )
        assert "12.8%" in full_text, "Advanced Materials growth rate not found"
        assert "7.5%" in full_text, "Logistics growth rate not found"


# ---------------------------------------------------------------------------
# Management Representation Letter
# ---------------------------------------------------------------------------


class TestMgmtRepLetter:
    """Verify management_rep_letter.docx content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_rep_letter.docx"
        assert path.exists()

    def test_asserts_8pct_growth(self) -> None:
        """Letter should claim approximately 8% growth."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_rep_letter.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "8%" in text, "Management letter should claim ~8% growth"

    def test_mentions_advanced_materials(self) -> None:
        """Letter should mention Advanced Materials as growth driver."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_rep_letter.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs).lower()
        assert "advanced" in text, (
            "Management letter should mention Advanced Materials"
        )

    def test_mentions_cascade_industries(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_rep_letter.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs).lower()
        assert "cascade" in text, "Management letter should mention Cascade Industries"


# ---------------------------------------------------------------------------
# Revenue model validation — 9.2% consolidated growth
# ---------------------------------------------------------------------------


class TestRevenueModelConsistency:
    """Verify canonical model produces the expected growth rates."""

    def test_consolidated_growth_is_9_2_pct(self) -> None:
        """Consolidated FY24→25 growth must be 9.2% per spec."""
        model, _, _, _, _ = _ensure_emitted()
        growth = validate_consolidated_growth(model.revenue_records)
        actual = float(growth.get("FY2025_growth", Decimal(0))) * 100
        assert abs(actual - 9.2) < 0.1, (
            f"Consolidated growth should be ~9.2%, got {actual:.1f}%"
        )

    def test_one_product_line_declines(self) -> None:
        """At least one product line should show negative YoY growth."""
        model, _, _, _, _ = _ensure_emitted()
        pl_growth = validate_product_line_growth(model.revenue_records)
        declining = [name for name, g in pl_growth.items() if g < 0]
        assert len(declining) >= 1, "At least one product line should decline"

    def test_specialty_coatings_declines_about_4pct(self) -> None:
        """Specialty Coatings should decline ~4% per spec."""
        model, _, _, _, _ = _ensure_emitted()
        pl_growth = validate_product_line_growth(model.revenue_records)
        sc_growth = float(pl_growth.get("Specialty Coatings", Decimal(0))) * 100
        assert sc_growth < 0, f"Specialty Coatings should decline, got {sc_growth:.1f}%"
        assert abs(sc_growth - (-4.0)) < 2.0, (
            f"Specialty Coatings should decline ~4%, got {sc_growth:.1f}%"
        )


# ---------------------------------------------------------------------------
# ERR-009 — Stale industry benchmark
# ---------------------------------------------------------------------------


class TestERR009PlantedError:
    """Verify ERR-009: stale data in industry benchmark PDF."""

    def test_err009_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-009" in errors.entries

    def test_err009_is_stale_data(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-009"]
        assert err.type == "stale_data"

    def test_err009_in_benchmark_pdf(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-009"]
        assert "industry_benchmark_report.pdf" in err.file

    def test_err009_references_page4(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-009"]
        assert "4" in err.location

    def test_err009_stale_value_in_pdf(self) -> None:
        """Page 4 should show 3.8% (stale FY2024 value) instead of 4.2%."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/industry_benchmark_report.pdf"
        reader = PdfReader(str(path))
        page4_text = reader.pages[3].extract_text()
        assert "3.8%" in page4_text, (
            "Page 4 should contain stale 3.8% value (ERR-009)"
        )


# ---------------------------------------------------------------------------
# ERR-018 — Classification error in March 2025
# ---------------------------------------------------------------------------


class TestERR018PlantedError:
    """Verify ERR-018: classification error in March 2025 revenue row."""

    def test_err018_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-018" in errors.entries

    def test_err018_is_classification_error(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-018"]
        assert err.type == "classification_error"

    def test_err018_in_revenue_file(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-018"]
        assert "revenue_by_product_monthly_fy2024_fy2025.xlsx" in err.file

    def test_err018_march_2025_misclassified(self) -> None:
        """March 2025 should show 'Advanced Composites' where
        'Specialty Coatings' should be."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Monthly Revenue"]
        # Find the March 2025 row with misclassified product line
        march_2025_pls = []
        for row in range(2, ws.max_row + 1):
            year = ws.cell(row=row, column=1).value
            month = ws.cell(row=row, column=2).value
            pl = ws.cell(row=row, column=3).value
            if year == 2025 and isinstance(month, str) and month.lower() == "march":
                march_2025_pls.append(pl)
        # Should have "Advanced Composites" appearing twice (once real, once ERR-018)
        ac_count = sum(1 for pl in march_2025_pls if pl == "Advanced Composites")
        assert ac_count >= 2, (
            f"Expected 'Advanced Composites' at least twice in March 2025 "
            f"(one real + one misclassified), got {ac_count}. PLs: {march_2025_pls}"
        )


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """Verify canary codes are embedded in files."""

    def test_all_canary_keys_assigned(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            code = canaries.canary_for(key)
            assert len(code) == 8, f"Canary for {key} should be 8 chars"

    def test_revenue_xlsx_canary(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc03_revenue_by_product")
        path = output / f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in revenue XLSX properties"

    def test_benchmark_pdf_canary(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc03_industry_benchmark")
        path = output / f"{_INPUT_DIR}/industry_benchmark_report.pdf"
        reader = PdfReader(str(path))
        author = reader.metadata.author or ""
        assert canary in author, f"Canary {canary} not in benchmark PDF metadata"

    def test_mgmt_rep_letter_canary(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc03_mgmt_rep_letter")
        path = output / f"{_INPUT_DIR}/management_rep_letter.docx"
        doc = Document(str(path))
        comments = doc.core_properties.comments or ""
        assert canary in comments, f"Canary {canary} not in mgmt rep letter properties"


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold standard JSON structure."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-03_gold.json"
        assert path.exists()

    def test_gold_has_consolidated_revenue(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-03_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "consolidated_revenue" in eo
        cr = eo["consolidated_revenue"]
        assert "fy2024" in cr
        assert "fy2025" in cr
        assert "yoy_growth_pct" in cr

    def test_gold_growth_is_9_2_pct(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-03_gold.json").read_text())
        growth = gold["expected_outputs"]["consolidated_revenue"]["yoy_growth_pct"]
        assert abs(growth - 9.2) < 0.5, f"Gold growth should be ~9.2%, got {growth}"

    def test_gold_has_management_discrepancy(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-03_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "management_rep_discrepancy" in eo
        disc = eo["management_rep_discrepancy"]
        assert disc["discrepancy_detected"] is True
        assert "8%" in disc["claimed_growth"]

    def test_gold_has_product_line_analysis(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-03_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "product_line_analysis" in eo
        pla = eo["product_line_analysis"]
        assert len(pla) >= 6, f"Expected ≥6 product lines in analysis, got {len(pla)}"

    def test_gold_has_industry_benchmarks(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-03_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "industry_benchmark_comparison" in eo
        ibc = eo["industry_benchmark_comparison"]
        assert "industrial_manufacturing" in ibc
        assert "advanced_materials" in ibc
        assert "logistics_distribution" in ibc

    def test_gold_has_growth_drivers(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-03_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "growth_drivers" in eo
        assert "declining_lines" in eo

    def test_gold_canary_verification(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-03_gold.json").read_text())
        cv = gold["canary_verification"]
        for key in ["read_revenue_data", "read_benchmark_report", "read_mgmt_rep_letter"]:
            assert key in cv, f"Missing canary verification key: {key}"

    def test_gold_scoring_hints(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-03_gold.json").read_text())
        hints = gold["scoring_hints"]
        for key in ["correctness", "completeness", "format_compliance", "robustness", "communication"]:
            assert key in hints

    def test_gold_has_output_files(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-03_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "output_files" in eo
        of = eo["output_files"]
        assert "memo" in of
        assert "analysis_workbook" in of


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt and expected behavior files are generated."""

    def test_prompt_md_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-03/prompt.md"
        assert path.exists()

    def test_prompt_mentions_analytical_procedures(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-03/prompt.md").read_text().lower()
        assert "analytical" in text

    def test_prompt_mentions_revenue(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-03/prompt.md").read_text().lower()
        assert "revenue" in text

    def test_prompt_mentions_growth(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-03/prompt.md").read_text().lower()
        assert "growth" in text

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-03/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_discrepancy(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-03/expected_behavior.md").read_text().lower()
        assert "8%" in text or "discrepancy" in text

    def test_expected_behavior_mentions_9_2_pct(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-03/expected_behavior.md").read_text()
        assert "9.2" in text, "Expected behavior should mention 9.2% actual growth"


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifest:
    """Verify files are registered in the manifest."""

    def test_manifest_has_tc03_entries(self) -> None:
        _, _, _, _, manifest = _ensure_emitted()
        tc03_count = sum(
            1 for v in manifest.entries.values()
            if "TC-03" in (v.test_cases or [])
        )
        # revenue xlsx + benchmark pdf + mgmt rep letter = 3
        assert tc03_count >= 3, f"Expected ≥3 TC-03 manifest entries, got {tc03_count}"


# ---------------------------------------------------------------------------
# Constants verification
# ---------------------------------------------------------------------------


class TestConstants:
    """Verify TC-03 constants are consistent with spec."""

    def test_6_product_lines_in_asp_table(self) -> None:
        assert len(_ASP_TABLE) == 6

    def test_asp_table_keys_match_product_lines(self) -> None:
        expected = {pl.name for pl in PRODUCT_LINES}
        assert set(_ASP_TABLE.keys()) == expected

    def test_3_industry_benchmarks(self) -> None:
        assert len(_INDUSTRY_BENCHMARKS) == 3

    def test_benchmark_pages_are_4_7_11(self) -> None:
        pages = sorted(bm["page"] for bm in _INDUSTRY_BENCHMARKS.values())
        assert pages == [4, 7, 11]
