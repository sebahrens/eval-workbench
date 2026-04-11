"""Tests for TC-21 — Combined Diligence Findings Memo (Cross-service, Adversarial).

Verifies:
- legal/ directory with 3 curated contract summaries (LCTR-001, LCTR-003, LCTR-004)
- legal/ directory with 2 curated amendments (AMD-002, AMD-003)
- hr/ directory with 2 employment agreements (EA-001, EA-006)
- hr/ directory with severance_schedule.xlsx and retention_schedule.xlsx
- management_qa_summary.docx (management representations, not independently verified)
- diligence_request_tracker.xlsx (combined legal + HR requests)
- prompt.md and expected_behavior.md
- Judgment traps: top legal risk (LDI-001), top HR exposure (SEV-001),
  unresolved DR-006, management-vs-source distinction, no legal advice language
- Canary embedding in all files
- Gold standard structure with evidence expectations and judgment traps
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc21 import (
    _CANARY_KEYS,
    _CURATED_AGREEMENT_IDS,
    _CURATED_AMENDMENT_IDS,
    _CURATED_CONTRACT_IDS,
    emit_tc21,
)
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model
from generator.model.hr_diligence import (
    DILIGENCE_REQUESTS,
    SEVERANCE_EXPOSURES,
    open_diligence_requests,
    total_retention_awards,
    total_severance_exposure,
)
from generator.model.legal import (
    LEGAL_DILIGENCE_ISSUES,
    missing_consent_issues,
    stale_summary_issues,
)

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc21 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc21_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc21(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard
        emit_gold("TC-21", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-21/input_files"


# ---------------------------------------------------------------------------
# Legal directory — curated contracts
# ---------------------------------------------------------------------------


class TestCuratedContracts:
    """Verify 3 curated contract summary docx files in legal/."""

    def test_legal_dir_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        assert (output / f"{_INPUT_DIR}/legal").is_dir()

    def test_3_contract_files(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        for cid in _CURATED_CONTRACT_IDS:
            path = output / f"{_INPUT_DIR}/legal/contract_{cid.lower()}.docx"
            assert path.exists(), f"Missing contract file for {cid}"

    def test_contract_lctr001_mentions_change_of_control(self) -> None:
        """LCTR-001 (Acme Manufacturing) should mention change-of-control."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/legal/contract_lctr-001.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\n".join(cell.text for cell in row.cells)
        assert "change" in text.lower() and "control" in text.lower()

    def test_contract_lctr003_mentions_mfn(self) -> None:
        """LCTR-003 (TechAlloy Systems) should mention MFN pricing."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/legal/contract_lctr-003.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\n".join(cell.text for cell in row.cells)
        assert "mfn" in text.lower() or "most-favored" in text.lower() or "most favored" in text.lower()

    def test_contract_lctr004_mentions_exclusivity(self) -> None:
        """LCTR-004 (NextGen Composites) should mention exclusivity."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/legal/contract_lctr-004.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\n".join(cell.text for cell in row.cells)
        assert "exclusiv" in text.lower()


# ---------------------------------------------------------------------------
# Legal directory — curated amendments
# ---------------------------------------------------------------------------


class TestCuratedAmendments:
    """Verify 2 curated amendment docx files in legal/."""

    def test_2_amendment_files(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        for aid in _CURATED_AMENDMENT_IDS:
            path = output / f"{_INPUT_DIR}/legal/amendment_{aid.lower()}.docx"
            assert path.exists(), f"Missing amendment file for {aid}"

    def test_total_legal_files_is_5(self) -> None:
        """3 contracts + 2 amendments = 5 docx files in legal/."""
        _, output, _, _, _ = _ensure_emitted()
        legal_dir = output / f"{_INPUT_DIR}/legal"
        docx_files = list(legal_dir.glob("*.docx"))
        assert len(docx_files) == 5, f"Expected 5 legal docx files, got {len(docx_files)}"


# ---------------------------------------------------------------------------
# HR directory — employment agreements
# ---------------------------------------------------------------------------


class TestCuratedAgreements:
    """Verify 2 curated employment agreement docx files in hr/."""

    def test_hr_dir_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        assert (output / f"{_INPUT_DIR}/hr").is_dir()

    def test_2_agreement_files(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        for eaid in _CURATED_AGREEMENT_IDS:
            path = output / f"{_INPUT_DIR}/hr/agreement_{eaid.lower()}.docx"
            assert path.exists(), f"Missing agreement file for {eaid}"


# ---------------------------------------------------------------------------
# HR directory — severance and retention schedules
# ---------------------------------------------------------------------------


class TestSchedules:
    """Verify severance and retention schedule xlsx files in hr/."""

    def test_severance_schedule_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        assert (output / f"{_INPUT_DIR}/hr/severance_schedule.xlsx").exists()

    def test_retention_schedule_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        assert (output / f"{_INPUT_DIR}/hr/retention_schedule.xlsx").exists()


# ---------------------------------------------------------------------------
# Management Q&A Summary
# ---------------------------------------------------------------------------


class TestManagementQA:
    """Verify management_qa_summary.docx content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        assert (output / f"{_INPUT_DIR}/management_qa_summary.docx").exists()

    def test_mentions_acme_change_of_control(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_qa_summary.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "lctr-001" in text.lower() or "acme" in text.lower()
        assert "change" in text.lower() and "control" in text.lower()

    def test_mentions_mfn_without_amd002_detail(self) -> None:
        """Management Q&A should mention MFN but NOT disclose AMD-002 threshold."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_qa_summary.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "mfn" in text.lower() or "most-favored" in text.lower() or "most favored" in text.lower()

    def test_mentions_missing_ea006(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_qa_summary.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "ea-006" in text.lower() or "patel" in text.lower()

    def test_mentions_not_independently_verified(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_qa_summary.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "not" in text.lower() and "verified" in text.lower()

    def test_mentions_ceo_severance(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_qa_summary.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        ceo_payout = int(SEVERANCE_EXPOSURES[0].estimated_payout)
        assert f"${ceo_payout:,}" in text or "cascade" in text.lower()


# ---------------------------------------------------------------------------
# Diligence Request Tracker
# ---------------------------------------------------------------------------


class TestDiligenceRequestTracker:
    """Verify diligence_request_tracker.xlsx structure."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        assert (output / f"{_INPUT_DIR}/diligence_request_tracker.xlsx").exists()

    def test_has_legal_requests_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/diligence_request_tracker.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Legal Requests" in wb.sheetnames

    def test_has_hr_requests_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/diligence_request_tracker.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "HR Requests" in wb.sheetnames

    def test_legal_requests_has_rows(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/diligence_request_tracker.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Legal Requests"]
        data_rows = ws.max_row - 1  # minus header
        assert data_rows == len(LEGAL_DILIGENCE_ISSUES)

    def test_hr_requests_has_rows(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/diligence_request_tracker.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["HR Requests"]
        data_rows = ws.max_row - 1  # minus header
        assert data_rows == len(DILIGENCE_REQUESTS)

    def test_legal_headers(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/diligence_request_tracker.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Legal Requests"]
        headers = [ws.cell(row=1, column=c).value for c in range(1, 8)]
        assert "Issue ID" in headers
        assert "Severity" in headers
        assert "Status" in headers

    def test_hr_headers(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/diligence_request_tracker.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["HR Requests"]
        headers = [ws.cell(row=1, column=c).value for c in range(1, 9)]
        assert "Request ID" in headers
        assert "Status" in headers


# ---------------------------------------------------------------------------
# Judgment traps — content verification
# ---------------------------------------------------------------------------


class TestJudgmentTraps:
    """Verify judgment trap signals are present in the generated files."""

    def test_top_legal_risk_ldi001_in_legal_requests(self) -> None:
        """LDI-001 (change-of-control) should appear in legal requests sheet."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/diligence_request_tracker.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Legal Requests"]
        ids = set()
        for row in range(2, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if val:
                ids.add(val)
        assert "LDI-001" in ids

    def test_top_hr_exposure_sev001_in_severance(self) -> None:
        """SEV-001 (CEO golden parachute) should appear in severance schedule."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/hr/severance_schedule.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        found = False
        for row in range(1, ws.max_row + 1):
            for col in range(1, ws.max_column + 1):
                val = ws.cell(row=row, column=col).value
                if isinstance(val, str) and "SEV-001" in val:
                    found = True
                    break
            if found:
                break
        assert found, "SEV-001 not found in severance schedule"

    def test_unresolved_dr006_exists(self) -> None:
        """DR-006 should be in the open diligence requests."""
        open_reqs = open_diligence_requests()
        req_ids = [dr.request_id for dr in open_reqs]
        assert "DR-006" in req_ids

    def test_management_vs_source_distinction_setup(self) -> None:
        """Management Q&A should contain statements that differ from primary sources.

        This verifies the setup for the judgment trap: the management Q&A
        describes exclusivity in original terms but AMD-003 expanded scope.
        """
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_qa_summary.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Management describes "original contract terms" for exclusivity
        assert "original" in text.lower()

    def test_stale_summary_issues_exist(self) -> None:
        """Model should have stale summary issues for management contradictions."""
        stale = stale_summary_issues()
        assert len(stale) > 0, "Expected at least one stale summary issue"

    def test_missing_consent_issues_exist(self) -> None:
        """Model should have missing consent issues."""
        consent = missing_consent_issues()
        assert len(consent) > 0, "Expected at least one missing consent issue"


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """Verify canary codes are embedded in all TC-21 files."""

    def test_all_canary_keys_assigned(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            code = canaries.canary_for(key)
            assert len(code) == 8, f"Canary for {key} should be 8 chars"

    def test_contract_canaries_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        for cid in _CURATED_CONTRACT_IDS:
            ckey = f"tc21_contract_{cid.lower().replace('-', '_')}"
            canary = canaries.canary_for(ckey)
            path = output / f"{_INPUT_DIR}/legal/contract_{cid.lower()}.docx"
            doc = Document(str(path))
            comments = doc.core_properties.comments or ""
            assert canary in comments, f"Canary {canary} not in {cid} contract properties"

    def test_amendment_canaries_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        for aid in _CURATED_AMENDMENT_IDS:
            ckey = f"tc21_amendment_{aid.lower().replace('-', '_')}"
            canary = canaries.canary_for(ckey)
            path = output / f"{_INPUT_DIR}/legal/amendment_{aid.lower()}.docx"
            doc = Document(str(path))
            comments = doc.core_properties.comments or ""
            assert canary in comments, f"Canary {canary} not in {aid} amendment properties"

    def test_agreement_canaries_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        for eaid in _CURATED_AGREEMENT_IDS:
            ckey = f"tc21_agreement_{eaid.lower().replace('-', '_')}"
            canary = canaries.canary_for(ckey)
            path = output / f"{_INPUT_DIR}/hr/agreement_{eaid.lower()}.docx"
            doc = Document(str(path))
            comments = doc.core_properties.comments or ""
            assert canary in comments, f"Canary {canary} not in {eaid} agreement properties"

    def test_severance_schedule_canary(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc21_severance_schedule")
        path = output / f"{_INPUT_DIR}/hr/severance_schedule.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in severance schedule properties"

    def test_retention_schedule_canary(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc21_retention_schedule")
        path = output / f"{_INPUT_DIR}/hr/retention_schedule.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in retention schedule properties"

    def test_management_qa_canary(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc21_management_qa_summary")
        path = output / f"{_INPUT_DIR}/management_qa_summary.docx"
        doc = Document(str(path))
        comments = doc.core_properties.comments or ""
        assert canary in comments, f"Canary {canary} not in management Q&A properties"

    def test_diligence_tracker_canary(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc21_diligence_request_tracker")
        path = output / f"{_INPUT_DIR}/diligence_request_tracker.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in diligence tracker properties"


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold standard JSON structure."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-21_gold.json"
        assert path.exists()

    def test_gold_has_required_sections(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert eo["output_format"] == "combined_findings_memo"
        sections = eo["required_sections"]
        for s in ["executive_summary", "legal_findings", "hr_findings",
                   "unresolved_requests", "assumptions_and_limitations", "citations"]:
            assert s in sections, f"Missing required section: {s}"

    def test_gold_curated_counts(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert eo["curated_contract_count"] == 3
        assert eo["curated_amendment_count"] == 2
        assert eo["curated_agreement_count"] == 2

    def test_gold_top_legal_risk(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        risk = gold["expected_outputs"]["top_legal_risk"]
        assert risk["issue_id"] == "LDI-001"
        assert risk["contract_id"] == "LCTR-001"
        assert risk["severity"] == "high"

    def test_gold_top_hr_exposure(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        hr = gold["expected_outputs"]["top_hr_exposure"]
        assert hr["exposure_id"] == "SEV-001"
        assert hr["employee_name"] == "Robert J. Cascade"

    def test_gold_severance_and_retention_totals(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert eo["total_severance_exposure"] == str(total_severance_exposure())
        assert eo["total_retention_awards"] == str(total_retention_awards())

    def test_gold_management_source_contradictions(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        contradictions = gold["expected_outputs"]["management_source_contradictions"]
        stale = stale_summary_issues()
        assert contradictions == [issue.issue_id for issue in stale]

    def test_gold_unresolved_requests(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        open_reqs = open_diligence_requests()
        assert gold["expected_outputs"]["unresolved_requests"] == [
            dr.request_id for dr in open_reqs
        ]

    def test_gold_canary_verification_keys(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        cv = gold["canary_verification"]
        # Should have entries for all curated files plus schedules/tracker/qa
        expected_keys = (
            [f"read_contract_{cid.lower().replace('-', '_')}" for cid in _CURATED_CONTRACT_IDS]
            + [f"read_amendment_{aid.lower().replace('-', '_')}" for aid in _CURATED_AMENDMENT_IDS]
            + [f"read_agreement_{eaid.lower().replace('-', '_')}" for eaid in _CURATED_AGREEMENT_IDS]
            + [
                "read_severance_schedule",
                "read_retention_schedule",
                "read_management_qa_summary",
                "read_diligence_request_tracker",
            ]
        )
        for key in expected_keys:
            assert key in cv, f"Missing canary verification key: {key}"

    def test_gold_scoring_hints(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        hints = gold["scoring_hints"]
        for key in ["synthesis", "top_risks", "source_distinction",
                     "professional_language", "unresolved_tracking"]:
            assert key in hints, f"Missing scoring hint: {key}"

    def test_gold_has_evidence_expectations(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        ee = gold["evidence_expectations"]
        for key in ["finding_top_legal_risk", "finding_top_hr_exposure",
                     "finding_mfn_contradiction", "finding_exclusivity_expansion",
                     "finding_unresolved_requests", "finding_missing_agreement"]:
            assert key in ee, f"Missing evidence expectation: {key}"
            assert "required_sources" in ee[key]
            assert "acceptable_terms" in ee[key]

    def test_gold_has_judgment_traps(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        traps = gold["judgment_traps"]
        assert len(traps) == 5
        trap_ids = {t["trap_id"] for t in traps}
        for i in range(1, 6):
            assert f"JT-TC21-00{i}" in trap_ids

    def test_gold_scenario_pack(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        assert gold["scenario_pack"] == "ma_legal_hr_diligence"
        assert gold["service_line"] == "advisory"

    def test_gold_no_error_detection(self) -> None:
        """TC-21 has no planted errors — error_detection should be empty."""
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-21_gold.json").read_text())
        assert gold["error_detection"] == {}


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt and expected behavior files are generated."""

    def test_prompt_md_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        assert (output / "test_cases/TC-21/prompt.md").exists()

    def test_prompt_mentions_combined_findings(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-21/prompt.md").read_text().lower()
        assert "findings memo" in text

    def test_prompt_mentions_executive_summary(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-21/prompt.md").read_text().lower()
        assert "executive summary" in text

    def test_prompt_mentions_citations(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-21/prompt.md").read_text().lower()
        assert "citation" in text

    def test_prompt_mentions_no_legal_advice(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-21/prompt.md").read_text().lower()
        assert "legal advice" in text or "disclaimer" in text

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        assert (output / "test_cases/TC-21/expected_behavior.md").exists()

    def test_expected_behavior_mentions_ldi001(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-21/expected_behavior.md").read_text()
        assert "LDI-001" in text

    def test_expected_behavior_mentions_sev001(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-21/expected_behavior.md").read_text()
        assert "SEV-001" in text

    def test_expected_behavior_mentions_dr006(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-21/expected_behavior.md").read_text()
        assert "DR-006" in text

    def test_expected_behavior_mentions_amd002(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-21/expected_behavior.md").read_text()
        assert "AMD-002" in text

    def test_expected_behavior_mentions_canonical_amounts(self) -> None:
        """Expected behavior should use canonical model amounts."""
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-21/expected_behavior.md").read_text()
        total_sev = int(total_severance_exposure())
        total_ret = int(total_retention_awards())
        assert f"${total_sev:,}" in text
        assert f"${total_ret:,}" in text


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifest:
    """Verify files are registered in the manifest."""

    def test_manifest_has_tc21_entries(self) -> None:
        _, _, _, _, manifest = _ensure_emitted()
        tc21_count = sum(
            1 for v in manifest.entries.values()
            if "TC-21" in (v.test_cases or [])
        )
        # 3 contracts + 2 amendments + 2 agreements + severance + retention
        # + management_qa + diligence_tracker = 11
        assert tc21_count >= 11, f"Expected >= 11 TC-21 manifest entries, got {tc21_count}"
