"""Tests for TC-10-EU — VAT and Cross-Border Tax Position Analysis.

Verifies:
- Intercompany sales XLSX (2 sheets, correct columns, ERR-EU-010 planted)
- VAT registrations XLSX (5 rows including pending Polish registration trap)
- VAT returns summary XLSX (16 rows: 4 entities x 4 quarters)
- EU VAT rules reference DOCX (exists, has canary, key sections)
- ERR-EU-010 planted error (registered, type "classification_error", severity "material")
- Canary embedding in all 4 files
- Gold standard JSON (structure, expected outputs, scoring hints)
- Prompt and expected behavior markdown files
- Model computation tests (IC sales, VAT registrations, VAT returns)
- Determinism (identical outputs on repeat)
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import docx
import openpyxl

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc10_eu import emit_tc10_eu
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model
from generator.model.vat_eu import (
    ALL_CANARY_KEYS_TC10EU,
    ERR_EU_010_AMOUNT,
    ERR_EU_010_VAT_CHARGED,
    VAT_RATES,
    generate_ic_sales_eu,
    generate_vat_registrations,
    generate_vat_returns,
)

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc10_eu once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc10eu_test_"))
        _CANARIES = build_canary_registry(ALL_CANARY_KEYS_TC10EU, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc10_eu(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)
        emit_gold("TC-10-EU", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-10-EU/input_files"


# ---------------------------------------------------------------------------
# Intercompany Sales XLSX
# ---------------------------------------------------------------------------


class TestIntercompanySales:
    """Verify tc10eu_intercompany_sales_fy2025.xlsx."""

    def test_xlsx_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_intercompany_sales_fy2025.xlsx"
        assert path.exists()

    def test_xlsx_has_canary(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_intercompany_sales_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        expected = canaries.canary_for("tc10eu_intercompany_sales_fy2025")
        desc = wb.properties.description or ""
        assert expected in desc, f"Canary {expected} not in IC sales xlsx"

    def test_has_intercompany_sales_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_intercompany_sales_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Intercompany Sales" in wb.sheetnames

    def test_has_third_party_sales_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_intercompany_sales_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Third-Party Sales Summary" in wb.sheetnames

    def test_ic_sales_columns(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_intercompany_sales_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Intercompany Sales"]
        headers = [ws.cell(row=1, column=c).value for c in range(1, 10)]
        assert "Seller Entity" in headers
        assert "VAT Treatment Applied" in headers
        assert "Proof of Dispatch" in headers

    def test_ic_sales_row_count(self) -> None:
        """Should have 24 IC sales rows (6 flows x 4 quarters)."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_intercompany_sales_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Intercompany Sales"]
        data_rows = sum(1 for r in range(2, ws.max_row + 1) if ws.cell(row=r, column=1).value)
        assert data_rows == 24, f"Expected 24 IC rows, got {data_rows}"

    def test_err_eu_010_row_present(self) -> None:
        """ERR-EU-010 row: CE→CM Q3 with 20% French VAT."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_intercompany_sales_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Intercompany Sales"]
        found = False
        for r in range(2, ws.max_row + 1):
            seller = ws.cell(row=r, column=1).value
            buyer = ws.cell(row=r, column=2).value
            vat_rate = ws.cell(row=r, column=6).value
            quarter = ws.cell(row=r, column=9).value
            if seller == "CE" and buyer == "CM" and vat_rate == "20%" and quarter == "Q3":
                found = True
                break
        assert found, "ERR-EU-010 row not found (CE→CM Q3 with 20% VAT)"

    def test_q2_partial_dispatch(self) -> None:
        """Q2 CP→CD should have 'Partial' proof of dispatch."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_intercompany_sales_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Intercompany Sales"]
        found = False
        for r in range(2, ws.max_row + 1):
            seller = ws.cell(row=r, column=1).value
            buyer = ws.cell(row=r, column=2).value
            pod = ws.cell(row=r, column=8).value
            quarter = ws.cell(row=r, column=9).value
            if seller == "CP" and buyer == "CD" and quarter == "Q2" and pod == "Partial":
                found = True
                break
        assert found, "Q2 CP→CD 'Partial' proof of dispatch not found"

    def test_cm_italian_sales_in_summary(self) -> None:
        """Third-party summary should mention CM's Italian sales."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_intercompany_sales_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Third-Party Sales Summary"]
        found = False
        for r in range(2, ws.max_row + 1):
            notes = ws.cell(row=r, column=6).value or ""
            if "380k" in notes or "Italy" in notes:
                found = True
                break
        assert found, "CM Italian sales reference not in third-party summary"


# ---------------------------------------------------------------------------
# VAT Registrations XLSX
# ---------------------------------------------------------------------------


class TestVATRegistrations:
    """Verify tc10eu_vat_registrations.xlsx."""

    def test_xlsx_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_vat_registrations.xlsx"
        assert path.exists()

    def test_xlsx_has_canary(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_vat_registrations.xlsx"
        wb = openpyxl.load_workbook(str(path))
        expected = canaries.canary_for("tc10eu_vat_registrations")
        desc = wb.properties.description or ""
        assert expected in desc

    def test_registration_count(self) -> None:
        """5 registrations: CE(NL), CP(DE), CP(PL-pending), CM(FR), CD(UK)."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_vat_registrations.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["VAT Registrations"]
        data_rows = sum(1 for r in range(2, ws.max_row + 1) if ws.cell(row=r, column=1).value)
        assert data_rows == 5, f"Expected 5 registrations, got {data_rows}"

    def test_polish_pending_registration(self) -> None:
        """CP should have a pending Polish VAT registration (missing data trap)."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_vat_registrations.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["VAT Registrations"]
        found = False
        for r in range(2, ws.max_row + 1):
            entity = ws.cell(row=r, column=1).value
            country = ws.cell(row=r, column=2).value
            status = ws.cell(row=r, column=9).value or ""
            if entity == "CP" and country == "Poland" and "Pending" in status:
                found = True
                break
        assert found, "CP Polish pending registration not found"

    def test_columns(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_vat_registrations.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["VAT Registrations"]
        headers = [ws.cell(row=1, column=c).value for c in range(1, 10)]
        assert "Entity Code" in headers
        assert "VAT ID" in headers
        assert "Status" in headers


# ---------------------------------------------------------------------------
# VAT Returns Summary XLSX
# ---------------------------------------------------------------------------


class TestVATReturns:
    """Verify tc10eu_vat_returns_summary_fy2025.xlsx."""

    def test_xlsx_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_vat_returns_summary_fy2025.xlsx"
        assert path.exists()

    def test_xlsx_has_canary(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_vat_returns_summary_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        expected = canaries.canary_for("tc10eu_vat_returns_summary_fy2025")
        desc = wb.properties.description or ""
        assert expected in desc

    def test_return_count(self) -> None:
        """16 VAT return rows: 4 entities x 4 quarters."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_vat_returns_summary_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["VAT Returns Summary"]
        data_rows = sum(1 for r in range(2, ws.max_row + 1) if ws.cell(row=r, column=1).value)
        assert data_rows == 16, f"Expected 16 return rows, got {data_rows}"

    def test_ce_q4_pending_assessment(self) -> None:
        """CE Q4 should show 'Filed — Pending Assessment' (trap)."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_vat_returns_summary_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["VAT Returns Summary"]
        found = False
        for r in range(2, ws.max_row + 1):
            entity = ws.cell(row=r, column=1).value
            quarter = ws.cell(row=r, column=2).value
            status = ws.cell(row=r, column=8).value or ""
            if entity == "CE" and quarter == "Q4" and "Pending" in status:
                found = True
                break
        assert found, "CE Q4 'Pending Assessment' not found in VAT returns"

    def test_columns(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_vat_returns_summary_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["VAT Returns Summary"]
        headers = [ws.cell(row=1, column=c).value for c in range(1, 9)]
        assert "Entity" in headers
        assert "Filing Status" in headers


# ---------------------------------------------------------------------------
# EU VAT Rules Reference DOCX
# ---------------------------------------------------------------------------


class TestVATRulesReference:
    """Verify tc10eu_eu_vat_rules_reference.docx."""

    def test_docx_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_eu_vat_rules_reference.docx"
        assert path.exists()

    def test_docx_has_canary(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_eu_vat_rules_reference.docx"
        d = docx.Document(str(path))
        expected = canaries.canary_for("tc10eu_eu_vat_rules_reference")
        comments = d.core_properties.comments or ""
        assert expected in comments, f"Canary {expected} not in DOCX comments"

    def test_has_key_sections(self) -> None:
        """DOCX should cover Art. 138, Art. 196, post-Brexit, and VAT PE."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_eu_vat_rules_reference.docx"
        d = docx.Document(str(path))
        full_text = "\n".join(p.text for p in d.paragraphs)
        assert "Article 138" in full_text
        assert "Article 196" in full_text
        assert "Brexit" in full_text or "third country" in full_text
        assert "fixed establishment" in full_text.lower() or "VAT PE" in full_text

    def test_vat_pe_income_tax_distinction(self) -> None:
        """DOCX must distinguish VAT PE from income tax PE."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "tc10eu_eu_vat_rules_reference.docx"
        d = docx.Document(str(path))
        full_text = "\n".join(p.text for p in d.paragraphs)
        assert "NOT the same" in full_text or "not the same" in full_text.lower()
        assert "Art. 5 OECD" in full_text or "OECD Model" in full_text


# ---------------------------------------------------------------------------
# ERR-EU-010 planted error
# ---------------------------------------------------------------------------


class TestErrEU010:
    """Verify ERR-EU-010 is registered correctly."""

    def test_error_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-EU-010" in errors.entries

    def test_error_type(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.get("ERR-EU-010")
        assert err.type == "classification_error"

    def test_error_severity(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.get("ERR-EU-010")
        assert err.severity == "material"

    def test_error_description_mentions_reverse_charge(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.get("ERR-EU-010")
        assert "reverse charge" in err.description.lower()

    def test_error_test_case(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.get("ERR-EU-010")
        assert "TC-10-EU" in err.which_test_cases_should_catch


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold_standards/TC-10-EU_gold.json."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-10-EU_gold.json"
        assert path.exists()

    def test_gold_has_expected_outputs(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-10-EU_gold.json"
        gold = json.loads(path.read_text())
        eo = gold["expected_outputs"]
        assert "transaction_vat_analysis" in eo
        assert "registration_assessment" in eo
        assert "missing_data_gaps" in eo
        assert "pe_risk" in eo
        assert "quantified_risks" in eo

    def test_gold_has_canary_verification(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-10-EU_gold.json"
        gold = json.loads(path.read_text())
        cv = gold["canary_verification"]
        assert "read_ic_sales" in cv
        assert "read_vat_registrations" in cv
        assert "read_vat_returns" in cv
        assert "read_vat_rules" in cv

    def test_gold_has_error_detection(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-10-EU_gold.json"
        gold = json.loads(path.read_text())
        assert "ERR-EU-010" in gold["error_detection"]

    def test_gold_has_scoring_hints(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-10-EU_gold.json"
        gold = json.loads(path.read_text())
        sh = gold["scoring_hints"]
        assert "correctness" in sh
        assert "completeness" in sh
        assert "eu_vat_framework" in sh
        assert "adversarial_elements" in sh

    def test_gold_transaction_analysis_covers_all_flows(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-10-EU_gold.json"
        gold = json.loads(path.read_text())
        tva = gold["expected_outputs"]["transaction_vat_analysis"]
        assert "cp_to_cm_raw_materials" in tva
        assert "cp_to_cd_finished_goods" in tva
        assert "ce_to_cp_mgmt_fee" in tva
        assert "ce_to_cm_mgmt_fee" in tva
        assert "ce_to_cd_mgmt_fee" in tva
        assert "cm_to_cp_royalty" in tva


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt.md and expected_behavior.md."""

    def test_prompt_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-10-EU/prompt.md"
        assert path.exists()

    def test_prompt_mentions_vat(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-10-EU/prompt.md"
        text = path.read_text()
        assert "VAT" in text
        assert "reverse charge" in text.lower() or "Reverse charge" in text

    def test_prompt_mentions_uk_post_brexit(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-10-EU/prompt.md"
        text = path.read_text()
        assert "post-Brexit" in text or "post Brexit" in text

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-10-EU/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_err_eu_010(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-10-EU/expected_behavior.md"
        text = path.read_text()
        assert "ERR-EU-010" in text

    def test_expected_behavior_mentions_missing_data(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-10-EU/expected_behavior.md"
        text = path.read_text()
        assert "Polish" in text
        assert "Italian" in text or "Italy" in text


# ---------------------------------------------------------------------------
# Model computation tests
# ---------------------------------------------------------------------------


class TestModelComputations:
    """Direct unit tests for vat_eu model functions."""

    def test_ic_sales_count(self) -> None:
        sales = generate_ic_sales_eu()
        assert len(sales) == 24, f"Expected 24 IC sales, got {len(sales)}"

    def test_ic_sales_has_error_row(self) -> None:
        sales = generate_ic_sales_eu()
        error_rows = [s for s in sales if s.is_error]
        assert len(error_rows) == 1
        assert error_rows[0].seller == "CE"
        assert error_rows[0].buyer == "CM"
        assert error_rows[0].invoice_vat_rate == "20%"

    def test_ic_sales_q2_partial(self) -> None:
        sales = generate_ic_sales_eu()
        q2_cp_cd = [s for s in sales if s.seller == "CP" and s.buyer == "CD" and s.quarter == "Q2"]
        assert len(q2_cp_cd) == 1
        assert q2_cp_cd[0].proof_of_dispatch == "Partial"

    def test_vat_registrations_count(self) -> None:
        regs = generate_vat_registrations()
        assert len(regs) == 5

    def test_vat_registrations_has_polish_pending(self) -> None:
        regs = generate_vat_registrations()
        pl = [r for r in regs if r.country == "Poland"]
        assert len(pl) == 1
        assert "Pending" in pl[0].status

    def test_vat_returns_count(self) -> None:
        returns = generate_vat_returns()
        assert len(returns) == 16

    def test_vat_returns_ce_q4_pending(self) -> None:
        returns = generate_vat_returns()
        ce_q4 = [r for r in returns if r.entity_code == "CE" and r.quarter == "Q4"]
        assert len(ce_q4) == 1
        assert "Pending" in ce_q4[0].filing_status

    def test_err_eu_010_amount(self) -> None:
        assert ERR_EU_010_AMOUNT > 0
        assert ERR_EU_010_VAT_CHARGED > 0
        assert ERR_EU_010_VAT_CHARGED == (ERR_EU_010_AMOUNT * VAT_RATES["FR"]).quantize(
            ERR_EU_010_AMOUNT.__class__("1"),
        )

    def test_canary_keys_sorted(self) -> None:
        assert ALL_CANARY_KEYS_TC10EU == sorted(ALL_CANARY_KEYS_TC10EU)

    def test_canary_keys_count(self) -> None:
        assert len(ALL_CANARY_KEYS_TC10EU) == 4


# ---------------------------------------------------------------------------
# Determinism
# ---------------------------------------------------------------------------


class TestDeterminism:
    """Verify identical outputs across runs."""

    def test_ic_sales_deterministic(self) -> None:
        run1 = generate_ic_sales_eu()
        run2 = generate_ic_sales_eu()
        assert len(run1) == len(run2)
        for a, b in zip(run1, run2):
            assert a == b

    def test_vat_registrations_deterministic(self) -> None:
        run1 = generate_vat_registrations()
        run2 = generate_vat_registrations()
        assert run1 == run2

    def test_vat_returns_deterministic(self) -> None:
        run1 = generate_vat_returns()
        run2 = generate_vat_returns()
        assert len(run1) == len(run2)
        for a, b in zip(run1, run2):
            assert a == b
