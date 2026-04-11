"""Tests for TC-16 — Engagement Letter Generation formatter.

Verifies:
- Fee calculation: audit $285,000 × 1.15 = $327,750; tax $95,000 + $12,000×3 = $131,000; total $458,750
- Merge field population (client_name, engagement_scope, fee_amount, payment_terms, start_date, partner_name)
- Template formatting preservation (engagement_letter_template.docx copied from templates/)
- Planted error ERR-013 (fee_schedule.xlsx audit total missing IPO multiplier)
- Canary embedding in all 3 input files
- Prompt and expected behavior markdown files
- Gold standard registration
"""

from __future__ import annotations

import json
import tempfile
from decimal import Decimal
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc16 import (
    _ADDITIONAL_ENTITIES,
    _AUDIT_BASE_FEE,
    _AUDIT_FEE,
    _CLIENT_NAME,
    _IPO_COMPLEXITY_MULTIPLIER,
    _PARTNER_NAME,
    _PAYMENT_TERMS,
    _PER_ENTITY_TAX_ADDER,
    _START_DATE,
    _TAX_BASE_FEE,
    _TAX_FEE,
    _TOTAL_FEE,
    emit_tc16,
)
from generator.formatters.templates import emit_templates
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc16 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None

_CANARY_KEYS = sorted([
    "tc16_client_profile",
    "tc16_fee_schedule",
    "tc16_engagement_template",
])


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc16_test_"))

        # TC-16 needs template canary keys too (emit_templates uses these)
        all_keys = sorted(set(_CANARY_KEYS) | {
            "engagement_letter_template",
            "cover_page_template",
            "formatting_guide",
        })
        _CANARIES = build_canary_registry(all_keys, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        # Templates must be emitted first (TC-16 copies engagement_letter_template.docx)
        emit_templates(_OUTPUT, _CANARIES, _MANIFEST)
        emit_tc16(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit TC-16 gold standard (registered via @register_gold)
        emit_gold("TC-16", _CANARIES, _ERRORS, _OUTPUT / "gold_standards")

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


# ---------------------------------------------------------------------------
# Constants verification (gold standard values from prompt.md)
# ---------------------------------------------------------------------------

_INPUT_DIR = "test_cases/TC-16/input_files"


class TestFeeConstants:
    """Verify that the formatter constants match prompt.md gold standard."""

    def test_audit_base_fee(self) -> None:
        assert _AUDIT_BASE_FEE == Decimal("285000")

    def test_ipo_multiplier(self) -> None:
        assert _IPO_COMPLEXITY_MULTIPLIER == Decimal("1.15")

    def test_audit_fee(self) -> None:
        assert _AUDIT_FEE == Decimal("285000") * Decimal("1.15")
        assert _AUDIT_FEE == Decimal("327750")

    def test_tax_base_fee(self) -> None:
        assert _TAX_BASE_FEE == Decimal("95000")

    def test_per_entity_adder(self) -> None:
        assert _PER_ENTITY_TAX_ADDER == Decimal("12000")

    def test_additional_entities(self) -> None:
        assert _ADDITIONAL_ENTITIES == 3

    def test_tax_fee(self) -> None:
        assert _TAX_FEE == Decimal("95000") + Decimal("12000") * 3
        assert _TAX_FEE == Decimal("131000")

    def test_total_fee(self) -> None:
        assert _TOTAL_FEE == Decimal("327750") + Decimal("131000")
        assert _TOTAL_FEE == Decimal("458750")


# ---------------------------------------------------------------------------
# Input file existence and format validity
# ---------------------------------------------------------------------------


class TestInputFilesExist:
    """All 3 input files must be generated."""

    def test_client_profile_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        assert (out / _INPUT_DIR / "client_profile.docx").exists()

    def test_fee_schedule_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        assert (out / _INPUT_DIR / "fee_schedule.xlsx").exists()

    def test_engagement_letter_template_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        assert (out / _INPUT_DIR / "engagement_letter_template.docx").exists()

    def test_client_profile_opens_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile.docx"))
        assert len(doc.paragraphs) > 0

    def test_fee_schedule_opens_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule.xlsx", data_only=True)
        assert wb.sheetnames

    def test_engagement_template_opens_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template.docx"))
        assert len(doc.paragraphs) > 0


# ---------------------------------------------------------------------------
# Client profile content
# ---------------------------------------------------------------------------


class TestClientProfile:
    """client_profile.docx must contain entity count, revenue tier, complexity factors."""

    def test_has_company_name(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert _CLIENT_NAME in text

    def test_has_revenue_tier(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "$100M" in text and "$500M" in text

    def test_has_entity_count(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "3 subsidiaries" in text

    def test_has_ipo_complexity(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "IPO" in text

    def test_has_key_contacts(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Robert Chen" in text
        assert "Maria Santos" in text

    def test_has_entity_table(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile.docx"))
        assert len(doc.tables) >= 1, "Client profile should have an entity table"
        # Table should have 5 rows (header + 4 entities)
        table = doc.tables[0]
        assert len(table.rows) == 5


# ---------------------------------------------------------------------------
# Fee schedule content
# ---------------------------------------------------------------------------


class TestFeeSchedule:
    """fee_schedule.xlsx must have fee matrix with correct structure and values."""

    def test_has_fee_schedule_sheet(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule.xlsx", data_only=True)
        assert "Fee Schedule" in wb.sheetnames

    def test_header_row(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        headers = [ws.cell(row=1, column=c).value for c in range(1, 7)]
        assert "Service Type" in headers
        assert "Revenue Tier" in headers
        assert "Base Fee" in headers
        assert "Per-Entity Adder" in headers
        assert "Complexity Multiplier" in headers

    def test_audit_row_for_cascade_tier(self) -> None:
        """Cascade's audit row: $100M–$500M tier, 4–6 entities, base $285,000."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        found = False
        for row in range(2, ws.max_row + 1):
            stype = ws.cell(row=row, column=1).value
            tier = ws.cell(row=row, column=2).value
            ec = ws.cell(row=row, column=3).value
            base = ws.cell(row=row, column=4).value
            if stype == "Audit" and tier == "$100M–$500M" and ec == "4–6":
                assert base == 285000, f"Audit base fee should be 285000, got {base}"
                found = True
                break
        assert found, "Missing Audit row for $100M–$500M / 4–6 entities"

    def test_tax_row_for_cascade_tier(self) -> None:
        """Cascade's tax row: $100M–$500M tier, 4–6 entities, base $95,000."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        found = False
        for row in range(2, ws.max_row + 1):
            stype = ws.cell(row=row, column=1).value
            tier = ws.cell(row=row, column=2).value
            ec = ws.cell(row=row, column=3).value
            base = ws.cell(row=row, column=4).value
            adder = ws.cell(row=row, column=5).value
            if stype and "Tax" in stype and tier == "$100M–$500M" and ec == "4–6":
                assert base == 95000, f"Tax base fee should be 95000, got {base}"
                assert adder == 12000, f"Tax per-entity adder should be 12000, got {adder}"
                found = True
                break
        assert found, "Missing Tax row for $100M–$500M / 4–6 entities"

    def test_has_ipo_complexity_note(self) -> None:
        """Fee schedule must document the 1.15x IPO complexity multiplier."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        all_values = []
        for row in range(1, ws.max_row + 1):
            for col in range(1, ws.max_column + 1):
                val = ws.cell(row=row, column=col).value
                if val:
                    all_values.append(str(val))
        text = " ".join(all_values)
        assert "1.15" in text, "Fee schedule should mention 1.15x IPO multiplier"
        assert "IPO" in text, "Fee schedule should mention IPO readiness"

    def test_cascade_fee_summary_section(self) -> None:
        """Fee schedule has a Cascade Industries summary section."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        all_values = []
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if val:
                all_values.append(str(val))
        text = " ".join(all_values)
        assert "Cascade Industries" in text, "Fee schedule should have Cascade summary section"


# ---------------------------------------------------------------------------
# Planted error ERR-013
# ---------------------------------------------------------------------------


class TestPlantedErrorERR013:
    """ERR-013: fee_schedule.xlsx audit total omits IPO complexity multiplier."""

    def test_error_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-013" in errors.entries, "ERR-013 must be registered"

    def test_error_type_is_formula_error(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-013"]
        assert err.type == "formula_error"

    def test_error_severity_is_material(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-013"]
        assert err.severity == "material"

    def test_error_file_points_to_fee_schedule(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-013"]
        assert "fee_schedule.xlsx" in err.file

    def test_error_catches_tc16(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-013"]
        assert "TC-16" in err.which_test_cases_should_catch

    def test_wrong_audit_total_in_summary(self) -> None:
        """The summary audit total should show $285,000 (wrong) not $327,750 (correct)."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        # Find the Cascade summary section and check the audit total
        wrong_total = int(_AUDIT_BASE_FEE)  # 285,000 (missing 1.15x)
        correct_total = int(_AUDIT_FEE)  # 327,750
        # Walk all cells looking for the wrong audit total in the summary
        found_wrong = False
        found_correct = False
        for row in range(1, ws.max_row + 1):
            for col in range(1, ws.max_column + 1):
                val = ws.cell(row=row, column=col).value
                if val == wrong_total:
                    # Check this is in the summary section (Audit row, total column)
                    row_label = ws.cell(row=row, column=1).value
                    if row_label == "Audit":
                        found_wrong = True
                if val == correct_total:
                    found_correct = True
        assert found_wrong, f"Summary should show wrong audit total ${wrong_total:,}"
        assert not found_correct, f"Summary should NOT contain correct audit total ${correct_total:,}"


# ---------------------------------------------------------------------------
# Engagement letter template (merge fields)
# ---------------------------------------------------------------------------


class TestEngagementTemplate:
    """engagement_letter_template.docx must have merge field placeholders."""

    def test_has_merge_field_placeholders(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Check for <<field_name>> format merge fields
        for field in ["client_name", "fee_amount", "payment_terms", "start_date", "partner_name"]:
            assert f"<<{field}>>" in text, f"Missing merge field <<{field}>>"

    def test_has_engagement_scope_field(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "<<engagement_scope>>" in text

    def test_template_sections_present(self) -> None:
        """Template must contain key engagement letter sections."""
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        for section in [
            "Engagement Scope",
            "Fees and Billing",
            "Engagement Period",
            "Responsibilities",
            "Confidentiality",
            "Limitation of Liability",
            "ACCEPTED AND AGREED",
        ]:
            assert section in text, f"Missing template section: {section}"


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """All 3 input files must have canaries registered."""

    def test_all_3_canary_keys_registered(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            assert key in canaries.entries, f"Missing canary key: {key}"

    def test_canary_values_are_8_chars(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            canary = canaries.entries[key].canary
            assert len(canary) == 8, f"Canary for {key} is {len(canary)} chars: {canary}"

    def test_canaries_are_unique(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        values = [canaries.entries[k].canary for k in _CANARY_KEYS]
        assert len(values) == len(set(values)), "Duplicate canary values found"


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifestRegistration:
    """All TC-16 files must appear in the manifest."""

    def test_manifest_has_tc16_entries(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        manifest_data = json.loads((out / "manifest.json").read_text())
        tc16_entries = [e for e in manifest_data if "TC-16" in e.get("test_cases", [])]
        # 3 input files: client_profile.docx, fee_schedule.xlsx, engagement_letter_template.docx
        assert len(tc16_entries) >= 3, (
            f"Expected at least 3 TC-16 manifest entries, got {len(tc16_entries)}"
        )


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """TC-16 must have prompt.md and expected_behavior.md."""

    def test_prompt_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-16/prompt.md"
        assert path.exists()
        text = path.read_text()
        assert "engagement letter" in text.lower()

    def test_prompt_mentions_fee_schedule(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16/prompt.md").read_text()
        assert "fee schedule" in text.lower()

    def test_prompt_mentions_ipo_complexity(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16/prompt.md").read_text()
        assert "1.15" in text

    def test_prompt_mentions_sarah_chen(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16/prompt.md").read_text()
        assert "Sarah Chen" in text

    def test_expected_behavior_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-16/expected_behavior.md"
        assert path.exists()
        text = path.read_text()
        assert len(text) > 100

    def test_expected_behavior_fee_calculation(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16/expected_behavior.md").read_text()
        assert "$327,750" in text, "Expected behavior should mention audit fee $327,750"
        assert "$131,000" in text, "Expected behavior should mention tax fee $131,000"
        assert "$458,750" in text, "Expected behavior should mention total $458,750"

    def test_expected_behavior_merge_fields(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16/expected_behavior.md").read_text()
        assert "client_name" in text
        assert "fee_amount" in text
        assert "partner_name" in text

    def test_expected_behavior_template_preservation(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16/expected_behavior.md").read_text()
        assert "template" in text.lower() or "formatting" in text.lower()


# ---------------------------------------------------------------------------
# Gold standard registration
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """TC-16 gold standard must be registered with correct fee values."""

    def test_gold_json_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "gold_standards/TC-16_gold.json"
        assert path.exists()

    def test_gold_has_fee_calculation(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-16_gold.json").read_text())
        fee_calc = gold["expected_outputs"]["fee_calculation"]
        assert fee_calc["audit_fee"] == "$327,750"
        assert fee_calc["tax_fee"] == "$131,000"
        assert fee_calc["total_engagement_fee"] == "$458,750"

    def test_gold_has_merge_fields(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-16_gold.json").read_text())
        merge = gold["expected_outputs"]["merge_fields"]
        assert merge["client_name"] == _CLIENT_NAME
        assert merge["partner_name"] == _PARTNER_NAME
        assert merge["start_date"] == _START_DATE
        assert merge["payment_terms"] == _PAYMENT_TERMS

    def test_gold_has_error_detection(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-16_gold.json").read_text())
        assert "ERR-013" in gold["error_detection"]

    def test_gold_has_canary_verification(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-16_gold.json").read_text())
        assert "read_client_profile" in gold["canary_verification"]
        assert "read_fee_schedule" in gold["canary_verification"]
        assert "read_engagement_template" in gold["canary_verification"]

    def test_gold_has_template_sections(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-16_gold.json").read_text())
        sections = gold["expected_outputs"]["template_sections_preserved"]
        assert "Engagement Scope" in sections
        assert "Fees and Billing" in sections
        assert "Acceptance" in sections
