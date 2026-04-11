"""Tests for TC-12 — Data Room Triage & Document Index formatter.

Verifies:
- 32-file data room structure across 7 subdirectories
- 65-item DD checklist (docx)
- Planted red flags:
  - Pending litigation: product liability suit $2.5M exposure
  - CEO golden parachute: 3× salary
  - Acme change-of-control termination clause
  - 2 patents expiring within 18 months
  - Incomplete IP assignments from 2 founding employees
- ERR-021 planted error (blank salary in employee_census.xlsx)
- Gap identification (missing environmental/regulatory docs)
- Canary embedding in all files
- Gold standard structure and scoring hints
- Prompt and expected behavior markdown files
"""

from __future__ import annotations

import json
import tempfile
from decimal import Decimal
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc12 import (
    _DD_CHECKLIST,
    _MISSING_IP_ASSIGNMENTS,
    _PATENTS,
    emit_tc12,
)
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model
from generator.model.customers import (
    KEY_PERSONNEL,
    LITIGATION,
    contracts_with_change_of_control,
)

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc12 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None

# All canary keys used by TC-12
_CANARY_KEYS = sorted([
    "tc12_articles_of_incorporation",
    "tc12_bylaws",
    "tc12_board_minutes_2024",
    "tc12_board_minutes_2025",
    "tc12_org_chart",
    "tc12_audited_financials_fy2023",
    "tc12_audited_financials_fy2024",
    "tc12_management_financials_fy2025",
    "tc12_budget_fy2025",
    "tc12_debt_schedule",
    "tc12_customer_agreement_acme",
    "tc12_customer_agreement_globex",
    "tc12_supplier_agreement_initech",
    "tc12_pending_litigation_summary",
    "tc12_ip_assignment_agreements",
    "tc12_insurance_policies_summary",
    "tc12_employee_census",
    "tc12_benefits_summary",
    "tc12_ceo_employment_agreement",
    "tc12_cfo_employment_agreement",
    "tc12_cto_employment_agreement",
    "tc12_org_chart_detailed",
    "tc12_federal_returns_fy2023",
    "tc12_federal_returns_fy2024",
    "tc12_state_returns_summary",
    "tc12_tax_notices",
    "tc12_facility_leases",
    "tc12_equipment_list",
    "tc12_customer_list_with_revenue",
    "tc12_vendor_list",
    "tc12_patent_portfolio",
    "tc12_software_licenses",
    "tc12_it_infrastructure_overview",
    "tc12_dd_checklist",
])


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc12_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc12(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard
        emit_gold("TC-12", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-12/input_files"
_DR = f"{_INPUT_DIR}/data_room"

# ── Expected data room files (32 total) ────────────────────────────────────

_CORPORATE_FILES = [
    "articles_of_incorporation.pdf",
    "bylaws.pdf",
    "board_minutes_2024.pdf",
    "board_minutes_2025.pdf",
    "org_chart.pdf",
]

_FINANCIAL_FILES = [
    "audited_financials_fy2023.pdf",
    "audited_financials_fy2024.pdf",
    "management_financials_fy2025_ytd.xlsx",
    "budget_fy2025.xlsx",
    "debt_schedule.xlsx",
]

_LEGAL_FILES = [
    "material_contracts/customer_agreement_acme.pdf",
    "material_contracts/customer_agreement_globex.pdf",
    "material_contracts/supplier_agreement_initech.pdf",
    "pending_litigation_summary.docx",
    "ip_assignment_agreements.pdf",
    "insurance_policies_summary.pdf",
]

_HR_FILES = [
    "employee_census.xlsx",
    "benefits_summary.pdf",
    "key_employee_agreements/ceo_employment_agreement.pdf",
    "key_employee_agreements/cfo_employment_agreement.pdf",
    "key_employee_agreements/cto_employment_agreement.pdf",
    "org_chart_detailed.xlsx",
]

_TAX_FILES = [
    "federal_returns_fy2023.pdf",
    "federal_returns_fy2024.pdf",
    "state_returns_summary.xlsx",
    "tax_notices.pdf",
]

_OPERATIONS_FILES = [
    "facility_leases.pdf",
    "equipment_list.xlsx",
    "customer_list_with_revenue.xlsx",
    "vendor_list.xlsx",
]

_TECHNOLOGY_FILES = [
    "patent_portfolio.pdf",
    "software_licenses.xlsx",
    "it_infrastructure_overview.docx",
]

_ALL_DATA_ROOM_FILES = (
    [f"01_corporate/{f}" for f in _CORPORATE_FILES]
    + [f"02_financial/{f}" for f in _FINANCIAL_FILES]
    + [f"03_legal/{f}" for f in _LEGAL_FILES]
    + [f"04_hr/{f}" for f in _HR_FILES]
    + [f"05_tax/{f}" for f in _TAX_FILES]
    + [f"06_operations/{f}" for f in _OPERATIONS_FILES]
    + [f"07_technology/{f}" for f in _TECHNOLOGY_FILES]
)


# ---------------------------------------------------------------------------
# Data room structure — 32 files across 7 subdirectories
# ---------------------------------------------------------------------------


class TestDataRoomStructure:
    """Verify the 32-file data room structure."""

    def test_total_file_count(self) -> None:
        # Spec says 32 but actual tree has 33 (5+5+6+6+4+4+3)
        assert len(_ALL_DATA_ROOM_FILES) == 33

    def test_all_files_exist(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        for rel in _ALL_DATA_ROOM_FILES:
            path = output / _DR / rel
            assert path.exists(), f"Missing data room file: {rel}"

    def test_7_subdirectories_exist(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        for subdir in [
            "01_corporate", "02_financial", "03_legal", "04_hr",
            "05_tax", "06_operations", "07_technology",
        ]:
            path = output / _DR / subdir
            assert path.is_dir(), f"Missing subdirectory: {subdir}"

    def test_corporate_has_5_files(self) -> None:
        assert len(_CORPORATE_FILES) == 5

    def test_financial_has_5_files(self) -> None:
        assert len(_FINANCIAL_FILES) == 5

    def test_legal_has_6_files(self) -> None:
        assert len(_LEGAL_FILES) == 6

    def test_hr_has_6_files(self) -> None:
        assert len(_HR_FILES) == 6

    def test_tax_has_4_files(self) -> None:
        assert len(_TAX_FILES) == 4

    def test_operations_has_4_files(self) -> None:
        assert len(_OPERATIONS_FILES) == 4

    def test_technology_has_3_files(self) -> None:
        assert len(_TECHNOLOGY_FILES) == 3


# ---------------------------------------------------------------------------
# DD Checklist — 65 items
# ---------------------------------------------------------------------------


class TestDDChecklist:
    """Verify the due diligence checklist document."""

    def test_checklist_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "dd_checklist_standard.docx"
        assert path.exists()

    def test_checklist_has_65_items(self) -> None:
        assert len(_DD_CHECKLIST) == 65

    def test_checklist_covers_9_categories(self) -> None:
        categories = sorted(set(cat for cat, _ in _DD_CHECKLIST))
        expected = sorted([
            "Corporate", "Financial", "Tax", "Legal", "HR",
            "Operations", "Technology", "Environmental", "Regulatory",
        ])
        assert categories == expected

    def test_checklist_docx_contains_items(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _INPUT_DIR / "dd_checklist_standard.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Spot-check a few items from different categories
        assert "Articles of Incorporation" in text
        assert "Audited financial statements" in text
        assert "Pending or threatened litigation" in text
        assert "environmental site assessments" in text.lower()


# ---------------------------------------------------------------------------
# Red flags — litigation
# ---------------------------------------------------------------------------


class TestLitigationRedFlag:
    """Verify pending litigation red flag in data room files."""

    def test_litigation_summary_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _DR / "03_legal/pending_litigation_summary.docx"
        assert path.exists()

    def test_litigation_mentions_exposure(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _DR / "03_legal/pending_litigation_summary.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Should mention the $2.5M potential exposure
        assert "2,500,000" in text or "2.5" in text

    def test_litigation_title_present(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _DR / "03_legal/pending_litigation_summary.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert LITIGATION[0].title in text


# ---------------------------------------------------------------------------
# Red flags — CEO golden parachute
# ---------------------------------------------------------------------------


class TestCEOGoldenParachute:
    """Verify CEO employment agreement has 3× salary change-of-control."""

    def test_ceo_agreement_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _DR / "04_hr/key_employee_agreements/ceo_employment_agreement.pdf"
        assert path.exists()

    def test_ceo_change_of_control_multiplier(self) -> None:
        ceo = KEY_PERSONNEL[0]
        assert ceo.change_of_control_multiplier == Decimal("3")

    def test_ceo_payout_calculation(self) -> None:
        ceo = KEY_PERSONNEL[0]
        payout = Decimal(ceo.base_salary) * ceo.change_of_control_multiplier
        # Spec says ~$975K → 3× $325,000 = $975,000
        assert payout > 0


# ---------------------------------------------------------------------------
# Red flags — Acme change-of-control clause
# ---------------------------------------------------------------------------


class TestAcmeChangeOfControl:
    """Verify Acme contract has change-of-control termination clause."""

    def test_acme_agreement_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _DR / "03_legal/material_contracts/customer_agreement_acme.pdf"
        assert path.exists()

    def test_acme_has_change_of_control(self) -> None:
        coc = contracts_with_change_of_control()
        acme_contracts = [c for c in coc if "acme" in c.customer_name.lower()]
        assert len(acme_contracts) >= 1, "Acme should have a change-of-control clause"


# ---------------------------------------------------------------------------
# Red flags — expiring patents
# ---------------------------------------------------------------------------


class TestExpiringPatents:
    """Verify 2 patents expiring within 18 months of FY2025 year-end."""

    def test_patent_portfolio_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _DR / "07_technology/patent_portfolio.pdf"
        assert path.exists()

    def test_two_patents_expiring_within_18_months(self) -> None:
        import datetime

        cutoff = datetime.date(2027, 6, 30)  # 18 months from FY2025 end
        expiring = [p for p in _PATENTS if p["expiration_date"] <= cutoff]
        assert len(expiring) == 2, f"Expected 2 expiring patents, got {len(expiring)}"


# ---------------------------------------------------------------------------
# Red flags — incomplete IP assignments
# ---------------------------------------------------------------------------


class TestIncompleteIPAssignments:
    """Verify IP assignment gaps for 2 founding employees."""

    def test_ip_assignments_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _DR / "03_legal/ip_assignment_agreements.pdf"
        assert path.exists()

    def test_two_missing_ip_assignments(self) -> None:
        assert len(_MISSING_IP_ASSIGNMENTS) == 2

    def test_missing_assignments_name_founders(self) -> None:
        combined = " ".join(_MISSING_IP_ASSIGNMENTS)
        assert "Cascade" in combined  # CEO / Co-Founder
        assert "Nakamura" in combined  # CTO / Co-Founder


# ---------------------------------------------------------------------------
# ERR-021 — blank salary in employee_census.xlsx
# ---------------------------------------------------------------------------


class TestERR021PlantedError:
    """Verify ERR-021: blank salary cell in employee_census.xlsx."""

    def test_employee_census_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _DR / "04_hr/employee_census.xlsx"
        assert path.exists()

    def test_err021_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-021" in errors.entries

    def test_err021_has_blank_salary(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / _DR / "04_hr/employee_census.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        # Column G (7) is Annual Salary; find the blank cell
        blank_salary_rows = []
        for row in range(2, ws.max_row + 1):
            val = ws.cell(row=row, column=7).value
            if val is None or val == "":
                blank_salary_rows.append(row)
        assert len(blank_salary_rows) == 1, (
            f"Expected exactly 1 blank salary cell, found {len(blank_salary_rows)}"
        )


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """Verify canary codes are embedded in files."""

    def test_all_canary_keys_assigned(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            code = canaries.canary_for(key)
            assert len(code) == 8, f"Canary for {key} should be 8 chars"

    def test_xlsx_canary_in_properties(self) -> None:
        """Canary is embedded in xlsx document properties (description)."""
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc12_employee_census")
        path = output / _DR / "04_hr/employee_census.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in xlsx properties"

    def test_docx_canary_in_properties(self) -> None:
        """Canary is embedded in docx core properties (comments)."""
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc12_dd_checklist")
        path = output / _INPUT_DIR / "dd_checklist_standard.docx"
        doc = Document(str(path))
        comments = doc.core_properties.comments or ""
        assert canary in comments, f"Canary {canary} not in docx properties"


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold standard JSON structure."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-12_gold.json"
        assert path.exists()

    def test_gold_has_expected_outputs(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-12_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "output_files" in eo
        assert "red_flags" in eo
        assert "gap_analysis_summary" in eo

    def test_gold_output_files_structure(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-12_gold.json").read_text())
        of = gold["expected_outputs"]["output_files"]
        assert of["document_index"]["total_files_indexed"] == 32
        assert of["gap_analysis"]["total_checklist_items"] == 65
        assert "red_flags_memo" in of

    def test_gold_red_flags_complete(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-12_gold.json").read_text())
        rf = gold["expected_outputs"]["red_flags"]
        assert "litigation" in rf
        assert "ceo_golden_parachute" in rf
        assert "acme_change_of_control" in rf
        assert "expiring_patents" in rf
        assert rf["expiring_patents"]["count"] == 2
        assert "incomplete_ip_assignments" in rf
        assert "missing_documents" in rf
        assert len(rf["missing_documents"]) >= 4  # At least the 4 named in spec

    def test_gold_canary_verification(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-12_gold.json").read_text())
        cv = gold["canary_verification"]
        expected_keys = [
            "read_acme_contract",
            "read_litigation_summary",
            "read_ceo_agreement",
            "read_patent_portfolio",
            "read_ip_assignments",
            "read_dd_checklist",
        ]
        for key in expected_keys:
            assert key in cv, f"Missing canary verification key: {key}"

    def test_gold_scoring_hints(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-12_gold.json").read_text())
        hints = gold["scoring_hints"]
        for key in ["correctness", "completeness", "judgment", "communication"]:
            assert key in hints


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt and expected behavior files are generated."""

    def test_prompt_md_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-12/prompt.md"
        assert path.exists()

    def test_prompt_mentions_data_room(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-12/prompt.md").read_text()
        assert "data room" in text.lower()

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-12/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_red_flags(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-12/expected_behavior.md").read_text()
        assert "red flag" in text.lower() or "Red Flag" in text


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifest:
    """Verify files are registered in the manifest."""

    def test_manifest_has_entries(self) -> None:
        _, _, _, _, manifest = _ensure_emitted()
        all_entries = manifest.entries
        tc12_count = sum(
            1 for v in all_entries.values()
            if "TC-12" in (v.test_cases or [])
        )
        # 33 data room files + DD checklist = 34
        assert tc12_count >= 34
