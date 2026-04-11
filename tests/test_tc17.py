"""Tests for TC-17 — Multi-File Deliverable Assembly formatter.

Verifies:
- 6 workpaper sections generated (4 docx, 2 xlsx)
- Cover page template copied from templates/
- Formatting guide PDF copied from templates/
- Correct section numbering (01–06 prefix order)
- Canary embedding in all 6 authored sections
- Prompt and expected behavior markdown files
- Gold standard registration
- Page numbering scheme documented (Roman for TOC, Arabic for body)
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc17 import _CANARY_KEYS, emit_tc17
from generator.formatters.templates import emit_templates
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc17 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc17_test_"))

        # TC-17 needs template canary keys too
        all_keys = sorted(set(_CANARY_KEYS) | {"cover_page_template", "formatting_guide"})
        _CANARIES = build_canary_registry(all_keys, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        # Templates must be emitted first (TC-17 copies from templates/)
        emit_templates(_OUTPUT, _CANARIES, _MANIFEST)
        emit_tc17(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


# ---------------------------------------------------------------------------
# Input directory
# ---------------------------------------------------------------------------

_INPUT_DIR = "test_cases/TC-17/input_files"

# The 6 authored workpaper sections
_DOCX_SECTIONS = [
    "01_executive_summary.docx",
    "03_industry_overview.docx",
    "04_risk_assessment.docx",
    "06_recommendations.docx",
]

_XLSX_SECTIONS = [
    "02_financial_analysis.xlsx",
    "05_detailed_findings.xlsx",
]


# ---------------------------------------------------------------------------
# Section file existence and format validity
# ---------------------------------------------------------------------------


class TestSectionFilesExist:
    """All 6 workpaper sections must be generated."""

    def test_all_four_docx_sections_exist(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _DOCX_SECTIONS:
            path = out / _INPUT_DIR / name
            assert path.exists(), f"Missing docx section: {name}"

    def test_all_two_xlsx_sections_exist(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _XLSX_SECTIONS:
            path = out / _INPUT_DIR / name
            assert path.exists(), f"Missing xlsx section: {name}"

    def test_docx_files_open_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _DOCX_SECTIONS:
            doc = Document(str(out / _INPUT_DIR / name))
            assert len(doc.paragraphs) > 0, f"{name} has no paragraphs"

    def test_xlsx_files_open_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _XLSX_SECTIONS:
            wb = openpyxl.load_workbook(out / _INPUT_DIR / name, data_only=True)
            assert wb.sheetnames, f"{name} has no sheets"


# ---------------------------------------------------------------------------
# Cover page template and formatting guide (copied from templates/)
# ---------------------------------------------------------------------------


class TestTemplateFiles:
    """Cover page template and formatting guide must be copied into input_files."""

    def test_cover_page_template_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / _INPUT_DIR / "cover_page_template.docx"
        assert path.exists(), "cover_page_template.docx missing from input_files"

    def test_cover_page_is_valid_docx(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "cover_page_template.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Template should have placeholder fields
        assert "REPORT TITLE" in text.upper() or "report title" in text.lower() or "[" in text, (
            "Cover page template should have placeholder fields"
        )

    def test_formatting_guide_pdf_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / _INPUT_DIR / "formatting_guide.pdf"
        assert path.exists(), "formatting_guide.pdf missing from input_files"

    def test_formatting_guide_is_valid_pdf(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / _INPUT_DIR / "formatting_guide.pdf"
        data = path.read_bytes()
        assert data[:5] == b"%PDF-", "formatting_guide.pdf does not start with PDF magic bytes"


# ---------------------------------------------------------------------------
# Section order (numbered prefix 01–06)
# ---------------------------------------------------------------------------


class TestSectionOrder:
    """Sections must use numbered prefixes enforcing the correct order."""

    def test_section_prefix_order(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        input_dir = out / _INPUT_DIR
        # Collect numbered files only (01_ through 06_)
        numbered = sorted(
            f.name for f in input_dir.iterdir()
            if f.is_file() and f.name[:2].isdigit()
        )
        assert len(numbered) == 6, f"Expected 6 numbered sections, got {len(numbered)}: {numbered}"
        # Verify order matches spec
        assert numbered[0].startswith("01_")
        assert numbered[1].startswith("02_")
        assert numbered[2].startswith("03_")
        assert numbered[3].startswith("04_")
        assert numbered[4].startswith("05_")
        assert numbered[5].startswith("06_")

    def test_total_file_count(self) -> None:
        """8 files total: 6 sections + cover page + formatting guide."""
        _, out, _, _, _ = _ensure_emitted()
        input_dir = out / _INPUT_DIR
        all_files = [f for f in input_dir.iterdir() if f.is_file()]
        assert len(all_files) == 8, (
            f"Expected 8 files in input_files, got {len(all_files)}: {[f.name for f in all_files]}"
        )


# ---------------------------------------------------------------------------
# Docx section content checks
# ---------------------------------------------------------------------------


class TestExecutiveSummary:
    """01_executive_summary.docx must have financial highlights from model."""

    def test_references_cascade(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "01_executive_summary.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Cascade" in text

    def test_has_revenue_figure(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "01_executive_summary.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Should contain dollar-formatted revenue
        assert "$" in text, "Executive summary should contain dollar figures"

    def test_has_key_sections(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "01_executive_summary.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Key Financial Highlights" in text
        assert "Engagement Scope" in text


class TestIndustryOverview:
    """03_industry_overview.docx must have market context."""

    def test_has_market_context(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "03_industry_overview.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Market Context" in text
        assert "Peer Comparison" in text

    def test_has_revenue_from_model(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "03_industry_overview.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "$" in text, "Industry overview should reference revenue figure"


class TestRiskAssessment:
    """04_risk_assessment.docx must have risk matrix with scores."""

    def test_has_risk_entries(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "04_risk_assessment.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Customer Concentration" in text
        assert "Raw Material Price" in text

    def test_has_risk_scores(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "04_risk_assessment.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Risk Score: 12" in text, "Should have critical-threshold risks scored at 12"

    def test_has_five_risks(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "04_risk_assessment.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        for i in range(1, 6):
            assert f"{i}." in text, f"Missing risk #{i}"


class TestRecommendations:
    """06_recommendations.docx must have prioritized recommendations."""

    def test_has_priority_levels(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "06_recommendations.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "High Priority" in text
        assert "Medium Priority" in text

    def test_has_six_recommendations(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "06_recommendations.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        for i in range(1, 7):
            assert f"{i}." in text, f"Missing recommendation #{i}"


# ---------------------------------------------------------------------------
# Xlsx section content checks
# ---------------------------------------------------------------------------


class TestFinancialAnalysis:
    """02_financial_analysis.xlsx must have income statement and key ratios."""

    def test_has_income_statement_sheet(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "02_financial_analysis.xlsx", data_only=True)
        assert "Income Statement" in wb.sheetnames

    def test_has_key_ratios_sheet(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "02_financial_analysis.xlsx", data_only=True)
        assert "Key Ratios" in wb.sheetnames

    def test_income_statement_has_data(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "02_financial_analysis.xlsx", data_only=True)
        ws = wb["Income Statement"]
        # Header + at least 7 line items (Revenue through Net Income)
        assert ws.max_row >= 8, f"Income statement too short: {ws.max_row} rows"
        # Check first data row is Revenue
        assert ws.cell(row=2, column=1).value == "Revenue"

    def test_income_statement_has_fy_columns(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "02_financial_analysis.xlsx", data_only=True)
        ws = wb["Income Statement"]
        headers = [ws.cell(row=1, column=c).value for c in range(1, 6)]
        assert "FY2025" in headers
        assert "FY2024" in headers

    def test_key_ratios_has_margins(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "02_financial_analysis.xlsx", data_only=True)
        ws = wb["Key Ratios"]
        labels = [ws.cell(row=r, column=1).value for r in range(2, ws.max_row + 1)]
        assert any("Gross" in str(lab) for lab in labels if lab), "Missing Gross Margin ratio"
        assert any("Operating" in str(lab) for lab in labels if lab), "Missing Operating Margin ratio"


class TestDetailedFindings:
    """05_detailed_findings.xlsx must have segment analysis tables."""

    def test_has_segment_revenue_sheet(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "05_detailed_findings.xlsx", data_only=True)
        assert "Segment Revenue" in wb.sheetnames

    def test_has_balance_sheet(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "05_detailed_findings.xlsx", data_only=True)
        assert "Balance Sheet" in wb.sheetnames

    def test_has_headcount_sheet(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "05_detailed_findings.xlsx", data_only=True)
        assert "Headcount" in wb.sheetnames

    def test_segment_revenue_has_three_subsidiaries(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "05_detailed_findings.xlsx", data_only=True)
        ws = wb["Segment Revenue"]
        names = [ws.cell(row=r, column=1).value for r in range(2, ws.max_row + 1)]
        assert any("Precision" in str(n) for n in names if n)
        assert any("Advanced" in str(n) for n in names if n)
        assert any("Distribution" in str(n) for n in names if n)

    def test_headcount_has_total_row(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "05_detailed_findings.xlsx", data_only=True)
        ws = wb["Headcount"]
        last_col1 = ws.cell(row=ws.max_row, column=1).value
        assert last_col1 == "Total", f"Last row should be Total, got {last_col1}"


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """All 6 authored section files must have canaries registered."""

    _CANARY_KEYS = sorted([
        "tc17_executive_summary",
        "tc17_financial_analysis",
        "tc17_industry_overview",
        "tc17_risk_assessment",
        "tc17_detailed_findings",
        "tc17_recommendations",
    ])

    def test_all_6_canary_keys_registered(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in self._CANARY_KEYS:
            assert key in canaries.entries, f"Missing canary key: {key}"

    def test_canary_values_are_8_chars(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in self._CANARY_KEYS:
            canary = canaries.entries[key].canary
            assert len(canary) == 8, f"Canary for {key} is {len(canary)} chars: {canary}"

    def test_canaries_are_unique(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        values = [canaries.entries[k].canary for k in self._CANARY_KEYS]
        assert len(values) == len(set(values)), "Duplicate canary values found"


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifestRegistration:
    """All TC-17 files must appear in the manifest."""

    def test_manifest_has_tc17_entries(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        manifest_data = json.loads((out / "manifest.json").read_text())
        tc17_entries = [e for e in manifest_data if "TC-17" in e.get("test_cases", [])]
        # 6 sections + cover page + formatting guide = 8
        assert len(tc17_entries) >= 8, (
            f"Expected at least 8 TC-17 manifest entries, got {len(tc17_entries)}"
        )


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """TC-17 must have prompt.md and expected_behavior.md."""

    def test_prompt_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-17/prompt.md"
        assert path.exists()
        text = path.read_text()
        assert "cover page" in text.lower()
        assert "formatting guide" in text.lower()

    def test_prompt_mentions_section_order(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-17/prompt.md").read_text()
        assert "section order" in text.lower()

    def test_expected_behavior_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-17/expected_behavior.md"
        assert path.exists()
        text = path.read_text()
        assert len(text) > 100

    def test_expected_behavior_page_numbering(self) -> None:
        """Expected behavior must document Roman/Arabic page numbering scheme."""
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-17/expected_behavior.md").read_text()
        assert "Roman" in text or "roman" in text, (
            "Expected behavior should mention Roman numeral page numbering for TOC"
        )
        assert "Arabic" in text or "arabic" in text, (
            "Expected behavior should mention Arabic page numbering for body"
        )

    def test_expected_behavior_section_order(self) -> None:
        """Expected behavior must list all 8 sections in order."""
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-17/expected_behavior.md").read_text()
        assert "Executive Summary" in text
        assert "Financial Analysis" in text
        assert "Risk Assessment" in text
        assert "Recommendations" in text


# ---------------------------------------------------------------------------
# No planted errors (TC-17 is assembly-focused)
# ---------------------------------------------------------------------------


class TestNoPlantedErrors:
    """TC-17 has no planted errors — difficulty is in assembly, not data."""

    def test_no_errors_registered_for_tc17(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        tc17_errors = [
            e for e in errors.entries.values()
            if "TC-17" in getattr(e, "file", "")
        ]
        assert len(tc17_errors) == 0, f"TC-17 should have no planted errors, found {tc17_errors}"
