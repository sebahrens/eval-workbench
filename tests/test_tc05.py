"""Tests for TC-05 — Audit Workpaper Memo — Accounts Receivable (Routine).

Verifies:
- ar_aging_fy2025.xlsx (aging by customer with current/30/60/90/120+ buckets)
- ar_confirmations_summary.xlsx (confirmation results with statuses)
- allowance_analysis.xlsx (bad debt history, reserve rates)
- workpaper_memo_template.docx (firm template copied and canary-embedded)
- ERR-016 planted error (date inconsistency in confirmations)
- ERR-020 planted error (missing data in aging)
- Canary embedding in all files
- Gold standard structure and scoring hints
- Prompt and expected behavior markdown files
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc05 import (
    _CONFIRMATION_STATUSES,
    emit_tc05,
)
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc05 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None

# All canary keys used by TC-05
_CANARY_KEYS = sorted([
    "ar_aging_fy2025",
    "ar_confirmations_summary",
    "allowance_analysis",
    "workpaper_memo_template",
])


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc05_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc05(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard
        emit_gold("TC-05", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-05/input_files"


# ---------------------------------------------------------------------------
# AR Aging Schedule
# ---------------------------------------------------------------------------


class TestARAgingSchedule:
    """Verify ar_aging_fy2025.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_aging_fy2025.xlsx"
        assert path.exists()

    def test_has_ar_aging_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_aging_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "AR Aging" in wb.sheetnames

    def test_has_header_row(self) -> None:
        """Header row should include all bucket columns."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_aging_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["AR Aging"]
        headers = []
        for col in range(1, ws.max_column + 1):
            val = ws.cell(row=5, column=col).value
            if isinstance(val, str):
                headers.append(val)
        assert "Customer ID" in headers
        assert "Customer Name" in headers
        assert "Total AR" in headers

    def test_has_customer_rows(self) -> None:
        """Should have multiple customer rows below the header."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_aging_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["AR Aging"]
        customer_count = 0
        for row in range(6, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str) and val.startswith("CUST-"):
                customer_count += 1
        assert customer_count >= 5, f"Expected ≥5 customers, got {customer_count}"

    def test_has_totals_row(self) -> None:
        """Should have a TOTAL row at the bottom."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_aging_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["AR Aging"]
        found_total = False
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=2).value
            if isinstance(val, str) and val == "TOTAL":
                found_total = True
                break
        assert found_total, "Missing TOTAL row in AR aging"

    def test_amounts_are_numeric(self) -> None:
        """Bucket amounts should be numeric values."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_aging_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["AR Aging"]
        numeric_count = 0
        for row in range(6, ws.max_row):
            for col in range(4, 10):  # Columns D through I (buckets + total)
                val = ws.cell(row=row, column=col).value
                if isinstance(val, (int, float)):
                    numeric_count += 1
        assert numeric_count >= 20, f"Expected ≥20 numeric cells, got {numeric_count}"

    def test_title_mentions_cascade(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_aging_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["AR Aging"]
        assert "Cascade" in str(ws["A1"].value)


# ---------------------------------------------------------------------------
# AR Confirmations Summary
# ---------------------------------------------------------------------------


class TestARConfirmationsSummary:
    """Verify ar_confirmations_summary.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_confirmations_summary.xlsx"
        assert path.exists()

    def test_has_confirmations_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_confirmations_summary.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Confirmations" in wb.sheetnames

    def test_has_customer_rows(self) -> None:
        """Should have multiple confirmation rows."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_confirmations_summary.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Confirmations"]
        row_count = 0
        for row in range(6, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str) and val.startswith("CUST-"):
                row_count += 1
        assert row_count >= 5, f"Expected ≥5 confirmation rows, got {row_count}"

    def test_all_statuses_present(self) -> None:
        """All three confirmation statuses should appear."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_confirmations_summary.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Confirmations"]
        statuses = set()
        for row in range(6, ws.max_row + 1):
            val = ws.cell(row=row, column=7).value  # Status column
            if isinstance(val, str) and val in _CONFIRMATION_STATUSES:
                statuses.add(val)
        assert statuses == set(_CONFIRMATION_STATUSES), (
            f"Expected all statuses {set(_CONFIRMATION_STATUSES)}, got {statuses}"
        )

    def test_has_summary_section(self) -> None:
        """Should have a Summary section with response rate."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_confirmations_summary.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Confirmations"]
        found_summary = False
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str) and val == "Summary":
                found_summary = True
                break
        assert found_summary, "Missing Summary section in confirmations"

    def test_confirmation_date_shown(self) -> None:
        """Title area should mention confirmation date."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_confirmations_summary.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Confirmations"]
        a3 = str(ws["A3"].value or "")
        assert "January 15, 2026" in a3 or "01/15/2026" in a3


# ---------------------------------------------------------------------------
# Allowance Analysis
# ---------------------------------------------------------------------------


class TestAllowanceAnalysis:
    """Verify allowance_analysis.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/allowance_analysis.xlsx"
        assert path.exists()

    def test_has_rollforward_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/allowance_analysis.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Allowance Rollforward" in wb.sheetnames

    def test_has_reserve_rates_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/allowance_analysis.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Reserve Rates" in wb.sheetnames

    def test_rollforward_has_data(self) -> None:
        """Rollforward sheet should have entity/year/balance rows."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/allowance_analysis.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Allowance Rollforward"]
        data_rows = 0
        for row in range(5, ws.max_row + 1):
            entity = ws.cell(row=row, column=1).value
            if isinstance(entity, str) and entity in ("PC", "AM", "DS"):
                data_rows += 1
        assert data_rows >= 3, f"Expected ≥3 entity rows, got {data_rows}"

    def test_reserve_rates_has_all_buckets(self) -> None:
        """Reserve rates sheet should have all 5 aging buckets."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/allowance_analysis.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Reserve Rates"]
        labels = set()
        for row in range(4, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str):
                labels.add(val)
        assert "Current (0-30 days)" in labels
        assert "31-60 Days" in labels
        assert "61-90 Days" in labels
        assert "91-120 Days" in labels
        assert "120+ Days" in labels

    def test_reserve_rates_are_populated(self) -> None:
        """Each bucket should have a rate percentage."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/allowance_analysis.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Reserve Rates"]
        rate_count = 0
        for row in range(4, ws.max_row + 1):
            val = ws.cell(row=row, column=2).value
            if isinstance(val, str) and "%" in val:
                rate_count += 1
        assert rate_count >= 5, f"Expected ≥5 rate entries, got {rate_count}"

    def test_consolidated_fy2025_total(self) -> None:
        """Should have a CONSOLIDATED FY2025 summary row."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/allowance_analysis.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Allowance Rollforward"]
        found = False
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str) and "CONSOLIDATED" in val:
                found = True
                break
        assert found, "Missing CONSOLIDATED FY2025 row"


# ---------------------------------------------------------------------------
# Workpaper Memo Template
# ---------------------------------------------------------------------------


class TestWorkpaperMemoTemplate:
    """Verify workpaper_memo_template.docx is copied and canary-embedded."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/workpaper_memo_template.docx"
        assert path.exists()

    def test_is_valid_docx(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/workpaper_memo_template.docx"
        doc = Document(str(path))
        # Should have at least some content
        assert len(doc.paragraphs) > 0

    def test_has_template_sections(self) -> None:
        """Template should mention key audit memo sections."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/workpaper_memo_template.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs).lower()
        assert "objective" in text
        assert "scope" in text or "procedures" in text


# ---------------------------------------------------------------------------
# ERR-016 — Date inconsistency in confirmations
# ---------------------------------------------------------------------------


class TestERR016PlantedError:
    """Verify ERR-016: date inconsistency in ar_confirmations_summary.xlsx."""

    def test_err016_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-016" in errors.entries

    def test_err016_is_date_inconsistency(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-016"]
        assert err.type == "date_inconsistency"

    def test_err016_in_confirmations_file(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-016"]
        assert "ar_confirmations_summary.xlsx" in err.file

    def test_err016_references_response_date(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-016"]
        assert "date" in err.description.lower() or "date" in err.location.lower()

    def test_err016_wrong_date_in_sheet(self) -> None:
        """The corrupted date should appear in the confirmations sheet."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_confirmations_summary.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Confirmations"]
        dates = []
        for row in range(6, ws.max_row + 1):
            val = ws.cell(row=row, column=5).value  # Response Date column
            if isinstance(val, str):
                dates.append(val)
        # At least one date should be the wrong date (02/15/2026)
        assert any("02/15/2026" in d for d in dates), (
            f"Expected corrupted date 02/15/2026 in confirmations, got dates: {dates}"
        )


# ---------------------------------------------------------------------------
# ERR-020 — Missing data in AR aging
# ---------------------------------------------------------------------------


class TestERR020PlantedError:
    """Verify ERR-020: missing data in ar_aging_fy2025.xlsx."""

    def test_err020_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-020" in errors.entries

    def test_err020_is_missing_data(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-020"]
        assert err.type == "missing_data"

    def test_err020_in_aging_file(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-020"]
        assert "ar_aging_fy2025.xlsx" in err.file

    def test_err020_blank_total_in_sheet(self) -> None:
        """The 3rd customer row (0-based idx 2) should have a blank Total AR."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/ar_aging_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["AR Aging"]
        # The target is row index 2 (0-based) → row 8 (1-based, since header=5, data starts at 6)
        target_row = 6 + 2  # row 8
        total_ar_val = ws.cell(row=target_row, column=9).value  # Column I = Total AR
        assert total_ar_val is None, (
            f"Expected blank Total AR for 3rd customer (row {target_row}), got {total_ar_val}"
        )


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """Verify canary codes are embedded in all TC-05 files."""

    def test_all_canary_keys_assigned(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            code = canaries.canary_for(key)
            assert len(code) == 8, f"Canary for {key} should be 8 chars"

    def test_aging_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("ar_aging_fy2025")
        path = output / f"{_INPUT_DIR}/ar_aging_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in AR aging properties"

    def test_confirmations_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("ar_confirmations_summary")
        path = output / f"{_INPUT_DIR}/ar_confirmations_summary.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in confirmations properties"

    def test_allowance_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("allowance_analysis")
        path = output / f"{_INPUT_DIR}/allowance_analysis.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in allowance properties"

    def test_template_canary_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("workpaper_memo_template")
        path = output / f"{_INPUT_DIR}/workpaper_memo_template.docx"
        doc = Document(str(path))
        comments = doc.core_properties.comments or ""
        assert canary in comments, f"Canary {canary} not in template properties"


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold standard JSON structure."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-05_gold.json"
        assert path.exists()

    def test_gold_has_expected_outputs(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-05_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "required_sections" in eo
        assert "ar_data" in eo
        assert "concentration" in eo
        assert "confirmations" in eo
        assert "allowance" in eo

    def test_gold_required_sections(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-05_gold.json").read_text())
        sections = gold["expected_outputs"]["required_sections"]
        for s in ["Objective", "Scope", "Procedures Performed", "Findings", "Conclusion"]:
            assert s in sections, f"Missing required section: {s}"

    def test_gold_ar_data_populated(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-05_gold.json").read_text())
        ar = gold["expected_outputs"]["ar_data"]
        assert ar["total_ar_balance"] > 0
        assert ar["customer_count"] >= 5

    def test_gold_concentration_populated(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-05_gold.json").read_text())
        conc = gold["expected_outputs"]["concentration"]
        assert "top_customer_name" in conc
        assert conc["top_customer_ar"] > 0
        assert conc["top_customer_pct_of_ar"] > 0

    def test_gold_confirmations_populated(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-05_gold.json").read_text())
        conf = gold["expected_outputs"]["confirmations"]
        assert conf["total_sent"] >= 5
        assert conf["agreed"] >= 1
        assert conf["response_rate_pct"] > 0

    def test_gold_file_type_docx(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-05_gold.json").read_text())
        assert gold["expected_outputs"]["file_type"] == "docx"

    def test_gold_canary_verification(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-05_gold.json").read_text())
        cv = gold["canary_verification"]
        for key in ["read_aging", "read_confirmations", "read_allowance", "read_template"]:
            assert key in cv, f"Missing canary verification key: {key}"

    def test_gold_error_detection(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-05_gold.json").read_text())
        assert "ERR-016" in gold["error_detection"]
        assert "ERR-020" in gold["error_detection"]

    def test_gold_scoring_hints(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-05_gold.json").read_text())
        hints = gold["scoring_hints"]
        for key in ["correctness", "completeness", "format_compliance", "communication"]:
            assert key in hints


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt and expected behavior files are generated."""

    def test_prompt_md_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-05/prompt.md"
        assert path.exists()

    def test_prompt_mentions_ar(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-05/prompt.md").read_text()
        assert "accounts receivable" in text.lower()

    def test_prompt_mentions_workpaper(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-05/prompt.md").read_text().lower()
        assert "workpaper" in text or "memo" in text

    def test_prompt_mentions_template(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-05/prompt.md").read_text().lower()
        assert "template" in text

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-05/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_template_usage(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-05/expected_behavior.md").read_text()
        assert "template" in text.lower()

    def test_expected_behavior_mentions_procedures(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-05/expected_behavior.md").read_text()
        assert "Procedures" in text

    def test_expected_behavior_mentions_conclusion(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-05/expected_behavior.md").read_text()
        assert "Conclusion" in text


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifest:
    """Verify files are registered in the manifest."""

    def test_manifest_has_tc05_entries(self) -> None:
        _, _, _, _, manifest = _ensure_emitted()
        tc05_count = sum(
            1 for v in manifest.entries.values()
            if "TC-05" in (v.test_cases or [])
        )
        # ar_aging + confirmations + allowance + template = 4
        assert tc05_count >= 4, f"Expected ≥4 TC-05 manifest entries, got {tc05_count}"
