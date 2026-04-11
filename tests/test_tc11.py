"""Tests for TC-11 — Quality of Earnings (Financial Due Diligence) formatter.

Verifies:
- 36 months of P&L data (FY2023–FY2025) with line-item detail
- Management adjustment categories (6 adjustments including 2 traps)
- Reported EBITDA ≈$28.4M (FY2025)
- Adjusted EBITDA ≈$29.3M (accepting legitimate adjustments only)
- 8 customer contracts as PDFs
- Top customer (Acme) at ~18% concentration
- ERR-011 planted error (recurring consulting fees misclassified as non-recurring)
- Canary embedding in all files
- Gold standard structure and scoring hints
- Prompt and expected behavior markdown files
"""

from __future__ import annotations

import json
import tempfile
from decimal import Decimal
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc11 import (
    _ADJUSTMENTS,
    _REPORTED_EBITDA_FY2025,
    _TOTAL_ADJUSTMENTS,
    emit_tc11,
)
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model
from generator.model.customers import (
    CONTRACTS,
    compute_customer_concentration,
    contracts_expiring_within,
)

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc11 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None

# All canary keys used by TC-11
_CANARY_KEYS = sorted(
    [
        "tc11_monthly_pl",
        "tc11_mgmt_adjustments",
        "tc11_interview_notes",
    ]
    + [f"tc11_contract_{i:03d}" for i in range(1, 9)]
)


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc11_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc11(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard
        emit_gold("TC-11", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-11/input_files"


# ---------------------------------------------------------------------------
# Monthly P&L — 36 months of data
# ---------------------------------------------------------------------------


class TestMonthlyPL:
    """Verify monthly P&L workbook structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/monthly_pl_fy2023_fy2024_fy2025.xlsx"
        assert path.exists()

    def test_36_months_of_data(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/monthly_pl_fy2023_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        # Data starts at row 5, year in column A
        months = set()
        for row in range(5, ws.max_row + 1):
            yr = ws.cell(row=row, column=1).value
            mo = ws.cell(row=row, column=2).value
            if isinstance(yr, int) and isinstance(mo, int):
                months.add((yr, mo))
        assert len(months) == 36, f"Expected 36 months, got {len(months)}"

    def test_covers_fy2023_through_fy2025(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/monthly_pl_fy2023_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        years = set()
        for row in range(5, ws.max_row + 1):
            yr = ws.cell(row=row, column=1).value
            if isinstance(yr, int):
                years.add(yr)
        assert {2023, 2024, 2025} <= years

    def test_has_line_item_detail(self) -> None:
        """P&L should have line-item detail, not just categories."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/monthly_pl_fy2023_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        # Read header row (row 4)
        headers = []
        for col in range(1, ws.max_column + 1):
            val = ws.cell(row=4, column=col).value
            if val:
                headers.append(val)
        # Should have detailed line items, not just totals
        assert "Product Sales" in headers
        assert "Services Revenue" in headers
        assert "Raw Materials" in headers
        assert "Direct Labor" in headers
        assert "Salaries & Benefits" in headers
        assert "EBITDA" in headers

    def test_reported_ebitda_fy2025_approx_28_4m(self) -> None:
        """FY2025 reported EBITDA should be approximately $28.4M."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/monthly_pl_fy2023_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        # Find EBITDA column
        ebitda_col = None
        for col in range(1, ws.max_column + 1):
            if ws.cell(row=4, column=col).value == "EBITDA":
                ebitda_col = col
                break
        assert ebitda_col is not None, "EBITDA column not found"

        # Sum FY2025 EBITDA
        total = 0
        for row in range(5, ws.max_row + 1):
            yr = ws.cell(row=row, column=1).value
            if yr == 2025:
                val = ws.cell(row=row, column=ebitda_col).value
                if isinstance(val, (int, float)):
                    total += val

        # Should be within 5% of $28.4M
        assert abs(total - 28_400_000) / 28_400_000 < 0.05, (
            f"FY2025 EBITDA {total:,.0f} not within 5% of $28.4M"
        )

    def test_subtotal_rows_present(self) -> None:
        """Should have annual subtotal rows."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/monthly_pl_fy2023_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        subtotals = []
        for row in range(5, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str) and "Total" in val:
                subtotals.append(val)
        assert len(subtotals) == 3, f"Expected 3 annual subtotals, got {len(subtotals)}"


# ---------------------------------------------------------------------------
# Management Adjustments — 6 proposed adjustments
# ---------------------------------------------------------------------------


class TestManagementAdjustments:
    """Verify management adjustments workbook."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_adjustments.xlsx"
        assert path.exists()

    def test_6_adjustments(self) -> None:
        assert len(_ADJUSTMENTS) == 6

    def test_total_adjustments_1_855m(self) -> None:
        assert _TOTAL_ADJUSTMENTS == Decimal("1_855_000")

    def test_adjustment_descriptions_match_spec(self) -> None:
        descriptions = [a["description"] for a in _ADJUSTMENTS]
        assert "Owner compensation above-market" in descriptions
        assert "One-time legal settlement" in descriptions
        assert "COVID-related PPP loan forgiveness" in descriptions
        assert "Non-recurring consulting fees" in descriptions
        assert "Run-rate adjustment for new Q4 customer" in descriptions
        assert "Facility relocation costs" in descriptions

    def test_adjustment_amounts_match_spec(self) -> None:
        amounts = {a["description"]: a["amount"] for a in _ADJUSTMENTS}
        assert amounts["Owner compensation above-market"] == Decimal("180_000")
        assert amounts["One-time legal settlement"] == Decimal("420_000")
        assert amounts["COVID-related PPP loan forgiveness"] == Decimal("250_000")
        assert amounts["Non-recurring consulting fees"] == Decimal("95_000")
        assert amounts["Run-rate adjustment for new Q4 customer"] == Decimal("600_000")
        assert amounts["Facility relocation costs"] == Decimal("310_000")

    def test_two_trap_adjustments(self) -> None:
        """Two adjustments should be traps (not truly non-recurring)."""
        traps = [a for a in _ADJUSTMENTS if not a["is_truly_nonrecurring"]]
        assert len(traps) == 2
        trap_ids = {a["id"] for a in traps}
        assert trap_ids == {"ADJ-004", "ADJ-005"}

    def test_xlsx_has_6_data_rows(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_adjustments.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        # Data starts at row 5, adjustment IDs in column A
        adj_ids = []
        for row in range(5, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str) and val.startswith("ADJ-"):
                adj_ids.append(val)
        assert len(adj_ids) == 6


# ---------------------------------------------------------------------------
# Customer Contracts — 8 PDFs
# ---------------------------------------------------------------------------


class TestCustomerContracts:
    """Verify 8 customer contract PDFs."""

    def test_8_contracts_defined(self) -> None:
        assert len(CONTRACTS) == 8

    def test_8_contract_pdfs_exist(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        contracts_dir = output / f"{_INPUT_DIR}/customer_contracts"
        assert contracts_dir.is_dir()
        pdfs = sorted(contracts_dir.glob("*.pdf"))
        assert len(pdfs) == 8

    def test_contract_filenames(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        contracts_dir = output / f"{_INPUT_DIR}/customer_contracts"
        for contract in CONTRACTS:
            path = contracts_dir / f"{contract.contract_id}.pdf"
            assert path.exists(), f"Missing contract: {contract.contract_id}"

    def test_top_customer_18_pct_concentration(self) -> None:
        """Acme should be ~18% of revenue per spec."""
        concentration = compute_customer_concentration(2025)
        top = concentration[0]
        pct = float(top.pct_of_consolidated) * 100
        assert 15 <= pct <= 21, f"Top customer at {pct:.1f}%, expected ~18%"
        assert "acme" in top.customer_name.lower()

    def test_contracts_expiring_within_12_months(self) -> None:
        """Per spec, contracts are expiring within 12 months."""
        expiring = contracts_expiring_within(12)
        assert len(expiring) >= 2, f"Expected ≥2 expiring contracts, got {len(expiring)}"

    def test_acme_has_change_of_control(self) -> None:
        acme = CONTRACTS[0]
        assert acme.change_of_control_clause is True
        assert "acme" in acme.customer_name.lower()


# ---------------------------------------------------------------------------
# Management Interview Notes — DOCX
# ---------------------------------------------------------------------------


class TestInterviewNotes:
    """Verify management interview notes document."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_interview_notes.docx"
        assert path.exists()

    def test_contains_key_topics(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_interview_notes.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Spec requires: business overview, growth drivers, customer
        # concentration, key personnel, and pending litigation
        assert "business overview" in text.lower()
        assert "growth" in text.lower()
        assert "concentration" in text.lower()
        assert "litigation" in text.lower() or "legal" in text.lower()

    def test_mentions_new_q4_customer_600k(self) -> None:
        """Should mention the Q4 new customer with $600K annualized."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_interview_notes.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "600" in text  # $600K or $600,000

    def test_mentions_acme_18_pct(self) -> None:
        """Should flag Acme at ~18% concentration."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_interview_notes.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "18%" in text or "18 percent" in text.lower()


# ---------------------------------------------------------------------------
# ERR-011 — recurring consulting fees misclassified
# ---------------------------------------------------------------------------


class TestERR011PlantedError:
    """Verify ERR-011: consulting fees classified as non-recurring but recurring."""

    def test_err011_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-011" in errors.entries

    def test_err011_is_classification_error(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-011"]
        assert err.type == "classification_error"

    def test_err011_references_adj004(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-011"]
        assert "ADJ-004" in err.description or "ADJ-004" in err.location

    def test_consulting_fee_visible_in_fy2024_and_fy2025(self) -> None:
        """The P&L should show professional fees bump in both FY2024 and FY2025,
        making the 'non-recurring' classification challengeable."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/monthly_pl_fy2023_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        # Find Professional Fees column
        prof_fees_col = None
        for col in range(1, ws.max_column + 1):
            if ws.cell(row=4, column=col).value == "Professional Fees":
                prof_fees_col = col
                break
        assert prof_fees_col is not None, "Professional Fees column not found"

        # Sum by year
        fees_by_year: dict[int, int] = {}
        for row in range(5, ws.max_row + 1):
            yr = ws.cell(row=row, column=1).value
            if isinstance(yr, int):
                val = ws.cell(row=row, column=prof_fees_col).value
                if isinstance(val, (int, float)):
                    fees_by_year[yr] = fees_by_year.get(yr, 0) + val

        # FY2024 and FY2025 should both show elevated professional fees
        # compared to FY2023 (the consulting fee bump)
        assert fees_by_year.get(2024, 0) > fees_by_year.get(2023, 0), (
            "FY2024 professional fees should be higher than FY2023 (recurring trap)"
        )
        assert fees_by_year.get(2025, 0) > fees_by_year.get(2023, 0), (
            "FY2025 professional fees should be higher than FY2023 (recurring trap)"
        )


# ---------------------------------------------------------------------------
# EBITDA gold standard values
# ---------------------------------------------------------------------------


class TestEBITDAGoldValues:
    """Verify spec-mandated EBITDA figures."""

    def test_reported_ebitda_target(self) -> None:
        assert _REPORTED_EBITDA_FY2025 == Decimal("28_400_000")

    def test_total_management_adjustments(self) -> None:
        assert _TOTAL_ADJUSTMENTS == Decimal("1_855_000")

    def test_adjusted_ebitda_approx_29_3m(self) -> None:
        """Properly adjusted EBITDA ≈$29.3M per spec gold standard.
        Accepted adjustments: ADJ-001 ($180K) + ADJ-002 ($420K) + ADJ-006 ($310K) = $910K.
        """
        # The actual EBITDA may differ slightly from exactly $28.4M due to
        # rounding, so we check the gold standard JSON for the precise figure
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-11_gold.json").read_text())
        adjusted = gold["expected_outputs"]["financial_metrics"]["adjusted_ebitda_fy2025"]
        # Should be around $29.3M (reported ~$28.4M + $910K accepted)
        assert abs(adjusted - 29_300_000) / 29_300_000 < 0.05, (
            f"Adjusted EBITDA {adjusted:,} not within 5% of $29.3M"
        )


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """Verify canary codes are embedded in files."""

    def test_all_canary_keys_assigned(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            code = canaries.canary_for(key)
            assert len(code) == 8, f"Canary for {key} should be 8 chars"

    def test_monthly_pl_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc11_monthly_pl")
        path = output / f"{_INPUT_DIR}/monthly_pl_fy2023_fy2024_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in monthly P&L properties"

    def test_mgmt_adjustments_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc11_mgmt_adjustments")
        path = output / f"{_INPUT_DIR}/management_adjustments.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in mgmt adjustments properties"

    def test_interview_notes_canary_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc11_interview_notes")
        path = output / f"{_INPUT_DIR}/management_interview_notes.docx"
        doc = Document(str(path))
        comments = doc.core_properties.comments or ""
        assert canary in comments, f"Canary {canary} not in interview notes properties"


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold standard JSON structure."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-11_gold.json"
        assert path.exists()

    def test_gold_has_expected_outputs(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-11_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "output_files" in eo
        assert "financial_metrics" in eo
        assert "adjustment_evaluation" in eo
        assert "customer_concentration" in eo
        assert "contract_renewal_risk" in eo

    def test_gold_output_files_structure(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-11_gold.json").read_text())
        of = gold["expected_outputs"]["output_files"]
        assert "analysis_workbook" in of
        assert "key_findings_memo" in of
        assert of["analysis_workbook"]["type"] == "xlsx"
        assert of["key_findings_memo"]["type"] == "docx"

    def test_gold_all_6_adjustments_evaluated(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-11_gold.json").read_text())
        adj_eval = gold["expected_outputs"]["adjustment_evaluation"]
        for adj_id in ["ADJ-001", "ADJ-002", "ADJ-003", "ADJ-004", "ADJ-005", "ADJ-006"]:
            assert adj_id in adj_eval, f"Missing adjustment evaluation: {adj_id}"

    def test_gold_adj004_rejected_as_recurring(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-11_gold.json").read_text())
        adj004 = gold["expected_outputs"]["adjustment_evaluation"]["ADJ-004"]
        assert adj004["action"] == "reject"
        assert "recurring" in adj004["reason"].lower()

    def test_gold_adj005_challenged_as_aggressive(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-11_gold.json").read_text())
        adj005 = gold["expected_outputs"]["adjustment_evaluation"]["ADJ-005"]
        assert adj005["action"] == "challenge"
        assert "aggressive" in adj005["reason"].lower() or "1 quarter" in adj005["reason"]

    def test_gold_canary_verification(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-11_gold.json").read_text())
        cv = gold["canary_verification"]
        for key in [
            "read_monthly_pl",
            "read_mgmt_adjustments",
            "read_interview_notes",
            "read_acme_contract",
        ]:
            assert key in cv, f"Missing canary verification key: {key}"

    def test_gold_error_detection(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-11_gold.json").read_text())
        assert "ERR-011" in gold["error_detection"]

    def test_gold_scoring_hints(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-11_gold.json").read_text())
        hints = gold["scoring_hints"]
        for key in ["correctness", "completeness", "judgment", "communication"]:
            assert key in hints


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt and expected behavior files are generated."""

    def test_prompt_md_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-11/prompt.md"
        assert path.exists()

    def test_prompt_mentions_quality_of_earnings(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-11/prompt.md").read_text()
        assert "quality of earnings" in text.lower()

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-11/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_adjustments(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-11/expected_behavior.md").read_text()
        assert "adjust" in text.lower()


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifest:
    """Verify files are registered in the manifest."""

    def test_manifest_has_tc11_entries(self) -> None:
        _, _, _, _, manifest = _ensure_emitted()
        tc11_count = sum(
            1 for v in manifest.entries.values()
            if "TC-11" in (v.test_cases or [])
        )
        # monthly_pl + mgmt_adjustments + 8 contracts + interview_notes = 11
        assert tc11_count >= 11, f"Expected ≥11 TC-11 manifest entries, got {tc11_count}"
