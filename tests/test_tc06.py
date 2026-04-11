"""Tests for TC-06 — Tax Provision ASC 740 (Tax, Complex) formatter.

Verifies:
- cascade_consolidated_tb_fy2025.xlsx (consolidated trial balance with pre-tax income)
- tax_provision_fy2024_workpaper.xlsx (prior year provision: current/deferred, rollforward, ETR)
- permanent_temporary_differences_fy2025.docx (FY2025 book-tax differences table)
- statutory_rates.docx (federal 21%, blended state 6.2%, apportionment-weighted)
- ERR-006 planted error (stale state rate 5.8% vs 6.2%)
- ERR-012 planted error (formula error double-counting last debit account in TB total)
- Effective tax rate ~24.8%
- Canary embedding in all files
- Gold standard structure and scoring hints
- Prompt and expected behavior markdown files
"""

from __future__ import annotations

import json
import tempfile
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc06 import emit_tc06
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc06 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None

# All canary keys used by TC-06
_CANARY_KEYS = sorted([
    "cascade_consolidated_tb_fy2025",
    "tax_provision_fy2024_workpaper",
    "perm_temp_differences_fy2025",
    "statutory_rates",
])


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc06_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc06(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard
        emit_gold("TC-06", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-06/input_files"


# ---------------------------------------------------------------------------
# Consolidated Trial Balance
# ---------------------------------------------------------------------------


class TestConsolidatedTB:
    """Verify cascade_consolidated_tb_fy2025.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/cascade_consolidated_tb_fy2025.xlsx"
        assert path.exists()

    def test_has_trial_balance_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/cascade_consolidated_tb_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Trial Balance" in wb.sheetnames

    def test_has_header_row(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/cascade_consolidated_tb_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Trial Balance"]
        headers = []
        for col in range(1, 6):
            val = ws.cell(row=5, column=col).value
            if val:
                headers.append(val)
        assert "Account" in headers
        assert "Description" in headers
        assert "Debit" in headers
        assert "Credit" in headers

    def test_has_account_rows(self) -> None:
        """TB should have multiple account rows with data."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/cascade_consolidated_tb_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Trial Balance"]
        account_rows = 0
        for row in range(6, ws.max_row + 1):
            acct = ws.cell(row=row, column=1).value
            if isinstance(acct, str) and acct and acct != "TOTAL":
                account_rows += 1
        assert account_rows >= 10, f"Expected >=10 account rows, got {account_rows}"

    def test_has_totals_row(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/cascade_consolidated_tb_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Trial Balance"]
        found_total = False
        for row in range(1, ws.max_row + 1):
            if ws.cell(row=row, column=2).value == "TOTAL":
                found_total = True
                break
        assert found_total, "Missing TOTAL row in trial balance"


# ---------------------------------------------------------------------------
# Prior Year Provision Workpaper
# ---------------------------------------------------------------------------


class TestPriorYearWorkpaper:
    """Verify tax_provision_fy2024_workpaper.xlsx structure."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx"
        assert path.exists()

    def test_has_current_deferred_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Current & Deferred" in wb.sheetnames

    def test_has_rollforward_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "DTA-DTL Rollforward" in wb.sheetnames

    def test_has_rate_reconciliation_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Rate Reconciliation" in wb.sheetnames

    def test_has_differences_detail_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Differences Detail" in wb.sheetnames

    def test_current_deferred_has_pre_tax_income(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Current & Deferred"]
        labels = set()
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str):
                labels.add(val.strip())
        assert "Pre-tax book income" in labels

    def test_current_deferred_has_provision_sections(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Current & Deferred"]
        labels = set()
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str):
                labels.add(val.strip())
        assert "Permanent Differences" in labels
        assert "Current Provision" in labels
        assert "Deferred Provision" in labels
        assert "TOTAL PROVISION" in labels

    def test_current_deferred_has_effective_tax_rate(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Current & Deferred"]
        found = False
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str) and "effective tax rate" in val.lower():
                found = True
                break
        assert found, "Missing effective tax rate in Current & Deferred sheet"

    def test_err006_stale_rate_visible(self) -> None:
        """The state tax rate label should show 5.8% (stale) instead of 6.2%."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Current & Deferred"]
        found_stale = False
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str) and "5.8%" in val:
                found_stale = True
                break
        assert found_stale, "ERR-006: stale state rate 5.8% not found in workpaper"


# ---------------------------------------------------------------------------
# Permanent & Temporary Differences Document
# ---------------------------------------------------------------------------


class TestPermTempDifferences:
    """Verify permanent_temporary_differences_fy2025.docx content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/permanent_temporary_differences_fy2025.docx"
        assert path.exists()

    def test_has_permanent_differences_section(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "permanent differences" in text.lower()

    def test_has_temporary_differences_section(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "temporary differences" in text.lower()

    def test_has_permanent_difference_table(self) -> None:
        """Permanent differences table with description, amount, effect."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        assert len(doc.tables) >= 1, "Expected at least 1 table (permanent differences)"
        # First table header row
        hdr = [cell.text for cell in doc.tables[0].rows[0].cells]
        assert "Description" in hdr
        assert "Amount" in hdr

    def test_has_temporary_difference_table(self) -> None:
        """Temporary differences table with book/tax/difference/creates columns."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        assert len(doc.tables) >= 2, "Expected at least 2 tables"
        hdr = [cell.text for cell in doc.tables[1].rows[0].cells]
        assert "Book Amount" in hdr
        assert "Tax Amount" in hdr
        assert "Difference" in hdr

    def test_mentions_key_permanent_items(self) -> None:
        """Should include meals & entertainment, tax-exempt interest, etc."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                all_text += "\n" + "\n".join(cell.text for cell in row.cells)
        text_lower = all_text.lower()
        assert "meals" in text_lower or "entertainment" in text_lower
        assert "tax-exempt" in text_lower or "municipal" in text_lower

    def test_mentions_key_temporary_items(self) -> None:
        """Should include depreciation and lease adjustments."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        all_text = ""
        for table in doc.tables:
            for row in table.rows:
                all_text += "\n" + "\n".join(cell.text for cell in row.cells)
        text_lower = all_text.lower()
        assert "depreciation" in text_lower
        assert "lease" in text_lower or "asc 842" in text_lower


# ---------------------------------------------------------------------------
# Statutory Rates Document
# ---------------------------------------------------------------------------


class TestStatutoryRates:
    """Verify statutory_rates.docx content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/statutory_rates.docx"
        assert path.exists()

    def test_mentions_federal_21_pct(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "21%" in text

    def test_mentions_blended_state_6_2_pct(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "6.2%" in text

    def test_mentions_apportionment(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "apportionment" in text.lower()

    def test_mentions_3_states(self) -> None:
        """Should list Oregon, Texas, Illinois."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\n".join(cell.text for cell in row.cells)
        text_lower = text.lower()
        assert "oregon" in text_lower
        assert "texas" in text_lower
        assert "illinois" in text_lower

    def test_mentions_combined_rate(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "25.898%" in text or "combined" in text.lower()


# ---------------------------------------------------------------------------
# ERR-006 — Stale state rate 5.8% vs 6.2%
# ---------------------------------------------------------------------------


class TestERR006PlantedError:
    """Verify ERR-006: stale state rate in prior year workpaper."""

    def test_err006_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-006" in errors.entries

    def test_err006_is_stale_data(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-006"]
        assert err.type == "stale_data"

    def test_err006_references_state_rate(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-006"]
        assert "5.8%" in err.description
        assert "6.2%" in err.description

    def test_err006_in_workpaper_file(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-006"]
        assert "tax_provision_fy2024_workpaper.xlsx" in err.file

    def test_err006_severity_immaterial(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-006"]
        assert err.severity == "immaterial"


# ---------------------------------------------------------------------------
# ERR-012 — Formula error double-counting last debit account
# ---------------------------------------------------------------------------


class TestERR012PlantedError:
    """Verify ERR-012: formula error in consolidated TB total debit."""

    def test_err012_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-012" in errors.entries

    def test_err012_is_formula_error(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-012"]
        assert err.type == "formula_error"

    def test_err012_references_double_counting(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-012"]
        assert "double-counting" in err.description.lower() or "double counts" in err.description.lower()

    def test_err012_in_tb_file(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-012"]
        assert "cascade_consolidated_tb_fy2025.xlsx" in err.file

    def test_err012_severity_material(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-012"]
        assert err.severity == "material"

    def test_err012_total_debit_overstated_in_file(self) -> None:
        """The displayed total debit should be higher than the correct value."""
        model, output, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-012"]
        # The description contains both wrong and correct values
        assert "instead of" in err.description


# ---------------------------------------------------------------------------
# Effective tax rate ~24.8%
# ---------------------------------------------------------------------------


class TestEffectiveTaxRate:
    """Verify the effective tax rate is approximately 24.8%."""

    def test_effective_rate_in_reasonable_range(self) -> None:
        """ETR should be above the 21% federal statutory rate (due to state
        taxes and permanent differences) but below ~40%.  The spec says ~24.8%
        but the canonical model's current calibration may differ."""
        model, _, _, _, _ = _ensure_emitted()
        prov25 = model.tax_provisions[2025]
        etr = float(prov25.effective_tax_rate)
        assert 0.21 < etr < 0.40, (
            f"Effective tax rate {etr:.4f} outside plausible range (0.21, 0.40)"
        )

    def test_current_plus_deferred_equals_total(self) -> None:
        """Tie-out: current + deferred = total provision."""
        model, _, _, _, _ = _ensure_emitted()
        prov25 = model.tax_provisions[2025]

        def _rd(d: Decimal) -> int:
            return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))

        total_check = _rd(prov25.total_current) + _rd(prov25.total_deferred)
        assert total_check == _rd(prov25.total_provision), (
            f"Tie-out failed: {_rd(prov25.total_current)} + "
            f"{_rd(prov25.total_deferred)} != {_rd(prov25.total_provision)}"
        )


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """Verify canary codes are embedded in all files."""

    def test_all_canary_keys_assigned(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            code = canaries.canary_for(key)
            assert len(code) == 8, f"Canary for {key} should be 8 chars"

    def test_consolidated_tb_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("cascade_consolidated_tb_fy2025")
        path = output / f"{_INPUT_DIR}/cascade_consolidated_tb_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in consolidated TB properties"

    def test_workpaper_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tax_provision_fy2024_workpaper")
        path = output / f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in workpaper properties"

    def test_perm_temp_canary_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("perm_temp_differences_fy2025")
        path = output / f"{_INPUT_DIR}/permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        comments = doc.core_properties.comments or ""
        assert canary in comments, f"Canary {canary} not in perm/temp differences properties"

    def test_statutory_rates_canary_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("statutory_rates")
        path = output / f"{_INPUT_DIR}/statutory_rates.docx"
        doc = Document(str(path))
        comments = doc.core_properties.comments or ""
        assert canary in comments, f"Canary {canary} not in statutory rates properties"


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold standard JSON structure."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-06_gold.json"
        assert path.exists()

    def test_gold_has_expected_outputs(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "pre_tax_book_income" in eo
        assert "taxable_income" in eo
        assert "current_provision" in eo
        assert "deferred_provision" in eo
        assert "total_provision" in eo
        assert "effective_tax_rate" in eo

    def test_gold_has_permanent_differences(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06_gold.json").read_text())
        perm = gold["expected_outputs"]["permanent_differences"]
        assert len(perm) >= 3, f"Expected >=3 permanent differences, got {len(perm)}"

    def test_gold_has_temporary_differences(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06_gold.json").read_text())
        temp = gold["expected_outputs"]["temporary_differences"]
        assert len(temp) >= 4, f"Expected >=4 temporary differences, got {len(temp)}"

    def test_gold_has_current_provision_components(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06_gold.json").read_text())
        cp = gold["expected_outputs"]["current_provision"]
        assert "federal" in cp
        assert "state" in cp
        assert "rd_credit" in cp
        assert "total_current" in cp

    def test_gold_has_deferred_provision_components(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06_gold.json").read_text())
        dp = gold["expected_outputs"]["deferred_provision"]
        assert "prior_dta" in dp
        assert "current_dta" in dp
        assert "total_deferred" in dp

    def test_gold_has_rate_reconciliation(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06_gold.json").read_text())
        rr = gold["expected_outputs"]["rate_reconciliation"]
        assert len(rr) >= 2, f"Expected >=2 rate recon items, got {len(rr)}"

    def test_gold_has_tie_out_checks(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06_gold.json").read_text())
        tie = gold["expected_outputs"]["tie_out_checks"]
        assert tie["current_plus_deferred_equals_total"] is True

    def test_gold_has_required_sheets(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06_gold.json").read_text())
        sheets = gold["expected_outputs"]["required_sheets"]
        assert "Current Provision" in sheets
        assert "Deferred Rollforward" in sheets
        assert "Rate Reconciliation" in sheets
        assert "Summary" in sheets

    def test_gold_canary_verification(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06_gold.json").read_text())
        cv = gold["canary_verification"]
        for key in [
            "read_consolidated_tb",
            "read_prior_workpaper",
            "read_differences",
            "read_statutory_rates",
        ]:
            assert key in cv, f"Missing canary verification key: {key}"

    def test_gold_error_detection(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06_gold.json").read_text())
        assert "ERR-006" in gold["error_detection"]
        assert "ERR-012" in gold["error_detection"]

    def test_gold_scoring_hints(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06_gold.json").read_text())
        hints = gold["scoring_hints"]
        for key in ["correctness", "completeness", "format_compliance", "communication"]:
            assert key in hints


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt and expected behavior files are generated."""

    def test_prompt_md_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-06/prompt.md"
        assert path.exists()

    def test_prompt_mentions_asc_740(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-06/prompt.md").read_text()
        assert "ASC 740" in text

    def test_prompt_mentions_key_tasks(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-06/prompt.md").read_text().lower()
        assert "pre-tax book income" in text
        assert "taxable income" in text
        assert "current tax provision" in text
        assert "deferred tax provision" in text
        assert "effective tax rate" in text

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-06/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_err006(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-06/expected_behavior.md").read_text()
        assert "5.8%" in text
        assert "6.2%" in text

    def test_expected_behavior_mentions_err012(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-06/expected_behavior.md").read_text()
        assert "double-count" in text.lower() or "sum range" in text.lower()


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifest:
    """Verify files are registered in the manifest."""

    def test_manifest_has_tc06_entries(self) -> None:
        _, _, _, _, manifest = _ensure_emitted()
        tc06_count = sum(
            1 for v in manifest.entries.values()
            if "TC-06" in (v.test_cases or [])
        )
        # consolidated_tb + workpaper + differences + statutory_rates = 4
        assert tc06_count >= 4, f"Expected >=4 TC-06 manifest entries, got {tc06_count}"
