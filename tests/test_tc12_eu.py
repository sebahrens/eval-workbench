"""Tests for TC-12-EU: European Diligence Data Room Variant."""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import openpyxl
import pytest
from docx import Document

from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc12_eu import ALL_CANARY_KEYS_TC12EU, emit_tc12_eu
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import build_model

# ---------------------------------------------------------------------------
# Module-level fixture: build once, share across all tests
# ---------------------------------------------------------------------------

_OUTPUT: Path | None = None
_ERRORS: ErrorRegistry | None = None


def _ensure_emitted() -> tuple[Path, ErrorRegistry]:
    global _OUTPUT, _ERRORS  # noqa: PLW0603
    if _OUTPUT is None:
        model = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc12eu_test_"))
        canaries = build_canary_registry(ALL_CANARY_KEYS_TC12EU, seed=42)
        _ERRORS = ErrorRegistry()
        manifest = Manifest(_OUTPUT)
        manifest.__enter__()
        emit_tc12_eu(model, _OUTPUT, canaries, _ERRORS, manifest)
        emit_gold("TC-12-EU", canaries, _ERRORS, _OUTPUT / "gold_standards", model=model)
        manifest.__exit__(None, None, None)
    return _OUTPUT, _ERRORS


_DR = "test_cases/TC-12-EU/input_files/data_room"


# ── 01_corporate ──────────────────────────────────────────────────────────


class TestCorporateFiles:

    @pytest.mark.parametrize("filename", [
        "ce_kvk_extract.pdf",
        "ce_articles_of_association.pdf",
        "cp_handelsregister_auszug.pdf",
        "cm_kbis_extract.pdf",
        "cd_companies_house_filing.pdf",
        "ce_board_minutes_2024.pdf",
        "ce_board_minutes_2025.pdf",
        "group_org_chart.pdf",
    ])
    def test_file_exists(self, filename: str) -> None:
        output, _ = _ensure_emitted()
        assert (output / _DR / "01_corporate" / filename).exists()


# ── 02_financial ──────────────────────────────────────────────────────────


class TestFinancialFiles:
    @pytest.mark.parametrize("filename", [
        "group_ifrs_financials_fy2023.pdf",
        "group_ifrs_financials_fy2024.pdf",
        "group_management_accounts_fy2025_ytd.xlsx",
        "group_budget_fy2025.xlsx",
        "group_debt_schedule.xlsx",
    ])
    def test_file_exists(self, filename: str) -> None:
        output, _ = _ensure_emitted()
        assert (output / _DR / "02_financial" / filename).exists()

    def test_management_accounts_has_canary(self) -> None:
        output, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(
            output / _DR / "02_financial/group_management_accounts_fy2025_ytd.xlsx",
        )
        desc = wb.properties.description or ""
        assert "CANARY:" in desc

    def test_debt_schedule_has_intercompany_loan(self) -> None:
        output, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(output / _DR / "02_financial/group_debt_schedule.xlsx")
        ws = wb.active
        found = False
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] and "Intercompany" in str(row[0]):
                found = True
                break
        assert found, "Intercompany loan not found in debt schedule"


# ── 03_legal ──────────────────────────────────────────────────────────────


class TestLegalFiles:
    @pytest.mark.parametrize("filename", [
        "material_contracts/customer_agreement_autohaus_mueller.pdf",
        "material_contracts/customer_agreement_renault_tier2.pdf",
        "material_contracts/supplier_agreement_thyssenkrupp_materials.pdf",
        "pending_litigation_summary.docx",
        "ip_assignment_agreements.pdf",
        "insurance_policies_summary.pdf",
    ])
    def test_file_exists(self, filename: str) -> None:
        output, _ = _ensure_emitted()
        assert (output / _DR / "03_legal" / filename).exists()

    def test_litigation_docx_has_canary(self) -> None:
        output, _ = _ensure_emitted()
        d = Document(output / _DR / "03_legal/pending_litigation_summary.docx")
        comments = d.core_properties.comments or ""
        assert "CANARY:" in comments

    def test_litigation_mentions_arbeitsgericht(self) -> None:
        output, _ = _ensure_emitted()
        d = Document(output / _DR / "03_legal/pending_litigation_summary.docx")
        text = "\n".join(p.text for p in d.paragraphs)
        assert "Arbeitsgericht" in text


# ── 04_hr ─────────────────────────────────────────────────────────────────


class TestHRFiles:
    @pytest.mark.parametrize("filename", [
        "group_employee_census.xlsx",
        "benefits_summary_by_country.pdf",
        "key_employee_agreements/managing_director_service_agreement_ce.pdf",
        "key_employee_agreements/finance_director_agreement_ce.pdf",
        "key_employee_agreements/technical_director_agreement_cm.pdf",
        "group_org_chart_detailed.xlsx",
    ])
    def test_file_exists(self, filename: str) -> None:
        output, _ = _ensure_emitted()
        assert (output / _DR / "04_hr" / filename).exists()

    def test_census_has_works_council_column(self) -> None:
        output, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(output / _DR / "04_hr/group_employee_census.xlsx")
        ws = wb.active
        headers = [ws.cell(row=1, column=c).value for c in range(1, 8)]
        assert "Works Council" in headers


# ── 05_tax ────────────────────────────────────────────────────────────────


class TestTaxFiles:
    @pytest.mark.parametrize("filename", [
        "cp_korperschaftsteuer_fy2023.pdf",
        "cp_korperschaftsteuer_fy2024.pdf",
        "cm_liasse_fiscale_fy2024.pdf",
        "cd_ct600_fy2024.pdf",
        "ce_vpb_aangifte_fy2024.pdf",
        "group_vat_returns_summary.xlsx",
        "tax_notices_and_assessments.pdf",
    ])
    def test_file_exists(self, filename: str) -> None:
        output, _ = _ensure_emitted()
        assert (output / _DR / "05_tax" / filename).exists()

    def test_vat_returns_has_canary(self) -> None:
        output, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(
            output / _DR / "05_tax/group_vat_returns_summary.xlsx",
        )
        desc = wb.properties.description or ""
        assert "CANARY:" in desc

    def test_vat_returns_row_count(self) -> None:
        output, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(
            output / _DR / "05_tax/group_vat_returns_summary.xlsx",
        )
        ws = wb.active
        data_rows = sum(1 for row in ws.iter_rows(min_row=2, values_only=True) if row[0])
        assert data_rows == 16, f"Expected 16 VAT return rows, got {data_rows}"

    def test_err_eu_012_present(self) -> None:
        """CP Q3 output VAT should be 847200 (the planted error)."""
        output, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(
            output / _DR / "05_tax/group_vat_returns_summary.xlsx",
        )
        ws = wb.active
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] == "CP" and row[1] == "Q3":
                assert row[2] == 847200, f"Expected planted error 847200, got {row[2]}"
                return
        pytest.fail("CP Q3 row not found in VAT returns")

    def test_cp_other_quarters_correct(self) -> None:
        """CP Q1/Q2/Q4 should all show 874200."""
        output, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(
            output / _DR / "05_tax/group_vat_returns_summary.xlsx",
        )
        ws = wb.active
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] == "CP" and row[1] in ("Q1", "Q2", "Q4"):
                assert row[2] == 874200, f"CP {row[1]} output VAT should be 874200, got {row[2]}"


# ── 06_operations ─────────────────────────────────────────────────────────


class TestOperationsFiles:
    @pytest.mark.parametrize("filename", [
        "facility_leases.pdf",
        "equipment_list.xlsx",
        "customer_list_with_revenue.xlsx",
        "vendor_list.xlsx",
    ])
    def test_file_exists(self, filename: str) -> None:
        output, _ = _ensure_emitted()
        assert (output / _DR / "06_operations" / filename).exists()


# ── 07_technology ─────────────────────────────────────────────────────────


class TestTechnologyFiles:
    @pytest.mark.parametrize("filename", [
        "patent_portfolio.pdf",
        "software_licenses.xlsx",
        "it_infrastructure_overview.docx",
    ])
    def test_file_exists(self, filename: str) -> None:
        output, _ = _ensure_emitted()
        assert (output / _DR / "07_technology" / filename).exists()

    def test_it_infra_has_canary(self) -> None:
        output, _ = _ensure_emitted()
        d = Document(output / _DR / "07_technology/it_infrastructure_overview.docx")
        comments = d.core_properties.comments or ""
        assert "CANARY:" in comments


# ── 08_compliance ─────────────────────────────────────────────────────────


class TestComplianceFiles:
    @pytest.mark.parametrize("filename", [
        "gdpr_data_processing_register.xlsx",
        "works_council_agreements_cp.pdf",
    ])
    def test_file_exists(self, filename: str) -> None:
        output, _ = _ensure_emitted()
        assert (output / _DR / "08_compliance" / filename).exists()

    def test_gdpr_register_has_missing_dpas(self) -> None:
        output, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(
            output / _DR / "08_compliance/gdpr_data_processing_register.xlsx",
        )
        ws = wb.active
        missing_count = 0
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row and len(row) >= 8 and row[7] and "NOT EXECUTED" in str(row[7]):
                missing_count += 1
        assert missing_count == 2, f"Expected 2 missing DPAs, got {missing_count}"


# ── DD Checklist ──────────────────────────────────────────────────────────


class TestDDChecklist:
    def test_checklist_exists(self) -> None:
        output, _ = _ensure_emitted()
        assert (output / "test_cases/TC-12-EU/input_files/dd_checklist_european.docx").exists()

    def test_checklist_has_canary(self) -> None:
        output, _ = _ensure_emitted()
        d = Document(output / "test_cases/TC-12-EU/input_files/dd_checklist_european.docx")
        comments = d.core_properties.comments or ""
        assert "CANARY:" in comments

    def test_checklist_has_72_items(self) -> None:
        output, _ = _ensure_emitted()
        d = Document(output / "test_cases/TC-12-EU/input_files/dd_checklist_european.docx")
        text = "\n".join(p.text for p in d.paragraphs)
        assert "72." in text


# ── Error Registration ────────────────────────────────────────────────────


class TestErrorRegistration:
    def test_err_eu_012_registered(self) -> None:
        _, errors = _ensure_emitted()
        assert "ERR-EU-012" in errors.entries

    def test_err_eu_012_type(self) -> None:
        _, errors = _ensure_emitted()
        assert errors.entries["ERR-EU-012"].type == "transposed_digits"

    def test_err_eu_012_severity(self) -> None:
        _, errors = _ensure_emitted()
        assert errors.entries["ERR-EU-012"].severity == "material"

    def test_err_eu_012_test_cases(self) -> None:
        _, errors = _ensure_emitted()
        assert "TC-12-EU" in errors.entries["ERR-EU-012"].which_test_cases_should_catch


# ── Gold Standard ─────────────────────────────────────────────────────────


class TestGoldStandard:
    def _load_gold(self) -> dict:
        output, _ = _ensure_emitted()
        path = output / "gold_standards/TC-12-EU_gold.json"
        assert path.exists(), "Gold standard JSON not generated"
        return json.loads(path.read_text())

    def test_gold_test_case(self) -> None:
        assert self._load_gold()["test_case"] == "TC-12-EU"

    def test_gold_has_36_files(self) -> None:
        gold = self._load_gold()
        assert gold["expected_outputs"]["output_files"]["document_index"]["total_files_indexed"] == 41

    def test_gold_has_72_checklist_items(self) -> None:
        gold = self._load_gold()
        assert gold["expected_outputs"]["output_files"]["gap_analysis"]["total_checklist_items"] == 72

    def test_gold_has_8_red_flags(self) -> None:
        flags = self._load_gold()["expected_outputs"]["red_flags"]
        assert "litigation" in flags
        assert "managing_director_coc" in flags
        assert "renault_change_of_control" in flags
        assert "missing_gdpr_dpas" in flags
        assert "ip_assignment_gaps" in flags
        assert "expiring_patents" in flags
        assert "betriebspruefung" in flags
        assert "works_council_consultation" in flags

    def test_gold_has_err_eu_012(self) -> None:
        assert "ERR-EU-012" in self._load_gold()["error_detection"]

    def test_gold_has_canary_verification(self) -> None:
        assert len(self._load_gold()["canary_verification"]) == 10

    def test_gold_has_scoring_hints(self) -> None:
        hints = self._load_gold()["scoring_hints"]
        assert len(hints) == 5
        assert "correctness" in hints

    def test_gold_scenario_pack(self) -> None:
        assert self._load_gold()["scenario_pack"] == "cascade_europe_ifrs"

    def test_gold_has_judgment_traps(self) -> None:
        assert len(self._load_gold()["judgment_traps"]) >= 6

    def test_gold_has_missing_documents(self) -> None:
        missing = self._load_gold()["expected_outputs"]["red_flags"]["missing_documents"]
        assert len(missing) == 6


# ── Prompt and Expected Behavior ──────────────────────────────────────────


class TestPromptAndExpectedBehavior:
    def test_prompt_exists(self) -> None:
        output, _ = _ensure_emitted()
        assert (output / "test_cases/TC-12-EU/prompt.md").exists()

    def test_prompt_mentions_european(self) -> None:
        output, _ = _ensure_emitted()
        text = (output / "test_cases/TC-12-EU/prompt.md").read_text()
        assert "European" in text

    def test_prompt_mentions_jurisdiction(self) -> None:
        output, _ = _ensure_emitted()
        text = (output / "test_cases/TC-12-EU/prompt.md").read_text()
        assert "Jurisdiction" in text

    def test_expected_behavior_exists(self) -> None:
        output, _ = _ensure_emitted()
        assert (output / "test_cases/TC-12-EU/expected_behavior.md").exists()

    def test_expected_behavior_mentions_err_eu_012(self) -> None:
        output, _ = _ensure_emitted()
        text = (output / "test_cases/TC-12-EU/expected_behavior.md").read_text()
        assert "ERR-EU-012" in text

    def test_expected_behavior_mentions_judgment_traps(self) -> None:
        output, _ = _ensure_emitted()
        text = (output / "test_cases/TC-12-EU/expected_behavior.md").read_text()
        assert "US-framing trap" in text
        assert "GDPR underweight" in text
        assert "Works council omission" in text


# ── File count ────────────────────────────────────────────────────────────


class TestFileCount:
    def test_total_data_room_files(self) -> None:
        """41 data room files across 8 categories.

        Design bead says 36 but detailed file inventory lists 41:
        01_corporate=8, 02_financial=5, 03_legal=6, 04_hr=6,
        05_tax=7, 06_operations=4, 07_technology=3, 08_compliance=2.
        """
        output, _ = _ensure_emitted()
        dr = output / _DR
        count = sum(1 for p in dr.rglob("*") if p.is_file())
        assert count == 41, f"Expected 41 data room files, got {count}"

    def test_total_input_files(self) -> None:
        """41 data room + 1 checklist = 42 input files."""
        output, _ = _ensure_emitted()
        input_dir = output / "test_cases/TC-12-EU/input_files"
        count = sum(1 for p in input_dir.rglob("*") if p.is_file())
        assert count == 42, f"Expected 42 input files, got {count}"


class TestColumnWidths:
    """Regression: spreadsheets must have explicit column widths to prevent ### overflow."""

    _AFFECTED_FILES = [
        ("02_financial/group_management_accounts_fy2025_ytd.xlsx", 7),
        ("02_financial/group_budget_fy2025.xlsx", 4),
        ("02_financial/group_debt_schedule.xlsx", 8),
        ("06_operations/customer_list_with_revenue.xlsx", 6),
        ("06_operations/vendor_list.xlsx", 6),
    ]

    @pytest.mark.parametrize("rel_path,expected_cols", _AFFECTED_FILES,
                             ids=[p for p, _ in _AFFECTED_FILES])
    def test_columns_have_explicit_widths(self, rel_path: str, expected_cols: int) -> None:
        """Every data column must have an explicit width > default (8)."""
        output, _ = _ensure_emitted()
        xlsx_path = output / _DR / rel_path
        wb = openpyxl.load_workbook(xlsx_path)
        ws = wb.active
        for col_idx in range(1, expected_cols + 1):
            letter = openpyxl.utils.get_column_letter(col_idx)
            dim = ws.column_dimensions.get(letter)
            assert dim is not None and dim.width is not None and dim.width > 8, (
                f"{rel_path} column {letter}: width not set or too narrow "
                f"(got {dim.width if dim else 'None'})"
            )
