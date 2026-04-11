"""Tests for TC-19 — Contract Risk Matrix (Legal Diligence, Advisory).

Verifies:
- 10 contract summary docx files in contracts/ directory
- 3 amendment docx files in amendments/ directory
- management_summary_memo.docx (deliberately stale re AMD-002)
- diligence_request_list.xlsx
- prompt.md and expected_behavior.md
- Judgment traps: change-of-control consent (LDI-001), MFN/exclusivity
  risk (LDI-002/LDI-003), management-vs-primary contradiction (LDI-002),
  missing amendment (LDI-003), source reference requirement (LDI-004)
- Canary embedding in all files
- Gold standard structure with evidence expectations and judgment traps
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc19 import (
    _CANARY_KEYS,
    emit_tc19,
)
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model
from generator.model.legal import (
    CONTRACT_AMENDMENTS,
    CONTRACT_CLAUSES,
    LEGAL_CONTRACTS,
    LEGAL_DILIGENCE_ISSUES,
    high_risk_clauses,
    missing_consent_issues,
    stale_summary_issues,
)

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc19 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc19_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc19(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard
        emit_gold("TC-19", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-19/input_files"


# ---------------------------------------------------------------------------
# Contract summaries — 10 docx files
# ---------------------------------------------------------------------------


class TestContractSummaries:
    """Verify 10 contract summary docx files in contracts/ directory."""

    def test_contracts_directory_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/contracts"
        assert path.is_dir()

    def test_10_contract_files_exist(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        for contract in LEGAL_CONTRACTS:
            fname = f"contract_{contract.contract_id.lower()}.docx"
            path = output / f"{_INPUT_DIR}/contracts/{fname}"
            assert path.exists(), f"Missing contract file: {fname}"

    def test_acme_contract_mentions_change_of_control(self) -> None:
        """LCTR-001 (Acme) should reference change-of-control provision."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/contracts/contract_lctr-001.docx"
        doc = Document(str(path))
        all_text = _extract_all_text(doc)
        assert "change" in all_text.lower() and "control" in all_text.lower(), (
            "LCTR-001 should mention change of control"
        )

    def test_techalloy_contract_mentions_mfn(self) -> None:
        """LCTR-003 (TechAlloy) should reference MFN clause."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/contracts/contract_lctr-003.docx"
        doc = Document(str(path))
        all_text = _extract_all_text(doc).lower()
        assert "most-favored-nation" in all_text or "mfn" in all_text

    def test_nextgen_contract_mentions_exclusivity(self) -> None:
        """LCTR-004 (NextGen) should reference exclusivity clause."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/contracts/contract_lctr-004.docx"
        doc = Document(str(path))
        all_text = _extract_all_text(doc).lower()
        assert "exclusivity" in all_text or "exclusive" in all_text

    def test_chemsource_contract_mentions_ip_license(self) -> None:
        """LCTR-010 (ChemSource) should reference IP license."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/contracts/contract_lctr-010.docx"
        doc = Document(str(path))
        all_text = _extract_all_text(doc).lower()
        assert "ip" in all_text or "license" in all_text or "intellectual property" in all_text


# ---------------------------------------------------------------------------
# Amendments — 3 docx files
# ---------------------------------------------------------------------------


class TestAmendments:
    """Verify 3 amendment docx files in amendments/ directory."""

    def test_amendments_directory_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/amendments"
        assert path.is_dir()

    def test_3_amendment_files_exist(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        for amd in CONTRACT_AMENDMENTS:
            fname = f"amendment_{amd.amendment_id.lower()}.docx"
            path = output / f"{_INPUT_DIR}/amendments/{fname}"
            assert path.exists(), f"Missing amendment file: {fname}"

    def test_amd002_mentions_mfn_threshold(self) -> None:
        """AMD-002 (TechAlloy side letter) should mention the 5% threshold."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/amendments/amendment_amd-002.docx"
        doc = Document(str(path))
        all_text = _extract_all_text(doc)
        assert "5%" in all_text, "AMD-002 should mention the 5% threshold"

    def test_amd003_mentions_thermal_barrier(self) -> None:
        """AMD-003 (NextGen scope expansion) should mention thermal barrier coatings."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/amendments/amendment_amd-003.docx"
        doc = Document(str(path))
        all_text = _extract_all_text(doc).lower()
        assert "thermal barrier" in all_text


# ---------------------------------------------------------------------------
# Management summary memo
# ---------------------------------------------------------------------------


class TestManagementSummaryMemo:
    """Verify management_summary_memo.docx — deliberately stale."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_summary_memo.docx"
        assert path.exists()

    def test_mentions_acme(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_summary_memo.docx"
        doc = Document(str(path))
        all_text = _extract_all_text(doc)
        assert "Acme" in all_text

    def test_stale_mfn_description(self) -> None:
        """Memo should describe MFN without the 5% threshold from AMD-002
        — stale relative to the current contractual position."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_summary_memo.docx"
        doc = Document(str(path))
        all_text = _extract_all_text(doc).lower()
        # Memo describes MFN clause but does NOT mention the 5% threshold
        assert "most-favored-nation" in all_text or "mfn" in all_text, (
            "Memo should mention MFN clause"
        )
        assert "5%" not in all_text, (
            "Memo should be stale — should NOT mention 5% threshold from AMD-002"
        )

    def test_does_not_mention_amd003(self) -> None:
        """Memo should NOT mention AMD-003 (exclusivity expansion) — stale."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/management_summary_memo.docx"
        doc = Document(str(path))
        all_text = _extract_all_text(doc)
        assert "AMD-003" not in all_text, "Memo should be stale — no mention of AMD-003"


# ---------------------------------------------------------------------------
# Diligence request list
# ---------------------------------------------------------------------------


class TestDiligenceRequestList:
    """Verify diligence_request_list.xlsx structure."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/diligence_request_list.xlsx"
        assert path.exists()

    def test_has_rows(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/diligence_request_list.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        assert ws.max_row >= 2, "Request list should have header + at least 1 row"


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt and expected behavior files are generated."""

    def test_prompt_md_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-19/prompt.md"
        assert path.exists()

    def test_prompt_mentions_risk_matrix(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-19/prompt.md").read_text().lower()
        assert "risk matrix" in text

    def test_prompt_mentions_change_of_control(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-19/prompt.md").read_text().lower()
        assert "consent" in text or "novation" in text or "change" in text

    def test_prompt_mentions_source_citation(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-19/prompt.md").read_text().lower()
        assert "source" in text

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-19/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_ldi_001(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-19/expected_behavior.md").read_text()
        assert "LDI-001" in text

    def test_expected_behavior_mentions_ldi_002(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-19/expected_behavior.md").read_text()
        assert "LDI-002" in text

    def test_expected_behavior_mentions_amd002(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-19/expected_behavior.md").read_text()
        assert "AMD-002" in text

    def test_expected_behavior_mentions_amd003(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-19/expected_behavior.md").read_text()
        assert "AMD-003" in text


# ---------------------------------------------------------------------------
# Judgment traps — verify canonical model data
# ---------------------------------------------------------------------------


class TestJudgmentTraps:
    """Verify judgment trap data from the canonical legal model."""

    def test_change_of_control_consent_ldi001(self) -> None:
        """LDI-001: Acme change-of-control, no consent obtained."""
        issue = LEGAL_DILIGENCE_ISSUES[0]
        assert issue.issue_id == "LDI-001"
        assert issue.contract_id == "LCTR-001"
        assert issue.clause_id == "CLS-001"
        assert issue.issue_type == "missing_consent"
        assert issue.severity == "high"

    def test_mfn_contradiction_ldi002(self) -> None:
        """LDI-002: MFN summary contradicts AMD-002 side letter."""
        issue = LEGAL_DILIGENCE_ISSUES[1]
        assert issue.issue_id == "LDI-002"
        assert issue.contract_id == "LCTR-003"
        assert issue.issue_type == "contradicts_summary"
        assert "AMD-002" in issue.source_refs

    def test_exclusivity_stale_ldi003(self) -> None:
        """LDI-003: Exclusivity scope expansion not in summary."""
        issue = LEGAL_DILIGENCE_ISSUES[2]
        assert issue.issue_id == "LDI-003"
        assert issue.contract_id == "LCTR-004"
        assert issue.issue_type == "stale_document"
        assert "AMD-003" in issue.source_refs

    def test_govt_assignment_ldi004(self) -> None:
        """LDI-004: Government subcontract assignment consent required."""
        issue = LEGAL_DILIGENCE_ISSUES[3]
        assert issue.issue_id == "LDI-004"
        assert issue.contract_id == "LCTR-004"
        assert issue.issue_type == "missing_consent"
        assert issue.severity == "high"

    def test_ip_dependency_ldi005(self) -> None:
        """LDI-005: ChemSource IP license strategic dependency."""
        issue = LEGAL_DILIGENCE_ISSUES[4]
        assert issue.issue_id == "LDI-005"
        assert issue.contract_id == "LCTR-010"
        assert issue.issue_type == "scope_boundary"

    def test_10_legal_contracts(self) -> None:
        assert len(LEGAL_CONTRACTS) == 10

    def test_3_amendments(self) -> None:
        assert len(CONTRACT_AMENDMENTS) == 3

    def test_5_diligence_issues(self) -> None:
        assert len(LEGAL_DILIGENCE_ISSUES) == 5

    def test_high_risk_clauses_include_coc_and_assignment(self) -> None:
        """CLS-001 (change-of-control) and CLS-006 (assignment) are high risk."""
        high = high_risk_clauses()
        ids = {c.clause_id for c in high}
        assert "CLS-001" in ids
        assert "CLS-006" in ids

    def test_missing_consent_issues_include_ldi001_ldi004(self) -> None:
        missing = missing_consent_issues()
        ids = {i.issue_id for i in missing}
        assert ids == {"LDI-001", "LDI-004"}

    def test_stale_summary_issues_include_ldi002_ldi003(self) -> None:
        stale = stale_summary_issues()
        ids = {i.issue_id for i in stale}
        assert ids == {"LDI-002", "LDI-003"}

    def test_amd002_changes_cls003(self) -> None:
        """AMD-002 should modify CLS-003 (MFN clause)."""
        amd = CONTRACT_AMENDMENTS[1]
        assert amd.amendment_id == "AMD-002"
        assert "CLS-003" in amd.changes_clause_ids

    def test_amd003_changes_cls005(self) -> None:
        """AMD-003 should modify CLS-005 (exclusivity clause)."""
        amd = CONTRACT_AMENDMENTS[2]
        assert amd.amendment_id == "AMD-003"
        assert "CLS-005" in amd.changes_clause_ids


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """Verify canary codes are embedded in all TC-19 files."""

    def test_all_canary_keys_assigned(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            code = canaries.canary_for(key)
            assert len(code) == 8, f"Canary for {key} should be 8 chars"

    def test_contract_canaries_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        for contract in LEGAL_CONTRACTS:
            ckey = f"tc19_contract_{contract.contract_id.lower().replace('-', '_')}"
            canary = canaries.canary_for(ckey)
            fname = f"contract_{contract.contract_id.lower()}.docx"
            path = output / f"{_INPUT_DIR}/contracts/{fname}"
            doc = Document(str(path))
            comments = doc.core_properties.comments or ""
            assert canary in comments, (
                f"Canary {canary} not in {fname} properties"
            )

    def test_amendment_canaries_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        for amd in CONTRACT_AMENDMENTS:
            ckey = f"tc19_amendment_{amd.amendment_id.lower().replace('-', '_')}"
            canary = canaries.canary_for(ckey)
            fname = f"amendment_{amd.amendment_id.lower()}.docx"
            path = output / f"{_INPUT_DIR}/amendments/{fname}"
            doc = Document(str(path))
            comments = doc.core_properties.comments or ""
            assert canary in comments, (
                f"Canary {canary} not in {fname} properties"
            )

    def test_memo_canary_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc19_management_summary_memo")
        path = output / f"{_INPUT_DIR}/management_summary_memo.docx"
        doc = Document(str(path))
        comments = doc.core_properties.comments or ""
        assert canary in comments, f"Canary {canary} not in memo properties"

    def test_request_list_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc19_diligence_request_list")
        path = output / f"{_INPUT_DIR}/diligence_request_list.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in request list properties"


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold standard JSON structure."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-19_gold.json"
        assert path.exists()

    def test_gold_has_expected_outputs(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-19_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert eo["contract_count"] == len(LEGAL_CONTRACTS)
        assert eo["clause_count"] == len(CONTRACT_CLAUSES)
        assert eo["amendment_count"] == len(CONTRACT_AMENDMENTS)

    def test_gold_risk_matrix_entries(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-19_gold.json").read_text())
        entries = gold["expected_outputs"]["risk_matrix_entries"]
        assert len(entries) == len(CONTRACT_CLAUSES)
        contract_ids = {e["contract_id"] for e in entries}
        # Should cover contracts that have clauses
        assert "LCTR-001" in contract_ids
        assert "LCTR-003" in contract_ids
        assert "LCTR-004" in contract_ids
        assert "LCTR-010" in contract_ids

    def test_gold_diligence_findings(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-19_gold.json").read_text())
        findings = gold["expected_outputs"]["diligence_findings"]
        assert len(findings) == len(LEGAL_DILIGENCE_ISSUES)
        ids = {f["issue_id"] for f in findings}
        for issue in LEGAL_DILIGENCE_ISSUES:
            assert issue.issue_id in ids

    def test_gold_consent_required_contracts(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-19_gold.json").read_text())
        consent = gold["expected_outputs"]["consent_required_contracts"]
        assert "LCTR-001" in consent
        assert "LCTR-004" in consent

    def test_gold_summary_contradictions(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-19_gold.json").read_text())
        contradictions = gold["expected_outputs"]["summary_contradictions"]
        assert "LDI-002" in contradictions
        assert "LDI-003" in contradictions

    def test_gold_canary_verification(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-19_gold.json").read_text())
        cv = gold["canary_verification"]
        # Should have entries for all 10 contracts + 3 amendments + memo + request list
        assert "read_contract_lctr_001" in cv
        assert "read_contract_lctr_010" in cv
        assert "read_amendment_amd_001" in cv
        assert "read_amendment_amd_003" in cv
        assert "read_management_summary_memo" in cv
        assert "read_diligence_request_list" in cv

    def test_gold_scoring_hints(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-19_gold.json").read_text())
        hints = gold["scoring_hints"]
        assert "correctness" in hints
        assert "completeness" in hints
        assert "source_reliance" in hints
        assert "professional_judgment" in hints

    def test_gold_judgment_traps(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-19_gold.json").read_text())
        traps = gold["judgment_traps"]
        trap_ids = {t["trap_id"] for t in traps}
        assert "JT-TC19-001" in trap_ids  # change-of-control consent
        assert "JT-TC19-002" in trap_ids  # MFN summary contradiction
        assert "JT-TC19-003" in trap_ids  # exclusivity stale summary
        assert "JT-TC19-004" in trap_ids  # govt assignment consent
        assert "JT-TC19-005" in trap_ids  # IP license dependency
        assert "JT-TC19-006" in trap_ids  # source citation requirement

    def test_gold_evidence_expectations(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-19_gold.json").read_text())
        ee = gold["evidence_expectations"]
        assert "risk_change_of_control" in ee
        assert "risk_mfn_contradiction" in ee
        assert "risk_exclusivity_expansion" in ee
        assert "risk_govt_assignment" in ee
        assert "risk_ip_dependency" in ee

    def test_gold_no_planted_errors(self) -> None:
        """TC-19 has no planted errors — error_detection should be empty."""
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-19_gold.json").read_text())
        assert gold["error_detection"] == {}


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifest:
    """Verify files are registered in the manifest."""

    def test_manifest_has_tc19_entries(self) -> None:
        _, _, _, _, manifest = _ensure_emitted()
        tc19_count = sum(
            1 for v in manifest.entries.values()
            if "TC-19" in (v.test_cases or [])
        )
        # 10 contracts + 3 amendments + 1 memo + 1 request list = 15
        assert tc19_count >= 15, f"Expected >= 15 TC-19 manifest entries, got {tc19_count}"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _extract_all_text(doc: Document) -> str:
    """Extract all text from paragraphs and tables."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)
