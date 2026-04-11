"""Tests for generator.writers.legal — legal document writer helpers."""

from __future__ import annotations

import openpyxl
import pytest
from docx import Document

from generator.canaries import build_registry
from generator.model.legal import (
    CONTRACT_AMENDMENTS,
    LEGAL_CONTRACTS,
    LEGAL_DILIGENCE_ISSUES,
    amendments_for_contract,
    clauses_for_contract,
)
from generator.writers.legal import (
    write_all_amendments,
    write_all_contract_summaries,
    write_amendment,
    write_contract_summary,
    write_diligence_request_list,
    write_management_summary_memo,
)

# ── Canary key fixtures ────────────────────────────────────────────────────

_CONTRACT_CANARY_KEYS = {
    c.contract_id: f"tc19_contract_{c.contract_id.lower()}"
    for c in LEGAL_CONTRACTS
}
_AMENDMENT_CANARY_KEYS = {
    a.amendment_id: f"tc19_amendment_{a.amendment_id.lower()}"
    for a in CONTRACT_AMENDMENTS
}
_MEMO_KEY = "tc19_mgmt_summary_memo"
_DILIGENCE_KEY = "tc19_diligence_requests"

_ALL_KEYS = sorted(
    list(_CONTRACT_CANARY_KEYS.values())
    + list(_AMENDMENT_CANARY_KEYS.values())
    + [_MEMO_KEY, _DILIGENCE_KEY]
)


@pytest.fixture
def canaries():
    return build_registry(_ALL_KEYS, seed=42)


# ── Contract summary tests ─────────────────────────────────────────────────


class TestContractSummary:
    def test_single_contract_creates_docx(self, tmp_path, canaries):
        contract = LEGAL_CONTRACTS[0]  # LCTR-001
        clauses = clauses_for_contract(contract.contract_id)
        amendments = amendments_for_contract(contract.contract_id)
        out = tmp_path / "contract_lctr-001.docx"
        key = _CONTRACT_CANARY_KEYS[contract.contract_id]

        loc = write_contract_summary(contract, clauses, amendments, out, canaries, key)

        assert out.exists()
        assert loc == "Core properties → comments"

    def test_canary_embedded_in_contract(self, tmp_path, canaries):
        contract = LEGAL_CONTRACTS[0]
        key = _CONTRACT_CANARY_KEYS[contract.contract_id]
        out = tmp_path / "contract.docx"
        write_contract_summary(
            contract, clauses_for_contract(contract.contract_id),
            amendments_for_contract(contract.contract_id),
            out, canaries, key,
        )

        doc = Document(str(out))
        assert canaries.canary_for(key) in doc.core_properties.comments

    def test_contract_contains_counterparty_info(self, tmp_path, canaries):
        contract = LEGAL_CONTRACTS[0]  # Acme
        key = _CONTRACT_CANARY_KEYS[contract.contract_id]
        out = tmp_path / "contract.docx"
        write_contract_summary(
            contract, clauses_for_contract(contract.contract_id),
            amendments_for_contract(contract.contract_id),
            out, canaries, key,
        )

        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text

        assert contract.counterparty_name in full_text
        assert contract.contract_id in full_text

    def test_all_contracts_written(self, tmp_path, canaries):
        locations = write_all_contract_summaries(
            tmp_path, canaries, _CONTRACT_CANARY_KEYS,
        )

        assert len(locations) == len(LEGAL_CONTRACTS)
        for contract in LEGAL_CONTRACTS:
            fname = f"contract_{contract.contract_id.lower()}.docx"
            assert (tmp_path / fname).exists()

    def test_deterministic_output(self, tmp_path, canaries):
        """Two runs produce byte-identical files."""
        dir1 = tmp_path / "run1"
        dir2 = tmp_path / "run2"
        dir1.mkdir()
        dir2.mkdir()

        canaries1 = build_registry(_ALL_KEYS, seed=42)
        canaries2 = build_registry(_ALL_KEYS, seed=42)

        contract = LEGAL_CONTRACTS[0]
        key = _CONTRACT_CANARY_KEYS[contract.contract_id]

        write_contract_summary(
            contract, clauses_for_contract(contract.contract_id),
            amendments_for_contract(contract.contract_id),
            dir1 / "c.docx", canaries1, key,
        )
        write_contract_summary(
            contract, clauses_for_contract(contract.contract_id),
            amendments_for_contract(contract.contract_id),
            dir2 / "c.docx", canaries2, key,
        )

        assert (dir1 / "c.docx").read_bytes() == (dir2 / "c.docx").read_bytes()

    def test_contract_with_clauses_includes_clause_info(self, tmp_path, canaries):
        contract = LEGAL_CONTRACTS[0]  # LCTR-001, has CLS-001, CLS-002
        key = _CONTRACT_CANARY_KEYS[contract.contract_id]
        clauses = clauses_for_contract(contract.contract_id)
        assert len(clauses) >= 1

        out = tmp_path / "c.docx"
        write_contract_summary(contract, clauses, [], out, canaries, key)

        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for cl in clauses:
            assert cl.clause_id in full_text

    def test_contract_with_amendments_includes_amendment_info(self, tmp_path, canaries):
        contract = LEGAL_CONTRACTS[0]  # LCTR-001, has AMD-001
        key = _CONTRACT_CANARY_KEYS[contract.contract_id]
        amendments = amendments_for_contract(contract.contract_id)
        assert len(amendments) >= 1

        out = tmp_path / "c.docx"
        write_contract_summary(contract, [], amendments, out, canaries, key)

        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for amd in amendments:
            assert amd.amendment_id in full_text


# ── Amendment tests ─────────────────────────────────────────────────────────


class TestAmendment:
    def test_single_amendment_creates_docx(self, tmp_path, canaries):
        amd = CONTRACT_AMENDMENTS[0]
        contract = next(c for c in LEGAL_CONTRACTS if c.contract_id == amd.contract_id)
        key = _AMENDMENT_CANARY_KEYS[amd.amendment_id]
        out = tmp_path / "amd.docx"

        loc = write_amendment(amd, contract, out, canaries, key)
        assert out.exists()
        assert loc == "Core properties → comments"

    def test_canary_embedded(self, tmp_path, canaries):
        amd = CONTRACT_AMENDMENTS[0]
        contract = next(c for c in LEGAL_CONTRACTS if c.contract_id == amd.contract_id)
        key = _AMENDMENT_CANARY_KEYS[amd.amendment_id]
        out = tmp_path / "amd.docx"

        write_amendment(amd, contract, out, canaries, key)
        doc = Document(str(out))
        assert canaries.canary_for(key) in doc.core_properties.comments

    def test_all_amendments_written(self, tmp_path, canaries):
        locations = write_all_amendments(tmp_path, canaries, _AMENDMENT_CANARY_KEYS)

        assert len(locations) == len(CONTRACT_AMENDMENTS)
        for amd in CONTRACT_AMENDMENTS:
            fname = f"amendment_{amd.amendment_id.lower()}.docx"
            assert (tmp_path / fname).exists()

    def test_amendment_contains_contract_reference(self, tmp_path, canaries):
        amd = CONTRACT_AMENDMENTS[0]
        contract = next(c for c in LEGAL_CONTRACTS if c.contract_id == amd.contract_id)
        key = _AMENDMENT_CANARY_KEYS[amd.amendment_id]
        out = tmp_path / "amd.docx"

        write_amendment(amd, contract, out, canaries, key)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert contract.contract_id in full_text
        assert contract.counterparty_name in full_text

    def test_deterministic_output(self, tmp_path, canaries):
        amd = CONTRACT_AMENDMENTS[1]  # AMD-002
        contract = next(c for c in LEGAL_CONTRACTS if c.contract_id == amd.contract_id)
        key = _AMENDMENT_CANARY_KEYS[amd.amendment_id]

        d1, d2 = tmp_path / "r1", tmp_path / "r2"
        d1.mkdir()
        d2.mkdir()

        c1 = build_registry(_ALL_KEYS, seed=42)
        c2 = build_registry(_ALL_KEYS, seed=42)

        write_amendment(amd, contract, d1 / "a.docx", c1, key)
        write_amendment(amd, contract, d2 / "a.docx", c2, key)
        assert (d1 / "a.docx").read_bytes() == (d2 / "a.docx").read_bytes()


# ── Management summary memo tests ──────────────────────────────────────────


class TestManagementSummaryMemo:
    def test_creates_docx(self, tmp_path, canaries):
        out = tmp_path / "memo.docx"
        loc = write_management_summary_memo(out, canaries, _MEMO_KEY)
        assert out.exists()
        assert loc == "Core properties → comments"

    def test_canary_embedded(self, tmp_path, canaries):
        out = tmp_path / "memo.docx"
        write_management_summary_memo(out, canaries, _MEMO_KEY)
        doc = Document(str(out))
        assert canaries.canary_for(_MEMO_KEY) in doc.core_properties.comments

    def test_contains_stale_mfn_description(self, tmp_path, canaries):
        """The memo must describe MFN in original terms (no AMD-002 threshold)."""
        out = tmp_path / "memo.docx"
        write_management_summary_memo(out, canaries, _MEMO_KEY)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # Should mention MFN and "any" (original terms)
        assert "most-favored-nation" in full_text.lower() or "MFN" in full_text
        # Should NOT mention the 5% threshold from AMD-002
        assert "5%" not in full_text

    def test_does_not_mention_amd003_scope_expansion(self, tmp_path, canaries):
        """The memo must not mention thermal barrier coatings from AMD-003."""
        out = tmp_path / "memo.docx"
        write_management_summary_memo(out, canaries, _MEMO_KEY)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "thermal barrier" not in full_text.lower()

    def test_deterministic(self, tmp_path, canaries):
        d1, d2 = tmp_path / "r1", tmp_path / "r2"
        d1.mkdir()
        d2.mkdir()
        c1 = build_registry(_ALL_KEYS, seed=42)
        c2 = build_registry(_ALL_KEYS, seed=42)

        write_management_summary_memo(d1 / "m.docx", c1, _MEMO_KEY)
        write_management_summary_memo(d2 / "m.docx", c2, _MEMO_KEY)
        assert (d1 / "m.docx").read_bytes() == (d2 / "m.docx").read_bytes()


# ── Diligence request list tests ───────────────────────────────────────────


class TestDiligenceRequestList:
    def test_creates_xlsx(self, tmp_path, canaries):
        out = tmp_path / "requests.xlsx"
        loc = write_diligence_request_list(out, canaries, _DILIGENCE_KEY)
        assert out.exists()
        assert loc == "Document properties → description"

    def test_canary_embedded(self, tmp_path, canaries):
        out = tmp_path / "requests.xlsx"
        write_diligence_request_list(out, canaries, _DILIGENCE_KEY)
        wb = openpyxl.load_workbook(str(out))
        assert canaries.canary_for(_DILIGENCE_KEY) in wb.properties.description

    def test_contains_all_issues(self, tmp_path, canaries):
        out = tmp_path / "requests.xlsx"
        write_diligence_request_list(out, canaries, _DILIGENCE_KEY)
        wb = openpyxl.load_workbook(str(out))
        ws = wb.active

        # Header + N data rows
        data_rows = list(ws.iter_rows(min_row=2, max_col=1, values_only=True))
        issue_ids = [r[0] for r in data_rows if r[0]]
        for issue in LEGAL_DILIGENCE_ISSUES:
            assert issue.issue_id in issue_ids

    def test_deterministic(self, tmp_path, canaries):
        d1, d2 = tmp_path / "r1", tmp_path / "r2"
        d1.mkdir()
        d2.mkdir()
        c1 = build_registry(_ALL_KEYS, seed=42)
        c2 = build_registry(_ALL_KEYS, seed=42)

        write_diligence_request_list(d1 / "r.xlsx", c1, _DILIGENCE_KEY)
        write_diligence_request_list(d2 / "r.xlsx", c2, _DILIGENCE_KEY)
        assert (d1 / "r.xlsx").read_bytes() == (d2 / "r.xlsx").read_bytes()
