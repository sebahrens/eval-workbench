"""Tests for generator.canaries — canary generation, registry, and quality gate.

The quality gate (TestCanaryQualityGate) runs the full generator and verifies
that every canary in canary_registry.json is findable in its target file.
Ref: bead synth-data-6f8, prompt.md §1.6.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

import pytest

from generate_test_suite import generate
from generator.canaries import (
    build_registry,
    embed_canary_csv_comment,
    embed_canary_docx,
    embed_canary_pdf_fpdf2,
    embed_canary_pdf_reportlab,
    embed_canary_xlsx,
    generate_canary,
)
from generator.config import load_config

# ---------------------------------------------------------------------------
# Basic canary generation
# ---------------------------------------------------------------------------

CANARY_RE = re.compile(r"^[A-Z0-9]{8}$")


class TestGenerateCanary:
    def test_format(self):
        """Canary is 8 uppercase-alphanumeric characters."""
        import random
        rng = random.Random(42)
        for _ in range(50):
            assert CANARY_RE.match(generate_canary(rng))

    def test_deterministic(self):
        """Same seed produces the same sequence."""
        import random
        a = [generate_canary(random.Random(99)) for _ in range(10)]
        b = [generate_canary(random.Random(99)) for _ in range(10)]
        assert a == b


# ---------------------------------------------------------------------------
# Registry builder
# ---------------------------------------------------------------------------

SAMPLE_KEYS = sorted([
    "cascade_tb_fy2025",
    "cascade_tb_fy2024",
    "employee_roster",
    "master_coa",
    "bank_statements",
])


class TestBuildRegistry:
    def test_all_keys_present(self):
        reg = build_registry(SAMPLE_KEYS)
        assert sorted(reg.entries) == SAMPLE_KEYS

    def test_canaries_unique(self):
        reg = build_registry(SAMPLE_KEYS)
        canaries = [e.canary for e in reg.entries.values()]
        assert len(set(canaries)) == len(canaries)

    def test_deterministic(self):
        r1 = build_registry(SAMPLE_KEYS, seed=42)
        r2 = build_registry(SAMPLE_KEYS, seed=42)
        for key in SAMPLE_KEYS:
            assert r1.canary_for(key) == r2.canary_for(key)

    def test_different_seed_different_canaries(self):
        r1 = build_registry(SAMPLE_KEYS, seed=42)
        r2 = build_registry(SAMPLE_KEYS, seed=99)
        # Extremely unlikely all match by chance
        assert any(
            r1.canary_for(k) != r2.canary_for(k) for k in SAMPLE_KEYS
        )


# ---------------------------------------------------------------------------
# Registry serialisation
# ---------------------------------------------------------------------------

class TestRegistryJson:
    def test_write_and_read(self, tmp_path: Path):
        reg = build_registry(SAMPLE_KEYS)
        reg.set_location("master_coa", "shared_data/master_coa.xlsx", "Properties → description")

        out = tmp_path / "canary_registry.json"
        reg.write_json(out)

        data = json.loads(out.read_text())
        assert isinstance(data, list)
        assert len(data) == len(SAMPLE_KEYS)
        # Must be sorted by file_key
        assert [d["file_key"] for d in data] == SAMPLE_KEYS

    def test_deterministic_json(self, tmp_path: Path):
        """Two writes from the same seed produce identical JSON."""
        r1 = build_registry(SAMPLE_KEYS, seed=42)
        r2 = build_registry(SAMPLE_KEYS, seed=42)

        p1 = tmp_path / "a.json"
        p2 = tmp_path / "b.json"
        r1.write_json(p1)
        r2.write_json(p2)
        assert p1.read_text() == p2.read_text()


# ---------------------------------------------------------------------------
# Embedding helpers
# ---------------------------------------------------------------------------

class TestEmbedXlsx:
    def test_embed(self):
        from openpyxl import Workbook
        wb = Workbook()
        loc = embed_canary_xlsx(wb, "AB12CD34")
        assert "CANARY: AB12CD34" in wb.properties.description
        assert loc  # non-empty description

    def test_findable_in_saved_file(self, tmp_path: Path):
        """Canary is findable when the workbook is saved and re-opened."""
        from openpyxl import Workbook, load_workbook
        wb = Workbook()
        embed_canary_xlsx(wb, "XK7P2M9Q")
        path = tmp_path / "test.xlsx"
        wb.save(path)

        wb2 = load_workbook(path)
        assert "CANARY: XK7P2M9Q" in (wb2.properties.description or "")


class TestEmbedDocx:
    def test_embed(self):
        from docx import Document
        doc = Document()
        loc = embed_canary_docx(doc, "EF56GH78")
        assert "CANARY: EF56GH78" in doc.core_properties.comments
        assert loc

    def test_findable_in_saved_file(self, tmp_path: Path):
        from docx import Document
        doc = Document()
        embed_canary_docx(doc, "LM3N8R2T")
        path = tmp_path / "test.docx"
        doc.save(path)

        doc2 = Document(str(path))
        assert "CANARY: LM3N8R2T" in (doc2.core_properties.comments or "")


class TestEmbedPdfReportlab:
    def test_embed(self, tmp_path: Path):
        from reportlab.pdfgen.canvas import Canvas
        path = tmp_path / "test.pdf"
        c = Canvas(str(path))
        loc = embed_canary_pdf_reportlab(c, "QW5E9Y1A")
        c.save()
        assert loc

        # Verify canary is in the raw PDF bytes
        raw = path.read_bytes()
        assert b"CANARY: QW5E9Y1A" in raw


class TestEmbedPdfFpdf2:
    def test_embed(self, tmp_path: Path):
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        loc = embed_canary_pdf_fpdf2(pdf, "ZT4R7K1W")
        path = tmp_path / "test.pdf"
        pdf.output(str(path))
        assert loc

        raw = path.read_bytes()
        assert b"CANARY: ZT4R7K1W" in raw


class TestEmbedCsv:
    def test_comment_format(self):
        line = embed_canary_csv_comment("NP8Q3V6X")
        assert line == "# CANARY: NP8Q3V6X\n"

    def test_findable_in_file(self, tmp_path: Path):
        path = tmp_path / "test.csv"
        line = embed_canary_csv_comment("NP8Q3V6X")
        path.write_text(line + "col1,col2\n1,2\n")
        content = path.read_text()
        assert "CANARY: NP8Q3V6X" in content


# ---------------------------------------------------------------------------
# Quality gate: every canary findable in its generated file
# ---------------------------------------------------------------------------

def _find_canary_in_file(file_path: Path, canary: str) -> bool:
    """Return True if *canary* is findable in *file_path*.

    Supports xlsx (openpyxl properties), docx (core_properties.comments),
    pdf (raw bytes), and csv/txt (raw text).
    """
    suffix = file_path.suffix.lower()
    needle = f"CANARY: {canary}"

    if suffix == ".xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(file_path)
        desc = wb.properties.description or ""
        return needle in desc

    if suffix == ".docx":
        from docx import Document
        doc = Document(str(file_path))
        comments = doc.core_properties.comments or ""
        return needle in comments

    if suffix == ".pdf":
        raw = file_path.read_bytes()
        return needle.encode() in raw

    if suffix in (".csv", ".txt", ".md"):
        return needle in file_path.read_text(errors="replace")

    # Fallback: try raw bytes
    return needle.encode() in file_path.read_bytes()


class TestCanaryQualityGate:
    """Quality gate: after generation, every canary must be findable."""

    @pytest.fixture(scope="class")
    def suite_dir(self, tmp_path_factory: pytest.TempPathFactory) -> Path:
        """Run the generator once and return the output directory."""
        out = tmp_path_factory.mktemp("canary_qg")
        config = load_config("config.yaml")
        generate(config, out)
        return out

    def test_registry_exists(self, suite_dir: Path) -> None:
        reg_path = suite_dir / "canary_registry.json"
        assert reg_path.exists(), "canary_registry.json not emitted"

    def test_all_entries_have_file_path(self, suite_dir: Path) -> None:
        reg = json.loads((suite_dir / "canary_registry.json").read_text())
        missing = [e["file_key"] for e in reg if not e.get("file_path")]
        assert not missing, f"Canary entries without file_path: {missing}"

    def test_all_files_exist(self, suite_dir: Path) -> None:
        reg = json.loads((suite_dir / "canary_registry.json").read_text())
        missing = [
            e["file_key"]
            for e in reg
            if not (suite_dir / e["file_path"]).exists()
        ]
        assert not missing, f"Canary files not found on disk: {missing}"

    def test_every_canary_findable(self, suite_dir: Path) -> None:
        """Core quality gate: every canary must be present in its file."""
        reg = json.loads((suite_dir / "canary_registry.json").read_text())
        failures: list[str] = []
        for entry in reg:
            fpath = suite_dir / entry["file_path"]
            if not fpath.exists():
                failures.append(f"{entry['file_key']}: file missing ({entry['file_path']})")
                continue
            if not _find_canary_in_file(fpath, entry["canary"]):
                failures.append(
                    f"{entry['file_key']}: canary {entry['canary']} not found "
                    f"in {entry['file_path']} (location: {entry.get('location', '?')})"
                )
        assert not failures, (
            f"{len(failures)} canary(ies) not findable:\n"
            + "\n".join(f"  - {f}" for f in failures)
        )
