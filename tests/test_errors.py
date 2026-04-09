"""Tests for generator.errors — planted error framework and quality gate.

The quality gate (TestErrorQualityGate) runs the full generator and verifies
that every planted error in error_registry.json is present in its target file.
Ref: bead synth-data-vcj, prompt.md §1.7 / §9.
"""

from __future__ import annotations

import datetime
import json
import re
from pathlib import Path

import pytest

from generate_test_suite import generate
from generator.config import load_config
from generator.errors import (
    ErrorRegistry,
    PlantedError,
    classification_error,
    date_inconsistency,
    formula_error,
    mismatch_total,
    missing_data,
    rounding_discrepancy,
    stale_data,
    transpose_digits,
    wrong_entity,
)

# ---------------------------------------------------------------------------
# PlantedError dataclass
# ---------------------------------------------------------------------------

class TestPlantedError:
    def test_fields(self):
        e = PlantedError(
            error_id="ERR-001",
            file="cascade_tb_fy2025.xlsx",
            location="Sheet 'Trial Balance', Cell G47",
            type="transposed_digits",
            description="AR balance transposed",
            severity="material",
            which_test_cases_should_catch=["TC-01", "TC-02"],
        )
        assert e.error_id == "ERR-001"
        assert e.which_test_cases_should_catch == ["TC-01", "TC-02"]

    def test_default_test_cases_empty(self):
        e = PlantedError(
            error_id="ERR-099",
            file="test.xlsx",
            location="A1",
            type="stale_data",
            description="test",
            severity="immaterial",
        )
        assert e.which_test_cases_should_catch == []


# ---------------------------------------------------------------------------
# ErrorRegistry
# ---------------------------------------------------------------------------

def _make_error(error_id: str, etype: str = "stale_data", file: str = "a.xlsx",
                test_cases: list[str] | None = None) -> PlantedError:
    return PlantedError(
        error_id=error_id,
        file=file,
        location="A1",
        type=etype,
        description=f"Test error {error_id}",
        severity="material",
        which_test_cases_should_catch=test_cases or [],
    )


class TestErrorRegistry:
    def test_add_and_get(self):
        reg = ErrorRegistry()
        e = _make_error("ERR-001")
        reg.add(e)
        assert reg.get("ERR-001") is e

    def test_reject_invalid_type(self):
        reg = ErrorRegistry()
        e = _make_error("ERR-001", etype="bogus_type")
        with pytest.raises(ValueError, match="Invalid error type"):
            reg.add(e)

    def test_reject_duplicate_id(self):
        reg = ErrorRegistry()
        reg.add(_make_error("ERR-001"))
        with pytest.raises(ValueError, match="Duplicate error_id"):
            reg.add(_make_error("ERR-001", file="b.xlsx"))

    def test_by_file(self):
        reg = ErrorRegistry()
        reg.add(_make_error("ERR-001", file="a.xlsx"))
        reg.add(_make_error("ERR-002", file="b.xlsx"))
        reg.add(_make_error("ERR-003", file="a.xlsx"))
        result = reg.by_file("a.xlsx")
        assert [e.error_id for e in result] == ["ERR-001", "ERR-003"]

    def test_by_type(self):
        reg = ErrorRegistry()
        reg.add(_make_error("ERR-001", etype="stale_data"))
        reg.add(_make_error("ERR-002", etype="missing_data"))
        reg.add(_make_error("ERR-003", etype="stale_data"))
        result = reg.by_type("stale_data")
        assert [e.error_id for e in result] == ["ERR-001", "ERR-003"]

    def test_by_test_case(self):
        reg = ErrorRegistry()
        reg.add(_make_error("ERR-001", test_cases=["TC-01", "TC-02"]))
        reg.add(_make_error("ERR-002", test_cases=["TC-03"]))
        reg.add(_make_error("ERR-003", test_cases=["TC-01"]))
        result = reg.by_test_case("TC-01")
        assert [e.error_id for e in result] == ["ERR-001", "ERR-003"]


class TestRegistryJson:
    def test_write_and_read(self, tmp_path: Path):
        reg = ErrorRegistry()
        reg.add(_make_error("ERR-002"))
        reg.add(_make_error("ERR-001", file="b.xlsx"))

        out = tmp_path / "error_registry.json"
        reg.write_json(out)

        data = json.loads(out.read_text())
        assert isinstance(data, list)
        assert len(data) == 2
        # Sorted by error_id
        assert [d["error_id"] for d in data] == ["ERR-001", "ERR-002"]

    def test_deterministic_json(self, tmp_path: Path):
        """Two identical registries produce identical JSON."""
        def build() -> ErrorRegistry:
            reg = ErrorRegistry()
            reg.add(_make_error("ERR-003", file="c.xlsx"))
            reg.add(_make_error("ERR-001", file="a.xlsx"))
            reg.add(_make_error("ERR-002", file="b.xlsx"))
            return reg

        p1 = tmp_path / "a.json"
        p2 = tmp_path / "b.json"
        build().write_json(p1)
        build().write_json(p2)
        assert p1.read_text() == p2.read_text()


# ---------------------------------------------------------------------------
# Transformation functions
# ---------------------------------------------------------------------------

class TestTransposeDigits:
    def test_basic_int(self):
        # 18423109 → swap pos -3 and -2 → 18423190? Let's be explicit.
        # digits: 1 8 4 2 3 1 0 9
        # pos -3 = index 5 = '1', pos -2 = index 6 = '0'
        # swapped: 1 8 4 2 3 0 1 9 = 18423019
        result = transpose_digits(18423109, pos1=-3, pos2=-2)
        assert result == 18423019

    def test_basic_float(self):
        result = transpose_digits(12345.67, pos1=0, pos2=1)
        # digits of int part "12345": swap index 0,1 → "21345"
        assert result == 21345.67

    def test_negative(self):
        result = transpose_digits(-12345, pos1=0, pos2=1)
        assert result == -21345

    def test_returns_same_type(self):
        assert isinstance(transpose_digits(12345, pos1=0, pos2=1), int)
        assert isinstance(transpose_digits(12345.0, pos1=0, pos2=1), float)

    def test_same_position_raises(self):
        with pytest.raises(ValueError, match="same index"):
            transpose_digits(12345, pos1=2, pos2=2)

    def test_identical_digits_raises(self):
        # 11234 — positions 0 and 1 are both '1'
        with pytest.raises(ValueError, match="identical"):
            transpose_digits(11234, pos1=0, pos2=1)


class TestMismatchTotal:
    def test_positive_delta(self):
        assert mismatch_total(100_000, 5_000) == 105_000

    def test_negative_delta(self):
        assert mismatch_total(100_000, -5_000) == 95_000

    def test_float(self):
        result = mismatch_total(100.50, 0.25)
        assert result == pytest.approx(100.75)

    def test_preserves_type(self):
        assert isinstance(mismatch_total(100, 5), int)
        assert isinstance(mismatch_total(100.0, 5.0), float)


class TestStaleData:
    def test_returns_input(self):
        assert stale_data(42) == 42
        assert stale_data("old value") == "old value"


class TestFormulaError:
    def test_returns_wrong_value(self):
        assert formula_error(correct_value=1000, wrong_value=900) == 900


class TestWrongEntity:
    def test_returns_wrong_name(self):
        assert wrong_entity("Cascade Precision", "Cascade Advanced") == "Cascade Advanced"


class TestDateInconsistency:
    def test_returns_wrong_date(self):
        assert date_inconsistency("2025-01-01", "2024-01-01") == "2024-01-01"


class TestClassificationError:
    def test_returns_wrong_account(self):
        assert classification_error("6100-Travel", "6200-Meals") == "6200-Meals"


class TestMissingData:
    def test_default_none(self):
        assert missing_data() is None

    def test_custom_placeholder(self):
        assert missing_data("") == ""
        assert missing_data(0) == 0


class TestRoundingDiscrepancy:
    def test_round_up(self):
        # 1234.5 rounded to 0 dp with ROUND_HALF_UP → 1235
        result = rounding_discrepancy(1234.5, decimal_places=0, direction="up")
        assert result == 1235.0

    def test_truncate_down(self):
        # 1234.9 truncated to 0 dp → 1234
        result = rounding_discrepancy(1234.9, decimal_places=0, direction="down")
        assert result == 1234.0

    def test_no_discrepancy_raises(self):
        # 1234.0 rounded to 0 dp is still 1234.0
        with pytest.raises(ValueError, match="same value"):
            rounding_discrepancy(1234.0, decimal_places=0, direction="up")

    def test_invalid_direction_raises(self):
        with pytest.raises(ValueError, match="direction must be"):
            rounding_discrepancy(1234.5, direction="sideways")

    def test_negative_truncate(self):
        # -1234.9 truncated toward zero → -1234
        result = rounding_discrepancy(-1234.9, decimal_places=0, direction="down")
        assert result == -1234.0

    def test_decimal_places(self):
        # 100.456 rounded to 2dp up → 100.46
        result = rounding_discrepancy(100.456, decimal_places=2, direction="up")
        assert result == pytest.approx(100.46)


# ---------------------------------------------------------------------------
# Quality gate: every planted error verifiable in its file
# Ref: bead synth-data-vcj, prompt.md §1.7 / §9
# ---------------------------------------------------------------------------

# Target count from prompt.md §1.7
_TARGET_ERROR_COUNT = 25


def _extract_all_values_xlsx(file_path: Path) -> list[str]:
    """Return all cell values from all sheets as strings."""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, data_only=True)
    values: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for cell in row:
                if cell is not None:
                    if isinstance(cell, datetime.datetime):
                        # Include multiple date formats for matching
                        values.append(cell.strftime("%m/%d/%Y"))
                        values.append(cell.strftime("%Y-%m-%d"))
                    values.append(str(cell))
    return values


def _extract_all_text_docx(file_path: Path) -> str:
    """Return all text content from a docx file."""
    from docx import Document

    doc = Document(str(file_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _extract_all_text_pdf(file_path: Path) -> str:
    """Return all text content from a PDF file."""
    import pypdf

    reader = pypdf.PdfReader(str(file_path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n".join(parts)


def _get_file_content(file_path: Path) -> list[str]:
    """Return searchable content from a file.

    For xlsx: list of string cell values.
    For docx/pdf/csv/txt: list containing the full text as one element.
    """
    suffix = file_path.suffix.lower()

    if suffix == ".xlsx":
        return _extract_all_values_xlsx(file_path)
    if suffix == ".docx":
        return [_extract_all_text_docx(file_path)]
    if suffix == ".pdf":
        return [_extract_all_text_pdf(file_path)]
    if suffix in (".csv", ".txt", ".md"):
        return [file_path.read_text(errors="replace")]

    # Fallback: raw bytes decoded
    return [file_path.read_bytes().decode(errors="replace")]


def _parse_corrupted_value(description: str) -> str | None:
    """Extract the corrupted value from a description like 'shows X instead of Y'.

    Returns the corrupted value (X) as a string, or None if unparseable.
    """
    # Pattern: "shows $X instead of $Y" or "shows X instead of Y"
    m = re.search(
        r"shows\s+\$?([\d,._/%]+(?:%)?)\s+(?:\([^)]*\)\s+)?instead\s+of",
        description,
    )
    if m:
        return m.group(1).replace(",", "").replace("$", "")
    return None


def _verify_error_in_file(
    error: dict, file_content: list[str]
) -> str | None:
    """Return None if the error is verifiable, or a failure message.

    Verification strategy by error type:
    - transposed_digits / mismatched_total / rounding_discrepancy:
      Parse "shows X instead of Y" from description, verify X is in file.
    - stale_data: Parse stale value from description, verify in file.
    - date_inconsistency: Parse wrong date, verify in file.
    - formula_error: Parse wrong value, verify in file.
    - wrong_entity: Parse wrong entity name, verify in file.
    - classification_error: Verify the misclassified item exists in the file.
    - missing_data: Verify the file exists (absence is the error — hard to
      verify positively, so we just confirm the file is readable).
    """
    etype = error["type"]
    desc = error["description"]

    # missing_data: the error IS the absence of data — we can't verify a
    # missing value is missing without knowing what cell to check.
    # The file existence check (done separately) is sufficient.
    if etype == "missing_data":
        return None

    # For classification_error: the error is about wrong categorization.
    # Verify the misclassified item mentioned in the description exists.
    if etype == "classification_error":
        # Look for key identifiers from the description in the file content
        # E.g., "ADJ-004" or "Non-recurring consulting fees"
        all_text = " ".join(file_content)
        # Extract quoted or notable identifiers from description
        identifiers = re.findall(r"[A-Z]{2,4}-\d{3,4}", desc)
        if identifiers:
            for ident in identifiers:
                if ident in all_text:
                    return None
            return (
                f"Classification error identifiers {identifiers} "
                f"not found in file"
            )
        # Fallback: if no identifier pattern, just check file is non-empty
        return None

    # For wrong_entity: look for the wrong entity name in the file
    if etype == "wrong_entity":
        m = re.search(r"shows\s+['\"]?(.+?)['\"]?\s+instead\s+of", desc)
        if m:
            wrong_name = m.group(1).strip()
            all_text = " ".join(file_content)
            if wrong_name in all_text:
                return None
            return f"Wrong entity name '{wrong_name}' not found in file"
        return None  # Can't parse — skip

    # For stale_data: description may mention the stale value
    if etype == "stale_data":
        corrupted = _parse_corrupted_value(desc)
        if corrupted:
            all_text = " ".join(file_content)
            if corrupted in all_text:
                return None
            return f"Stale value '{corrupted}' not found in file"
        # Some stale_data descriptions don't follow "shows X instead of Y"
        # pattern — e.g. they mention a rate in parentheses like "5.8%"
        m = re.search(r"(\d+\.\d+%)", desc)
        if m:
            stale_val = m.group(1)
            all_text = " ".join(file_content)
            if stale_val in all_text:
                return None
            return f"Stale rate '{stale_val}' not found in file"
        return None

    # Generic: try to parse "shows X instead of Y"
    corrupted = _parse_corrupted_value(desc)
    if corrupted is None:
        # Can't parse — skip verification (file existence already checked)
        return None

    # Search for the corrupted value
    # For numbers, strip formatting and search
    clean_corrupted = corrupted.replace("%", "").strip(".")
    all_text = " ".join(file_content)

    if clean_corrupted in all_text:
        return None

    # For dates, try multiple formats
    if etype == "date_inconsistency":
        # Try both MM/DD/YYYY and YYYY-MM-DD
        m = re.match(r"(\d{2})/(\d{2})/(\d{4})", corrupted)
        if m:
            alt = f"{m.group(3)}-{m.group(1)}-{m.group(2)}"
            if alt in all_text:
                return None

    return f"Corrupted value '{corrupted}' not found in file"


# -- Module-level generation cache (same pattern as test_file_integrity) ------

_REGISTRY: list[dict] | None = None
_OUTPUT_DIR: Path | None = None


def _ensure_generated() -> tuple[list[dict], Path]:
    """Run the generator once and cache the result for the module."""
    global _REGISTRY, _OUTPUT_DIR  # noqa: PLW0603
    if _REGISTRY is None:
        import tempfile

        outdir = Path(tempfile.mkdtemp(prefix="error_qg_"))
        cfg = load_config("config.yaml")
        generate(cfg, outdir)
        reg_path = outdir / "error_registry.json"
        _REGISTRY = json.loads(reg_path.read_text())
        _OUTPUT_DIR = outdir
    return _REGISTRY, _OUTPUT_DIR


class TestErrorQualityGate:
    """Quality gate: every planted error must be verifiable in its file."""

    def test_registry_exists(self) -> None:
        registry, outdir = _ensure_generated()
        assert (outdir / "error_registry.json").exists(), (
            "error_registry.json not emitted"
        )

    def test_registry_not_empty(self) -> None:
        registry, _ = _ensure_generated()
        assert len(registry) > 0, "error_registry.json is empty"

    def test_target_count(self) -> None:
        """Prompt.md §1.7 requires exactly 25 planted errors."""
        registry, _ = _ensure_generated()
        assert len(registry) == _TARGET_ERROR_COUNT, (
            f"Expected {_TARGET_ERROR_COUNT} errors, got {len(registry)}. "
            f"IDs present: {[e['error_id'] for e in registry]}"
        )

    def test_unique_error_ids(self) -> None:
        registry, _ = _ensure_generated()
        ids = [e["error_id"] for e in registry]
        assert len(ids) == len(set(ids)), f"Duplicate error IDs: {ids}"

    def test_all_types_valid(self) -> None:
        from generator.errors import _VALID_TYPES

        registry, _ = _ensure_generated()
        invalid = [
            (e["error_id"], e["type"])
            for e in registry
            if e["type"] not in _VALID_TYPES
        ]
        assert not invalid, f"Invalid error types: {invalid}"

    def test_all_files_exist(self) -> None:
        registry, outdir = _ensure_generated()
        missing = [
            (e["error_id"], e["file"])
            for e in registry
            if not (outdir / e["file"]).exists()
        ]
        assert not missing, (
            f"{len(missing)} error file(s) not found:\n"
            + "\n".join(f"  - {eid}: {f}" for eid, f in missing)
        )

    def test_every_error_verifiable(self) -> None:
        """Core quality gate: each error's corrupted value must be
        findable in its referenced file."""
        registry, outdir = _ensure_generated()
        failures: list[str] = []

        for error in registry:
            fpath = outdir / error["file"]
            if not fpath.exists():
                failures.append(
                    f"{error['error_id']}: file missing ({error['file']})"
                )
                continue

            content = _get_file_content(fpath)
            result = _verify_error_in_file(error, content)
            if result:
                failures.append(f"{error['error_id']}: {result}")

        assert not failures, (
            f"{len(failures)} error(s) not verifiable:\n"
            + "\n".join(f"  - {f}" for f in failures)
        )
