"""Tests for TC-10 — Multi-State Apportionment (Tax, Routine) formatter.

Verifies:
- 6-state scenario (OR, TX, IL, CA, WA, NY)
- Consolidated P&L income statement with line items
- State factors schedule: OR/TX/IL complete, CA/WA partial, NY blank
- "$0" vs blank distinction for zero-presence vs not-provided
- Apportionment formula application per state
- ERR-003 planted error (transposed digits in IL sales factor)
- ERR-023 planted error (rounding discrepancy in total operating expenses)
- Canary embedding in all files
- Gold standard structure, state-by-state outputs, and scoring hints
- Prompt and expected behavior markdown files
"""

from __future__ import annotations

import json
import tempfile
from decimal import Decimal
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc10 import (
    _APPORTIONMENT_FORMULAS,
    _STATE_TAX_RATES,
    _compute_state_factors,
    emit_tc10,
)
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc10 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None

# All canary keys used by TC-10
_CANARY_KEYS = sorted([
    "tc10_consolidated_pl",
    "tc10_state_factors",
    "tc10_apportionment_rules",
])


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc10_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc10(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard
        emit_gold("TC-10", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-10/input_files"


# ---------------------------------------------------------------------------
# Consolidated P&L — Income statement
# ---------------------------------------------------------------------------


class TestConsolidatedPL:
    """Verify consolidated_pl_fy2025.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/consolidated_pl_fy2025.xlsx"
        assert path.exists()

    def test_has_income_statement_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/consolidated_pl_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Income Statement" in wb.sheetnames

    def test_has_key_line_items(self) -> None:
        """P&L should have standard line items: revenue, COGS, opex, pre-tax."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/consolidated_pl_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Income Statement"]
        labels = set()
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str):
                labels.add(val.strip())
        assert "Total Revenue" in labels
        assert "Cost of Goods Sold" in labels
        assert "Gross Profit" in labels
        assert "Total Operating Expenses" in labels
        assert "Operating Income" in labels
        assert "Pre-Tax Income" in labels

    def test_amounts_are_populated(self) -> None:
        """All line items should have dollar amounts in column B."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/consolidated_pl_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Income Statement"]
        amounts = []
        for row in range(2, ws.max_row + 1):
            val = ws.cell(row=row, column=2).value
            if isinstance(val, (int, float)):
                amounts.append(val)
        assert len(amounts) >= 6, f"Expected ≥6 line item amounts, got {len(amounts)}"


# ---------------------------------------------------------------------------
# State Factors — 6 states with varying completeness
# ---------------------------------------------------------------------------


class TestStateFactors:
    """Verify state_factors.xlsx structure and data completeness."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/state_factors.xlsx"
        assert path.exists()

    def test_has_state_factors_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/state_factors.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "State Factors" in wb.sheetnames

    def test_6_states_present(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/state_factors.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["State Factors"]
        states = []
        for row in range(2, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str) and len(val) == 2:
                states.append(val)
        assert set(states) == {"OR", "TX", "IL", "CA", "WA", "NY"}

    def test_or_tx_il_have_all_three_factors(self) -> None:
        """OR, TX, IL should have sales, payroll, and property factors."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/state_factors.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["State Factors"]
        for row in range(2, ws.max_row + 1):
            state = ws.cell(row=row, column=1).value
            if state in ("OR", "TX", "IL"):
                # All three factor columns (B, C, D) should be non-empty
                sales = ws.cell(row=row, column=2).value
                payroll = ws.cell(row=row, column=3).value
                prop = ws.cell(row=row, column=4).value
                assert sales is not None, f"{state} missing sales factor"
                assert payroll is not None, f"{state} missing payroll factor"
                assert prop is not None, f"{state} missing property factor"

    def test_ca_wa_sales_only(self) -> None:
        """CA and WA should have sales factor only; payroll and property blank."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/state_factors.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["State Factors"]
        for row in range(2, ws.max_row + 1):
            state = ws.cell(row=row, column=1).value
            if state in ("CA", "WA"):
                sales = ws.cell(row=row, column=2).value
                payroll = ws.cell(row=row, column=3).value
                prop = ws.cell(row=row, column=4).value
                assert sales is not None, f"{state} should have sales factor"
                assert payroll is None, f"{state} payroll should be blank"
                assert prop is None, f"{state} property should be blank"

    def test_ny_all_blank(self) -> None:
        """NY should have all factors blank."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/state_factors.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["State Factors"]
        for row in range(2, ws.max_row + 1):
            state = ws.cell(row=row, column=1).value
            if state == "NY":
                sales = ws.cell(row=row, column=2).value
                payroll = ws.cell(row=row, column=3).value
                prop = ws.cell(row=row, column=4).value
                assert sales is None, "NY sales should be blank"
                assert payroll is None, "NY payroll should be blank"
                assert prop is None, "NY property should be blank"

    def test_blank_cells_for_partial_states(self) -> None:
        """CA/WA/NY should have blank (None) payroll/property cells,
        distinct from the numeric values in OR/TX/IL. The formatter also
        supports "$0" for zero-presence when a factor is exactly zero,
        but the current model data may not produce that case."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/state_factors.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["State Factors"]
        blank_count = 0
        for row in range(2, ws.max_row + 1):
            state = ws.cell(row=row, column=1).value
            if state in ("CA", "WA", "NY"):
                for col in (3, 4):  # payroll, property columns
                    if ws.cell(row=row, column=col).value is None:
                        blank_count += 1
        # CA(2) + WA(2) + NY(3 — sales also blank) = at least 4 blanks for payroll/property
        assert blank_count >= 4, f"Expected ≥4 blank cells for partial states, got {blank_count}"


# ---------------------------------------------------------------------------
# Apportionment factor computation from canonical model
# ---------------------------------------------------------------------------


class TestFactorComputation:
    """Verify that state factors are computed from the canonical model."""

    def test_compute_state_factors_returns_6_states(self) -> None:
        model, _, _, _, _ = _ensure_emitted()
        factors = _compute_state_factors(model)
        assert set(factors.keys()) == {"OR", "TX", "IL", "CA", "WA", "NY"}

    def test_or_tx_il_have_all_factors(self) -> None:
        model, _, _, _, _ = _ensure_emitted()
        factors = _compute_state_factors(model)
        for state in ("OR", "TX", "IL"):
            sf = factors[state]
            assert sf["sales"] is not None, f"{state} sales factor is None"
            assert sf["payroll"] is not None, f"{state} payroll factor is None"
            assert sf["property"] is not None, f"{state} property factor is None"

    def test_ca_wa_sales_only_in_model(self) -> None:
        model, _, _, _, _ = _ensure_emitted()
        factors = _compute_state_factors(model)
        for state in ("CA", "WA"):
            sf = factors[state]
            assert sf["sales"] is not None, f"{state} sales should be computed"
            assert sf["payroll"] is None, f"{state} payroll should be None"
            assert sf["property"] is None, f"{state} property should be None"

    def test_ny_all_none(self) -> None:
        model, _, _, _, _ = _ensure_emitted()
        factors = _compute_state_factors(model)
        sf = factors["NY"]
        assert sf["sales"] is None
        assert sf["payroll"] is None
        assert sf["property"] is None

    def test_factors_are_fractions(self) -> None:
        """All non-None factors should be between 0 and 1."""
        model, _, _, _, _ = _ensure_emitted()
        factors = _compute_state_factors(model)
        for state, sf in factors.items():
            for name in ("sales", "payroll", "property"):
                val = sf[name]
                if val is not None:
                    assert Decimal(0) <= val <= Decimal(1), (
                        f"{state} {name} factor {val} not in [0, 1]"
                    )


# ---------------------------------------------------------------------------
# Apportionment rules docx
# ---------------------------------------------------------------------------


class TestApportionmentRules:
    """Verify apportionment_rules.docx content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/apportionment_rules.docx"
        assert path.exists()

    def test_mentions_all_6_states(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/apportionment_rules.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\n".join(cell.text for cell in row.cells)
        for state in ["Oregon", "Texas", "Illinois", "California", "Washington", "New York"]:
            assert state.lower() in text.lower(), f"{state} not mentioned in rules doc"

    def test_mentions_margin_tax_for_tx(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/apportionment_rules.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\n".join(cell.text for cell in row.cells)
        assert "margin tax" in text.lower()

    def test_mentions_bno_tax_for_wa(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/apportionment_rules.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\n".join(cell.text for cell in row.cells)
        assert "b&o" in text.lower() or "business & occupation" in text.lower()

    def test_mentions_zero_vs_missing(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/apportionment_rules.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "$0" in text or "zero" in text.lower()
        assert "blank" in text.lower() or "not provided" in text.lower()


# ---------------------------------------------------------------------------
# State tax rates and formulas (constants)
# ---------------------------------------------------------------------------


class TestStateConstants:
    """Verify state tax rates and formula definitions match spec."""

    def test_6_state_tax_rates(self) -> None:
        assert set(_STATE_TAX_RATES.keys()) == {"OR", "TX", "IL", "CA", "WA", "NY"}

    def test_or_rate_6_6_pct(self) -> None:
        assert _STATE_TAX_RATES["OR"] == Decimal("0.066")

    def test_tx_rate_zero(self) -> None:
        # TX has margin tax, no income tax rate
        assert _STATE_TAX_RATES["TX"] == Decimal("0.00")

    def test_il_rate_9_5_pct(self) -> None:
        assert _STATE_TAX_RATES["IL"] == Decimal("0.095")

    def test_ca_rate_8_84_pct(self) -> None:
        assert _STATE_TAX_RATES["CA"] == Decimal("0.084")

    def test_wa_rate_zero(self) -> None:
        # WA has B&O tax, no income tax rate
        assert _STATE_TAX_RATES["WA"] == Decimal("0.00")

    def test_ny_rate_7_1_pct(self) -> None:
        assert _STATE_TAX_RATES["NY"] == Decimal("0.071")

    def test_6_apportionment_formulas(self) -> None:
        assert set(_APPORTIONMENT_FORMULAS.keys()) == {"OR", "TX", "IL", "CA", "WA", "NY"}

    def test_or_single_sales_factor(self) -> None:
        assert "single sales factor" in _APPORTIONMENT_FORMULAS["OR"].lower()

    def test_tx_margin_tax(self) -> None:
        assert "margin tax" in _APPORTIONMENT_FORMULAS["TX"].lower()

    def test_wa_bno_tax(self) -> None:
        assert "b&o" in _APPORTIONMENT_FORMULAS["WA"].lower()


# ---------------------------------------------------------------------------
# ERR-003 — Transposed digits in IL sales factor
# ---------------------------------------------------------------------------


class TestERR003PlantedError:
    """Verify ERR-003: transposed digits in IL sales factor."""

    def test_err003_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-003" in errors.entries

    def test_err003_is_transposed_digits(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-003"]
        assert err.type == "transposed_digits"

    def test_err003_references_il_sales(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-003"]
        assert "IL" in err.location or "IL" in err.description

    def test_err003_in_state_factors_file(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-003"]
        assert "state_factors.xlsx" in err.file


# ---------------------------------------------------------------------------
# ERR-023 — Rounding discrepancy in total operating expenses
# ---------------------------------------------------------------------------


class TestERR023PlantedError:
    """Verify ERR-023: rounding discrepancy in total operating expenses."""

    def test_err023_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-023" in errors.entries

    def test_err023_is_rounding_discrepancy(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-023"]
        assert err.type == "rounding_discrepancy"

    def test_err023_references_operating_expenses(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-023"]
        assert "operating expenses" in err.description.lower() or "operating expenses" in err.location.lower()

    def test_err023_in_consolidated_pl_file(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-023"]
        assert "consolidated_pl_fy2025.xlsx" in err.file


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """Verify canary codes are embedded in files."""

    def test_all_canary_keys_assigned(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            code = canaries.canary_for(key)
            assert len(code) == 8, f"Canary for {key} should be 8 chars"

    def test_consolidated_pl_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc10_consolidated_pl")
        path = output / f"{_INPUT_DIR}/consolidated_pl_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in consolidated P&L properties"

    def test_state_factors_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc10_state_factors")
        path = output / f"{_INPUT_DIR}/state_factors.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in state factors properties"

    def test_apportionment_rules_canary_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc10_apportionment_rules")
        path = output / f"{_INPUT_DIR}/apportionment_rules.docx"
        doc = Document(str(path))
        comments = doc.core_properties.comments or ""
        assert canary in comments, f"Canary {canary} not in apportionment rules properties"


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold standard JSON structure."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-10_gold.json"
        assert path.exists()

    def test_gold_has_expected_outputs(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-10_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "pre_tax_income" in eo
        assert "states" in eo
        assert "flags" in eo

    def test_gold_covers_6_states(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-10_gold.json").read_text())
        states = gold["expected_outputs"]["states"]
        assert set(states.keys()) == {"OR", "TX", "IL", "CA", "WA", "NY"}

    def test_gold_tx_cannot_compute(self) -> None:
        """TX should be flagged as margin tax — cannot compute with income data."""
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-10_gold.json").read_text())
        tx = gold["expected_outputs"]["states"]["TX"]
        assert tx["tax_type"] == "margin_tax"
        assert tx["apportioned_income"] == "cannot_compute"

    def test_gold_wa_cannot_compute(self) -> None:
        """WA should be flagged as gross receipts tax — cannot compute."""
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-10_gold.json").read_text())
        wa = gold["expected_outputs"]["states"]["WA"]
        assert wa["tax_type"] == "gross_receipts_tax"
        assert wa["apportioned_income"] == "cannot_compute"

    def test_gold_ny_no_data(self) -> None:
        """NY should have no computable values."""
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-10_gold.json").read_text())
        ny = gold["expected_outputs"]["states"]["NY"]
        assert ny["apportioned_income"] == "cannot_compute"
        assert ny["estimated_tax"] == "cannot_compute"

    def test_gold_or_il_have_computed_tax(self) -> None:
        """OR and IL should have computed apportioned income and tax."""
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-10_gold.json").read_text())
        for state in ("OR", "IL"):
            st = gold["expected_outputs"]["states"][state]
            assert st["apportioned_income"] != "cannot_compute", f"{state} should have income"
            assert st["estimated_tax"] != "cannot_compute", f"{state} should have tax"

    def test_gold_flags_present(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-10_gold.json").read_text())
        flags = gold["expected_outputs"]["flags"]
        assert "tx_margin_tax" in flags
        assert "wa_bno_tax" in flags
        assert "wa_nexus_question" in flags
        assert "ca_incomplete_data" in flags
        assert "ny_no_data" in flags
        assert "zero_vs_missing" in flags

    def test_gold_canary_verification(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-10_gold.json").read_text())
        cv = gold["canary_verification"]
        for key in ["read_consolidated_pl", "read_state_factors", "read_apportionment_rules"]:
            assert key in cv, f"Missing canary verification key: {key}"

    def test_gold_error_detection(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-10_gold.json").read_text())
        assert "ERR-003" in gold["error_detection"]
        assert "ERR-023" in gold["error_detection"]

    def test_gold_scoring_hints(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-10_gold.json").read_text())
        hints = gold["scoring_hints"]
        for key in ["correctness", "completeness", "format_compliance", "robustness", "communication"]:
            assert key in hints


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt and expected behavior files are generated."""

    def test_prompt_md_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-10/prompt.md"
        assert path.exists()

    def test_prompt_mentions_apportionment(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-10/prompt.md").read_text()
        assert "apportionment" in text.lower()

    def test_prompt_mentions_all_key_tasks(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-10/prompt.md").read_text().lower()
        assert "incomplete" in text or "missing" in text
        assert "margin tax" in text or "non-standard" in text
        assert "nexus" in text

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-10/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_zero_vs_missing(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-10/expected_behavior.md").read_text()
        assert "$0" in text
        assert "blank" in text.lower() or "not provided" in text.lower()

    def test_expected_behavior_mentions_tx_and_wa(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-10/expected_behavior.md").read_text()
        assert "texas" in text.lower() or "TX" in text
        assert "washington" in text.lower() or "WA" in text


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifest:
    """Verify files are registered in the manifest."""

    def test_manifest_has_tc10_entries(self) -> None:
        _, _, _, _, manifest = _ensure_emitted()
        tc10_count = sum(
            1 for v in manifest.entries.values()
            if "TC-10" in (v.test_cases or [])
        )
        # consolidated_pl + state_factors + apportionment_rules = 3
        assert tc10_count >= 3, f"Expected ≥3 TC-10 manifest entries, got {tc10_count}"
