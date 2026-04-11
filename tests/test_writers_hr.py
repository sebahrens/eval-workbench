"""Tests for generator.writers.hr — HR document writer helpers."""

from __future__ import annotations

import openpyxl
import pytest
from docx import Document

from generator.canaries import build_registry
from generator.model.hr_diligence import (
    CONTRACTOR_CLASSIFICATION_SIGNALS,
    DILIGENCE_REQUESTS,
    EMPLOYMENT_AGREEMENTS,
    RETENTION_AWARDS,
    SEVERANCE_EXPOSURES,
)
from generator.writers.hr import (
    write_all_employment_agreements,
    write_contractor_roster,
    write_employee_census,
    write_employment_agreement,
    write_hr_diligence_requests,
    write_retention_schedule,
    write_severance_schedule,
)

# ── Canary key fixtures ────────────────────────────────────────────────────

_EA_CANARY_KEYS = {
    ea.agreement_id: f"tc20_agreement_{ea.agreement_id.lower()}"
    for ea in EMPLOYMENT_AGREEMENTS
}
_CENSUS_KEY = "tc20_employee_census"
_SEVERANCE_KEY = "tc20_severance_schedule"
_RETENTION_KEY = "tc20_retention_schedule"
_CONTRACTOR_KEY = "tc20_contractor_roster"
_DILIGENCE_KEY = "tc20_diligence_requests"

_ALL_KEYS = sorted(
    list(_EA_CANARY_KEYS.values())
    + [_CENSUS_KEY, _SEVERANCE_KEY, _RETENTION_KEY, _CONTRACTOR_KEY, _DILIGENCE_KEY]
)


@pytest.fixture
def canaries():
    return build_registry(_ALL_KEYS, seed=42)


# ── Employment agreement tests ─────────────────────────────────────────────


class TestEmploymentAgreement:
    def test_single_agreement_creates_docx(self, tmp_path, canaries):
        ea = EMPLOYMENT_AGREEMENTS[0]  # EA-001 CEO
        key = _EA_CANARY_KEYS[ea.agreement_id]
        out = tmp_path / "ea.docx"

        loc = write_employment_agreement(ea, out, canaries, key)
        assert out.exists()
        assert loc == "Core properties → comments"

    def test_canary_embedded(self, tmp_path, canaries):
        ea = EMPLOYMENT_AGREEMENTS[0]
        key = _EA_CANARY_KEYS[ea.agreement_id]
        out = tmp_path / "ea.docx"

        write_employment_agreement(ea, out, canaries, key)
        doc = Document(str(out))
        assert canaries.canary_for(key) in doc.core_properties.comments

    def test_draft_agreement_marked_draft(self, tmp_path, canaries):
        """EA-006 (Dr. Patel) is not executed — should say DRAFT."""
        ea = next(e for e in EMPLOYMENT_AGREEMENTS if not e.executed)
        key = _EA_CANARY_KEYS[ea.agreement_id]
        out = tmp_path / "ea.docx"

        write_employment_agreement(ea, out, canaries, key)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "DRAFT" in full_text

    def test_executed_agreement_not_marked_draft(self, tmp_path, canaries):
        ea = EMPLOYMENT_AGREEMENTS[0]  # EA-001 — executed
        key = _EA_CANARY_KEYS[ea.agreement_id]
        out = tmp_path / "ea.docx"

        write_employment_agreement(ea, out, canaries, key)
        doc = Document(str(out))
        # Check headings only (title contains DRAFT or not)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert not any("DRAFT" in h for h in headings)

    def test_contains_compensation_info(self, tmp_path, canaries):
        ea = EMPLOYMENT_AGREEMENTS[0]  # CEO, $325,000
        key = _EA_CANARY_KEYS[ea.agreement_id]
        out = tmp_path / "ea.docx"

        write_employment_agreement(ea, out, canaries, key)
        doc = Document(str(out))
        all_text = _extract_all_text(doc)
        assert "$325,000" in all_text
        assert ea.employee_name in all_text

    def test_all_agreements_written(self, tmp_path, canaries):
        locations = write_all_employment_agreements(
            tmp_path, canaries, _EA_CANARY_KEYS,
        )
        assert len(locations) == len(EMPLOYMENT_AGREEMENTS)
        for ea in EMPLOYMENT_AGREEMENTS:
            fname = f"agreement_{ea.agreement_id.lower()}.docx"
            assert (tmp_path / fname).exists()

    def test_deterministic_output(self, tmp_path, canaries):
        ea = EMPLOYMENT_AGREEMENTS[0]
        key = _EA_CANARY_KEYS[ea.agreement_id]
        d1, d2 = tmp_path / "r1", tmp_path / "r2"
        d1.mkdir()
        d2.mkdir()

        c1 = build_registry(_ALL_KEYS, seed=42)
        c2 = build_registry(_ALL_KEYS, seed=42)

        write_employment_agreement(ea, d1 / "ea.docx", c1, key)
        write_employment_agreement(ea, d2 / "ea.docx", c2, key)
        assert (d1 / "ea.docx").read_bytes() == (d2 / "ea.docx").read_bytes()


# ── Employee census tests ──────────────────────────────────────────────────


class TestEmployeeCensus:
    def test_creates_xlsx(self, tmp_path, canaries):
        out = tmp_path / "census.xlsx"
        loc = write_employee_census(out, canaries, _CENSUS_KEY)
        assert out.exists()
        assert loc == "Document properties → description"

    def test_canary_embedded(self, tmp_path, canaries):
        out = tmp_path / "census.xlsx"
        write_employee_census(out, canaries, _CENSUS_KEY)
        wb = openpyxl.load_workbook(str(out))
        assert canaries.canary_for(_CENSUS_KEY) in wb.properties.description

    def test_contains_all_executives(self, tmp_path, canaries):
        out = tmp_path / "census.xlsx"
        write_employee_census(out, canaries, _CENSUS_KEY)
        wb = openpyxl.load_workbook(str(out))
        ws = wb.active

        names = [row[1] for row in ws.iter_rows(min_row=2, max_col=2, values_only=True) if row[1]]
        for ea in EMPLOYMENT_AGREEMENTS:
            assert ea.employee_name in names

    def test_missing_agreement_flagged(self, tmp_path, canaries):
        """EA-006 (Dr. Patel) should show 'Draft — not on file'."""
        out = tmp_path / "census.xlsx"
        write_employee_census(out, canaries, _CENSUS_KEY)
        wb = openpyxl.load_workbook(str(out))
        ws = wb.active

        statuses = {
            row[0]: row[7]
            for row in ws.iter_rows(min_row=2, values_only=True)
            if row[0]
        }
        assert "Draft" in statuses.get("EA-006", "")

    def test_deterministic(self, tmp_path, canaries):
        d1, d2 = tmp_path / "r1", tmp_path / "r2"
        d1.mkdir()
        d2.mkdir()
        c1 = build_registry(_ALL_KEYS, seed=42)
        c2 = build_registry(_ALL_KEYS, seed=42)

        write_employee_census(d1 / "c.xlsx", c1, _CENSUS_KEY)
        write_employee_census(d2 / "c.xlsx", c2, _CENSUS_KEY)
        assert (d1 / "c.xlsx").read_bytes() == (d2 / "c.xlsx").read_bytes()


# ── Severance schedule tests ───────────────────────────────────────────────


class TestSeveranceSchedule:
    def test_creates_xlsx(self, tmp_path, canaries):
        out = tmp_path / "severance.xlsx"
        loc = write_severance_schedule(out, canaries, _SEVERANCE_KEY)
        assert out.exists()
        assert loc == "Document properties → description"

    def test_canary_embedded(self, tmp_path, canaries):
        out = tmp_path / "severance.xlsx"
        write_severance_schedule(out, canaries, _SEVERANCE_KEY)
        wb = openpyxl.load_workbook(str(out))
        assert canaries.canary_for(_SEVERANCE_KEY) in wb.properties.description

    def test_contains_all_exposures(self, tmp_path, canaries):
        out = tmp_path / "severance.xlsx"
        write_severance_schedule(out, canaries, _SEVERANCE_KEY)
        wb = openpyxl.load_workbook(str(out))
        ws = wb.active

        ids = [
            row[0] for row in ws.iter_rows(min_row=2, max_col=1, values_only=True)
            if row[0] and row[0] != "TOTAL"
        ]
        for sev in SEVERANCE_EXPOSURES:
            assert sev.exposure_id in ids

    def test_total_row_present(self, tmp_path, canaries):
        out = tmp_path / "severance.xlsx"
        write_severance_schedule(out, canaries, _SEVERANCE_KEY)
        wb = openpyxl.load_workbook(str(out))
        ws = wb.active

        all_col1 = [
            row[0] for row in ws.iter_rows(min_row=2, max_col=1, values_only=True)
        ]
        assert "TOTAL" in all_col1

    def test_deterministic(self, tmp_path, canaries):
        d1, d2 = tmp_path / "r1", tmp_path / "r2"
        d1.mkdir()
        d2.mkdir()
        c1 = build_registry(_ALL_KEYS, seed=42)
        c2 = build_registry(_ALL_KEYS, seed=42)

        write_severance_schedule(d1 / "s.xlsx", c1, _SEVERANCE_KEY)
        write_severance_schedule(d2 / "s.xlsx", c2, _SEVERANCE_KEY)
        assert (d1 / "s.xlsx").read_bytes() == (d2 / "s.xlsx").read_bytes()


# ── Retention schedule tests ───────────────────────────────────────────────


class TestRetentionSchedule:
    def test_creates_xlsx(self, tmp_path, canaries):
        out = tmp_path / "retention.xlsx"
        loc = write_retention_schedule(out, canaries, _RETENTION_KEY)
        assert out.exists()
        assert loc == "Document properties → description"

    def test_canary_embedded(self, tmp_path, canaries):
        out = tmp_path / "retention.xlsx"
        write_retention_schedule(out, canaries, _RETENTION_KEY)
        wb = openpyxl.load_workbook(str(out))
        assert canaries.canary_for(_RETENTION_KEY) in wb.properties.description

    def test_contains_all_awards(self, tmp_path, canaries):
        out = tmp_path / "retention.xlsx"
        write_retention_schedule(out, canaries, _RETENTION_KEY)
        wb = openpyxl.load_workbook(str(out))
        ws = wb.active

        ids = [
            row[0] for row in ws.iter_rows(min_row=2, max_col=1, values_only=True)
            if row[0] and row[0] != "TOTAL"
        ]
        for ret in RETENTION_AWARDS:
            assert ret.award_id in ids

    def test_deterministic(self, tmp_path, canaries):
        d1, d2 = tmp_path / "r1", tmp_path / "r2"
        d1.mkdir()
        d2.mkdir()
        c1 = build_registry(_ALL_KEYS, seed=42)
        c2 = build_registry(_ALL_KEYS, seed=42)

        write_retention_schedule(d1 / "r.xlsx", c1, _RETENTION_KEY)
        write_retention_schedule(d2 / "r.xlsx", c2, _RETENTION_KEY)
        assert (d1 / "r.xlsx").read_bytes() == (d2 / "r.xlsx").read_bytes()


# ── Contractor roster tests ────────────────────────────────────────────────


class TestContractorRoster:
    def test_creates_xlsx(self, tmp_path, canaries):
        out = tmp_path / "contractors.xlsx"
        loc = write_contractor_roster(out, canaries, _CONTRACTOR_KEY)
        assert out.exists()
        assert loc == "Document properties → description"

    def test_canary_embedded(self, tmp_path, canaries):
        out = tmp_path / "contractors.xlsx"
        write_contractor_roster(out, canaries, _CONTRACTOR_KEY)
        wb = openpyxl.load_workbook(str(out))
        assert canaries.canary_for(_CONTRACTOR_KEY) in wb.properties.description

    def test_contains_all_contractors(self, tmp_path, canaries):
        out = tmp_path / "contractors.xlsx"
        write_contractor_roster(out, canaries, _CONTRACTOR_KEY)
        wb = openpyxl.load_workbook(str(out))
        ws = wb.active

        ids = [
            row[0] for row in ws.iter_rows(min_row=2, max_col=1, values_only=True)
            if row[0]
        ]
        for ccs in CONTRACTOR_CLASSIFICATION_SIGNALS:
            assert ccs.signal_id in ids

    def test_deterministic(self, tmp_path, canaries):
        d1, d2 = tmp_path / "r1", tmp_path / "r2"
        d1.mkdir()
        d2.mkdir()
        c1 = build_registry(_ALL_KEYS, seed=42)
        c2 = build_registry(_ALL_KEYS, seed=42)

        write_contractor_roster(d1 / "c.xlsx", c1, _CONTRACTOR_KEY)
        write_contractor_roster(d2 / "c.xlsx", c2, _CONTRACTOR_KEY)
        assert (d1 / "c.xlsx").read_bytes() == (d2 / "c.xlsx").read_bytes()


# ── HR diligence request tests ─────────────────────────────────────────────


class TestHRDiligenceRequests:
    def test_creates_xlsx(self, tmp_path, canaries):
        out = tmp_path / "requests.xlsx"
        loc = write_hr_diligence_requests(out, canaries, _DILIGENCE_KEY)
        assert out.exists()
        assert loc == "Document properties → description"

    def test_canary_embedded(self, tmp_path, canaries):
        out = tmp_path / "requests.xlsx"
        write_hr_diligence_requests(out, canaries, _DILIGENCE_KEY)
        wb = openpyxl.load_workbook(str(out))
        assert canaries.canary_for(_DILIGENCE_KEY) in wb.properties.description

    def test_contains_all_requests(self, tmp_path, canaries):
        out = tmp_path / "requests.xlsx"
        write_hr_diligence_requests(out, canaries, _DILIGENCE_KEY)
        wb = openpyxl.load_workbook(str(out))
        ws = wb.active

        ids = [
            row[0] for row in ws.iter_rows(min_row=2, max_col=1, values_only=True)
            if row[0]
        ]
        for dr in DILIGENCE_REQUESTS:
            assert dr.request_id in ids

    def test_not_received_shows_correct_status(self, tmp_path, canaries):
        """DR-006 has status not_received and no received_date."""
        out = tmp_path / "requests.xlsx"
        write_hr_diligence_requests(out, canaries, _DILIGENCE_KEY)
        wb = openpyxl.load_workbook(str(out))
        ws = wb.active

        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] == "DR-006":
                assert row[4] == "not_received"
                assert row[6] == "Not received"
                break
        else:
            pytest.fail("DR-006 not found in output")

    def test_deterministic(self, tmp_path, canaries):
        d1, d2 = tmp_path / "r1", tmp_path / "r2"
        d1.mkdir()
        d2.mkdir()
        c1 = build_registry(_ALL_KEYS, seed=42)
        c2 = build_registry(_ALL_KEYS, seed=42)

        write_hr_diligence_requests(d1 / "r.xlsx", c1, _DILIGENCE_KEY)
        write_hr_diligence_requests(d2 / "r.xlsx", c2, _DILIGENCE_KEY)
        assert (d1 / "r.xlsx").read_bytes() == (d2 / "r.xlsx").read_bytes()


# ── Helpers ────────────────────────────────────────────────────────────────

def _extract_all_text(doc: Document) -> str:
    """Extract all text from paragraphs and tables."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)
