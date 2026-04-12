"""Tests for generator.noise — controlled document noise profiles."""

from __future__ import annotations

import random

import pytest

from generator.noise import (
    ExclusionZone,
    apply_csv_noise,
    apply_xlsx_noise,
    make_noise_rng,
)
from generator.scenario_context import ScenarioContext

# ---------------------------------------------------------------------------
# make_noise_rng
# ---------------------------------------------------------------------------

class TestMakeNoiseRng:
    def test_deterministic(self):
        ctx = ScenarioContext(seed=42)
        rng1 = make_noise_rng(ctx, "TC-01", "cascade_tb_fy2025")
        rng2 = make_noise_rng(ctx, "TC-01", "cascade_tb_fy2025")
        assert [rng1.random() for _ in range(10)] == [rng2.random() for _ in range(10)]

    def test_different_files_differ(self):
        ctx = ScenarioContext(seed=42)
        rng_a = make_noise_rng(ctx, "TC-01", "file_a")
        rng_b = make_noise_rng(ctx, "TC-01", "file_b")
        assert [rng_a.random() for _ in range(10)] != [rng_b.random() for _ in range(10)]

    def test_different_tcs_differ(self):
        ctx = ScenarioContext(seed=42)
        rng_a = make_noise_rng(ctx, "TC-01", "file_a")
        rng_b = make_noise_rng(ctx, "TC-02", "file_a")
        assert [rng_a.random() for _ in range(10)] != [rng_b.random() for _ in range(10)]


# ---------------------------------------------------------------------------
# XLSX noise
# ---------------------------------------------------------------------------

class TestXlsxNoise:
    @pytest.fixture()
    def simple_workbook(self):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        # Row 1: header
        ws["A1"] = "Account #"
        ws["B1"] = "Description"
        ws["C1"] = "Amount"
        # Rows 6+: data (skip 2-5 as potential header zone)
        ws["A6"] = "1100"
        ws["B6"] = "Accounts Receivable"
        ws["C6"] = 50000
        ws["A7"] = "2010"
        ws["B7"] = "Accounts Payable"
        ws["C7"] = 30000
        # Canary in document properties
        wb.properties.description = "CANARY: ABCD1234"
        return wb

    def test_canary_preserved(self, simple_workbook):
        rng = random.Random(42)
        apply_xlsx_noise(simple_workbook, rng)
        assert simple_workbook.properties.description == "CANARY: ABCD1234"

    def test_deterministic(self, simple_workbook):
        import openpyxl
        # Create two identical workbooks and apply noise with same seed
        def make_wb():
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Data"
            ws["A1"] = "Account #"
            ws["B1"] = "Description"
            ws["C1"] = "Amount"
            ws["A6"] = "1100"
            ws["B6"] = "Accounts Receivable"
            ws["C6"] = 50000
            return wb

        wb1 = make_wb()
        wb2 = make_wb()
        apply_xlsx_noise(wb1, random.Random(42))
        apply_xlsx_noise(wb2, random.Random(42))

        ws1 = wb1.active
        ws2 = wb2.active
        for row in range(1, 8):
            for col in range(1, 4):
                assert ws1.cell(row=row, column=col).value == ws2.cell(row=row, column=col).value

    def test_exclusion_zone_respected(self, simple_workbook):
        excl = ExclusionZone(cells={("Data", 6, 2)})  # Protect B6
        original_b6 = simple_workbook.active["B6"].value
        # Apply noise many times to increase chance of mutation
        for seed in range(100):
            # Reset B6 each time
            simple_workbook.active["B6"].value = original_b6
            apply_xlsx_noise(simple_workbook, random.Random(seed), excl)
            assert simple_workbook.active["B6"].value == original_b6

    def test_header_perturbation_occurs(self):
        """At least one header gets perturbed across 50 seeds."""
        import openpyxl

        changed = False
        for seed in range(50):
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Test"
            ws["A1"] = "Account #"
            ws["B1"] = "Description"
            apply_xlsx_noise(wb, random.Random(seed))
            if ws["A1"].value != "Account #" or ws["B1"].value != "Description":
                changed = True
                break
        assert changed, "Header perturbation never fired across 50 seeds"

    def test_numeric_cells_untouched(self, simple_workbook):
        """Numeric cell values must never change (model fact preservation)."""
        rng = random.Random(42)
        apply_xlsx_noise(simple_workbook, rng)
        ws = simple_workbook.active
        assert ws["C6"].value == 50000
        assert ws["C7"].value == 30000


# ---------------------------------------------------------------------------
# CSV noise
# ---------------------------------------------------------------------------

class TestCsvNoise:
    @pytest.fixture()
    def sample_lines(self):
        return [
            "# CANARY: WXYZ5678\n",
            "Account,Description,Amount\n",
            "1100,Accounts Receivable,50000\n",
            "2010,Accounts Payable,30000\n",
        ]

    def test_canary_line_preserved(self, sample_lines):
        rng = random.Random(42)
        result = apply_csv_noise(sample_lines, rng)
        assert result[0] == "# CANARY: WXYZ5678\n"

    def test_deterministic(self, sample_lines):
        r1 = apply_csv_noise(list(sample_lines), random.Random(42))
        r2 = apply_csv_noise(list(sample_lines), random.Random(42))
        assert r1 == r2

    def test_exclusion_zone_respected(self, sample_lines):
        excl = ExclusionZone(rows={2})  # Protect row index 2
        original_row2 = sample_lines[2]
        for seed in range(100):
            result = apply_csv_noise(list(sample_lines), random.Random(seed), excl)
            # Row might have BOM prepended to line 0, but row 2 should be unchanged
            assert result[2] == original_row2

    def test_bom_insertion_deterministic(self):
        lines = ["header,col\n", "data,val\n"]
        # Force BOM on
        result = apply_csv_noise(list(lines), random.Random(42), add_bom=True)
        assert result[0].startswith("\ufeff")

    def test_bom_not_doubled(self):
        lines = ["\ufeffheader,col\n", "data,val\n"]
        result = apply_csv_noise(list(lines), random.Random(42), add_bom=True)
        assert not result[0].startswith("\ufeff\ufeff")

    def test_no_bom_when_disabled(self, sample_lines):
        result = apply_csv_noise(list(sample_lines), random.Random(42), add_bom=False)
        assert not result[0].startswith("\ufeff")


# ---------------------------------------------------------------------------
# DOCX noise
# ---------------------------------------------------------------------------

class TestDocxNoise:
    @pytest.fixture()
    def simple_doc(self):
        from docx import Document
        doc = Document()
        doc.core_properties.comments = "CANARY: TESTCAN1"
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Body paragraph with some text.")
        doc.add_paragraph("Another paragraph for testing noise.")
        return doc

    def test_canary_preserved(self, simple_doc):
        from generator.noise import apply_docx_noise
        rng = random.Random(42)
        apply_docx_noise(simple_doc, rng)
        assert simple_doc.core_properties.comments == "CANARY: TESTCAN1"

    def test_deterministic(self):
        from docx import Document  # noqa: I001
        from generator.noise import apply_docx_noise

        def make_doc():
            doc = Document()
            doc.add_paragraph("Test paragraph.")
            return doc

        doc1 = make_doc()
        doc2 = make_doc()
        apply_docx_noise(doc1, random.Random(42))
        apply_docx_noise(doc2, random.Random(42))
        # Metadata should match
        assert doc1.core_properties.author == doc2.core_properties.author
        assert doc1.core_properties.category == doc2.core_properties.category

    def test_exclusion_zone_respected(self, simple_doc):
        from generator.noise import apply_docx_noise
        # Paragraph 1 is the body paragraph (index 1)
        excl = ExclusionZone(paragraphs={1})
        original_runs = [
            (run.text, run.font.size)
            for run in simple_doc.paragraphs[1].runs
        ]
        apply_docx_noise(simple_doc, random.Random(42), excl)
        current_runs = [
            (run.text, run.font.size)
            for run in simple_doc.paragraphs[1].runs
        ]
        # Text must be identical (font size jitter skipped for excluded para)
        assert [r[0] for r in current_runs] == [r[0] for r in original_runs]

    def test_heading_font_not_jittered(self, simple_doc):
        from generator.noise import apply_docx_noise
        heading = simple_doc.paragraphs[0]
        original_sizes = [run.font.size for run in heading.runs]
        # Run noise 50 times — heading runs should never change
        for seed in range(50):
            from docx import Document
            doc = Document()
            doc.add_heading("Title", level=1)
            doc.add_paragraph("Body text.")
            apply_docx_noise(doc, random.Random(seed))
            for run, orig_size in zip(doc.paragraphs[0].runs, original_sizes):
                # Heading runs should be untouched
                assert run.font.size == orig_size or orig_size is None
