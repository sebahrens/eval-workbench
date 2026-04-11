"""Tests for TC-20 — HR Diligence Exposure Summary (Advisory, Routine).

Verifies:
- 7 employment agreement docx files in agreements/ directory
- employee_census.xlsx with all 7 executives
- severance_schedule.xlsx with all 7 exposures and TOTAL row
- retention_plan.xlsx with all 4 retention awards
- contractor_roster.xlsx with all 4 contractor signals
- prompt.md and expected_behavior.md
- Judgment traps: severance exposure tied to salary/multiplier,
  retention/severance double-count avoidance (Dr. Patel greater-of),
  missing EA-006 executed agreement, contractor classification as
  signals not conclusions, census/agreement source citations
- Canary embedding in all files
- Gold standard structure with evidence expectations and judgment traps
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc20 import (
    _CANARY_KEYS,
    emit_tc20,
)
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model
from generator.model.hr_diligence import (
    CONTRACTOR_CLASSIFICATION_SIGNALS,
    EMPLOYMENT_AGREEMENTS,
    RETENTION_AWARDS,
    SEVERANCE_EXPOSURES,
    high_risk_contractors,
    missing_executed_agreements,
    total_retention_awards,
    total_severance_exposure,
)

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc20 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc20_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc20(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard
        emit_gold("TC-20", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-20/input_files"


# ---------------------------------------------------------------------------
# Employment agreements — 7 docx files
# ---------------------------------------------------------------------------


class TestEmploymentAgreements:
    """Verify 7 employment agreement docx files in agreements/ directory."""

    def test_agreements_directory_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/agreements"
        assert path.is_dir()

    def test_7_agreement_files_exist(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        for ea in EMPLOYMENT_AGREEMENTS:
            fname = f"agreement_{ea.agreement_id.lower()}.docx"
            path = output / f"{_INPUT_DIR}/agreements/{fname}"
            assert path.exists(), f"Missing agreement file: {fname}"

    def test_ea006_draft_marked(self) -> None:
        """EA-006 (Dr. Patel) is not executed — should say DRAFT."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/agreements/agreement_ea-006.docx"
        doc = Document(str(path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "DRAFT" in full_text

    def test_executed_agreements_not_draft(self) -> None:
        """EA-001 through EA-005 and EA-007 are executed — no DRAFT heading."""
        _, output, _, _, _ = _ensure_emitted()
        executed_ids = [ea.agreement_id for ea in EMPLOYMENT_AGREEMENTS if ea.executed]
        for aid in executed_ids:
            fname = f"agreement_{aid.lower()}.docx"
            path = output / f"{_INPUT_DIR}/agreements/{fname}"
            doc = Document(str(path))
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
            assert not any("DRAFT" in h for h in headings), f"{aid} should not be DRAFT"

    def test_ceo_agreement_contains_compensation(self) -> None:
        """EA-001 (CEO) should contain $325,000 salary."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/agreements/agreement_ea-001.docx"
        doc = Document(str(path))
        all_text = _extract_all_text(doc)
        assert "$325,000" in all_text
        assert "Robert J. Cascade" in all_text


# ---------------------------------------------------------------------------
# Employee census
# ---------------------------------------------------------------------------


class TestEmployeeCensus:
    """Verify employee_census.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/employee_census.xlsx"
        assert path.exists()

    def test_contains_all_executives(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/employee_census.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        names = [
            row[1] for row in ws.iter_rows(min_row=2, max_col=2, values_only=True)
            if row[1]
        ]
        for ea in EMPLOYMENT_AGREEMENTS:
            assert ea.employee_name in names, f"{ea.employee_name} not in census"

    def test_ea006_flagged_as_draft(self) -> None:
        """EA-006 (Dr. Patel) should show 'Draft' in agreement status column."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/employee_census.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        statuses = {
            row[0]: row[7]
            for row in ws.iter_rows(min_row=2, values_only=True)
            if row[0]
        }
        assert "Draft" in statuses.get("EA-006", "")


# ---------------------------------------------------------------------------
# Severance schedule
# ---------------------------------------------------------------------------


class TestSeveranceSchedule:
    """Verify severance_schedule.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/severance_schedule.xlsx"
        assert path.exists()

    def test_contains_all_exposures(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/severance_schedule.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        ids = [
            row[0] for row in ws.iter_rows(min_row=2, max_col=1, values_only=True)
            if row[0] and row[0] != "TOTAL"
        ]
        for sev in SEVERANCE_EXPOSURES:
            assert sev.exposure_id in ids, f"{sev.exposure_id} not in severance schedule"

    def test_total_row_present(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/severance_schedule.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        col1 = [
            row[0] for row in ws.iter_rows(min_row=2, max_col=1, values_only=True)
        ]
        assert "TOTAL" in col1


# ---------------------------------------------------------------------------
# Retention plan
# ---------------------------------------------------------------------------


class TestRetentionPlan:
    """Verify retention_plan.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/retention_plan.xlsx"
        assert path.exists()

    def test_contains_all_awards(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/retention_plan.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        ids = [
            row[0] for row in ws.iter_rows(min_row=2, max_col=1, values_only=True)
            if row[0] and row[0] != "TOTAL"
        ]
        for ret in RETENTION_AWARDS:
            assert ret.award_id in ids, f"{ret.award_id} not in retention plan"


# ---------------------------------------------------------------------------
# Contractor roster
# ---------------------------------------------------------------------------


class TestContractorRoster:
    """Verify contractor_roster.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/contractor_roster.xlsx"
        assert path.exists()

    def test_contains_all_contractors(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/contractor_roster.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb.active
        ids = [
            row[0] for row in ws.iter_rows(min_row=2, max_col=1, values_only=True)
            if row[0]
        ]
        for ccs in CONTRACTOR_CLASSIFICATION_SIGNALS:
            assert ccs.signal_id in ids, f"{ccs.signal_id} not in contractor roster"


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt and expected behavior files are generated."""

    def test_prompt_md_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-20/prompt.md"
        assert path.exists()

    def test_prompt_mentions_key_tasks(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-20/prompt.md").read_text().lower()
        assert "exposure" in text
        assert "severance" in text
        assert "retention" in text
        assert "contractor" in text
        assert "missing" in text or "executed" in text

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-20/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_double_count(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-20/expected_behavior.md").read_text()
        assert "double-count" in text.lower() or "double count" in text.lower()

    def test_expected_behavior_mentions_patel(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-20/expected_behavior.md").read_text()
        assert "Patel" in text

    def test_expected_behavior_mentions_ea006(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-20/expected_behavior.md").read_text()
        assert "EA-006" in text

    def test_expected_behavior_mentions_greater_of(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-20/expected_behavior.md").read_text()
        assert "greater-of" in text.lower() or "greater of" in text.lower()


# ---------------------------------------------------------------------------
# Judgment traps — verify canonical model data
# ---------------------------------------------------------------------------


class TestJudgmentTraps:
    """Verify judgment trap data from the canonical model."""

    def test_ceo_golden_parachute(self) -> None:
        """SEV-001: CEO 3x = $975,000."""
        sev = SEVERANCE_EXPOSURES[0]
        assert sev.exposure_id == "SEV-001"
        assert sev.base_salary == 325_000
        assert sev.severance_multiplier == 3
        assert sev.estimated_payout == 975_000

    def test_patel_severance_retention_overlap(self) -> None:
        """SEV-006 ($255k) and RET-004 ($150k) — greater-of, not additive."""
        sev = SEVERANCE_EXPOSURES[5]
        ret = RETENTION_AWARDS[3]
        assert sev.exposure_id == "SEV-006"
        assert ret.award_id == "RET-004"
        assert sev.employee_name == "Dr. Anika Patel"
        assert ret.employee_name == "Dr. Anika Patel"
        assert sev.estimated_payout == 255_000
        assert ret.award_amount == 150_000
        # Greater-of means net exposure is the larger amount
        assert sev.estimated_payout > ret.award_amount

    def test_missing_executed_agreement(self) -> None:
        """EA-006 should be the only missing executed agreement."""
        missing = missing_executed_agreements()
        assert len(missing) == 1
        assert missing[0].agreement_id == "EA-006"
        assert missing[0].employee_name == "Dr. Anika Patel"

    def test_high_risk_contractors(self) -> None:
        """CCS-001 and CCS-003 should be high risk."""
        high = high_risk_contractors()
        ids = {c.signal_id for c in high}
        assert ids == {"CCS-001", "CCS-003"}

    def test_severance_computations_match(self) -> None:
        """All severance payouts should equal base_salary * multiplier."""
        for sev in SEVERANCE_EXPOSURES:
            expected = sev.base_salary * sev.severance_multiplier
            assert sev.estimated_payout == expected, (
                f"{sev.exposure_id}: {sev.base_salary} x {sev.severance_multiplier} "
                f"= {expected}, got {sev.estimated_payout}"
            )


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """Verify canary codes are embedded in all TC-20 files."""

    def test_all_canary_keys_assigned(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            code = canaries.canary_for(key)
            assert len(code) == 8, f"Canary for {key} should be 8 chars"

    def test_agreement_canaries_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        for ea in EMPLOYMENT_AGREEMENTS:
            ckey = f"tc20_agreement_{ea.agreement_id.lower().replace('-', '_')}"
            canary = canaries.canary_for(ckey)
            fname = f"agreement_{ea.agreement_id.lower()}.docx"
            path = output / f"{_INPUT_DIR}/agreements/{fname}"
            doc = Document(str(path))
            comments = doc.core_properties.comments or ""
            assert canary in comments, (
                f"Canary {canary} not in {fname} properties"
            )

    def test_census_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc20_employee_census")
        path = output / f"{_INPUT_DIR}/employee_census.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in census properties"

    def test_severance_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc20_severance_schedule")
        path = output / f"{_INPUT_DIR}/severance_schedule.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in severance schedule properties"

    def test_retention_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc20_retention_plan")
        path = output / f"{_INPUT_DIR}/retention_plan.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in retention plan properties"

    def test_contractor_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc20_contractor_roster")
        path = output / f"{_INPUT_DIR}/contractor_roster.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in contractor roster properties"


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold standard JSON structure."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-20_gold.json"
        assert path.exists()

    def test_gold_has_expected_outputs(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert eo["agreement_count"] == len(EMPLOYMENT_AGREEMENTS)
        assert eo["severance_exposure_count"] == len(SEVERANCE_EXPOSURES)
        assert eo["retention_award_count"] == len(RETENTION_AWARDS)
        assert eo["contractor_count"] == len(CONTRACTOR_CLASSIFICATION_SIGNALS)

    def test_gold_totals_match_model(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert eo["total_severance_exposure"] == str(total_severance_exposure())
        assert eo["total_retention_awards"] == str(total_retention_awards())

    def test_gold_severance_entries(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        entries = gold["expected_outputs"]["severance_entries"]
        assert len(entries) == len(SEVERANCE_EXPOSURES)
        ids = {e["exposure_id"] for e in entries}
        for sev in SEVERANCE_EXPOSURES:
            assert sev.exposure_id in ids

    def test_gold_retention_entries(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        entries = gold["expected_outputs"]["retention_entries"]
        assert len(entries) == len(RETENTION_AWARDS)

    def test_gold_contractor_entries(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        entries = gold["expected_outputs"]["contractor_entries"]
        assert len(entries) == len(CONTRACTOR_CLASSIFICATION_SIGNALS)

    def test_gold_double_count_risk(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        dc = gold["expected_outputs"]["double_count_risk"]
        assert dc["employee"] == "Dr. Anika Patel"
        assert dc["severance_id"] == "SEV-006"
        assert dc["retention_id"] == "RET-004"
        assert dc["interaction"] == "greater_of_not_additive"

    def test_gold_missing_agreements(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        missing = gold["expected_outputs"]["missing_executed_agreements"]
        assert missing == ["EA-006"]

    def test_gold_high_risk_contractors(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        high = gold["expected_outputs"]["high_risk_contractors"]
        assert set(high) == {"CCS-001", "CCS-003"}

    def test_gold_canary_verification(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        cv = gold["canary_verification"]
        # Should have entries for all 7 agreements + census + severance + retention + contractor
        assert "read_employee_census" in cv
        assert "read_severance_schedule" in cv
        assert "read_retention_plan" in cv
        assert "read_contractor_roster" in cv
        assert "read_agreement_ea_001" in cv
        assert "read_agreement_ea_006" in cv

    def test_gold_scoring_hints(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        hints = gold["scoring_hints"]
        assert "correctness" in hints
        assert "completeness" in hints
        assert "double_count_avoidance" in hints
        assert "professional_judgment" in hints

    def test_gold_judgment_traps(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        traps = gold["judgment_traps"]
        trap_ids = {t["trap_id"] for t in traps}
        assert "JT-TC20-001" in trap_ids  # missing executed agreement
        assert "JT-TC20-002" in trap_ids  # double-count avoidance
        assert "JT-TC20-003" in trap_ids  # contractor signals not conclusions
        assert "JT-TC20-004" in trap_ids  # source citations
        assert "JT-TC20-005" in trap_ids  # severance computation match

    def test_gold_evidence_expectations(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-20_gold.json").read_text())
        ee = gold["evidence_expectations"]
        assert "exposure_ceo_golden_parachute" in ee
        assert "exposure_double_count_risk" in ee
        assert "exposure_missing_agreement" in ee
        assert "exposure_contractor_classification" in ee
        assert "exposure_source_citations" in ee


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifest:
    """Verify files are registered in the manifest."""

    def test_manifest_has_tc20_entries(self) -> None:
        _, _, _, _, manifest = _ensure_emitted()
        tc20_count = sum(
            1 for v in manifest.entries.values()
            if "TC-20" in (v.test_cases or [])
        )
        # 7 agreements + census + severance + retention + contractor = 11
        assert tc20_count >= 11, f"Expected >= 11 TC-20 manifest entries, got {tc20_count}"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _extract_all_text(doc: Document) -> str:
    """Extract all text from paragraphs and tables."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)
