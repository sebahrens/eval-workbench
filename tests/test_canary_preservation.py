"""Canary-preservation tests for controlled document noise.

Verifies that canaries remain findable after every supported noise transform
across the formatter-family matrix (xlsx_openpyxl, csv_stdlib, docx_python_docx).

Each test class covers one formatter family.  Parametrization over seeds
exercises the full probability surface of noise helpers (header perturbation,
trailing whitespace, date format variation, BOM insertion, font jitter, etc.).

PDF families (pdf_reportlab, pdf_fpdf2) are intentionally excluded — noise
helpers for those families are deferred to the second wave per the
formatter-family matrix.
"""

from __future__ import annotations

import random

import pytest

from generator.canaries import (
    embed_canary_csv_comment,
    embed_canary_docx,
    embed_canary_xlsx,
)
from generator.noise import (
    ExclusionZone,
    apply_csv_noise,
    apply_docx_noise,
    apply_xlsx_noise,
    make_noise_rng,
)
from generator.scenario_context import ScenarioContext

# 50 seeds covers the full probability surface of all noise helpers:
# - header perturbation fires at 40% → P(never in 50) ≈ 0
# - trailing whitespace at 15% → P(never in 50) ≈ 0
# - date format variation at 30%, csv quoting at 8%, BOM at 25%
_SEEDS = list(range(50))


# ---------------------------------------------------------------------------
# XLSX family — canary in document properties → description
# ---------------------------------------------------------------------------

class TestXlsxCanaryPreservation:
    """Canary in wb.properties.description must survive all xlsx noise transforms."""

    @staticmethod
    def _make_workbook(canary_code: str = "TB01XL99"):
        """Build a minimal workbook mimicking TC-01 trial balance structure."""
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Trial Balance"

        # Header row (noise target: header perturbation)
        headers = ["Account #", "Account Name", "Debit", "Credit", "Net Balance"]
        for col, h in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=h)

        # Data rows (noise target: trailing whitespace, date format)
        data = [
            ("1100", "Cash and Equivalents", 1_250_000, 0, 1_250_000),
            ("1200", "Accounts Receivable", 3_800_000, 0, 3_800_000),
            ("2100", "Accounts Payable", 0, 2_100_000, -2_100_000),
            ("3000", "Retained Earnings", 0, 5_400_000, -5_400_000),
        ]
        for r, row_data in enumerate(data, 6):
            for c, val in enumerate(row_data, 1):
                ws.cell(row=r, column=c, value=val)

        # Date column to exercise date-format noise
        ws.cell(row=1, column=6, value="Date")
        import datetime
        for r in range(6, 10):
            cell = ws.cell(row=r, column=6, value=datetime.date(2025, 3, 31))
            cell.number_format = "MM/DD/YYYY"

        embed_canary_xlsx(wb, canary_code)
        return wb

    @pytest.mark.parametrize("seed", _SEEDS, ids=[f"seed-{s}" for s in _SEEDS])
    def test_canary_findable_after_noise(self, seed: int):
        """Canary in document properties must be intact after xlsx noise."""
        canary = "TB01XL99"
        wb = self._make_workbook(canary)
        apply_xlsx_noise(wb, random.Random(seed))
        assert wb.properties.description == f"CANARY: {canary}", (
            f"XLSX canary lost after noise with seed={seed}"
        )

    @pytest.mark.parametrize("seed", _SEEDS, ids=[f"seed-{s}" for s in _SEEDS])
    def test_canary_findable_with_exclusion_zones(self, seed: int):
        """Canary survives even when exclusion zones redirect noise elsewhere."""
        canary = "TB01XL99"
        wb = self._make_workbook(canary)
        excl = ExclusionZone(cells={("Trial Balance", 6, 1), ("Trial Balance", 7, 2)})
        apply_xlsx_noise(wb, random.Random(seed), excl)
        assert wb.properties.description == f"CANARY: {canary}", (
            f"XLSX canary lost with exclusion zones, seed={seed}"
        )

    def test_canary_survives_make_noise_rng_integration(self):
        """End-to-end: make_noise_rng → apply_xlsx_noise preserves canary."""
        canary = "INTEG001"
        wb = self._make_workbook(canary)
        ctx = ScenarioContext(seed=42)
        rng = make_noise_rng(ctx, "TC-01", "cascade_tb_fy2025")
        apply_xlsx_noise(wb, rng)
        assert wb.properties.description == f"CANARY: {canary}"

    def test_multi_sheet_canary_preserved(self):
        """Canary survives when workbook has multiple sheets (common in real TCs)."""
        import openpyxl

        canary = "MULTI001"
        wb = openpyxl.Workbook()
        for name in ["Summary", "Detail", "Adjustments"]:
            ws = wb.create_sheet(name)
            ws["A1"] = "Account #"
            ws["B1"] = "Description"
            ws["A6"] = "Test data"
        # Remove default sheet
        del wb["Sheet"]
        embed_canary_xlsx(wb, canary)

        for seed in range(20):
            wb_copy = openpyxl.Workbook()
            for name in ["Summary", "Detail", "Adjustments"]:
                ws = wb_copy.create_sheet(name)
                ws["A1"] = "Account #"
                ws["B1"] = "Description"
                ws["A6"] = "Test data"
            del wb_copy["Sheet"]
            embed_canary_xlsx(wb_copy, canary)
            apply_xlsx_noise(wb_copy, random.Random(seed))
            assert wb_copy.properties.description == f"CANARY: {canary}", (
                f"Multi-sheet XLSX canary lost, seed={seed}"
            )


# ---------------------------------------------------------------------------
# CSV family — canary as first-line comment: # CANARY: {code}
# ---------------------------------------------------------------------------

class TestCsvCanaryPreservation:
    """Canary comment line must survive all csv noise transforms."""

    @staticmethod
    def _make_lines(canary_code: str = "CSV08REC") -> list[str]:
        """Build CSV lines mimicking TC-08 time records."""
        canary_line = embed_canary_csv_comment(canary_code)
        return [
            canary_line,
            "Employee,Date,Hours,Project,Rate\n",
            "E001,2025-01-15,8.0,RD-001,125.00\n",
            "E002,2025-01-15,7.5,RD-002,110.00\n",
            "E003,2025-01-16,8.0,RD-001,95.00\n",
            "E001,2025-01-16,6.0,RD-003,125.00\n",
            "E004,2025-01-17,8.5,RD-002,140.00\n",
        ]

    @pytest.mark.parametrize("seed", _SEEDS, ids=[f"seed-{s}" for s in _SEEDS])
    def test_canary_line_intact_after_noise(self, seed: int):
        """First-line canary comment must be verbatim after csv noise."""
        canary = "CSV08REC"
        lines = self._make_lines(canary)
        result = apply_csv_noise(lines, random.Random(seed))
        # Strip BOM if present — canary search must be BOM-tolerant
        first = result[0].lstrip("\ufeff")
        assert first == f"# CANARY: {canary}\n", (
            f"CSV canary lost after noise with seed={seed}, got: {result[0]!r}"
        )

    @pytest.mark.parametrize("seed", _SEEDS, ids=[f"seed-{s}" for s in _SEEDS])
    def test_canary_findable_with_bom(self, seed: int):
        """Canary remains findable even when BOM is prepended."""
        canary = "CSV08REC"
        lines = self._make_lines(canary)
        result = apply_csv_noise(lines, random.Random(seed), add_bom=True)
        # The canary text must be present regardless of BOM
        assert f"CANARY: {canary}" in result[0], (
            f"CSV canary not findable through BOM, seed={seed}"
        )

    @pytest.mark.parametrize("seed", _SEEDS, ids=[f"seed-{s}" for s in _SEEDS])
    def test_canary_with_exclusion_zones(self, seed: int):
        """Canary survives when exclusion zones protect data rows."""
        canary = "CSV08REC"
        lines = self._make_lines(canary)
        excl = ExclusionZone(rows={2, 4})  # Protect some data rows
        result = apply_csv_noise(lines, random.Random(seed), exclusions=excl)
        first = result[0].lstrip("\ufeff")
        assert first == f"# CANARY: {canary}\n", (
            f"CSV canary lost with exclusion zones, seed={seed}"
        )

    def test_canary_survives_make_noise_rng_integration(self):
        """End-to-end: make_noise_rng → apply_csv_noise preserves canary."""
        canary = "INTCSV01"
        lines = self._make_lines(canary)
        ctx = ScenarioContext(seed=42)
        rng = make_noise_rng(ctx, "TC-08", "rd_employee_time_records")
        result = apply_csv_noise(lines, rng)
        first = result[0].lstrip("\ufeff")
        assert f"CANARY: {canary}" in first

    def test_canary_not_in_non_canary_line(self):
        """Sanity: noise doesn't accidentally inject canary text elsewhere."""
        canary = "CSV08REC"
        lines = self._make_lines(canary)
        result = apply_csv_noise(lines, random.Random(42))
        for line in result[1:]:
            assert "CANARY:" not in line


# ---------------------------------------------------------------------------
# DOCX family — canary in core_properties.comments
# ---------------------------------------------------------------------------

class TestDocxCanaryPreservation:
    """Canary in doc.core_properties.comments must survive all docx noise."""

    @staticmethod
    def _make_document(canary_code: str = "DOC08RD1"):
        """Build a minimal docx Document mimicking TC-08 R&D project description."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading("Project RD-001: Advanced Materials Research", level=1)
        doc.add_paragraph("Project Lead: Dr. Sarah Chen")
        doc.add_paragraph(
            "This research initiative focuses on developing next-generation "
            "composite materials for aerospace applications. The project spans "
            "three phases over 24 months."
        )
        doc.add_heading("Budget Summary", level=2)
        doc.add_paragraph("Total budget: $1,250,000")
        doc.add_paragraph("Personnel costs: $875,000")
        doc.add_paragraph("Equipment and supplies: $375,000")

        # Set font sizes on body paragraphs so font jitter has something to work with
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size is None:
                    run.font.size = Pt(11)

        # Set paragraph spacing so spacing noise has something to modify
        for para in doc.paragraphs:
            pf = para.paragraph_format
            if pf.space_before is None:
                pf.space_before = Pt(6)
            if pf.space_after is None:
                pf.space_after = Pt(6)

        embed_canary_docx(doc, canary_code)
        return doc

    @pytest.mark.parametrize("seed", _SEEDS, ids=[f"seed-{s}" for s in _SEEDS])
    def test_canary_findable_after_noise(self, seed: int):
        """Canary in core_properties.comments must be intact after docx noise."""
        canary = "DOC08RD1"
        doc = self._make_document(canary)
        apply_docx_noise(doc, random.Random(seed))
        assert doc.core_properties.comments == f"CANARY: {canary}", (
            f"DOCX canary lost after noise with seed={seed}"
        )

    @pytest.mark.parametrize("seed", _SEEDS, ids=[f"seed-{s}" for s in _SEEDS])
    def test_canary_findable_with_exclusion_zones(self, seed: int):
        """Canary survives with paragraph exclusion zones active."""
        canary = "DOC08RD1"
        doc = self._make_document(canary)
        excl = ExclusionZone(paragraphs={1, 3})
        apply_docx_noise(doc, random.Random(seed), excl)
        assert doc.core_properties.comments == f"CANARY: {canary}", (
            f"DOCX canary lost with exclusion zones, seed={seed}"
        )

    def test_canary_survives_make_noise_rng_integration(self):
        """End-to-end: make_noise_rng → apply_docx_noise preserves canary."""
        canary = "INTDOC01"
        doc = self._make_document(canary)
        ctx = ScenarioContext(seed=42)
        rng = make_noise_rng(ctx, "TC-08", "rd_project_RD-001")
        apply_docx_noise(doc, rng)
        assert doc.core_properties.comments == f"CANARY: {canary}"

    @pytest.mark.parametrize("seed", _SEEDS, ids=[f"seed-{s}" for s in _SEEDS])
    def test_metadata_clutter_does_not_touch_comments(self, seed: int):
        """Metadata clutter sets author/category but must not alter comments."""
        canary = "DOC08RD1"
        doc = self._make_document(canary)
        apply_docx_noise(doc, random.Random(seed))
        # author should be changed (metadata clutter)
        # but comments (canary location) must be untouched
        assert doc.core_properties.comments == f"CANARY: {canary}"
        # author should be one of the plausible values (not empty, not canary)
        assert doc.core_properties.author is not None
        assert "CANARY" not in (doc.core_properties.author or "")

    def test_paragraph_text_unchanged_by_noise(self):
        """Font jitter and spacing noise must not alter paragraph text content."""
        canary = "DOC08RD1"
        doc = self._make_document(canary)
        original_texts = [p.text for p in doc.paragraphs]
        apply_docx_noise(doc, random.Random(42))
        current_texts = [p.text for p in doc.paragraphs]
        assert original_texts == current_texts, (
            "DOCX noise altered paragraph text content (should only touch "
            "font sizes and spacing)"
        )


# ---------------------------------------------------------------------------
# Cross-family: determinism regression
# ---------------------------------------------------------------------------

class TestCrossFamily:
    """Cross-family regression tests for determinism after noise application."""

    def test_xlsx_noise_deterministic_across_runs(self):
        """Two identical workbooks + same seed → identical results."""
        import openpyxl

        def make():
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Data"
            ws["A1"] = "Account #"
            ws["B1"] = "Description"
            ws["C1"] = "Amount"
            ws["A6"] = "1100"
            ws["B6"] = "Accounts Receivable"
            ws["C6"] = 50000
            embed_canary_xlsx(wb, "DETXLSX1")
            return wb

        wb1, wb2 = make(), make()
        apply_xlsx_noise(wb1, random.Random(42))
        apply_xlsx_noise(wb2, random.Random(42))

        ws1, ws2 = wb1.active, wb2.active
        for row in range(1, ws1.max_row + 1):
            for col in range(1, ws1.max_column + 1):
                c1 = ws1.cell(row=row, column=col)
                c2 = ws2.cell(row=row, column=col)
                assert c1.value == c2.value, (
                    f"XLSX non-deterministic at ({row},{col}): "
                    f"{c1.value!r} vs {c2.value!r}"
                )
        assert wb1.properties.description == wb2.properties.description

    def test_csv_noise_deterministic_across_runs(self):
        """Two identical line lists + same seed → identical results."""
        lines = [
            "# CANARY: DETCSV01\n",
            "Col1,Col2,Col3\n",
            "a,b,100\n",
            "d,e,200\n",
        ]
        r1 = apply_csv_noise(list(lines), random.Random(42))
        r2 = apply_csv_noise(list(lines), random.Random(42))
        assert r1 == r2, "CSV noise is non-deterministic"

    def test_docx_noise_deterministic_across_runs(self):
        """Two identical documents + same seed → identical metadata."""
        from docx import Document
        from docx.shared import Pt

        def make():
            doc = Document()
            doc.add_paragraph("Test paragraph.")
            for run in doc.paragraphs[0].runs:
                run.font.size = Pt(11)
            embed_canary_docx(doc, "DETDOCX1")
            return doc

        doc1, doc2 = make(), make()
        apply_docx_noise(doc1, random.Random(42))
        apply_docx_noise(doc2, random.Random(42))

        assert doc1.core_properties.comments == doc2.core_properties.comments
        assert doc1.core_properties.author == doc2.core_properties.author
        assert doc1.core_properties.category == doc2.core_properties.category
