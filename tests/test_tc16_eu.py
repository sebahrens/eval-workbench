"""Tests for TC-16-EU — European Engagement Letter Generation formatter.

Verifies:
- Fee calculation: Audit EUR 210,000 x 1.20 = EUR 252,000; CIT EUR 72,000 + EUR 15,000x3 = EUR 117,000;
  VAT EUR 35,000 + EUR 8,000x3 = EUR 59,000; TP EUR 48,000 + EUR 12,000x3 = EUR 84,000; Total EUR 512,000
- Merge field population (client_name, engagement_scope, fee_amount, payment_terms, start_date, partner_name)
- European template with ISA/IFRS/OECD language and NL law limitation of liability
- Planted error ERR-EU-016 (fee_schedule_eu.xlsx audit total missing IFRS multiplier)
- Canary embedding in all 3 input files
- Prompt and expected behavior markdown files
- Gold standard registration
"""

from __future__ import annotations

import json
import tempfile
from decimal import Decimal
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc16_eu import (
    _ADDITIONAL_ENTITIES,
    _AUDIT_BASE_FEE,
    _AUDIT_FEE,
    _CIT_BASE_FEE,
    _CIT_FEE,
    _CIT_PER_ENTITY_ADDER,
    _CLIENT_NAME,
    _IFRS_COMPLEXITY_MULTIPLIER,
    _PARTNER_NAME,
    _PAYMENT_TERMS,
    _START_DATE,
    _TOTAL_FEE,
    _TP_BASE_FEE,
    _TP_FEE,
    _TP_PER_ENTITY_ADDER,
    _VAT_BASE_FEE,
    _VAT_FEE,
    _VAT_PER_ENTITY_ADDER,
    emit_tc16_eu,
)
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc16_eu once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None

_CANARY_KEYS = sorted([
    "tc16eu_client_profile",
    "tc16eu_fee_schedule",
    "tc16eu_engagement_template",
])


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc16eu_test_"))

        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc16_eu(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit TC-16-EU gold standard (registered via @register_gold)
        emit_gold("TC-16-EU", _CANARIES, _ERRORS, _OUTPUT / "gold_standards")

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


# ---------------------------------------------------------------------------
# Constants verification (gold standard values from design bead)
# ---------------------------------------------------------------------------

_INPUT_DIR = "test_cases/TC-16-EU/input_files"


class TestFeeConstants:
    """Verify that the formatter constants match design gold standard."""

    def test_audit_base_fee(self) -> None:
        assert _AUDIT_BASE_FEE == Decimal("210000")

    def test_ifrs_multiplier(self) -> None:
        assert _IFRS_COMPLEXITY_MULTIPLIER == Decimal("1.20")

    def test_audit_fee(self) -> None:
        assert _AUDIT_FEE == Decimal("210000") * Decimal("1.20")
        assert _AUDIT_FEE == Decimal("252000.00")

    def test_cit_base_fee(self) -> None:
        assert _CIT_BASE_FEE == Decimal("72000")

    def test_cit_per_entity_adder(self) -> None:
        assert _CIT_PER_ENTITY_ADDER == Decimal("15000")

    def test_cit_fee(self) -> None:
        assert _CIT_FEE == Decimal("72000") + Decimal("15000") * 3
        assert _CIT_FEE == Decimal("117000")

    def test_vat_base_fee(self) -> None:
        assert _VAT_BASE_FEE == Decimal("35000")

    def test_vat_per_entity_adder(self) -> None:
        assert _VAT_PER_ENTITY_ADDER == Decimal("8000")

    def test_vat_fee(self) -> None:
        assert _VAT_FEE == Decimal("35000") + Decimal("8000") * 3
        assert _VAT_FEE == Decimal("59000")

    def test_tp_base_fee(self) -> None:
        assert _TP_BASE_FEE == Decimal("48000")

    def test_tp_per_entity_adder(self) -> None:
        assert _TP_PER_ENTITY_ADDER == Decimal("12000")

    def test_tp_fee(self) -> None:
        assert _TP_FEE == Decimal("48000") + Decimal("12000") * 3
        assert _TP_FEE == Decimal("84000")

    def test_additional_entities(self) -> None:
        assert _ADDITIONAL_ENTITIES == 3

    def test_total_fee(self) -> None:
        assert _TOTAL_FEE == Decimal("252000.00") + Decimal("117000") + Decimal("59000") + Decimal("84000")
        assert _TOTAL_FEE == Decimal("512000.00")


# ---------------------------------------------------------------------------
# Input file existence and format validity
# ---------------------------------------------------------------------------


class TestInputFilesExist:
    """All 3 input files must be generated."""

    def test_client_profile_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        assert (out / _INPUT_DIR / "client_profile_eu.docx").exists()

    def test_fee_schedule_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        assert (out / _INPUT_DIR / "fee_schedule_eu.xlsx").exists()

    def test_engagement_letter_template_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        assert (out / _INPUT_DIR / "engagement_letter_template_eu.docx").exists()

    def test_client_profile_opens_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile_eu.docx"))
        assert len(doc.paragraphs) > 0

    def test_fee_schedule_opens_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule_eu.xlsx", data_only=True)
        assert wb.sheetnames

    def test_engagement_template_opens_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template_eu.docx"))
        assert len(doc.paragraphs) > 0


# ---------------------------------------------------------------------------
# Client profile content
# ---------------------------------------------------------------------------


class TestClientProfile:
    """client_profile_eu.docx must contain EU entity details, revenue tier, complexity factors."""

    def test_has_company_name(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert _CLIENT_NAME in text

    def test_has_revenue_tier(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "\u20ac75M" in text and "\u20ac250M" in text

    def test_has_entity_count(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "3 subsidiaries" in text

    def test_has_ifrs_adoption(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "IFRS" in text
        assert "first-time" in text.lower() or "First-Time" in text

    def test_has_key_contacts(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Hans van der Berg" in text
        assert "Isabelle Moreau" in text

    def test_has_entity_table(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile_eu.docx"))
        assert len(doc.tables) >= 1, "Client profile should have an entity table"
        table = doc.tables[0]
        assert len(table.rows) == 5  # header + 4 entities

    def test_has_jurisdictions(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "client_profile_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Netherlands" in text
        assert "Germany" in text
        assert "France" in text
        assert "United Kingdom" in text


# ---------------------------------------------------------------------------
# Fee schedule content
# ---------------------------------------------------------------------------


class TestFeeSchedule:
    """fee_schedule_eu.xlsx must have EUR fee matrix with 4 service lines."""

    def test_has_fee_schedule_sheet(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule_eu.xlsx", data_only=True)
        assert "Fee Schedule" in wb.sheetnames

    def test_header_row(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule_eu.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        headers = [ws.cell(row=1, column=c).value for c in range(1, 7)]
        assert "Service Type" in headers
        assert "Revenue Tier" in headers
        assert "Complexity Multiplier" in headers

    def test_audit_row_for_cascade_tier(self) -> None:
        """Cascade's audit row: EUR 75M-250M tier, 4-6 entities, base EUR 210,000."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule_eu.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        found = False
        for row in range(2, ws.max_row + 1):
            stype = ws.cell(row=row, column=1).value
            tier = ws.cell(row=row, column=2).value
            ec = ws.cell(row=row, column=3).value
            base = ws.cell(row=row, column=4).value
            if stype and "Audit" in stype and tier and tier == "\u20ac75M\u2013\u20ac250M" and ec == "4\u20136":
                assert base == 210000, f"Audit base fee should be 210000, got {base}"
                found = True
                break
        assert found, "Missing Statutory Audit row for \u20ac75M\u2013\u20ac250M / 4\u20136 entities"

    def test_cit_row_for_cascade_tier(self) -> None:
        """Cascade's CIT row: EUR 75M-250M tier, 4-6 entities, base EUR 72,000."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule_eu.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        found = False
        for row in range(2, ws.max_row + 1):
            stype = ws.cell(row=row, column=1).value
            tier = ws.cell(row=row, column=2).value
            ec = ws.cell(row=row, column=3).value
            base = ws.cell(row=row, column=4).value
            if stype and "CIT" in stype and tier and tier == "\u20ac75M\u2013\u20ac250M" and ec == "4\u20136":
                assert base == 72000, f"CIT base fee should be 72000, got {base}"
                found = True
                break
        assert found, "Missing CIT Compliance row for \u20ac75M\u2013\u20ac250M / 4\u20136 entities"

    def test_vat_row_for_cascade_tier(self) -> None:
        """Cascade's VAT row: EUR 75M-250M tier, 4-6 entities, base EUR 35,000."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule_eu.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        found = False
        for row in range(2, ws.max_row + 1):
            stype = ws.cell(row=row, column=1).value
            tier = ws.cell(row=row, column=2).value
            ec = ws.cell(row=row, column=3).value
            base = ws.cell(row=row, column=4).value
            if stype and "VAT" in stype and tier and tier == "\u20ac75M\u2013\u20ac250M" and ec == "4\u20136":
                assert base == 35000, f"VAT base fee should be 35000, got {base}"
                found = True
                break
        assert found, "Missing VAT Compliance row for \u20ac75M\u2013\u20ac250M / 4\u20136 entities"

    def test_tp_row_for_cascade_tier(self) -> None:
        """Cascade's TP row: EUR 75M-250M tier, 4-6 entities, base EUR 48,000."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule_eu.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        found = False
        for row in range(2, ws.max_row + 1):
            stype = ws.cell(row=row, column=1).value
            tier = ws.cell(row=row, column=2).value
            ec = ws.cell(row=row, column=3).value
            base = ws.cell(row=row, column=4).value
            if stype and "Transfer" in stype and tier and tier == "\u20ac75M\u2013\u20ac250M" and ec == "4\u20136":
                assert base == 48000, f"TP base fee should be 48000, got {base}"
                found = True
                break
        assert found, "Missing Transfer Pricing row for \u20ac75M\u2013\u20ac250M / 4\u20136 entities"

    def test_has_ifrs_complexity_note(self) -> None:
        """Fee schedule must document the 1.20x IFRS first-time adoption multiplier."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule_eu.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        all_values = []
        for row in range(1, ws.max_row + 1):
            for col in range(1, ws.max_column + 1):
                val = ws.cell(row=row, column=col).value
                if val:
                    all_values.append(str(val))
        text = " ".join(all_values)
        assert "1.20" in text, "Fee schedule should mention 1.20x IFRS multiplier"
        assert "IFRS" in text, "Fee schedule should mention IFRS"

    def test_cascade_fee_summary_section(self) -> None:
        """Fee schedule has a Cascade Europe summary section."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule_eu.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        all_values = []
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if val:
                all_values.append(str(val))
        text = " ".join(all_values)
        assert "Cascade Europe" in text, "Fee schedule should have Cascade Europe summary section"


# ---------------------------------------------------------------------------
# Planted error ERR-EU-016
# ---------------------------------------------------------------------------


class TestPlantedErrorEU016:
    """ERR-EU-016: fee_schedule_eu.xlsx audit total omits IFRS complexity multiplier."""

    def test_error_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-EU-016" in errors.entries, "ERR-EU-016 must be registered"

    def test_error_type_is_formula_error(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-016"]
        assert err.type == "formula_error"

    def test_error_severity_is_material(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-016"]
        assert err.severity == "material"

    def test_error_file_points_to_fee_schedule(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-016"]
        assert "fee_schedule_eu.xlsx" in err.file

    def test_error_catches_tc16_eu(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-016"]
        assert "TC-16-EU" in err.which_test_cases_should_catch

    def test_wrong_audit_total_in_summary(self) -> None:
        """The summary audit total should show EUR 210,000 (wrong) not EUR 252,000 (correct)."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "fee_schedule_eu.xlsx", data_only=True)
        ws = wb["Fee Schedule"]
        wrong_total = int(_AUDIT_BASE_FEE)  # 210,000 (missing 1.20x)
        correct_total = int(_AUDIT_FEE)  # 252,000
        found_wrong = False
        found_correct = False
        for row in range(1, ws.max_row + 1):
            for col in range(1, ws.max_column + 1):
                val = ws.cell(row=row, column=col).value
                if val == wrong_total:
                    row_label = ws.cell(row=row, column=1).value
                    if row_label and "Audit" in str(row_label):
                        found_wrong = True
                if val == correct_total:
                    found_correct = True
        assert found_wrong, f"Summary should show wrong audit total \u20ac{wrong_total:,}"
        assert not found_correct, f"Summary should NOT contain correct audit total \u20ac{correct_total:,}"


# ---------------------------------------------------------------------------
# Engagement letter template (merge fields, EU language)
# ---------------------------------------------------------------------------


class TestEngagementTemplate:
    """engagement_letter_template_eu.docx must have merge fields and EU language."""

    def test_has_merge_field_placeholders(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        for field in ["client_name", "fee_amount", "payment_terms", "start_date", "partner_name"]:
            assert f"<<{field}>>" in text, f"Missing merge field <<{field}>>"

    def test_has_engagement_scope_field(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "<<engagement_scope>>" in text

    def test_template_sections_present(self) -> None:
        """Template must contain key engagement letter sections."""
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        for section in [
            "Engagement Scope",
            "Fees and Billing",
            "Engagement Period",
            "Responsibilities",
            "Confidentiality",
            "Limitation of Liability",
            "ACCEPTED AND AGREED",
        ]:
            assert section in text, f"Missing template section: {section}"

    def test_has_european_firm_name(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Mitchell & Partners International LLP" in text

    def test_has_amsterdam_address(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Amsterdam" in text

    def test_has_isa_language(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "International Standards on Auditing" in text or "ISA" in text

    def test_has_oecd_language(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "OECD" in text

    def test_has_netherlands_law(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _INPUT_DIR / "engagement_letter_template_eu.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Netherlands" in text


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """All 3 input files must have canaries registered."""

    def test_all_3_canary_keys_registered(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            assert key in canaries.entries, f"Missing canary key: {key}"

    def test_canary_values_are_8_chars(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            canary = canaries.entries[key].canary
            assert len(canary) == 8, f"Canary for {key} is {len(canary)} chars: {canary}"

    def test_canaries_are_unique(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        values = [canaries.entries[k].canary for k in _CANARY_KEYS]
        assert len(values) == len(set(values)), "Duplicate canary values found"


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifestRegistration:
    """All TC-16-EU files must appear in the manifest."""

    def test_manifest_has_tc16_eu_entries(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        manifest_data = json.loads((out / "manifest.json").read_text())
        tc16_eu_entries = [e for e in manifest_data if "TC-16-EU" in e.get("test_cases", [])]
        assert len(tc16_eu_entries) >= 3, (
            f"Expected at least 3 TC-16-EU manifest entries, got {len(tc16_eu_entries)}"
        )


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """TC-16-EU must have prompt.md and expected_behavior.md."""

    def test_prompt_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-16-EU/prompt.md"
        assert path.exists()
        text = path.read_text()
        assert "engagement letter" in text.lower()

    def test_prompt_mentions_fee_schedule(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16-EU/prompt.md").read_text()
        assert "fee schedule" in text.lower()

    def test_prompt_mentions_ifrs_complexity(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16-EU/prompt.md").read_text()
        assert "1.20" in text

    def test_prompt_mentions_pieter_de_jong(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16-EU/prompt.md").read_text()
        assert "Pieter de Jong" in text

    def test_prompt_mentions_4_service_lines(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16-EU/prompt.md").read_text()
        assert "statutory audit" in text.lower()
        assert "CIT" in text
        assert "VAT" in text
        assert "transfer pricing" in text.lower()

    def test_expected_behavior_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-16-EU/expected_behavior.md"
        assert path.exists()
        text = path.read_text()
        assert len(text) > 100

    def test_expected_behavior_fee_calculation(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16-EU/expected_behavior.md").read_text()
        assert "\u20ac252,000" in text, "Expected behavior should mention audit fee \u20ac252,000"
        assert "\u20ac117,000" in text, "Expected behavior should mention CIT fee \u20ac117,000"
        assert "\u20ac59,000" in text, "Expected behavior should mention VAT fee \u20ac59,000"
        assert "\u20ac84,000" in text, "Expected behavior should mention TP fee \u20ac84,000"
        assert "\u20ac512,000" in text, "Expected behavior should mention total \u20ac512,000"

    def test_expected_behavior_merge_fields(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16-EU/expected_behavior.md").read_text()
        assert "client_name" in text
        assert "fee_amount" in text
        assert "partner_name" in text

    def test_expected_behavior_eu_language(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-16-EU/expected_behavior.md").read_text()
        assert "ISA" in text or "IFRS" in text
        assert "Mitchell & Partners International LLP" in text


# ---------------------------------------------------------------------------
# Gold standard registration
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """TC-16-EU gold standard must be registered with correct fee values."""

    def test_gold_json_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "gold_standards/TC-16-EU_gold.json"
        assert path.exists()

    def test_gold_has_fee_calculation(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-16-EU_gold.json").read_text())
        fee_calc = gold["expected_outputs"]["fee_calculation"]
        assert fee_calc["audit_fee"] == "\u20ac252,000"
        assert fee_calc["cit_fee"] == "\u20ac117,000"
        assert fee_calc["vat_fee"] == "\u20ac59,000"
        assert fee_calc["tp_fee"] == "\u20ac84,000"
        assert fee_calc["total_engagement_fee"] == "\u20ac512,000"

    def test_gold_has_merge_fields(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-16-EU_gold.json").read_text())
        merge = gold["expected_outputs"]["merge_fields"]
        assert merge["client_name"] == _CLIENT_NAME
        assert merge["partner_name"] == _PARTNER_NAME
        assert merge["start_date"] == _START_DATE
        assert merge["payment_terms"] == _PAYMENT_TERMS

    def test_gold_has_error_detection(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-16-EU_gold.json").read_text())
        assert "ERR-EU-016" in gold["error_detection"]

    def test_gold_has_canary_verification(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-16-EU_gold.json").read_text())
        assert "read_client_profile" in gold["canary_verification"]
        assert "read_fee_schedule" in gold["canary_verification"]
        assert "read_engagement_template" in gold["canary_verification"]

    def test_gold_has_template_sections(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-16-EU_gold.json").read_text())
        sections = gold["expected_outputs"]["template_sections_preserved"]
        assert "Engagement Scope" in sections
        assert "Fees and Billing" in sections
        assert "ACCEPTED AND AGREED" in sections
