"""Tests for TC-18 — Prior Year Workpaper Rollforward formatter.

Verifies:
- 6 prior-year xlsx workpapers (revenue, expenses, balance_sheet, cash, fixed_assets, leases)
- 4 prior-year docx memos (planning, risk_assessment, summary, management_letter)
- 5 current-year data files with format changes (TB now CSV, projections now docx)
- New goodwill_impairment_analysis.xlsx (not in prior year)
- Format change detection (xlsx→CSV for TB, xlsx→docx for projections)
- ERR-010 planted error (stale data in revenue workpaper)
- Canary embedding in all 15 files
- Prompt and expected behavior markdown files
"""

from __future__ import annotations

import csv
import json
import tempfile
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc18 import _CANARY_KEYS, emit_tc18
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc18 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc18_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()
        emit_tc18(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)
        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


# ---------------------------------------------------------------------------
# Prior Year Workpapers — 6 xlsx files
# ---------------------------------------------------------------------------

_PY_DIR = "test_cases/TC-18/input_files/prior_year_workpapers"

_XLSX_WORKPAPERS = [
    "wp_revenue_fy2024.xlsx",
    "wp_expenses_fy2024.xlsx",
    "wp_balance_sheet_fy2024.xlsx",
    "wp_cash_fy2024.xlsx",
    "wp_fixed_assets_fy2024.xlsx",
    "wp_leases_fy2024.xlsx",
]


class TestPriorYearXlsxWorkpapers:
    """6 prior-year xlsx workpapers must exist and be valid."""

    def test_all_six_xlsx_exist(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _XLSX_WORKPAPERS:
            path = out / _PY_DIR / name
            assert path.exists(), f"Missing prior-year workpaper: {name}"

    def test_xlsx_files_open_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _XLSX_WORKPAPERS:
            path = out / _PY_DIR / name
            wb = openpyxl.load_workbook(path, data_only=True)
            assert wb.sheetnames, f"{name} has no sheets"

    def test_revenue_wp_has_data_rows(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_revenue_fy2024.xlsx", data_only=True)
        ws = wb.active
        # Header row + at least some revenue accounts
        assert ws.max_row >= 3, "Revenue workpaper should have header + data rows"

    def test_expenses_wp_has_cogs_and_opex(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_expenses_fy2024.xlsx", data_only=True)
        ws = wb.active
        # Should have header + expense rows + subtotals
        assert ws.max_row >= 5, "Expense workpaper too small"
        # Check subtotal labels exist
        labels = [ws.cell(row=r, column=2).value for r in range(1, ws.max_row + 1)]
        assert "Total COGS" in labels
        assert "Total OpEx" in labels

    def test_balance_sheet_wp_has_classifications(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_balance_sheet_fy2024.xlsx", data_only=True)
        ws = wb.active
        # Header should include Classification column
        assert ws.cell(row=1, column=4).value == "Classification"

    def test_fixed_assets_wp_has_depreciation(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_fixed_assets_fy2024.xlsx", data_only=True)
        ws = wb.active
        headers = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
        assert "Accum. Depr." in headers or "Accumulated Depreciation" in headers or any(
            h and "Depr" in h for h in headers if h
        ), f"Fixed assets WP missing depreciation column. Headers: {headers}"

    def test_leases_wp_has_lease_rows(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_leases_fy2024.xlsx", data_only=True)
        ws = wb.active
        # Should have header + lease rows
        assert ws.max_row >= 3, "Lease workpaper should have multiple lease entries"


# ---------------------------------------------------------------------------
# Prior Year Workpapers — 4 docx memos
# ---------------------------------------------------------------------------

_DOCX_MEMOS = [
    "memo_planning_fy2024.docx",
    "memo_risk_assessment_fy2024.docx",
    "memo_summary_fy2024.docx",
    "memo_management_letter_fy2024.docx",
]


class TestPriorYearDocxMemos:
    """4 prior-year docx memos must exist and be valid."""

    def test_all_four_docx_exist(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _DOCX_MEMOS:
            path = out / _PY_DIR / name
            assert path.exists(), f"Missing prior-year memo: {name}"

    def test_docx_files_open_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _DOCX_MEMOS:
            path = out / _PY_DIR / name
            doc = Document(str(path))
            assert len(doc.paragraphs) > 0, f"{name} has no paragraphs"

    def test_planning_memo_has_content(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _PY_DIR / "memo_planning_fy2024.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Cascade" in text, "Planning memo should reference Cascade Industries"

    def test_management_letter_has_content(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _PY_DIR / "memo_management_letter_fy2024.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert len(text) > 100, "Management letter should have substantial content"


# ---------------------------------------------------------------------------
# Current Year Data — 5 files with format changes
# ---------------------------------------------------------------------------

_CY_DIR = "test_cases/TC-18/input_files/current_year_data"


class TestCurrentYearData:
    """5 current-year data files with deliberate format changes."""

    def test_trial_balance_is_csv(self) -> None:
        """TB format change: prior year was xlsx, current year is CSV."""
        _, out, _, _, _ = _ensure_emitted()
        csv_path = out / _CY_DIR / "trial_balance_fy2025.csv"
        assert csv_path.exists(), "Current year TB should be CSV (format change)"
        # Verify it's a valid CSV
        with open(csv_path, newline="") as f:
            lines = f.readlines()
        # First line may be canary comment
        data_lines = [line for line in lines if not line.startswith("#")]
        reader = csv.reader(data_lines)
        header = next(reader)
        assert "Account" in header
        rows = list(reader)
        assert len(rows) >= 5, "TB CSV should have multiple account rows"

    def test_bank_statements_csv_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / _CY_DIR / "bank_statements_fy2025.csv"
        assert path.exists()
        with open(path, newline="") as f:
            lines = [line for line in f.readlines() if not line.startswith("#")]
        reader = csv.reader(lines)
        header = next(reader)
        assert "Statement Balance" in header

    def test_lease_schedule_xlsx_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / _CY_DIR / "lease_schedule_fy2025.xlsx"
        assert path.exists()
        wb = openpyxl.load_workbook(path, data_only=True)
        ws = wb.active
        assert ws.max_row >= 3, "Lease schedule should have multiple leases"

    def test_lease_schedule_has_new_leases(self) -> None:
        """FY2025 lease schedule should have leases not in the FY2024 workpaper."""
        _, out, _, _, _ = _ensure_emitted()
        # Read FY2024 lease IDs
        wb_2024 = openpyxl.load_workbook(
            out / _PY_DIR / "wp_leases_fy2024.xlsx", data_only=True,
        )
        ws_2024 = wb_2024.active
        fy2024_ids = set()
        for row in range(2, ws_2024.max_row + 1):
            val = ws_2024.cell(row=row, column=1).value
            if val and str(val).startswith("LS-"):
                fy2024_ids.add(str(val))

        # Read FY2025 lease IDs
        wb_2025 = openpyxl.load_workbook(
            out / _CY_DIR / "lease_schedule_fy2025.xlsx", data_only=True,
        )
        ws_2025 = wb_2025.active
        fy2025_ids = set()
        for row in range(2, ws_2025.max_row + 1):
            val = ws_2025.cell(row=row, column=1).value
            if val and str(val).startswith("LS-"):
                fy2025_ids.add(str(val))

        new_leases = fy2025_ids - fy2024_ids
        assert len(new_leases) >= 2, (
            f"Expected at least 2 new leases in FY2025, found {len(new_leases)}: {new_leases}"
        )

    def test_mgmt_projections_is_docx(self) -> None:
        """Projections format change: prior year was xlsx, current year is docx."""
        _, out, _, _, _ = _ensure_emitted()
        path = out / _CY_DIR / "management_projections_fy2025.docx"
        assert path.exists(), "Current year projections should be docx (format change)"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "FY2026" in text, "Projections should reference FY2026"
        assert "Revenue" in text

    def test_goodwill_impairment_is_new(self) -> None:
        """goodwill_impairment_analysis.xlsx is NEW — no prior year equivalent."""
        _, out, _, _, _ = _ensure_emitted()
        path = out / _CY_DIR / "goodwill_impairment_analysis.xlsx"
        assert path.exists()
        wb = openpyxl.load_workbook(path, data_only=True)
        ws = wb.active
        assert ws.title == "Goodwill Impairment"
        # Should have reporting unit rows
        assert ws.max_row >= 3

    def test_goodwill_has_reporting_units(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(
            out / _CY_DIR / "goodwill_impairment_analysis.xlsx", data_only=True,
        )
        ws = wb.active
        units = []
        for row in range(2, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if val and "Cascade" in str(val):
                units.append(val)
        assert len(units) >= 2, "Should have at least 2 reporting units"


# ---------------------------------------------------------------------------
# File count totals
# ---------------------------------------------------------------------------

class TestFileCountTotals:
    """TC-18 must produce exactly the right number of files."""

    def test_prior_year_file_count(self) -> None:
        """10 prior year workpapers: 6 xlsx + 4 docx."""
        _, out, _, _, _ = _ensure_emitted()
        py_dir = out / _PY_DIR
        xlsx_files = list(py_dir.glob("*.xlsx"))
        docx_files = list(py_dir.glob("*.docx"))
        assert len(xlsx_files) == 6, f"Expected 6 xlsx, got {len(xlsx_files)}: {[f.name for f in xlsx_files]}"
        assert len(docx_files) == 4, f"Expected 4 docx, got {len(docx_files)}: {[f.name for f in docx_files]}"

    def test_current_year_file_count(self) -> None:
        """5 current year data files: 2 csv + 1 docx + 2 xlsx."""
        _, out, _, _, _ = _ensure_emitted()
        cy_dir = out / _CY_DIR
        all_files = [f for f in cy_dir.iterdir() if f.is_file()]
        assert len(all_files) == 5, (
            f"Expected 5 current year files, got {len(all_files)}: {[f.name for f in all_files]}"
        )


# ---------------------------------------------------------------------------
# ERR-010: stale_data in revenue workpaper
# ---------------------------------------------------------------------------

class TestPlantedErrorERR010:
    """ERR-010 must be planted in wp_revenue_fy2024.xlsx."""

    def test_err010_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err_ids = [e.error_id for e in errors.entries.values()]
        assert "ERR-010" in err_ids, f"ERR-010 not in error registry: {err_ids}"

    def test_err010_targets_revenue_workpaper(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = next(e for e in errors.entries.values() if e.error_id == "ERR-010")
        assert "wp_revenue_fy2024.xlsx" in err.file
        assert err.type == "stale_data"

    def test_err010_is_stale_fy2023_value(self) -> None:
        """The planted error should show a FY2023 value instead of FY2024."""
        _, _, _, errors, _ = _ensure_emitted()
        err = next(e for e in errors.entries.values() if e.error_id == "ERR-010")
        assert "FY2023" in err.description, (
            f"ERR-010 description should mention FY2023 stale value: {err.description}"
        )


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------

class TestCanaryEmbedding:
    """All 15 TC-18 files must have canaries registered."""

    _CANARY_KEYS = sorted([
        "tc18_wp_revenue_fy2024",
        "tc18_wp_expenses_fy2024",
        "tc18_wp_balance_sheet_fy2024",
        "tc18_wp_cash_fy2024",
        "tc18_wp_fixed_assets_fy2024",
        "tc18_wp_leases_fy2024",
        "tc18_memo_planning_fy2024",
        "tc18_memo_risk_assessment_fy2024",
        "tc18_memo_summary_fy2024",
        "tc18_memo_management_letter_fy2024",
        "tc18_cy_trial_balance_fy2025",
        "tc18_cy_bank_statements_fy2025",
        "tc18_cy_lease_schedule_fy2025",
        "tc18_cy_mgmt_projections_fy2025",
        "tc18_cy_goodwill_impairment_fy2025",
    ])

    def test_all_15_canary_keys_registered(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in self._CANARY_KEYS:
            assert key in canaries.entries, f"Missing canary key: {key}"

    def test_canary_values_are_8_chars(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in self._CANARY_KEYS:
            canary = canaries.entries[key].canary
            assert len(canary) == 8, f"Canary for {key} is {len(canary)} chars: {canary}"

    def test_canaries_are_unique(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        values = [canaries.entries[k].canary for k in self._CANARY_KEYS]
        assert len(values) == len(set(values)), "Duplicate canary values found"


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------

class TestManifestRegistration:
    """All TC-18 files must appear in the manifest."""

    def test_manifest_has_tc18_entries(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        manifest_data = json.loads((out / "manifest.json").read_text())
        tc18_entries = [e for e in manifest_data if "TC-18" in e.get("test_cases", [])]
        # 15 input files
        assert len(tc18_entries) >= 15, (
            f"Expected at least 15 TC-18 manifest entries, got {len(tc18_entries)}"
        )


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------

class TestPromptAndExpectedBehavior:
    """TC-18 must have prompt.md and expected_behavior.md."""

    def test_prompt_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-18/prompt.md"
        assert path.exists()
        text = path.read_text()
        assert "Prior Year" in text

    def test_expected_behavior_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-18/expected_behavior.md"
        assert path.exists()
        text = path.read_text()
        assert len(text) > 100


# ---------------------------------------------------------------------------
# Format change detection markers
# ---------------------------------------------------------------------------

class TestFormatChangeFlags:
    """Expected behavior should call out format changes for manager attention."""

    def test_expected_behavior_mentions_format_changes(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-18/expected_behavior.md").read_text()
        # Two format changes: TB xlsx→CSV, projections xlsx→docx
        assert "format" in text.lower() or "CSV" in text, (
            "Expected behavior should mention format changes"
        )

    def test_expected_behavior_mentions_goodwill_new_area(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-18/expected_behavior.md").read_text()
        assert "goodwill" in text.lower() or "new" in text.lower(), (
            "Expected behavior should flag goodwill impairment as new audit area"
        )
