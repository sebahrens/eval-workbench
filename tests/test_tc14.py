"""Tests for TC-14 — 13-Week Cash Flow Forecast formatter.

Verifies:
- Balance sheet current xlsx (cash, AR, AP, covenant note)
- AR aging report xlsx (customer rows with DSO and collection weeks)
- AP aging report xlsx (vendor rows with payment terms and due weeks)
- Committed/discretionary expenses docx (three categories with amounts)
- Planted error ERR-008 (mismatched_total: AP total row omits one invoice, $1,800 delta)
- Planted error ERR-024 (rounding_discrepancy: Total Current Assets off by $1)
- Gold standard: cash trough Week 8 at $1,340,000 (below $2M covenant)
- Deferral scenario: marketing $168,750/week + training $35,000/week from Week 5
  → trough $2,155,000
- Sensitivity scenario: 1-week collection delay → trough $780,000 with deferrals
- Canary embedding in all 4 input files
- Prompt and expected behavior markdown files
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc14 import (
    _COMMITTED_EXPENSES,
    _DISC_EXPENSES,
    _OPENING_CASH,
    _SEMI_DISC_EXPENSES,
    _WEEKLY_AP_PAYMENTS,
    _WEEKLY_AR_COLLECTIONS,
    _compute_deferred_cash_flows,
    _compute_sensitivity_cash_flows,
    _compute_weekly_cash_flows,
    emit_tc14,
)
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc14 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None

_CANARY_KEYS = sorted([
    "tc14_balance_sheet_current",
    "tc14_ar_aging_report",
    "tc14_ap_aging_report",
    "tc14_committed_expenses",
])


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc14_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc14(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard (registered via @register_gold)
        emit_gold("TC-14", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


# ---------------------------------------------------------------------------
# Input directory
# ---------------------------------------------------------------------------

_INPUT_DIR = "test_cases/TC-14/input_files"


# ---------------------------------------------------------------------------
# Constants verification (from prompt.md §5 TC-14)
# ---------------------------------------------------------------------------


class TestCashFlowConstants:
    """Verify formatter constants match prompt.md gold standard."""

    def test_opening_cash(self) -> None:
        assert _OPENING_CASH == 5_200_000

    def test_thirteen_weeks_ar(self) -> None:
        assert len(_WEEKLY_AR_COLLECTIONS) == 13

    def test_thirteen_weeks_ap(self) -> None:
        assert len(_WEEKLY_AP_PAYMENTS) == 13

    def test_committed_expense_categories(self) -> None:
        names = [e[0] for e in _COMMITTED_EXPENSES]
        assert "Payroll" in names
        assert "Rent" in names
        assert "Insurance" in names
        assert "Utilities" in names

    def test_discretionary_expense_categories(self) -> None:
        names = [e[0] for e in _DISC_EXPENSES]
        assert "Marketing" in names
        assert "Training" in names
        assert "Travel" in names
        assert "Bonuses" in names

    def test_marketing_amount(self) -> None:
        marketing = next(e for e in _DISC_EXPENSES if e[0] == "Marketing")
        assert marketing[2] == 168_750

    def test_training_amount(self) -> None:
        training = next(e for e in _DISC_EXPENSES if e[0] == "Training")
        assert training[2] == 35_000

    def test_semi_disc_categories(self) -> None:
        names = [e[0] for e in _SEMI_DISC_EXPENSES]
        assert "Maintenance" in names
        assert "Professional Fees" in names


# ---------------------------------------------------------------------------
# Cash flow model verification — gold standard numbers
# ---------------------------------------------------------------------------


class TestBaseScenario:
    """Verify the base 13-week cash flow forecast matches gold standard."""

    def test_trough_week_is_8(self) -> None:
        flows = _compute_weekly_cash_flows()
        trough = min(flows, key=lambda w: w["closing_balance"])
        assert trough["week"] == 8

    def test_trough_balance_1_340_000(self) -> None:
        flows = _compute_weekly_cash_flows()
        trough = min(flows, key=lambda w: w["closing_balance"])
        assert trough["closing_balance"] == 1_340_000

    def test_covenant_breached(self) -> None:
        flows = _compute_weekly_cash_flows()
        trough = min(flows, key=lambda w: w["closing_balance"])
        assert trough["closing_balance"] < 2_000_000

    def test_opening_balance_week1(self) -> None:
        flows = _compute_weekly_cash_flows()
        assert flows[0]["opening_balance"] == _OPENING_CASH

    def test_thirteen_weeks_computed(self) -> None:
        flows = _compute_weekly_cash_flows()
        assert len(flows) == 13
        assert flows[0]["week"] == 1
        assert flows[-1]["week"] == 13

    def test_closing_equals_opening_plus_net(self) -> None:
        """Verify balance continuity across all weeks."""
        flows = _compute_weekly_cash_flows()
        for w in flows:
            assert w["closing_balance"] == w["opening_balance"] + w["net_cash_flow"]

    def test_balance_continuity(self) -> None:
        """Closing balance of week N = opening balance of week N+1."""
        flows = _compute_weekly_cash_flows()
        for i in range(len(flows) - 1):
            assert flows[i]["closing_balance"] == flows[i + 1]["opening_balance"]


class TestDeferralScenario:
    """Verify deferral scenario: marketing + training deferred from Week 5."""

    def test_deferral_trough_week_8(self) -> None:
        flows = _compute_deferred_cash_flows()
        trough = min(flows, key=lambda w: w["closing_balance"])
        assert trough["week"] == 8

    def test_deferral_trough_balance_2_155_000(self) -> None:
        flows = _compute_deferred_cash_flows()
        trough = min(flows, key=lambda w: w["closing_balance"])
        assert trough["closing_balance"] == 2_155_000

    def test_deferral_avoids_covenant_breach(self) -> None:
        flows = _compute_deferred_cash_flows()
        trough = min(flows, key=lambda w: w["closing_balance"])
        assert trough["closing_balance"] > 2_000_000


class TestSensitivityScenario:
    """Verify sensitivity scenario: 1-week collection delay with deferrals."""

    def test_sensitivity_trough_780_000(self) -> None:
        flows = _compute_sensitivity_cash_flows()
        trough = min(flows, key=lambda w: w["closing_balance"])
        assert trough["closing_balance"] == 780_000

    def test_sensitivity_still_breaches_covenant(self) -> None:
        flows = _compute_sensitivity_cash_flows()
        trough = min(flows, key=lambda w: w["closing_balance"])
        assert trough["closing_balance"] < 2_000_000

    def test_sensitivity_week1_zero_collections(self) -> None:
        """With 1-week delay, Week 1 should have zero AR collections."""
        flows = _compute_sensitivity_cash_flows()
        # Net cash flow in week 1 should reflect zero inflows
        base_flows = _compute_weekly_cash_flows()
        # Sensitivity Week 1 net should be much worse than base Week 1
        assert flows[0]["net_cash_flow"] < base_flows[0]["net_cash_flow"]


# ---------------------------------------------------------------------------
# Balance sheet xlsx
# ---------------------------------------------------------------------------


class TestBalanceSheet:
    """Verify balance_sheet_current.xlsx."""

    def test_file_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / _INPUT_DIR / "balance_sheet_current.xlsx"
        assert path.exists()

    def test_sheet_name(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "balance_sheet_current.xlsx")
        assert "Balance Sheet" in wb.sheetnames

    def test_title_row(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "balance_sheet_current.xlsx")
        ws = wb["Balance Sheet"]
        assert "Cascade Industries" in str(ws["A1"].value)

    def test_cash_balance_matches_opening(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "balance_sheet_current.xlsx")
        ws = wb["Balance Sheet"]
        # Find Cash and Cash Equivalents row
        for row in ws.iter_rows(min_row=1, max_col=2, values_only=False):
            if row[0].value and "Cash" in str(row[0].value):
                assert row[1].value == _OPENING_CASH
                return
        raise AssertionError("Cash and Cash Equivalents row not found")

    def test_covenant_note_present(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "balance_sheet_current.xlsx")
        ws = wb["Balance Sheet"]
        found = False
        for row in ws.iter_rows(values_only=True):
            for cell in row:
                if cell and "2,000,000" in str(cell) and "covenant" in str(cell).lower():
                    found = True
                    break
        assert found, "Minimum liquidity covenant note not found"


# ---------------------------------------------------------------------------
# AR aging report xlsx
# ---------------------------------------------------------------------------


class TestARAgingReport:
    """Verify ar_aging_report.xlsx."""

    def test_file_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / _INPUT_DIR / "ar_aging_report.xlsx"
        assert path.exists()

    def test_sheet_name(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "ar_aging_report.xlsx")
        assert "AR Aging" in wb.sheetnames

    def test_headers_present(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "ar_aging_report.xlsx")
        ws = wb["AR Aging"]
        headers = [ws.cell(row=1, column=c).value for c in range(1, 12)]
        assert "Customer ID" in headers
        assert "Customer Name" in headers
        assert "DSO (days)" in headers
        assert "Expected Collection (weeks)" in headers
        assert "Total AR" in headers

    def test_has_customer_rows(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "ar_aging_report.xlsx")
        ws = wb["AR Aging"]
        # At least one data row
        assert ws.cell(row=2, column=1).value is not None

    def test_dso_column_positive(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "ar_aging_report.xlsx")
        ws = wb["AR Aging"]
        for row in range(2, ws.max_row + 1):
            dso = ws.cell(row=row, column=10).value
            if dso is not None:
                assert dso > 0


# ---------------------------------------------------------------------------
# AP aging report xlsx
# ---------------------------------------------------------------------------


class TestAPAgingReport:
    """Verify ap_aging_report.xlsx."""

    def test_file_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / _INPUT_DIR / "ap_aging_report.xlsx"
        assert path.exists()

    def test_sheet_name(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "ap_aging_report.xlsx")
        assert "AP Aging" in wb.sheetnames

    def test_headers_present(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "ap_aging_report.xlsx")
        ws = wb["AP Aging"]
        headers = [ws.cell(row=1, column=c).value for c in range(1, 11)]
        assert "Vendor ID" in headers
        assert "Vendor Name" in headers
        assert "Total AP" in headers
        assert "Payment Terms" in headers
        assert "Due (weeks)" in headers

    def test_has_vendor_rows(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "ap_aging_report.xlsx")
        ws = wb["AP Aging"]
        assert ws.cell(row=2, column=1).value is not None

    def test_total_row_present(self) -> None:
        """The TOTAL row should exist (even though it has an error)."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "ap_aging_report.xlsx")
        ws = wb["AP Aging"]
        found_total = False
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if val and str(val).strip() == "TOTAL":
                found_total = True
                break
        assert found_total, "TOTAL row not found in AP aging report"


# ---------------------------------------------------------------------------
# Planted errors
# ---------------------------------------------------------------------------


class TestPlantedErrors:
    """Verify planted errors ERR-008 and ERR-024."""

    def test_err008_ap_total_mismatch(self) -> None:
        """ERR-008: AP aging total row should be $1,800 less than actual sum."""
        _, out, _, errors, _ = _ensure_emitted()
        err = errors.get("ERR-008")
        assert err is not None
        assert err.type == "mismatched_total"
        assert "ap_aging_report.xlsx" in err.file
        assert err.severity == "material"

        # Verify the delta in the actual file
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "ap_aging_report.xlsx")
        ws = wb["AP Aging"]

        # Sum the individual vendor totals
        vendor_totals = []
        for row in range(2, ws.max_row + 1):
            val_a = ws.cell(row=row, column=1).value
            if val_a and str(val_a).strip() == "TOTAL":
                break
            total_val = ws.cell(row=row, column=8).value
            if total_val is not None:
                vendor_totals.append(total_val)

        actual_sum = sum(vendor_totals)

        # Find the TOTAL row's value
        for row in range(2, ws.max_row + 1):
            val_a = ws.cell(row=row, column=1).value
            if val_a and str(val_a).strip() == "TOTAL":
                shown_total = ws.cell(row=row, column=8).value
                break

        assert actual_sum - shown_total == 1_800

    def test_err024_balance_sheet_rounding(self) -> None:
        """ERR-024: Total Current Assets should differ by $1 from component sum."""
        _, out, _, errors, _ = _ensure_emitted()
        err = errors.get("ERR-024")
        assert err is not None
        assert err.type == "rounding_discrepancy"
        assert "balance_sheet_current.xlsx" in err.file
        assert err.severity == "immaterial"

    def test_err008_test_case_attribution(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.get("ERR-008")
        assert "TC-14" in err.which_test_cases_should_catch

    def test_err024_test_case_attribution(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.get("ERR-024")
        assert "TC-14" in err.which_test_cases_should_catch


# ---------------------------------------------------------------------------
# Expenses docx
# ---------------------------------------------------------------------------


class TestExpensesDocx:
    """Verify committed_discretionary_expenses.docx."""

    def test_file_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / _INPUT_DIR / "committed_discretionary_expenses.docx"
        assert path.exists()

    def test_has_three_expense_categories(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(out / _INPUT_DIR / "committed_discretionary_expenses.docx")
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        heading_text = " ".join(headings).lower()
        assert "committed" in heading_text
        assert "semi-discretionary" in heading_text
        assert "discretionary" in heading_text

    def test_covenant_section_present(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(out / _INPUT_DIR / "committed_discretionary_expenses.docx")
        text = " ".join(p.text for p in doc.paragraphs)
        assert "2,000,000" in text
        assert "covenant" in text.lower()

    def test_tables_have_expense_rows(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(out / _INPUT_DIR / "committed_discretionary_expenses.docx")
        # Should have 3 tables (committed, semi-disc, discretionary)
        assert len(doc.tables) == 3
        # Each table should have header + data rows
        for table in doc.tables:
            assert len(table.rows) > 1  # at least header + one data row


# ---------------------------------------------------------------------------
# Canary verification
# ---------------------------------------------------------------------------


class TestCanaries:
    """Verify canary embedding in all TC-14 input files."""

    def test_balance_sheet_canary(self) -> None:
        _, out, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc14_balance_sheet_current")
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "balance_sheet_current.xlsx")
        assert canary in (wb.properties.description or ""), (
            f"Canary {canary} not in document properties"
        )

    def test_ar_aging_canary(self) -> None:
        _, out, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc14_ar_aging_report")
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "ar_aging_report.xlsx")
        assert canary in (wb.properties.description or ""), (
            f"Canary {canary} not in document properties"
        )

    def test_ap_aging_canary(self) -> None:
        _, out, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc14_ap_aging_report")
        wb = openpyxl.load_workbook(out / _INPUT_DIR / "ap_aging_report.xlsx")
        assert canary in (wb.properties.description or ""), (
            f"Canary {canary} not in document properties"
        )

    def test_expenses_docx_canary(self) -> None:
        _, out, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc14_committed_expenses")
        doc = Document(out / _INPUT_DIR / "committed_discretionary_expenses.docx")
        assert canary in (doc.core_properties.comments or ""), (
            f"Canary {canary} not in core properties comments"
        )


# ---------------------------------------------------------------------------
# Prompt and expected behavior markdown
# ---------------------------------------------------------------------------


class TestMarkdownFiles:
    """Verify prompt.md and expected_behavior.md."""

    def test_prompt_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-14/prompt.md"
        assert path.exists()

    def test_prompt_mentions_13_week(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-14/prompt.md").read_text()
        assert "13-week" in text

    def test_prompt_mentions_covenant(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-14/prompt.md").read_text()
        assert "2,000,000" in text

    def test_prompt_mentions_sensitivity(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-14/prompt.md").read_text()
        assert "sensitivity" in text.lower() or "collections slow" in text.lower()

    def test_expected_behavior_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-14/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_trough(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-14/expected_behavior.md").read_text()
        assert "1,340,000" in text

    def test_expected_behavior_mentions_deferral(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-14/expected_behavior.md").read_text()
        assert "2,155,000" in text

    def test_expected_behavior_mentions_sensitivity(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-14/expected_behavior.md").read_text()
        assert "780,000" in text


# ---------------------------------------------------------------------------
# Gold standard JSON
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify TC-14 gold standard JSON."""

    def test_gold_file_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "gold_standards/TC-14_gold.json"
        assert path.exists()

    def test_gold_base_scenario(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-14_gold.json").read_text())
        base = gold["expected_outputs"]["base_scenario"]
        assert base["trough_week"] == 8
        assert base["trough_balance"] == "$1,340,000"
        assert base["covenant_breached"] is True

    def test_gold_deferral_scenario(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-14_gold.json").read_text())
        deferral = gold["expected_outputs"]["deferral_scenario"]
        assert deferral["trough_balance"] == "$2,155,000"
        assert deferral["covenant_breached"] is False
        assert deferral["deferral_start_week"] == 5
        assert "Marketing" in str(deferral["deferred_expenses"])
        assert "Training" in str(deferral["deferred_expenses"])

    def test_gold_sensitivity_scenario(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-14_gold.json").read_text())
        sens = gold["expected_outputs"]["sensitivity_scenario"]
        assert sens["trough_balance"] == "$780,000"
        assert sens["covenant_breached"] is True

    def test_gold_opening_cash(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-14_gold.json").read_text())
        assert gold["expected_outputs"]["opening_cash_balance"] == "$5,200,000"

    def test_gold_covenant_threshold(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-14_gold.json").read_text())
        assert gold["expected_outputs"]["covenant_threshold"] == "$2,000,000"

    def test_gold_weekly_detail_thirteen_weeks(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-14_gold.json").read_text())
        detail = gold["expected_outputs"]["base_scenario"]["weekly_detail"]
        assert len(detail) == 13

    def test_gold_canary_verification(self) -> None:
        _, out, canaries, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-14_gold.json").read_text())
        cv = gold["canary_verification"]
        assert cv["read_balance_sheet"] == canaries.canary_for("tc14_balance_sheet_current")
        assert cv["read_ar_aging"] == canaries.canary_for("tc14_ar_aging_report")
        assert cv["read_ap_aging"] == canaries.canary_for("tc14_ap_aging_report")
        assert cv["read_expense_schedule"] == canaries.canary_for("tc14_committed_expenses")

    def test_gold_error_detection_keys(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-14_gold.json").read_text())
        assert "ERR-008" in gold["error_detection"]
        assert "ERR-024" in gold["error_detection"]

    def test_gold_scoring_hints(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards/TC-14_gold.json").read_text())
        hints = gold["scoring_hints"]
        assert "correctness" in hints
        assert "completeness" in hints
        assert "format_compliance" in hints
