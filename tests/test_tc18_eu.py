"""Tests for TC-18-EU — IFRS/ISA Prior-Year Workpaper Rollforward formatter.

Verifies:
- 6 prior-year xlsx workpapers with IFRS terminology
- 4 prior-year docx memos with ISA framework references
- 5 current-year data files with format changes (TB now CSV, projections now docx)
- New goodwill_impairment_analysis_ifrs.xlsx (IAS 36, not in prior year)
- Multi-currency handling (EUR + GBP)
- ERR-EU-018 planted error (stale CP product revenue)
- Canary embedding in all 15 files
- IFRS/ISA terminology preservation
- Pillar Two threshold assessment in projections
- Gold standard with judgment traps
"""

from __future__ import annotations

import csv
import json
import tempfile
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc18_eu import _CANARY_KEYS, emit_tc18_eu
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc18_eu once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc18eu_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()
        emit_tc18_eu(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)
        emit_gold("TC-18-EU", _CANARIES, _ERRORS, _OUTPUT / "gold_standards")
        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


# ---------------------------------------------------------------------------
# Prior Year Workpapers — 6 xlsx files
# ---------------------------------------------------------------------------

_PY_DIR = "test_cases/TC-18-EU/input_files/prior_year_workpapers"

_XLSX_WORKPAPERS = [
    "wp_revenue_fy2024.xlsx",
    "wp_operating_expenses_fy2024.xlsx",
    "wp_balance_sheet_fy2024.xlsx",
    "wp_cash_fy2024.xlsx",
    "wp_fixed_assets_fy2024.xlsx",
    "wp_leases_fy2024.xlsx",
]


class TestPriorYearXlsxWorkpapers:
    """6 prior-year xlsx workpapers must exist and be valid."""

    def test_all_six_xlsx_exist(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _XLSX_WORKPAPERS:
            path = out / _PY_DIR / name
            assert path.exists(), f"Missing prior-year workpaper: {name}"

    def test_xlsx_files_open_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _XLSX_WORKPAPERS:
            path = out / _PY_DIR / name
            wb = openpyxl.load_workbook(path, data_only=True)
            assert wb.sheetnames, f"{name} has no sheets"

    def test_revenue_wp_has_entity_rows(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_revenue_fy2024.xlsx", data_only=True)
        ws = wb.active
        assert ws.max_row >= 8, "Revenue workpaper should have multiple entity revenue rows"

    def test_revenue_wp_has_ifrs15_category(self) -> None:
        """Revenue workpaper should reference IFRS 15."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_revenue_fy2024.xlsx", data_only=True)
        ws = wb.active
        header = ws.cell(row=1, column=4).value
        assert "IFRS 15" in str(header), f"Expected IFRS 15 reference in header, got: {header}"

    def test_revenue_wp_has_ic_eliminations(self) -> None:
        """Revenue workpaper should include intercompany elimination rows."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_revenue_fy2024.xlsx", data_only=True)
        ws = wb.active
        all_text = " ".join(
            str(ws.cell(row=r, column=c).value or "")
            for r in range(1, ws.max_row + 1)
            for c in range(1, ws.max_column + 1)
        )
        assert "Elimination" in all_text or "elimination" in all_text, (
            "Revenue WP should include intercompany eliminations"
        )

    def test_opex_wp_classification_by_nature(self) -> None:
        """Operating expenses should be classified by nature per IAS 1."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_operating_expenses_fy2024.xlsx", data_only=True)
        ws = wb.active
        header = ws.cell(row=1, column=1).value
        assert "IAS 1" in str(header), f"Expected IAS 1 reference, got: {header}"
        # Check for nature-based categories
        labels = [ws.cell(row=r, column=1).value for r in range(2, ws.max_row + 1)]
        nature_keywords = ["Raw materials", "Employee benefits", "Depreciation"]
        found = [kw for kw in nature_keywords if any(kw in str(lbl) for lbl in labels if lbl)]
        assert len(found) >= 2, f"Expected nature-based expense categories, found: {found}"

    def test_balance_sheet_ifrs_terminology(self) -> None:
        """Balance sheet should use 'Statement of Financial Position' per IFRS."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_balance_sheet_fy2024.xlsx", data_only=True)
        ws = wb.active
        assert ws.title == "Statement of Financial Position", (
            f"Sheet title should be IFRS terminology, got: {ws.title}"
        )

    def test_balance_sheet_has_rou_assets(self) -> None:
        """Balance sheet should show Right-of-use assets (IFRS 16)."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_balance_sheet_fy2024.xlsx", data_only=True)
        ws = wb.active
        all_text = " ".join(
            str(ws.cell(row=r, column=1).value or "")
            for r in range(1, ws.max_row + 1)
        )
        assert "Right-of-use" in all_text, "BS should show Right-of-use assets per IFRS 16"

    def test_balance_sheet_has_revaluation_reserve(self) -> None:
        """Balance sheet should include IAS 16 revaluation reserve."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_balance_sheet_fy2024.xlsx", data_only=True)
        ws = wb.active
        all_text = " ".join(
            str(ws.cell(row=r, column=1).value or "")
            for r in range(1, ws.max_row + 1)
        )
        assert "Revaluation reserve" in all_text, "BS should include IAS 16 revaluation reserve"

    def test_cash_wp_multi_currency(self) -> None:
        """Cash workpaper should show both EUR and GBP accounts."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_cash_fy2024.xlsx", data_only=True)
        ws = wb.active
        currencies = set()
        for r in range(2, ws.max_row + 1):
            val = ws.cell(row=r, column=4).value
            if val:
                currencies.add(str(val))
        assert "EUR" in currencies and "GBP" in currencies, (
            f"Cash WP should show EUR and GBP, found: {currencies}"
        )

    def test_fixed_assets_has_revaluation_column(self) -> None:
        """Fixed assets workpaper should include revaluation column for IAS 16."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_fixed_assets_fy2024.xlsx", data_only=True)
        ws = wb.active
        headers = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
        assert any("Revaluation" in str(h) for h in headers if h), (
            f"Fixed assets WP should have Revaluation column. Headers: {headers}"
        )

    def test_leases_wp_ifrs16_single_model(self) -> None:
        """Lease workpaper should reference IFRS 16 single model."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_leases_fy2024.xlsx", data_only=True)
        ws = wb.active
        assert "IFRS 16" in str(ws.title), f"Sheet title should reference IFRS 16, got: {ws.title}"

    def test_leases_wp_has_ibr_column(self) -> None:
        """Lease workpaper should show incremental borrowing rates."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_leases_fy2024.xlsx", data_only=True)
        ws = wb.active
        headers = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
        assert any("IBR" in str(h) for h in headers if h), (
            f"Lease WP should have IBR column. Headers: {headers}"
        )


# ---------------------------------------------------------------------------
# Prior Year Workpapers — 4 docx memos
# ---------------------------------------------------------------------------

_DOCX_MEMOS = [
    "memo_planning_fy2024.docx",
    "memo_risk_assessment_fy2024.docx",
    "memo_summary_fy2024.docx",
    "memo_management_letter_fy2024.docx",
]


class TestPriorYearDocxMemos:
    """4 prior-year docx memos must exist with ISA references."""

    def test_all_four_docx_exist(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _DOCX_MEMOS:
            path = out / _PY_DIR / name
            assert path.exists(), f"Missing prior-year memo: {name}"

    def test_docx_files_open_cleanly(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        for name in _DOCX_MEMOS:
            path = out / _PY_DIR / name
            doc = Document(str(path))
            assert len(doc.paragraphs) > 0, f"{name} has no paragraphs"

    def test_planning_memo_references_isa(self) -> None:
        """Planning memo should reference ISA standards, not PCAOB."""
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _PY_DIR / "memo_planning_fy2024.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "ISA 300" in text, "Planning memo should reference ISA 300"
        assert "ISA 315" in text, "Planning memo should reference ISA 315"
        assert "ISA 600" in text, "Planning memo should reference ISA 600"
        assert "Cascade Europe" in text, "Planning memo should reference Cascade Europe"

    def test_planning_memo_has_component_auditors(self) -> None:
        """Planning memo should describe component auditor arrangements."""
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _PY_DIR / "memo_planning_fy2024.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "component auditor" in text.lower(), "Planning memo should discuss component auditors"

    def test_risk_assessment_references_isa315(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _PY_DIR / "memo_risk_assessment_fy2024.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "ISA 315" in text, "Risk assessment should reference ISA 315"

    def test_summary_memo_references_kam(self) -> None:
        """Summary memo should reference Key Audit Matters (ISA 701), not CAM."""
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _PY_DIR / "memo_summary_fy2024.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Key Audit Matters" in text or "KAM" in text, (
            "Summary memo should reference KAM (ISA 701)"
        )
        assert "ISA 701" in text, "Summary memo should reference ISA 701"

    def test_management_letter_eu_content(self) -> None:
        """Management letter should reference EU-specific matters."""
        _, out, _, _, _ = _ensure_emitted()
        doc = Document(str(out / _PY_DIR / "memo_management_letter_fy2024.docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert len(text) > 200, "Management letter should have substantial content"
        assert "IFRS" in text or "transfer pricing" in text.lower() or "GDPR" in text, (
            "Management letter should reference EU-specific matters"
        )


# ---------------------------------------------------------------------------
# Current Year Data — 5 files with format changes
# ---------------------------------------------------------------------------

_CY_DIR = "test_cases/TC-18-EU/input_files/current_year_data"


class TestCurrentYearData:
    """5 current-year data files with deliberate format changes."""

    def test_trial_balance_is_csv(self) -> None:
        """TB format change: prior year was xlsx, current year is CSV."""
        _, out, _, _, _ = _ensure_emitted()
        csv_path = out / _CY_DIR / "trial_balance_fy2025.csv"
        assert csv_path.exists(), "Current year TB should be CSV (format change)"
        with open(csv_path, newline="") as f:
            lines = f.readlines()
        data_lines = [line for line in lines if not line.startswith("#")]
        reader = csv.reader(data_lines)
        header = next(reader)
        assert "IFRS" in " ".join(header), "TB CSV should use IFRS account descriptions"
        rows = list(reader)
        assert len(rows) >= 10, "TB CSV should have multiple account rows"

    def test_trial_balance_has_entity_column(self) -> None:
        """TB should include entity identifiers for multi-entity group."""
        _, out, _, _, _ = _ensure_emitted()
        with open(out / _CY_DIR / "trial_balance_fy2025.csv", newline="") as f:
            lines = [line for line in f.readlines() if not line.startswith("#")]
        reader = csv.reader(lines)
        header = next(reader)
        assert "Entity" in header, f"TB CSV should have Entity column. Header: {header}"

    def test_bank_statements_multi_currency(self) -> None:
        """Bank statements should show EUR and GBP accounts."""
        _, out, _, _, _ = _ensure_emitted()
        path = out / _CY_DIR / "bank_statements_fy2025.csv"
        assert path.exists()
        with open(path, newline="") as f:
            lines = [line for line in f.readlines() if not line.startswith("#")]
        reader = csv.reader(lines)
        header = next(reader)
        assert "Currency" in header
        rows = list(reader)
        currencies = {r[header.index("Currency")] for r in rows}
        assert "EUR" in currencies and "GBP" in currencies, (
            f"Bank statements should have EUR and GBP, found: {currencies}"
        )

    def test_lease_schedule_has_new_leases(self) -> None:
        """FY2025 lease schedule should have leases not in FY2024 workpaper."""
        _, out, _, _, _ = _ensure_emitted()
        # Read FY2024 lease IDs
        wb_2024 = openpyxl.load_workbook(
            out / _PY_DIR / "wp_leases_fy2024.xlsx", data_only=True,
        )
        ws_2024 = wb_2024.active
        fy2024_ids = set()
        for row in range(2, ws_2024.max_row + 1):
            val = ws_2024.cell(row=row, column=1).value
            if val and str(val).startswith("EU-LS-"):
                fy2024_ids.add(str(val))

        # Read FY2025 lease IDs
        wb_2025 = openpyxl.load_workbook(
            out / _CY_DIR / "lease_schedule_fy2025.xlsx", data_only=True,
        )
        ws_2025 = wb_2025.active
        fy2025_ids = set()
        for row in range(2, ws_2025.max_row + 1):
            val = ws_2025.cell(row=row, column=1).value
            if val and str(val).startswith("EU-LS-"):
                fy2025_ids.add(str(val))

        new_leases = fy2025_ids - fy2024_ids
        assert len(new_leases) >= 2, (
            f"Expected at least 2 new leases in FY2025, found {len(new_leases)}: {new_leases}"
        )

    def test_mgmt_projections_is_docx_with_pillar_two(self) -> None:
        """Projections format change: prior year was xlsx, current year is docx.
        Must include Pillar Two threshold assessment.
        """
        _, out, _, _, _ = _ensure_emitted()
        path = out / _CY_DIR / "management_projections_fy2025.docx"
        assert path.exists(), "Current year projections should be docx (format change)"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Pillar Two" in text, "Projections should include Pillar Two assessment"
        assert "750" in text, "Projections should reference €750M threshold"
        assert "does NOT" in text or "not apply" in text.lower(), (
            "Projections should state Pillar Two does not apply"
        )

    def test_goodwill_impairment_is_ias36(self) -> None:
        """goodwill_impairment_analysis_ifrs.xlsx — IAS 36 CGU-based."""
        _, out, _, _, _ = _ensure_emitted()
        path = out / _CY_DIR / "goodwill_impairment_analysis_ifrs.xlsx"
        assert path.exists()
        wb = openpyxl.load_workbook(path, data_only=True)
        ws = wb.active
        assert "IAS 36" in str(ws.title), f"Sheet should reference IAS 36, got: {ws.title}"

    def test_goodwill_has_cgus(self) -> None:
        """Goodwill impairment should list cash-generating units."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(
            out / _CY_DIR / "goodwill_impairment_analysis_ifrs.xlsx", data_only=True,
        )
        ws = wb.active
        header = ws.cell(row=1, column=1).value
        assert "CGU" in str(header), f"Expected CGU in header, got: {header}"
        # Should have at least 3 CGUs
        cgus = []
        for row in range(2, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if val and "(" in str(val):
                cgus.append(val)
        assert len(cgus) >= 3, f"Expected at least 3 CGUs, found: {cgus}"

    def test_goodwill_uses_pretax_discount_rate(self) -> None:
        """IAS 36 requires pre-tax discount rates."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(
            out / _CY_DIR / "goodwill_impairment_analysis_ifrs.xlsx", data_only=True,
        )
        ws = wb.active
        all_text = " ".join(
            str(ws.cell(row=r, column=c).value or "")
            for r in range(1, ws.max_row + 1)
            for c in range(1, ws.max_column + 1)
        )
        assert "pre-tax" in all_text.lower(), "IAS 36 impairment should use pre-tax discount rates"


# ---------------------------------------------------------------------------
# File count totals
# ---------------------------------------------------------------------------

class TestFileCountTotals:
    """TC-18-EU must produce exactly the right number of files."""

    def test_prior_year_file_count(self) -> None:
        """10 prior year workpapers: 6 xlsx + 4 docx."""
        _, out, _, _, _ = _ensure_emitted()
        py_dir = out / _PY_DIR
        xlsx_files = list(py_dir.glob("*.xlsx"))
        docx_files = list(py_dir.glob("*.docx"))
        assert len(xlsx_files) == 6, f"Expected 6 xlsx, got {len(xlsx_files)}: {[f.name for f in xlsx_files]}"
        assert len(docx_files) == 4, f"Expected 4 docx, got {len(docx_files)}: {[f.name for f in docx_files]}"

    def test_current_year_file_count(self) -> None:
        """5 current year data files: 2 csv + 1 docx + 2 xlsx."""
        _, out, _, _, _ = _ensure_emitted()
        cy_dir = out / _CY_DIR
        all_files = [f for f in cy_dir.iterdir() if f.is_file()]
        assert len(all_files) == 5, (
            f"Expected 5 current year files, got {len(all_files)}: {[f.name for f in all_files]}"
        )


# ---------------------------------------------------------------------------
# ERR-EU-018: stale_data in revenue workpaper
# ---------------------------------------------------------------------------

class TestPlantedErrorEU018:
    """ERR-EU-018 must be planted in wp_revenue_fy2024.xlsx."""

    def test_err_eu_018_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err_ids = [e.error_id for e in errors.entries.values()]
        assert "ERR-EU-018" in err_ids, f"ERR-EU-018 not in error registry: {err_ids}"

    def test_err_eu_018_targets_revenue_workpaper(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = next(e for e in errors.entries.values() if e.error_id == "ERR-EU-018")
        assert "wp_revenue_fy2024.xlsx" in err.file
        assert err.type == "stale_data"
        assert err.severity == "material"

    def test_err_eu_018_is_stale_fy2023_value(self) -> None:
        """The planted error should show FY2023 CP product revenue value."""
        _, _, _, errors, _ = _ensure_emitted()
        err = next(e for e in errors.entries.values() if e.error_id == "ERR-EU-018")
        assert "38,200,000" in err.description, (
            f"ERR-EU-018 should mention stale €38,200,000: {err.description}"
        )
        assert "40,100,000" in err.description, (
            f"ERR-EU-018 should mention correct €40,100,000: {err.description}"
        )

    def test_err_eu_018_visible_in_xlsx(self) -> None:
        """The stale value should actually appear in the generated xlsx."""
        _, out, _, _, _ = _ensure_emitted()
        wb = openpyxl.load_workbook(out / _PY_DIR / "wp_revenue_fy2024.xlsx", data_only=True)
        ws = wb.active
        values = []
        for r in range(2, ws.max_row + 1):
            val = ws.cell(row=r, column=3).value
            if val is not None:
                values.append(val)
        assert 38_200_000 in values, (
            f"Stale value €38,200,000 should appear in revenue WP. Values: {values}"
        )


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------

class TestCanaryEmbedding:
    """All 15 TC-18-EU files must have canaries registered."""

    def test_all_15_canary_keys_registered(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            assert key in canaries.entries, f"Missing canary key: {key}"

    def test_canary_values_are_8_chars(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            canary = canaries.entries[key].canary
            assert len(canary) == 8, f"Canary for {key} is {len(canary)} chars: {canary}"

    def test_canaries_are_unique(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        values = [canaries.entries[k].canary for k in _CANARY_KEYS]
        assert len(values) == len(set(values)), "Duplicate canary values found"


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------

class TestManifestRegistration:
    """All TC-18-EU files must appear in the manifest."""

    def test_manifest_has_tc18_eu_entries(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        manifest_data = json.loads((out / "manifest.json").read_text())
        tc18eu_entries = [e for e in manifest_data if "TC-18-EU" in e.get("test_cases", [])]
        assert len(tc18eu_entries) >= 15, (
            f"Expected at least 15 TC-18-EU manifest entries, got {len(tc18eu_entries)}"
        )


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------

class TestPromptAndExpectedBehavior:
    """TC-18-EU must have prompt.md and expected_behavior.md."""

    def test_prompt_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-18-EU/prompt.md"
        assert path.exists()
        text = path.read_text()
        assert "IFRS" in text
        assert "ISA" in text
        assert "Pillar Two" in text

    def test_expected_behavior_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "test_cases/TC-18-EU/expected_behavior.md"
        assert path.exists()
        text = path.read_text()
        assert len(text) > 200

    def test_expected_behavior_mentions_ifrs_preservation(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-18-EU/expected_behavior.md").read_text()
        assert "IFRS" in text, "Expected behavior should mention IFRS terminology preservation"
        assert "ISA" in text, "Expected behavior should mention ISA references"

    def test_expected_behavior_mentions_judgment_traps(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        text = (out / "test_cases/TC-18-EU/expected_behavior.md").read_text()
        assert "Pillar Two" in text, "Expected behavior should mention Pillar Two trap"
        assert "revaluation" in text.lower(), "Expected behavior should mention revaluation model"
        assert "IAS 7" in text, "Expected behavior should mention IAS 7 interest classification"


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------

class TestGoldStandard:
    """TC-18-EU gold standard must exist with correct structure."""

    def test_gold_json_exists(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        path = out / "gold_standards" / "TC-18-EU_gold.json"
        assert path.exists()

    def test_gold_has_expected_outputs(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards" / "TC-18-EU_gold.json").read_text())
        assert "expected_outputs" in gold
        eo = gold["expected_outputs"]
        assert "files_to_update" in eo
        assert "fy2024_financials" in eo
        assert "fy2025_financials" in eo
        assert "format_changes" in eo
        assert "new_audit_area" in eo
        assert "pillar_two_assessment" in eo

    def test_gold_has_error_detection(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards" / "TC-18-EU_gold.json").read_text())
        assert "ERR-EU-018" in gold["error_detection"]

    def test_gold_has_canary_verification(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards" / "TC-18-EU_gold.json").read_text())
        assert "canary_verification" in gold
        assert len(gold["canary_verification"]) == 15

    def test_gold_has_judgment_traps(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards" / "TC-18-EU_gold.json").read_text())
        assert "judgment_traps" in gold
        traps = gold["judgment_traps"]
        assert len(traps) == 7, f"Expected 7 judgment traps, got {len(traps)}"
        trap_ids = [t["trap_id"] for t in traps]
        assert "JT-EU-18-01" in trap_ids, "Should have US GAAP terminology trap"
        assert "JT-EU-18-04" in trap_ids, "Should have Pillar Two over-application trap"

    def test_gold_pillar_two_not_applicable(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards" / "TC-18-EU_gold.json").read_text())
        p2 = gold["expected_outputs"]["pillar_two_assessment"]
        assert p2["applicable"] is False
        assert p2["threshold"] == 750_000_000

    def test_gold_scoring_hints(self) -> None:
        _, out, _, _, _ = _ensure_emitted()
        gold = json.loads((out / "gold_standards" / "TC-18-EU_gold.json").read_text())
        hints = gold["scoring_hints"]
        assert "correctness" in hints
        assert "completeness" in hints
        assert "format_compliance" in hints
        assert "robustness" in hints
        assert "communication" in hints
        # IFRS-specific language in hints
        assert "IFRS" in hints["format_compliance"]
