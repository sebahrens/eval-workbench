"""Tests for TC-06-EU — IAS 12 Income Tax Provision (European Group) formatter.

Verifies:
- tc06eu_consolidated_tb_fy2025.xlsx (consolidated TB by entity in EUR)
- tc06eu_tax_provision_fy2024_workpaper.xlsx (prior year: current by jurisdiction,
  deferred rollforward, ETR reconciliation, temporary differences)
- tc06eu_permanent_temporary_differences_fy2025.docx (FY2025 book-tax diffs by entity)
- tc06eu_statutory_rates.docx (NL/DE/FR/UK rates, R&D incentives, Pillar Two)
- ERR-EU-002 planted error (stale Munich Gewerbesteuer Hebesatz 480% vs 490%)
- ERR-EU-003 planted error (formula error double-counting in TB total debit)
- Multi-jurisdiction current tax at correct local rates
- Canary embedding in all files
- Gold standard structure and scoring hints
- Prompt and expected behavior markdown files
"""

from __future__ import annotations

import json
import tempfile
from decimal import Decimal
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc06_eu import emit_tc06_eu
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc06_eu once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None

# All canary keys used by TC-06-EU
_CANARY_KEYS = sorted([
    "tc06eu_consolidated_tb_fy2025",
    "tc06eu_tax_provision_fy2024_workpaper",
    "tc06eu_perm_temp_differences_fy2025",
    "tc06eu_statutory_rates",
])


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc06eu_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc06_eu(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard
        emit_gold("TC-06-EU", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-06-EU/input_files"


# ---------------------------------------------------------------------------
# Consolidated Trial Balance
# ---------------------------------------------------------------------------


class TestConsolidatedTB:
    """Verify tc06eu_consolidated_tb_fy2025.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_consolidated_tb_fy2025.xlsx"
        assert path.exists()

    def test_has_trial_balance_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_consolidated_tb_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Trial Balance" in wb.sheetnames

    def test_has_header_row(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_consolidated_tb_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Trial Balance"]
        headers = []
        for col in range(1, 7):
            val = ws.cell(row=5, column=col).value
            if val:
                headers.append(val)
        assert "Account" in headers
        assert "Entity" in headers
        assert "Debit (EUR)" in headers
        assert "Credit (EUR)" in headers

    def test_has_entity_rows(self) -> None:
        """TB should have rows for CE, CP, CM, CD."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_consolidated_tb_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Trial Balance"]
        entities = set()
        for row in range(6, ws.max_row + 1):
            val = ws.cell(row=row, column=3).value
            if isinstance(val, str) and val:
                entities.add(val)
        for entity in ("CE", "CP", "CM", "CD"):
            assert entity in entities, f"Missing entity {entity} in TB"

    def test_has_totals_row(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_consolidated_tb_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Trial Balance"]
        found_total = False
        for row in range(1, ws.max_row + 1):
            if ws.cell(row=row, column=2).value == "TOTAL":
                found_total = True
                break
        assert found_total, "Missing TOTAL row in trial balance"

    def test_has_elimination_row(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_consolidated_tb_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Trial Balance"]
        found_elim = False
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=3).value
            if isinstance(val, str) and "ELIM" in val:
                found_elim = True
                break
        assert found_elim, "Missing intercompany elimination row"


# ---------------------------------------------------------------------------
# Prior Year Provision Workpaper
# ---------------------------------------------------------------------------


class TestPriorYearWorkpaper:
    """Verify tc06eu_tax_provision_fy2024_workpaper.xlsx structure."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_tax_provision_fy2024_workpaper.xlsx"
        assert path.exists()

    def test_has_current_tax_by_jurisdiction_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Current Tax by Jurisdiction" in wb.sheetnames

    def test_has_deferred_rollforward_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Deferred Tax Rollforward" in wb.sheetnames

    def test_has_rate_reconciliation_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Rate Reconciliation" in wb.sheetnames

    def test_has_temporary_differences_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Temporary Differences" in wb.sheetnames

    def test_current_tax_has_all_jurisdictions(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Current Tax by Jurisdiction"]
        jurisdictions = set()
        for row in range(1, ws.max_row + 1):
            val = ws.cell(row=row, column=1).value
            if isinstance(val, str) and val in ("NL", "DE", "FR", "UK"):
                jurisdictions.add(val)
        assert jurisdictions == {"NL", "DE", "FR", "UK"}

    def test_err_eu_002_stale_rate_visible(self) -> None:
        """The DE statutory rate should show ~29.58% (stale) instead of 29.9%."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Current Tax by Jurisdiction"]
        found_stale = False
        for row in range(1, ws.max_row + 1):
            jur = ws.cell(row=row, column=1).value
            rate = ws.cell(row=row, column=5).value
            if jur == "DE" and rate is not None:
                # Should be ~0.2958 (stale) not 0.299
                if isinstance(rate, (int, float)) and abs(rate - 0.2958) < 0.005:
                    found_stale = True
                    break
        assert found_stale, "ERR-EU-002: stale DE rate not found in workpaper"


# ---------------------------------------------------------------------------
# Permanent & Temporary Differences Document
# ---------------------------------------------------------------------------


class TestPermTempDifferences:
    """Verify tc06eu_permanent_temporary_differences_fy2025.docx content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_permanent_temporary_differences_fy2025.docx"
        assert path.exists()

    def test_has_permanent_differences_section(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "permanent differences" in text.lower()

    def test_has_temporary_differences_section(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "temporary differences" in text.lower()

    def test_has_entity_column_in_perm_table(self) -> None:
        """Permanent differences table should have Entity column."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        assert len(doc.tables) >= 1
        hdr = [cell.text for cell in doc.tables[0].rows[0].cells]
        assert "Entity" in hdr

    def test_mentions_bewirtungskosten(self) -> None:
        """Should mention German entertainment expense rule."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                all_text += "\n" + "\n".join(cell.text for cell in row.cells)
        assert "bewirtungskosten" in all_text.lower()

    def test_mentions_dgccrf(self) -> None:
        """Should mention French regulatory fine."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                all_text += "\n" + "\n".join(cell.text for cell in row.cells)
        assert "dgccrf" in all_text.lower()

    def test_mentions_ias_21_translation(self) -> None:
        """Should note GBP→EUR translation per IAS 21."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "ias 21" in text.lower() or "1.17" in text


# ---------------------------------------------------------------------------
# Statutory Rates Document
# ---------------------------------------------------------------------------


class TestStatutoryRates:
    """Verify tc06eu_statutory_rates.docx content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_statutory_rates.docx"
        assert path.exists()

    def test_mentions_nl_25_8_pct(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\n".join(cell.text for cell in row.cells)
        assert "25.8%" in text

    def test_mentions_de_29_9_pct(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\n".join(cell.text for cell in row.cells)
        assert "29.9%" in text

    def test_mentions_hebesatz_490(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "490%" in text or "490" in text

    def test_mentions_cir(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "cir" in text.lower() or "crédit" in text.lower()

    def test_mentions_forschungszulage(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "forschungszulage" in text.lower()

    def test_mentions_pillar_two_not_applicable(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "750" in text  # €750M threshold
        assert "do not apply" in text.lower() or "not apply" in text.lower()

    def test_mentions_gbp_eur_rate(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/tc06eu_statutory_rates.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "1.17" in text


# ---------------------------------------------------------------------------
# ERR-EU-002 — Stale Munich Gewerbesteuer Hebesatz
# ---------------------------------------------------------------------------


class TestERREU002PlantedError:
    """Verify ERR-EU-002: stale Hebesatz in prior year workpaper."""

    def test_err_eu_002_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-EU-002" in errors.entries

    def test_err_eu_002_is_stale_data(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-002"]
        assert err.type == "stale_data"

    def test_err_eu_002_references_hebesatz(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-002"]
        assert "480" in err.description
        assert "490" in err.description

    def test_err_eu_002_in_workpaper_file(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-002"]
        assert "tc06eu_tax_provision_fy2024_workpaper.xlsx" in err.file

    def test_err_eu_002_severity_significant(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-002"]
        assert err.severity == "significant"


# ---------------------------------------------------------------------------
# ERR-EU-003 — Formula error double-counting in TB
# ---------------------------------------------------------------------------


class TestERREU003PlantedError:
    """Verify ERR-EU-003: formula error in consolidated TB total debit."""

    def test_err_eu_003_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-EU-003" in errors.entries

    def test_err_eu_003_is_formula_error(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-003"]
        assert err.type == "formula_error"

    def test_err_eu_003_references_double_counting(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-003"]
        assert "double-counting" in err.description.lower() or "double counts" in err.description.lower()

    def test_err_eu_003_in_tb_file(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-003"]
        assert "tc06eu_consolidated_tb_fy2025.xlsx" in err.file

    def test_err_eu_003_severity_material(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-EU-003"]
        assert err.severity == "material"


# ---------------------------------------------------------------------------
# Tax model validation
# ---------------------------------------------------------------------------


class TestTaxModel:
    """Verify the European tax provision model produces valid results."""

    def test_current_plus_deferred_equals_total(self) -> None:
        """Tie-out: current + deferred = total provision."""
        from generator.model.tax_eu import compute_eu_tax_provision

        prov25 = compute_eu_tax_provision(2025)

        def _rd(d: Decimal) -> int:
            from decimal import ROUND_HALF_UP
            return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))

        total_check = _rd(prov25.total_current_tax) + _rd(prov25.total_deferred_movement)
        assert total_check == _rd(prov25.total_provision), (
            f"Tie-out failed: {_rd(prov25.total_current_tax)} + "
            f"{_rd(prov25.total_deferred_movement)} != {_rd(prov25.total_provision)}"
        )

    def test_effective_rate_in_reasonable_range(self) -> None:
        """ETR should be between 15% and 35%."""
        from generator.model.tax_eu import compute_eu_tax_provision

        prov25 = compute_eu_tax_provision(2025)
        etr = float(prov25.effective_tax_rate)
        assert 0.15 < etr < 0.35, f"ETR {etr:.4f} outside plausible range"

    def test_all_four_entities_have_current_tax(self) -> None:
        from generator.model.tax_eu import compute_eu_tax_provision

        prov25 = compute_eu_tax_provision(2025)
        entities = {ct.entity_code for ct in prov25.current_tax_by_entity}
        assert entities == {"CE", "CP", "CM", "CD"}

    def test_pillar_two_not_applicable(self) -> None:
        from generator.model.tax_eu import compute_eu_tax_provision

        prov25 = compute_eu_tax_provision(2025)
        assert prov25.pillar_two_applicable is False

    def test_cir_credit_reduces_french_tax(self) -> None:
        from generator.model.tax_eu import compute_eu_tax_provision

        prov25 = compute_eu_tax_provision(2025)
        cm_tax = [ct for ct in prov25.current_tax_by_entity if ct.entity_code == "CM"][0]
        assert cm_tax.tax_credits_eur > 0, "CIR credit should be positive"
        assert cm_tax.current_tax_eur < cm_tax.gross_tax_eur, "CIR should reduce net tax"

    def test_weighted_statutory_rate_computed(self) -> None:
        from generator.model.tax_eu import compute_eu_tax_provision

        prov25 = compute_eu_tax_provision(2025)
        # Weighted rate should be between min (25%) and max (29.9%) rates
        wsr = float(prov25.weighted_statutory_rate)
        assert 0.25 < wsr < 0.30, f"Weighted statutory rate {wsr:.4f} outside expected range"


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """Verify canary codes are embedded in all files."""

    def test_all_canary_keys_assigned(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            code = canaries.canary_for(key)
            assert len(code) == 8, f"Canary for {key} should be 8 chars"

    def test_consolidated_tb_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc06eu_consolidated_tb_fy2025")
        path = output / f"{_INPUT_DIR}/tc06eu_consolidated_tb_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in consolidated TB properties"

    def test_workpaper_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc06eu_tax_provision_fy2024_workpaper")
        path = output / f"{_INPUT_DIR}/tc06eu_tax_provision_fy2024_workpaper.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in workpaper properties"

    def test_perm_temp_canary_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc06eu_perm_temp_differences_fy2025")
        path = output / f"{_INPUT_DIR}/tc06eu_permanent_temporary_differences_fy2025.docx"
        doc = Document(str(path))
        comments = doc.core_properties.comments or ""
        assert canary in comments, f"Canary {canary} not in perm/temp differences properties"

    def test_statutory_rates_canary_in_docx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for("tc06eu_statutory_rates")
        path = output / f"{_INPUT_DIR}/tc06eu_statutory_rates.docx"
        doc = Document(str(path))
        comments = doc.core_properties.comments or ""
        assert canary in comments, f"Canary {canary} not in statutory rates properties"


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold standard JSON structure."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-06-EU_gold.json"
        assert path.exists()

    def test_gold_has_expected_outputs(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06-EU_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "entity_pre_tax_income_eur" in eo
        assert "consolidated_pre_tax_income_eur" in eo
        assert "current_tax_by_entity" in eo
        assert "total_current_tax_eur" in eo
        assert "total_provision_eur" in eo
        assert "effective_tax_rate" in eo
        assert "weighted_statutory_rate" in eo

    def test_gold_has_all_entity_taxes(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06-EU_gold.json").read_text())
        ct = gold["expected_outputs"]["current_tax_by_entity"]
        assert set(ct.keys()) == {"CE", "CP", "CM", "CD"}

    def test_gold_has_permanent_differences(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06-EU_gold.json").read_text())
        perm = gold["expected_outputs"]["permanent_differences"]
        assert len(perm) >= 4, f"Expected >=4 permanent differences, got {len(perm)}"

    def test_gold_has_temporary_differences(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06-EU_gold.json").read_text())
        temp = gold["expected_outputs"]["temporary_differences"]
        assert len(temp) >= 8, f"Expected >=8 temporary differences, got {len(temp)}"

    def test_gold_has_deferred_rollforward(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06-EU_gold.json").read_text())
        dr = gold["expected_outputs"]["deferred_rollforward"]
        assert len(dr) >= 8, f"Expected >=8 deferred items, got {len(dr)}"

    def test_gold_has_rate_reconciliation(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06-EU_gold.json").read_text())
        rr = gold["expected_outputs"]["rate_reconciliation"]
        assert len(rr) >= 3, f"Expected >=3 rate recon items, got {len(rr)}"

    def test_gold_has_pillar_two_answer(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06-EU_gold.json").read_text())
        assert gold["expected_outputs"]["pillar_two_applicable"] is False

    def test_gold_has_tie_out_checks(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06-EU_gold.json").read_text())
        tie = gold["expected_outputs"]["tie_out_checks"]
        assert tie["current_plus_deferred_equals_total"] is True

    def test_gold_has_judgment_traps(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06-EU_gold.json").read_text())
        assert "judgment_traps" in gold
        traps = gold["judgment_traps"]
        assert len(traps) >= 5

    def test_gold_has_error_detection(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-06-EU_gold.json").read_text())
        ed = gold["error_detection"]
        assert "ERR-EU-002" in ed
        assert "ERR-EU-003" in ed


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt and expected behavior files exist and have key content."""

    def test_prompt_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-06-EU/prompt.md"
        assert path.exists()

    def test_prompt_mentions_ias_12(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-06-EU/prompt.md"
        text = path.read_text()
        assert "IAS 12" in text

    def test_prompt_mentions_pillar_two(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-06-EU/prompt.md"
        text = path.read_text()
        assert "Pillar Two" in text or "GloBE" in text

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-06-EU/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_valuation_allowance_trap(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-06-EU/expected_behavior.md"
        text = path.read_text()
        assert "valuation allowance" in text.lower()
