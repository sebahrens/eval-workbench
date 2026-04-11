"""Tests for TC-08 — R&D Tax Credit Study, Section 41 (Tax, Complex) formatter.

Verifies:
- rd_employee_time_records.csv (~2,340 rows for 45 R&D-eligible employees)
- 12 project description docx files in rd_project_descriptions/
- payroll_data_fy2025.xlsx with payroll register for Advanced Materials
- rd_supply_expenses.xlsx with R&D supply costs by project
- ERR-005 planted error (mismatched payroll total — omits one employee)
- ERR-019 planted error (classification error — expense under wrong project)
- Credit ~$185,000 via Alternative Simplified Credit method
- Canary embedding in all files
- Gold standard structure, project details, QRE breakdown, scoring hints
- Prompt and expected behavior markdown files
"""

from __future__ import annotations

import json
import tempfile
from decimal import Decimal
from pathlib import Path

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry
from generator.canaries import build_registry as build_canary_registry
from generator.errors import ErrorRegistry
from generator.formatters.tc08 import (
    _KEY_PAYROLL,
    _KEY_SUPPLY,
    _KEY_TIME_RECORDS,
    _PROJECT_DOC_KEYS,
    emit_tc08,
)
from generator.golds.framework import emit_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel, build_model
from generator.model.rd import (
    PRIOR_YEAR_QRES,
    RD_EMPLOYEE_COUNT,
    RD_PROJECTS,
    TARGET_CREDIT,
    generate_rd_projects,
)

# ---------------------------------------------------------------------------
# Module-level fixture: build the model and run emit_tc08 once
# ---------------------------------------------------------------------------

_MODEL: CascadeModel | None = None
_OUTPUT: Path | None = None
_CANARIES: CanaryRegistry | None = None
_ERRORS: ErrorRegistry | None = None
_MANIFEST: Manifest | None = None

# All canary keys used by TC-08
_CANARY_KEYS = sorted(
    [_KEY_TIME_RECORDS, _KEY_PAYROLL, _KEY_SUPPLY] + _PROJECT_DOC_KEYS
)


def _ensure_emitted() -> tuple[CascadeModel, Path, CanaryRegistry, ErrorRegistry, Manifest]:
    global _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST  # noqa: PLW0603
    if _MODEL is None:
        _MODEL = build_model(seed=42)
        _OUTPUT = Path(tempfile.mkdtemp(prefix="tc08_test_"))
        _CANARIES = build_canary_registry(_CANARY_KEYS, seed=42)
        _ERRORS = ErrorRegistry()
        _MANIFEST = Manifest(_OUTPUT)
        _MANIFEST.__enter__()

        emit_tc08(_MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST)

        # Emit gold standard
        emit_gold("TC-08", _CANARIES, _ERRORS, _OUTPUT / "gold_standards", model=_MODEL)

        _MANIFEST.__exit__(None, None, None)
    return _MODEL, _OUTPUT, _CANARIES, _ERRORS, _MANIFEST


_INPUT_DIR = "test_cases/TC-08/input_files"
_PROJ_DESC_DIR = f"{_INPUT_DIR}/rd_project_descriptions"


# ---------------------------------------------------------------------------
# Time Records CSV — ~2,340 rows for 45 employees
# ---------------------------------------------------------------------------


class TestTimeRecordsCSV:
    """Verify rd_employee_time_records.csv structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/rd_employee_time_records.csv"
        assert path.exists()

    def test_row_count_approximate(self) -> None:
        """Should have ~2,340 data rows (45 emp × 52 weeks)."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/rd_employee_time_records.csv"
        lines = path.read_text().splitlines()
        # First line is canary comment, second is header, rest are data
        data_lines = [ln for ln in lines if ln and not ln.startswith("#")]
        data_row_count = len(data_lines) - 1  # subtract header
        assert 2200 <= data_row_count <= 2400, (
            f"Expected ~2,340 data rows, got {data_row_count}"
        )

    def test_has_expected_columns(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/rd_employee_time_records.csv"
        lines = path.read_text().splitlines()
        # Skip canary comment line(s)
        header_line = next(ln for ln in lines if not ln.startswith("#"))
        headers = header_line.split(",")
        assert "employee_id" in headers
        assert "week_ending" in headers
        assert "project_code" in headers
        assert "hours" in headers
        assert "activity_description" in headers

    def test_45_unique_employees(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/rd_employee_time_records.csv"
        lines = path.read_text().splitlines()
        non_comment = [ln for ln in lines if ln and not ln.startswith("#")]
        # Parse employee_id from first column
        emp_ids = set()
        for line in non_comment[1:]:  # skip header
            emp_id = line.split(",")[0]
            emp_ids.add(emp_id)
        assert len(emp_ids) == RD_EMPLOYEE_COUNT, (
            f"Expected {RD_EMPLOYEE_COUNT} unique employees, got {len(emp_ids)}"
        )

    def test_has_rd_and_overhead_codes(self) -> None:
        """Records should include both R&D project and GEN-xxx overhead codes."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/rd_employee_time_records.csv"
        text = path.read_text()
        assert "RD-" in text, "No R&D project codes found"
        assert "GEN-" in text, "No overhead codes found"


# ---------------------------------------------------------------------------
# Project Descriptions — 12 docx files
# ---------------------------------------------------------------------------


class TestProjectDescriptions:
    """Verify 12 R&D project description docx files."""

    def test_directory_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        assert (output / _PROJ_DESC_DIR).is_dir()

    def test_12_docx_files_exist(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        projects = generate_rd_projects()
        for proj in projects:
            path = output / _PROJ_DESC_DIR / f"{proj.code}.docx"
            assert path.exists(), f"Missing {proj.code}.docx"

    def test_no_extra_files(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        docx_files = list((output / _PROJ_DESC_DIR).glob("*.docx"))
        assert len(docx_files) == 12

    def test_docx_contains_project_name(self) -> None:
        """Each docx should contain its project code in the content."""
        _, output, _, _, _ = _ensure_emitted()
        projects = generate_rd_projects()
        for proj in projects:
            path = output / _PROJ_DESC_DIR / f"{proj.code}.docx"
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs)
            assert proj.code in text, f"{proj.code} not found in its docx"

    def test_docx_has_required_sections(self) -> None:
        """Each docx should have Objective, Technical Uncertainty, Methodology."""
        _, output, _, _, _ = _ensure_emitted()
        projects = generate_rd_projects()
        # Check first project as representative
        path = output / _PROJ_DESC_DIR / f"{projects[0].code}.docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs).lower()
        assert "objective" in text
        assert "technical uncertainty" in text or "uncertainty" in text
        assert "methodology" in text

    def test_qualification_mix(self) -> None:
        """Projects should have 8 qualifying, 2 borderline, 2 non-qualifying."""
        projects = generate_rd_projects()
        quals = [p.qualifies for p in projects]
        assert quals.count("yes") == 8
        assert quals.count("borderline") == 2
        assert quals.count("no") == 2


# ---------------------------------------------------------------------------
# Payroll Data XLSX
# ---------------------------------------------------------------------------


class TestPayrollData:
    """Verify payroll_data_fy2025.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/payroll_data_fy2025.xlsx"
        assert path.exists()

    def test_has_payroll_register_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/payroll_data_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "Payroll Register FY2025" in wb.sheetnames

    def test_has_header_row(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/payroll_data_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Payroll Register FY2025"]
        headers = [ws.cell(row=4, column=c).value for c in range(1, 9)]
        assert "Employee ID" in headers
        assert "W-2 Wages (Annual Salary)" in headers
        assert "Total Compensation" in headers

    def test_has_am_employees(self) -> None:
        """Should contain AM entity employees with salary data."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/payroll_data_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Payroll Register FY2025"]
        row_count = 0
        for row in range(5, ws.max_row + 1):
            emp_id = ws.cell(row=row, column=1).value
            if emp_id and emp_id != "":
                wages = ws.cell(row=row, column=5).value
                if isinstance(wages, (int, float)) and wages > 0:
                    row_count += 1
        assert row_count >= RD_EMPLOYEE_COUNT, (
            f"Expected ≥{RD_EMPLOYEE_COUNT} employee rows, got {row_count}"
        )

    def test_has_total_row(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/payroll_data_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Payroll Register FY2025"]
        found_total = False
        for row in range(5, ws.max_row + 1):
            val = ws.cell(row=row, column=2).value
            if val == "TOTAL":
                found_total = True
                break
        assert found_total, "Missing TOTAL row in payroll register"


# ---------------------------------------------------------------------------
# Supply Expenses XLSX
# ---------------------------------------------------------------------------


class TestSupplyExpenses:
    """Verify rd_supply_expenses.xlsx structure and content."""

    def test_file_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/rd_supply_expenses.xlsx"
        assert path.exists()

    def test_has_supply_expenses_sheet(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/rd_supply_expenses.xlsx"
        wb = openpyxl.load_workbook(str(path))
        assert "R&D Supply Expenses" in wb.sheetnames

    def test_has_header_columns(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/rd_supply_expenses.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["R&D Supply Expenses"]
        headers = [ws.cell(row=4, column=c).value for c in range(1, 7)]
        assert "Date" in headers
        assert "Amount" in headers
        assert "Project Code" in headers

    def test_has_expense_rows(self) -> None:
        """Should have 96-144 expense rows."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/rd_supply_expenses.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["R&D Supply Expenses"]
        row_count = 0
        for row in range(5, ws.max_row + 1):
            if ws.cell(row=row, column=1).value is not None:
                row_count += 1
        assert 96 <= row_count <= 144, f"Expected 96-144 expense rows, got {row_count}"


# ---------------------------------------------------------------------------
# ERR-005 — Mismatched payroll total
# ---------------------------------------------------------------------------


class TestERR005PlantedError:
    """Verify ERR-005: payroll total wages omits one employee."""

    def test_err005_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-005" in errors.entries

    def test_err005_is_mismatched_total(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-005"]
        assert err.type == "mismatched_total"

    def test_err005_in_payroll_file(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-005"]
        assert "payroll_data_fy2025.xlsx" in err.file

    def test_err005_severity_material(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-005"]
        assert err.severity == "material"

    def test_err005_references_tc08(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-005"]
        assert "TC-08" in err.which_test_cases_should_catch

    def test_err005_total_mismatch_in_file(self) -> None:
        """The TOTAL row in the payroll file should NOT match sum of employee wages."""
        _, output, _, _, _ = _ensure_emitted()
        path = output / f"{_INPUT_DIR}/payroll_data_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        ws = wb["Payroll Register FY2025"]

        total_from_rows = 0
        total_row_value = None

        for row in range(5, ws.max_row + 1):
            col2 = ws.cell(row=row, column=2).value
            wages = ws.cell(row=row, column=5).value
            if col2 == "TOTAL":
                total_row_value = wages
            elif isinstance(wages, (int, float)) and wages > 0:
                total_from_rows += wages

        assert total_row_value is not None, "No TOTAL row found"
        assert total_row_value != total_from_rows, (
            "TOTAL row matches sum of employee wages — ERR-005 not planted"
        )


# ---------------------------------------------------------------------------
# ERR-019 — Classification error in supply expenses
# ---------------------------------------------------------------------------


class TestERR019PlantedError:
    """Verify ERR-019: supply expense under wrong project code."""

    def test_err019_registered(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        assert "ERR-019" in errors.entries

    def test_err019_is_classification_error(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-019"]
        assert err.type == "classification_error"

    def test_err019_in_supply_file(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-019"]
        assert "rd_supply_expenses.xlsx" in err.file

    def test_err019_severity_immaterial(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-019"]
        assert err.severity == "immaterial"

    def test_err019_references_tc08(self) -> None:
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-019"]
        assert "TC-08" in err.which_test_cases_should_catch

    def test_err019_mentions_rd011(self) -> None:
        """ERR-019 should reference RD-011 (non-qualifying market research)."""
        _, _, _, errors, _ = _ensure_emitted()
        err = errors.entries["ERR-019"]
        assert "RD-011" in err.description


# ---------------------------------------------------------------------------
# Credit computation
# ---------------------------------------------------------------------------


class TestCreditComputation:
    """Verify R&D credit amount from gold standard."""

    def test_credit_within_tolerance(self) -> None:
        """ASC credit must be within $500 of $185,000."""
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-08_gold.json").read_text())
        credit = Decimal(gold["expected_outputs"]["asc_credit"])
        assert abs(credit - TARGET_CREDIT) <= 500, (
            f"Credit {credit} not within $500 of {TARGET_CREDIT}"
        )

    def test_prior_year_qres_correct(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-08_gold.json").read_text())
        prior = gold["expected_outputs"]["prior_year_qres"]
        assert Decimal(prior["fy2023"]) == PRIOR_YEAR_QRES[2023]
        assert Decimal(prior["fy2024"]) == PRIOR_YEAR_QRES[2024]


# ---------------------------------------------------------------------------
# Canary embedding
# ---------------------------------------------------------------------------


class TestCanaryEmbedding:
    """Verify canary codes are embedded in TC-08 files."""

    def test_all_canary_keys_assigned(self) -> None:
        _, _, canaries, _, _ = _ensure_emitted()
        for key in _CANARY_KEYS:
            code = canaries.canary_for(key)
            assert len(code) == 8, f"Canary for {key} should be 8 chars"

    def test_time_records_canary_in_csv(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for(_KEY_TIME_RECORDS)
        path = output / f"{_INPUT_DIR}/rd_employee_time_records.csv"
        text = path.read_text()
        assert canary in text, f"Canary {canary} not in time records CSV"

    def test_payroll_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for(_KEY_PAYROLL)
        path = output / f"{_INPUT_DIR}/payroll_data_fy2025.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in payroll xlsx properties"

    def test_supply_canary_in_xlsx(self) -> None:
        _, output, canaries, _, _ = _ensure_emitted()
        canary = canaries.canary_for(_KEY_SUPPLY)
        path = output / f"{_INPUT_DIR}/rd_supply_expenses.xlsx"
        wb = openpyxl.load_workbook(str(path))
        desc = wb.properties.description or ""
        assert canary in desc, f"Canary {canary} not in supply xlsx properties"

    def test_project_docx_canaries(self) -> None:
        """Each project description docx should have its canary."""
        _, output, canaries, _, _ = _ensure_emitted()
        projects = generate_rd_projects()
        for i, proj in enumerate(projects):
            key = _PROJECT_DOC_KEYS[i]
            canary = canaries.canary_for(key)
            path = output / _PROJ_DESC_DIR / f"{proj.code}.docx"
            doc = Document(str(path))
            comments = doc.core_properties.comments or ""
            assert canary in comments, (
                f"Canary {canary} not in {proj.code}.docx properties"
            )


# ---------------------------------------------------------------------------
# Gold standard
# ---------------------------------------------------------------------------


class TestGoldStandard:
    """Verify gold standard JSON structure."""

    def test_gold_json_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "gold_standards" / "TC-08_gold.json"
        assert path.exists()

    def test_gold_has_expected_outputs(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-08_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "projects_total" in eo
        assert eo["projects_total"] == 12
        assert eo["projects_qualifying"] == 8
        assert eo["projects_borderline"] == 2
        assert eo["projects_disqualified"] == 2

    def test_gold_has_qre_breakdown(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-08_gold.json").read_text())
        eo = gold["expected_outputs"]
        assert "wage_qres" in eo
        assert "supply_qres" in eo
        assert "total_qres_fy2025" in eo

    def test_gold_wage_plus_supply_equals_total(self) -> None:
        """Arithmetic tie-out in gold: wage + supply = total QREs."""
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-08_gold.json").read_text())
        eo = gold["expected_outputs"]
        wage = Decimal(eo["wage_qres"])
        supply = Decimal(eo["supply_qres"])
        total = Decimal(eo["total_qres_fy2025"])
        assert wage + supply == total

    def test_gold_has_project_details(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-08_gold.json").read_text())
        details = gold["expected_outputs"]["project_details"]
        assert len(details) == 12

    def test_gold_has_per_project_qres(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-08_gold.json").read_text())
        qres = gold["expected_outputs"]["qres_by_project"]
        assert len(qres) > 0
        # Non-qualifying projects should not appear
        for code in qres:
            proj = next(p for p in RD_PROJECTS if p.code == code)
            assert proj.qualifies != "no", (
                f"Non-qualifying project {code} in QRE breakdown"
            )

    def test_gold_canary_verification(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-08_gold.json").read_text())
        cv = gold["canary_verification"]
        assert "read_time_records" in cv
        assert "read_payroll" in cv
        assert "read_supply_expenses" in cv
        # 12 project descriptions
        for proj in RD_PROJECTS:
            assert f"read_project_{proj.code}" in cv

    def test_gold_error_detection(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-08_gold.json").read_text())
        assert "ERR-005" in gold["error_detection"]
        assert "ERR-019" in gold["error_detection"]

    def test_gold_scoring_hints(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        gold = json.loads((output / "gold_standards" / "TC-08_gold.json").read_text())
        hints = gold["scoring_hints"]
        for key in ["correctness", "completeness", "format_compliance", "robustness", "communication"]:
            assert key in hints


# ---------------------------------------------------------------------------
# Prompt and expected behavior
# ---------------------------------------------------------------------------


class TestPromptAndExpectedBehavior:
    """Verify prompt and expected behavior files are generated."""

    def test_prompt_md_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-08/prompt.md"
        assert path.exists()

    def test_prompt_mentions_rd_credit(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-08/prompt.md").read_text()
        assert "R&D" in text
        assert "credit" in text.lower()

    def test_prompt_mentions_four_part_test(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-08/prompt.md").read_text().lower()
        assert "four-part" in text or "4-part" in text

    def test_prompt_mentions_asc_method(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-08/prompt.md").read_text()
        assert "Alternative Simplified Credit" in text or "ASC" in text

    def test_prompt_mentions_prior_year_qres(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-08/prompt.md").read_text()
        assert "$3.1M" in text
        assert "$3.4M" in text

    def test_expected_behavior_exists(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        path = output / "test_cases/TC-08/expected_behavior.md"
        assert path.exists()

    def test_expected_behavior_mentions_qualification(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-08/expected_behavior.md").read_text()
        assert "qualif" in text.lower()
        assert "borderline" in text.lower()

    def test_expected_behavior_mentions_credit_amount(self) -> None:
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-08/expected_behavior.md").read_text()
        assert "$185,000" in text

    def test_expected_behavior_mentions_non_qualifying(self) -> None:
        """Should call out RD-011 and RD-012 as non-qualifying."""
        _, output, _, _, _ = _ensure_emitted()
        text = (output / "test_cases/TC-08/expected_behavior.md").read_text()
        assert "RD-011" in text
        assert "RD-012" in text


# ---------------------------------------------------------------------------
# Manifest registration
# ---------------------------------------------------------------------------


class TestManifest:
    """Verify files are registered in the manifest."""

    def test_manifest_has_tc08_entries(self) -> None:
        _, _, _, _, manifest = _ensure_emitted()
        # Check that TC-08 input files are registered in the manifest
        # (the formatter registers paths but doesn't tag test_cases=["TC-08"])
        tc08_paths = [
            k for k in manifest.entries
            if "TC-08" in k
        ]
        # time_records.csv + payroll.xlsx + supply.xlsx + 12 docx = 15
        assert len(tc08_paths) >= 15, (
            f"Expected ≥15 TC-08 manifest entries, got {len(tc08_paths)}"
        )
