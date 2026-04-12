"""Canary-preservation tests for controlled document noise (synth-data-8ok.9).

Verifies that noise transforms across the formatter-family matrix cannot
erase, corrupt, or make canaries unfindable.  Covers the three implemented
families: xlsx_openpyxl, csv_stdlib, docx_python_docx.

Each test is parametrized so failures identify the specific family, TC,
file key, and transform that broke the canary.
"""

from __future__ import annotations

import random

import pytest

from generator.canaries import (
    build_registry,
    embed_canary_csv_comment,
    embed_canary_docx,
    embed_canary_xlsx,
)
from generator.noise import (
    ExclusionZone,
    apply_csv_noise,
    apply_docx_noise,
    apply_xlsx_noise,
    make_noise_rng,
)
from generator.scenario_context import ScenarioContext

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

# Representative (tc_id, file_key) pairs from the formatter-family matrix.
XLSX_REPRESENTATIVES = [
    ("TC-01", "cascade_tb_fy2025"),
    ("TC-02", "cascade_gl_cash_dec2025"),
    ("TC-08", "payroll_data_fy2025"),
]

CSV_REPRESENTATIVES = [
    ("TC-08", "rd_employee_time_records"),
    ("TC-02", "bank_statement_dec2025"),
]

DOCX_REPRESENTATIVES = [
    ("TC-08", "rd_project_RD-001"),
    ("TC-05", "workpaper_memo_template"),
]

# Seeds to test — using multiple seeds exercises different noise paths.
SEED_RANGE = list(range(50))


@pytest.fixture()
def canary_registry():
    """Build a registry with canaries for all representative file keys."""
    all_keys = sorted(
        {fk for _, fk in XLSX_REPRESENTATIVES + CSV_REPRESENTATIVES + DOCX_REPRESENTATIVES}
    )
    return build_registry(all_keys, seed=42)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_xlsx_workbook(canary: str):
    """Create a minimal openpyxl Workbook with canary and realistic data."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    # Headers in row 1
    ws["A1"] = "Account #"
    ws["B1"] = "Description"
    ws["C1"] = "Amount"
    ws["D1"] = "Date"
    # Data rows (starting at row 6 to exercise header-zone scanning)
    ws["A6"] = "1100"
    ws["B6"] = "Accounts Receivable"
    ws["C6"] = 50000
    ws["A7"] = "2010"
    ws["B7"] = "Accounts Payable"
    ws["C7"] = 30000
    # Embed canary
    embed_canary_xlsx(wb, canary)
    return wb


def _make_csv_lines(canary: str) -> list[str]:
    """Create minimal CSV lines with canary comment."""
    return [
        embed_canary_csv_comment(canary),
        "Account,Description,Amount\n",
        "1100,Accounts Receivable,50000\n",
        "2010,Accounts Payable,30000\n",
        "3000,Revenue,120000\n",
        "4000,Cost of Goods Sold,80000\n",
    ]


def _make_docx_document(canary: str):
    """Create a minimal python-docx Document with canary."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Test Document", level=1)
    p = doc.add_paragraph("Body paragraph with content.")
    for run in p.runs:
        run.font.size = Pt(11)
    doc.add_paragraph("Second paragraph for additional noise surface.")
    for run in doc.paragraphs[-1].runs:
        run.font.size = Pt(11)
    embed_canary_docx(doc, canary)
    return doc


def _canary_findable_xlsx(wb, canary: str) -> bool:
    """Check canary is findable in XLSX document properties."""
    desc = wb.properties.description or ""
    return canary in desc


def _canary_findable_csv(lines: list[str], canary: str) -> bool:
    """Check canary is findable in CSV lines (including BOM tolerance)."""
    text = "".join(lines)
    # Strip BOM for search — canary search must be BOM-tolerant
    text_clean = text.lstrip("\ufeff")
    return canary in text_clean


def _canary_findable_docx(doc, canary: str) -> bool:
    """Check canary is findable in DOCX core properties."""
    comments = doc.core_properties.comments or ""
    return canary in comments


# ---------------------------------------------------------------------------
# XLSX canary-preservation tests
# ---------------------------------------------------------------------------

class TestXlsxCanaryPreservation:
    """Canary survives all xlsx noise transforms across seeds."""

    @pytest.mark.parametrize("tc_id,file_key", XLSX_REPRESENTATIVES, ids=lambda x: x)
    @pytest.mark.parametrize("seed", SEED_RANGE, ids=lambda s: f"seed={s}")
    def test_canary_survives_noise(self, canary_registry, tc_id, file_key, seed):
        canary = canary_registry.canary_for(file_key)
        wb = _make_xlsx_workbook(canary)
        rng = random.Random(seed)
        apply_xlsx_noise(wb, rng)
        assert _canary_findable_xlsx(wb, canary), (
            f"Canary {canary!r} not findable in {tc_id}/{file_key} after xlsx noise (seed={seed})"
        )

    @pytest.mark.parametrize("tc_id,file_key", XLSX_REPRESENTATIVES, ids=lambda x: x)
    def test_canary_survives_with_exclusion_zones(self, canary_registry, tc_id, file_key):
        """Canary preserved even when exclusion zones are in use."""
        canary = canary_registry.canary_for(file_key)
        # Exclude some data cells (simulates planted-error protection)
        excl = ExclusionZone(cells={("Data", 6, 1), ("Data", 7, 3)})
        for seed in range(20):
            wb_copy = _make_xlsx_workbook(canary)
            rng = random.Random(seed)
            apply_xlsx_noise(wb_copy, rng, excl)
            assert _canary_findable_xlsx(wb_copy, canary), (
                f"Canary lost in {tc_id}/{file_key} with exclusion zones (seed={seed})"
            )

    @pytest.mark.parametrize("tc_id,file_key", XLSX_REPRESENTATIVES, ids=lambda x: x)
    def test_canary_exact_value_preserved(self, canary_registry, tc_id, file_key):
        """The description property is byte-identical before and after noise."""
        canary = canary_registry.canary_for(file_key)
        wb = _make_xlsx_workbook(canary)
        expected = wb.properties.description
        rng = random.Random(42)
        apply_xlsx_noise(wb, rng)
        assert wb.properties.description == expected


# ---------------------------------------------------------------------------
# CSV canary-preservation tests
# ---------------------------------------------------------------------------

class TestCsvCanaryPreservation:
    """Canary survives all csv noise transforms across seeds."""

    @pytest.mark.parametrize("tc_id,file_key", CSV_REPRESENTATIVES, ids=lambda x: x)
    @pytest.mark.parametrize("seed", SEED_RANGE, ids=lambda s: f"seed={s}")
    def test_canary_survives_noise(self, canary_registry, tc_id, file_key, seed):
        canary = canary_registry.canary_for(file_key)
        lines = _make_csv_lines(canary)
        rng = random.Random(seed)
        result = apply_csv_noise(list(lines), rng)
        assert _canary_findable_csv(result, canary), (
            f"Canary {canary!r} not findable in {tc_id}/{file_key} after csv noise (seed={seed})"
        )

    @pytest.mark.parametrize("tc_id,file_key", CSV_REPRESENTATIVES, ids=lambda x: x)
    def test_canary_line_byte_identical(self, canary_registry, tc_id, file_key):
        """The canary comment line is never modified."""
        canary = canary_registry.canary_for(file_key)
        lines = _make_csv_lines(canary)
        expected_line = lines[0]
        for seed in range(50):
            result = apply_csv_noise(list(lines), random.Random(seed))
            # Canary line may have BOM prepended — check the core content
            actual = result[0].lstrip("\ufeff")
            assert actual == expected_line, (
                f"Canary line modified in {tc_id}/{file_key} (seed={seed}): "
                f"{result[0]!r} != {expected_line!r}"
            )

    @pytest.mark.parametrize("tc_id,file_key", CSV_REPRESENTATIVES, ids=lambda x: x)
    def test_canary_survives_bom_insertion(self, canary_registry, tc_id, file_key):
        """Canary findable even when BOM is forcefully inserted."""
        canary = canary_registry.canary_for(file_key)
        lines = _make_csv_lines(canary)
        result = apply_csv_noise(list(lines), random.Random(42), add_bom=True)
        assert _canary_findable_csv(result, canary), (
            f"Canary {canary!r} lost after BOM insertion in {tc_id}/{file_key}"
        )

    @pytest.mark.parametrize("tc_id,file_key", CSV_REPRESENTATIVES, ids=lambda x: x)
    def test_canary_survives_with_exclusion_zones(self, canary_registry, tc_id, file_key):
        """Canary preserved when row exclusion zones are active."""
        canary = canary_registry.canary_for(file_key)
        lines = _make_csv_lines(canary)
        excl = ExclusionZone(rows={1, 2})
        for seed in range(20):
            result = apply_csv_noise(list(lines), random.Random(seed), excl)
            assert _canary_findable_csv(result, canary), (
                f"Canary lost in {tc_id}/{file_key} with row exclusions (seed={seed})"
            )


# ---------------------------------------------------------------------------
# DOCX canary-preservation tests
# ---------------------------------------------------------------------------

class TestDocxCanaryPreservation:
    """Canary survives all docx noise transforms across seeds."""

    @pytest.mark.parametrize("tc_id,file_key", DOCX_REPRESENTATIVES, ids=lambda x: x)
    @pytest.mark.parametrize("seed", SEED_RANGE, ids=lambda s: f"seed={s}")
    def test_canary_survives_noise(self, canary_registry, tc_id, file_key, seed):
        canary = canary_registry.canary_for(file_key)
        doc = _make_docx_document(canary)
        rng = random.Random(seed)
        apply_docx_noise(doc, rng)
        assert _canary_findable_docx(doc, canary), (
            f"Canary {canary!r} not findable in {tc_id}/{file_key} after docx noise (seed={seed})"
        )

    @pytest.mark.parametrize("tc_id,file_key", DOCX_REPRESENTATIVES, ids=lambda x: x)
    def test_canary_comments_byte_identical(self, canary_registry, tc_id, file_key):
        """The comments property is never modified by noise."""
        canary = canary_registry.canary_for(file_key)
        doc = _make_docx_document(canary)
        expected = doc.core_properties.comments
        for seed in range(50):
            doc_copy = _make_docx_document(canary)
            apply_docx_noise(doc_copy, random.Random(seed))
            assert doc_copy.core_properties.comments == expected, (
                f"Comments property changed in {tc_id}/{file_key} (seed={seed})"
            )

    @pytest.mark.parametrize("tc_id,file_key", DOCX_REPRESENTATIVES, ids=lambda x: x)
    def test_canary_survives_with_exclusion_zones(self, canary_registry, tc_id, file_key):
        """Canary preserved when paragraph exclusion zones are active."""
        canary = canary_registry.canary_for(file_key)
        excl = ExclusionZone(paragraphs={0, 1})
        for seed in range(20):
            doc = _make_docx_document(canary)
            apply_docx_noise(doc, random.Random(seed), excl)
            assert _canary_findable_docx(doc, canary), (
                f"Canary lost in {tc_id}/{file_key} with paragraph exclusions (seed={seed})"
            )


# ---------------------------------------------------------------------------
# Cross-family property test: make_noise_rng isolation
# ---------------------------------------------------------------------------

class TestNoiseRngCanaryIsolation:
    """Noise RNG derived from ScenarioContext doesn't perturb canary generation."""

    def test_canary_registry_stable_across_noise_usage(self):
        """Creating noise RNGs does not affect canary registry output.

        This guards against accidental shared-state corruption where noise
        RNG creation could consume from the canary RNG stream.
        """
        file_keys = sorted(["file_a", "file_b", "file_c"])

        # Build registry without any noise RNG usage
        reg1 = build_registry(file_keys, seed=42)

        # Build registry, then create noise RNGs (should be independent)
        reg2 = build_registry(file_keys, seed=42)
        ctx = ScenarioContext(seed=42)
        for fk in file_keys:
            _ = make_noise_rng(ctx, "TC-01", fk)

        for fk in file_keys:
            assert reg1.canary_for(fk) == reg2.canary_for(fk), (
                f"Canary for {fk} changed after noise RNG creation"
            )

    @pytest.mark.parametrize(
        "family,tc_id,file_key",
        [
            ("xlsx", "TC-01", "cascade_tb_fy2025"),
            ("csv", "TC-08", "rd_employee_time_records"),
            ("docx", "TC-08", "rd_project_RD-001"),
        ],
        ids=["xlsx", "csv", "docx"],
    )
    def test_noise_rng_deterministic_per_family(self, family, tc_id, file_key):
        """Same (tc_id, file_key) always produces the same noise sequence."""
        ctx1 = ScenarioContext(seed=42)
        ctx2 = ScenarioContext(seed=42)
        rng1 = make_noise_rng(ctx1, tc_id, file_key)
        rng2 = make_noise_rng(ctx2, tc_id, file_key)
        seq1 = [rng1.random() for _ in range(100)]
        seq2 = [rng2.random() for _ in range(100)]
        assert seq1 == seq2
