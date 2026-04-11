"""Legal document writer helpers for TC-19 / TC-21.

Deterministic generators for:
- Contract summaries (docx) — one per LegalContract
- Amendment/side letters (docx) — one per ContractAmendment
- Management summary memo (docx)
- Diligence request list (xlsx)

All functions accept model data from generator.model.legal and embed
canaries via the CanaryRegistry.  They never invent facts outside the
canonical model.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import openpyxl
from docx import Document
from docx.shared import Pt

from generator.canaries import CanaryRegistry, embed_canary_docx, embed_canary_xlsx
from generator.model.legal import (
    CONTRACT_AMENDMENTS,
    LEGAL_CONTRACTS,
    LEGAL_DILIGENCE_ISSUES,
    ContractAmendment,
    ContractClause,
    LegalContract,
    amendments_for_contract,
    clauses_for_contract,
)
from generator.writers import (
    new_workbook,
    save_docx_deterministic,
    save_xlsx_deterministic,
    write_header_row,
)

# ── Contract summary documents (docx) ──────────────────────────────────────


def write_contract_summary(
    contract: LegalContract,
    clauses: list[ContractClause],
    amendments: list[ContractAmendment],
    output_path: Path,
    canaries: CanaryRegistry,
    canary_key: str,
) -> str:
    """Write a single contract summary as a docx file.

    Returns the canary embedding location description.
    """
    doc = Document()
    canary = canaries.canary_for(canary_key)
    location = embed_canary_docx(doc, canary)

    # Title
    title = doc.add_heading(f"Contract Summary — {contract.contract_id}", level=1)
    title.runs[0].font.size = Pt(14)

    # Contract details table
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    _add_row(table, "Contract ID", contract.contract_id)
    _add_row(table, "Counterparty", f"{contract.counterparty_name} ({contract.counterparty_id})")
    _add_row(table, "Entity", contract.entity_code)
    _add_row(table, "Type", contract.contract_type.title())
    _add_row(table, "Effective Date", contract.effective_date.isoformat())
    _add_row(table, "Expiration Date", contract.expiration_date.isoformat())
    _add_row(table, "Annual Value", f"${contract.annual_value:,.0f}")
    _add_row(table, "Governing Law", contract.governing_law)
    _add_row(table, "Auto-Renew", "Yes" if contract.auto_renew else "No")

    if contract.notes:
        doc.add_heading("Notes", level=2)
        doc.add_paragraph(contract.notes)

    # Key clauses
    if clauses:
        doc.add_heading("Key Clauses", level=2)
        for cl in clauses:
            p = doc.add_paragraph()
            run = p.add_run(f"{cl.clause_id} — {cl.clause_type.replace('_', ' ').title()}")
            run.bold = True
            run.font.size = Pt(11)
            doc.add_paragraph(f"Source: {cl.source_section}", style="List Bullet")
            doc.add_paragraph(f"Risk Level: {cl.risk_level.title()}", style="List Bullet")
            doc.add_paragraph(f"Business Impact: {cl.business_impact.replace('_', ' ').title()}", style="List Bullet")
            doc.add_paragraph(cl.summary)
            if cl.notes:
                doc.add_paragraph(f"Note: {cl.notes}", style="List Bullet 2")

    # Amendments
    if amendments:
        doc.add_heading("Amendments & Side Letters", level=2)
        for amd in amendments:
            p = doc.add_paragraph()
            run = p.add_run(f"{amd.amendment_id} — Effective {amd.effective_date.isoformat()}")
            run.bold = True
            doc.add_paragraph(amd.description)
            if amd.changes_clause_ids:
                doc.add_paragraph(
                    f"Affects clauses: {', '.join(amd.changes_clause_ids)}",
                    style="List Bullet",
                )
            if amd.notes:
                doc.add_paragraph(f"Note: {amd.notes}", style="List Bullet 2")

    save_docx_deterministic(doc, output_path)
    return location


def write_all_contract_summaries(
    output_dir: Path,
    canaries: CanaryRegistry,
    canary_keys: dict[str, str],
) -> dict[str, str]:
    """Write one docx per LegalContract.

    Parameters
    ----------
    canary_keys:
        Mapping of contract_id → canary file key.

    Returns
    -------
    dict mapping contract_id → canary embedding location.
    """
    locations: dict[str, str] = {}
    for contract in LEGAL_CONTRACTS:
        ckey = canary_keys[contract.contract_id]
        clauses = clauses_for_contract(contract.contract_id)
        amendments = amendments_for_contract(contract.contract_id)
        fname = f"contract_{contract.contract_id.lower()}.docx"
        loc = write_contract_summary(
            contract, clauses, amendments,
            output_dir / fname,
            canaries, ckey,
        )
        locations[contract.contract_id] = loc
    return locations


# ── Amendment / side letter documents (docx) ────────────────────────────────


def write_amendment(
    amendment: ContractAmendment,
    contract: LegalContract,
    output_path: Path,
    canaries: CanaryRegistry,
    canary_key: str,
) -> str:
    """Write a single amendment or side letter as a docx file.

    Returns the canary embedding location description.
    """
    doc = Document()
    canary = canaries.canary_for(canary_key)
    location = embed_canary_docx(doc, canary)

    doc.add_heading(f"Amendment {amendment.amendment_id}", level=1)
    doc.add_paragraph(
        f"To: {contract.contract_id} — {contract.counterparty_name}"
    )
    doc.add_paragraph(f"Effective Date: {amendment.effective_date.isoformat()}")

    doc.add_heading("Description of Changes", level=2)
    doc.add_paragraph(amendment.description)

    if amendment.changes_clause_ids:
        doc.add_heading("Affected Clauses", level=2)
        for cid in amendment.changes_clause_ids:
            doc.add_paragraph(cid, style="List Bullet")

    if amendment.supersedes_original:
        doc.add_paragraph(
            "This amendment supersedes the original clause(s) in their entirety."
        )
    else:
        doc.add_paragraph(
            "This amendment modifies the referenced clause(s) as described above. "
            "All other terms of the original agreement remain in full force and effect."
        )

    if amendment.notes:
        doc.add_heading("Notes", level=2)
        doc.add_paragraph(amendment.notes)

    save_docx_deterministic(doc, output_path)
    return location


def write_all_amendments(
    output_dir: Path,
    canaries: CanaryRegistry,
    canary_keys: dict[str, str],
) -> dict[str, str]:
    """Write one docx per ContractAmendment.

    Parameters
    ----------
    canary_keys:
        Mapping of amendment_id → canary file key.

    Returns
    -------
    dict mapping amendment_id → canary embedding location.
    """
    contract_map = {c.contract_id: c for c in LEGAL_CONTRACTS}
    locations: dict[str, str] = {}
    for amd in CONTRACT_AMENDMENTS:
        ckey = canary_keys[amd.amendment_id]
        contract = contract_map[amd.contract_id]
        fname = f"amendment_{amd.amendment_id.lower()}.docx"
        loc = write_amendment(amd, contract, output_dir / fname, canaries, ckey)
        locations[amd.amendment_id] = loc
    return locations


# ── Management summary memo (docx) ─────────────────────────────────────────


def write_management_summary_memo(
    output_path: Path,
    canaries: CanaryRegistry,
    canary_key: str,
) -> str:
    """Write the management summary memo that deliberately omits AMD-002.

    This memo is a TC-19 trap: it describes the TechAlloy MFN clause using
    the original terms, not the narrowed threshold from the side letter.
    The agent must identify the contradiction.

    Returns the canary embedding location description.
    """
    doc = Document()
    canary = canaries.canary_for(canary_key)
    location = embed_canary_docx(doc, canary)

    doc.add_heading("Management Summary — Contract Portfolio", level=1)
    doc.add_paragraph(
        "Prepared by Cascade Industries management for due diligence purposes. "
        "This memo summarizes key contracts and notable provisions."
    )

    doc.add_heading("Key Customer Contracts", level=2)

    # Acme — correct summary
    doc.add_heading("Acme Manufacturing Corp (LCTR-001)", level=3)
    doc.add_paragraph(
        "Master supply agreement representing approximately 18% of consolidated "
        "revenue. Contract includes a change-of-control provision (Section 14.2) "
        "allowing Acme to terminate without penalty upon change of ownership. "
        "No waiver has been obtained."
    )

    # TechAlloy — STALE summary (does not reflect AMD-002 side letter)
    doc.add_heading("TechAlloy Systems (LCTR-003)", level=3)
    doc.add_paragraph(
        "Supply agreement with most-favored-nation pricing clause (Section 8.1). "
        "Cascade must offer TechAlloy pricing no less favorable than any comparable "
        "customer for equivalent volume and alloy grades. Notification required "
        "within 30 days of any better terms offered to another customer."
    )
    # NOTE: The above deliberately omits AMD-002 which raised the MFN threshold
    # to 5% per-kg differences. This is the contradicts_summary trap (LDI-002).

    # NextGen — partial summary (omits AMD-003 scope expansion)
    doc.add_heading("NextGen Composites LLC (LCTR-004)", level=3)
    doc.add_paragraph(
        "Government subcontract subject to DFARS flow-down provisions. "
        "Exclusivity clause (Section 6.4) restricts Advanced Materials from "
        "supplying competing defense contractors for identical alloy specifications."
    )
    # NOTE: Does not mention AMD-003 expanding exclusivity to thermal barrier
    # coatings. This is the stale_document trap (LDI-003).

    # Other contracts — brief summaries
    doc.add_heading("Other Notable Contracts", level=2)
    for contract in LEGAL_CONTRACTS:
        if contract.contract_id in ("LCTR-001", "LCTR-003", "LCTR-004"):
            continue
        doc.add_paragraph(
            f"{contract.counterparty_name} ({contract.contract_id}): "
            f"{contract.contract_type.title()} agreement, "
            f"${contract.annual_value:,.0f}/year, "
            f"{'auto-renew' if contract.auto_renew else 'fixed term'}, "
            f"expires {contract.expiration_date.isoformat()}."
        )

    doc.add_heading("Key Vendor Relationships", level=2)
    for contract in LEGAL_CONTRACTS:
        if contract.counterparty_id.startswith("VEND"):
            doc.add_paragraph(
                f"{contract.counterparty_name} ({contract.contract_id}): "
                f"{contract.contract_type.title()}, ${contract.annual_value:,.0f}/year."
            )
            if contract.notes:
                doc.add_paragraph(contract.notes, style="List Bullet")

    save_docx_deterministic(doc, output_path)
    return location


# ── Diligence request list (xlsx) ──────────────────────────────────────────


def write_diligence_request_list(
    output_path: Path,
    canaries: CanaryRegistry,
    canary_key: str,
) -> str:
    """Write a diligence request tracker as an xlsx file.

    Includes all diligence issues as request items with status tracking.

    Returns the canary embedding location description.
    """
    wb = new_workbook()
    canary = canaries.canary_for(canary_key)
    location = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Diligence Requests"

    headers = [
        "Issue ID", "Contract ID", "Clause ID", "Issue Type",
        "Severity", "Description", "Recommended Action", "Source References",
    ]
    write_header_row(ws, headers)

    for row_idx, issue in enumerate(LEGAL_DILIGENCE_ISSUES, 2):
        ws.cell(row=row_idx, column=1, value=issue.issue_id)
        ws.cell(row=row_idx, column=2, value=issue.contract_id)
        ws.cell(row=row_idx, column=3, value=issue.clause_id or "")
        ws.cell(row=row_idx, column=4, value=issue.issue_type)
        ws.cell(row=row_idx, column=5, value=issue.severity)
        ws.cell(row=row_idx, column=6, value=issue.description)
        ws.cell(row=row_idx, column=7, value=issue.recommended_action)
        ws.cell(row=row_idx, column=8, value=", ".join(issue.source_refs))

    # Auto-fit column widths (approximate)
    for col_idx, header in enumerate(headers, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = max(
            len(header) + 2, 15
        )

    save_xlsx_deterministic(wb, output_path)
    return location


# ── Helpers ─────────────────────────────────────────────────────────────────


def _add_row(table: Any, label: str, value: str) -> None:
    """Add a label-value row to a docx table."""
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
