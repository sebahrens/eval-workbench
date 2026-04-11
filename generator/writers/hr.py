"""HR document writer helpers for TC-20 / TC-21.

Deterministic generators for:
- Employment agreements (docx) — one per EmploymentAgreement
- Employee census / HR schedule (xlsx)
- Severance exposure schedule (xlsx)
- Retention award schedule (xlsx)
- Contractor roster (xlsx)
- Diligence request tracker (xlsx)

All functions accept model data from generator.model.hr_diligence and
embed canaries via the CanaryRegistry.  They never invent facts outside
the canonical model.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry, embed_canary_docx, embed_canary_xlsx
from generator.model.hr_diligence import (
    CONTRACTOR_CLASSIFICATION_SIGNALS,
    DILIGENCE_REQUESTS,
    EMPLOYMENT_AGREEMENTS,
    RETENTION_AWARDS,
    SEVERANCE_EXPOSURES,
    EmploymentAgreement,
)
from generator.writers import (
    BOLD_FONT,
    new_workbook,
    save_docx_deterministic,
    save_xlsx_deterministic,
    write_header_row,
)

# ── Employment agreement documents (docx) ──────────────────────────────────


def write_employment_agreement(
    agreement: EmploymentAgreement,
    output_path: Path,
    canaries: CanaryRegistry,
    canary_key: str,
) -> str:
    """Write a single employment agreement as a docx file.

    For agreements where executed=False, the document is marked as DRAFT.

    Returns the canary embedding location description.
    """
    doc = Document()
    canary = canaries.canary_for(canary_key)
    location = embed_canary_docx(doc, canary)

    # Draft watermark in title if not executed
    status = "DRAFT — " if not agreement.executed else ""
    doc.add_heading(
        f"{status}Employment Agreement — {agreement.agreement_id}", level=1
    )

    # Agreement details table
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    _add_row(table, "Agreement ID", agreement.agreement_id)
    _add_row(table, "Employee", agreement.employee_name)
    _add_row(table, "Title", agreement.employee_title)
    _add_row(table, "Entity", agreement.entity_code)
    _add_row(table, "Effective Date", agreement.effective_date.isoformat())
    _add_row(table, "Status", "Executed" if agreement.executed else "Draft — executed copy not on file")

    doc.add_heading("Compensation", level=2)
    comp_table = doc.add_table(rows=0, cols=2)
    comp_table.style = "Table Grid"
    _add_row(comp_table, "Base Salary", f"${agreement.base_salary:,}")
    _add_row(comp_table, "Bonus Target", f"{agreement.bonus_target_pct * 100:.0f}% of base salary")

    doc.add_heading("Severance & Change of Control", level=2)
    sev_table = doc.add_table(rows=0, cols=2)
    sev_table.style = "Table Grid"
    _add_row(sev_table, "Severance Multiplier", f"{agreement.severance_multiplier}x base salary")
    _add_row(sev_table, "Severance Notice Period", f"{agreement.severance_months} months")
    _add_row(sev_table, "Change of Control Multiplier", f"{agreement.change_of_control_multiplier}x base salary")

    doc.add_heading("Restrictive Covenants", level=2)
    cov_table = doc.add_table(rows=0, cols=2)
    cov_table.style = "Table Grid"
    _add_row(cov_table, "Non-Compete Period", f"{agreement.non_compete_months} months")
    _add_row(cov_table, "IP Assignment", "Yes" if agreement.ip_assignment else "No")

    if agreement.notes:
        doc.add_heading("Notes", level=2)
        doc.add_paragraph(agreement.notes)

    save_docx_deterministic(doc, output_path)
    return location


def write_all_employment_agreements(
    output_dir: Path,
    canaries: CanaryRegistry,
    canary_keys: dict[str, str],
) -> dict[str, str]:
    """Write one docx per EmploymentAgreement.

    Parameters
    ----------
    canary_keys:
        Mapping of agreement_id → canary file key.

    Returns
    -------
    dict mapping agreement_id → canary embedding location.
    """
    locations: dict[str, str] = {}
    for ea in EMPLOYMENT_AGREEMENTS:
        ckey = canary_keys[ea.agreement_id]
        fname = f"agreement_{ea.agreement_id.lower()}.docx"
        loc = write_employment_agreement(ea, output_dir / fname, canaries, ckey)
        locations[ea.agreement_id] = loc
    return locations


# ── Employee census (xlsx) ──────────────────────────────────────────────────


def write_employee_census(
    output_path: Path,
    canaries: CanaryRegistry,
    canary_key: str,
) -> str:
    """Write the executive employee census as an xlsx file.

    Covers all employees with employment agreements (the executive layer).

    Returns the canary embedding location description.
    """
    wb = new_workbook()
    canary = canaries.canary_for(canary_key)
    location = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Executive Census"

    headers = [
        "Agreement ID", "Employee Name", "Title", "Entity",
        "Base Salary", "Bonus Target %", "Effective Date", "Status",
    ]
    write_header_row(ws, headers)

    for row_idx, ea in enumerate(EMPLOYMENT_AGREEMENTS, 2):
        ws.cell(row=row_idx, column=1, value=ea.agreement_id)
        ws.cell(row=row_idx, column=2, value=ea.employee_name)
        ws.cell(row=row_idx, column=3, value=ea.employee_title)
        ws.cell(row=row_idx, column=4, value=ea.entity_code)
        cell = ws.cell(row=row_idx, column=5, value=ea.base_salary)
        cell.number_format = "#,##0"
        ws.cell(row=row_idx, column=6, value=float(ea.bonus_target_pct * 100))
        ws.cell(row=row_idx, column=7, value=ea.effective_date.isoformat())
        ws.cell(
            row=row_idx, column=8,
            value="Executed" if ea.executed else "Draft — not on file",
        )

    _autofit_columns(ws, headers)
    save_xlsx_deterministic(wb, output_path)
    return location


# ── Severance exposure schedule (xlsx) ──────────────────────────────────────


def write_severance_schedule(
    output_path: Path,
    canaries: CanaryRegistry,
    canary_key: str,
) -> str:
    """Write the severance exposure schedule as an xlsx file.

    Returns the canary embedding location description.
    """
    wb = new_workbook()
    canary = canaries.canary_for(canary_key)
    location = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Severance Exposure"

    headers = [
        "Exposure ID", "Employee Name", "Title", "Entity",
        "Base Salary", "Multiplier", "Estimated Payout", "Trigger", "Accrued",
    ]
    write_header_row(ws, headers)

    for row_idx, sev in enumerate(SEVERANCE_EXPOSURES, 2):
        ws.cell(row=row_idx, column=1, value=sev.exposure_id)
        ws.cell(row=row_idx, column=2, value=sev.employee_name)
        ws.cell(row=row_idx, column=3, value=sev.employee_title)
        ws.cell(row=row_idx, column=4, value=sev.entity_code)
        cell = ws.cell(row=row_idx, column=5, value=sev.base_salary)
        cell.number_format = "#,##0"
        ws.cell(row=row_idx, column=6, value=float(sev.severance_multiplier))
        cell = ws.cell(row=row_idx, column=7, value=int(sev.estimated_payout))
        cell.number_format = "#,##0"
        ws.cell(row=row_idx, column=8, value=sev.trigger.replace("_", " ").title())
        ws.cell(row=row_idx, column=9, value="Yes" if sev.accrued else "No")

    # Total row
    total_row = len(SEVERANCE_EXPOSURES) + 2
    ws.cell(row=total_row, column=1, value="TOTAL").font = BOLD_FONT
    total_cell = ws.cell(
        row=total_row, column=7,
        value=int(sum(s.estimated_payout for s in SEVERANCE_EXPOSURES)),
    )
    total_cell.number_format = "#,##0"
    total_cell.font = BOLD_FONT

    if any(s for s in SEVERANCE_EXPOSURES if s.notes):
        notes_row = total_row + 2
        ws.cell(row=notes_row, column=1, value="Notes:").font = BOLD_FONT
        for i, sev in enumerate(SEVERANCE_EXPOSURES):
            if sev.notes:
                ws.cell(row=notes_row + 1 + i, column=1, value=f"{sev.exposure_id}: {sev.notes}")

    _autofit_columns(ws, headers)
    save_xlsx_deterministic(wb, output_path)
    return location


# ── Retention award schedule (xlsx) ─────────────────────────────────────────


def write_retention_schedule(
    output_path: Path,
    canaries: CanaryRegistry,
    canary_key: str,
) -> str:
    """Write the retention award schedule as an xlsx file.

    Returns the canary embedding location description.
    """
    wb = new_workbook()
    canary = canaries.canary_for(canary_key)
    location = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Retention Awards"

    headers = [
        "Award ID", "Employee Name", "Title", "Entity",
        "Award Amount", "Vesting Date", "Retention Period (Months)",
        "Forfeiture Conditions",
    ]
    write_header_row(ws, headers)

    for row_idx, ret in enumerate(RETENTION_AWARDS, 2):
        ws.cell(row=row_idx, column=1, value=ret.award_id)
        ws.cell(row=row_idx, column=2, value=ret.employee_name)
        ws.cell(row=row_idx, column=3, value=ret.employee_title)
        ws.cell(row=row_idx, column=4, value=ret.entity_code)
        cell = ws.cell(row=row_idx, column=5, value=int(ret.award_amount))
        cell.number_format = "#,##0"
        ws.cell(row=row_idx, column=6, value=ret.vesting_date.isoformat())
        ws.cell(row=row_idx, column=7, value=ret.retention_period_months)
        ws.cell(row=row_idx, column=8, value=ret.forfeiture_conditions)

    # Total row
    total_row = len(RETENTION_AWARDS) + 2
    ws.cell(row=total_row, column=1, value="TOTAL").font = BOLD_FONT
    total_cell = ws.cell(
        row=total_row, column=5,
        value=int(sum(r.award_amount for r in RETENTION_AWARDS)),
    )
    total_cell.number_format = "#,##0"
    total_cell.font = BOLD_FONT

    if any(r for r in RETENTION_AWARDS if r.notes):
        notes_row = total_row + 2
        ws.cell(row=notes_row, column=1, value="Notes:").font = BOLD_FONT
        for i, ret in enumerate(RETENTION_AWARDS):
            if ret.notes:
                ws.cell(row=notes_row + 1 + i, column=1, value=f"{ret.award_id}: {ret.notes}")

    _autofit_columns(ws, headers)
    save_xlsx_deterministic(wb, output_path)
    return location


# ── Contractor roster (xlsx) ────────────────────────────────────────────────


def write_contractor_roster(
    output_path: Path,
    canaries: CanaryRegistry,
    canary_key: str,
) -> str:
    """Write the contractor classification roster as an xlsx file.

    Returns the canary embedding location description.
    """
    wb = new_workbook()
    canary = canaries.canary_for(canary_key)
    location = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Contractor Roster"

    headers = [
        "Signal ID", "Contractor Name", "Entity", "Role Description",
        "Tenure (Months)", "Exclusive", "Company Equipment", "Set Hours",
        "Risk Level", "Notes",
    ]
    write_header_row(ws, headers)

    for row_idx, ccs in enumerate(CONTRACTOR_CLASSIFICATION_SIGNALS, 2):
        ws.cell(row=row_idx, column=1, value=ccs.signal_id)
        ws.cell(row=row_idx, column=2, value=ccs.contractor_name)
        ws.cell(row=row_idx, column=3, value=ccs.entity_code)
        ws.cell(row=row_idx, column=4, value=ccs.role_description)
        ws.cell(row=row_idx, column=5, value=ccs.tenure_months)
        ws.cell(row=row_idx, column=6, value="Yes" if ccs.exclusive_engagement else "No")
        ws.cell(row=row_idx, column=7, value="Yes" if ccs.uses_company_equipment else "No")
        ws.cell(row=row_idx, column=8, value="Yes" if ccs.has_set_hours else "No")
        ws.cell(row=row_idx, column=9, value=ccs.risk_level.title())
        ws.cell(row=row_idx, column=10, value=ccs.notes)

    _autofit_columns(ws, headers)
    save_xlsx_deterministic(wb, output_path)
    return location


# ── HR diligence request tracker (xlsx) ─────────────────────────────────────


def write_hr_diligence_requests(
    output_path: Path,
    canaries: CanaryRegistry,
    canary_key: str,
) -> str:
    """Write the HR diligence request tracker as an xlsx file.

    Returns the canary embedding location description.
    """
    wb = new_workbook()
    canary = canaries.canary_for(canary_key)
    location = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Diligence Requests"

    headers = [
        "Request ID", "Category", "Description", "Requested From",
        "Status", "Due Date", "Received Date", "Source References",
    ]
    write_header_row(ws, headers)

    for row_idx, dr in enumerate(DILIGENCE_REQUESTS, 2):
        ws.cell(row=row_idx, column=1, value=dr.request_id)
        ws.cell(row=row_idx, column=2, value=dr.category)
        ws.cell(row=row_idx, column=3, value=dr.description)
        ws.cell(row=row_idx, column=4, value=dr.requested_from)
        ws.cell(row=row_idx, column=5, value=dr.status)
        ws.cell(row=row_idx, column=6, value=dr.due_date.isoformat())
        ws.cell(
            row=row_idx, column=7,
            value=dr.received_date.isoformat() if dr.received_date else "Not received",
        )
        ws.cell(row=row_idx, column=8, value=", ".join(dr.source_refs))

    _autofit_columns(ws, headers)
    save_xlsx_deterministic(wb, output_path)
    return location


# ── Helpers ─────────────────────────────────────────────────────────────────


def _add_row(table: Any, label: str, value: str) -> None:
    """Add a label-value row to a docx table."""
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value


def _autofit_columns(ws: Any, headers: list[str]) -> None:
    """Set approximate column widths based on header length."""
    for col_idx, header in enumerate(headers, 1):
        ws.column_dimensions[
            openpyxl.utils.get_column_letter(col_idx)
        ].width = max(len(header) + 2, 15)
