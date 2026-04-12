"""Formatter: TC-10-EU — VAT and Cross-Border Tax Position Analysis.

Emits:
- test_cases/TC-10-EU/input_files/tc10eu_intercompany_sales_fy2025.xlsx
  Intercompany and third-party sales with VAT treatment (~24 IC rows + summary).
- test_cases/TC-10-EU/input_files/tc10eu_vat_registrations.xlsx
  VAT registrations per entity (5 rows).
- test_cases/TC-10-EU/input_files/tc10eu_vat_returns_summary_fy2025.xlsx
  Quarterly VAT return summaries (16 rows: 4 entities x 4 quarters).
- test_cases/TC-10-EU/input_files/tc10eu_eu_vat_rules_reference.docx
  Reference document covering EU VAT rules.
- test_cases/TC-10-EU/prompt.md
- test_cases/TC-10-EU/expected_behavior.md
- gold_standards/TC-10-EU_gold.json

Planted error:
  ERR-EU-010: vat_treatment_error — CE (Netherlands) invoiced CM (France)
  for Q3 management fees with 20% French VAT. CE is not FR-registered;
  reverse charge under Art. 196 should apply.

Uses deterministic European VAT model — never hardcodes numbers that should
come from the model.
"""

from __future__ import annotations

import datetime
import io
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from generator.canaries import (
    CanaryRegistry,
    embed_canary_docx,
    embed_canary_xlsx,
)
from generator.errors import (
    ErrorRegistry,
    PlantedError,
)
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel
from generator.model.vat_eu import (
    ALL_CANARY_KEYS_TC10EU,
    CE_TO_CM_MGMT_FEE,
    CE_TO_CD_MGMT_FEE,
    CE_TO_CP_MGMT_FEE,
    CM_ITALIAN_SALES,
    CM_TO_CP_ROYALTY,
    CP_TO_CD_FINISHED_GOODS,
    CP_TO_CD_Q2_TOTAL,
    CP_TO_CD_Q2_PER_SHIPMENT,
    CP_TO_CD_Q2_UNDOCUMENTED,
    CP_TO_CM_RAW_MATERIALS,
    ENTITY_NAMES_VAT,
    ENTITY_JURISDICTIONS_VAT,
    ERR_EU_010_AMOUNT,
    ERR_EU_010_QUARTER,
    ERR_EU_010_VAT_CHARGED,
    MGMT_FEE_PCT,
    ROYALTY_PCT,
    VAT_IDS,
    VAT_RATES,
    generate_ic_sales_eu,
    generate_vat_registrations,
    generate_vat_returns,
)

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-10-EU"
_INPUT_DIR = f"test_cases/{_TC}/input_files"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)
_FIXED_ZIP_DT = (2025, 3, 15, 9, 0, 0)


# ── Deterministic save helpers ───────────────────────────────────────────────


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)
    wb.properties.created = _FIXED_DATETIME
    wb.properties.modified = _FIXED_DATETIME
    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _whole_euros(d: Decimal) -> int:
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


# ── Excel styling helpers ────────────────────────────────────────────────────

_HEADER_FONT = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
_HEADER_FILL = PatternFill("solid", fgColor="2F5496")
_HEADER_ALIGN = Alignment(horizontal="center", wrap_text=True)
_DATA_FONT = Font(name="Calibri", size=10)
_THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)


def _hdr_cell(ws, row, col, value, width=None):
    cell = ws.cell(row=row, column=col, value=value)
    cell.font = _HEADER_FONT
    cell.fill = _HEADER_FILL
    cell.alignment = _HEADER_ALIGN
    cell.border = _THIN_BORDER
    if width:
        ws.column_dimensions[cell.column_letter].width = width
    return cell


def _data_cell(ws, row, col, value, fmt=None):
    cell = ws.cell(row=row, column=col, value=value)
    cell.font = _DATA_FONT
    cell.border = _THIN_BORDER
    if fmt:
        cell.number_format = fmt
    return cell


# ── 1. Intercompany Sales XLSX ───────────────────────────────────────────────


def _write_ic_sales_xlsx(
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write intercompany sales with VAT treatment and third-party summary."""
    ic_rows = generate_ic_sales_eu()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Intercompany Sales"

    # Title
    ws.cell(row=1, column=1, value="Cascade Europe \u2014 Intercompany Sales FY2025")
    ws.cell(row=1, column=1).font = Font(name="Calibri", size=14, bold=True)

    # Headers at row 3
    headers = [
        ("Seller Entity", 22),
        ("Buyer Entity", 22),
        ("Description", 55),
        ("Amount (EUR)", 18),
        ("VAT Treatment Applied", 35),
        ("Invoice VAT Rate", 18),
        ("Incoterms", 18),
        ("Proof of Dispatch", 18),
    ]
    for col, (hdr, width) in enumerate(headers, 1):
        _hdr_cell(ws, 3, col, hdr, width)

    # Data rows starting at row 4
    for i, row in enumerate(ic_rows, 4):
        _data_cell(ws, i, 1, ENTITY_NAMES_VAT.get(row.seller, row.seller))
        _data_cell(ws, i, 2, ENTITY_NAMES_VAT.get(row.buyer, row.buyer))
        _data_cell(ws, i, 3, row.description)
        _data_cell(ws, i, 4, _whole_euros(row.amount_eur), "#,##0")
        _data_cell(ws, i, 5, row.vat_treatment)
        _data_cell(ws, i, 6, row.invoice_vat_rate)
        _data_cell(ws, i, 7, row.incoterms)
        _data_cell(ws, i, 8, row.proof_of_dispatch)

    # Third-Party Sales Summary section
    tp_start = len(ic_rows) + 4 + 2  # 2 blank rows after IC data
    ws.cell(row=tp_start, column=1, value="Third-Party Sales Summary").font = Font(
        name="Calibri", size=12, bold=True,
    )
    tp_headers = [
        ("Entity", 22),
        ("Domestic Sales (EUR)", 22),
        ("Intra-EU Sales (EUR)", 22),
        ("Non-EU Sales (EUR)", 22),
        ("Notes", 45),
    ]
    tp_hdr_row = tp_start + 1
    for col, (hdr, width) in enumerate(tp_headers, 1):
        _hdr_cell(ws, tp_hdr_row, col, hdr, width)

    # Third-party data from design
    tp_data = [
        (
            ENTITY_NAMES_VAT["CE"], 2_500_000, 0, 0,
            "NL domestic advisory services only",
        ),
        (
            ENTITY_NAMES_VAT["CP"], 28_000_000, 6_000_000, 0,
            "DE domestic + intra-EU (non-group) sales",
        ),
        (
            ENTITY_NAMES_VAT["CM"], 20_000_000, _whole_euros(CM_ITALIAN_SALES), 0,
            "FR domestic + Italian customers (\u20ac380k) \u2014 no IT VAT registration",
        ),
        (
            ENTITY_NAMES_VAT["CD"], 21_000_000, 0, 0,
            "UK domestic (\u00a318M \u2248 \u20ac21M)",
        ),
    ]
    for j, (entity, dom, intra, non_eu, notes) in enumerate(tp_data, tp_hdr_row + 1):
        _data_cell(ws, j, 1, entity)
        _data_cell(ws, j, 2, dom, "#,##0")
        _data_cell(ws, j, 3, intra, "#,##0")
        _data_cell(ws, j, 4, non_eu, "#,##0")
        _data_cell(ws, j, 5, notes)

    # Canary
    canary = canaries.canary_for("tc10eu_intercompany_sales_fy2025")
    loc = embed_canary_xlsx(wb, canary)
    canaries.set_location(
        "tc10eu_intercompany_sales_fy2025",
        f"{_INPUT_DIR}/tc10eu_intercompany_sales_fy2025.xlsx",
        loc,
    )

    file_path = output_dir / _INPUT_DIR / "tc10eu_intercompany_sales_fy2025.xlsx"
    file_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, file_path)
    manifest.register(
        f"{_INPUT_DIR}/tc10eu_intercompany_sales_fy2025.xlsx",
        "xlsx",
        canary=canary,
        test_cases=[_TC],
    )

    # Register planted error ERR-EU-010
    errors.add(PlantedError(
        error_id="ERR-EU-010",
        file=f"{_INPUT_DIR}/tc10eu_intercompany_sales_fy2025.xlsx",
        location="Sheet 'Intercompany Sales', CE\u2192CM Q3 management fee row",
        type="vat_treatment_error",
        description=(
            "CE (Netherlands) invoiced CM (France) for Q3 management fees "
            f"(\u20ac{_whole_euros(ERR_EU_010_AMOUNT):,}) with 20% French VAT "
            f"(\u20ac{_whole_euros(ERR_EU_010_VAT_CHARGED):,}). CE is not registered for "
            "VAT in France \u2014 reverse charge under Art. 196 should apply. "
            f"CM has an invalid invoice; deduction of \u20ac{_whole_euros(ERR_EU_010_VAT_CHARGED):,} "
            "input VAT is invalid."
        ),
        severity="material",
        which_test_cases_should_catch=["TC-10-EU"],
    ))


# ── 2. VAT Registrations XLSX ───────────────────────────────────────────────


def _write_vat_registrations_xlsx(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write VAT registrations per entity."""
    regs = generate_vat_registrations()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "VAT Registrations"

    # Title
    ws.cell(row=1, column=1, value="Cascade Europe \u2014 VAT Registrations")
    ws.cell(row=1, column=1).font = Font(name="Calibri", size=14, bold=True)

    # Headers at row 3
    headers = [
        ("Entity Code", 14),
        ("Country", 18),
        ("VAT ID", 22),
        ("Registration Date", 18),
        ("VAT Group (Y/N)", 16),
        ("Fiscal Representative", 22),
        ("EC Sales List Filed (Y/N)", 24),
        ("Intrastat Threshold Exceeded (Y/N)", 30),
        ("Status", 28),
    ]
    for col, (hdr, width) in enumerate(headers, 1):
        _hdr_cell(ws, 3, col, hdr, width)

    # Data rows starting at row 4
    for i, reg in enumerate(regs, 4):
        _data_cell(ws, i, 1, reg.entity_code)
        _data_cell(ws, i, 2, reg.country)
        _data_cell(ws, i, 3, reg.vat_id)
        _data_cell(ws, i, 4, reg.registration_date)
        _data_cell(ws, i, 5, reg.vat_group)
        _data_cell(ws, i, 6, reg.fiscal_representative or "\u2014")
        _data_cell(ws, i, 7, reg.ecsl_filed)
        _data_cell(ws, i, 8, reg.intrastat_exceeded)
        _data_cell(ws, i, 9, reg.status)

    # Canary
    canary = canaries.canary_for("tc10eu_vat_registrations")
    loc = embed_canary_xlsx(wb, canary)
    canaries.set_location(
        "tc10eu_vat_registrations",
        f"{_INPUT_DIR}/tc10eu_vat_registrations.xlsx",
        loc,
    )

    file_path = output_dir / _INPUT_DIR / "tc10eu_vat_registrations.xlsx"
    file_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, file_path)
    manifest.register(
        f"{_INPUT_DIR}/tc10eu_vat_registrations.xlsx",
        "xlsx",
        canary=canary,
        test_cases=[_TC],
    )


# ── 3. VAT Returns Summary XLSX ─────────────────────────────────────────────


def _write_vat_returns_xlsx(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write quarterly VAT return summaries."""
    returns = generate_vat_returns()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Quarterly VAT Returns"

    # Title
    ws.cell(row=1, column=1, value="Cascade Europe \u2014 Quarterly VAT Returns FY2025")
    ws.cell(row=1, column=1).font = Font(name="Calibri", size=14, bold=True)

    # Headers at row 3
    headers = [
        ("Entity", 18),
        ("Quarter", 10),
        ("Output VAT (Domestic)", 22),
        ("Output VAT (Intra-EU/Export at 0%)", 30),
        ("Input VAT (Domestic)", 22),
        ("Input VAT (Reverse Charge Self-Assessed)", 35),
        ("VAT Payable/Refundable", 22),
        ("Filing Status", 28),
    ]
    for col, (hdr, width) in enumerate(headers, 1):
        _hdr_cell(ws, 3, col, hdr, width)

    # Data rows starting at row 4
    for i, ret in enumerate(returns, 4):
        _data_cell(ws, i, 1, ENTITY_NAMES_VAT.get(ret.entity_code, ret.entity_code))
        _data_cell(ws, i, 2, ret.quarter)
        _data_cell(ws, i, 3, _whole_euros(ret.output_vat_domestic), "#,##0")
        _data_cell(ws, i, 4, _whole_euros(ret.output_vat_intra_eu_export), "#,##0")
        _data_cell(ws, i, 5, _whole_euros(ret.input_vat_domestic), "#,##0")
        _data_cell(ws, i, 6, _whole_euros(ret.input_vat_reverse_charge), "#,##0")
        _data_cell(ws, i, 7, _whole_euros(ret.vat_payable), "#,##0")
        _data_cell(ws, i, 8, ret.filing_status)

    # Canary
    canary = canaries.canary_for("tc10eu_vat_returns_summary_fy2025")
    loc = embed_canary_xlsx(wb, canary)
    canaries.set_location(
        "tc10eu_vat_returns_summary_fy2025",
        f"{_INPUT_DIR}/tc10eu_vat_returns_summary_fy2025.xlsx",
        loc,
    )

    file_path = output_dir / _INPUT_DIR / "tc10eu_vat_returns_summary_fy2025.xlsx"
    file_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, file_path)
    manifest.register(
        f"{_INPUT_DIR}/tc10eu_vat_returns_summary_fy2025.xlsx",
        "xlsx",
        canary=canary,
        test_cases=[_TC],
    )


# ── 4. EU VAT Rules Reference DOCX ──────────────────────────────────────────


def _write_vat_rules_docx(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write EU VAT rules reference document."""
    from docx import Document

    doc = Document()

    # Title
    doc.add_heading(
        "EU VAT Rules \u2014 Cross-Border Transaction Reference", level=0,
    )

    # Section 1: Intra-Community Supply of Goods
    doc.add_heading(
        "1. Intra-Community Supply of Goods (Article 138 VAT Directive)", level=1,
    )
    doc.add_paragraph(
        "Conditions for the zero-rate (0%) on intra-Community supplies of goods:"
    )
    for bullet in [
        "The acquirer must hold a valid VAT identification number issued by "
        "another member state.",
        "The goods must be physically transported from one member state to another.",
        "The supplier must retain proof of dispatch or transport (e.g., CMR "
        "waybill, bill of lading, carrier confirmation).",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")
    doc.add_paragraph("Documentation required:")
    for bullet in [
        "Commercial invoice referencing the acquirer\u2019s VAT ID and Art. 138",
        "Transport documents (CMR waybill, bill of lading, or equivalent)",
        "Proof of dispatch: signed CMR, carrier confirmation, or arrival certificate",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    # Section 2: Reverse Charge on Cross-Border Services
    doc.add_heading(
        "2. Reverse Charge on Cross-Border Services "
        "(Article 196 VAT Directive)", level=1,
    )
    doc.add_paragraph(
        "For B2B services, the general rule is that the place of supply is "
        "where the customer is established (Article 44 VAT Directive). "
        "The supplier does not charge local VAT on its invoice."
    )
    doc.add_paragraph(
        "The customer self-assesses output VAT in its own member state at the "
        "local rate, and simultaneously claims input VAT deduction (net effect "
        "is zero if fully recoverable). The supplier\u2019s invoice must state "
        "\u2018Reverse charge \u2014 Article 196 VAT Directive\u2019 and must not include VAT."
    )

    # Section 3: UK Post-Brexit Treatment
    doc.add_heading("3. UK Post-Brexit Treatment", level=1)
    doc.add_paragraph(
        "Since 1 January 2021, the United Kingdom is treated as a third "
        "country for EU VAT purposes. Key consequences:"
    )
    for bullet in [
        "Exports from EU to UK are zero-rated (treated as exports to a non-EU "
        "country) but require customs declarations (export/import).",
        "Imports from UK into the EU are subject to import VAT at the member "
        "state of importation\u2019s standard rate.",
        "Intra-Community acquisition treatment no longer applies to UK transactions.",
        "Postponed VAT accounting may apply in the UK for imports.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    # Section 4: Triangulation Simplification
    doc.add_heading(
        "4. Triangulation Simplification (Article 141 VAT Directive)", level=1,
    )
    doc.add_paragraph(
        "Triangulation (or ABC transaction simplification) applies when goods "
        "move directly from party A in one member state to party C in another "
        "member state, but the invoicing goes A \u2192 B \u2192 C with B in a third "
        "member state."
    )
    doc.add_paragraph(
        "Under the simplification, the intermediate party B can avoid VAT "
        "registration in the country of arrival (C\u2019s member state). B issues "
        "a zero-rated invoice to C with a reference to Article 141. Strict "
        "conditions apply: B must hold a VAT number in a member state different "
        "from both A\u2019s and C\u2019s member states, and the goods must be shipped "
        "directly from A to C."
    )

    # Section 5: Call-off Stock Arrangements
    doc.add_heading(
        "5. Call-off Stock Arrangements (Article 17a VAT Directive)", level=1,
    )
    doc.add_paragraph(
        "Under call-off stock simplification, goods may be transferred to "
        "another member state for a known customer without triggering a deemed "
        "intra-Community supply at the time of transport. Instead, the supply "
        "is deemed to occur when the customer takes the goods from the stock."
    )
    doc.add_paragraph(
        "Requirements: the supplier must register the goods in a call-off "
        "stock register, and the goods must be called off within 12 months. "
        "The customer must be identified by VAT number before the transport "
        "begins."
    )

    # Section 6: Permanent Establishment for VAT Purposes
    doc.add_heading(
        "6. Permanent Establishment for VAT Purposes", level=1,
    )
    doc.add_paragraph(
        "A VAT fixed establishment is not the same as a corporate tax PE "
        "under Art. 5 OECD Model Convention."
    )
    doc.add_paragraph(
        "A fixed establishment for VAT purposes requires sufficient human "
        "and technical resources to make or receive taxable supplies "
        "independently. The test is different from and independent of the "
        "income tax PE test under double tax treaties."
    )
    doc.add_paragraph(
        "An entity may have a VAT fixed establishment in a country without "
        "having an income tax PE there, and vice versa. The CJEU case law "
        "(e.g., Welmory C-605/12, Titanium C-931/19) provides further "
        "guidance on the criteria."
    )

    # Set core properties
    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME

    # Canary
    canary = canaries.canary_for("tc10eu_eu_vat_rules_reference")
    loc = embed_canary_docx(doc, canary)
    canaries.set_location(
        "tc10eu_eu_vat_rules_reference",
        f"{_INPUT_DIR}/tc10eu_eu_vat_rules_reference.docx",
        loc,
    )

    file_path = output_dir / _INPUT_DIR / "tc10eu_eu_vat_rules_reference.docx"
    file_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, file_path)
    manifest.register(
        f"{_INPUT_DIR}/tc10eu_eu_vat_rules_reference.docx",
        "docx",
        canary=canary,
        test_cases=[_TC],
    )


# ── 5. Prompt ────────────────────────────────────────────────────────────────


def _write_prompt(output_dir: Path) -> None:
    text = """\
Analyze the VAT and cross-border tax position for the Cascade Europe
Holdings B.V. group for FY2025. The group comprises four entities:

- CE: Cascade Europe Holdings B.V. (Netherlands) \u2014 holding company
- CP: Cascade Pr\u00e4zisionsteile GmbH (Germany) \u2014 licensed manufacturer
- CM: Cascade Mat\u00e9riaux Avanc\u00e9s SAS (France) \u2014 R&D centre and IP developer
- CD: Cascade Distribution Services Ltd (United Kingdom) \u2014 distributor

Using the intercompany sales data, VAT registrations, quarterly VAT return
summaries, and EU VAT rules reference provided:

1. Analyze each intercompany transaction for correct VAT treatment:
   - Intra-Community supplies of goods (Art. 138): verify zero-rating
     conditions and proof of dispatch
   - Cross-border services (Art. 196): verify reverse charge applied correctly
   - UK post-Brexit: verify export/import treatment
   - Management fees: verify VAT treatment per entity pair

2. Review VAT registrations for completeness:
   - Are all required registrations in place?
   - Are there any pending registrations that affect compliance?
   - Are there any missing registrations that should exist?

3. Reconcile quarterly VAT returns:
   - Cross-check output VAT and input VAT amounts against transaction data
   - Identify any anomalies in reverse-charge self-assessment
   - Flag any quarters with unusual filing status

4. Assess VAT permanent establishment risk:
   - Distinguish between VAT fixed establishment and income tax PE
   - Identify any entities with potential VAT PE exposure

5. Identify data gaps and quantify risks:
   - Flag missing documentation (proof of dispatch, registrations)
   - Quantify potential VAT exposure from documentation gaps
   - Recommend corrective actions

Export:
- Transaction-by-transaction VAT analysis as Excel with assessment per row
- VAT registration gap analysis as a summary table
- Quarterly VAT reconciliation as Excel
- Risk summary as Word document with quantified exposures
- Recommendations for remediation
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── 6. Expected Behavior ────────────────────────────────────────────────────


def _write_expected_behavior(output_dir: Path) -> None:
    q2_undoc = _whole_euros(CP_TO_CD_Q2_UNDOCUMENTED)
    q2_vat_risk = _whole_euros(CP_TO_CD_Q2_UNDOCUMENTED * VAT_RATES["DE"])
    err_amount = _whole_euros(ERR_EU_010_AMOUNT)
    err_vat = _whole_euros(ERR_EU_010_VAT_CHARGED)
    italian_sales = _whole_euros(CM_ITALIAN_SALES)

    text = f"""\
# TC-10-EU: VAT and Cross-Border Tax Position \u2014 Expected Behavior

## Key Findings the Agent Should Produce

1. **Transaction-by-transaction VAT analysis**: Each intercompany transaction
   must be assessed for correct VAT treatment against the EU VAT rules
   reference. The analysis must cover:
   - Goods flows (CP\u2192CM, CP\u2192CD): zero-rating conditions, proof of dispatch
   - Services (CE management fees, CM\u2192CP royalty): reverse charge under Art. 196
   - UK transactions: post-Brexit third-country treatment

2. **ERR-EU-010 detection**: CE (Netherlands) invoiced CM (France) for Q3
   management fees (\u20ac{err_amount:,}) with 20% French VAT (\u20ac{err_vat:,}) charged
   on the invoice. CE is not registered for VAT in France \u2014 reverse charge
   under Art. 196 should apply. The \u20ac{err_vat:,} VAT charged is incorrect;
   CM cannot validly deduct this amount as input VAT.

3. **Missing data identification** (3 gaps):
   - **Q2 proof of dispatch**: CP\u2192CD Q2 finished goods shipments show "Partial"
     proof of dispatch \u2014 2 of 5 shipments (~\u20ac{q2_undoc:,}) lack
     documentation. Zero-rating is at risk without proof.
   - **Polish VAT registration**: CP has a pending Polish VAT registration
     (applied 2025-03-15). If CP makes supplies in Poland before registration
     completes, compliance risk arises.
   - **Italian sales without registration**: CM (France) has \u20ac{italian_sales:,}
     in sales to Italian customers but holds no Italian VAT registration. If
     these are supplies of goods with installation or local supplies, an
     Italian registration may be required.

4. **VAT PE vs income tax PE distinction**: The EU VAT rules reference notes
   that a VAT fixed establishment is not the same as an income tax PE under
   Art. 5 OECD Model Convention. Agent must demonstrate awareness of this
   distinction when assessing PE risk.

5. **UK post-Brexit treatment**: CD (UK) transactions must be treated as
   exports to/imports from a third country. No intra-Community acquisition
   treatment. Customs declarations required.

6. **CE Q4 pending assessment flag**: CE\u2019s Q4 VAT return shows "Filed \u2014
   Pending Assessment" (all other quarters show "Filed \u2014 Assessed"). Agent
   should flag this anomaly.

7. **Quantified risk**: The Q2 documentation gap for CP\u2192CD shipments puts
   ~\u20ac{q2_undoc:,} of zero-rated supplies at risk. If DE tax
   authorities deny zero-rating, 19% VAT = ~\u20ac{q2_vat_risk:,} exposure.

## Data Challenges

- **ERR-EU-010**: The error is subtle \u2014 it looks like a normal invoice line
  but violates Art. 196 reverse charge. Agents that do not cross-reference
  the VAT rules reference will miss it.
- **Three missing data traps**: Q2 dispatch docs, Polish registration,
  Italian sales \u2014 each requires different analysis.
- **VAT PE vs tax PE**: A conceptual distinction that many models conflate.
- **CE Q4 anomaly**: A filing status change that could indicate a VAT audit.
- **Multiple VAT regimes**: NL, DE, FR, UK \u2014 each with different rates
  and procedures.
- **No factor-based allocation**: VAT is a transaction tax. Apportioning
  income across jurisdictions (as in income tax) is the wrong methodology.

## Expected Output Structure

### Transaction Analysis (Excel):
- One row per IC transaction with VAT treatment assessment
- Flag column for errors/issues found
- Corrected treatment column where applicable

### VAT Registration Gap Analysis:
- Current registrations vs. required registrations
- Risk assessment for each gap

### Quarterly Reconciliation (Excel):
- Reconciliation of VAT returns against transaction data
- Anomaly flags (CE Q4 pending assessment)

### Risk Summary (Word):
- Quantified exposures (ERR-EU-010: \u20ac{err_vat:,}, Q2 dispatch: ~\u20ac{q2_vat_risk:,})
- Priority ranking of remediation actions
- Timeline recommendations
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── 7. Gold Standard ────────────────────────────────────────────────────────


@register_gold(_TC)
def _tc10_eu_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """TC-10-EU gold standard: VAT and cross-border tax position analysis."""
    q2_vat_exposure = _whole_euros(CP_TO_CD_Q2_UNDOCUMENTED * VAT_RATES["DE"])

    return GoldStandard(
        test_case=_TC,
        expected_outputs={
            "output_files": {
                "vat_analysis_workbook": {
                    "type": "xlsx",
                    "required_sheets": [
                        "Transaction-by-Transaction VAT Analysis",
                        "VAT Registration Assessment",
                        "Quarterly VAT Reconciliation",
                        "PE Risk Assessment",
                        "Risk Summary and Recommendations",
                    ],
                },
            },
            "transaction_analysis": {
                "ic_transaction_count": len(generate_ic_sales_eu()),
                "flows_analyzed": [
                    {
                        "flow": "CP\u2192CM raw materials",
                        "vat_treatment": "Zero-rated intra-EU supply (Art. 138)",
                        "conditions": [
                            "Valid VAT ID of acquirer (CM: FR12345678901)",
                            "Goods physically transported DE\u2192FR",
                            "Proof of dispatch on file",
                        ],
                        "annual_amount_eur": _whole_euros(CP_TO_CM_RAW_MATERIALS),
                        "status": "Compliant",
                    },
                    {
                        "flow": "CP\u2192CD finished goods",
                        "vat_treatment": "Zero-rated export to third country (post-Brexit)",
                        "conditions": [
                            "Customs export declaration",
                            "Proof of dispatch/transport",
                        ],
                        "annual_amount_eur": _whole_euros(CP_TO_CD_FINISHED_GOODS),
                        "status": "Partial compliance \u2014 Q2 documentation gap",
                        "q2_issue": {
                            "total_q2_eur": _whole_euros(CP_TO_CD_Q2_TOTAL),
                            "undocumented_eur": _whole_euros(CP_TO_CD_Q2_UNDOCUMENTED),
                            "shipments_missing_docs": 2,
                            "shipments_total": 5,
                            "vat_exposure_eur": q2_vat_exposure,
                        },
                    },
                    {
                        "flow": "CE\u2192CP management fees",
                        "vat_treatment": "Reverse charge (Art. 196, B2B services)",
                        "annual_amount_eur": _whole_euros(CE_TO_CP_MGMT_FEE),
                        "status": "Compliant",
                    },
                    {
                        "flow": "CE\u2192CM management fees",
                        "vat_treatment": "Should be reverse charge (Art. 196)",
                        "annual_amount_eur": _whole_euros(CE_TO_CM_MGMT_FEE),
                        "status": "ERR-EU-010 in Q3",
                        "error": {
                            "quarter": ERR_EU_010_QUARTER,
                            "amount_eur": _whole_euros(ERR_EU_010_AMOUNT),
                            "vat_incorrectly_charged_eur": _whole_euros(
                                ERR_EU_010_VAT_CHARGED,
                            ),
                            "issue": (
                                "CE charged 20% French VAT \u2014 CE is not registered "
                                "in FR; reverse charge under Art. 196 should apply"
                            ),
                            "impact": (
                                f"CM cannot validly deduct \u20ac{_whole_euros(ERR_EU_010_VAT_CHARGED):,} "
                                "input VAT \u2014 invalid invoice"
                            ),
                        },
                    },
                    {
                        "flow": "CE\u2192CD management fees",
                        "vat_treatment": "Outside scope of EU VAT (UK reverse charge)",
                        "annual_amount_eur": _whole_euros(CE_TO_CD_MGMT_FEE),
                        "status": "Compliant",
                    },
                    {
                        "flow": "CM\u2192CP R&D royalty",
                        "vat_treatment": "Reverse charge (Art. 196, B2B services)",
                        "annual_amount_eur": _whole_euros(CM_TO_CP_ROYALTY),
                        "status": "Compliant",
                    },
                ],
            },
            "vat_registrations": {
                "current_registrations": [
                    {"entity": "CE", "country": "NL", "vat_id": VAT_IDS["CE"], "status": "Active"},
                    {"entity": "CP", "country": "DE", "vat_id": VAT_IDS["CP"], "status": "Active"},
                    {"entity": "CP", "country": "PL", "vat_id": "PL9876543210", "status": "Pending"},
                    {"entity": "CM", "country": "FR", "vat_id": VAT_IDS["CM"], "status": "Active"},
                    {"entity": "CD", "country": "UK", "vat_id": VAT_IDS["CD"], "status": "Active"},
                ],
                "gaps_identified": [
                    {
                        "entity": "CM",
                        "missing_country": "Italy",
                        "reason": (
                            f"CM has \u20ac{_whole_euros(CM_ITALIAN_SALES):,} in sales to Italian "
                            "customers \u2014 may require IT VAT registration depending on "
                            "the nature of supplies (goods with installation, local supplies)"
                        ),
                    },
                    {
                        "entity": "CP",
                        "pending_country": "Poland",
                        "reason": (
                            "CP has a pending Polish VAT registration (applied 2025-03-15). "
                            "Any supplies in PL before completion create compliance risk."
                        ),
                    },
                ],
            },
            "pe_risk_assessment": {
                "vat_pe_vs_income_tax_pe": (
                    "A VAT fixed establishment requires sufficient human and technical "
                    "resources to make or receive supplies independently. This test is "
                    "different from and independent of the income tax PE test under "
                    "Art. 5 OECD Model Convention."
                ),
                "entities_assessed": {
                    "cp_in_fr": "Low \u2014 goods shipped, not installed",
                    "cm_in_de": "Low \u2014 IP licensed, no personnel in DE",
                    "ce_in_subs": "Medium \u2014 depends on travel patterns (data gap)",
                    "cd_in_eu": "Low \u2014 UK-only operations",
                },
            },
            "quarterly_reconciliation": {
                "entities": ["CE", "CP", "CM", "CD"],
                "quarters": ["Q1", "Q2", "Q3", "Q4"],
                "anomalies": [
                    {
                        "entity": "CE",
                        "quarter": "Q4",
                        "issue": (
                            "Filing status is 'Filed \u2014 Pending Assessment' "
                            "(all other quarters show 'Filed \u2014 Assessed')"
                        ),
                        "implication": (
                            "May indicate a VAT audit or query from "
                            "Dutch tax authorities"
                        ),
                    },
                ],
            },
            "risk_summary": {
                "total_identified_risks": 4,
                "risks": [
                    {
                        "risk_id": "ERR-EU-010",
                        "category": "Incorrect VAT treatment",
                        "exposure_eur": _whole_euros(ERR_EU_010_VAT_CHARGED),
                        "severity": "Material",
                        "description": (
                            "CE charged French VAT on CE\u2192CM Q3 management fees; "
                            "reverse charge should apply"
                        ),
                    },
                    {
                        "risk_id": "Q2-DISPATCH-GAP",
                        "category": "Documentation gap",
                        "exposure_eur": q2_vat_exposure,
                        "severity": "Material",
                        "description": (
                            f"2 of 5 CP\u2192CD Q2 shipments (\u20ac{_whole_euros(CP_TO_CD_Q2_UNDOCUMENTED):,}) "
                            "lack proof of dispatch \u2014 zero-rating at risk"
                        ),
                    },
                    {
                        "risk_id": "CM-IT-REGISTRATION",
                        "category": "Missing VAT registration",
                        "exposure_eur": _whole_euros(CM_ITALIAN_SALES),
                        "severity": "Moderate",
                        "description": (
                            f"CM has \u20ac{_whole_euros(CM_ITALIAN_SALES):,} Italian sales "
                            "without Italian VAT registration"
                        ),
                    },
                    {
                        "risk_id": "CE-Q4-PENDING",
                        "category": "Filing anomaly",
                        "exposure_eur": 0,
                        "severity": "Low \u2014 monitoring",
                        "description": (
                            "CE Q4 return pending assessment \u2014 "
                            "possible audit trigger"
                        ),
                    },
                ],
            },
        },
        canary_verification={
            "read_ic_sales": canaries.canary_for("tc10eu_intercompany_sales_fy2025"),
            "read_vat_registrations": canaries.canary_for("tc10eu_vat_registrations"),
            "read_vat_returns": canaries.canary_for("tc10eu_vat_returns_summary_fy2025"),
            "read_vat_rules_reference": canaries.canary_for("tc10eu_eu_vat_rules_reference"),
        },
        error_detection={
            "ERR-EU-010": (
                "CE (Netherlands) invoiced CM (France) for Q3 management fees "
                f"(\u20ac{_whole_euros(ERR_EU_010_AMOUNT):,}) with 20% French VAT "
                f"(\u20ac{_whole_euros(ERR_EU_010_VAT_CHARGED):,}). CE is not registered for "
                "VAT in France \u2014 reverse charge under Art. 196 should apply. "
                "CM\u2019s input VAT deduction is invalid."
            ),
        },
        scoring_hints={
            "correctness": (
                "Each IC transaction must be assessed against correct VAT treatment. "
                "ERR-EU-010 must be identified \u2014 CE cannot charge French VAT without "
                f"FR registration. Q2 dispatch gap quantified at ~\u20ac{q2_vat_exposure:,} risk. "
                "UK transactions correctly treated as third-country exports."
            ),
            "completeness": (
                "All 6 IC flows analyzed with correct VAT articles cited. "
                "3 missing data gaps identified (Q2 dispatch, PL registration, IT sales). "
                "VAT PE vs income tax PE distinction addressed. "
                "CE Q4 filing anomaly flagged. "
                "Third-party sales reviewed for registration requirements."
            ),
            "format_compliance": (
                "Transaction analysis as Excel with per-row assessment. "
                "Registration gap analysis as summary table. "
                "Quarterly reconciliation as Excel. "
                "Risk summary as Word with quantified exposures."
            ),
            "robustness": (
                "Agent must detect ERR-EU-010 by cross-referencing Art. 196 rules. "
                "Agent must distinguish VAT PE from income tax PE. "
                "Agent must quantify Q2 documentation risk. "
                "Agent must flag CE Q4 pending assessment. "
                "Agent must identify CM Italian sales registration gap."
            ),
            "communication": (
                "Clear per-transaction VAT assessment with article citations. "
                "Professional EU VAT terminology (intra-Community supply, reverse charge). "
                "Actionable remediation recommendations with priority ranking."
            ),
        },
        scenario_pack="cascade_europe_ifrs",
        judgment_traps=[
            {
                "trap": "ERR-EU-010 \u2014 VAT treatment error",
                "description": (
                    "CE charges 20% French VAT on CE\u2192CM Q3 management fees. "
                    "CE is not registered in France; reverse charge under Art. 196 "
                    "should apply. An agent that does not cross-reference the VAT "
                    "rules reference or check CE\u2019s registration status will miss this."
                ),
            },
            {
                "trap": "Q2 proof of dispatch gap",
                "description": (
                    "CP\u2192CD Q2 finished goods show \u2018Partial\u2019 proof of dispatch. "
                    "2 of 5 shipments lack documentation. Without proof, German tax "
                    "authorities can deny zero-rating and assess 19% VAT."
                ),
            },
            {
                "trap": "CM Italian sales without IT registration",
                "description": (
                    f"CM has \u20ac{_whole_euros(CM_ITALIAN_SALES):,} in sales to Italian "
                    "customers but no Italian VAT registration. Agent must assess "
                    "whether an IT registration is required based on the nature of "
                    "the supplies."
                ),
            },
            {
                "trap": "VAT PE vs income tax PE conflation",
                "description": (
                    "The VAT rules reference explicitly states these are different "
                    "tests. An agent that treats them as equivalent fails to "
                    "demonstrate proper EU VAT knowledge."
                ),
            },
            {
                "trap": "CE Q4 pending assessment anomaly",
                "description": (
                    "All CE returns except Q4 show \u2018Filed \u2014 Assessed\u2019. The Q4 "
                    "\u2018Pending Assessment\u2019 status may indicate a VAT audit or query. "
                    "An agent that ignores filing status misses a compliance signal."
                ),
            },
        ],
    )


# ── Public entry point ───────────────────────────────────────────────────────


def emit_tc10_eu(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
    **kwargs: object,
) -> None:
    """Emit all TC-10-EU files."""
    _write_ic_sales_xlsx(output_dir, canaries, errors, manifest)
    _write_vat_registrations_xlsx(output_dir, canaries, manifest)
    _write_vat_returns_xlsx(output_dir, canaries, manifest)
    _write_vat_rules_docx(output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
