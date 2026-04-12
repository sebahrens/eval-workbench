"""Formatter: TC-17 — Multi-File Deliverable Assembly (Cross-service, Complex).

Emits:
- test_cases/TC-17/input_files/01_executive_summary.docx
- test_cases/TC-17/input_files/02_financial_analysis.xlsx
- test_cases/TC-17/input_files/03_industry_overview.docx
- test_cases/TC-17/input_files/04_risk_assessment.docx
- test_cases/TC-17/input_files/05_detailed_findings.xlsx
- test_cases/TC-17/input_files/06_recommendations.docx
- test_cases/TC-17/input_files/cover_page_template.docx  (copied from templates/)
- test_cases/TC-17/input_files/formatting_guide.pdf      (copied from templates/)
- test_cases/TC-17/prompt.md
- test_cases/TC-17/expected_behavior.md
- gold_standards/TC-17_gold.json

No planted errors — difficulty is in assembly, not data correctness.
Uses the canonical model — never hardcodes numbers.
"""

from __future__ import annotations

import datetime
import io
import shutil
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import openpyxl
from docx import Document
from docx.shared import Pt
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.writer.excel import ExcelWriter

from generator.canaries import (
    CanaryRegistry,
    embed_canary_docx,
    embed_canary_xlsx,
)
from generator.errors import ErrorRegistry
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel
from generator.model.consolidation import (
    build_balance_sheet,
    build_income_statement,
)

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-17"
_INPUT_DIR = f"test_cases/{_TC}/input_files"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)
_FIXED_ZIP_DT = (2025, 3, 15, 9, 0, 0)

_CANARY_KEYS: list[str] = sorted([
    "tc17_executive_summary",
    "tc17_financial_analysis",
    "tc17_industry_overview",
    "tc17_risk_assessment",
    "tc17_detailed_findings",
    "tc17_recommendations",
])

# Styling constants
_HEADER_FILL = PatternFill(start_color="1A3C6E", end_color="1A3C6E", fill_type="solid")
_HEADER_FONT = Font(bold=True, color="FFFFFF", size=10)
_BOLD_FONT = Font(bold=True, size=10)
_NORMAL_FONT = Font(size=10)
_THIN_BORDER = Border(
    bottom=Side(style="thin"),
    top=Side(style="thin"),
    left=Side(style="thin"),
    right=Side(style="thin"),
)
_ALT_FILL = PatternFill(start_color="F2F6FA", end_color="F2F6FA", fill_type="solid")


# ── Deterministic save helpers ───────────────────────────────────────────────


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    path = Path(path)
    wb.properties.modified = _FIXED_DATETIME

    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _whole_dollars(d: Decimal | int) -> int:
    """Round a Decimal to whole dollars."""
    if isinstance(d, int):
        return d
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


def _fmt_dollars(amount: int) -> str:
    """Format an integer as $X,XXX,XXX."""
    return f"${amount:,.0f}"


def _fmt_pct(value: float) -> str:
    """Format a float as X.X%."""
    return f"{value:.1f}%"


# ── Docx helpers ─────────────────────────────────────────────────────────────


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a styled heading to a docx."""
    doc.add_heading(text, level=level)


def _add_body(doc: Document, text: str) -> None:
    """Add a body paragraph to a docx."""
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)


def _add_bullet(doc: Document, text: str) -> None:
    """Add a bullet-point paragraph."""
    doc.add_paragraph(text, style="List Bullet")


# ── xlsx table helper ────────────────────────────────────────────────────────


def _write_table_header(ws: Any, row: int, headers: list[str]) -> None:
    """Write a formatted header row to a worksheet."""
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col_idx, value=header)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = _THIN_BORDER


def _write_data_row(
    ws: Any, row: int, values: list[Any], *, alt: bool = False,
) -> None:
    """Write a data row with optional alternating shading."""
    for col_idx, val in enumerate(values, 1):
        cell = ws.cell(row=row, column=col_idx, value=val)
        cell.font = _NORMAL_FONT
        cell.border = _THIN_BORDER
        if alt:
            cell.fill = _ALT_FILL
        if isinstance(val, (int, float)):
            cell.alignment = Alignment(horizontal="right")
            cell.number_format = "#,##0"


# ── Section 1: Executive Summary (docx) ─────────────────────────────────────


def _write_executive_summary(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Create 01_executive_summary.docx with Cascade financial highlights."""
    is_2025 = build_income_statement(model.ledger, 2025)
    is_2024 = build_income_statement(model.ledger, 2024)
    bs = build_balance_sheet(model.ledger, datetime.date(2025, 12, 31))

    revenue_growth = float(
        (is_2025.total_revenue - is_2024.total_revenue) / is_2024.total_revenue * 100
    ) if is_2024.total_revenue else 0
    gross_margin = float(is_2025.gross_profit / is_2025.total_revenue * 100) if is_2025.total_revenue else 0

    doc = Document()
    canary_code = canaries.canary_for("tc17_executive_summary")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "tc17_executive_summary",
        f"{_INPUT_DIR}/01_executive_summary.docx",
        location,
    )

    _add_heading(doc, "Executive Summary", level=1)

    _add_body(doc, (
        "This report presents the findings and recommendations from our comprehensive "
        "financial advisory engagement with Cascade Industries, Inc. for the fiscal year "
        "ended December 31, 2025."
    ))

    _add_heading(doc, "Key Financial Highlights", level=2)
    _add_bullet(doc, f"Consolidated revenue: {_fmt_dollars(_whole_dollars(is_2025.total_revenue))}")
    _add_bullet(doc, f"Revenue growth (YoY): {_fmt_pct(revenue_growth)}")
    _add_bullet(doc, f"Gross profit margin: {_fmt_pct(gross_margin)}")
    _add_bullet(doc, f"Net income: {_fmt_dollars(_whole_dollars(is_2025.net_income))}")
    _add_bullet(doc, f"Total assets: {_fmt_dollars(_whole_dollars(bs.total_assets))}")

    _add_heading(doc, "Engagement Scope", level=2)
    _add_body(doc, (
        "Our engagement covered the following areas: financial statement analysis, "
        "industry benchmarking, risk assessment, and strategic recommendations. "
        "The analysis encompasses all three operating subsidiaries: Cascade Precision "
        "Components LLC, Cascade Advanced Materials, Inc., and Cascade Distribution "
        "Services LLC."
    ))

    _add_heading(doc, "Summary of Findings", level=2)
    _add_body(doc, (
        "Cascade Industries demonstrated solid revenue growth driven primarily by the "
        "Advanced Materials segment. The company's diversified business model provides "
        "resilience, though concentration risk in certain product lines warrants attention. "
        "Detailed findings and actionable recommendations are presented in the subsequent "
        "sections of this report."
    ))

    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME

    dest = output_dir / _INPUT_DIR / "01_executive_summary.docx"
    dest.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, dest)

    manifest.register(
        f"{_INPUT_DIR}/01_executive_summary.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Section 2: Financial Analysis (xlsx) ─────────────────────────────────────


def _write_financial_analysis(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Create 02_financial_analysis.xlsx with income statement and key ratios."""
    is_2025 = build_income_statement(model.ledger, 2025)
    is_2024 = build_income_statement(model.ledger, 2024)

    wb = openpyxl.Workbook()
    wb.properties.created = _FIXED_DATETIME
    canary_code = canaries.canary_for("tc17_financial_analysis")
    location = embed_canary_xlsx(wb, canary_code)
    canaries.set_location(
        "tc17_financial_analysis",
        f"{_INPUT_DIR}/02_financial_analysis.xlsx",
        location,
    )

    # ── Sheet 1: Income Statement Summary ──
    ws = wb.active
    ws.title = "Income Statement"
    ws.sheet_properties.tabColor = "1A3C6E"

    headers = ["Line Item", "FY2025", "FY2024", "Change ($)", "Change (%)"]
    _write_table_header(ws, 1, headers)

    rows_data = [
        ("Revenue", is_2025.total_revenue, is_2024.total_revenue),
        ("Cost of Goods Sold", is_2025.total_cogs, is_2024.total_cogs),
        ("Gross Profit", is_2025.gross_profit, is_2024.gross_profit),
        ("Operating Expenses", is_2025.total_opex, is_2024.total_opex),
        ("Operating Income", is_2025.operating_income, is_2024.operating_income),
        ("Pre-Tax Income", is_2025.pre_tax_income, is_2024.pre_tax_income),
        ("Net Income", is_2025.net_income, is_2024.net_income),
    ]

    for i, (label, fy25, fy24) in enumerate(rows_data, 2):
        fy25_int = _whole_dollars(fy25)
        fy24_int = _whole_dollars(fy24)
        change = fy25_int - fy24_int
        pct = float(change / fy24_int * 100) if fy24_int else 0
        _write_data_row(ws, i, [label, fy25_int, fy24_int, change, round(pct, 1)], alt=(i % 2 == 0))

    for col in range(1, 6):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 22

    # ── Sheet 2: Key Ratios ──
    ws2 = wb.create_sheet("Key Ratios")
    ws2.sheet_properties.tabColor = "1A3C6E"
    _write_table_header(ws2, 1, ["Ratio", "FY2025", "FY2024"])

    gross_margin_25 = float(is_2025.gross_profit / is_2025.total_revenue * 100) if is_2025.total_revenue else 0
    gross_margin_24 = float(is_2024.gross_profit / is_2024.total_revenue * 100) if is_2024.total_revenue else 0
    op_margin_25 = float(is_2025.operating_income / is_2025.total_revenue * 100) if is_2025.total_revenue else 0
    op_margin_24 = float(is_2024.operating_income / is_2024.total_revenue * 100) if is_2024.total_revenue else 0
    net_margin_25 = float(is_2025.net_income / is_2025.total_revenue * 100) if is_2025.total_revenue else 0
    net_margin_24 = float(is_2024.net_income / is_2024.total_revenue * 100) if is_2024.total_revenue else 0

    ratios = [
        ("Gross Margin (%)", round(gross_margin_25, 1), round(gross_margin_24, 1)),
        ("Operating Margin (%)", round(op_margin_25, 1), round(op_margin_24, 1)),
        ("Net Margin (%)", round(net_margin_25, 1), round(net_margin_24, 1)),
        (
            "Revenue Growth (%)",
            round(float((is_2025.total_revenue - is_2024.total_revenue)
                        / is_2024.total_revenue * 100), 1) if is_2024.total_revenue else 0,
            "N/A",
        ),
    ]
    for i, (label, fy25, fy24) in enumerate(ratios, 2):
        _write_data_row(ws2, i, [label, fy25, fy24], alt=(i % 2 == 0))

    for col in range(1, 4):
        ws2.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 25

    dest = output_dir / _INPUT_DIR / "02_financial_analysis.xlsx"
    dest.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, dest)

    manifest.register(
        f"{_INPUT_DIR}/02_financial_analysis.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Section 3: Industry Overview (docx) ──────────────────────────────────────


def _write_industry_overview(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Create 03_industry_overview.docx with market context."""
    is_2025 = build_income_statement(model.ledger, 2025)

    doc = Document()
    canary_code = canaries.canary_for("tc17_industry_overview")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "tc17_industry_overview",
        f"{_INPUT_DIR}/03_industry_overview.docx",
        location,
    )

    _add_heading(doc, "Industry Overview", level=1)

    _add_heading(doc, "Market Context", level=2)
    _add_body(doc, (
        "Cascade Industries operates in the mid-market manufacturing sector, "
        "with diversified operations spanning precision components, advanced "
        "materials, and distribution services. The U.S. manufacturing sector "
        "experienced moderate growth in FY2025, driven by reshoring trends "
        "and infrastructure investment."
    ))

    _add_heading(doc, "Peer Comparison", level=2)
    _add_body(doc, (
        f"With consolidated revenue of {_fmt_dollars(_whole_dollars(is_2025.total_revenue))}, "
        "Cascade is positioned in the upper quartile of mid-market manufacturers. "
        "Industry average gross margins for comparable manufacturers range from "
        "28% to 38%, with specialty materials companies typically commanding "
        "higher margins due to proprietary formulations."
    ))

    _add_heading(doc, "Key Industry Trends", level=2)
    _add_bullet(doc, "Supply chain diversification: shift from single-source to multi-source procurement")
    _add_bullet(doc, "Automation investment: increasing capex in robotics and process optimization")
    _add_bullet(doc, "ESG reporting: growing regulatory requirements for sustainability disclosures")
    _add_bullet(doc, "Skilled labor shortage: wage pressure in manufacturing and logistics roles")
    _add_bullet(doc, "Raw material price volatility: commodity hedging strategies gaining importance")

    _add_heading(doc, "Competitive Position", level=2)
    _add_body(doc, (
        "Cascade's three-subsidiary structure provides diversification benefits. "
        "The Advanced Materials segment's higher margins offset the thinner margins "
        "in Distribution Services. The company's vertically integrated supply chain "
        "between Precision Components and Advanced Materials creates a competitive "
        "moat, though it also introduces intercompany complexity."
    ))

    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME

    dest = output_dir / _INPUT_DIR / "03_industry_overview.docx"
    dest.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, dest)

    manifest.register(
        f"{_INPUT_DIR}/03_industry_overview.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Section 4: Risk Assessment (docx) ────────────────────────────────────────


def _write_risk_assessment(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Create 04_risk_assessment.docx with risk matrix."""
    doc = Document()
    canary_code = canaries.canary_for("tc17_risk_assessment")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "tc17_risk_assessment",
        f"{_INPUT_DIR}/04_risk_assessment.docx",
        location,
    )

    _add_heading(doc, "Risk Assessment", level=1)

    _add_heading(doc, "Methodology", level=2)
    _add_body(doc, (
        "Risks were assessed using a likelihood-impact framework on a scale of 1 "
        "(low) to 5 (high). Each risk is categorized by domain and assigned a risk "
        "score (likelihood x impact). Risks scoring 12 or above require immediate "
        "mitigation plans."
    ))

    _add_heading(doc, "Key Risks Identified", level=2)

    # Risk 1
    _add_heading(doc, "1. Customer Concentration Risk", level=3)
    _add_body(doc, "Likelihood: 3 | Impact: 4 | Risk Score: 12")
    _add_body(doc, (
        "Top 10 customers represent a significant portion of consolidated revenue. "
        "Loss of a major customer would materially impact the Precision Components segment."
    ))

    # Risk 2
    _add_heading(doc, "2. Raw Material Price Volatility", level=3)
    _add_body(doc, "Likelihood: 4 | Impact: 3 | Risk Score: 12")
    _add_body(doc, (
        "Advanced Materials relies on specialty inputs with limited supplier alternatives. "
        "Price spikes could compress margins if not passed through to customers."
    ))

    # Risk 3
    _add_heading(doc, "3. Intercompany Transfer Pricing", level=3)
    _add_body(doc, "Likelihood: 2 | Impact: 4 | Risk Score: 8")
    _add_body(doc, (
        "The cost-plus-8% transfer pricing arrangement between Precision Components "
        "and Advanced Materials may attract tax authority scrutiny if not properly documented."
    ))

    # Risk 4
    _add_heading(doc, "4. Key Personnel Dependency", level=3)
    _add_body(doc, "Likelihood: 3 | Impact: 3 | Risk Score: 9")
    _add_body(doc, (
        "Several senior technical roles in Advanced Materials R&D have no identified "
        "successors. An 8% annual turnover rate creates ongoing retention risk."
    ))

    # Risk 5
    _add_heading(doc, "5. Distribution Margin Compression", level=3)
    _add_body(doc, "Likelihood: 3 | Impact: 3 | Risk Score: 9")
    _add_body(doc, (
        "Distribution Services operates on thin margins (~18% gross). Rising freight "
        "and warehousing costs could push the segment toward breakeven."
    ))

    _add_heading(doc, "Risk Heat Map Summary", level=2)
    _add_body(doc, (
        "Two risks score at the critical threshold (12): customer concentration and "
        "raw material volatility. These require active monitoring and mitigation "
        "strategies as detailed in the Recommendations section."
    ))

    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME

    dest = output_dir / _INPUT_DIR / "04_risk_assessment.docx"
    dest.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, dest)

    manifest.register(
        f"{_INPUT_DIR}/04_risk_assessment.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Section 5: Detailed Findings (xlsx) ──────────────────────────────────────


def _write_detailed_findings(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Create 05_detailed_findings.xlsx with segment analysis tables."""
    bs = build_balance_sheet(model.ledger, datetime.date(2025, 12, 31))

    wb = openpyxl.Workbook()
    wb.properties.created = _FIXED_DATETIME
    canary_code = canaries.canary_for("tc17_detailed_findings")
    location = embed_canary_xlsx(wb, canary_code)
    canaries.set_location(
        "tc17_detailed_findings",
        f"{_INPUT_DIR}/05_detailed_findings.xlsx",
        location,
    )

    # ── Sheet 1: Segment Revenue Breakdown ──
    ws = wb.active
    ws.title = "Segment Revenue"
    ws.sheet_properties.tabColor = "1A3C6E"

    headers = ["Subsidiary", "Revenue FY2025", "% of Total", "Gross Margin"]
    _write_table_header(ws, 1, headers)

    # Compute segment revenues from revenue records
    seg_totals: dict[str, Decimal] = {}
    for r in model.revenue_records:
        if r.year == 2025:
            seg_totals[r.entity_code] = seg_totals.get(r.entity_code, Decimal(0)) + r.revenue

    total_rev = sum(seg_totals.values())

    seg_names = {
        "PC": "Precision Components",
        "AM": "Advanced Materials",
        "DS": "Distribution Services",
    }
    seg_margins = {"PC": 0.35, "AM": 0.52, "DS": 0.18}

    row_idx = 2
    for code in sorted(seg_totals.keys()):
        rev = seg_totals[code]
        pct = float(rev / total_rev * 100) if total_rev else 0
        margin = seg_margins.get(code, 0)
        _write_data_row(
            ws, row_idx,
            [seg_names.get(code, code), _whole_dollars(rev), round(pct, 1), round(margin * 100, 1)],
            alt=(row_idx % 2 == 0),
        )
        row_idx += 1

    # Total row
    for col_idx, val in enumerate([
        "Total", _whole_dollars(total_rev), 100.0, "",
    ], 1):
        cell = ws.cell(row=row_idx, column=col_idx, value=val)
        cell.font = _BOLD_FONT
        cell.border = _THIN_BORDER

    for col in range(1, 5):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 22

    # ── Sheet 2: Balance Sheet Summary ──
    ws2 = wb.create_sheet("Balance Sheet")
    ws2.sheet_properties.tabColor = "1A3C6E"

    _write_table_header(ws2, 1, ["Category", "Amount ($)"])

    bs_rows = [
        ("Total Assets", _whole_dollars(bs.total_assets)),
        ("Total Liabilities", _whole_dollars(bs.total_liabilities)),
        ("Total Equity", _whole_dollars(bs.total_equity)),
    ]
    for i, (label, amount) in enumerate(bs_rows, 2):
        _write_data_row(ws2, i, [label, amount], alt=(i % 2 == 0))

    ws2.column_dimensions["A"].width = 25
    ws2.column_dimensions["B"].width = 20

    # ── Sheet 3: Employee Headcount ──
    ws3 = wb.create_sheet("Headcount")
    ws3.sheet_properties.tabColor = "1A3C6E"

    _write_table_header(ws3, 1, ["Subsidiary", "Headcount", "% of Total"])

    entity_counts: dict[str, int] = {}
    for emp in model.employees:
        entity_counts[emp.entity_code] = entity_counts.get(emp.entity_code, 0) + 1

    total_emps = len(model.employees)
    row_idx = 2
    for code in sorted(entity_counts.keys()):
        count = entity_counts[code]
        pct = round(count / total_emps * 100, 1) if total_emps else 0
        _write_data_row(
            ws3, row_idx,
            [seg_names.get(code, code), count, pct],
            alt=(row_idx % 2 == 0),
        )
        row_idx += 1

    for col_idx, val in enumerate(["Total", total_emps, 100.0], 1):
        cell = ws3.cell(row=row_idx, column=col_idx, value=val)
        cell.font = _BOLD_FONT
        cell.border = _THIN_BORDER

    ws3.column_dimensions["A"].width = 25
    ws3.column_dimensions["B"].width = 15
    ws3.column_dimensions["C"].width = 15

    dest = output_dir / _INPUT_DIR / "05_detailed_findings.xlsx"
    dest.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, dest)

    manifest.register(
        f"{_INPUT_DIR}/05_detailed_findings.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Section 6: Recommendations (docx) ────────────────────────────────────────


def _write_recommendations(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Create 06_recommendations.docx with actionable recommendations."""
    doc = Document()
    canary_code = canaries.canary_for("tc17_recommendations")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "tc17_recommendations",
        f"{_INPUT_DIR}/06_recommendations.docx",
        location,
    )

    _add_heading(doc, "Recommendations", level=1)

    _add_body(doc, (
        "Based on our analysis, we recommend the following actions prioritized "
        "by expected impact on Cascade Industries' financial performance and "
        "risk posture."
    ))

    _add_heading(doc, "High Priority", level=2)

    _add_heading(doc, "1. Diversify Customer Base (Precision Components)", level=3)
    _add_body(doc, (
        "Develop a targeted business development plan to reduce concentration "
        "in the top 10 customers. Target: reduce top-10 revenue concentration "
        "below 40% within 18 months through new customer acquisition in adjacent "
        "manufacturing verticals."
    ))

    _add_heading(doc, "2. Commodity Hedging Program (Advanced Materials)", level=3)
    _add_body(doc, (
        "Implement a formal commodity hedging strategy for key raw material "
        "inputs. Engage a derivatives advisor to structure forward contracts "
        "covering 60-80% of projected annual consumption for critical inputs."
    ))

    _add_heading(doc, "Medium Priority", level=2)

    _add_heading(doc, "3. Transfer Pricing Documentation", level=3)
    _add_body(doc, (
        "Commission a contemporaneous transfer pricing study to document the "
        "arm's-length nature of the cost-plus-8% intercompany arrangement. "
        "Update annually as part of the tax compliance cycle."
    ))

    _add_heading(doc, "4. Succession Planning (Advanced Materials R&D)", level=3)
    _add_body(doc, (
        "Identify and develop successors for the 5 most critical technical "
        "roles in the R&D function. Implement retention incentives (deferred "
        "compensation, equity participation) for key personnel."
    ))

    _add_heading(doc, "5. Distribution Services Margin Improvement", level=3)
    _add_body(doc, (
        "Conduct a comprehensive logistics cost review. Evaluate route "
        "optimization, warehouse consolidation, and technology investments "
        "to improve the segment's gross margin from 18% toward the 22-25% "
        "industry benchmark."
    ))

    _add_heading(doc, "Lower Priority", level=2)

    _add_heading(doc, "6. ESG Reporting Framework", level=3)
    _add_body(doc, (
        "Begin developing sustainability reporting capabilities in anticipation "
        "of forthcoming SEC climate disclosure requirements. Conduct a baseline "
        "carbon footprint assessment across all three subsidiaries."
    ))

    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME

    dest = output_dir / _INPUT_DIR / "06_recommendations.docx"
    dest.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, dest)

    manifest.register(
        f"{_INPUT_DIR}/06_recommendations.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Copy templates into TC-17 input_files ────────────────────────────────────


def _copy_cover_page(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Copy deliverable_cover_page.docx from templates/ into TC-17 input_files.

    The cover page template is emitted by templates.py into output_dir/templates/.
    We copy it into TC-17's input_files so the agent sees it as an input.
    """
    src = output_dir / "templates" / "deliverable_cover_page.docx"
    dest = output_dir / _INPUT_DIR / "cover_page_template.docx"
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(str(src), str(dest))

    manifest.register(
        f"{_INPUT_DIR}/cover_page_template.docx",
        "docx",
        canary=canaries.canary_for("cover_page_template"),
        test_cases=[_TC],
    )


def _copy_formatting_guide(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Copy formatting_guide.pdf from templates/ into TC-17 input_files."""
    src = output_dir / "templates" / "formatting_guide.pdf"
    dest = output_dir / _INPUT_DIR / "formatting_guide.pdf"
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(str(src), str(dest))

    manifest.register(
        f"{_INPUT_DIR}/formatting_guide.pdf",
        "pdf",
        canary=canaries.canary_for("formatting_guide"),
        test_cases=[_TC],
    )


# ── Prompt and expected behavior ─────────────────────────────────────────────


def _write_prompt(output_dir: Path) -> None:
    """Write test_cases/TC-17/prompt.md per spec."""
    text = """\
Assemble the 6 workpaper sections into a single client deliverable.

1. Use the cover page template (populate with "Cascade Industries — FY2025
   Financial Advisory Report" and today's date).
2. Follow the section order specified in the formatting guide.
3. Add a table of contents after the cover page.
4. Apply consistent formatting per the formatting guide (fonts, headers, spacing).
5. Add page numbers (Roman numerals for TOC, Arabic starting at 1 for Section 1).
6. For the Excel sections, extract key tables and charts and embed them
   in the appropriate location in the document flow.
7. Export as a single PDF.

The final deliverable should look like something you'd hand to a C-suite client.
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write test_cases/TC-17/expected_behavior.md per spec."""
    text = """\
# TC-17: Multi-File Deliverable Assembly — Expected Behavior

## Assembly Requirements
- All 6 workpaper sections must be assembled in the correct order as specified
  in the formatting guide:
  1. Cover Page
  2. Table of Contents
  3. Executive Summary
  4. Financial Analysis
  5. Industry Overview
  6. Risk Assessment
  7. Detailed Findings
  8. Recommendations

## Cover Page
- The cover page template must be populated with:
  - Report title: "Cascade Industries — FY2025 Financial Advisory Report"
  - Client name: "Cascade Industries, Inc."
  - Date: today's date
  - Prepared by field completed
- The "[COMPANY LOGO]" placeholder may remain or be styled appropriately.

## Table of Contents
- A TOC must be generated listing all sections with page numbers.
- TOC entries must match the actual section headings and page numbers.

## Page Numbering
- Cover page: no page number.
- Table of Contents: Roman numerals (i, ii, ...).
- Body sections (Executive Summary onward): Arabic numerals starting at 1.
- Format: centered in footer per the formatting guide.

## Formatting Consistency
- Fonts must follow the formatting guide specifications:
  - Section headings: Helvetica Bold 16pt, #1A3C6E
  - Body text: Helvetica Regular 11pt, black
  - Table headers: Helvetica Bold 10pt, white on #1A3C6E
- Headers on every page except cover: "Cascade Industries - CONFIDENTIAL"
- Footers with page numbers and date.
- 1-inch margins on all sides.

## Excel Content Embedding
- Tables from 02_financial_analysis.xlsx and 05_detailed_findings.xlsx must be
  embedded as properly formatted tables (not screenshots).
- Income Statement Summary, Key Ratios, Segment Revenue, Balance Sheet Summary,
  and Headcount tables should all appear in the appropriate sections.
- Tables must have captions and source references.

## Output
- Output must be a single PDF file.
- The PDF must open without errors.
- Professional appearance suitable for C-suite presentation.

## Scoring Focus
This is primarily a format and assembly test. Grade on:
- Correct section order
- TOC accuracy
- Consistent formatting
- Page number scheme (Roman/Arabic)
- Professional appearance
- Successful Excel table embedding
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold Standard ────────────────────────────────────────────────────────────


@register_gold("TC-17")
def _tc17_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """Build the TC-17 gold standard."""
    return GoldStandard(
        test_case="TC-17",
        expected_outputs={
            "output_format": "pdf",
            "section_order": [
                "Cover Page",
                "Table of Contents",
                "Executive Summary",
                "Financial Analysis",
                "Industry Overview",
                "Risk Assessment",
                "Detailed Findings",
                "Recommendations",
            ],
            "page_numbering": {
                "cover_page": "none",
                "toc": "roman_numerals",
                "body_sections": "arabic_starting_at_1",
            },
            "required_tables": [
                "Income Statement Summary",
                "Key Ratios",
                "Segment Revenue Breakdown",
                "Balance Sheet Summary",
                "Employee Headcount",
            ],
            "cover_page_fields": {
                "report_title": "Cascade Industries — FY2025 Financial Advisory Report",
                "client_name": "Cascade Industries, Inc.",
                "date": "populated",
            },
            "formatting": {
                "section_headings": "Helvetica Bold 16pt #1A3C6E",
                "body_text": "Helvetica Regular 11pt black",
                "table_headers": "Helvetica Bold 10pt white on #1A3C6E",
                "margins": "1 inch all sides",
                "headers": "Cascade Industries - CONFIDENTIAL",
            },
        },
        canary_verification={
            "read_executive_summary": canaries.canary_for("tc17_executive_summary"),
            "read_financial_analysis": canaries.canary_for("tc17_financial_analysis"),
            "read_industry_overview": canaries.canary_for("tc17_industry_overview"),
            "read_risk_assessment": canaries.canary_for("tc17_risk_assessment"),
            "read_detailed_findings": canaries.canary_for("tc17_detailed_findings"),
            "read_recommendations": canaries.canary_for("tc17_recommendations"),
            "read_cover_template": canaries.canary_for("cover_page_template"),
            "read_formatting_guide": canaries.canary_for("formatting_guide"),
        },
        error_detection={},
        scoring_hints={
            "correctness": "Section order and TOC must match formatting guide spec exactly",
            "completeness": "All 6 sections + cover + TOC assembled; all Excel tables embedded",
            "format_compliance": (
                "Professional PDF output; correct page numbering scheme; "
                "consistent fonts/headers/footers per formatting guide"
            ),
            "communication": (
                "Cover page properly populated; source references for embedded tables"
            ),
        },
    )


# ── Public entry point ──────────────────────────────────────────────────────


def emit_tc17(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
    **kwargs: object,
) -> None:
    """Write all TC-17 files to *output_dir*.

    IMPORTANT: emit_templates() must be called before emit_tc17() so that
    the cover page template and formatting guide PDF exist in output_dir/templates/
    for copying into TC-17 input_files.
    """
    _write_executive_summary(model, output_dir, canaries, manifest)
    _write_financial_analysis(model, output_dir, canaries, manifest)
    _write_industry_overview(model, output_dir, canaries, manifest)
    _write_risk_assessment(model, output_dir, canaries, manifest)
    _write_detailed_findings(model, output_dir, canaries, manifest)
    _write_recommendations(output_dir, canaries, manifest)
    _copy_cover_page(output_dir, canaries, manifest)
    _copy_formatting_guide(output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
