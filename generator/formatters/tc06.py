"""Formatter: TC-06 — Tax Provision under ASC 740 (Tax, Complex).

Emits:
- test_cases/TC-06/input_files/cascade_consolidated_tb_fy2025.xlsx
  Consolidated trial balance with pre-tax book income
- test_cases/TC-06/input_files/tax_provision_fy2024_workpaper.xlsx
  Prior year provision workpaper (current/deferred, rollforward, ETR recon)
- test_cases/TC-06/input_files/permanent_temporary_differences_fy2025.docx
  Table of FY2025 book-tax differences with dollar amounts
- test_cases/TC-06/input_files/statutory_rates.docx
  Federal 21%, blended state 6.2%, apportionment-weighted calculation
- test_cases/TC-06/prompt.md
- test_cases/TC-06/expected_behavior.md
- gold_standards/TC-06_gold.json

Planted errors:
  ERR-006: stale_data — prior year state rate (5.8%) not updated to 6.2%
           in the FY2024 workpaper state rate cell

Uses the canonical model — never hardcodes numbers.
"""

from __future__ import annotations

import datetime
import io
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from generator.canaries import CanaryRegistry, embed_canary_docx, embed_canary_xlsx
from generator.errors import ErrorRegistry, PlantedError, stale_data
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel
from generator.model.views import consolidated_trial_balance

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-06"
_INPUT_DIR = f"test_cases/{_TC}/input_files"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)
_FIXED_ZIP_DT = (2025, 3, 15, 9, 0, 0)


# ── Deterministic save helpers ───────────────────────────────────────────────

def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)
    wb.properties.modified = _FIXED_DATETIME

    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _pin_xlsx_dates(wb: openpyxl.Workbook) -> None:
    wb.properties.created = _FIXED_DATETIME


def _fmt_dollars(d: Decimal | int) -> int:
    """Round to whole dollars."""
    if isinstance(d, int):
        return d
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


# ── Styling constants ────────────────────────────────────────────────────────

_HEADER_FONT = Font(bold=True, size=10, color="FFFFFF")
_HEADER_FILL = PatternFill("solid", fgColor="1A3C6E")
_DATA_FONT = Font(size=10)
_BOLD_FONT = Font(bold=True, size=10)
_MONEY_FMT = "#,##0"
_PCT_FMT = "0.00%"
_THIN_BORDER = Border(
    bottom=Side(style="thin"),
)


def _write_header_row(ws: Any, row: int, headers: list[str]) -> None:
    """Write a styled header row."""
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col, value=header)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = Alignment(horizontal="center")


# ── 1. Consolidated Trial Balance ────────────────────────────────────────────


def _write_consolidated_tb(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write cascade_consolidated_tb_fy2025.xlsx — consolidated TB for tax."""
    from generator.model.coa import ACCOUNTS_BY_NUMBER

    as_of = datetime.date(2025, 12, 31)
    tb = consolidated_trial_balance(model.ledger, as_of)

    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    ws = wb.active
    ws.title = "Trial Balance"

    canary_code = canaries.canary_for("cascade_consolidated_tb_fy2025")
    location = embed_canary_xlsx(wb, canary_code)
    canaries.set_location(
        "cascade_consolidated_tb_fy2025",
        f"{_INPUT_DIR}/cascade_consolidated_tb_fy2025.xlsx",
        location,
    )

    ws["A1"] = "Cascade Industries, Inc."
    ws["A1"].font = Font(bold=True, size=12)
    ws["A2"] = "Consolidated Trial Balance"
    ws["A2"].font = Font(bold=True, size=11)
    ws["A3"] = "As of December 31, 2025"
    ws["A3"].font = Font(italic=True, size=10)

    header_row = 5
    _write_header_row(ws, header_row, [
        "Account", "Description", "Debit", "Credit", "Net Balance",
    ])

    row = header_row + 1
    total_debit = Decimal(0)
    total_credit = Decimal(0)

    for acct_num in sorted(tb.keys()):
        balance = tb[acct_num]
        acct_info = ACCOUNTS_BY_NUMBER.get(acct_num)
        desc = acct_info.name if acct_info else acct_num

        debit = _fmt_dollars(balance) if balance > 0 else 0
        credit = _fmt_dollars(abs(balance)) if balance < 0 else 0

        ws.cell(row=row, column=1, value=acct_num).font = _DATA_FONT
        ws.cell(row=row, column=2, value=desc).font = _DATA_FONT
        c3 = ws.cell(row=row, column=3, value=debit if debit else "")
        c3.font = _DATA_FONT
        c3.number_format = _MONEY_FMT
        c4 = ws.cell(row=row, column=4, value=credit if credit else "")
        c4.font = _DATA_FONT
        c4.number_format = _MONEY_FMT
        c5 = ws.cell(row=row, column=5, value=_fmt_dollars(balance))
        c5.font = _DATA_FONT
        c5.number_format = _MONEY_FMT

        total_debit += max(balance, Decimal(0))
        total_credit += abs(min(balance, Decimal(0)))
        row += 1

    # Totals row
    ws.cell(row=row, column=2, value="TOTAL").font = _BOLD_FONT
    c3t = ws.cell(row=row, column=3, value=_fmt_dollars(total_debit))
    c3t.font = _BOLD_FONT
    c3t.number_format = _MONEY_FMT
    c4t = ws.cell(row=row, column=4, value=_fmt_dollars(total_credit))
    c4t.font = _BOLD_FONT
    c4t.number_format = _MONEY_FMT

    # Column widths
    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 40
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 18
    ws.column_dimensions["E"].width = 18

    path = output_dir / _INPUT_DIR / "cascade_consolidated_tb_fy2025.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, path)

    manifest.register(
        f"{_INPUT_DIR}/cascade_consolidated_tb_fy2025.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── 2. Prior Year Provision Workpaper ────────────────────────────────────────

# ERR-006: stale state rate — FY2024 workpaper shows 5.8% instead of 6.2%.
_STALE_STATE_RATE = Decimal("0.058")


def _write_prior_year_workpaper(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write tax_provision_fy2024_workpaper.xlsx — prior year provision."""
    prov24 = model.tax_provisions[2024]

    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)

    canary_code = canaries.canary_for("tax_provision_fy2024_workpaper")
    location = embed_canary_xlsx(wb, canary_code)
    canaries.set_location(
        "tax_provision_fy2024_workpaper",
        f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx",
        location,
    )

    # ── Sheet 1: Current & Deferred Provision ────────────────────────
    ws1 = wb.active
    ws1.title = "Current & Deferred"

    ws1["A1"] = "Cascade Industries, Inc."
    ws1["A1"].font = Font(bold=True, size=12)
    ws1["A2"] = "Income Tax Provision — FY2024"
    ws1["A2"].font = Font(bold=True, size=11)

    row = 4
    ws1.cell(row=row, column=1, value="Pre-tax book income").font = _BOLD_FONT
    ws1.cell(row=row, column=2, value=_fmt_dollars(prov24.pre_tax_book_income)).font = _BOLD_FONT
    ws1.cell(row=row, column=2).number_format = _MONEY_FMT
    row += 2

    # Permanent differences
    ws1.cell(row=row, column=1, value="Permanent Differences").font = _BOLD_FONT
    row += 1
    for pd in prov24.permanent_differences:
        ws1.cell(row=row, column=1, value=f"  {pd.description}").font = _DATA_FONT
        c = ws1.cell(row=row, column=2, value=_fmt_dollars(pd.amount))
        c.font = _DATA_FONT
        c.number_format = _MONEY_FMT
        row += 1
    ws1.cell(row=row, column=1, value="Total permanent differences").font = _BOLD_FONT
    ws1.cell(row=row, column=2, value=_fmt_dollars(prov24.total_permanent)).font = _BOLD_FONT
    ws1.cell(row=row, column=2).number_format = _MONEY_FMT
    row += 2

    # Temporary differences
    ws1.cell(row=row, column=1, value="Temporary Differences (Current Year Change)").font = _BOLD_FONT
    row += 1
    for td in prov24.temporary_differences:
        ws1.cell(row=row, column=1, value=f"  {td.description}").font = _DATA_FONT
        c = ws1.cell(row=row, column=2, value=_fmt_dollars(td.difference))
        c.font = _DATA_FONT
        c.number_format = _MONEY_FMT
        row += 1
    ws1.cell(row=row, column=1, value="Total temporary differences").font = _BOLD_FONT
    ws1.cell(row=row, column=2, value=_fmt_dollars(prov24.total_temporary_change)).font = _BOLD_FONT
    ws1.cell(row=row, column=2).number_format = _MONEY_FMT
    row += 2

    # Taxable income
    ws1.cell(row=row, column=1, value="Taxable income").font = _BOLD_FONT
    ws1.cell(row=row, column=2, value=_fmt_dollars(prov24.taxable_income)).font = _BOLD_FONT
    ws1.cell(row=row, column=2).number_format = _MONEY_FMT
    row += 2

    # Current provision
    ws1.cell(row=row, column=1, value="Current Provision").font = _BOLD_FONT
    row += 1
    ws1.cell(row=row, column=1, value="  Federal (21%)").font = _DATA_FONT
    ws1.cell(row=row, column=2, value=_fmt_dollars(prov24.federal_current)).font = _DATA_FONT
    ws1.cell(row=row, column=2).number_format = _MONEY_FMT
    row += 1

    # ERR-006: Stale state rate — show 5.8% instead of 6.2%
    stale_rate_str = stale_data("5.8%")  # Should be "6.2%"
    err_006 = PlantedError(
        error_id="ERR-006",
        file=f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx",
        location="Sheet 'Current & Deferred', state tax rate label",
        type="stale_data",
        description=(
            f"State tax rate shows {stale_rate_str} (FY2023 rate) instead of "
            "6.2% (FY2024 blended rate). The dollar amount is correct per the "
            "6.2% rate but the displayed rate percentage was not updated."
        ),
        severity="immaterial",
        which_test_cases_should_catch=[_TC],
    )
    errors.add(err_006)

    ws1.cell(row=row, column=1, value=f"  State ({stale_rate_str})").font = _DATA_FONT
    ws1.cell(row=row, column=2, value=_fmt_dollars(prov24.state_current)).font = _DATA_FONT
    ws1.cell(row=row, column=2).number_format = _MONEY_FMT
    row += 1
    ws1.cell(row=row, column=1, value="  R&D credit (Section 41)").font = _DATA_FONT
    c = ws1.cell(row=row, column=2, value=-_fmt_dollars(prov24.rd_credit))
    c.font = _DATA_FONT
    c.number_format = _MONEY_FMT
    row += 1
    ws1.cell(row=row, column=1, value="Total current provision").font = _BOLD_FONT
    ws1.cell(row=row, column=2, value=_fmt_dollars(prov24.total_current)).font = _BOLD_FONT
    ws1.cell(row=row, column=2).number_format = _MONEY_FMT
    row += 2

    # Deferred provision
    ws1.cell(row=row, column=1, value="Deferred Provision").font = _BOLD_FONT
    row += 1
    ws1.cell(row=row, column=1, value="Total deferred provision").font = _BOLD_FONT
    ws1.cell(row=row, column=2, value=_fmt_dollars(prov24.total_deferred)).font = _BOLD_FONT
    ws1.cell(row=row, column=2).number_format = _MONEY_FMT
    row += 2

    # Total provision
    ws1.cell(row=row, column=1, value="TOTAL PROVISION").font = Font(bold=True, size=11)
    ws1.cell(row=row, column=2, value=_fmt_dollars(prov24.total_provision)).font = Font(bold=True, size=11)
    ws1.cell(row=row, column=2).number_format = _MONEY_FMT
    row += 1
    ws1.cell(row=row, column=1, value="Effective tax rate").font = _BOLD_FONT
    ws1.cell(row=row, column=2, value=float(prov24.effective_tax_rate)).font = _BOLD_FONT
    ws1.cell(row=row, column=2).number_format = _PCT_FMT

    ws1.column_dimensions["A"].width = 45
    ws1.column_dimensions["B"].width = 20

    # ── Sheet 2: Deferred Tax Rollforward ────────────────────────────
    ws2 = wb.create_sheet("DTA-DTL Rollforward")

    ws2["A1"] = "Deferred Tax Asset / Liability Rollforward — FY2024"
    ws2["A1"].font = Font(bold=True, size=11)

    header_row = 3
    _write_header_row(ws2, header_row, [
        "Description", "Type", "Beginning", "Change", "Ending",
    ])

    row = header_row + 1
    prov23 = model.tax_provisions[2023]
    for i, item in enumerate(prov24.deferred_items):
        prior_item = prov23.deferred_items[i] if i < len(prov23.deferred_items) else None
        beg = _fmt_dollars(prior_item.deferred_tax) if prior_item else 0
        end = _fmt_dollars(item.deferred_tax)
        change = end - beg

        ws2.cell(row=row, column=1, value=item.description).font = _DATA_FONT
        ws2.cell(row=row, column=2, value=item.item_type).font = _DATA_FONT
        for col, val in [(3, beg), (4, change), (5, end)]:
            c = ws2.cell(row=row, column=col, value=val)
            c.font = _DATA_FONT
            c.number_format = _MONEY_FMT
        row += 1

    # Totals
    ws2.cell(row=row, column=1, value="Net DTA").font = _BOLD_FONT
    ws2.cell(row=row, column=5, value=_fmt_dollars(
        prov24.current_dta_total
    )).font = _BOLD_FONT
    ws2.cell(row=row, column=5).number_format = _MONEY_FMT
    row += 1
    ws2.cell(row=row, column=1, value="Net DTL").font = _BOLD_FONT
    ws2.cell(row=row, column=5, value=_fmt_dollars(
        prov24.current_dtl_total
    )).font = _BOLD_FONT
    ws2.cell(row=row, column=5).number_format = _MONEY_FMT

    ws2.column_dimensions["A"].width = 40
    ws2.column_dimensions["B"].width = 8
    ws2.column_dimensions["C"].width = 16
    ws2.column_dimensions["D"].width = 16
    ws2.column_dimensions["E"].width = 16

    # ── Sheet 3: ETR Reconciliation ──────────────────────────────────
    ws3 = wb.create_sheet("Rate Reconciliation")

    ws3["A1"] = "Effective Tax Rate Reconciliation — FY2024"
    ws3["A1"].font = Font(bold=True, size=11)

    header_row = 3
    _write_header_row(ws3, header_row, [
        "Description", "Amount", "Rate Impact",
    ])

    row = header_row + 1
    for desc, amt, rate in prov24.rate_reconciliation:
        ws3.cell(row=row, column=1, value=desc).font = _DATA_FONT
        c2 = ws3.cell(row=row, column=2, value=_fmt_dollars(amt))
        c2.font = _DATA_FONT
        c2.number_format = _MONEY_FMT
        c3 = ws3.cell(row=row, column=3, value=float(rate))
        c3.font = _DATA_FONT
        c3.number_format = _PCT_FMT
        row += 1

    # Total
    ws3.cell(row=row, column=1, value="Effective tax rate").font = _BOLD_FONT
    ws3.cell(row=row, column=2, value=_fmt_dollars(
        prov24.total_provision
    )).font = _BOLD_FONT
    ws3.cell(row=row, column=2).number_format = _MONEY_FMT
    ws3.cell(row=row, column=3, value=float(
        prov24.effective_tax_rate
    )).font = _BOLD_FONT
    ws3.cell(row=row, column=3).number_format = _PCT_FMT

    ws3.column_dimensions["A"].width = 45
    ws3.column_dimensions["B"].width = 18
    ws3.column_dimensions["C"].width = 14

    # ── Sheet 4: Perm/Temp Differences List ──────────────────────────
    ws4 = wb.create_sheet("Differences Detail")

    ws4["A1"] = "Permanent and Temporary Differences — FY2024"
    ws4["A1"].font = Font(bold=True, size=11)

    header_row = 3
    _write_header_row(ws4, header_row, [
        "Type", "Description", "Amount",
    ])

    row = header_row + 1
    for pd in prov24.permanent_differences:
        ws4.cell(row=row, column=1, value="Permanent").font = _DATA_FONT
        ws4.cell(row=row, column=2, value=pd.description).font = _DATA_FONT
        c = ws4.cell(row=row, column=3, value=_fmt_dollars(pd.amount))
        c.font = _DATA_FONT
        c.number_format = _MONEY_FMT
        row += 1

    for td in prov24.temporary_differences:
        ws4.cell(row=row, column=1, value="Temporary").font = _DATA_FONT
        ws4.cell(row=row, column=2, value=td.description).font = _DATA_FONT
        c = ws4.cell(row=row, column=3, value=_fmt_dollars(td.difference))
        c.font = _DATA_FONT
        c.number_format = _MONEY_FMT
        row += 1

    ws4.column_dimensions["A"].width = 14
    ws4.column_dimensions["B"].width = 45
    ws4.column_dimensions["C"].width = 18

    # ── Save ─────────────────────────────────────────────────────────
    path = output_dir / _INPUT_DIR / "tax_provision_fy2024_workpaper.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, path)

    manifest.register(
        f"{_INPUT_DIR}/tax_provision_fy2024_workpaper.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── 3. Permanent & Temporary Differences Document ───────────────────────────


def _write_perm_temp_differences(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write permanent_temporary_differences_fy2025.docx."""
    from docx import Document
    prov25 = model.tax_provisions[2025]
    doc = Document()

    canary_code = canaries.canary_for("perm_temp_differences_fy2025")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "perm_temp_differences_fy2025",
        f"{_INPUT_DIR}/permanent_temporary_differences_fy2025.docx",
        location,
    )

    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME

    # Title
    doc.add_heading("Cascade Industries, Inc.", level=1)
    doc.add_heading("Book-Tax Differences — FY2025", level=2)
    doc.add_paragraph("")

    # Permanent differences table
    doc.add_heading("Permanent Differences", level=3)
    doc.add_paragraph(
        "The following items create permanent differences between book income "
        "and taxable income. These differences do not reverse in future periods."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Description"
    hdr[1].text = "Amount"
    hdr[2].text = "Effect on Taxable Income"

    for pd in prov25.permanent_differences:
        row = table.add_row().cells
        row[0].text = pd.description
        row[1].text = f"${abs(_fmt_dollars(pd.amount)):,}"
        row[2].text = "Increases" if pd.amount > 0 else "Decreases"

    total_row = table.add_row().cells
    total_row[0].text = "Total Permanent Differences"
    total_perm = _fmt_dollars(prov25.total_permanent)
    total_row[1].text = f"${abs(total_perm):,}"
    total_row[2].text = "Net increase" if total_perm > 0 else "Net decrease"

    doc.add_paragraph("")

    # Temporary differences table
    doc.add_heading("Temporary Differences", level=3)
    doc.add_paragraph(
        "The following items create temporary differences that will reverse "
        "in future periods, giving rise to deferred tax assets or liabilities."
    )

    table2 = doc.add_table(rows=1, cols=5)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Description"
    hdr2[1].text = "Book Amount"
    hdr2[2].text = "Tax Amount"
    hdr2[3].text = "Difference"
    hdr2[4].text = "Creates"

    for td in prov25.temporary_differences:
        row = table2.add_row().cells
        row[0].text = td.description
        row[1].text = f"${_fmt_dollars(td.book_amount):,}"
        row[2].text = f"${_fmt_dollars(td.tax_amount):,}"
        diff = _fmt_dollars(td.difference)
        row[3].text = f"${abs(diff):,}" if diff >= 0 else f"(${abs(diff):,})"
        row[4].text = "DTL" if td.difference > 0 else "DTA"

    total_row2 = table2.add_row().cells
    total_row2[0].text = "Total Temporary Differences"
    total_temp = _fmt_dollars(prov25.total_temporary_change)
    total_row2[3].text = f"${abs(total_temp):,}" if total_temp >= 0 else f"(${abs(total_temp):,})"

    path = output_dir / _INPUT_DIR / "permanent_temporary_differences_fy2025.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, path)

    manifest.register(
        f"{_INPUT_DIR}/permanent_temporary_differences_fy2025.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── 4. Statutory Rates Document ─────────────────────────────────────────────


def _write_statutory_rates(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write statutory_rates.docx — federal 21%, blended state 6.2%."""
    from docx import Document

    doc = Document()

    canary_code = canaries.canary_for("statutory_rates")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "statutory_rates",
        f"{_INPUT_DIR}/statutory_rates.docx",
        location,
    )

    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME

    doc.add_heading("Cascade Industries, Inc.", level=1)
    doc.add_heading("Statutory Tax Rates — FY2025", level=2)
    doc.add_paragraph("")

    doc.add_heading("Federal Corporate Income Tax Rate", level=3)
    doc.add_paragraph(
        "The federal corporate income tax rate for FY2025 is 21%, as established "
        "by the Tax Cuts and Jobs Act of 2017 (IRC §11(b))."
    )

    doc.add_heading("Blended State Income Tax Rate", level=3)
    doc.add_paragraph(
        "Cascade Industries operates in multiple states through its subsidiaries. "
        "The blended state income tax rate is computed using an apportionment-weighted "
        "average of each state's corporate tax rate:"
    )

    # State rate table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "State"
    hdr[1].text = "Entity"
    hdr[2].text = "State Rate"
    hdr[3].text = "Apportionment %"

    state_data = [
        ("Oregon", "Precision Components (PC)", "6.6%", "47.5%"),
        ("Texas", "Advanced Materials (AM)", "0.0%", "32.5%"),
        ("Illinois", "Distribution Services (DS)", "9.5%", "20.0%"),
    ]
    for state, entity, rate, apportion in state_data:
        row = table.add_row().cells
        row[0].text = state
        row[1].text = entity
        row[2].text = rate
        row[3].text = apportion

    doc.add_paragraph("")
    doc.add_paragraph(
        "Blended state rate calculation:\n"
        "  Oregon:   6.6% × 47.5% = 3.135%\n"
        "  Texas:    0.0% × 32.5% = 0.000%\n"
        "  Illinois: 9.5% × 20.0% = 1.900%\n"
        "  Other (remote workers in CA, WA, NY): ~1.165%\n"
        "  ─────────────────────────────\n"
        "  Blended state rate: 6.2%"
    )

    doc.add_paragraph("")
    doc.add_heading("Combined Effective Statutory Rate", level=3)
    doc.add_paragraph(
        "For deferred tax calculation purposes, the combined rate is:\n"
        "  Federal: 21.0%\n"
        "  State (net of federal benefit): 6.2% × (1 − 21%) = 4.898%\n"
        "  Combined rate: 25.898%\n\n"
        "Note: For current provision, federal and state are computed separately, "
        "with state tax deductible for federal purposes."
    )

    path = output_dir / _INPUT_DIR / "statutory_rates.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, path)

    manifest.register(
        f"{_INPUT_DIR}/statutory_rates.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Prompt & Expected Behavior ──────────────────────────────────────────────


def _write_prompt(output_dir: Path) -> None:
    """Write test_cases/TC-06/prompt.md per spec."""
    text = """\
Compute the income tax provision for Cascade Industries for FY2025 under ASC 740.

1. Calculate pre-tax book income from the trial balance.
2. Compute taxable income by applying all permanent and temporary differences.
3. Calculate the current tax provision (federal and state).
4. Calculate the deferred tax provision by computing the change in deferred tax
   assets and liabilities from prior year.
5. Roll forward the deferred tax balance sheet from FY2024 to FY2025.
6. Prepare the effective tax rate reconciliation (statutory rate to effective rate).
7. Identify the total provision (current + deferred) and the effective tax rate.

Export as an Excel workbook with separate sheets for:
- Current Provision
- Deferred Rollforward
- Rate Reconciliation
- Summary

Verify that everything ties: current + deferred = total provision, and the
rate reconciliation explains the difference between statutory and effective rates.
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write test_cases/TC-06/expected_behavior.md per spec."""
    text = """\
# TC-06: Tax Provision (ASC 740) — Expected Behavior

## Current Provision
- The agent should correctly extract pre-tax book income from the consolidated
  trial balance by summing all revenue, COGS, operating expense, and other
  income/expense accounts.
- Permanent differences must be correctly applied: meals & entertainment (50%
  non-deductible), tax-exempt municipal bond interest, stock compensation
  excess, fines & penalties.
- Temporary differences must be identified and quantified: depreciation
  (MACRS vs. straight-line), ASC 842 lease adjustments, warranty reserve,
  inventory obsolescence reserve, accrued bonuses, bad debt reserve.
- Federal current provision at 21% and state current provision at 6.2%.
- State tax should be treated as deductible for federal tax purposes.
- R&D tax credit (Section 41) should reduce the current provision.

## Deferred Tax Rollforward
- The agent should roll forward DTA/DTL balances from the FY2024 workpaper
  to FY2025 based on the temporary difference changes.
- Each deferred item should be correctly classified as DTA or DTL.
- The net deferred provision should equal the change in (DTL − DTA) from
  FY2024 to FY2025.

## Effective Tax Rate Reconciliation
- Start from the federal statutory rate (21%).
- Bridge to the effective rate through: state taxes (net of federal benefit),
  permanent differences, R&D credit, and deferred tax provision.
- The reconciliation must tie to the total provision.

## Output Quality
- Excel workbook with four sheets: Current Provision, Deferred Rollforward,
  Rate Reconciliation, Summary.
- All numbers must tie: current + deferred = total provision.
- The rate reconciliation must explain the full difference between statutory
  and effective rates.

## Error Detection
- The agent should notice that the FY2024 workpaper displays the state tax
  rate as 5.8% rather than 6.2%. While the dollar amount is correct, the
  displayed rate percentage is stale from a prior year.
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold Standard ────────────────────────────────────────────────────────────


@register_gold("TC-06")
def _tc06_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """Build the TC-06 gold standard from the canonical model."""
    model: CascadeModel = model_kwargs["model"]
    prov25 = model.tax_provisions[2025]

    return GoldStandard(
        test_case="TC-06",
        expected_outputs={
            "file_type": "xlsx",
            "required_sheets": [
                "Current Provision",
                "Deferred Rollforward",
                "Rate Reconciliation",
                "Summary",
            ],
            "pre_tax_book_income": _fmt_dollars(prov25.pre_tax_book_income),
            "permanent_differences": {
                pd.description: _fmt_dollars(pd.amount)
                for pd in prov25.permanent_differences
            },
            "total_permanent": _fmt_dollars(prov25.total_permanent),
            "temporary_differences": {
                td.description: {
                    "book": _fmt_dollars(td.book_amount),
                    "tax": _fmt_dollars(td.tax_amount),
                    "difference": _fmt_dollars(td.difference),
                }
                for td in prov25.temporary_differences
            },
            "total_temporary_change": _fmt_dollars(prov25.total_temporary_change),
            "taxable_income": _fmt_dollars(prov25.taxable_income),
            "current_provision": {
                "federal": _fmt_dollars(prov25.federal_current),
                "state": _fmt_dollars(prov25.state_current),
                "rd_credit": _fmt_dollars(prov25.rd_credit),
                "total_current": _fmt_dollars(prov25.total_current),
            },
            "deferred_items": {
                di.description: {
                    "cumulative_difference": _fmt_dollars(di.cumulative_difference),
                    "deferred_tax": _fmt_dollars(di.deferred_tax),
                    "type": di.item_type,
                }
                for di in prov25.deferred_items
            },
            "deferred_provision": {
                "prior_dta": _fmt_dollars(prov25.prior_dta_total),
                "prior_dtl": _fmt_dollars(prov25.prior_dtl_total),
                "current_dta": _fmt_dollars(prov25.current_dta_total),
                "current_dtl": _fmt_dollars(prov25.current_dtl_total),
                "total_deferred": _fmt_dollars(prov25.total_deferred),
            },
            "total_provision": _fmt_dollars(prov25.total_provision),
            "effective_tax_rate": float(prov25.effective_tax_rate),
            "rate_reconciliation": {
                desc: {
                    "amount": _fmt_dollars(amt),
                    "rate": float(rate),
                }
                for desc, amt, rate in prov25.rate_reconciliation
            },
            "tie_out_checks": {
                "current_plus_deferred_equals_total": (
                    _fmt_dollars(prov25.total_current) +
                    _fmt_dollars(prov25.total_deferred) ==
                    _fmt_dollars(prov25.total_provision)
                ),
            },
        },
        canary_verification={
            "read_consolidated_tb": canaries.canary_for("cascade_consolidated_tb_fy2025"),
            "read_prior_workpaper": canaries.canary_for("tax_provision_fy2024_workpaper"),
            "read_differences": canaries.canary_for("perm_temp_differences_fy2025"),
            "read_statutory_rates": canaries.canary_for("statutory_rates"),
        },
        error_detection={
            "ERR-006": (
                "FY2024 workpaper shows state tax rate as 5.8% instead of 6.2% "
                "(stale data from prior year, dollar amount is correct)"
            ),
        },
        scoring_hints={
            "correctness": (
                "Pre-tax income, taxable income, current/deferred provision, "
                "and ETR must match gold standard values"
            ),
            "completeness": (
                "All 4 sheets present; all permanent and temporary differences "
                "identified; rate reconciliation ties"
            ),
            "format_compliance": "Valid xlsx with separate named sheets",
            "communication": (
                "Flagged stale state rate in prior workpaper (ERR-006); "
                "verified tie-outs between current + deferred = total"
            ),
        },
    )


# ── Public entry point ──────────────────────────────────────────────────────


def emit_tc06(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write all TC-06 files to *output_dir*."""
    _write_consolidated_tb(model, output_dir, canaries, manifest)
    _write_prior_year_workpaper(model, output_dir, canaries, errors, manifest)
    _write_perm_temp_differences(model, output_dir, canaries, manifest)
    _write_statutory_rates(output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
