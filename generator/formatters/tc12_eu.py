"""Formatter: TC-12-EU — European Diligence Data Room Variant.

Emits a 36-file European data room across 8 categories plus a 72-item
DD checklist for the potential acquisition of Cascade Europe Holdings B.V.:

  test_cases/TC-12-EU/input_files/
  ├── data_room/
  │   ├── 01_corporate/   (5 PDFs + 3 entity-specific PDFs)
  │   ├── 02_financial/   (2 PDFs + 3 XLSX)
  │   ├── 03_legal/       (3 contract PDFs + 1 DOCX + 2 PDFs)
  │   ├── 04_hr/          (1 XLSX + 1 PDF + 3 key-employee PDFs + 1 XLSX)
  │   ├── 05_tax/         (4 entity CIT PDFs + 1 VAT XLSX + 1 notices PDF)
  │   ├── 06_operations/  (1 PDF + 3 XLSX)
  │   ├── 07_technology/  (1 PDF + 1 XLSX + 1 DOCX)
  │   └── 08_compliance/  (1 XLSX + 1 PDF)
  └── dd_checklist_european.docx

Planted errors:
- ERR-EU-012 (transposed_digits) in group_vat_returns_summary.xlsx —
  CP Q3 output VAT €847,200 should be €874,200

Gold standard red flags (8):
1. Munich labor court claim (€180K exposure)
2. Managing director CoC severance (2.5× compensation)
3. Renault CoC termination clause
4. Missing GDPR DPAs (2 of 4 processors)
5. IP assignment gaps (2 French researchers)
6. Expiring EPO patents (2 within 18 months)
7. Open Betriebsprüfung for CP (FY2022-2023)
8. Works council consultation requirement (§111-113 BetrVG)

Missing document categories (6):
- Environmental impact assessments (CP Munich, German Immissionsschutz)
- Transfer pricing documentation (Master File/Local File)
- UK post-Brexit customs/trade compliance
- French CSE consultation records (CM)
- Insurance claim history
- Real property surveys
"""

from __future__ import annotations

import datetime
import io
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import docx
import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from generator.canaries import (
    CanaryRegistry,
    embed_canary_docx,
    embed_canary_xlsx,
)
from generator.errors import ErrorRegistry, PlantedError
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel

# ── Canary keys (public — used by pack and tests) ───────────────────────────

ALL_CANARY_KEYS_TC12EU: list[str] = sorted([
    "tc12eu_benefits_summary_by_country",
    "tc12eu_cd_companies_house_filing",
    "tc12eu_cd_ct600_fy2024",
    "tc12eu_ce_articles_of_association",
    "tc12eu_ce_board_minutes_2024",
    "tc12eu_ce_board_minutes_2025",
    "tc12eu_ce_kvk_extract",
    "tc12eu_ce_vpb_aangifte_fy2024",
    "tc12eu_cm_kbis_extract",
    "tc12eu_cm_liasse_fiscale_fy2024",
    "tc12eu_cp_handelsregister_auszug",
    "tc12eu_cp_korperschaftsteuer_fy2023",
    "tc12eu_cp_korperschaftsteuer_fy2024",
    "tc12eu_customer_agreement_autohaus",
    "tc12eu_customer_agreement_renault",
    "tc12eu_customer_list_with_revenue",
    "tc12eu_dd_checklist_european",
    "tc12eu_equipment_list",
    "tc12eu_facility_leases",
    "tc12eu_finance_director_agreement_ce",
    "tc12eu_gdpr_data_processing_register",
    "tc12eu_group_budget_fy2025",
    "tc12eu_group_debt_schedule",
    "tc12eu_group_employee_census",
    "tc12eu_group_ifrs_financials_fy2023",
    "tc12eu_group_ifrs_financials_fy2024",
    "tc12eu_group_management_accounts_ytd",
    "tc12eu_group_org_chart",
    "tc12eu_group_org_chart_detailed",
    "tc12eu_group_vat_returns_summary",
    "tc12eu_insurance_policies_summary",
    "tc12eu_ip_assignment_agreements",
    "tc12eu_it_infrastructure_overview",
    "tc12eu_managing_director_agreement_ce",
    "tc12eu_patent_portfolio",
    "tc12eu_pending_litigation_summary",
    "tc12eu_software_licenses",
    "tc12eu_supplier_agreement_thyssenkrupp",
    "tc12eu_tax_notices_and_assessments",
    "tc12eu_technical_director_agreement_cm",
    "tc12eu_vendor_list",
    "tc12eu_works_council_agreements_cp",
])

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-12-EU"
_INPUT_DIR = f"test_cases/{_TC}/input_files"
_DR = f"{_INPUT_DIR}/data_room"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)
_FIXED_ZIP_DT = (2025, 3, 15, 9, 0, 0)

# ── European entity data ─────────────────────────────────────────────────────

_ENTITIES = {
    "CE": {
        "name": "Cascade Europe Holdings B.V.",
        "jurisdiction": "Netherlands",
        "legal_form": "Besloten vennootschap",
        "registration": "KvK 12345678",
        "city": "Amsterdam",
        "role": "European holding company",
        "employees": 25,
    },
    "CP": {
        "name": "Cascade Präzisionsteile GmbH",
        "jurisdiction": "Germany",
        "legal_form": "Gesellschaft mit beschränkter Haftung",
        "registration": "HRB 98765, AG München",
        "city": "Munich",
        "role": "Licensed manufacturer",
        "employees": 310,
    },
    "CM": {
        "name": "Cascade Matériaux Avancés SAS",
        "jurisdiction": "France",
        "legal_form": "Société par actions simplifiée",
        "registration": "RCS Lyon 987 654 321",
        "city": "Lyon",
        "role": "R&D centre, IP developer",
        "employees": 85,
    },
    "CD": {
        "name": "Cascade Distribution Services Ltd",
        "jurisdiction": "United Kingdom",
        "legal_form": "Private limited company",
        "registration": "Companies House 12345678",
        "city": "Birmingham",
        "role": "Post-Brexit distributor",
        "employees": 45,
    },
}

# ── European patent portfolio ────────────────────────────────────────────────

_EU_PATENTS = [
    {
        "number": "EP 2,345,678",
        "title": "High-Precision Bearing Assembly for Automotive Applications",
        "filing_date": datetime.date(2012, 5, 20),
        "grant_date": datetime.date(2015, 3, 8),
        "expiration_date": datetime.date(2032, 5, 20),
        "applicant": "Cascade Matériaux Avancés SAS",
        "inventors": "Dr. Pierre Lefèvre, Dr. Marie Dupont",
        "designated_states": "DE, FR, NL, UK, IT, ES",
        "status": "Active",
    },
    {
        "number": "EP 2,789,012",
        "title": "Composite Layup Process for Aerospace Structural Components",
        "filing_date": datetime.date(2014, 9, 15),
        "grant_date": datetime.date(2017, 7, 22),
        "expiration_date": datetime.date(2034, 9, 15),
        "applicant": "Cascade Matériaux Avancés SAS",
        "inventors": "Dr. Pierre Lefèvre, Dr. Stefan Braun",
        "designated_states": "DE, FR, NL, UK, IT",
        "status": "Active",
    },
    {
        "number": "EP 3,012,345",
        "title": "Thermal Spray Coating System for Industrial Bearings",
        "filing_date": datetime.date(2015, 11, 30),
        "grant_date": datetime.date(2019, 2, 14),
        "expiration_date": datetime.date(2035, 11, 30),
        "applicant": "Cascade Präzisionsteile GmbH",
        "inventors": "Dr. Stefan Braun, Klaus Schneider",
        "designated_states": "DE, FR, NL, AT, CH",
        "status": "Active",
    },
    {
        # RED FLAG: Expiring within 18 months of FY2025 year-end
        "number": "EP 1,890,456",
        "title": "Surface Treatment Process for Precision Metal Components",
        "filing_date": datetime.date(2007, 8, 12),
        "grant_date": datetime.date(2010, 11, 5),
        "expiration_date": datetime.date(2027, 8, 12),
        "applicant": "Cascade Präzisionsteile GmbH",
        "inventors": "Klaus Schneider, Dr. Stefan Braun",
        "designated_states": "DE, FR, NL, UK",
        "status": "Active — expiring within 18 months",
    },
    {
        # RED FLAG: Expiring within 18 months of FY2025 year-end
        "number": "EP 1,678,901",
        "title": "Advanced Polymer Matrix Reinforcement for Lightweight Structures",
        "filing_date": datetime.date(2006, 4, 3),
        "grant_date": datetime.date(2009, 12, 18),
        "expiration_date": datetime.date(2026, 4, 3),
        "applicant": "Cascade Matériaux Avancés SAS",
        "inventors": "Dr. Pierre Lefèvre",
        "designated_states": "DE, FR, NL, UK, IT, ES",
        "status": "Active — expiring within 18 months",
    },
    {
        "number": "DE 10 2019 123 456",
        "title": "Präzisionsausrichtungsvorrichtung für Mehrachsenbearbeitung",
        "filing_date": datetime.date(2019, 6, 1),
        "grant_date": datetime.date(2022, 8, 20),
        "expiration_date": datetime.date(2039, 6, 1),
        "applicant": "Cascade Präzisionsteile GmbH",
        "inventors": "Klaus Schneider",
        "designated_states": "DE",
        "status": "Active",
    },
]

# IP assignments — deliberately incomplete (missing 2 French researchers)
_EU_IP_ASSIGNMENTS = [
    {
        "employee": "Klaus Schneider",
        "entity": "CP",
        "title": "Senior Manufacturing Engineer",
        "date": datetime.date(2015, 3, 1),
        "jurisdiction": "Germany",
        "scope": "All inventions related to precision metal components and surface treatment",
        "status": "Complete — §§ 5-19 ArbnErfG compliant",
    },
    {
        "employee": "Dr. Stefan Braun",
        "entity": "CP",
        "title": "Head of R&D (Manufacturing)",
        "date": datetime.date(2014, 8, 15),
        "jurisdiction": "Germany",
        "scope": "All inventions related to thermal spray and coating processes",
        "status": "Complete — §§ 5-19 ArbnErfG compliant",
    },
    {
        "employee": "Dr. Marie Dupont",
        "entity": "CM",
        "title": "Senior Researcher (Composites)",
        "date": datetime.date(2014, 10, 1),
        "jurisdiction": "France",
        "scope": "All inventions related to composite materials and layup processes",
        "status": "Complete — déclaration d'invention filed",
    },
    {
        "employee": "Jean-Paul Martin",
        "entity": "CM",
        "title": "Process Engineer",
        "date": datetime.date(2018, 3, 15),
        "jurisdiction": "France",
        "scope": "All inventions related to polymer reinforcement processes",
        "status": "Complete — déclaration d'invention filed",
    },
]

# RED FLAG: Missing IP assignments from 2 founding CM researchers
_MISSING_IP_ASSIGNMENTS_EU = [
    (
        "Dr. Pierre Lefèvre (CM — Chief Scientist / Co-Founder"
        " — named inventor on 3 EP patents, déclaration d'invention missing)"
    ),
    (
        "Dr. Élodie Moreau (CM — Senior Researcher / Founding Team"
        " — named inventor on 1 EP patent, déclaration d'invention missing)"
    ),
]

# ── European contracts ────────────────────────────────────────────────────────

_AUTOHAUS_CONTRACT = {
    "name": "Autohaus Müller GmbH",
    "type": "Customer Agreement — Supply of Precision Components",
    "entity": "CP",
    "jurisdiction": "Germany",
    "effective": datetime.date(2022, 1, 1),
    "expiration": datetime.date(2027, 12, 31),
    "value": Decimal("12_500_000"),
    "terms": "Fixed pricing with annual escalator tied to German PPI; DDP Munich",
    "payment": "Net 30, bank transfer to Deutsche Bank account",
    "governing_law": "German law, Munich courts",
}

_RENAULT_CONTRACT = {
    "name": "Renault S.A. (Tier 2 Supply Agreement)",
    "type": "Customer Agreement — Advanced Materials Supply",
    "entity": "CM",
    "jurisdiction": "France",
    "effective": datetime.date(2023, 4, 1),
    "expiration": datetime.date(2028, 3, 31),
    "value": Decimal("8_200_000"),
    "terms": "Volume-based pricing; DAP Lyon; quality per IATF 16949",
    "payment": "Net 60, SEPA transfer",
    "governing_law": "French law, Lyon Commercial Court",
    # RED FLAG: change-of-control termination clause
    "coc_clause": (
        "Section 14.3: Either party may terminate upon 90 days' written notice "
        "if there is a change of indirect control of the other party. "
        "Renault may also terminate without notice if the change of control "
        "results in a competitor of Renault acquiring control."
    ),
}

_THYSSENKRUPP_CONTRACT = {
    "name": "thyssenkrupp Materials Services GmbH",
    "type": "Supplier Agreement — Raw Materials",
    "entity": "CP",
    "jurisdiction": "Germany",
    "effective": datetime.date(2024, 1, 1),
    "expiration": datetime.date(2026, 12, 31),
    "value": Decimal("15_000_000"),
    "terms": "Quarterly volume commitments; CIF Munich; DIN EN standards",
    "payment": "Net 45, SEPA transfer",
    "governing_law": "German law, Essen courts",
}

# ── Litigation ────────────────────────────────────────────────────────────────

_EU_LITIGATION = {
    "title": "Former CP employee unfair dismissal claim",
    "court": "Arbeitsgericht München (Munich Labor Court)",
    "case_ref": "ArbG München 12 Ca 4567/25",
    "filing_date": datetime.date(2025, 2, 15),
    "plaintiff": "Hans Weber (former Production Manager)",
    "defendant": "Cascade Präzisionsteile GmbH",
    "claim_amount": Decimal("180_000"),
    "description": (
        "Former production manager Hans Weber filed an unfair dismissal claim "
        "(Kündigungsschutzklage) at the Munich Labor Court on 15 February 2025. "
        "Weber was terminated for alleged repeated safety protocol violations. "
        "Weber claims the dismissal was pretextual and related to his works council "
        "candidacy. Under German labor law (KSchG), if the court finds the dismissal "
        "was not socially justified, remedies include reinstatement or severance of up "
        "to 12 monthly salaries (§10 KSchG). Weber's annual salary was €95,000. "
        "Potential exposure estimated at €180,000 including legal costs. "
        "First hearing scheduled for May 2025."
    ),
    "status": "Filed — first hearing pending",
    "accrued": Decimal("45_000"),
}

# French commercial court case for context (not a primary red flag)
_EU_LITIGATION_FR = {
    "title": "Supplier quality dispute — SteelTech Industries",
    "court": "Tribunal de Commerce de Lyon",
    "case_ref": "TC Lyon 2024/1234",
    "filing_date": datetime.date(2024, 9, 10),
    "plaintiff": "Cascade Matériaux Avancés SAS",
    "defendant": "SteelTech Industries SARL",
    "claim_amount": Decimal("320_000"),
    "description": (
        "CM filed a claim for €320,000 against SteelTech for delivery of substandard "
        "raw materials that caused a production line shutdown. SteelTech has "
        "counterclaimed €85,000 for unpaid invoices. Settlement discussions ongoing."
    ),
    "status": "Active — settlement discussions",
    "accrued": Decimal("0"),
}

# ── Key employee data ─────────────────────────────────────────────────────────

_KEY_EMPLOYEES = [
    {
        "title": "Managing Director (Bestuurder)",
        "name": "Erik van der Berg",
        "entity": "CE",
        "jurisdiction": "Netherlands",
        "base_salary": Decimal("280_000"),
        "coc_multiplier": Decimal("2.5"),
        "notice_period": "6 months",
        "non_compete": "12 months post-termination, EU-wide",
        "filename": "managing_director_service_agreement_ce.pdf",
        "canary_key": "tc12eu_managing_director_agreement_ce",
    },
    {
        "title": "Finance Director",
        "name": "Anna de Vries",
        "entity": "CE",
        "jurisdiction": "Netherlands",
        "base_salary": Decimal("220_000"),
        "coc_multiplier": Decimal("2.0"),
        "notice_period": "3 months",
        "non_compete": "12 months post-termination, Netherlands",
        "filename": "finance_director_agreement_ce.pdf",
        "canary_key": "tc12eu_finance_director_agreement_ce",
    },
    {
        "title": "Technical Director (Directeur Technique)",
        "name": "Dr. Pierre Lefèvre",
        "entity": "CM",
        "jurisdiction": "France",
        "base_salary": Decimal("195_000"),
        "coc_multiplier": Decimal("1.5"),
        "notice_period": "3 months (Convention Collective Métallurgie)",
        "non_compete": "18 months post-termination, EU-wide — requires indemnité compensatrice",
        "filename": "technical_director_agreement_cm.pdf",
        "canary_key": "tc12eu_technical_director_agreement_cm",
    },
]

# ── Betriebsprüfung (German tax audit) ────────────────────────────────────────

_BETRIEBSPRUEFUNG = {
    "entity": "CP",
    "entity_name": "Cascade Präzisionsteile GmbH",
    "audit_notice_date": datetime.date(2025, 1, 20),
    "audit_period": "FY2022–FY2023",
    "tax_types": "Körperschaftsteuer, Gewerbesteuer, Umsatzsteuer",
    "auditor": "Finanzamt München III, Betriebsprüfungsstelle",
    "status": "In progress — on-site review commenced March 2025",
    "focus_areas": (
        "Transfer pricing of intercompany transactions (CE→CP management fees, "
        "CM→CP royalties); VAT treatment of cross-border supplies; "
        "Gewerbesteuer Hinzurechnungen (trade tax add-backs) on license payments"
    ),
    "estimated_exposure": "Unquantified — no preliminary findings issued yet",
}

# ── GDPR data ─────────────────────────────────────────────────────────────────

_GDPR_PROCESSORS = [
    {
        "processor": "CloudServe B.V.",
        "service": "Cloud hosting and infrastructure",
        "data_types": "Employee data, customer data, financial records",
        "dpa_status": "Not executed",
        "country": "Netherlands",
    },
    {
        "processor": "PayrollPlus GmbH",
        "service": "External HR/payroll processing",
        "data_types": "Employee personal data, salary data, tax IDs",
        "dpa_status": "Not executed",
        "country": "Germany",
    },
    {
        "processor": "SecureDoc SAS",
        "service": "Document management and archival",
        "data_types": "Corporate documents, contracts",
        "dpa_status": "Executed — dated 2023-06-15",
        "country": "France",
    },
    {
        "processor": "CyberGuard Ltd",
        "service": "IT security monitoring and incident response",
        "data_types": "Network logs, access logs, employee device data",
        "dpa_status": "Executed — dated 2024-01-10",
        "country": "United Kingdom",
    },
]

# ── Works council data ────────────────────────────────────────────────────────

_WORKS_COUNCIL = {
    "entity": "CP",
    "entity_name": "Cascade Präzisionsteile GmbH",
    "established": datetime.date(2018, 5, 1),
    "members": 9,
    "chair": "Wolfgang Richter",
    "agreement_date": datetime.date(2019, 3, 15),
    "agreement_scope": (
        "Working hours, overtime compensation, health & safety committee, "
        "annual bonus framework, employee data protection"
    ),
    "coc_coverage": (
        "The current Betriebsvereinbarung (works council agreement) does NOT "
        "contain provisions addressing change-of-control scenarios, restructuring, "
        "or acquisition-related employee protections. Under §111-113 BetrVG, "
        "CP must inform and consult the works council before any 'fundamental change' "
        "(Betriebsänderung), which includes a change of ownership. Failure to consult "
        "does not invalidate the transaction but may result in claims for a "
        "Sozialplan (social plan / redundancy plan) with significant financial exposure."
    ),
}

# ── Planted error: ERR-EU-012 ─────────────────────────────────────────────────
# CP Q3 output VAT: correct = €874,200, planted = €847,200 (transposed digits)
_ERR_EU_012_CORRECT = Decimal("874200")
_ERR_EU_012_PLANTED = Decimal("847200")

# ── VAT returns summary data (deterministic, TC-12-EU specific) ──────────────
# These are simplified data room values — not the full TC-10-EU VAT model.

_VAT_RETURNS = [
    # CE Netherlands: advisory/holding — small domestic output
    ("CE", "Q1", 52000, 0, 44200, 0, 7800, "Filed — Assessed"),
    ("CE", "Q2", 52000, 0, 44200, 0, 7800, "Filed — Assessed"),
    ("CE", "Q3", 52000, 0, 44200, 0, 7800, "Filed — Assessed"),
    ("CE", "Q4", 52000, 0, 44200, 0, 7800, "Filed — Pending Assessment"),
    # CP Germany: manufacturing — larger domestic output
    ("CP", "Q1", 874200, 0, 742570, 0, 131630, "Filed — Assessed"),
    ("CP", "Q2", 874200, 0, 742570, 0, 131630, "Filed — Assessed"),
    # ERR-EU-012: Q3 output VAT transposed — 847200 instead of 874200
    ("CP", "Q3", 847200, 0, 742570, 0, 104630, "Filed — Assessed"),
    ("CP", "Q4", 874200, 0, 742570, 0, 131630, "Filed — Assessed"),
    # CM France: R&D — moderate domestic
    ("CM", "Q1", 245000, 0, 196000, 24000, 25000, "Filed — Assessed"),
    ("CM", "Q2", 245000, 0, 196000, 24000, 25000, "Filed — Assessed"),
    ("CM", "Q3", 245000, 0, 196000, 24000, 25000, "Filed — Assessed"),
    ("CM", "Q4", 245000, 0, 196000, 24000, 25000, "Filed — Assessed"),
    # CD UK: distribution — GBP converted to EUR at 1.17
    ("CD", "Q1", 175500, 0, 149175, 0, 26325, "Filed — Assessed"),
    ("CD", "Q2", 175500, 0, 149175, 0, 26325, "Filed — Assessed"),
    ("CD", "Q3", 175500, 0, 149175, 0, 26325, "Filed — Assessed"),
    ("CD", "Q4", 175500, 0, 149175, 0, 26325, "Filed — Assessed"),
]

# ── DD Checklist (72 items across 10 categories) ─────────────────────────────

_DD_CHECKLIST_EU = [
    # Corporate / Constitutional (9 items)
    ("Corporate/Constitutional", "Chamber of Commerce / commercial register extracts (all entities)"),
    ("Corporate/Constitutional", "Articles of association / statuten / Satzung (all entities)"),
    ("Corporate/Constitutional", "Board of directors / management board meeting minutes (last 3 years)"),
    ("Corporate/Constitutional", "Shareholder / general meeting minutes (last 3 years)"),
    ("Corporate/Constitutional", "Group organizational chart (legal entity structure with legal forms)"),
    ("Corporate/Constitutional", "List of statutory directors / officers per entity"),
    ("Corporate/Constitutional", "Good standing / certificate of incumbency (all jurisdictions)"),
    ("Corporate/Constitutional", "Capitalization table / shareholder register per entity"),
    ("Corporate/Constitutional", "Powers of attorney and signatory authorities"),
    # Financial / IFRS (9 items)
    ("Financial/IFRS", "IFRS consolidated audited financial statements (last 2 fiscal years)"),
    ("Financial/IFRS", "Management accounts / interim financial statements (current YTD)"),
    ("Financial/IFRS", "Annual operating budget (current year)"),
    ("Financial/IFRS", "Revenue backlog and pipeline reports"),
    ("Financial/IFRS", "Accounts receivable aging schedule"),
    ("Financial/IFRS", "Accounts payable aging schedule"),
    ("Financial/IFRS", "Debt schedule (all outstanding obligations including intercompany)"),
    ("Financial/IFRS", "Capital expenditure plans and forecasts"),
    ("Financial/IFRS", "Working capital analysis"),
    # Tax / VAT (9 items)
    ("Tax/VAT", "Corporate income tax returns per entity (last 2 years)"),
    ("Tax/VAT", "VAT returns summary (all entities, all jurisdictions)"),
    ("Tax/VAT", "Tax notices, assessments, and audit correspondence"),
    ("Tax/VAT", "Transfer pricing documentation (Master File / Local File per OECD)"),
    ("Tax/VAT", "VAT registration certificates (all jurisdictions)"),
    ("Tax/VAT", "Withholding tax compliance records (dividends, royalties, interest)"),
    ("Tax/VAT", "Tax sharing / tax consolidation agreements (intercompany)"),
    ("Tax/VAT", "Gewerbesteuer / trade tax records (German entities)"),
    ("Tax/VAT", "CIR / R&D tax credit documentation (French entity)"),
    # Legal (8 items)
    ("Legal", "Material contracts — customer agreements"),
    ("Legal", "Material contracts — supplier agreements"),
    ("Legal", "Pending or threatened litigation summary (all jurisdictions)"),
    ("Legal", "Settlement agreements (last 5 years)"),
    ("Legal", "Intellectual property assignments and licenses"),
    ("Legal", "Insurance policies and coverage summaries"),
    ("Legal", "Consent decrees or regulatory orders"),
    ("Legal", "Insurance claim history (last 5 years)"),
    # Employment / Labor (9 items)
    ("Employment/Labor", "Employee census / headcount by entity and department"),
    ("Employment/Labor", "Employee benefits summary by country"),
    ("Employment/Labor", "Key employee service agreements / employment contracts"),
    ("Employment/Labor", "Non-compete and non-solicitation agreements"),
    ("Employment/Labor", "Detailed organizational chart (with reporting lines)"),
    ("Employment/Labor", "Works council / Betriebsrat agreements (German entities)"),
    ("Employment/Labor", "Comité social et économique (CSE) records (French entities)"),
    ("Employment/Labor", "Collective bargaining / sector agreements per jurisdiction"),
    ("Employment/Labor", "Social insurance and pension scheme documentation"),
    # Operations (7 items)
    ("Operations", "Facility leases and real property agreements"),
    ("Operations", "Equipment and machinery list (fixed assets register)"),
    ("Operations", "Customer list with revenue by customer"),
    ("Operations", "Vendor list (top suppliers by spend)"),
    ("Operations", "Quality certifications (ISO, IATF 16949, AS9100, etc.)"),
    ("Operations", "Business continuity / disaster recovery plans"),
    ("Operations", "Supply chain concentration analysis"),
    # Technology / IP (6 items)
    ("Technology/IP", "Patent portfolio summary (EPO and national patents)"),
    ("Technology/IP", "Software license inventory"),
    ("Technology/IP", "IT infrastructure overview and architecture"),
    ("Technology/IP", "Cybersecurity audit or assessment reports"),
    ("Technology/IP", "Source code escrow agreements"),
    ("Technology/IP", "Technology roadmap and R&D pipeline"),
    # Environmental (5 items) — ALL MISSING from data room
    ("Environmental", "Environmental impact assessments (all facilities)"),
    ("Environmental", "Environmental permits and compliance records"),
    ("Environmental", "Hazardous waste disposal records"),
    ("Environmental", "Environmental remediation obligations"),
    ("Environmental", "German Immissionsschutz (emissions) permits for CP Munich"),
    # Regulatory / Compliance (5 items)
    ("Regulatory/Compliance", "Government permits and licenses (all jurisdictions)"),
    ("Regulatory/Compliance", "Export control / dual-use goods compliance (EU Regulation 2021/821)"),
    ("Regulatory/Compliance", "UK post-Brexit customs and trade compliance documentation"),
    ("Regulatory/Compliance", "Real property surveys and title reports (all facilities)"),
    ("Regulatory/Compliance", "Antitrust / competition law compliance documentation"),
    # Data Protection / GDPR (5 items)
    ("Data Protection/GDPR", "GDPR Article 30 processing register"),
    ("Data Protection/GDPR", "Data processing agreements (DPAs) with all processors"),
    ("Data Protection/GDPR", "Data protection impact assessments (DPIAs)"),
    ("Data Protection/GDPR", "Data breach notification records"),
    ("Data Protection/GDPR", "Data protection officer (DPO) appointment and reports"),
]

assert len(_DD_CHECKLIST_EU) == 72, f"Expected 72 DD items, got {len(_DD_CHECKLIST_EU)}"


# ── Deterministic save helpers ───────────────────────────────────────────────


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)
    wb.properties.created = _FIXED_DATETIME
    wb.properties.modified = _FIXED_DATETIME
    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _whole_euros(d: Decimal) -> int:
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


# ── Shared reportlab styles ─────────────────────────────────────────────────

def _pdf_styles() -> tuple[
    ParagraphStyle, ParagraphStyle, ParagraphStyle, ParagraphStyle,
]:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle", parent=styles["Title"],
        fontSize=16, spaceAfter=20, alignment=1,
    )
    heading_style = ParagraphStyle(
        "DocHeading", parent=styles["Heading2"],
        fontSize=12, spaceBefore=12, spaceAfter=6,
    )
    subheading_style = ParagraphStyle(
        "DocSubheading", parent=styles["Heading3"],
        fontSize=10, spaceBefore=8, spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "DocBody", parent=styles["Normal"],
        fontSize=10, leading=14, spaceAfter=6,
    )
    return title_style, heading_style, subheading_style, body_style


def _build_simple_pdf(
    full_path: Path,
    canary_code: str,
    title_text: str,
    elements: list,
) -> str:
    """Build a reportlab PDF with canary in author metadata. Returns canary location."""
    doc = SimpleDocTemplate(
        str(full_path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=title_text,
        author=f"CANARY: {canary_code}",
        creator="Cascade Europe Document System",
        invariant=True,
    )
    doc.build(elements)
    return "PDF metadata → Author"


# ── XLSX style constants ────────────────────────────────────────────────────

_HEADER_FONT = Font(bold=True, size=11)
_HEADER_FILL = PatternFill(
    start_color="2F5496", end_color="2F5496", fill_type="solid",
)
_HEADER_FONT_WHITE = Font(bold=True, size=11, color="FFFFFF")
_THIN_BORDER = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin"),
)
_MONEY_FMT = '#,##0'
_DATE_FMT = 'YYYY-MM-DD'


def _style_header_row(ws: Any, row: int, col_count: int) -> None:
    for col in range(1, col_count + 1):
        cell = ws.cell(row=row, column=col)
        cell.font = _HEADER_FONT_WHITE
        cell.fill = _HEADER_FILL
        cell.border = _THIN_BORDER
        cell.alignment = Alignment(horizontal="center")


# ═══════════════════════════════════════════════════════════════════════════
# 01_corporate — 8 PDFs
# ═══════════════════════════════════════════════════════════════════════════


def _write_kvk_extract(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """CE — KvK (Chamber of Commerce) extract."""
    title_s, heading_s, _, body_s = _pdf_styles()
    elements = [
        Paragraph("Kamer van Koophandel — Uittreksel", title_s),
        Spacer(1, 12),
        Paragraph("Handelsregister Extract", heading_s),
        Paragraph(f"KvK number: {_ENTITIES['CE']['registration'].split()[1]}", body_s),
        Paragraph(f"Legal name: {_ENTITIES['CE']['name']}", body_s),
        Paragraph(f"Legal form: {_ENTITIES['CE']['legal_form']} (B.V.)", body_s),
        Paragraph(f"Registered office: {_ENTITIES['CE']['city']}", body_s),
        Paragraph("Date of incorporation: 15 June 2018", body_s),
        Paragraph("Activities: Holding company — management and strategic direction of European subsidiaries", body_s),
        Spacer(1, 12),
        Paragraph("Statutory Directors (Bestuurders)", heading_s),
        Paragraph("Erik van der Berg — Managing Director (since 1 July 2018)", body_s),
        Paragraph("Anna de Vries — Finance Director (since 15 September 2019)", body_s),
        Spacer(1, 12),
        Paragraph("Authorized Signatories", heading_s),
        Paragraph("Erik van der Berg — sole signatory authority up to €500,000; joint authority above", body_s),
        Spacer(1, 12),
        Paragraph("Subsidiaries", heading_s),
        Paragraph(f"100% — {_ENTITIES['CP']['name']} ({_ENTITIES['CP']['jurisdiction']})", body_s),
        Paragraph(f"100% — {_ENTITIES['CM']['name']} ({_ENTITIES['CM']['jurisdiction']})", body_s),
        Paragraph(f"100% — {_ENTITIES['CD']['name']} ({_ENTITIES['CD']['jurisdiction']})", body_s),
        Spacer(1, 12),
        Paragraph("This extract was generated from the KvK Trade Register on 15 March 2025.", body_s),
    ]

    rel = f"{_DR}/01_corporate/ce_kvk_extract.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_ce_kvk_extract")
    loc = _build_simple_pdf(full, canary_code, "KvK Extract — Cascade Europe Holdings B.V.", elements)
    canaries.set_location("tc12eu_ce_kvk_extract", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_articles_of_association(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """CE — Articles of Association (Statuten)."""
    title_s, heading_s, sub_s, body_s = _pdf_styles()
    elements = [
        Paragraph("Statuten — Cascade Europe Holdings B.V.", title_s),
        Spacer(1, 12),
        Paragraph("Articles of Association", heading_s),
        Paragraph("Adopted by notarial deed on 15 June 2018, Amsterdam.", body_s),
        Paragraph("Last amended: 1 March 2022.", body_s),
        Spacer(1, 8),
        Paragraph("Article 1 — Name and Registered Office", sub_s),
        Paragraph(
            "The company is named Cascade Europe Holdings B.V. "
            "and has its registered office in Amsterdam.", body_s,
        ),
        Paragraph("Article 2 — Objects", sub_s),
        Paragraph(
            "To act as holding company, to manage and finance subsidiaries, "
            "and to provide management and advisory services.", body_s,
        ),
        Paragraph("Article 3 — Share Capital", sub_s),
        Paragraph(
            "Authorized capital: 10,000 ordinary shares at €1 nominal. "
            "Issued: 10,000 shares, fully paid.", body_s,
        ),
        Paragraph("Article 4 — Management Board", sub_s),
        Paragraph(
            "The company is managed by one or more managing directors "
            "(bestuurders) appointed by the general meeting.", body_s,
        ),
        Paragraph("Article 5 — Representation", sub_s),
        Paragraph(
            "The company is represented by a managing director acting solely "
            "for transactions up to €500,000; jointly for amounts above.",
            body_s,
        ),
        Paragraph("Article 6 — Financial Year", sub_s),
        Paragraph("The financial year corresponds to the calendar year.", body_s),
        Paragraph("Article 7 — Profit Distribution", sub_s),
        Paragraph(
            "Profits are distributed upon resolution of the general meeting, "
            "subject to applicable reserves.", body_s,
        ),
        Spacer(1, 8),
        Paragraph("Note: Full text runs to 25 articles. Key governance provisions summarized above.", body_s),
    ]

    rel = f"{_DR}/01_corporate/ce_articles_of_association.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_ce_articles_of_association")
    loc = _build_simple_pdf(full, canary_code, "Articles of Association — CE", elements)
    canaries.set_location("tc12eu_ce_articles_of_association", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_handelsregister(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """CP — Handelsregisterauszug."""
    title_s, heading_s, _, body_s = _pdf_styles()
    elements = [
        Paragraph("Handelsregisterauszug", title_s),
        Spacer(1, 12),
        Paragraph("Amtsgericht München — Registergericht", heading_s),
        Paragraph(f"Registernummer: {_ENTITIES['CP']['registration']}", body_s),
        Paragraph(f"Firma: {_ENTITIES['CP']['name']}", body_s),
        Paragraph(f"Rechtsform: {_ENTITIES['CP']['legal_form']} (GmbH)", body_s),
        Paragraph(f"Sitz: {_ENTITIES['CP']['city']}", body_s),
        Paragraph("Stammkapital: €500,000", body_s),
        Paragraph("Gründungsdatum: 1. September 2018", body_s),
        Spacer(1, 8),
        Paragraph("Geschäftsführer", heading_s),
        Paragraph("Thomas Richter — Geschäftsführer (seit 1. Oktober 2018)", body_s),
        Spacer(1, 8),
        Paragraph("Gesellschafter", heading_s),
        Paragraph("Cascade Europe Holdings B.V. — 100%", body_s),
        Spacer(1, 8),
        Paragraph("Unternehmensgegenstand", heading_s),
        Paragraph(
            "Herstellung und Vertrieb von Präzisionsbauteilen "
            "für die Automobil- und Luftfahrtindustrie.", body_s,
        ),
        Spacer(1, 12),
        Paragraph("Dieser Auszug wurde am 15. März 2025 erstellt.", body_s),
    ]

    rel = f"{_DR}/01_corporate/cp_handelsregister_auszug.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_cp_handelsregister_auszug")
    loc = _build_simple_pdf(full, canary_code, "Handelsregisterauszug — CP", elements)
    canaries.set_location("tc12eu_cp_handelsregister_auszug", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_kbis_extract(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """CM — K-bis extract."""
    title_s, heading_s, _, body_s = _pdf_styles()
    elements = [
        Paragraph("Extrait K-bis", title_s),
        Spacer(1, 12),
        Paragraph("Greffe du Tribunal de Commerce de Lyon", heading_s),
        Paragraph(f"SIREN: {_ENTITIES['CM']['registration'].split()[-1]}", body_s),
        Paragraph(f"Dénomination: {_ENTITIES['CM']['name']}", body_s),
        Paragraph(f"Forme juridique: {_ENTITIES['CM']['legal_form']} (SAS)", body_s),
        Paragraph(f"Siège social: {_ENTITIES['CM']['city']}", body_s),
        Paragraph("Capital social: €200,000", body_s),
        Paragraph("Date d'immatriculation: 1er octobre 2018", body_s),
        Spacer(1, 8),
        Paragraph("Président", heading_s),
        Paragraph("Cascade Europe Holdings B.V. (personne morale)", body_s),
        Paragraph("Directeur Général", heading_s),
        Paragraph("Dr. Pierre Lefèvre — Directeur Technique (depuis le 15 octobre 2018)", body_s),
        Spacer(1, 8),
        Paragraph("Activité", heading_s),
        Paragraph(
            "Recherche, développement et fabrication de matériaux avancés "
            "pour l'industrie aéronautique et automobile.", body_s,
        ),
        Spacer(1, 12),
        Paragraph("Cet extrait a été délivré le 15 mars 2025.", body_s),
    ]

    rel = f"{_DR}/01_corporate/cm_kbis_extract.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_cm_kbis_extract")
    loc = _build_simple_pdf(full, canary_code, "Extrait K-bis — CM", elements)
    canaries.set_location("tc12eu_cm_kbis_extract", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_companies_house(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """CD — Companies House filing."""
    title_s, heading_s, _, body_s = _pdf_styles()
    elements = [
        Paragraph("Companies House — Company Overview", title_s),
        Spacer(1, 12),
        Paragraph("Company Details", heading_s),
        Paragraph(f"Company number: {_ENTITIES['CD']['registration'].split()[-1]}", body_s),
        Paragraph(f"Company name: {_ENTITIES['CD']['name']}", body_s),
        Paragraph(f"Company type: {_ENTITIES['CD']['legal_form']}", body_s),
        Paragraph(f"Registered office: {_ENTITIES['CD']['city']}", body_s),
        Paragraph("Incorporated on: 15 November 2018", body_s),
        Paragraph("Company status: Active", body_s),
        Spacer(1, 8),
        Paragraph("Officers", heading_s),
        Paragraph("James Whitfield — Director (appointed 1 December 2018)", body_s),
        Paragraph("Cascade Europe Holdings B.V. — Corporate Director", body_s),
        Spacer(1, 8),
        Paragraph("Persons with Significant Control", heading_s),
        Paragraph("Cascade Europe Holdings B.V. — owns 100% of shares, voting rights 75% or more", body_s),
        Spacer(1, 8),
        Paragraph("SIC Codes", heading_s),
        Paragraph("46720 — Wholesale of metals and metal ores", body_s),
        Paragraph("46690 — Wholesale of other machinery and equipment", body_s),
        Spacer(1, 12),
        Paragraph("This information was last updated on 15 March 2025.", body_s),
    ]

    rel = f"{_DR}/01_corporate/cd_companies_house_filing.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_cd_companies_house_filing")
    loc = _build_simple_pdf(full, canary_code, "Companies House Filing — CD", elements)
    canaries.set_location("tc12eu_cd_companies_house_filing", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_board_minutes_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
    year: int, canary_key: str,
) -> None:
    """CE board minutes for a given year."""
    title_s, heading_s, _, body_s = _pdf_styles()
    elements = [
        Paragraph(f"Board Minutes — Cascade Europe Holdings B.V. — {year}", title_s),
        Spacer(1, 12),
    ]

    if year == 2024:
        elements.extend([
            Paragraph("Meeting of 15 March 2024", heading_s),
            Paragraph("Present: Erik van der Berg (Chair), Anna de Vries", body_s),
            Paragraph("1. Review of FY2023 consolidated IFRS results. Group revenue €98.2M.", body_s),
            Paragraph("2. Approval of dividend distribution of €2.5M to parent.", body_s),
            Paragraph("3. Discussion of CP expansion plans — Munich facility capacity increase.", body_s),
            Paragraph("4. Note: Betriebsprüfung notification received for CP (FY2022-2023). "
                       "Management to cooperate fully with Finanzamt München III.", body_s),
            Paragraph("5. Review of patent portfolio — 2 EPO patents approaching expiration "
                       "(EP 1,890,456 and EP 1,678,901). Management instructed to assess renewal "
                       "strategy and alternative IP protection.", body_s),
            Spacer(1, 8),
            Paragraph("Meeting of 20 September 2024", heading_s),
            Paragraph("Present: Erik van der Berg (Chair), Anna de Vries", body_s),
            Paragraph("1. H1 2024 results review — on track for FY2024 targets.", body_s),
            Paragraph("2. CM R&D pipeline update — 3 new patent applications filed at EPO.", body_s),
            Paragraph("3. CD post-Brexit supply chain review — customs procedures established.", body_s),
        ])
    else:  # 2025
        elements.extend([
            Paragraph("Meeting of 15 January 2025", heading_s),
            Paragraph("Present: Erik van der Berg (Chair), Anna de Vries", body_s),
            Paragraph("1. Approval of FY2024 consolidated IFRS financial statements.", body_s),
            Paragraph("2. Discussion of strategic options including potential sale of group.", body_s),
            Paragraph("3. Approval of managing director service agreement amendments — "
                       "change-of-control severance provision ratified (2.5× annual compensation "
                       "for managing director).", body_s),
            Paragraph("4. Update on CP labor dispute (Weber unfair dismissal claim). "
                       "External counsel engaged. Provision of €45,000 booked.", body_s),
            Paragraph("5. Works council (Betriebsrat) at CP — noted that any change of ownership "
                       "will require consultation under §111 BetrVG.", body_s),
            Spacer(1, 8),
            Paragraph("Meeting of 10 March 2025", heading_s),
            Paragraph("Present: Erik van der Berg (Chair), Anna de Vries", body_s),
            Paragraph("1. Preparation for potential data room opening.", body_s),
            Paragraph("2. GDPR compliance review — noted that DPAs with 2 processors are outstanding.", body_s),
            Paragraph("3. Review of Renault supply agreement — noted change-of-control clause.", body_s),
        ])

    rel = f"{_DR}/01_corporate/ce_board_minutes_{year}.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for(canary_key)
    loc = _build_simple_pdf(full, canary_code, f"Board Minutes {year} — CE", elements)
    canaries.set_location(canary_key, rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_group_org_chart(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """Group org chart showing EU legal structure."""
    title_s, heading_s, _, body_s = _pdf_styles()
    data = [
        ["Entity", "Legal Form", "Jurisdiction", "Ownership", "Role"],
    ]
    for code, ent in _ENTITIES.items():
        own = "Parent" if code == "CE" else "100% owned by CE"
        data.append([ent["name"], ent["legal_form"], ent["jurisdiction"], own, ent["role"]])

    tbl = Table(data, colWidths=[150, 100, 80, 100, 100])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5496")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))

    elements = [
        Paragraph("Cascade Europe Holdings B.V. — Group Structure", title_s),
        Spacer(1, 12),
        tbl,
        Spacer(1, 12),
        Paragraph("Total group headcount: 465 employees across 4 jurisdictions.", body_s),
    ]

    rel = f"{_DR}/01_corporate/group_org_chart.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_group_org_chart")
    loc = _build_simple_pdf(full, canary_code, "Group Org Chart — CE Group", elements)
    canaries.set_location("tc12eu_group_org_chart", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# 02_financial — 2 PDFs + 3 XLSX
# ═══════════════════════════════════════════════════════════════════════════


def _write_ifrs_financials_pdf(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
    year: int, canary_key: str,
) -> None:
    """Group IFRS consolidated financials for a given year."""
    title_s, heading_s, sub_s, body_s = _pdf_styles()

    # Deterministic financial data per year
    if year == 2023:
        revenue, cogs, gp = "98,200,000", "68,740,000", "29,460,000"
        opex, ebit = "18,500,000", "10,960,000"
        finance, pbt, tax, pat = "1,200,000", "9,760,000", "2,635,000", "7,125,000"
        assets, equity, liab = "85,300,000", "42,150,000", "43,150,000"
    else:  # 2024
        revenue, cogs, gp = "105,800,000", "73,560,000", "32,240,000"
        opex, ebit = "19,800,000", "12,440,000"
        finance, pbt, tax, pat = "1,350,000", "11,090,000", "2,994,000", "8,096,000"
        assets, equity, liab = "92,500,000", "48,750,000", "43,750,000"

    is_data = [
        ["", f"FY{year} (€'000)"],
        ["Revenue", revenue],
        ["Cost of sales", f"({cogs})"],
        ["Gross profit", gp],
        ["Operating expenses", f"({opex})"],
        ["EBIT", ebit],
        ["Finance costs", f"({finance})"],
        ["Profit before tax", pbt],
        ["Income tax expense", f"({tax})"],
        ["Profit after tax", pat],
    ]

    is_tbl = Table(is_data, colWidths=[200, 120])
    is_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5496")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("LINEABOVE", (0, 3), (-1, 3), 1, colors.black),
        ("LINEABOVE", (0, 5), (-1, 5), 1, colors.black),
    ]))

    elements = [
        Paragraph(f"Cascade Europe Holdings B.V. — Consolidated Financial Statements FY{year}", title_s),
        Paragraph("Prepared in accordance with International Financial Reporting Standards (IFRS)", body_s),
        Paragraph("Audited by Kramer & Partners, Amsterdam", body_s),
        Spacer(1, 12),
        Paragraph("Consolidated Income Statement", heading_s),
        is_tbl,
        Spacer(1, 12),
        Paragraph("Consolidated Balance Sheet (Summary)", heading_s),
        Paragraph(f"Total assets: €{assets}", body_s),
        Paragraph(f"Total equity: €{equity}", body_s),
        Paragraph(f"Total liabilities: €{liab}", body_s),
        Spacer(1, 12),
        Paragraph("Basis of preparation: Consolidated under IFRS 10. All subsidiaries are 100% owned.", body_s),
        Paragraph("Functional currencies: EUR (CE, CP, CM), GBP (CD — translated at closing rate per IAS 21).", body_s),
    ]

    rel = f"{_DR}/02_financial/group_ifrs_financials_fy{year}.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for(canary_key)
    loc = _build_simple_pdf(full, canary_code, f"IFRS Financials FY{year}", elements)
    canaries.set_location(canary_key, rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_management_accounts_xlsx(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Group management accounts FY2025 YTD."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Management Accounts"

    headers = ["Entity", "Revenue YTD (€)", "COGS YTD (€)", "GP YTD (€)",
               "OpEx YTD (€)", "EBIT YTD (€)", "Budget Var (%)"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    _style_header_row(ws, 1, len(headers))

    rows = [
        ("CE", 2800000, 0, 2800000, 2100000, 700000, -2.1),
        ("CP", 11250000, 7875000, 3375000, 2100000, 1275000, 3.5),
        ("CM", 4500000, 2700000, 1800000, 1350000, 450000, -1.8),
        ("CD", 4050000, 3240000, 810000, 540000, 270000, 5.2),
        ("Eliminations", -1920000, 0, -1920000, 0, -1920000, 0),
    ]
    for r, row in enumerate(rows, 2):
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            if c >= 2 and c <= 6:
                cell.number_format = _MONEY_FMT

    # Total row
    ws.cell(row=7, column=1, value="GROUP TOTAL")
    ws.cell(row=7, column=2, value=20680000).number_format = _MONEY_FMT
    ws.cell(row=7, column=3, value=13815000).number_format = _MONEY_FMT
    ws.cell(row=7, column=4, value=6865000).number_format = _MONEY_FMT
    ws.cell(row=7, column=5, value=6090000).number_format = _MONEY_FMT
    ws.cell(row=7, column=6, value=775000).number_format = _MONEY_FMT

    rel = f"{_DR}/02_financial/group_management_accounts_fy2025_ytd.xlsx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_group_management_accounts_ytd")

    loc = embed_canary_xlsx(wb, canary_code)

    canaries.set_location("tc12eu_group_management_accounts_ytd", rel, loc)
    _save_xlsx_deterministic(wb, full)
    manifest.register(rel, "xlsx", canary=canaries.canary_for("tc12eu_group_management_accounts_ytd"), test_cases=[_TC])


def _write_budget_xlsx(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Group budget FY2025."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Budget FY2025"

    headers = ["Entity", "FY2024 Actual (€)", "FY2025 Budget (€)", "Growth (%)"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    _style_header_row(ws, 1, len(headers))

    rows = [
        ("CE", 11500000, 12200000, 6.1),
        ("CP", 45000000, 48600000, 8.0),
        ("CM", 18000000, 19400000, 7.8),
        ("CD", 16200000, 17800000, 9.9),
        ("Eliminations", -7800000, -8400000, None),
        ("GROUP TOTAL", 82900000, 89600000, 8.1),
    ]
    for r, row in enumerate(rows, 2):
        for c, val in enumerate(row, 1):
            if val is not None:
                cell = ws.cell(row=r, column=c, value=val)
                if c in (2, 3):
                    cell.number_format = _MONEY_FMT

    rel = f"{_DR}/02_financial/group_budget_fy2025.xlsx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_group_budget_fy2025")

    loc = embed_canary_xlsx(wb, canary_code)

    canaries.set_location("tc12eu_group_budget_fy2025", rel, loc)
    _save_xlsx_deterministic(wb, full)
    manifest.register(rel, "xlsx", canary=canaries.canary_for("tc12eu_group_budget_fy2025"), test_cases=[_TC])


def _write_debt_schedule_xlsx(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Group debt schedule including intercompany loan."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Debt Schedule"

    headers = ["Facility", "Borrower", "Lender", "Type", "Outstanding (€)",
               "Interest Rate", "Maturity", "Security"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    _style_header_row(ws, 1, len(headers))

    facilities = [
        ("Revolving Credit Facility", "CE", "ABN AMRO Bank N.V.",
         "Revolver", 8500000, "EURIBOR + 2.25%", "2027-06-30", "Group guarantee"),
        ("Term Loan A", "CP", "Commerzbank AG",
         "Term", 12000000, "3.75% fixed", "2028-12-31", "CP assets"),
        ("Equipment Finance", "CM", "BNP Paribas SA",
         "Asset finance", 3200000, "EURIBOR + 1.50%", "2026-09-30", "CM equipment"),
        ("Intercompany Loan", "CM", "CE",
         "Shareholder loan", 5000000, "4.0% fixed", "2027-12-31", "Unsecured"),
        ("Overdraft Facility", "CD", "Barclays Bank PLC",
         "Overdraft", 1500000, "Bank rate + 3.0%", "Reviewed annually", "CD receivables"),
    ]
    for r, fac in enumerate(facilities, 2):
        for c, val in enumerate(fac, 1):
            cell = ws.cell(row=r, column=c, value=val)
            if c == 5:
                cell.number_format = _MONEY_FMT

    rel = f"{_DR}/02_financial/group_debt_schedule.xlsx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_group_debt_schedule")

    loc = embed_canary_xlsx(wb, canary_code)

    canaries.set_location("tc12eu_group_debt_schedule", rel, loc)
    _save_xlsx_deterministic(wb, full)
    manifest.register(rel, "xlsx", canary=canaries.canary_for("tc12eu_group_debt_schedule"), test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# 03_legal — 3 contract PDFs + 1 DOCX + 2 PDFs
# ═══════════════════════════════════════════════════════════════════════════


def _write_customer_agreement_autohaus(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Customer agreement — Autohaus Müller (DE)."""
    title_s, heading_s, _, body_s = _pdf_styles()
    c = _AUTOHAUS_CONTRACT
    elements = [
        Paragraph(f"Customer Agreement — {c['name']}", title_s),
        Spacer(1, 12),
        Paragraph("Contract Summary", heading_s),
        Paragraph(f"Parties: {c['name']} and {_ENTITIES['CP']['name']}", body_s),
        Paragraph(f"Type: {c['type']}", body_s),
        Paragraph(f"Effective: {c['effective']} to {c['expiration']}", body_s),
        Paragraph(f"Annual value: €{_whole_euros(c['value']):,}", body_s),
        Paragraph(f"Terms: {c['terms']}", body_s),
        Paragraph(f"Payment: {c['payment']}", body_s),
        Paragraph(f"Governing law: {c['governing_law']}", body_s),
        Spacer(1, 8),
        Paragraph("Key Clauses", heading_s),
        Paragraph("Quality standards: DIN EN ISO 9001:2015 and IATF 16949:2016", body_s),
        Paragraph("Warranty: 24 months from delivery", body_s),
        Paragraph("Liability cap: Annual contract value", body_s),
        Paragraph("Termination: 6 months' notice; immediate for material breach", body_s),
        Paragraph("No change-of-control provision.", body_s),
    ]

    rel = f"{_DR}/03_legal/material_contracts/customer_agreement_autohaus_mueller.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_customer_agreement_autohaus")
    loc = _build_simple_pdf(full, canary_code, "Customer Agreement — Autohaus Müller", elements)
    canaries.set_location("tc12eu_customer_agreement_autohaus", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_customer_agreement_renault(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Customer agreement — Renault Tier 2 (FR). Contains CoC clause (RED FLAG)."""
    title_s, heading_s, _, body_s = _pdf_styles()
    c = _RENAULT_CONTRACT
    elements = [
        Paragraph(f"Customer Agreement — {c['name']}", title_s),
        Spacer(1, 12),
        Paragraph("Contract Summary", heading_s),
        Paragraph(f"Parties: Renault S.A. and {_ENTITIES['CM']['name']}", body_s),
        Paragraph(f"Type: {c['type']}", body_s),
        Paragraph(f"Effective: {c['effective']} to {c['expiration']}", body_s),
        Paragraph(f"Annual value: €{_whole_euros(c['value']):,}", body_s),
        Paragraph(f"Terms: {c['terms']}", body_s),
        Paragraph(f"Payment: {c['payment']}", body_s),
        Paragraph(f"Governing law: {c['governing_law']}", body_s),
        Spacer(1, 8),
        Paragraph("Key Clauses", heading_s),
        Paragraph("Quality: IATF 16949:2016 required", body_s),
        Paragraph("Volume commitment: Minimum 500 units/quarter", body_s),
        Paragraph("Exclusivity: CM is sole supplier for specified composite parts", body_s),
        Spacer(1, 8),
        Paragraph("Section 14.3 — Change of Control", heading_s),
        Paragraph(c["coc_clause"], body_s),
    ]

    rel = f"{_DR}/03_legal/material_contracts/customer_agreement_renault_tier2.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_customer_agreement_renault")
    loc = _build_simple_pdf(full, canary_code, "Customer Agreement — Renault", elements)
    canaries.set_location("tc12eu_customer_agreement_renault", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_supplier_agreement_thyssenkrupp(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Supplier agreement — thyssenkrupp Materials (DE)."""
    title_s, heading_s, _, body_s = _pdf_styles()
    c = _THYSSENKRUPP_CONTRACT
    elements = [
        Paragraph(f"Supplier Agreement — {c['name']}", title_s),
        Spacer(1, 12),
        Paragraph("Contract Summary", heading_s),
        Paragraph(f"Parties: {_ENTITIES['CP']['name']} and {c['name']}", body_s),
        Paragraph(f"Type: {c['type']}", body_s),
        Paragraph(f"Effective: {c['effective']} to {c['expiration']}", body_s),
        Paragraph(f"Annual value: €{_whole_euros(c['value']):,}", body_s),
        Paragraph(f"Terms: {c['terms']}", body_s),
        Paragraph(f"Payment: {c['payment']}", body_s),
        Paragraph(f"Governing law: {c['governing_law']}", body_s),
        Spacer(1, 8),
        Paragraph("Key Clauses", heading_s),
        Paragraph("Material specifications: DIN EN 10025, DIN EN 10088", body_s),
        Paragraph("Force majeure: Standard ICC clause", body_s),
        Paragraph("Termination: 3 months' notice; immediate for insolvency or material breach", body_s),
    ]

    rel = f"{_DR}/03_legal/material_contracts/supplier_agreement_thyssenkrupp_materials.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_supplier_agreement_thyssenkrupp")
    loc = _build_simple_pdf(full, canary_code, "Supplier Agreement — thyssenkrupp", elements)
    canaries.set_location("tc12eu_supplier_agreement_thyssenkrupp", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_pending_litigation_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Pending litigation summary — DOCX. Contains German labor court red flag."""
    d = docx.Document()
    d.add_heading("Pending Litigation Summary — Cascade Europe Group", level=1)
    d.add_paragraph("Prepared: 15 March 2025")

    # Case 1: German labor court (RED FLAG)
    d.add_heading("1. Weber v. Cascade Präzisionsteile GmbH", level=2)
    lit = _EU_LITIGATION
    d.add_paragraph(f"Court: {lit['court']}")
    d.add_paragraph(f"Case reference: {lit['case_ref']}")
    d.add_paragraph(f"Filed: {lit['filing_date']}")
    d.add_paragraph(f"Status: {lit['status']}")
    d.add_paragraph(lit["description"])
    d.add_paragraph(f"Potential exposure: €{_whole_euros(lit['claim_amount']):,}")
    d.add_paragraph(f"Accrued provision: €{_whole_euros(lit['accrued']):,}")

    # Case 2: French commercial court
    d.add_heading("2. Cascade Matériaux v. SteelTech Industries", level=2)
    lit_fr = _EU_LITIGATION_FR
    d.add_paragraph(f"Court: {lit_fr['court']}")
    d.add_paragraph(f"Case reference: {lit_fr['case_ref']}")
    d.add_paragraph(f"Filed: {lit_fr['filing_date']}")
    d.add_paragraph(f"Status: {lit_fr['status']}")
    d.add_paragraph(lit_fr["description"])
    d.add_paragraph(f"Claim amount: €{_whole_euros(lit_fr['claim_amount']):,}")
    d.add_paragraph("Management assessment: probable recovery; no provision booked.")

    d.add_heading("3. Other Matters", level=2)
    d.add_paragraph(
        "No other pending or threatened litigation across the group. "
        "CD has one potential warranty claim (estimated below materiality threshold). "
        "No regulatory proceedings pending in any jurisdiction."
    )

    rel = f"{_DR}/03_legal/pending_litigation_summary.docx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_pending_litigation_summary")

    loc = embed_canary_docx(d, canary_code)

    canaries.set_location("tc12eu_pending_litigation_summary", rel, loc)
    _save_docx_deterministic(d, full)
    manifest.register(rel, "docx", canary=canaries.canary_for("tc12eu_pending_litigation_summary"), test_cases=[_TC])


def _write_ip_assignments_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """IP assignment agreements — PDF. Contains red flag: 2 missing French assignments."""
    title_s, heading_s, sub_s, body_s = _pdf_styles()
    elements = [
        Paragraph("Intellectual Property Assignment Agreements", title_s),
        Paragraph("Cascade Europe Group — All Entities", body_s),
        Spacer(1, 12),
    ]

    # Listed assignments
    for i, a in enumerate(_EU_IP_ASSIGNMENTS, 1):
        elements.extend([
            Paragraph(f"Assignment {i}", heading_s),
            Paragraph(f"Employee: {a['employee']}", body_s),
            Paragraph(f"Entity: {_ENTITIES[a['entity']]['name']} ({a['jurisdiction']})", body_s),
            Paragraph(f"Title: {a['title']}", body_s),
            Paragraph(f"Date: {a['date']}", body_s),
            Paragraph(f"Scope: {a['scope']}", body_s),
            Paragraph(f"Status: {a['status']}", body_s),
            Spacer(1, 6),
        ])

    elements.extend([
        Paragraph("Note on French IP Law", heading_s),
        Paragraph(
            "Under the French Code de la propriété intellectuelle (Articles L.611-7 "
            "and R.611-1 to R.611-7), employee inventions require a formal déclaration "
            "d'invention process. The employer must classify each invention (mission "
            "invention or hors mission attributable) and file the appropriate assignment. "
            "Failure to comply may result in the employee retaining rights to the invention.",
            body_s,
        ),
        Spacer(1, 8),
        Paragraph(
            "Note: Assignments on file cover 4 of 6 key researchers across the group. "
            "Two founding researchers at CM — Dr. Pierre Lefèvre (Chief Scientist, "
            "named inventor on 3 EP patents) and Dr. Élodie Moreau (Senior Researcher, "
            "named inventor on 1 EP patent) — do not have executed déclaration d'invention "
            "filings or IP assignment agreements on record.",
            body_s,
        ),
    ])

    rel = f"{_DR}/03_legal/ip_assignment_agreements.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_ip_assignment_agreements")
    loc = _build_simple_pdf(full, canary_code, "IP Assignments — EU Group", elements)
    canaries.set_location("tc12eu_ip_assignment_agreements", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_insurance_policies_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Insurance policies summary — PDF."""
    title_s, heading_s, _, body_s = _pdf_styles()

    policies = [
        ("Directors & Officers (D&O)", "Allianz SE", "€10,000,000", "All entities", "2025-12-31"),
        ("General Liability", "AXA S.A.", "€5,000,000", "CP, CM", "2025-09-30"),
        ("Product Liability", "Zurich Insurance Group", "€15,000,000", "All entities", "2025-06-30"),
        ("Property & Equipment", "HDI Global SE", "€25,000,000", "CP (Munich)", "2025-12-31"),
        ("Professional Indemnity", "Hiscox Ltd", "£2,000,000", "CD (UK)", "2025-11-30"),
        ("Cyber & Data Breach", "Chubb Europe", "€3,000,000", "All entities", "2025-08-31"),
    ]

    data = [["Policy Type", "Insurer", "Coverage Limit", "Covered Entities", "Renewal"]]
    for p in policies:
        data.append(list(p))

    tbl = Table(data, colWidths=[110, 90, 80, 80, 70])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5496")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))

    elements = [
        Paragraph("Insurance Policies Summary — Cascade Europe Group", title_s),
        Spacer(1, 12),
        tbl,
        Spacer(1, 12),
        Paragraph(
            "Note: Insurance claim history is maintained by each entity's local broker "
            "and has not been consolidated for this data room. Claim history records "
            "should be requested separately from each insurer.",
            body_s,
        ),
    ]

    rel = f"{_DR}/03_legal/insurance_policies_summary.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_insurance_policies_summary")
    loc = _build_simple_pdf(full, canary_code, "Insurance Summary — EU Group", elements)
    canaries.set_location("tc12eu_insurance_policies_summary", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# 04_hr — 1 XLSX + 1 PDF + 3 key-employee PDFs + 1 XLSX
# ═══════════════════════════════════════════════════════════════════════════


def _write_employee_census_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Group employee census by entity."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Employee Census"

    headers = ["Entity", "Jurisdiction", "Department", "Headcount",
               "Avg Salary (€)", "Works Council", "Social Insurance System"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    _style_header_row(ws, 1, len(headers))

    rows = [
        ("CE", "Netherlands", "Management", 5, 220000, "N/A", "Dutch social insurance (ZVW/WW)"),
        ("CE", "Netherlands", "Finance", 8, 85000, "N/A", "Dutch social insurance (ZVW/WW)"),
        ("CE", "Netherlands", "Legal & Compliance", 6, 95000, "N/A", "Dutch social insurance (ZVW/WW)"),
        ("CE", "Netherlands", "IT", 6, 90000, "N/A", "Dutch social insurance (ZVW/WW)"),
        ("CP", "Germany", "Production", 180, 52000, "Yes — Betriebsrat", "German social insurance (KV/RV/AV/PV)"),
        ("CP", "Germany", "Quality", 35, 58000, "Yes — Betriebsrat", "German social insurance (KV/RV/AV/PV)"),
        ("CP", "Germany", "Engineering", 45, 72000, "Yes — Betriebsrat", "German social insurance (KV/RV/AV/PV)"),
        ("CP", "Germany", "Administration", 25, 55000, "Yes — Betriebsrat", "German social insurance (KV/RV/AV/PV)"),
        ("CP", "Germany", "Management", 10, 120000, "Yes — Betriebsrat", "German social insurance (KV/RV/AV/PV)"),
        ("CP", "Germany", "Logistics", 15, 48000, "Yes — Betriebsrat", "German social insurance (KV/RV/AV/PV)"),
        ("CM", "France", "R&D", 40, 68000, "N/A (< 50 — no CSE)", "French social security (URSSAF)"),
        ("CM", "France", "Production", 25, 45000, "N/A (< 50 — no CSE)", "French social security (URSSAF)"),
        ("CM", "France", "Administration", 10, 52000, "N/A (< 50 — no CSE)", "French social security (URSSAF)"),
        ("CM", "France", "Management", 5, 110000, "N/A (< 50 — no CSE)", "French social security (URSSAF)"),
        ("CM", "France", "Quality", 5, 55000, "N/A (< 50 — no CSE)", "French social security (URSSAF)"),
        ("CD", "United Kingdom", "Sales", 20, 42000, "N/A", "UK NIC/PAYE"),
        ("CD", "United Kingdom", "Logistics", 12, 35000, "N/A", "UK NIC/PAYE"),
        ("CD", "United Kingdom", "Administration", 8, 45000, "N/A", "UK NIC/PAYE"),
        ("CD", "United Kingdom", "Management", 5, 85000, "N/A", "UK NIC/PAYE"),
    ]
    for r, row_data in enumerate(rows, 2):
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val)
            if c == 5:
                cell.number_format = _MONEY_FMT

    rel = f"{_DR}/04_hr/group_employee_census.xlsx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_group_employee_census")

    loc = embed_canary_xlsx(wb, canary_code)

    canaries.set_location("tc12eu_group_employee_census", rel, loc)
    _save_xlsx_deterministic(wb, full)
    manifest.register(rel, "xlsx", canary=canaries.canary_for("tc12eu_group_employee_census"), test_cases=[_TC])


def _write_benefits_summary_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Benefits summary by country — PDF."""
    title_s, heading_s, sub_s, body_s = _pdf_styles()
    elements = [
        Paragraph("Employee Benefits Summary by Country", title_s),
        Paragraph("Cascade Europe Group — All Entities", body_s),
        Spacer(1, 12),

        Paragraph("Netherlands (CE — 25 employees)", heading_s),
        Paragraph("Pension: Defined contribution scheme via Nationale-Nederlanden (5% employer contribution)", body_s),
        Paragraph("Health: Mandatory ZVW (Zorgverzekeringswet) — employer reimburses employee portion", body_s),
        Paragraph("Leave: 25 days annual leave + 8 public holidays", body_s),
        Spacer(1, 8),

        Paragraph("Germany (CP — 310 employees)", heading_s),
        Paragraph("Pension: Betriebliche Altersvorsorge (bAV) via Allianz Pensionskasse — 3% employer match", body_s),
        Paragraph("Health: Statutory health insurance (gesetzliche Krankenversicherung) — 50/50 split", body_s),
        Paragraph("Leave: 30 days annual leave + regional public holidays (Bavaria)", body_s),
        Paragraph("Works council bonus: Annual profit-sharing bonus per Betriebsvereinbarung", body_s),
        Spacer(1, 8),

        Paragraph("France (CM — 85 employees)", heading_s),
        Paragraph("Health: Mutuelle complémentaire via AXA (50% employer funded, per ANI 2016)", body_s),
        Paragraph("Pension: AGIRC-ARRCO mandatory supplementary pension", body_s),
        Paragraph("Leave: 25 days annual leave + RTT days (35h week agreement) + 11 public holidays", body_s),
        Paragraph("Profit-sharing: Participation et intéressement per collective agreement", body_s),
        Spacer(1, 8),

        Paragraph("United Kingdom (CD — 45 employees)", heading_s),
        Paragraph("Pension: Auto-enrollment workplace pension via NEST (5% employer, 3% employee)", body_s),
        Paragraph("Health: Private medical insurance via Bupa (optional, employer-funded)", body_s),
        Paragraph("Leave: 28 days including bank holidays", body_s),
    ]

    rel = f"{_DR}/04_hr/benefits_summary_by_country.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_benefits_summary_by_country")
    loc = _build_simple_pdf(full, canary_code, "Benefits Summary — EU Group", elements)
    canaries.set_location("tc12eu_benefits_summary_by_country", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_key_employee_agreement_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
    emp_idx: int,
) -> None:
    """Key employee service agreement PDF."""
    emp = _KEY_EMPLOYEES[emp_idx]
    title_s, heading_s, _, body_s = _pdf_styles()

    elements = [
        Paragraph(f"Service Agreement — {emp['name']}", title_s),
        Spacer(1, 12),
        Paragraph("Position Details", heading_s),
        Paragraph(f"Title: {emp['title']}", body_s),
        Paragraph(f"Entity: {_ENTITIES[emp['entity']]['name']}", body_s),
        Paragraph(f"Jurisdiction: {emp['jurisdiction']}", body_s),
        Spacer(1, 8),
        Paragraph("Compensation", heading_s),
        Paragraph(f"Base salary: €{_whole_euros(emp['base_salary']):,} per annum", body_s),
        Paragraph("Bonus: Up to 30% of base salary, discretionary", body_s),
        Spacer(1, 8),
        Paragraph("Notice Period", heading_s),
        Paragraph(f"Notice: {emp['notice_period']}", body_s),
        Spacer(1, 8),
        Paragraph("Non-Compete", heading_s),
        Paragraph(f"Post-termination restriction: {emp['non_compete']}", body_s),
        Spacer(1, 8),
        Paragraph("Change of Control", heading_s),
        Paragraph(
            f"Section 8.2: In the event of a change of control (defined as acquisition "
            f"of more than 50% of the voting shares or assets of the employing entity or "
            f"its ultimate parent), {emp['name']} shall be entitled to terminate this "
            f"agreement and receive a severance payment equal to {emp['coc_multiplier']}× "
            f"annual base salary (€{_whole_euros(emp['base_salary'] * emp['coc_multiplier']):,}).",
            body_s,
        ),
    ]

    rel = f"{_DR}/04_hr/key_employee_agreements/{emp['filename']}"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for(emp["canary_key"])
    loc = _build_simple_pdf(full, canary_code, f"Service Agreement — {emp['name']}", elements)
    canaries.set_location(emp["canary_key"], rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_org_chart_detailed_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Detailed org chart XLSX with reporting lines across entities."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Detailed Org Chart"

    headers = ["Entity", "Department", "Role", "Name", "Reports To", "Location"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    _style_header_row(ws, 1, len(headers))

    rows = [
        ("CE", "Management", "Managing Director", "Erik van der Berg", "Board", "Amsterdam"),
        ("CE", "Finance", "Finance Director", "Anna de Vries", "Erik van der Berg", "Amsterdam"),
        ("CE", "Legal", "General Counsel", "Marieke Jansen", "Erik van der Berg", "Amsterdam"),
        ("CP", "Management", "Geschäftsführer", "Thomas Richter", "Erik van der Berg", "Munich"),
        ("CP", "Production", "Head of Production", "Franz Müller", "Thomas Richter", "Munich"),
        ("CP", "Engineering", "Head of R&D (Manufacturing)", "Dr. Stefan Braun", "Thomas Richter", "Munich"),
        ("CP", "Quality", "Head of Quality", "Sabine Weber", "Thomas Richter", "Munich"),
        ("CM", "Management", "Directeur Général / Technical Director",
         "Dr. Pierre Lefèvre", "Erik van der Berg", "Lyon"),
        ("CM", "R&D", "Chief Scientist", "Dr. Pierre Lefèvre", "Board (CM)", "Lyon"),
        ("CM", "R&D", "Senior Researcher", "Dr. Marie Dupont", "Dr. Pierre Lefèvre", "Lyon"),
        ("CM", "Production", "Production Manager", "Jean-Paul Martin", "Dr. Pierre Lefèvre", "Lyon"),
        ("CD", "Management", "Director", "James Whitfield", "Erik van der Berg", "Birmingham"),
        ("CD", "Sales", "Head of Sales", "Sarah Thompson", "James Whitfield", "Birmingham"),
        ("CD", "Logistics", "Logistics Manager", "David Clarke", "James Whitfield", "Birmingham"),
    ]
    for r, row_data in enumerate(rows, 2):
        for c, val in enumerate(row_data, 1):
            ws.cell(row=r, column=c, value=val)

    rel = f"{_DR}/04_hr/group_org_chart_detailed.xlsx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_group_org_chart_detailed")

    loc = embed_canary_xlsx(wb, canary_code)

    canaries.set_location("tc12eu_group_org_chart_detailed", rel, loc)
    _save_xlsx_deterministic(wb, full)
    manifest.register(rel, "xlsx", canary=canaries.canary_for("tc12eu_group_org_chart_detailed"), test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# 05_tax — 4 entity CIT PDFs + 1 VAT XLSX + 1 notices PDF
# ═══════════════════════════════════════════════════════════════════════════


def _write_cit_return_cp_2023(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """CP Körperschaftsteuererklärung FY2023."""
    title_s, heading_s, _, body_s = _pdf_styles()
    elements = [
        Paragraph("Körperschaftsteuererklärung FY2023", title_s),
        Paragraph("Cascade Präzisionsteile GmbH", body_s),
        Paragraph("Finanzamt: München III", body_s),
        Paragraph("Steuernummer: 143/456/78901", body_s),
        Spacer(1, 12),
        Paragraph("Tax Computation", heading_s),
        Paragraph("Revenue: €42,800,000", body_s),
        Paragraph("Cost of sales: (€29,960,000)", body_s),
        Paragraph("Operating expenses: (€8,200,000)", body_s),
        Paragraph("Profit before tax: €4,640,000", body_s),
        Spacer(1, 8),
        Paragraph("Körperschaftsteuer (15%): €696,000", body_s),
        Paragraph("Solidaritätszuschlag (5.5% of KSt): €38,280", body_s),
        Paragraph("Gewerbesteuer (approx 14.35%): €665,840", body_s),
        Paragraph("Total tax: €1,400,120", body_s),
        Paragraph("Effective rate: 30.2%", body_s),
    ]

    rel = f"{_DR}/05_tax/cp_korperschaftsteuer_fy2023.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_cp_korperschaftsteuer_fy2023")
    loc = _build_simple_pdf(full, canary_code, "KSt FY2023 — CP", elements)
    canaries.set_location("tc12eu_cp_korperschaftsteuer_fy2023", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_cit_return_cp_2024(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """CP Körperschaftsteuererklärung FY2024."""
    title_s, heading_s, _, body_s = _pdf_styles()
    elements = [
        Paragraph("Körperschaftsteuererklärung FY2024", title_s),
        Paragraph("Cascade Präzisionsteile GmbH", body_s),
        Paragraph("Finanzamt: München III", body_s),
        Paragraph("Steuernummer: 143/456/78901", body_s),
        Spacer(1, 12),
        Paragraph("Tax Computation", heading_s),
        Paragraph("Revenue: €45,000,000", body_s),
        Paragraph("Cost of sales: (€31,500,000)", body_s),
        Paragraph("Operating expenses: (€8,640,000)", body_s),
        Paragraph("Profit before tax: €4,860,000", body_s),
        Spacer(1, 8),
        Paragraph("Körperschaftsteuer (15%): €729,000", body_s),
        Paragraph("Solidaritätszuschlag (5.5% of KSt): €40,095", body_s),
        Paragraph("Gewerbesteuer (approx 14.35%): €697,410", body_s),
        Paragraph("Total tax: €1,466,505", body_s),
        Paragraph("Effective rate: 30.2%", body_s),
    ]

    rel = f"{_DR}/05_tax/cp_korperschaftsteuer_fy2024.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_cp_korperschaftsteuer_fy2024")
    loc = _build_simple_pdf(full, canary_code, "KSt FY2024 — CP", elements)
    canaries.set_location("tc12eu_cp_korperschaftsteuer_fy2024", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_liasse_fiscale_cm(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """CM Liasse Fiscale FY2024."""
    title_s, heading_s, _, body_s = _pdf_styles()
    elements = [
        Paragraph("Liasse Fiscale — Exercice 2024", title_s),
        Paragraph("Cascade Matériaux Avancés SAS", body_s),
        Paragraph("Service des Impôts des Entreprises: Lyon 3ème", body_s),
        Paragraph("SIRET: 987 654 321 00012", body_s),
        Spacer(1, 12),
        Paragraph("Résultat Fiscal", heading_s),
        Paragraph("Chiffre d'affaires: €18,000,000", body_s),
        Paragraph("Charges d'exploitation: (€15,300,000)", body_s),
        Paragraph("Résultat d'exploitation: €2,700,000", body_s),
        Paragraph("Résultat financier: (€350,000)", body_s),
        Paragraph("Résultat courant: €2,350,000", body_s),
        Spacer(1, 8),
        Paragraph("Impôt sur les sociétés (25%): €587,500", body_s),
        Paragraph("Crédit d'Impôt Recherche (CIR): (€142,000)", body_s),
        Paragraph("IS net: €445,500", body_s),
        Paragraph("Taux effectif: 19.0%", body_s),
        Spacer(1, 8),
        Paragraph(
            "Note: CIR (Crédit d'Impôt Recherche) claimed for eligible R&D expenditure "
            "under Article 244 quater B du Code Général des Impôts. 30% rate on first "
            "€100M of eligible expenditure.",
            body_s,
        ),
    ]

    rel = f"{_DR}/05_tax/cm_liasse_fiscale_fy2024.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_cm_liasse_fiscale_fy2024")
    loc = _build_simple_pdf(full, canary_code, "Liasse Fiscale FY2024 — CM", elements)
    canaries.set_location("tc12eu_cm_liasse_fiscale_fy2024", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_ct600_cd(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """CD CT600 FY2024."""
    title_s, heading_s, _, body_s = _pdf_styles()
    elements = [
        Paragraph("Corporation Tax Return (CT600) — FY2024", title_s),
        Paragraph("Cascade Distribution Services Ltd", body_s),
        Paragraph("HMRC UTR: 1234567890", body_s),
        Spacer(1, 12),
        Paragraph("Tax Computation", heading_s),
        Paragraph("Turnover: £13,846,000 (€16,200,000 at avg rate 1.17)", body_s),
        Paragraph("Cost of sales: (£9,692,000)", body_s),
        Paragraph("Gross profit: £4,154,000", body_s),
        Paragraph("Administrative expenses: (£2,846,000)", body_s),
        Paragraph("Trading profit: £1,308,000", body_s),
        Spacer(1, 8),
        Paragraph("Corporation tax at 25%: £327,000", body_s),
        Paragraph("(€382,590 at avg rate 1.17)", body_s),
        Spacer(1, 8),
        Paragraph(
            "Note: Since 1 January 2021, CD is treated as a standalone UK taxpayer. "
            "No EU tax consolidation applies post-Brexit. UK-EU Withdrawal Agreement "
            "provisions govern transitional arrangements.",
            body_s,
        ),
    ]

    rel = f"{_DR}/05_tax/cd_ct600_fy2024.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_cd_ct600_fy2024")
    loc = _build_simple_pdf(full, canary_code, "CT600 FY2024 — CD", elements)
    canaries.set_location("tc12eu_cd_ct600_fy2024", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_vpb_aangifte_ce(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """CE Vennootschapsbelasting Aangifte FY2024."""
    title_s, heading_s, _, body_s = _pdf_styles()
    elements = [
        Paragraph("Aangifte Vennootschapsbelasting — Boekjaar 2024", title_s),
        Paragraph("Cascade Europe Holdings B.V.", body_s),
        Paragraph("Belastingdienst: Amsterdam", body_s),
        Paragraph("RSIN: 123456789", body_s),
        Spacer(1, 12),
        Paragraph("Fiscale Winst", heading_s),
        Paragraph("Opbrengsten: €11,500,000", body_s),
        Paragraph("Management fee income: €7,680,000", body_s),
        Paragraph("Dividend income: €3,820,000 (participation exemption applies)", body_s),
        Spacer(1, 8),
        Paragraph("Kosten: (€9,800,000)", body_s),
        Paragraph("Personnel: (€4,200,000)", body_s),
        Paragraph("Other operating: (€5,600,000)", body_s),
        Spacer(1, 8),
        Paragraph("Belastbare winst: €1,700,000", body_s),
        Paragraph("(After deelnemingsvrijstelling on dividend income)", body_s),
        Spacer(1, 8),
        Paragraph("VPB (25.8%): €438,600", body_s),
        Spacer(1, 8),
        Paragraph(
            "Note: Deelnemingsvrijstelling (participation exemption) applied to "
            "€3,820,000 dividend income from qualifying subsidiaries per "
            "Article 13 Wet VPB 1969.",
            body_s,
        ),
    ]

    rel = f"{_DR}/05_tax/ce_vpb_aangifte_fy2024.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_ce_vpb_aangifte_fy2024")
    loc = _build_simple_pdf(full, canary_code, "VPB Aangifte FY2024 — CE", elements)
    canaries.set_location("tc12eu_ce_vpb_aangifte_fy2024", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_vat_returns_summary_eu(
    output_dir: Path, canaries: CanaryRegistry, errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Group VAT returns summary — XLSX. Contains ERR-EU-012."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Quarterly VAT Returns"

    headers = [
        "Entity", "Quarter", "Output VAT — Domestic (€)",
        "Output VAT — Intra-EU/Export (€)", "Input VAT — Domestic (€)",
        "Input VAT — Reverse Charge (€)", "VAT Payable / (Refundable) (€)",
        "Filing Status",
    ]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    _style_header_row(ws, 1, len(headers))

    for r, row_data in enumerate(_VAT_RETURNS, 2):
        entity, qtr, out_dom, out_eu, in_dom, in_rc, payable, status = row_data
        ws.cell(row=r, column=1, value=entity)
        ws.cell(row=r, column=2, value=qtr)
        ws.cell(row=r, column=3, value=out_dom).number_format = _MONEY_FMT
        ws.cell(row=r, column=4, value=out_eu).number_format = _MONEY_FMT
        ws.cell(row=r, column=5, value=in_dom).number_format = _MONEY_FMT
        ws.cell(row=r, column=6, value=in_rc).number_format = _MONEY_FMT
        ws.cell(row=r, column=7, value=payable).number_format = _MONEY_FMT
        ws.cell(row=r, column=8, value=status)

    # Register the planted error
    errors.add(PlantedError(
        error_id="ERR-EU-012",
        file=f"{_DR}/05_tax/group_vat_returns_summary.xlsx",
        location="Sheet 'Quarterly VAT Returns', CP Q3 row, Column C (Output VAT — Domestic)",
        type="transposed_digits",
        description=(
            "CP (Germany) Q3 2025 output VAT is recorded as €847,200 but should be "
            "€874,200 (transposed digits '47' ↔ '74'). Discoverable by comparing Q3 "
            "against other quarters (Q1, Q2, Q4 all show €874,200) or by cross-referencing "
            "against intercompany sales data."
        ),
        severity="material",
        which_test_cases_should_catch=["TC-12-EU"],
    ))

    rel = f"{_DR}/05_tax/group_vat_returns_summary.xlsx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_group_vat_returns_summary")

    loc = embed_canary_xlsx(wb, canary_code)

    canaries.set_location("tc12eu_group_vat_returns_summary", rel, loc)
    _save_xlsx_deterministic(wb, full)
    manifest.register(rel, "xlsx", canary=canaries.canary_for("tc12eu_group_vat_returns_summary"), test_cases=[_TC])


def _write_tax_notices_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Tax notices and assessments — PDF. Contains Betriebsprüfung red flag."""
    title_s, heading_s, sub_s, body_s = _pdf_styles()
    bp = _BETRIEBSPRUEFUNG
    elements = [
        Paragraph("Tax Notices and Assessments — Cascade Europe Group", title_s),
        Spacer(1, 12),

        Paragraph("1. Betriebsprüfung — Cascade Präzisionsteile GmbH", heading_s),
        Paragraph(f"Entity: {bp['entity_name']}", body_s),
        Paragraph(f"Notice date: {bp['audit_notice_date']}", body_s),
        Paragraph(f"Audit period: {bp['audit_period']}", body_s),
        Paragraph(f"Tax types: {bp['tax_types']}", body_s),
        Paragraph(f"Conducting authority: {bp['auditor']}", body_s),
        Paragraph(f"Status: {bp['status']}", body_s),
        Paragraph(f"Focus areas: {bp['focus_areas']}", body_s),
        Paragraph(f"Estimated exposure: {bp['estimated_exposure']}", body_s),
        Spacer(1, 8),
        Paragraph(
            "Management note: The Betriebsprüfung is a routine comprehensive audit "
            "(Außenprüfung) covering Körperschaftsteuer, Gewerbesteuer, and Umsatzsteuer. "
            "However, the focus on transfer pricing of intercompany transactions and "
            "Gewerbesteuer Hinzurechnungen on license payments could result in material "
            "adjustments. No preliminary findings have been communicated yet.",
            body_s,
        ),

        Spacer(1, 12),
        Paragraph("2. Other Tax Notices", heading_s),
        Paragraph(
            "CE (Netherlands): No outstanding assessments or disputes with Belastingdienst.",
            body_s,
        ),
        Paragraph(
            "CM (France): CIR (Crédit d'Impôt Recherche) claim for FY2023 accepted without "
            "adjustment by the Direction Générale des Finances Publiques.",
            body_s,
        ),
        Paragraph(
            "CD (United Kingdom): HMRC enquiry into FY2023 CT600 closed with no adjustments.",
            body_s,
        ),
    ]

    rel = f"{_DR}/05_tax/tax_notices_and_assessments.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_tax_notices_and_assessments")
    loc = _build_simple_pdf(full, canary_code, "Tax Notices — EU Group", elements)
    canaries.set_location("tc12eu_tax_notices_and_assessments", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# 06_operations — 1 PDF + 3 XLSX
# ═══════════════════════════════════════════════════════════════════════════


def _write_facility_leases_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Facility leases — PDF."""
    title_s, heading_s, _, body_s = _pdf_styles()

    facilities = [
        ("CE", "Amsterdam", "WTC Amsterdam, Strawinskylaan 1627", "Office", "Lease", "€180,000/yr", "2027-12-31"),
        ("CP", "Munich", "Industriestr. 45, 80939 München", "Factory + Office", "Lease", "€1,200,000/yr", "2030-06-30"),
        ("CP", "Munich", "Industriestr. 47, 80939 München", "Warehouse", "Lease", "€320,000/yr", "2028-12-31"),
        ("CM", "Lyon", "Zone Industrielle, 69007 Lyon", "R&D Lab + Office", "Lease", "€450,000/yr", "2029-03-31"),
        ("CD", "Birmingham", "Unit 12, Castle Bromwich Business Park",
         "Warehouse + Office", "Lease", "£180,000/yr", "2027-09-30"),
    ]

    data = [["Entity", "City", "Address", "Type", "Tenure", "Annual Rent", "Expiry"]]
    for f in facilities:
        data.append(list(f))

    tbl = Table(data, colWidths=[30, 55, 120, 60, 35, 60, 55])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5496")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))

    elements = [
        Paragraph("Facility Leases Summary — Cascade Europe Group", title_s),
        Spacer(1, 12),
        tbl,
        Spacer(1, 12),
        Paragraph("All leases accounted for under IFRS 16. Right-of-use assets recognized on balance sheet.", body_s),
    ]

    rel = f"{_DR}/06_operations/facility_leases.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_facility_leases")
    loc = _build_simple_pdf(full, canary_code, "Facility Leases — EU Group", elements)
    canaries.set_location("tc12eu_facility_leases", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_equipment_list_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Equipment list — XLSX."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Equipment List"

    headers = ["Asset ID", "Entity", "Description", "Location",
               "Cost (€)", "Accum. Depr. (€)", "NBV (€)", "Useful Life"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    _style_header_row(ws, 1, len(headers))

    assets = [
        ("CP-EQ-001", "CP", "CNC Machining Centre (5-axis)", "Munich", 1200000, 480000, 720000, "10 years"),
        ("CP-EQ-002", "CP", "Thermal Spray Coating System", "Munich", 850000, 340000, 510000, "10 years"),
        ("CP-EQ-003", "CP", "Precision Grinding Machine", "Munich", 650000, 195000, 455000, "10 years"),
        ("CP-EQ-004", "CP", "Quality Inspection CMM", "Munich", 380000, 152000, 228000, "10 years"),
        ("CP-EQ-005", "CP", "Industrial Robot (Assembly)", "Munich", 420000, 84000, 336000, "10 years"),
        ("CM-EQ-001", "CM", "R&D Autoclave System", "Lyon", 980000, 294000, 686000, "10 years"),
        ("CM-EQ-002", "CM", "Materials Testing Lab Equipment", "Lyon", 560000, 224000, 336000, "10 years"),
        ("CM-EQ-003", "CM", "Polymer Processing Line", "Lyon", 720000, 216000, 504000, "10 years"),
        ("CD-EQ-001", "CD", "Warehouse Racking System", "Birmingham", 180000, 72000, 108000, "10 years"),
        ("CD-EQ-002", "CD", "Forklift Fleet (4 units)", "Birmingham", 120000, 48000, 72000, "5 years"),
    ]
    for r, asset in enumerate(assets, 2):
        for c, val in enumerate(asset, 1):
            cell = ws.cell(row=r, column=c, value=val)
            if c in (5, 6, 7):
                cell.number_format = _MONEY_FMT

    rel = f"{_DR}/06_operations/equipment_list.xlsx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_equipment_list")

    loc = embed_canary_xlsx(wb, canary_code)

    canaries.set_location("tc12eu_equipment_list", rel, loc)
    _save_xlsx_deterministic(wb, full)
    manifest.register(rel, "xlsx", canary=canaries.canary_for("tc12eu_equipment_list"), test_cases=[_TC])


def _write_customer_list_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Customer list with revenue — XLSX."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Customer Revenue"

    headers = ["Customer", "Entity", "Country", "FY2024 Revenue (€)",
               "Revenue Share (%)", "Currency"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    _style_header_row(ws, 1, len(headers))

    customers = [
        ("Autohaus Müller GmbH", "CP", "Germany", 12500000, 14.9, "EUR"),
        ("BMW AG", "CP", "Germany", 8200000, 9.8, "EUR"),
        ("Airbus Defence & Space", "CP", "Germany", 6800000, 8.1, "EUR"),
        ("Volkswagen AG", "CP", "Germany", 5400000, 6.4, "EUR"),
        ("Renault S.A.", "CM", "France", 8200000, 9.8, "EUR"),
        ("Safran S.A.", "CM", "France", 4500000, 5.4, "EUR"),
        ("Stellantis N.V.", "CM", "France", 3200000, 3.8, "EUR"),
        ("Rolls-Royce plc", "CD", "UK", 5800000, 6.9, "GBP→EUR"),
        ("BAE Systems plc", "CD", "UK", 4200000, 5.0, "GBP→EUR"),
        ("Other customers (various)", "Group", "Various", 24500000, 29.2, "EUR/GBP"),
    ]
    for r, cust in enumerate(customers, 2):
        for c, val in enumerate(cust, 1):
            cell = ws.cell(row=r, column=c, value=val)
            if c == 4:
                cell.number_format = _MONEY_FMT

    # Total row
    ws.cell(row=12, column=1, value="TOTAL").font = Font(bold=True)
    ws.cell(row=12, column=4, value=83300000).number_format = _MONEY_FMT
    ws.cell(row=12, column=5, value=100.0)

    rel = f"{_DR}/06_operations/customer_list_with_revenue.xlsx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_customer_list_with_revenue")

    loc = embed_canary_xlsx(wb, canary_code)

    canaries.set_location("tc12eu_customer_list_with_revenue", rel, loc)
    _save_xlsx_deterministic(wb, full)
    manifest.register(rel, "xlsx", canary=canaries.canary_for("tc12eu_customer_list_with_revenue"), test_cases=[_TC])


def _write_vendor_list_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Vendor list — XLSX."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Vendor List"

    headers = ["Vendor", "Entity", "Category", "Country",
               "FY2024 Spend (€)", "Payment Terms"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    _style_header_row(ws, 1, len(headers))

    vendors = [
        ("thyssenkrupp Materials Services", "CP", "Raw Materials", "Germany", 15000000, "Net 45"),
        ("Salzgitter AG", "CP", "Steel", "Germany", 5200000, "Net 30"),
        ("SteelTech Industries SARL", "CM", "Raw Materials", "France", 3800000, "Net 60"),
        ("Air Liquide S.A.", "CP/CM", "Industrial Gases", "France", 2100000, "Net 30"),
        ("Siemens AG", "CP", "Automation Equipment", "Germany", 1800000, "Net 60"),
        ("RS Components Ltd", "CD", "MRO Supplies", "UK", 450000, "Net 30"),
        ("DHL Supply Chain", "CD", "Logistics", "Germany", 1200000, "Monthly"),
        ("SAP SE", "Group", "ERP Software", "Germany", 380000, "Annual"),
    ]
    for r, vendor in enumerate(vendors, 2):
        for c, val in enumerate(vendor, 1):
            cell = ws.cell(row=r, column=c, value=val)
            if c == 5:
                cell.number_format = _MONEY_FMT

    rel = f"{_DR}/06_operations/vendor_list.xlsx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_vendor_list")

    loc = embed_canary_xlsx(wb, canary_code)

    canaries.set_location("tc12eu_vendor_list", rel, loc)
    _save_xlsx_deterministic(wb, full)
    manifest.register(rel, "xlsx", canary=canaries.canary_for("tc12eu_vendor_list"), test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# 07_technology — 1 PDF + 1 XLSX + 1 DOCX
# ═══════════════════════════════════════════════════════════════════════════


def _write_patent_portfolio_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Patent portfolio — PDF. Contains red flag: 2 expiring patents."""
    title_s, heading_s, _, body_s = _pdf_styles()

    data = [["Number", "Title", "Filing", "Expiry", "Applicant", "Status"]]
    for p in _EU_PATENTS:
        data.append([
            p["number"],
            p["title"][:45] + "..." if len(p["title"]) > 45 else p["title"],
            str(p["filing_date"]),
            str(p["expiration_date"]),
            p["applicant"].split()[1],  # Short name
            p["status"],
        ])

    tbl = Table(data, colWidths=[75, 110, 55, 55, 65, 75])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5496")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))

    elements = [
        Paragraph("Patent Portfolio — Cascade Europe Group", title_s),
        Spacer(1, 12),
        Paragraph("European Patent Office (EPO) and National Patents", heading_s),
        tbl,
        Spacer(1, 12),
        Paragraph(
            "Note: EP 1,890,456 (surface treatment — core manufacturing process used by CP) "
            "and EP 1,678,901 (polymer reinforcement) expire within 18 months of FY2025 "
            "year-end. Management should assess renewal/continuation strategy and "
            "alternative IP protection.",
            body_s,
        ),
        Spacer(1, 8),
        Paragraph(
            "IP ownership: EPO patents are held by the entity whose researchers developed "
            "the invention. See IP assignment agreements for employee-invention assignment status.",
            body_s,
        ),
    ]

    rel = f"{_DR}/07_technology/patent_portfolio.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_patent_portfolio")
    loc = _build_simple_pdf(full, canary_code, "Patent Portfolio — EU Group", elements)
    canaries.set_location("tc12eu_patent_portfolio", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


def _write_software_licenses_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Software licenses — XLSX."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Software Licenses"

    headers = ["Software", "Vendor", "License Type", "Entities",
               "Annual Cost (€)", "Renewal Date"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    _style_header_row(ws, 1, len(headers))

    licenses = [
        ("SAP S/4HANA", "SAP SE", "Enterprise", "Group", 280000, "2026-01-01"),
        ("Microsoft 365", "Microsoft", "Subscription", "Group", 95000, "2025-07-01"),
        ("CATIA V5", "Dassault Systèmes", "Named user", "CM", 120000, "2025-12-01"),
        ("Siemens NX", "Siemens", "Floating", "CP", 85000, "2025-09-01"),
        ("SolidWorks", "Dassault Systèmes", "Named user", "CP", 45000, "2025-11-01"),
        ("Salesforce", "Salesforce", "Subscription", "CD", 35000, "2025-06-01"),
        ("CyberArk PAM", "CyberArk", "Subscription", "Group", 42000, "2026-03-01"),
        ("Jira / Confluence", "Atlassian", "Cloud", "Group", 18000, "2025-08-01"),
    ]
    for r, lic in enumerate(licenses, 2):
        for c, val in enumerate(lic, 1):
            cell = ws.cell(row=r, column=c, value=val)
            if c == 5:
                cell.number_format = _MONEY_FMT

    rel = f"{_DR}/07_technology/software_licenses.xlsx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_software_licenses")

    loc = embed_canary_xlsx(wb, canary_code)

    canaries.set_location("tc12eu_software_licenses", rel, loc)
    _save_xlsx_deterministic(wb, full)
    manifest.register(rel, "xlsx", canary=canaries.canary_for("tc12eu_software_licenses"), test_cases=[_TC])


def _write_it_infrastructure_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """IT infrastructure overview — DOCX."""
    d = docx.Document()
    d.add_heading("IT Infrastructure Overview — Cascade Europe Group", level=1)

    d.add_heading("Network Architecture", level=2)
    d.add_paragraph(
        "Primary data centre: Frankfurt (AWS eu-central-1). "
        "Disaster recovery: Dublin (AWS eu-west-1). "
        "All production systems hosted within EU to comply with GDPR data residency. "
        "CD (UK) connects via dedicated VPN; no local data centre."
    )

    d.add_heading("ERP and Core Systems", level=2)
    d.add_paragraph(
        "SAP S/4HANA (on-premise converted 2023): Central ERP for CE, CP, CM. "
        "CD uses a separate SAP instance connected via CPI integration middleware."
    )

    d.add_heading("Cybersecurity", level=2)
    d.add_paragraph(
        "ISO 27001 certification: CE and CP certified (last audit Nov 2024). "
        "CM and CD in progress for certification (expected Q3 2025). "
        "Annual penetration test conducted by external firm (NCC Group). "
        "No material findings in FY2024 assessment."
    )

    d.add_heading("Disaster Recovery", level=2)
    d.add_paragraph(
        "RPO: 1 hour. RTO: 4 hours. "
        "DR tested quarterly. Last successful test: January 2025."
    )

    d.add_heading("IT Headcount", level=2)
    d.add_paragraph("CE: 6 FTE (central IT function). CP: 4 FTE (local support). "
                     "CM: 2 FTE. CD: 1 FTE.")

    rel = f"{_DR}/07_technology/it_infrastructure_overview.docx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_it_infrastructure_overview")

    loc = embed_canary_docx(d, canary_code)

    canaries.set_location("tc12eu_it_infrastructure_overview", rel, loc)
    _save_docx_deterministic(d, full)
    manifest.register(rel, "docx", canary=canaries.canary_for("tc12eu_it_infrastructure_overview"), test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# 08_compliance — 1 XLSX + 1 PDF (EU-specific category)
# ═══════════════════════════════════════════════════════════════════════════


def _write_gdpr_register(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """GDPR Article 30 data processing register — XLSX. Contains missing DPA red flag."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "GDPR Processing Register"

    headers = [
        "Processing Activity", "Controller", "Processor",
        "Data Categories", "Data Subjects", "Legal Basis",
        "Retention Period", "DPA Status", "Transfer Outside EU",
    ]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    _style_header_row(ws, 1, len(headers))

    activities = [
        ("Employee payroll processing", "CP/CM/CE", "PayrollPlus GmbH",
         "Personal data, salary, tax IDs", "Employees", "Art. 6(1)(b) Contract",
         "10 years post-employment", "NOT EXECUTED", "No"),
        ("Cloud hosting", "Group", "CloudServe B.V.",
         "All business data", "Employees, customers", "Art. 6(1)(f) Legitimate interest",
         "Duration of service", "NOT EXECUTED", "No"),
        ("Document management", "Group", "SecureDoc SAS",
         "Corporate documents", "Employees", "Art. 6(1)(f) Legitimate interest",
         "Duration of service + 5 years", "Executed — 2023-06-15", "No"),
        ("IT security monitoring", "Group", "CyberGuard Ltd",
         "Network/access logs", "Employees", "Art. 6(1)(f) Legitimate interest",
         "12 months", "Executed — 2024-01-10", "UK (adequacy decision)"),
        ("Customer CRM", "CD", "Salesforce Inc.",
         "Customer contact data", "Customers", "Art. 6(1)(b) Contract",
         "Duration of relationship", "Executed — 2023-09-01", "US (SCCs + TIA)"),
        ("HR recruitment", "Group", "In-house",
         "Candidate personal data", "Job applicants", "Art. 6(1)(a) Consent",
         "6 months post-application", "N/A (in-house)", "No"),
    ]
    for r, act in enumerate(activities, 2):
        for c, val in enumerate(act, 1):
            ws.cell(row=r, column=c, value=val)

    # Add note about missing DPAs
    ws.cell(row=9, column=1, value="COMPLIANCE NOTE:")
    ws.cell(row=9, column=2, value=(
        "2 of 4 external processors do not have executed Data Processing Agreements (DPAs). "
        "CloudServe B.V. and PayrollPlus GmbH are operating without DPAs — this represents "
        "a material GDPR compliance gap (Art. 28 GDPR requires a binding DPA with each processor)."
    ))

    rel = f"{_DR}/08_compliance/gdpr_data_processing_register.xlsx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_gdpr_data_processing_register")

    loc = embed_canary_xlsx(wb, canary_code)

    canaries.set_location("tc12eu_gdpr_data_processing_register", rel, loc)
    _save_xlsx_deterministic(wb, full)
    manifest.register(rel, "xlsx", canary=canaries.canary_for("tc12eu_gdpr_data_processing_register"), test_cases=[_TC])


def _write_works_council_agreements(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """Works council (Betriebsrat) agreements — PDF. Contains consultation requirement red flag."""
    title_s, heading_s, sub_s, body_s = _pdf_styles()
    wc = _WORKS_COUNCIL
    elements = [
        Paragraph(f"Betriebsvereinbarung — {wc['entity_name']}", title_s),
        Spacer(1, 12),
        Paragraph("Works Council Details", heading_s),
        Paragraph(f"Entity: {wc['entity_name']}", body_s),
        Paragraph(f"Established: {wc['established']}", body_s),
        Paragraph(f"Members: {wc['members']} (threshold: 5+ for 200-500 employees per §9 BetrVG)", body_s),
        Paragraph(f"Chair: {wc['chair']}", body_s),
        Spacer(1, 8),
        Paragraph("Current Agreement", heading_s),
        Paragraph(f"Date: {wc['agreement_date']}", body_s),
        Paragraph(f"Scope: {wc['agreement_scope']}", body_s),
        Spacer(1, 8),
        Paragraph("Change of Control Provisions", heading_s),
        Paragraph(wc["coc_coverage"], body_s),
        Spacer(1, 8),
        Paragraph("§111 BetrVG — Betriebsänderung (Fundamental Change)", sub_s),
        Paragraph(
            "A change of ownership constitutes a Betriebsänderung under §111 BetrVG. "
            "The employer must inform and consult the works council. Key requirements:\n"
            "• Timely and comprehensive information about the planned transaction\n"
            "• Consultation on the impact on employees\n"
            "• Negotiation of a Sozialplan (social plan) if the change results in significant "
            "disadvantages for the workforce\n"
            "• Interessenausgleich (reconciliation of interests) attempt required\n"
            "The works council cannot block the transaction, but failure to consult may result "
            "in claims under §113 BetrVG (Nachteilsausgleich).",
            body_s,
        ),
    ]

    rel = f"{_DR}/08_compliance/works_council_agreements_cp.pdf"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_works_council_agreements_cp")
    loc = _build_simple_pdf(full, canary_code, "Works Council — CP", elements)
    canaries.set_location("tc12eu_works_council_agreements_cp", rel, loc)
    manifest.register(rel, "pdf", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# DD Checklist
# ═══════════════════════════════════════════════════════════════════════════


def _write_dd_checklist_eu(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """European DD checklist — 72 items across 10 categories."""
    d = docx.Document()
    d.add_heading("European Due Diligence Checklist", level=1)
    d.add_paragraph(
        "Cascade Europe Holdings B.V. — Potential Acquisition\n"
        "72 line items across 10 categories. Adapted for multi-jurisdiction "
        "European due diligence (NL, DE, FR, UK)."
    )

    current_category = None
    item_num = 0
    for category, item in _DD_CHECKLIST_EU:
        if category != current_category:
            d.add_heading(category, level=2)
            current_category = category
        item_num += 1
        d.add_paragraph(f"{item_num}. {item}", style="List Number")

    d.add_heading("Notes", level=2)
    d.add_paragraph(
        "This checklist has been expanded from the standard 65-item US checklist "
        "to 72 items to cover EU-specific requirements including GDPR compliance, "
        "works council documentation, VAT registrations, transfer pricing (OECD), "
        "and jurisdiction-specific regulatory requirements."
    )

    rel = f"{_INPUT_DIR}/dd_checklist_european.docx"
    full = output_dir / rel
    full.parent.mkdir(parents=True, exist_ok=True)
    canary_code = canaries.canary_for("tc12eu_dd_checklist_european")

    loc = embed_canary_docx(d, canary_code)

    canaries.set_location("tc12eu_dd_checklist_european", rel, loc)
    _save_docx_deterministic(d, full)
    manifest.register(rel, "docx", canary=canaries.canary_for("tc12eu_dd_checklist_european"), test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# Prompt and Expected Behavior
# ═══════════════════════════════════════════════════════════════════════════


def _write_prompt(output_dir: Path) -> None:
    text = """\
# TC-12-EU: European Data Room Triage & Document Index

You have access to a deal data room for the potential acquisition of \
Cascade Europe Holdings B.V. and its European subsidiaries (Germany, France, UK).

## Input Files

- `data_room/` — 41 files organized into 8 categories:
  - `01_corporate/` — KvK extract, articles of association, Handelsregister, \
K-bis, Companies House filing, board minutes, org chart
  - `02_financial/` — IFRS consolidated financials, management accounts, budget, \
debt schedule
  - `03_legal/` — Material contracts (Autohaus Müller, Renault, thyssenkrupp), \
pending litigation, IP assignments, insurance
  - `04_hr/` — Employee census, benefits by country, key employee agreements, \
org chart
  - `05_tax/` — Entity CIT returns (KSt, Liasse Fiscale, CT600, VPB), VAT \
returns summary, tax notices
  - `06_operations/` — Facility leases, equipment list, customer/vendor lists
  - `07_technology/` — Patent portfolio (EPO), software licenses, IT overview
  - `08_compliance/` — GDPR data processing register, works council agreements
- `dd_checklist_european.docx` — European due diligence checklist with 72 line items

## Tasks

1. Create a complete document index: for each file, provide:
   - File path
   - Document type/category
   - Jurisdiction (NL, DE, FR, UK, or Group)
   - Date (if identifiable)
   - One-line summary of contents
   - Key data points or red flags noted

2. Cross-reference against the European due diligence checklist:
   - Mark which checklist items have corresponding documents
   - Flag checklist items with NO corresponding document (gaps)
   - Note any documents in the data room not covered by the checklist

3. Identify red flags or items requiring immediate attention, such as:
   - Pending litigation or regulatory proceedings
   - Contracts with change-of-control provisions
   - Key employee agreements with unusual terms
   - Missing critical documents
   - GDPR compliance gaps
   - Works council / employee representative body consultation requirements
   - Open tax audits or assessments

4. Prioritize the gaps: which missing documents are deal-critical vs. nice-to-have?
   Consider jurisdiction-specific requirements (e.g., German works council \
consultation is a legal requirement, not optional).

## Deliverables

Export:
- **Document index** as Excel (with columns: Path, Category, Jurisdiction, \
Date, Summary, Red Flags)
- **Gap analysis** as Excel (Checklist Item, Status, Jurisdiction, Priority, Notes)
- **Red flags summary** as Word memo
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    text = """\
# TC-12-EU: European Data Room Triage — Expected Behavior

## Key Red Flags (agent must identify all 8)

| # | Red Flag | Location | Severity |
|---|----------|----------|----------|
| 1 | Munich labor court claim — €180K exposure | `03_legal/pending_litigation_summary.docx` | High |
| 2 | Managing director CoC severance — 2.5× (€700K) | `04_hr/.../managing_director_service_agreement_ce.pdf` | High |
| 3 | Renault CoC termination clause | `03_legal/material_contracts/customer_agreement_renault_tier2.pdf` | High |
| 4 | Missing GDPR DPAs (2 of 4 processors) | `08_compliance/gdpr_data_processing_register.xlsx` | High |
| 5 | IP assignment gaps (2 French researchers) | `03_legal/ip_assignment_agreements.pdf` | High |
| 6 | 2 EPO patents expiring within 18 months | `07_technology/patent_portfolio.pdf` | Medium |
| 7 | Open Betriebsprüfung for CP (FY2022-2023) | `05_tax/tax_notices_and_assessments.pdf` | High |
| 8 | Works council consultation requirement | `08_compliance/works_council_agreements_cp.pdf` | High |

## Planted Error

ERR-EU-012: CP Q3 output VAT in group_vat_returns_summary.xlsx shows €847,200 \
instead of €874,200 (transposed digits). Discoverable by comparing against other \
quarters or cross-referencing intercompany sales data.

## Critical Document Gaps (missing from data room)

| Category | Missing Item | Priority |
|----------|-------------|----------|
| Tax/VAT | Transfer pricing documentation (Master File/Local File) | Deal-critical |
| Employment/Labor | CSE consultation records (CM, France) | Deal-critical |
| Data Protection/GDPR | Missing DPAs for 2 processors | Deal-critical |
| Environmental | Environmental impact assessments (CP Munich) | High |
| Environmental | German Immissionsschutz permits | High |
| Regulatory | UK post-Brexit customs/trade compliance | High |
| Legal | Insurance claim history | High |
| Regulatory | Real property surveys | Medium |

## Judgment Traps

1. **US-framing trap**: Agent treats KvK extract as "articles of incorporation" \
or looks for "bylaws" — applies US corporate document taxonomy.
2. **GDPR underweight**: Agent identifies missing DPAs but does not flag as \
deal-critical (potential fines up to 4% of global turnover).
3. **Works council omission**: Agent fails to flag Betriebsrat consultation — \
this is a mandatory legal step, not optional.
4. **Betriebsprüfung dismissal**: Agent treats open tax audit as routine — \
unquantified liability covering 2 fiscal years.
5. **French IP law ignorance**: Agent flags missing IP assignments without noting \
French-specific déclaration d'invention requirements.
6. **UK post-Brexit blind spot**: Agent does not flag absence of customs/trade \
compliance documentation for CD.

## Output Quality

- Document index should cover all 41 files with accurate jurisdiction tags.
- Gap analysis should cover all 72 checklist items with clear status.
- Red flags memo should be professional, jurisdiction-aware, and prioritized.
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ═══════════════════════════════════════════════════════════════════════════
# Gold Standard
# ═══════════════════════════════════════════════════════════════════════════


@register_gold("TC-12-EU")
def _tc12_eu_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """Build the TC-12-EU gold standard."""
    # Patents expiring within 18 months of FY2025 year-end
    cutoff_18m = datetime.date(2027, 6, 30)
    expiring_patents = [p for p in _EU_PATENTS if p["expiration_date"] <= cutoff_18m]

    return GoldStandard(
        test_case=_TC,
        expected_outputs={
            "output_files": {
                "document_index": {
                    "type": "xlsx",
                    "required_columns": [
                        "Path", "Category", "Jurisdiction", "Date", "Summary", "Red Flags",
                    ],
                    "total_files_indexed": 41,
                },
                "gap_analysis": {
                    "type": "xlsx",
                    "required_columns": [
                        "Checklist Item", "Status", "Jurisdiction", "Priority", "Notes",
                    ],
                    "total_checklist_items": 72,
                },
                "red_flags_memo": {
                    "type": "docx",
                    "required_sections": [
                        "Pending Litigation",
                        "Change-of-Control Provisions",
                        "Key Employee Agreements",
                        "IP and Patent Risks",
                        "GDPR Compliance Gaps",
                        "Works Council / Labor Requirements",
                        "Tax Audit Exposure",
                        "Missing Critical Documents",
                    ],
                },
            },
            "red_flags": {
                "litigation": {
                    "title": _EU_LITIGATION["title"],
                    "court": _EU_LITIGATION["court"],
                    "potential_exposure_eur": _whole_euros(_EU_LITIGATION["claim_amount"]),
                    "accrued_provision_eur": _whole_euros(_EU_LITIGATION["accrued"]),
                    "jurisdiction_note": (
                        "German labor courts strongly favor employees; "
                        "KSchG remedies include reinstatement or up to 12 months' severance"
                    ),
                },
                "managing_director_coc": {
                    "name": _KEY_EMPLOYEES[0]["name"],
                    "multiplier": float(_KEY_EMPLOYEES[0]["coc_multiplier"]),
                    "payout_eur": _whole_euros(
                        _KEY_EMPLOYEES[0]["base_salary"] * _KEY_EMPLOYEES[0]["coc_multiplier"],
                    ),
                },
                "renault_change_of_control": {
                    "contract": _RENAULT_CONTRACT["name"],
                    "annual_value_eur": _whole_euros(_RENAULT_CONTRACT["value"]),
                    "clause": "Section 14.3 — termination on change of indirect control",
                },
                "missing_gdpr_dpas": {
                    "processors_without_dpa": [
                        p["processor"] for p in _GDPR_PROCESSORS if "Not executed" in p["dpa_status"]
                    ],
                    "total_processors": len([p for p in _GDPR_PROCESSORS if p["dpa_status"] != "N/A (in-house)"]),
                    "processors_with_dpa": 2,
                    "compliance_gap": "Art. 28 GDPR requires binding DPA with each processor",
                },
                "ip_assignment_gaps": {
                    "missing_persons": _MISSING_IP_ASSIGNMENTS_EU,
                    "french_law_note": (
                        "Code de la propriété intellectuelle requires déclaration d'invention; "
                        "missing filings may leave invention rights with employees"
                    ),
                },
                "expiring_patents": {
                    "count": len(expiring_patents),
                    "patent_numbers": [p["number"] for p in expiring_patents],
                    "note": "EP 1,890,456 covers core manufacturing process used by CP",
                },
                "betriebspruefung": {
                    "entity": _BETRIEBSPRUEFUNG["entity_name"],
                    "audit_period": _BETRIEBSPRUEFUNG["audit_period"],
                    "status": "In progress — no preliminary findings",
                    "focus": "Transfer pricing and Gewerbesteuer Hinzurechnungen",
                    "exposure": "Unquantified",
                },
                "works_council_consultation": {
                    "entity": _WORKS_COUNCIL["entity_name"],
                    "employees": _ENTITIES["CP"]["employees"],
                    "legal_basis": "§111-113 BetrVG",
                    "requirement": (
                        "Change of ownership requires works council consultation; "
                        "current agreement (2019) does not cover acquisition scenarios"
                    ),
                },
                "missing_documents": [
                    "Environmental impact assessments (CP Munich — German Immissionsschutz)",
                    "Transfer pricing documentation (Master File / Local File per OECD)",
                    "UK post-Brexit customs / trade compliance documentation (CD)",
                    "French CSE consultation records (CM)",
                    "Insurance claim history across all entities",
                    "Real property surveys for all facilities",
                ],
            },
            "gap_analysis_summary": {
                "items_with_documents": 42,
                "items_without_documents": 30,
                "deal_critical_gaps": [
                    "Transfer pricing documentation",
                    "Works council consultation confirmation",
                    "GDPR DPAs (2 missing)",
                    "French CSE records",
                ],
            },
        },
        canary_verification={
            "read_kvk_extract": canaries.canary_for("tc12eu_ce_kvk_extract"),
            "read_renault_contract": canaries.canary_for("tc12eu_customer_agreement_renault"),
            "read_litigation_summary": canaries.canary_for("tc12eu_pending_litigation_summary"),
            "read_managing_director_agreement": canaries.canary_for("tc12eu_managing_director_agreement_ce"),
            "read_patent_portfolio": canaries.canary_for("tc12eu_patent_portfolio"),
            "read_ip_assignments": canaries.canary_for("tc12eu_ip_assignment_agreements"),
            "read_gdpr_register": canaries.canary_for("tc12eu_gdpr_data_processing_register"),
            "read_works_council": canaries.canary_for("tc12eu_works_council_agreements_cp"),
            "read_tax_notices": canaries.canary_for("tc12eu_tax_notices_and_assessments"),
            "read_dd_checklist": canaries.canary_for("tc12eu_dd_checklist_european"),
        },
        error_detection={
            "ERR-EU-012": (
                "CP Q3 output VAT is €847,200 but should be €874,200 "
                "(transposed digits — discoverable by comparing against Q1/Q2/Q4 "
                "which all show €874,200)"
            ),
        },
        scoring_hints={
            "correctness": (
                "All 41 files indexed with correct jurisdiction tags; "
                "all 8 red flags identified with jurisdiction-specific context"
            ),
            "completeness": (
                "All 72 checklist items mapped; all 6 missing document categories "
                "flagged; ERR-EU-012 detected by cross-referencing VAT returns"
            ),
            "format_compliance": (
                "Document index includes Jurisdiction column; gap analysis includes "
                "Jurisdiction and deal-criticality; memo is jurisdiction-aware"
            ),
            "robustness": (
                "Jurisdiction-specific characterization of each red flag "
                "(German labor law, French IP law, GDPR, BetrVG); "
                "missing documents prioritized by legal requirement vs. nice-to-have"
            ),
            "communication": (
                "Professional memo with clear prioritization; "
                "jurisdiction-specific recommendations; "
                "distinguishes mandatory requirements from best practices"
            ),
        },
        scenario_pack="cascade_europe_ifrs",
        judgment_traps=[
            {
                "trap": "US-framing trap",
                "description": (
                    "Agent treats KvK extract as 'articles of incorporation' or "
                    "tries to find 'bylaws' — applying US corporate document taxonomy"
                ),
            },
            {
                "trap": "GDPR underweight",
                "description": (
                    "Agent identifies missing DPAs but does not flag as deal-critical "
                    "(potential fines up to 4% of global turnover under Art. 83 GDPR)"
                ),
            },
            {
                "trap": "Works council omission",
                "description": (
                    "Agent fails to flag Betriebsrat consultation under §111-113 BetrVG "
                    "as a mandatory deal process requirement"
                ),
            },
            {
                "trap": "Betriebsprüfung dismissal",
                "description": (
                    "Agent treats open German tax audit as routine and non-material — "
                    "open audit covering 2 years represents unquantified tax liability"
                ),
            },
            {
                "trap": "French IP law ignorance",
                "description": (
                    "Agent flags missing IP assignments without noting French-specific "
                    "déclaration d'invention framework under Code de la propriété intellectuelle"
                ),
            },
            {
                "trap": "UK post-Brexit blind spot",
                "description": (
                    "Agent does not flag absence of customs/trade compliance documentation "
                    "for CD, or treats CD as an EU entity for VAT/customs purposes"
                ),
            },
        ],
    )


# ═══════════════════════════════════════════════════════════════════════════
# Public entry point
# ═══════════════════════════════════════════════════════════════════════════


def emit_tc12_eu(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
    **kwargs: object,
) -> None:
    """Write all TC-12-EU files to *output_dir*."""
    # 01_corporate (8 files)
    _write_kvk_extract(output_dir, canaries, manifest)
    _write_articles_of_association(output_dir, canaries, manifest)
    _write_handelsregister(output_dir, canaries, manifest)
    _write_kbis_extract(output_dir, canaries, manifest)
    _write_companies_house(output_dir, canaries, manifest)
    _write_board_minutes_eu(output_dir, canaries, manifest, 2024, "tc12eu_ce_board_minutes_2024")
    _write_board_minutes_eu(output_dir, canaries, manifest, 2025, "tc12eu_ce_board_minutes_2025")
    _write_group_org_chart(output_dir, canaries, manifest)

    # 02_financial (5 files)
    _write_ifrs_financials_pdf(output_dir, canaries, manifest, 2023, "tc12eu_group_ifrs_financials_fy2023")
    _write_ifrs_financials_pdf(output_dir, canaries, manifest, 2024, "tc12eu_group_ifrs_financials_fy2024")
    _write_management_accounts_xlsx(output_dir, canaries, manifest)
    _write_budget_xlsx(output_dir, canaries, manifest)
    _write_debt_schedule_xlsx(output_dir, canaries, manifest)

    # 03_legal (6 files)
    _write_customer_agreement_autohaus(output_dir, canaries, manifest)
    _write_customer_agreement_renault(output_dir, canaries, manifest)
    _write_supplier_agreement_thyssenkrupp(output_dir, canaries, manifest)
    _write_pending_litigation_eu(output_dir, canaries, manifest)
    _write_ip_assignments_eu(output_dir, canaries, manifest)
    _write_insurance_policies_eu(output_dir, canaries, manifest)

    # 04_hr (6 files)
    _write_employee_census_eu(output_dir, canaries, manifest)
    _write_benefits_summary_eu(output_dir, canaries, manifest)
    _write_key_employee_agreement_eu(output_dir, canaries, manifest, 0)
    _write_key_employee_agreement_eu(output_dir, canaries, manifest, 1)
    _write_key_employee_agreement_eu(output_dir, canaries, manifest, 2)
    _write_org_chart_detailed_eu(output_dir, canaries, manifest)

    # 05_tax (7 files)
    _write_cit_return_cp_2023(output_dir, canaries, manifest)
    _write_cit_return_cp_2024(output_dir, canaries, manifest)
    _write_liasse_fiscale_cm(output_dir, canaries, manifest)
    _write_ct600_cd(output_dir, canaries, manifest)
    _write_vpb_aangifte_ce(output_dir, canaries, manifest)
    _write_vat_returns_summary_eu(output_dir, canaries, errors, manifest)
    _write_tax_notices_eu(output_dir, canaries, manifest)

    # 06_operations (4 files)
    _write_facility_leases_eu(output_dir, canaries, manifest)
    _write_equipment_list_eu(output_dir, canaries, manifest)
    _write_customer_list_eu(output_dir, canaries, manifest)
    _write_vendor_list_eu(output_dir, canaries, manifest)

    # 07_technology (3 files)
    _write_patent_portfolio_eu(output_dir, canaries, manifest)
    _write_software_licenses_eu(output_dir, canaries, manifest)
    _write_it_infrastructure_eu(output_dir, canaries, manifest)

    # 08_compliance (2 files)
    _write_gdpr_register(output_dir, canaries, manifest)
    _write_works_council_agreements(output_dir, canaries, manifest)

    # DD Checklist
    _write_dd_checklist_eu(output_dir, canaries, manifest)

    # Prompt and expected behavior
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
