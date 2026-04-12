"""Formatter: TC-21 — Combined Diligence Findings Memo.

Emits:
- test_cases/TC-21/input_files/legal/
    3 high-risk contract summaries (curated from TC-19)
    2 amendment/side letters (curated from TC-19)
- test_cases/TC-21/input_files/hr/
    2 docx employment agreements (key executives with issues)
    1 xlsx severance schedule
    1 xlsx retention schedule
- test_cases/TC-21/input_files/management_qa_summary.docx
    Management Q&A summary
- test_cases/TC-21/input_files/diligence_request_tracker.xlsx
    Combined legal + HR diligence request tracker
- test_cases/TC-21/prompt.md
- test_cases/TC-21/expected_behavior.md

Judgment traps:
  - Top legal risk identification (change-of-control: LDI-001)
  - Top HR exposure identification (CEO golden parachute: SEV-001)
  - Unresolved diligence request (EA-006 missing executed copy: DR-006)
  - Management-statement vs source-backed finding distinction
  - No definitive legal advice language

Uses the canonical model — never hardcodes numbers.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import openpyxl
from docx import Document

from generator.canaries import CanaryRegistry, embed_canary_docx, embed_canary_xlsx
from generator.errors import ErrorRegistry
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel
from generator.model.hr_diligence import (
    DILIGENCE_REQUESTS,
    EMPLOYMENT_AGREEMENTS,
    SEVERANCE_EXPOSURES,
    open_diligence_requests,
    total_retention_awards,
    total_severance_exposure,
)
from generator.model.legal import (
    CONTRACT_AMENDMENTS,
    LEGAL_CONTRACTS,
    LEGAL_DILIGENCE_ISSUES,
    amendments_for_contract,
    clauses_for_contract,
    issues_by_severity,
    missing_consent_issues,
    stale_summary_issues,
)
from generator.writers import (
    new_workbook,
    save_docx_deterministic,
    save_xlsx_deterministic,
    write_header_row,
)
from generator.writers.hr import (
    write_employment_agreement,
    write_retention_schedule,
    write_severance_schedule,
)
from generator.writers.legal import (
    write_amendment,
    write_contract_summary,
)

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-21"
_INPUT_DIR = f"test_cases/{_TC}/input_files"

# Curated contract IDs — the high-risk contracts most relevant to a findings memo
_CURATED_CONTRACT_IDS = ["LCTR-001", "LCTR-003", "LCTR-004"]

# Curated amendment IDs — amendments that create contradictions or scope changes
_CURATED_AMENDMENT_IDS = ["AMD-002", "AMD-003"]

# Curated employment agreement IDs — key executives with notable issues
_CURATED_AGREEMENT_IDS = ["EA-001", "EA-006"]

# Canary file keys
_CANARY_KEYS: list[str] = sorted([
    # Curated legal contracts
    "tc21_contract_lctr_001",
    "tc21_contract_lctr_003",
    "tc21_contract_lctr_004",
    # Curated amendments
    "tc21_amendment_amd_002",
    "tc21_amendment_amd_003",
    # Curated employment agreements
    "tc21_agreement_ea_001",
    "tc21_agreement_ea_006",
    # HR schedules
    "tc21_severance_schedule",
    "tc21_retention_schedule",
    # Management Q&A summary
    "tc21_management_qa_summary",
    # Combined diligence request tracker
    "tc21_diligence_request_tracker",
])


# ── File writers ─────────────────────────────────────────────────────────────


def _write_curated_contracts(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write curated high-risk contract summaries into TC-21/input_files/legal/."""
    legal_dir = output_dir / _INPUT_DIR / "legal"
    legal_dir.mkdir(parents=True, exist_ok=True)

    contract_map = {c.contract_id: c for c in LEGAL_CONTRACTS}
    for cid in _CURATED_CONTRACT_IDS:
        contract = contract_map[cid]
        ckey = f"tc21_contract_{cid.lower().replace('-', '_')}"
        clauses = clauses_for_contract(cid)
        amendments = amendments_for_contract(cid)
        fname = f"contract_{cid.lower()}.docx"
        location = write_contract_summary(
            contract, clauses, amendments,
            legal_dir / fname, canaries, ckey,
        )
        rel_path = f"{_INPUT_DIR}/legal/{fname}"
        canaries.set_location(ckey, rel_path, location)
        manifest.register(rel_path, "docx", test_cases=[_TC])


def _write_curated_amendments(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write curated amendments into TC-21/input_files/legal/."""
    legal_dir = output_dir / _INPUT_DIR / "legal"
    legal_dir.mkdir(parents=True, exist_ok=True)

    contract_map = {c.contract_id: c for c in LEGAL_CONTRACTS}
    for amd in CONTRACT_AMENDMENTS:
        if amd.amendment_id not in _CURATED_AMENDMENT_IDS:
            continue
        ckey = f"tc21_amendment_{amd.amendment_id.lower().replace('-', '_')}"
        contract = contract_map[amd.contract_id]
        fname = f"amendment_{amd.amendment_id.lower()}.docx"
        location = write_amendment(amd, contract, legal_dir / fname, canaries, ckey)
        rel_path = f"{_INPUT_DIR}/legal/{fname}"
        canaries.set_location(ckey, rel_path, location)
        manifest.register(rel_path, "docx", test_cases=[_TC])


def _write_curated_agreements(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write curated employment agreements into TC-21/input_files/hr/."""
    hr_dir = output_dir / _INPUT_DIR / "hr"
    hr_dir.mkdir(parents=True, exist_ok=True)

    agreement_map = {ea.agreement_id: ea for ea in EMPLOYMENT_AGREEMENTS}
    for eaid in _CURATED_AGREEMENT_IDS:
        ea = agreement_map[eaid]
        ckey = f"tc21_agreement_{eaid.lower().replace('-', '_')}"
        fname = f"agreement_{eaid.lower()}.docx"
        location = write_employment_agreement(ea, hr_dir / fname, canaries, ckey)
        rel_path = f"{_INPUT_DIR}/hr/{fname}"
        canaries.set_location(ckey, rel_path, location)
        manifest.register(rel_path, "docx", test_cases=[_TC])


def _write_severance(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write severance schedule into TC-21/input_files/hr/."""
    xlsx_path = output_dir / _INPUT_DIR / "hr" / "severance_schedule.xlsx"
    xlsx_path.parent.mkdir(parents=True, exist_ok=True)

    ckey = "tc21_severance_schedule"
    location = write_severance_schedule(xlsx_path, canaries, ckey)
    rel_path = f"{_INPUT_DIR}/hr/severance_schedule.xlsx"
    canaries.set_location(ckey, rel_path, location)
    manifest.register(rel_path, "xlsx", test_cases=[_TC])


def _write_retention(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write retention schedule into TC-21/input_files/hr/."""
    xlsx_path = output_dir / _INPUT_DIR / "hr" / "retention_schedule.xlsx"
    xlsx_path.parent.mkdir(parents=True, exist_ok=True)

    ckey = "tc21_retention_schedule"
    location = write_retention_schedule(xlsx_path, canaries, ckey)
    rel_path = f"{_INPUT_DIR}/hr/retention_schedule.xlsx"
    canaries.set_location(ckey, rel_path, location)
    manifest.register(rel_path, "xlsx", test_cases=[_TC])


def _write_management_qa(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write management Q&A summary document.

    This is a new document for TC-21 that summarizes management responses
    to diligence questions. It includes some management statements that
    should be distinguished from source-backed findings.
    """
    doc_path = output_dir / _INPUT_DIR / "management_qa_summary.docx"
    doc_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    ckey = "tc21_management_qa_summary"
    canary = canaries.canary_for(ckey)
    location = embed_canary_docx(doc, canary)

    doc.add_heading("Management Q&A Summary", level=1)
    doc.add_paragraph(
        "The following summarizes management's responses to selected due "
        "diligence questions raised during the legal and HR workstreams. "
        "These responses have not been independently verified against "
        "primary source documents."
    )

    # Legal Q&A
    doc.add_heading("Legal Diligence Responses", level=2)

    doc.add_heading("Q: Change-of-control provisions", level=3)
    doc.add_paragraph(
        "Management confirms that the Acme Manufacturing contract (LCTR-001) "
        "includes a change-of-control termination right in Section 14.2. "
        "Management states that informal discussions with Acme suggest they "
        "would not exercise the termination right, but no written waiver or "
        "consent has been obtained."
    )

    doc.add_heading("Q: MFN pricing obligations", level=3)
    doc.add_paragraph(
        "Management states that the TechAlloy Systems (LCTR-003) MFN clause "
        "requires Cascade to match any better pricing offered to comparable "
        "customers. Management characterizes this as a standard provision "
        "that has not been triggered to date."
    )
    # NOTE: This understates the current position — AMD-002 raised the
    # threshold to 5% per-kg, which management does not mention here.

    doc.add_heading("Q: Exclusivity arrangements", level=3)
    doc.add_paragraph(
        "Management confirms that NextGen Composites (LCTR-004) has an "
        "exclusivity arrangement covering alloy specifications for defense "
        "applications. Management describes the scope as limited to the "
        "original contract terms."
    )
    # NOTE: AMD-003 expanded exclusivity to thermal barrier coatings,
    # which management does not disclose here.

    doc.add_heading("Q: Outstanding legal matters", level=3)
    doc.add_paragraph(
        "Management represents that there are no material pending or "
        "threatened litigation matters. Management notes that all material "
        "contracts are in good standing."
    )

    # HR Q&A
    doc.add_heading("HR Diligence Responses", level=2)

    doc.add_heading("Q: Executive severance obligations", level=3)
    doc.add_paragraph(
        f"Management confirms that the CEO (Robert J. Cascade) has a "
        f"change-of-control provision providing for a payment of "
        f"${int(SEVERANCE_EXPOSURES[0].estimated_payout):,} "
        f"({SEVERANCE_EXPOSURES[0].severance_multiplier}x base salary). "
        f"Management characterizes this as a board-approved retention "
        f"mechanism consistent with market practice."
    )

    doc.add_heading("Q: Missing employment agreements", level=3)
    doc.add_paragraph(
        "Management acknowledges that the executed copy of Dr. Anika Patel's "
        "employment agreement (EA-006) cannot be located. Management states "
        "that the agreement was signed and that the terms in the draft "
        "version are accurate. Management is continuing to search for "
        "the original."
    )

    doc.add_heading("Q: Contractor classification", level=3)
    doc.add_paragraph(
        "Management states that all contractor engagements have been reviewed "
        "by outside counsel within the past 18 months and are appropriately "
        "classified. No formal written opinions were obtained."
    )

    doc.add_heading("Limitations", level=2)
    doc.add_paragraph(
        "The above responses are based on representations by Cascade Industries "
        "management and have not been independently verified. Reliance on these "
        "statements without cross-referencing primary source documents is not "
        "recommended for purposes of the diligence findings memo."
    )

    save_docx_deterministic(doc, doc_path)
    canaries.set_location(ckey, f"{_INPUT_DIR}/management_qa_summary.docx", location)
    manifest.register(f"{_INPUT_DIR}/management_qa_summary.docx", "docx", test_cases=[_TC])


def _write_combined_request_tracker(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write combined legal + HR diligence request tracker."""
    xlsx_path = output_dir / _INPUT_DIR / "diligence_request_tracker.xlsx"
    xlsx_path.parent.mkdir(parents=True, exist_ok=True)

    wb = new_workbook()
    ckey = "tc21_diligence_request_tracker"
    canary = canaries.canary_for(ckey)
    location = embed_canary_xlsx(wb, canary)

    # Sheet 1: Legal requests (from LEGAL_DILIGENCE_ISSUES)
    ws_legal = wb.active
    ws_legal.title = "Legal Requests"

    legal_headers = [
        "Issue ID", "Contract ID", "Issue Type", "Severity",
        "Description", "Recommended Action", "Status",
    ]
    write_header_row(ws_legal, legal_headers)

    for row_idx, issue in enumerate(LEGAL_DILIGENCE_ISSUES, 2):
        ws_legal.cell(row=row_idx, column=1, value=issue.issue_id)
        ws_legal.cell(row=row_idx, column=2, value=issue.contract_id)
        ws_legal.cell(row=row_idx, column=3, value=issue.issue_type)
        ws_legal.cell(row=row_idx, column=4, value=issue.severity)
        ws_legal.cell(row=row_idx, column=5, value=issue.description)
        ws_legal.cell(row=row_idx, column=6, value=issue.recommended_action)
        ws_legal.cell(row=row_idx, column=7, value="Open")

    for col_idx, header in enumerate(legal_headers, 1):
        ws_legal.column_dimensions[
            openpyxl.utils.get_column_letter(col_idx)
        ].width = max(len(header) + 2, 15)

    # Sheet 2: HR requests (from DILIGENCE_REQUESTS)
    ws_hr = wb.create_sheet("HR Requests")

    hr_headers = [
        "Request ID", "Category", "Description", "Requested From",
        "Status", "Due Date", "Received Date", "Source References",
    ]
    write_header_row(ws_hr, hr_headers)

    for row_idx, dr in enumerate(DILIGENCE_REQUESTS, 2):
        ws_hr.cell(row=row_idx, column=1, value=dr.request_id)
        ws_hr.cell(row=row_idx, column=2, value=dr.category)
        ws_hr.cell(row=row_idx, column=3, value=dr.description)
        ws_hr.cell(row=row_idx, column=4, value=dr.requested_from)
        ws_hr.cell(row=row_idx, column=5, value=dr.status)
        ws_hr.cell(row=row_idx, column=6, value=dr.due_date.isoformat())
        ws_hr.cell(
            row=row_idx, column=7,
            value=dr.received_date.isoformat() if dr.received_date else "Not received",
        )
        ws_hr.cell(row=row_idx, column=8, value=", ".join(dr.source_refs))

    for col_idx, header in enumerate(hr_headers, 1):
        ws_hr.column_dimensions[
            openpyxl.utils.get_column_letter(col_idx)
        ].width = max(len(header) + 2, 15)

    save_xlsx_deterministic(wb, xlsx_path)
    rel_path = f"{_INPUT_DIR}/diligence_request_tracker.xlsx"
    canaries.set_location(ckey, rel_path, location)
    manifest.register(rel_path, "xlsx", test_cases=[_TC])


# ── Prompt & Expected Behavior ───────────────────────────────────────────────


def _write_prompt(output_dir: Path) -> None:
    """Write test_cases/TC-21/prompt.md per spec."""
    text = """\
Prepare a client-ready combined diligence findings memo for Cascade
Industries in connection with a potential acquisition.

Using the curated legal contracts, amendments, employment agreements,
severance and retention schedules, management Q&A summary, and
diligence request tracker provided:

1. **Executive Summary**: Provide a concise overview of the key findings
   from both the legal and HR diligence workstreams. Highlight the
   most material risks and their potential impact on the transaction.

2. **Legal Findings**: Summarize the key legal risks identified in the
   contract portfolio, including:
   - Change-of-control provisions and consent requirements
   - Contradictions between management representations and primary
     source documents (contracts, amendments)
   - Scope of exclusivity arrangements and any undisclosed expansions
   - Government subcontract assignment and novation requirements

3. **HR Findings**: Summarize the key HR-related risks, including:
   - Executive severance exposures and change-of-control provisions
   - Retention/severance interaction and double-count avoidance
   - Missing executed employment agreements and their implications
   - Contractor classification risk signals

4. **Unresolved Requests**: List all diligence requests that remain
   open, partially fulfilled, or not received. Note the implications
   of each outstanding item for the completeness of the findings.

5. **Assumptions and Limitations**: State all assumptions made in
   preparing the memo. Note which findings are based on management
   representations versus independently verified source documents.
   Include a clear disclaimer that this memo does not constitute
   legal advice.

6. **Citations**: For every material finding, cite the specific source
   document (contract ID, amendment ID, agreement ID, schedule
   reference, or management Q&A response). Distinguish between
   findings supported by primary source documents and those based
   solely on management statements.
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write test_cases/TC-21/expected_behavior.md per spec."""
    # Pull numbers from canonical model
    ceo_sev = SEVERANCE_EXPOSURES[0]  # SEV-001
    total_sev = total_severance_exposure()
    total_ret = total_retention_awards()
    open_reqs = open_diligence_requests()

    text = f"""\
# TC-21: Combined Diligence Findings Memo — Expected Behavior

## Executive Summary Requirements
- The memo must synthesize findings from both legal (TC-19) and HR
  (TC-20) diligence workstreams into a unified narrative.
- The executive summary should highlight the top legal risk and top
  HR exposure by materiality.

## Top Legal Risk: Change-of-Control (LDI-001)
- Acme Manufacturing (LCTR-001) has a change-of-control termination
  clause in Section 14.2 — highest-severity legal risk.
- Revenue at risk: approximately $36M/year (~18% of consolidated revenue).
- No written waiver or consent has been obtained.
- Management's informal indication that Acme would not exercise the
  right is not a substitute for a written waiver.
- The agent must cite the contract (LCTR-001) as primary source, not
  just the management Q&A response.

## Top HR Exposure: CEO Golden Parachute (SEV-001)
- Robert J. Cascade has a 3x change-of-control provision.
- Base salary ${ceo_sev.base_salary:,} x {ceo_sev.severance_multiplier}x = \
${int(ceo_sev.estimated_payout):,}.
- Total severance exposure across all executives: ${int(total_sev):,}.
- Total retention awards: ${int(total_ret):,}.

## Management Statements vs Source-Backed Findings
- The management Q&A understates the MFN clause position: AMD-002
  raised the MFN threshold to 5% per-kg, which management omits.
- The management Q&A describes exclusivity scope using original terms;
  AMD-003 expanded scope to thermal barrier coatings.
- The agent must distinguish between management representations and
  findings verified against primary source documents.
- Findings relying solely on management statements should be flagged
  with appropriate caveats.

## Unresolved Diligence Requests
- {len(open_reqs)} diligence request(s) remain unresolved:
  DR-002 (EA-006 executed copy — partial), DR-006 (EA-006 follow-up —
  not received).
- The missing executed agreement (EA-006, Dr. Anika Patel) affects IP
  assignment enforceability and non-compete validity.
- The agent must flag unresolved requests and their implications.

## No Legal Advice Language
- The memo must include a clear disclaimer that it does not constitute
  legal advice.
- Findings should use language such as "appears to," "may indicate,"
  and "subject to further review" rather than definitive legal conclusions.
- Contractor classification signals must be presented as requiring
  further investigation, not as final determinations.

## Source Citation Requirements
- Every material finding must cite a specific source document.
- Legal findings must reference contract IDs (LCTR-NNN) and amendment
  IDs (AMD-NNN), not just the management Q&A summary.
- HR findings must reference agreement IDs (EA-NNN), severance IDs
  (SEV-NNN), and retention IDs (RET-NNN).
- The agent must clearly indicate when a finding is based on management
  representations versus independently verified sources.

## Scoring Focus
Grade on:
- Synthesis quality (unified narrative across legal and HR)
- Identification of top legal risk (LDI-001) and top HR exposure (SEV-001)
- Detection of management statement / source document contradictions
- Unresolved request tracking and implications
- Professional language (no definitive legal advice)
- Source citation completeness and accuracy
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold Standard ────────────────────────────────────────────────────────────


@register_gold("TC-21")
def _tc21_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """Build the TC-21 gold standard."""
    # Build canary verification from all TC-21 file keys
    canary_verification: dict[str, str] = {}

    # Curated legal contracts
    for cid in _CURATED_CONTRACT_IDS:
        ckey = f"tc21_contract_{cid.lower().replace('-', '_')}"
        label = f"read_contract_{cid.lower().replace('-', '_')}"
        canary_verification[label] = canaries.canary_for(ckey)

    # Curated amendments
    for aid in _CURATED_AMENDMENT_IDS:
        ckey = f"tc21_amendment_{aid.lower().replace('-', '_')}"
        label = f"read_amendment_{aid.lower().replace('-', '_')}"
        canary_verification[label] = canaries.canary_for(ckey)

    # Curated employment agreements
    for eaid in _CURATED_AGREEMENT_IDS:
        ckey = f"tc21_agreement_{eaid.lower().replace('-', '_')}"
        label = f"read_agreement_{eaid.lower().replace('-', '_')}"
        canary_verification[label] = canaries.canary_for(ckey)

    # Schedules and trackers
    canary_verification["read_severance_schedule"] = canaries.canary_for(
        "tc21_severance_schedule"
    )
    canary_verification["read_retention_schedule"] = canaries.canary_for(
        "tc21_retention_schedule"
    )
    canary_verification["read_management_qa_summary"] = canaries.canary_for(
        "tc21_management_qa_summary"
    )
    canary_verification["read_diligence_request_tracker"] = canaries.canary_for(
        "tc21_diligence_request_tracker"
    )

    # Compute expected values from canonical model
    open_reqs = open_diligence_requests()
    high_issues = issues_by_severity("high")
    stale_issues = stale_summary_issues()
    consent_issues = missing_consent_issues()

    return GoldStandard(
        test_case=_TC,
        expected_outputs={
            "output_format": "combined_findings_memo",
            "required_sections": [
                "executive_summary",
                "legal_findings",
                "hr_findings",
                "unresolved_requests",
                "assumptions_and_limitations",
                "citations",
            ],
            "curated_contract_count": len(_CURATED_CONTRACT_IDS),
            "curated_amendment_count": len(_CURATED_AMENDMENT_IDS),
            "curated_agreement_count": len(_CURATED_AGREEMENT_IDS),
            "total_severance_exposure": str(total_severance_exposure()),
            "total_retention_awards": str(total_retention_awards()),
            "top_legal_risk": {
                "issue_id": "LDI-001",
                "contract_id": "LCTR-001",
                "description": "Change-of-control termination clause",
                "severity": "high",
            },
            "top_hr_exposure": {
                "exposure_id": "SEV-001",
                "employee_name": "Robert J. Cascade",
                "estimated_payout": str(SEVERANCE_EXPOSURES[0].estimated_payout),
            },
            "management_source_contradictions": [
                issue.issue_id for issue in stale_issues
            ],
            "unresolved_requests": [dr.request_id for dr in open_reqs],
            "consent_required": [
                issue.contract_id for issue in consent_issues
            ],
            "high_severity_legal_issues": [
                issue.issue_id for issue in high_issues
            ],
        },
        canary_verification=canary_verification,
        error_detection={},
        scoring_hints={
            "synthesis": (
                "Memo must integrate legal and HR findings into a unified "
                "narrative with a coherent executive summary"
            ),
            "top_risks": (
                "Top legal risk (LDI-001: change-of-control, ~$36M revenue) "
                "and top HR exposure (SEV-001: CEO golden parachute, "
                f"${int(SEVERANCE_EXPOSURES[0].estimated_payout):,}) "
                "must be prominently identified"
            ),
            "source_distinction": (
                "Management statements must be distinguished from "
                "source-backed findings; MFN and exclusivity contradictions "
                "must be flagged"
            ),
            "professional_language": (
                "No definitive legal advice language; findings use "
                "'appears to,' 'may indicate,' 'subject to further review'; "
                "contractor classification presented as signals"
            ),
            "unresolved_tracking": (
                "All unresolved diligence requests listed with implications; "
                "missing EA-006 executed copy flagged as affecting "
                "IP assignment and non-compete enforceability"
            ),
        },
        scenario_pack="ma_legal_hr_diligence",
        service_line="advisory",
        evidence_expectations={
            "finding_top_legal_risk": {
                "required_sources": ["tc21_contract_lctr_001"],
                "primary_source_required": True,
                "acceptable_terms": [
                    "change of control", "termination", "consent",
                    "revenue at risk",
                ],
            },
            "finding_top_hr_exposure": {
                "required_sources": ["tc21_agreement_ea_001"],
                "primary_source_required": True,
                "acceptable_terms": [
                    "golden parachute", "3x base salary", "change of control",
                    "severance",
                ],
            },
            "finding_mfn_contradiction": {
                "required_sources": [
                    "tc21_contract_lctr_003",
                    "tc21_amendment_amd_002",
                    "tc21_management_qa_summary",
                ],
                "primary_source_required": True,
                "acceptable_terms": [
                    "most-favored-nation", "MFN", "5%", "threshold",
                    "contradiction",
                ],
            },
            "finding_exclusivity_expansion": {
                "required_sources": [
                    "tc21_contract_lctr_004",
                    "tc21_amendment_amd_003",
                    "tc21_management_qa_summary",
                ],
                "primary_source_required": True,
                "acceptable_terms": [
                    "exclusivity", "thermal barrier coatings",
                    "scope expansion", "undisclosed",
                ],
            },
            "finding_unresolved_requests": {
                "required_sources": ["tc21_diligence_request_tracker"],
                "primary_source_required": True,
                "acceptable_terms": [
                    "unresolved", "outstanding", "not received",
                    "partial", "EA-006",
                ],
            },
            "finding_missing_agreement": {
                "required_sources": ["tc21_agreement_ea_006"],
                "primary_source_required": True,
                "acceptable_terms": [
                    "missing", "not executed", "draft", "IP assignment",
                    "non-compete",
                ],
            },
        },
        judgment_traps=[
            {
                "trap_id": "JT-TC21-001",
                "trap_type": "missing_evidence",
                "expected_response": "flag",
                "description": (
                    "The top legal risk is the Acme Manufacturing (LCTR-001) "
                    "change-of-control termination clause (LDI-001). Revenue "
                    "at risk is ~$36M/year (~18% consolidated). The agent "
                    "must identify this as the highest-severity legal finding "
                    "and cite the contract as primary source, not just the "
                    "management Q&A."
                ),
            },
            {
                "trap_id": "JT-TC21-002",
                "trap_type": "missing_evidence",
                "expected_response": "flag",
                "description": (
                    "The top HR exposure is the CEO golden parachute "
                    f"(SEV-001: ${int(SEVERANCE_EXPOSURES[0].estimated_payout):,}). "
                    "The agent must identify this as the largest single "
                    "severance exposure and cite EA-001."
                ),
            },
            {
                "trap_id": "JT-TC21-003",
                "trap_type": "missing_evidence",
                "expected_response": "flag",
                "description": (
                    "At least one diligence request is unresolved: DR-006 "
                    "(executed copy of EA-006, Dr. Patel). The agent must "
                    "list unresolved requests and note the implications for "
                    "IP assignment and non-compete enforceability."
                ),
            },
            {
                "trap_id": "JT-TC21-004",
                "trap_type": "summary_contradiction",
                "expected_response": "flag",
                "description": (
                    "The management Q&A understates the MFN position "
                    "(LCTR-003) and omits the exclusivity expansion "
                    "(LCTR-004, AMD-003). The agent must distinguish "
                    "management statements from source-backed findings and "
                    "flag the discrepancies."
                ),
            },
            {
                "trap_id": "JT-TC21-005",
                "trap_type": "overconfident_conclusion",
                "expected_response": "caveat",
                "description": (
                    "The findings memo must not contain definitive legal "
                    "advice. The agent should use language such as 'appears "
                    "to,' 'may indicate,' and 'subject to further review.' "
                    "Contractor classification signals must be presented as "
                    "requiring investigation, not as final determinations."
                ),
            },
        ],
        source_requirements={
            "minimum_sources_per_finding": 1,
            "primary_source_required_for_high_risk": True,
            "management_summary_not_sufficient_alone": True,
        },
    )


# ── Public entry point ──────────────────────────────────────────────────────


def emit_tc21(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
    **kwargs: object,
) -> None:
    """Write all TC-21 files to *output_dir*."""
    _write_curated_contracts(output_dir, canaries, manifest)
    _write_curated_amendments(output_dir, canaries, manifest)
    _write_curated_agreements(output_dir, canaries, manifest)
    _write_severance(output_dir, canaries, manifest)
    _write_retention(output_dir, canaries, manifest)
    _write_management_qa(output_dir, canaries, manifest)
    _write_combined_request_tracker(output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
