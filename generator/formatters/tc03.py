"""Formatter: TC-03 — Substantive Analytical Procedures — Revenue (Audit, Complex).

Emits:
- test_cases/TC-03/input_files/revenue_by_product_monthly_fy2024_fy2025.xlsx
  24 months of revenue by product line (6 lines) with unit volumes and ASPs
- test_cases/TC-03/input_files/industry_benchmark_report.pdf
  12-page synthetic industry report; benchmarks on pages 4, 7, 11
- test_cases/TC-03/input_files/management_rep_letter.docx
  Management representation asserting "~8% growth" (actual: 9.2%)
- test_cases/TC-03/prompt.md
- test_cases/TC-03/expected_behavior.md
- gold_standards/TC-03_gold.json

Planted errors:
- ERR-009 (stale_data): industry benchmark PDF shows FY2024 growth rate
  where FY2025's should be (Industrial Manufacturing sector, page 4)
- ERR-018 (classification_error): one revenue row labels Specialty Coatings
  as Advanced Composites (March 2025)
Uses the canonical model — never hardcodes numbers.
"""

from __future__ import annotations

import datetime
import io
import random
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import docx
import openpyxl
from docx.shared import Pt
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from generator.canaries import CanaryRegistry, embed_canary_docx, embed_canary_xlsx
from generator.errors import (
    ErrorRegistry,
    PlantedError,
    classification_error,
    stale_data,
)
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel
from generator.model.revenue import (
    PRODUCT_LINES,
    validate_consolidated_growth,
    validate_product_line_growth,
)

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-03"
_INPUT_DIR = f"test_cases/{_TC}/input_files"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)

# Synthetic unit volumes and ASPs per product line for FY2025.
# Revenue = units × ASP; we back-derive units from actual revenue.
# These multipliers produce reasonable-looking volumes.
_ASP_TABLE: dict[str, dict[str, Any]] = {
    "Industrial Parts": {"base_asp": Decimal("2500"), "asp_unit": "per batch"},
    "Custom Machining": {"base_asp": Decimal("8500"), "asp_unit": "per order"},
    "Advanced Composites": {"base_asp": Decimal("12000"), "asp_unit": "per shipment"},
    "Specialty Coatings": {"base_asp": Decimal("4500"), "asp_unit": "per lot"},
    "Warehousing Services": {"base_asp": Decimal("1800"), "asp_unit": "per contract-month"},
    "Freight & Logistics": {"base_asp": Decimal("950"), "asp_unit": "per shipment"},
}

# Industry benchmark growth rates — these are what the PDF will show
# on pages 4, 7, and 11.  Deliberately close but not identical to
# Cascade's actual rates, to force the agent to compare carefully.
_INDUSTRY_BENCHMARKS: dict[str, dict[str, Any]] = {
    "Industrial Manufacturing": {
        "growth_rate": "4.2%",
        "margin_range": "28%–38%",
        "market_size": "$142B",
        "page": 4,
    },
    "Advanced Materials & Composites": {
        "growth_rate": "12.8%",
        "margin_range": "42%–55%",
        "market_size": "$67B",
        "page": 7,
    },
    "Logistics & Distribution": {
        "growth_rate": "7.5%",
        "margin_range": "14%–22%",
        "market_size": "$298B",
        "page": 11,
    },
}


# ── Helpers ──────────────────────────────────────────────────────────────────


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)

    wb.properties.created = _FIXED_DATETIME
    wb.properties.modified = _FIXED_DATETIME

    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    fixed_date_time = (2025, 3, 15, 9, 0, 0)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=fixed_date_time)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)

    buf = io.BytesIO()
    doc.save(buf)

    fixed_date_time = (2025, 3, 15, 9, 0, 0)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=fixed_date_time)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _whole_dollars(d: Decimal) -> int:
    """Round a Decimal to whole dollars."""
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


def _compute_unit_volumes(
    revenue: Decimal,
    base_asp: Decimal,
    rng: random.Random,
) -> tuple[int, Decimal]:
    """Derive unit volume and actual ASP from revenue and base ASP.

    Applies ±5% perturbation to the ASP, then computes units = revenue / asp.
    Returns (units, actual_asp).
    """
    perturbation = Decimal(str(1 + rng.uniform(-0.05, 0.05)))
    actual_asp = (base_asp * perturbation).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    if actual_asp <= 0:
        actual_asp = base_asp
    units = int(revenue / actual_asp)
    if units <= 0:
        units = 1
    return units, actual_asp


# ── Revenue XLSX ─────────────────────────────────────────────────────────────


def _write_revenue_xlsx(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write revenue_by_product_monthly_fy2024_fy2025.xlsx."""
    file_key = "tc03_revenue_by_product"
    canary_code = canaries.canary_for(file_key)

    wb = openpyxl.Workbook()
    loc = embed_canary_xlsx(wb, canary_code)

    # ── Gather revenue data for FY2024 and FY2025 ────────────────────
    rng = random.Random(42 + 303)  # TC-03 specific sub-seed for ASP perturbation

    # Collect by (product_line, year, month)
    rev_data: dict[tuple[str, int, int], Decimal] = {}
    for rec in model.revenue_records:
        if rec.year in (2024, 2025):
            rev_data[(rec.product_line, rec.year, rec.month)] = rec.revenue

    # Get sorted unique product line names
    pl_names = sorted({rec.product_line for rec in model.revenue_records})

    # ── Sheet: Monthly Revenue ───────────────────────────────────────
    ws = wb.active
    ws.title = "Monthly Revenue"

    header_fill = PatternFill("solid", fgColor="1A3C6E")
    header_font_white = Font(bold=True, size=11, color="FFFFFF")
    number_fmt = '#,##0'
    dollar_fmt = '$#,##0'
    asp_fmt = '$#,##0.00'
    border = Border(
        bottom=Side(style="thin", color="CCCCCC"),
    )

    # Title row
    ws["A1"] = "Cascade Industries, Inc."
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "Revenue by Product Line — Monthly Detail (FY2024–FY2025)"
    ws["A2"].font = Font(bold=True, size=12, color="666666")
    ws.merge_cells("A1:H1")
    ws.merge_cells("A2:H2")

    # Headers in row 4
    headers = ["Year", "Month", "Product Line", "Entity", "Revenue ($)",
               "Unit Volume", "Avg Selling Price", "COGS ($)"]
    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col_idx, value=h)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    # Data rows — sorted deterministically
    row = 5
    for year in (2024, 2025):
        for month in range(1, 13):
            for pl_name in pl_names:
                key = (pl_name, year, month)
                revenue = rev_data.get(key)
                if revenue is None:
                    continue

                # Find the matching MonthlyRevenue for COGS
                cogs = Decimal(0)
                entity_code = ""
                for rec in model.revenue_records:
                    if (rec.product_line == pl_name and rec.year == year
                            and rec.month == month):
                        cogs = rec.cogs
                        entity_code = rec.entity_code
                        break

                # Compute unit volumes and ASP
                asp_info = _ASP_TABLE[pl_name]
                units, actual_asp = _compute_unit_volumes(
                    revenue, asp_info["base_asp"], rng,
                )

                month_name = datetime.date(year, month, 1).strftime("%B")

                # ERR-018: classification_error — misclassify one row
                display_pl_name = pl_name
                if (year == 2025 and month == 3
                        and pl_name == "Specialty Coatings"):
                    correct_pl = "Specialty Coatings"
                    wrong_pl = "Advanced Composites"
                    display_pl_name = classification_error(correct_pl, wrong_pl)
                    errors.add(PlantedError(
                        error_id="ERR-018",
                        file=f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx",
                        location=(
                            "Sheet 'Monthly Revenue', March 2025 row for "
                            "Specialty Coatings, Column C (Product Line)"
                        ),
                        type="classification_error",
                        description=(
                            f"shows {wrong_pl} instead of {correct_pl} for "
                            "March 2025 revenue line item"
                        ),
                        severity="immaterial",
                        which_test_cases_should_catch=["TC-03"],
                    ))

                ws.cell(row=row, column=1, value=year)
                ws.cell(row=row, column=2, value=month_name)
                ws.cell(row=row, column=3, value=display_pl_name)
                ws.cell(row=row, column=4, value=entity_code)

                rev_cell = ws.cell(row=row, column=5, value=_whole_dollars(revenue))
                rev_cell.number_format = dollar_fmt

                vol_cell = ws.cell(row=row, column=6, value=units)
                vol_cell.number_format = number_fmt

                asp_cell = ws.cell(row=row, column=7, value=float(actual_asp))
                asp_cell.number_format = asp_fmt

                cogs_cell = ws.cell(row=row, column=8, value=_whole_dollars(cogs))
                cogs_cell.number_format = dollar_fmt

                for c in range(1, 9):
                    ws.cell(row=row, column=c).border = border

                row += 1

    # Adjust column widths
    widths = [8, 12, 22, 8, 15, 13, 18, 15]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    # ── Sheet: Annual Summary ────────────────────────────────────────
    ws2 = wb.create_sheet("Annual Summary")

    ws2["A1"] = "Annual Revenue Summary by Product Line"
    ws2["A1"].font = Font(bold=True, size=13)
    ws2.merge_cells("A1:F1")

    sum_headers = ["Product Line", "Entity", "FY2024 Revenue",
                   "FY2025 Revenue", "YoY Growth ($)", "YoY Growth (%)"]
    for col_idx, h in enumerate(sum_headers, 1):
        cell = ws2.cell(row=3, column=col_idx, value=h)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    srow = 4
    total_fy24 = Decimal(0)
    total_fy25 = Decimal(0)

    for pl_name in pl_names:
        fy24 = sum(
            rev_data.get((pl_name, 2024, m), Decimal(0))
            for m in range(1, 13)
        )
        fy25 = sum(
            rev_data.get((pl_name, 2025, m), Decimal(0))
            for m in range(1, 13)
        )
        yoy_dollars = fy25 - fy24
        yoy_pct = (yoy_dollars / fy24 * 100) if fy24 != 0 else Decimal(0)

        # Find entity code
        entity_code = ""
        for pl in PRODUCT_LINES:
            if pl.name == pl_name:
                entity_code = pl.entity_code
                break

        ws2.cell(row=srow, column=1, value=pl_name)
        ws2.cell(row=srow, column=2, value=entity_code)

        c = ws2.cell(row=srow, column=3, value=_whole_dollars(fy24))
        c.number_format = dollar_fmt

        c = ws2.cell(row=srow, column=4, value=_whole_dollars(fy25))
        c.number_format = dollar_fmt

        c = ws2.cell(row=srow, column=5, value=_whole_dollars(yoy_dollars))
        c.number_format = dollar_fmt

        c = ws2.cell(row=srow, column=6, value=round(float(yoy_pct), 1))
        c.number_format = '0.0"%"'

        total_fy24 += fy24
        total_fy25 += fy25
        srow += 1

    # Consolidated total row
    ws2.cell(row=srow, column=1, value="CONSOLIDATED TOTAL").font = Font(bold=True)
    ws2.cell(row=srow, column=2, value="ALL")

    c = ws2.cell(row=srow, column=3, value=_whole_dollars(total_fy24))
    c.number_format = dollar_fmt
    c.font = Font(bold=True)

    c = ws2.cell(row=srow, column=4, value=_whole_dollars(total_fy25))
    c.number_format = dollar_fmt
    c.font = Font(bold=True)

    yoy_total = total_fy25 - total_fy24
    c = ws2.cell(row=srow, column=5, value=_whole_dollars(yoy_total))
    c.number_format = dollar_fmt
    c.font = Font(bold=True)

    consol_pct = (yoy_total / total_fy24 * 100) if total_fy24 != 0 else Decimal(0)
    c = ws2.cell(row=srow, column=6, value=round(float(consol_pct), 1))
    c.number_format = '0.0"%"'
    c.font = Font(bold=True)

    sum_widths = [22, 8, 18, 18, 18, 16]
    for i, w in enumerate(sum_widths, 1):
        ws2.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    # ── Save ─────────────────────────────────────────────────────────
    path = output_dir / _INPUT_DIR / "revenue_by_product_monthly_fy2024_fy2025.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, path)

    canaries.set_location(
        file_key,
        f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx",
        loc,
    )
    manifest.register(
        f"{_INPUT_DIR}/revenue_by_product_monthly_fy2024_fy2025.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Industry Benchmark PDF ───────────────────────────────────────────────────


def _write_benchmark_pdf(
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write industry_benchmark_report.pdf — 12-page synthetic report.

    Benchmarks are embedded on pages 4, 7, and 11 per spec.
    """
    file_key = "tc03_industry_benchmark"
    canary_code = canaries.canary_for(file_key)

    path = output_dir / _INPUT_DIR / "industry_benchmark_report.pdf"
    path.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BMTitle", parent=styles["Title"],
        fontSize=18, spaceAfter=12, textColor=colors.HexColor("#1A3C6E"),
    )
    heading_style = ParagraphStyle(
        "BMHeading", parent=styles["Heading1"],
        fontSize=14, spaceAfter=8, textColor=colors.HexColor("#1A3C6E"),
    )
    subheading_style = ParagraphStyle(
        "BMSubHeading", parent=styles["Heading2"],
        fontSize=12, spaceAfter=6, textColor=colors.HexColor("#2D5F8A"),
    )
    body_style = ParagraphStyle(
        "BMBody", parent=styles["Normal"],
        fontSize=10, spaceAfter=6, leading=14,
    )
    small_style = ParagraphStyle(
        "BMSmall", parent=styles["Normal"],
        fontSize=8, textColor=colors.HexColor("#999999"),
    )
    toc_style = ParagraphStyle(
        "BMTOC", parent=styles["Normal"],
        fontSize=10, spaceAfter=4, leftIndent=20,
    )

    story: list = []

    # ── Page 1: Title page ───────────────────────────────────────────
    story.append(Spacer(1, 2 * inch))
    story.append(Paragraph(
        "INDUSTRY BENCHMARK REPORT",
        title_style,
    ))
    story.append(Paragraph(
        "Annual Market Analysis &amp; Performance Benchmarks",
        subheading_style,
    ))
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph(
        "Prepared by: Meridian Research Group<br/>"
        "Publication Date: March 2026<br/>"
        "Report Period: Calendar Year 2025",
        body_style,
    ))
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph(
        "This report provides industry-level growth rates, margin benchmarks, "
        "and market sizing data for key sectors in the diversified industrial "
        "and materials space. Data sourced from public filings, trade association "
        "surveys, and proprietary Meridian models.",
        body_style,
    ))
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph(
        f"Report Reference: MRG-2026-{canary_code}",
        small_style,
    ))

    # ── Page 2: Table of Contents ────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("TABLE OF CONTENTS", heading_style))
    story.append(Spacer(1, 0.3 * inch))

    toc_entries = [
        ("1.", "Executive Summary", "3"),
        ("2.", "Industrial Manufacturing Sector", "4"),
        ("3.", "Methodology &amp; Data Sources", "5"),
        ("4.", "Economic Outlook — Macroeconomic Factors", "6"),
        ("5.", "Advanced Materials &amp; Composites Sector", "7"),
        ("6.", "Supply Chain &amp; Raw Materials Analysis", "8"),
        ("7.", "Labor Market &amp; Workforce Trends", "9"),
        ("8.", "Regulatory &amp; Compliance Environment", "10"),
        ("9.", "Logistics &amp; Distribution Sector", "11"),
        ("10.", "Emerging Technologies &amp; Industry 4.0", "12"),
    ]
    for num, title, pg in toc_entries:
        story.append(Paragraph(
            f"{num} {title} {'.' * (50 - len(title))} {pg}",
            toc_style,
        ))

    # ── Page 3: Executive Summary ────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("1. Executive Summary", heading_style))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "The diversified industrial sector experienced moderate growth in 2025, "
        "driven by continued demand for advanced materials and sustained "
        "infrastructure investment. Key findings include:",
        body_style,
    ))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(
        "&#8226; Traditional manufacturing grew 4.2% year-over-year, in line with "
        "GDP growth and consistent with the prior two-year trend.<br/>"
        "&#8226; The advanced materials segment outpaced the broader market with "
        "12.8% growth, fueled by aerospace, EV, and renewable energy demand.<br/>"
        "&#8226; Logistics and distribution grew 7.5%, supported by e-commerce "
        "volumes and warehouse automation investment.<br/>"
        "&#8226; Overall, companies in the diversified industrial space averaged "
        "6.8% consolidated revenue growth, with wide dispersion between "
        "segments.",
        body_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "Companies reporting consolidated growth above 9% should expect "
        "increased audit scrutiny on revenue recognition policies, particularly "
        "if growth is concentrated in a single segment or driven by non-recurring "
        "contract wins.",
        body_style,
    ))

    # ── Page 4: Industrial Manufacturing ─────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph(
        "2. Industrial Manufacturing Sector",
        heading_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "The industrial manufacturing sector, encompassing precision components, "
        "custom machining, and metal fabrication, recorded steady growth in 2025.",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))

    # Benchmark table — this is page 4 data
    bm_mfg = _INDUSTRY_BENCHMARKS["Industrial Manufacturing"]
    # ERR-009: stale_data — show FY2024's growth rate where FY2025's should be
    correct_mfg_growth = bm_mfg["growth_rate"]  # "4.2%"
    stale_mfg_growth = stale_data("3.8%")        # FY2024 value
    errors.add(PlantedError(
        error_id="ERR-009",
        file=f"{_INPUT_DIR}/industry_benchmark_report.pdf",
        location="Page 4, Industrial Manufacturing table, 'Revenue Growth (YoY)' row, '2025 Value' column",
        type="stale_data",
        description=(
            f"shows {stale_mfg_growth} instead of {correct_mfg_growth} for "
            "Industrial Manufacturing 2025 revenue growth"
        ),
        severity="immaterial",
        which_test_cases_should_catch=["TC-03"],
    ))
    mfg_data = [
        ["Metric", "2025 Value", "2024 Value", "Trend"],
        ["Revenue Growth (YoY)", stale_mfg_growth, "3.8%", "Improving"],
        ["Gross Margin (Median)", "33%", "32%", "Stable"],
        ["Gross Margin (Range)", bm_mfg["margin_range"], "27%–36%", "Widening"],
        ["Market Size (US)", bm_mfg["market_size"], "$136B", "Growing"],
        ["Capacity Utilization", "78.3%", "76.1%", "Improving"],
    ]
    mfg_table = Table(mfg_data, colWidths=[2.2 * inch, 1.5 * inch, 1.2 * inch, 1.2 * inch])
    mfg_table.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 9),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 9),
        ("LINEBELOW", (0, 0), (-1, 0), 1, colors.HexColor("#1A3C6E")),
        ("LINEBELOW", (0, -1), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF4")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(mfg_table)
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph(
        "Companies in the precision components sub-segment generally reported "
        "growth of 3–6%, with higher-margin custom work growing faster than "
        "commodity parts. Firms with strong automation programs saw margin "
        "expansion of 100–200 bps.",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "Key risk factors include raw material price volatility (steel, aluminum), "
        "labor shortages in skilled manufacturing, and increasing competition from "
        "offshore producers. Companies reporting growth significantly above the "
        "4.2% industry average should demonstrate clear operational drivers.",
        body_style,
    ))

    # ── Page 5: Methodology ──────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("3. Methodology &amp; Data Sources", heading_style))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "This report draws on multiple data sources to construct industry benchmarks:",
        body_style,
    ))
    story.append(Paragraph(
        "&#8226; Public financial filings of 342 companies across targeted sectors<br/>"
        "&#8226; Trade association surveys (National Association of Manufacturers, "
        "Material Handling Industry, Council of Supply Chain Management Professionals)<br/>"
        "&#8226; Federal Reserve economic data and Census Bureau manufacturing surveys<br/>"
        "&#8226; Proprietary Meridian Research Group models calibrated against "
        "10 years of historical data",
        body_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "Growth rates are computed on a revenue basis, excluding acquisitions and "
        "divestitures. Margin benchmarks reflect GAAP gross margins unless "
        "otherwise noted. Market size estimates use a bottom-up approach based on "
        "reported revenue of participants weighted by market coverage.",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "Sector classification follows the Meridian Industry Taxonomy (MIT), "
        "which groups companies by primary revenue source. Conglomerates are "
        "segmented where segment data is available; otherwise they are classified "
        "by their largest revenue segment.",
        body_style,
    ))

    # ── Page 6: Economic Outlook ─────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph(
        "4. Economic Outlook — Macroeconomic Factors",
        heading_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "The macroeconomic environment in 2025 was characterized by moderating "
        "inflation, stable interest rates, and continued (if slowing) GDP growth. "
        "Key macroeconomic indicators relevant to industrial companies:",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))

    macro_data = [
        ["Indicator", "2025", "2024", "Change"],
        ["US GDP Growth", "2.3%", "2.8%", "-0.5pp"],
        ["CPI Inflation", "2.6%", "3.1%", "-0.5pp"],
        ["Fed Funds Rate (Year-End)", "4.25%", "4.50%", "-25bps"],
        ["ISM Manufacturing PMI (Avg)", "52.1", "49.8", "+2.3"],
        ["Industrial Production Growth", "1.8%", "0.9%", "+0.9pp"],
    ]
    macro_table = Table(macro_data, colWidths=[2.5 * inch, 1.2 * inch, 1.2 * inch, 1.2 * inch])
    macro_table.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 9),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 9),
        ("LINEBELOW", (0, 0), (-1, 0), 1, colors.HexColor("#1A3C6E")),
        ("LINEBELOW", (0, -1), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF4")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(macro_table)
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "The manufacturing sector benefited from the PMI recovery above 50 "
        "(expansion territory) after a challenging 2024. However, trade policy "
        "uncertainty and potential tariff changes remain a risk factor for "
        "companies with significant import/export exposure.",
        body_style,
    ))

    # ── Page 7: Advanced Materials ───────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph(
        "5. Advanced Materials &amp; Composites Sector",
        heading_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "The advanced materials sector was the strongest performer among "
        "diversified industrials in 2025, driven by structural demand from "
        "aerospace, electric vehicle manufacturing, and renewable energy "
        "infrastructure.",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))

    bm_am = _INDUSTRY_BENCHMARKS["Advanced Materials & Composites"]
    am_data = [
        ["Metric", "2025 Value", "2024 Value", "Trend"],
        ["Revenue Growth (YoY)", bm_am["growth_rate"], "10.2%", "Accelerating"],
        ["Gross Margin (Median)", "48%", "46%", "Expanding"],
        ["Gross Margin (Range)", bm_am["margin_range"], "40%–52%", "Widening"],
        ["Market Size (US)", bm_am["market_size"], "$59B", "Growing"],
        ["R&D Intensity (% Revenue)", "6.2%", "5.8%", "Increasing"],
    ]
    am_table = Table(am_data, colWidths=[2.2 * inch, 1.5 * inch, 1.2 * inch, 1.2 * inch])
    am_table.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 9),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 9),
        ("LINEBELOW", (0, 0), (-1, 0), 1, colors.HexColor("#1A3C6E")),
        ("LINEBELOW", (0, -1), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF4")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(am_table)
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph(
        "Sub-segment analysis reveals significant dispersion. Advanced composites "
        "companies (carbon fiber, engineered polymers) grew 15–45% depending on "
        "end-market exposure, while specialty coatings and surface treatments "
        "declined 2–6% due to automotive OEM destocking and pricing pressure.",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "Companies with high composites concentration should expect margin "
        "pressure if aerospace order rates normalize. Firms reporting growth "
        "above 20% in this segment should demonstrate sustainable demand drivers "
        "rather than one-time contract wins.",
        body_style,
    ))

    # ── Page 8: Supply Chain ─────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph(
        "6. Supply Chain &amp; Raw Materials Analysis",
        heading_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "Raw material costs stabilized in 2025 after two years of volatility. "
        "Key commodity price movements relevant to industrial companies:",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "&#8226; Steel (HRC): -3.2% YoY, averaging $780/ton<br/>"
        "&#8226; Aluminum: +1.8% YoY, averaging $2,340/mt<br/>"
        "&#8226; Carbon fiber: -5.1% YoY as new capacity came online<br/>"
        "&#8226; Specialty chemicals: +4.2% YoY driven by regulatory costs<br/>"
        "&#8226; Packaging materials: +2.1% YoY",
        body_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "Supply chain lead times improved across most categories, with the "
        "average manufacturing lead time dropping from 8.2 weeks to 6.7 weeks. "
        "However, specialty materials (high-performance composites, rare earth "
        "elements) continued to face extended lead times due to geographic "
        "concentration of supply.",
        body_style,
    ))

    # ── Page 9: Labor Market ─────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph(
        "7. Labor Market &amp; Workforce Trends",
        heading_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "Labor market conditions for industrial companies remained tight in 2025, "
        "though showing signs of moderation. Average wage growth for manufacturing "
        "workers was 3.8%, down from 4.5% in 2024.",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "Key labor statistics for the sector:<br/>"
        "&#8226; Average turnover rate: 8.2% (down from 9.1% in 2024)<br/>"
        "&#8226; Unfilled positions rate: 4.3% of total headcount<br/>"
        "&#8226; Overtime hours: 12.1% above standard (vs. 11.4% prior year)<br/>"
        "&#8226; Average fully-loaded labor cost: $68,400/employee (up 3.4%)",
        body_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "Companies reporting significant revenue growth without corresponding "
        "headcount increases should be evaluated for automation-driven "
        "productivity gains or potential revenue quality concerns.",
        body_style,
    ))

    # ── Page 10: Regulatory ──────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph(
        "8. Regulatory &amp; Compliance Environment",
        heading_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "The regulatory environment for industrial companies saw several notable "
        "developments in 2025 that may impact revenue recognition, cost "
        "structures, and reporting requirements:",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "&#8226; FASB issued ASU 2025-03, clarifying performance obligation "
        "identification for bundled manufacturing and service contracts<br/>"
        "&#8226; EPA finalized updated PFAS regulations affecting specialty "
        "coatings manufacturers<br/>"
        "&#8226; DOL updated overtime rules, expanding eligibility for "
        "manufacturing employees<br/>"
        "&#8226; Trade policy: Section 301 tariffs on certain composites imports "
        "were extended, benefiting domestic producers",
        body_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "The PFAS regulatory changes are particularly relevant for specialty "
        "coatings companies, with estimated compliance costs of 1–3% of revenue "
        "for affected products. Several companies have begun reformulating "
        "product lines, with transition costs impacting near-term margins.",
        body_style,
    ))

    # ── Page 11: Logistics & Distribution ────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph(
        "9. Logistics &amp; Distribution Sector",
        heading_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "The logistics and distribution sector continued its post-pandemic "
        "normalization in 2025, with growth moderating from the exceptional "
        "rates seen in 2021–2023 but remaining healthy.",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))

    bm_ds = _INDUSTRY_BENCHMARKS["Logistics & Distribution"]
    ds_data = [
        ["Metric", "2025 Value", "2024 Value", "Trend"],
        ["Revenue Growth (YoY)", bm_ds["growth_rate"], "6.2%", "Improving"],
        ["Gross Margin (Median)", "18%", "17%", "Stable"],
        ["Gross Margin (Range)", bm_ds["margin_range"], "13%–21%", "Stable"],
        ["Market Size (US)", bm_ds["market_size"], "$278B", "Growing"],
        ["Warehouse Automation Rate", "34%", "28%", "Accelerating"],
    ]
    ds_table = Table(ds_data, colWidths=[2.2 * inch, 1.5 * inch, 1.2 * inch, 1.2 * inch])
    ds_table.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 9),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 9),
        ("LINEBELOW", (0, 0), (-1, 0), 1, colors.HexColor("#1A3C6E")),
        ("LINEBELOW", (0, -1), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF4")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(ds_table)
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph(
        "Warehousing services grew faster (8–12%) than pure freight and "
        "logistics (3–5%), reflecting the shift toward value-added warehousing "
        "and fulfillment services. Companies investing in automation and "
        "technology integration reported higher margins and faster growth.",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "Freight rates softened in Q3–Q4 2025 as overcapacity emerged in "
        "the trucking segment. Companies with high freight exposure should "
        "demonstrate volume growth to offset rate pressure.",
        body_style,
    ))

    # ── Page 12: Emerging Technologies ───────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph(
        "10. Emerging Technologies &amp; Industry 4.0",
        heading_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "Adoption of Industry 4.0 technologies continued to accelerate across "
        "the industrial sector in 2025. Key adoption metrics:",
        body_style,
    ))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "&#8226; IoT sensor deployment: 67% of surveyed manufacturers (up from 54%)<br/>"
        "&#8226; AI/ML in quality control: 31% of companies (up from 22%)<br/>"
        "&#8226; Digital twin adoption: 18% of large manufacturers<br/>"
        "&#8226; Predictive maintenance: 42% of companies (up from 35%)<br/>"
        "&#8226; Robotic process automation in back-office: 28% adoption",
        body_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "Companies that have invested in digital transformation reported, on "
        "average, 2.1 percentage points higher revenue growth and 150 bps higher "
        "operating margins compared to peers. However, the capital investment "
        "required remains a barrier for mid-market companies.",
        body_style,
    ))
    story.append(Spacer(1, 1 * inch))
    story.append(Paragraph(
        "<i>Disclaimer: This report is prepared by Meridian Research Group for "
        "informational purposes only. It does not constitute investment advice. "
        "All data is believed to be reliable but is not guaranteed.</i>",
        small_style,
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        f"&#169; 2026 Meridian Research Group. Report Ref: MRG-2026-{canary_code}",
        small_style,
    ))

    # ── Build PDF ────────────────────────────────────────────────────
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=inch,
        rightMargin=inch,
        title="Industry Benchmark Report — Meridian Research Group",
        author=f"CANARY: {canary_code}",
        creator="Meridian Research Group",
        invariant=True,
    )
    doc.build(story)

    canaries.set_location(
        file_key,
        f"{_INPUT_DIR}/industry_benchmark_report.pdf",
        "PDF metadata → Author; also in Report Reference on pages 1 and 12",
    )
    manifest.register(
        f"{_INPUT_DIR}/industry_benchmark_report.pdf",
        "pdf",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Management Representation Letter ─────────────────────────────────────────


def _write_mgmt_rep_letter(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write management_rep_letter.docx asserting ~8% growth."""
    file_key = "tc03_mgmt_rep_letter"
    canary_code = canaries.canary_for(file_key)

    path = output_dir / _INPUT_DIR / "management_rep_letter.docx"
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()

    # Set core properties
    doc.core_properties.author = "Cascade Industries Management"
    doc.core_properties.title = "Management Representation Letter — FY2025"
    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME
    embed_canary_docx(doc, canary_code)

    # ── Letterhead ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run("CASCADE INDUSTRIES, INC.")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    run = p.add_run(
        "500 NW Industrial Boulevard\n"
        "Portland, Oregon 97209\n"
        "Tel: (503) 555-0100"
    )
    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ── Date and addressee ───────────────────────────────────────────
    doc.add_paragraph("January 31, 2026")
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Mitchell & Associates LLP\n")
    run.bold = True
    p.add_run(
        "Attn: Audit Engagement Team\n"
        "1200 NW Couch Street, Suite 800\n"
        "Portland, Oregon 97209"
    )

    doc.add_paragraph()

    # ── Subject line ─────────────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run("Re: Management Representation Letter — Fiscal Year 2025 Audit")
    run.bold = True

    doc.add_paragraph()

    # ── Body ─────────────────────────────────────────────────────────
    doc.add_paragraph(
        "Dear Mitchell & Associates LLP:"
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "In connection with your audit of the consolidated financial statements "
        "of Cascade Industries, Inc. and its subsidiaries (\"the Company\") for "
        "the fiscal year ended December 31, 2025, we make the following "
        "representations to you. These representations are based on our knowledge "
        "of the Company's operations and financial position."
    )

    # Revenue section — the critical misrepresentation
    p = doc.add_paragraph()
    run = p.add_run("Revenue Performance")
    run.bold = True

    doc.add_paragraph(
        "Consolidated revenue grew approximately 8% year-over-year, driven "
        "primarily by Advanced Materials. The growth reflects continued strong "
        "demand for our advanced composites products in the aerospace and "
        "electric vehicle markets. Our Precision Components and Distribution "
        "Services segments contributed steady, consistent growth in line with "
        "their respective industry sectors."
    )

    doc.add_paragraph(
        "Revenue has been recognized in accordance with ASC 606, with all "
        "performance obligations satisfied at the point of transfer of control. "
        "No significant changes were made to revenue recognition policies "
        "during the year."
    )

    # Other standard sections
    p = doc.add_paragraph()
    run = p.add_run("Completeness of Information")
    run.bold = True

    doc.add_paragraph(
        "We have provided you with access to all financial records and related "
        "data, including minutes of meetings of stockholders, the board of "
        "directors, and committees of the board. There are no material "
        "transactions that have not been properly recorded in the accounting "
        "records underlying the financial statements."
    )

    p = doc.add_paragraph()
    run = p.add_run("Estimates and Judgments")
    run.bold = True

    doc.add_paragraph(
        "The methods, significant assumptions, and data used in making "
        "accounting estimates and their related disclosures are appropriate "
        "to achieve recognition, measurement, and disclosure that are in "
        "conformity with U.S. GAAP."
    )

    p = doc.add_paragraph()
    run = p.add_run("Subsequent Events")
    run.bold = True

    doc.add_paragraph(
        "There have been no events subsequent to the balance sheet date that "
        "would require adjustment to, or disclosure in, the financial "
        "statements for the year ended December 31, 2025."
    )

    p = doc.add_paragraph()
    run = p.add_run("Litigation and Claims")
    run.bold = True

    doc.add_paragraph(
        "We are not aware of any pending or threatened litigation, claims, "
        "or assessments that are not disclosed in the financial statements. "
        "There are no unasserted claims or assessments that our legal counsel "
        "has advised us are probable of assertion."
    )

    doc.add_paragraph()

    # Signatures
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Robert J. Cascade")
    run.bold = True

    doc.add_paragraph("Chief Executive Officer")
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Patricia A. Thornton")
    run.bold = True

    doc.add_paragraph("Chief Financial Officer")
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(f"Ref: CI-MRL-2025-{canary_code}")
    run.font.size = Pt(8)
    run.font.color.rgb = docx.shared.RGBColor(0x99, 0x99, 0x99)

    # ── Save ─────────────────────────────────────────────────────────
    _save_docx_deterministic(doc, path)

    canaries.set_location(
        file_key,
        f"{_INPUT_DIR}/management_rep_letter.docx",
        "Core properties → comments; also in Ref footer",
    )
    manifest.register(
        f"{_INPUT_DIR}/management_rep_letter.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Prompt & Expected Behavior ──────────────────────────────────────────────


def _write_prompt(output_dir: Path) -> None:
    """Write test_cases/TC-03/prompt.md per spec."""
    text = """\
Perform substantive analytical procedures on Cascade Industries' revenue for FY2025.

1. Analyze revenue trends by product line — monthly and annual.
2. Compute year-over-year growth rates by product line and in aggregate.
3. Compare growth rates to the industry benchmarks in the provided report.
4. Assess whether management's representation of ~8% growth is supported by the data.
5. Identify any product lines with unusual patterns (seasonality shifts, trend breaks,
   or growth significantly above/below industry benchmarks).
6. Draft an analytical procedures memo documenting:
   - Scope and objective
   - Data sources used
   - Methodology
   - Findings (with supporting data tables)
   - Conclusion and any follow-up procedures recommended

Output the memo as a Word document and the supporting analysis as an Excel workbook.
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write test_cases/TC-03/expected_behavior.md per spec."""
    text = """\
# TC-03: Substantive Analytical Procedures — Revenue — Expected Behavior

## Key Findings the Agent Should Identify

1. **Consolidated growth is 9.2%, not ~8%**: Management's representation letter
   claims "approximately 8% year-over-year" growth. The actual consolidated
   revenue growth from the data is 9.2%. The agent should detect this discrepancy
   and flag it — this is the central test of the analytical procedure.

2. **Two product lines drove growth**:
   - Advanced Composites grew ~45% YoY (well above the 12.8% industry benchmark
     for advanced materials — the agent should flag this as unusual)
   - Warehousing Services grew ~10% YoY (above the 7.5% industry benchmark)

3. **One product line declined**:
   - Specialty Coatings declined ~4% YoY (consistent with industry trends of
     -2% to -6% noted in the benchmark report on page 7)

4. **Industry benchmark comparison**: The agent must find benchmarks on pages
   4, 7, and 11 of the PDF — not just the executive summary. The relevant
   benchmarks are:
   - Industrial Manufacturing: 4.2% growth
   - Advanced Materials & Composites: 12.8% growth
   - Logistics & Distribution: 7.5% growth

## Data Challenges

- **Benchmark data is spread across pages**: The PDF is 12 pages; relevant
  benchmarks appear on pages 4, 7, and 11. An agent that only reads page 1
  or the executive summary will miss the detailed sector data.
- **Management letter is close but wrong**: "Approximately 8%" is close enough
  to be plausible, testing whether the agent actually computes the number vs.
  accepting management's assertion.
- **Advanced Composites growth outlier**: ~45% growth is far above the 12.8%
  industry average. The agent should flag this as requiring further investigation
  (e.g., is it driven by a single large contract?).

## Expected Output Structure

### Memo (Word document):
- Scope and objective statement
- List of data sources (revenue spreadsheet, benchmark report, rep letter)
- Methodology description
- Findings with supporting data tables
- Explicit identification of the management rep discrepancy
- Conclusion and recommended follow-up procedures

### Analysis Workbook (Excel):
- Monthly trend analysis
- Annual summary by product line
- Year-over-year growth calculations
- Industry benchmark comparison
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold Standard ────────────────────────────────────────────────────────────


@register_gold("TC-03")
def _tc03_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """TC-03 gold standard: revenue analytical procedures."""
    model: CascadeModel = model_kwargs["model"]

    # Compute growth rates from the canonical model
    consol_growth = validate_consolidated_growth(model.revenue_records)
    pl_growth = validate_product_line_growth(model.revenue_records)

    # Compute annual revenue by product line
    by_pl_year: dict[tuple[str, int], Decimal] = {}
    for rec in model.revenue_records:
        if rec.year in (2024, 2025):
            key = (rec.product_line, rec.year)
            by_pl_year[key] = by_pl_year.get(key, Decimal(0)) + rec.revenue

    pl_names = sorted({rec.product_line for rec in model.revenue_records})

    product_line_details = {}
    for name in pl_names:
        fy24 = by_pl_year.get((name, 2024), Decimal(0))
        fy25 = by_pl_year.get((name, 2025), Decimal(0))
        growth = pl_growth.get(name, Decimal(0))
        product_line_details[name] = {
            "fy2024_revenue": int(fy24.quantize(Decimal("1"), rounding=ROUND_HALF_UP)),
            "fy2025_revenue": int(fy25.quantize(Decimal("1"), rounding=ROUND_HALF_UP)),
            "yoy_growth_pct": round(float(growth * 100), 1),
        }

    fy24_total = sum(
        by_pl_year.get((name, 2024), Decimal(0)) for name in pl_names
    )
    fy25_total = sum(
        by_pl_year.get((name, 2025), Decimal(0)) for name in pl_names
    )

    return GoldStandard(
        test_case="TC-03",
        expected_outputs={
            "output_files": {
                "memo": {"type": "docx", "required_sections": [
                    "Scope and Objective",
                    "Data Sources",
                    "Methodology",
                    "Findings",
                    "Conclusion",
                ]},
                "analysis_workbook": {"type": "xlsx"},
            },
            "consolidated_revenue": {
                "fy2024": int(fy24_total.quantize(Decimal("1"), rounding=ROUND_HALF_UP)),
                "fy2025": int(fy25_total.quantize(Decimal("1"), rounding=ROUND_HALF_UP)),
                "yoy_growth_pct": round(
                    float(consol_growth.get("FY2025_growth", Decimal(0)) * 100), 1,
                ),
            },
            "management_rep_discrepancy": {
                "claimed_growth": "approximately 8%",
                "actual_growth_pct": round(
                    float(consol_growth.get("FY2025_growth", Decimal(0)) * 100), 1,
                ),
                "discrepancy_detected": True,
                "explanation": (
                    "Management claims ~8% growth; actual consolidated growth is "
                    "9.2%. The difference of ~1.2pp may indicate understatement "
                    "rather than rounding."
                ),
            },
            "product_line_analysis": product_line_details,
            "growth_drivers": [
                "Advanced Composites",
                "Warehousing Services",
            ],
            "declining_lines": [
                "Specialty Coatings",
            ],
            "industry_benchmark_comparison": {
                "industrial_manufacturing": {
                    "benchmark_growth": "4.2%",
                    "cascade_comparable_lines": [
                        "Industrial Parts",
                        "Custom Machining",
                    ],
                },
                "advanced_materials": {
                    "benchmark_growth": "12.8%",
                    "cascade_comparable_lines": [
                        "Advanced Composites",
                        "Specialty Coatings",
                    ],
                },
                "logistics_distribution": {
                    "benchmark_growth": "7.5%",
                    "cascade_comparable_lines": [
                        "Warehousing Services",
                        "Freight & Logistics",
                    ],
                },
            },
        },
        canary_verification={
            "read_revenue_data": canaries.canary_for("tc03_revenue_by_product"),
            "read_benchmark_report": canaries.canary_for("tc03_industry_benchmark"),
            "read_mgmt_rep_letter": canaries.canary_for("tc03_mgmt_rep_letter"),
        },
        error_detection={},
        scoring_hints={
            "correctness": (
                "Growth rate must be 9.2% (within rounding). "
                "Product-line growth rates must match gold standard within 0.5pp."
            ),
            "completeness": (
                "All 6 product lines analyzed. All 3 benchmark sectors compared. "
                "Management discrepancy explicitly identified."
            ),
            "format_compliance": (
                "Memo as Word document with required sections. "
                "Analysis as Excel workbook."
            ),
            "robustness": (
                "Agent must read benchmark PDF beyond page 1 to find sector data. "
                "Agent must compute actual growth rather than accepting management's claim."
            ),
            "communication": (
                "Clearly flags the ~8% vs 9.2% discrepancy. Notes Advanced Composites "
                "growth of ~45% as an outlier requiring investigation."
            ),
        },
    )


# ── Public entry point ───────────────────────────────────────────────────────


def emit_tc03(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Emit all TC-03 files."""
    _write_revenue_xlsx(model, output_dir, canaries, errors, manifest)
    _write_benchmark_pdf(output_dir, canaries, errors, manifest)
    _write_mgmt_rep_letter(model, output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
