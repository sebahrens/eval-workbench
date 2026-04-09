"""Formatter: static template files used by test cases.

Emits:
- templates/deliverable_cover_page.docx  (TC-17)
- templates/formatting_guide.pdf         (TC-17)
"""

from __future__ import annotations

import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF

from generator.canaries import CanaryRegistry, embed_canary_docx, embed_canary_pdf_fpdf2
from generator.manifest import Manifest

# ---------------------------------------------------------------------------
# Cover page template (docx)
# ---------------------------------------------------------------------------

def _write_cover_page(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """Create templates/deliverable_cover_page.docx with branded placeholders."""
    doc = Document()

    # -- Canary
    canary_code = canaries.canary_for("cover_page_template")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "cover_page_template",
        "templates/deliverable_cover_page.docx",
        location,
    )

    # -- Page margins
    for section in doc.sections:
        section.top_margin = Inches(2.0)
        section.bottom_margin = Inches(1.5)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.5)

    # -- Company logo placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[COMPANY LOGO]")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.space_after = Pt(72)

    # -- Title placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{REPORT_TITLE}}")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    p.space_after = Pt(24)

    # -- Subtitle / client name placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{CLIENT_NAME}}")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.space_after = Pt(12)

    # -- Date placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{DATE}}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.space_after = Pt(48)

    # -- Horizontal rule (via a thin border paragraph)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    p.space_after = Pt(24)

    # -- Prepared by / confidentiality notice
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared by: {{PREPARED_BY}}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL — FOR AUTHORIZED USE ONLY")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # -- Save
    path = output_dir / "templates" / "deliverable_cover_page.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))

    manifest.register(
        "templates/deliverable_cover_page.docx",
        "docx",
        canary=canary_code,
        test_cases=["TC-17"],
    )


# ---------------------------------------------------------------------------
# Formatting guide (PDF)
# ---------------------------------------------------------------------------

# Fixed creation date for determinism
_FIXED_DATE = datetime.datetime(2025, 1, 15, 9, 0, 0)


def _write_formatting_guide(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """Create templates/formatting_guide.pdf with formatting specs for TC-17."""
    pdf = FPDF()
    pdf.set_creation_date(_FIXED_DATE)
    pdf.set_auto_page_break(auto=True, margin=25)

    # -- Canary
    canary_code = canaries.canary_for("formatting_guide")
    location = embed_canary_pdf_fpdf2(pdf, canary_code)
    canaries.set_location(
        "formatting_guide",
        "templates/formatting_guide.pdf",
        location,
    )

    # ── Page 1: Title + Section Order ──
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 15, "Deliverable Formatting Guide", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, "Cascade Industries - Advisory Deliverables", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(10)

    # Section order
    _section_heading(pdf, "1. Section Order")
    _body(pdf, "All client deliverables must follow this section order:")
    pdf.ln(3)
    sections = [
        "1. Cover Page - Branded cover with report title, client name, and date.",
        "2. Table of Contents - Auto-generated; lists all sections with page numbers.",
        "3. Executive Summary - High-level overview of findings and recommendations.",
        "4. Financial Analysis - Quantitative analysis including key financial tables.",
        "5. Industry Overview - Market context and peer comparison.",
        "6. Risk Assessment - Identified risks ranked by likelihood and impact.",
        "7. Detailed Findings - In-depth analysis with supporting data tables.",
        "8. Recommendations - Actionable next steps prioritized by impact.",
    ]
    for section in sections:
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, f"  {section}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # ── Page numbering rules ──
    _section_heading(pdf, "2. Page Numbering")
    numbering_rules = [
        "Cover page: no page number displayed.",
        "Table of Contents: Roman numerals (i, ii, iii, ...).",
        "Body sections (Executive Summary onward): Arabic numerals starting at 1.",
        "Page numbers appear in the footer, centered, in 10pt font.",
        "Format: 'Page X of Y' for body sections; 'Page i' style for TOC.",
    ]
    for rule in numbering_rules:
        _bullet(pdf, rule)
    pdf.ln(5)

    # ── Header/footer requirements ──
    _section_heading(pdf, "3. Headers and Footers")
    _body(pdf, "Headers:")
    header_rules = [
        "All pages except the cover page must have a header.",
        "Left-aligned: 'Cascade Industries - CONFIDENTIAL'",
        "Right-aligned: Section title (e.g., 'Executive Summary').",
        "Font: Helvetica 9pt, dark gray (#666666).",
        "A thin horizontal rule (0.5pt) separates the header from body text.",
    ]
    for rule in header_rules:
        _bullet(pdf, rule)
    pdf.ln(3)

    _body(pdf, "Footers:")
    footer_rules = [
        "Centered page number (see numbering rules above).",
        "Right-aligned: Date of report generation.",
        "Font: Helvetica 9pt, dark gray (#666666).",
        "A thin horizontal rule (0.5pt) separates body text from the footer.",
    ]
    for rule in footer_rules:
        _bullet(pdf, rule)
    pdf.ln(5)

    # ── Page 2: Font standards + table formatting ──
    pdf.add_page()
    _section_heading(pdf, "4. Font Standards")
    font_specs = [
        ("Report title (cover page)", "Helvetica Bold, 28pt, #1A3C6E"),
        ("Section headings (H1)", "Helvetica Bold, 16pt, #1A3C6E"),
        ("Sub-headings (H2)", "Helvetica Bold, 13pt, #333333"),
        ("Body text", "Helvetica Regular, 11pt, #000000"),
        ("Table headers", "Helvetica Bold, 10pt, #FFFFFF on #1A3C6E background"),
        ("Table body", "Helvetica Regular, 10pt, #000000"),
        ("Captions / footnotes", "Helvetica Italic, 9pt, #666666"),
        ("Headers / footers", "Helvetica Regular, 9pt, #666666"),
    ]
    for element, spec in font_specs:
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, f"  {element}: {spec}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    _section_heading(pdf, "5. Spacing and Margins")
    spacing_rules = [
        "Page margins: 1 inch on all sides.",
        "Line spacing: 1.15 for body text, single for tables.",
        "Paragraph spacing: 6pt after each paragraph.",
        "Section breaks: Each major section starts on a new page.",
        "Table padding: 4pt cell padding on all sides.",
    ]
    for rule in spacing_rules:
        _bullet(pdf, rule)
    pdf.ln(5)

    _section_heading(pdf, "6. Table Formatting")
    table_rules = [
        "Header row: white text on #1A3C6E background, bold.",
        "Alternating row shading: white / #F2F6FA.",
        "Border: thin (0.5pt) #CCCCCC gridlines.",
        "Numeric columns: right-aligned with consistent decimal places.",
        "Currency: use dollar sign with commas (e.g., $1,234,567.89).",
        "Percentages: one decimal place (e.g., 12.3%).",
        "Tables must have a caption (e.g., 'Exhibit 1: Revenue by Segment').",
    ]
    for rule in table_rules:
        _bullet(pdf, rule)
    pdf.ln(5)

    _section_heading(pdf, "7. Embedded Excel Content")
    excel_rules = [
        "Key tables from Excel workpaper sections must be embedded as formatted tables, "
        "not as screenshots or images.",
        "Preserve column headers and data types from the source worksheet.",
        "Apply the table formatting standards defined in Section 6.",
        "Include a source reference below each table (e.g., 'Source: 02_financial_analysis.xlsx, "
        "Sheet: Summary').",
    ]
    for rule in excel_rules:
        _bullet(pdf, rule)

    # -- Save
    path = output_dir / "templates" / "formatting_guide.pdf"
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(path))

    manifest.register(
        "templates/formatting_guide.pdf",
        "pdf",
        canary=canary_code,
        test_cases=["TC-17"],
    )


# ---------------------------------------------------------------------------
# PDF text helpers
# ---------------------------------------------------------------------------

def _section_heading(pdf: FPDF, text: str) -> None:
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(0x1A, 0x3C, 0x6E)
    pdf.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)


def _body(pdf: FPDF, text: str) -> None:
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, text, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)


def _bullet(pdf: FPDF, text: str) -> None:
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 6, f"  - {text}", new_x="LMARGIN", new_y="NEXT")


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def emit_templates(output_dir: Path, canaries: CanaryRegistry, manifest: Manifest) -> None:
    """Write all template files to *output_dir*/templates/."""
    _write_cover_page(output_dir, canaries, manifest)
    _write_formatting_guide(output_dir, canaries, manifest)
