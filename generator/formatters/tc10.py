"""Formatter: TC-10 — Multi-State Apportionment (Tax, Routine).

Emits:
- test_cases/TC-10/input_files/consolidated_pl_fy2025.xlsx
  Consolidated income statement for FY2025.
- test_cases/TC-10/input_files/state_factors.xlsx
  State-by-state factors schedule: Sales, Payroll, Property factors.
  OR/TX/IL complete, CA/WA partial (sales only), NY blank.
  Some cells "$0" vs blank to test zero-vs-missing distinction.
- test_cases/TC-10/input_files/apportionment_rules.docx
  Reference table of each state's apportionment formula.
- test_cases/TC-10/prompt.md
- test_cases/TC-10/expected_behavior.md
- gold_standards/TC-10_gold.json

Planted errors:
- ERR-003 (transposed_digits): one state's sales factor has transposed digits
- ERR-023 (rounding_discrepancy): total expenses shows a rounding discrepancy
Uses the canonical model — never hardcodes numbers.
"""

from __future__ import annotations

import datetime
import io
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from generator.canaries import CanaryRegistry, embed_canary_docx, embed_canary_xlsx
from generator.errors import (
    ErrorRegistry,
    PlantedError,
    rounding_discrepancy,
    transpose_digits,
)
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel
from generator.model.entities import ENTITIES
from generator.model.views import build_income_statement

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-10"
_INPUT_DIR = f"test_cases/{_TC}/input_files"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)
_FIXED_ZIP_DT = (2025, 3, 15, 9, 0, 0)

# ── State apportionment data ────────────────────────────────────────────────
# States with operations or nexus. Factor data completeness varies per spec.
#
# OR, TX, IL: all three factors populated (from entity data).
# CA, WA: sales factor only (from remote employees generating nexus).
# NY: all blank (nexus questionable, no data provided).
#
# For OR/TX/IL the factors come from the canonical model:
# - Sales factor: entity revenue / consolidated revenue
# - Payroll factor: entity payroll / consolidated payroll
# - Property factor: entity net PP&E / consolidated net PP&E

# State tax rates (for gold standard computation)
_STATE_TAX_RATES: dict[str, Decimal] = {
    "OR": Decimal("0.066"),   # Oregon Corporate Excise Tax 6.6%
    "TX": Decimal("0.00"),    # Texas has margin tax, not income tax (computed differently)
    "IL": Decimal("0.095"),   # Illinois 7.0% + 2.5% replacement tax = 9.5%
    "CA": Decimal("0.084"),   # California 8.84%
    "WA": Decimal("0.00"),    # Washington has B&O tax, not income tax
    "NY": Decimal("0.071"),   # New York 7.1%
}

# Apportionment formulas per state
_APPORTIONMENT_FORMULAS: dict[str, str] = {
    "OR": "Single sales factor",
    "TX": "Margin tax — different base (gross revenue minus cost of goods sold or compensation)",
    "IL": "Single sales factor",
    "CA": "Single sales factor, market-based sourcing",
    "WA": "B&O tax (gross receipts) — not an income tax",
    "NY": "Single sales factor with customer-based sourcing",
}


# ── Deterministic save helpers ──────────────────────────────────────────────


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)
    wb.properties.modified = _FIXED_DATETIME

    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _pin_xlsx_dates(wb: openpyxl.Workbook) -> None:
    wb.properties.created = _FIXED_DATETIME
    wb.properties.modified = _FIXED_DATETIME


def _whole_dollars(d: Decimal) -> int:
    """Round a Decimal to whole dollars."""
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


def _fmt_dollars(d: Decimal | int) -> str:
    """Format as dollar string for gold standard."""
    if isinstance(d, int):
        return f"${d:,}"
    return f"${_whole_dollars(d):,}"


# ── Styling helpers ─────────────────────────────────────────────────────────

_HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
_HEADER_FILL = PatternFill("solid", fgColor="4472C4")
_HEADER_ALIGN = Alignment(horizontal="center", wrap_text=True)
_THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)
_NUMBER_FMT = '#,##0'
_PCT_FMT = '0.0%'
_MONEY_FMT = '$#,##0'


def _style_header(ws: Any, row: int, col_count: int) -> None:
    """Apply header styling to a row."""
    for col in range(1, col_count + 1):
        cell = ws.cell(row=row, column=col)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = _HEADER_ALIGN
        cell.border = _THIN_BORDER


def _style_data_cell(cell: Any, fmt: str = "") -> None:
    """Apply data cell styling."""
    cell.border = _THIN_BORDER
    if fmt:
        cell.number_format = fmt


# ── Factor computation from canonical model ─────────────────────────────────


def _compute_state_factors(
    model: CascadeModel,
) -> dict[str, dict[str, Decimal | None]]:
    """Compute sales, payroll, and property factors by state.

    Returns a dict: state -> {"sales": Decimal|None, "payroll": Decimal|None,
    "property": Decimal|None}.

    OR/TX/IL get all three factors from entity data.
    CA/WA get sales only (from remote employee revenue attribution).
    NY gets all None (blank).
    """
    # ── Sales factor: entity FY2025 revenue / consolidated revenue ──
    consolidated_revenue = Decimal(0)
    revenue_by_entity: dict[str, Decimal] = {}
    for r in model.revenue_records:
        if r.year == 2025:
            revenue_by_entity[r.entity_code] = (
                revenue_by_entity.get(r.entity_code, Decimal(0)) + r.revenue
            )
            consolidated_revenue += r.revenue

    # Map entity code to state
    entity_state = {e.code: e.state for e in ENTITIES.values() if not e.is_parent}

    revenue_by_state: dict[str, Decimal] = {}
    for code, rev in sorted(revenue_by_entity.items()):
        st = entity_state.get(code)
        if st:
            revenue_by_state[st] = revenue_by_state.get(st, Decimal(0)) + rev

    # ── Payroll factor: FY2025 salaries by entity state ──
    # Only count active employees (no termination date or terminated after 2025)
    consolidated_payroll = Decimal(0)
    payroll_by_state: dict[str, Decimal] = {}
    for emp in model.employees:
        if emp.termination_date and emp.termination_date.year < 2025:
            continue
        sal = Decimal(emp.annual_salary)
        # Employee state is their work state
        payroll_by_state[emp.state] = payroll_by_state.get(emp.state, Decimal(0)) + sal
        consolidated_payroll += sal

    # ── Property factor: net book value of PP&E by entity state ──
    import datetime as dt
    as_of = dt.date(2025, 12, 31)
    consolidated_property = Decimal(0)
    property_by_state: dict[str, Decimal] = {}
    for asset in model.assets:
        if asset.disposed and asset.disposal_date and asset.disposal_date <= as_of:
            continue
        # Compute net book value: cost - accumulated depreciation
        years_held = (as_of - asset.acquisition_date).days / Decimal("365.25")
        if years_held < 0:
            continue
        annual_depr = (asset.cost - asset.salvage_value) / Decimal(str(asset.book_life_years))
        accum = min(annual_depr * years_held, asset.cost - asset.salvage_value)
        nbv = asset.cost - accum.quantize(Decimal("1"), rounding=ROUND_HALF_UP)
        if nbv < 0:
            nbv = Decimal(0)

        st = entity_state.get(asset.entity_code)
        if st:
            property_by_state[st] = property_by_state.get(st, Decimal(0)) + nbv
            consolidated_property += nbv

    # ── Build factor dict ──
    factors: dict[str, dict[str, Decimal | None]] = {}

    # OR, TX, IL: all three factors
    for state in ["OR", "TX", "IL"]:
        sales = (
            revenue_by_state.get(state, Decimal(0)) / consolidated_revenue
            if consolidated_revenue
            else Decimal(0)
        )
        payroll = (
            payroll_by_state.get(state, Decimal(0)) / consolidated_payroll
            if consolidated_payroll
            else Decimal(0)
        )
        prop = (
            property_by_state.get(state, Decimal(0)) / consolidated_property
            if consolidated_property
            else Decimal(0)
        )
        factors[state] = {
            "sales": sales.quantize(Decimal("0.0001"), rounding=ROUND_HALF_UP),
            "payroll": payroll.quantize(Decimal("0.0001"), rounding=ROUND_HALF_UP),
            "property": prop.quantize(Decimal("0.0001"), rounding=ROUND_HALF_UP),
        }

    # CA, WA: sales factor only (partial data)
    # Remote employees create sales nexus; represent as small sales fractions
    # based on payroll presence (employees in those states imply some sales activity).
    for state in ["CA", "WA"]:
        # Use payroll proportion as a proxy for sales factor
        if consolidated_payroll and state in payroll_by_state:
            sales = (
                payroll_by_state[state] / consolidated_payroll
            ).quantize(Decimal("0.0001"), rounding=ROUND_HALF_UP)
        else:
            sales = Decimal("0.0000")
        factors[state] = {
            "sales": sales,
            "payroll": None,  # Not provided
            "property": None,  # Not provided
        }

    # NY: all blank
    factors["NY"] = {
        "sales": None,
        "payroll": None,
        "property": None,
    }

    return factors


def _compute_apportioned_income(
    model: CascadeModel,
    factors: dict[str, dict[str, Decimal | None]],
) -> dict[str, dict[str, Any]]:
    """Compute apportioned income by state for the gold standard.

    Returns state -> {
        "apportionment_pct": Decimal or None,
        "apportioned_income": Decimal or None,
        "tax_type": str,
        "estimated_tax": Decimal or None,
        "notes": str,
    }
    """
    is_stmt = build_income_statement(model.ledger, 2025)
    pre_tax = is_stmt.pre_tax_income

    results: dict[str, dict[str, Any]] = {}

    for state in ["OR", "TX", "IL", "CA", "WA", "NY"]:
        sf = factors[state]

        if state in ("OR", "IL"):
            # Single sales factor states
            pct = sf["sales"]
            apportioned = (pre_tax * pct).quantize(Decimal("1"), rounding=ROUND_HALF_UP)
            tax = (apportioned * _STATE_TAX_RATES[state]).quantize(
                Decimal("1"), rounding=ROUND_HALF_UP
            )
            results[state] = {
                "formula": _APPORTIONMENT_FORMULAS[state],
                "apportionment_pct": pct,
                "apportioned_income": apportioned,
                "tax_type": "income_tax",
                "estimated_tax": tax,
                "notes": "",
            }

        elif state == "TX":
            # Texas margin tax — different base, cannot compute with income data alone
            pct = sf["sales"]
            results[state] = {
                "formula": _APPORTIONMENT_FORMULAS[state],
                "apportionment_pct": pct,
                "apportioned_income": None,
                "tax_type": "margin_tax",
                "estimated_tax": None,
                "notes": (
                    "Texas imposes a margin tax (franchise tax), not an income tax. "
                    "The tax base is total revenue minus the greater of: (1) cost of "
                    "goods sold, (2) compensation, or (3) 30% of total revenue. "
                    "Additional data needed: Texas-sourced gross revenue and detailed "
                    "COGS/compensation breakdown for the Texas entity."
                ),
            }

        elif state == "CA":
            # Single sales factor but data incomplete (sales only, no payroll/property)
            pct = sf["sales"]
            apportioned = (pre_tax * pct).quantize(Decimal("1"), rounding=ROUND_HALF_UP)
            tax = (apportioned * _STATE_TAX_RATES[state]).quantize(
                Decimal("1"), rounding=ROUND_HALF_UP
            )
            results[state] = {
                "formula": _APPORTIONMENT_FORMULAS[state],
                "apportionment_pct": pct,
                "apportioned_income": apportioned,
                "tax_type": "income_tax",
                "estimated_tax": tax,
                "notes": (
                    "Incomplete data: only sales factor provided. "
                    "California uses single sales factor with market-based sourcing. "
                    "Need customer location data to properly source sales."
                ),
            }

        elif state == "WA":
            # B&O tax — gross receipts, not income. Nexus questionable.
            pct = sf["sales"]
            results[state] = {
                "formula": _APPORTIONMENT_FORMULAS[state],
                "apportionment_pct": pct,
                "apportioned_income": None,
                "tax_type": "gross_receipts_tax",
                "estimated_tax": None,
                "notes": (
                    "Washington imposes a Business & Occupation (B&O) tax on gross "
                    "receipts, not net income. A separate computation is required. "
                    "Nexus question: only sales data present with no physical "
                    "presence indicators (no payroll or property). Economic nexus "
                    "thresholds should be evaluated."
                ),
            }

        elif state == "NY":
            # All data blank — cannot compute
            results[state] = {
                "formula": _APPORTIONMENT_FORMULAS[state],
                "apportionment_pct": None,
                "apportioned_income": None,
                "tax_type": "income_tax",
                "estimated_tax": None,
                "notes": (
                    "No factor data provided. Cannot compute apportioned income. "
                    "Need sales, payroll, and property data for New York. "
                    "New York uses single sales factor with customer-based sourcing."
                ),
            }

    return results


# ── Consolidated P&L xlsx ───────────────────────────────────────────────────


def _write_consolidated_pl(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write consolidated_pl_fy2025.xlsx — consolidated income statement."""
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    ws = wb.active
    ws.title = "Income Statement"

    canary_code = canaries.canary_for("tc10_consolidated_pl")
    location = embed_canary_xlsx(wb, canary_code)
    canaries.set_location(
        "tc10_consolidated_pl",
        f"{_INPUT_DIR}/consolidated_pl_fy2025.xlsx",
        location,
    )

    is_stmt = build_income_statement(model.ledger, 2025)

    # Headers
    headers = ["Description", "Amount"]
    for col, hdr in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=hdr)
    _style_header(ws, 1, len(headers))

    ws.column_dimensions["A"].width = 45
    ws.column_dimensions["B"].width = 20

    # Build P&L line items from income statement detail
    row = 2

    def _write_line(label: str, amount: Decimal, bold: bool = False, indent: int = 0) -> int:
        nonlocal row
        cell_a = ws.cell(row=row, column=1, value=("  " * indent) + label)
        cell_b = ws.cell(row=row, column=2, value=_whole_dollars(amount))
        _style_data_cell(cell_a)
        _style_data_cell(cell_b, _MONEY_FMT)
        if bold:
            cell_a.font = Font(bold=True)
            cell_b.font = Font(bold=True)
        row += 1
        return row

    # Revenue section
    _write_line("Total Revenue", is_stmt.total_revenue, bold=True)

    # COGS
    _write_line("Cost of Goods Sold", is_stmt.total_cogs, indent=1)

    # Gross profit
    _write_line("Gross Profit", is_stmt.gross_profit, bold=True)

    # Operating expenses — ERR-023: rounding discrepancy
    # Introduce a fractional amount so truncation vs rounding produces a $1 diff
    correct_opex = _whole_dollars(is_stmt.total_opex)
    # Add $0.60 so truncation (down) yields correct_opex, but the "shown" value
    # uses ROUND_HALF_UP yielding correct_opex + 1
    nudged_opex = float(correct_opex) + 0.6
    corrupt_opex = int(rounding_discrepancy(nudged_opex, decimal_places=0, direction="up"))
    errors.add(PlantedError(
        error_id="ERR-023",
        file=f"{_INPUT_DIR}/consolidated_pl_fy2025.xlsx",
        location="Sheet 'Income Statement', Row 'Total Operating Expenses'",
        type="rounding_discrepancy",
        description=(
            f"Total Operating Expenses shows ${corrupt_opex:,} "
            f"instead of ${correct_opex:,}"
        ),
        severity="immaterial",
        which_test_cases_should_catch=["TC-10"],
    ))
    _write_line("Total Operating Expenses", Decimal(corrupt_opex), indent=1)

    # Operating income
    _write_line("Operating Income", is_stmt.operating_income, bold=True)

    # Other income/expense
    _write_line("Other Income / (Expense)", is_stmt.total_other, indent=1)

    # Pre-tax income
    _write_line("Pre-Tax Income", is_stmt.pre_tax_income, bold=True)

    path = output_dir / _INPUT_DIR / "consolidated_pl_fy2025.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, path)

    manifest.register(
        f"{_INPUT_DIR}/consolidated_pl_fy2025.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── State factors xlsx ──────────────────────────────────────────────────────


def _write_state_factors(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> dict[str, dict[str, Decimal | None]]:
    """Write state_factors.xlsx and return the computed factors dict."""
    factors = _compute_state_factors(model)

    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    ws = wb.active
    ws.title = "State Factors"

    canary_code = canaries.canary_for("tc10_state_factors")
    location = embed_canary_xlsx(wb, canary_code)
    canaries.set_location(
        "tc10_state_factors",
        f"{_INPUT_DIR}/state_factors.xlsx",
        location,
    )

    # Headers
    headers = ["State", "Sales Factor", "Payroll Factor", "Property Factor"]
    for col, hdr in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=hdr)
    _style_header(ws, 1, len(headers))

    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 16

    # ── ERR-003: Transpose digits in IL sales factor ────────────────
    _ERR003_STATE = "IL"
    correct_sales = factors[_ERR003_STATE]["sales"]
    assert correct_sales is not None
    # Work with integer representation (factor * 10000) to transpose digits
    correct_int = int(correct_sales * 10000)
    # Try several position pairs until we find one with different digits
    for _p1, _p2 in [(-1, -2), (-2, -3), (0, 1), (1, 2)]:
        try:
            corrupt_int = transpose_digits(correct_int, pos1=_p1, pos2=_p2)
            break
        except ValueError:
            continue
    else:
        raise ValueError(f"Cannot transpose any digit pair in {correct_int}")
    corrupt_sales = Decimal(corrupt_int) / Decimal(10000)
    corrupt_sales_float = float(corrupt_sales)
    correct_sales_float = float(correct_sales)
    errors.add(PlantedError(
        error_id="ERR-003",
        file=f"{_INPUT_DIR}/state_factors.xlsx",
        location=f"Sheet 'State Factors', {_ERR003_STATE} Sales Factor",
        type="transposed_digits",
        description=(
            f"{_ERR003_STATE} sales apportionment factor shows "
            f"{corrupt_sales_float} instead of {correct_sales_float}"
        ),
        severity="material",
        which_test_cases_should_catch=["TC-10"],
    ))

    # State order per spec
    state_order = ["OR", "TX", "IL", "CA", "WA", "NY"]
    for i, state in enumerate(state_order):
        row = i + 2
        sf = factors[state]

        cell_a = ws.cell(row=row, column=1, value=state)
        _style_data_cell(cell_a)

        # Sales factor — use corrupted value for ERR-003 target state
        cell_b = ws.cell(row=row, column=2)
        if sf["sales"] is not None:
            display_sales = float(corrupt_sales) if state == _ERR003_STATE else float(sf["sales"])
            cell_b.value = display_sales
            _style_data_cell(cell_b, _PCT_FMT)
        else:
            # NY: leave blank (None → empty cell)
            _style_data_cell(cell_b)

        # Payroll factor
        cell_c = ws.cell(row=row, column=3)
        if state in ("OR", "TX", "IL"):
            # Complete data: could be zero or non-zero
            val = sf["payroll"]
            if val is not None and val == Decimal(0):
                # "$0" — zero presence (the agent must distinguish from blank)
                cell_c.value = "$0"
                _style_data_cell(cell_c)
            elif val is not None:
                cell_c.value = float(val)
                _style_data_cell(cell_c, _PCT_FMT)
            else:
                _style_data_cell(cell_c)
        else:
            # CA, WA, NY: blank (not provided)
            _style_data_cell(cell_c)

        # Property factor
        cell_d = ws.cell(row=row, column=4)
        if state in ("OR", "TX", "IL"):
            val = sf["property"]
            if val is not None and val == Decimal(0):
                cell_d.value = "$0"
                _style_data_cell(cell_d)
            elif val is not None:
                cell_d.value = float(val)
                _style_data_cell(cell_d, _PCT_FMT)
            else:
                _style_data_cell(cell_d)
        else:
            _style_data_cell(cell_d)

    path = output_dir / _INPUT_DIR / "state_factors.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, path)

    manifest.register(
        f"{_INPUT_DIR}/state_factors.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=[_TC],
    )

    return factors


# ── Apportionment rules docx ───────────────────────────────────────────────


def _write_apportionment_rules(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write apportionment_rules.docx — reference table of state formulas."""
    from docx import Document

    doc = Document()

    canary_code = canaries.canary_for("tc10_apportionment_rules")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "tc10_apportionment_rules",
        f"{_INPUT_DIR}/apportionment_rules.docx",
        location,
    )

    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME

    doc.add_heading("Cascade Industries, Inc.", level=1)
    doc.add_heading("Multi-State Apportionment Rules Reference", level=2)
    doc.add_paragraph("")

    doc.add_paragraph(
        "This document summarizes the apportionment formulas applicable to each "
        "state where Cascade Industries has operations or potential nexus. "
        "Refer to the state_factors.xlsx schedule for the actual factor data."
    )
    doc.add_paragraph("")

    # Reference table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "State"
    hdr[1].text = "Tax Type"
    hdr[2].text = "Apportionment Formula"
    hdr[3].text = "Notes"

    state_data = [
        (
            "Oregon (OR)",
            "Corporate Excise Tax (6.6%)",
            "Single sales factor",
            "Standard income tax apportionment.",
        ),
        (
            "Texas (TX)",
            "Margin Tax (Franchise Tax)",
            "Gross receipts less deductions",
            "NOT an income tax. Tax base is total revenue minus the greater of: "
            "(1) COGS, (2) compensation, or (3) 30% of total revenue, "
            "then apportioned by gross receipts ratio. Rate: 0.75% for most entities.",
        ),
        (
            "Illinois (IL)",
            "Corporate Income Tax (7.0% + 2.5% replacement)",
            "Single sales factor",
            "Combined rate 9.5%. Market-based sourcing for services.",
        ),
        (
            "California (CA)",
            "Corporate Tax (8.84%)",
            "Single sales factor, market-based sourcing",
            "Requires customer location data for proper sales sourcing. "
            "Minimum franchise tax of $800 applies.",
        ),
        (
            "Washington (WA)",
            "Business & Occupation (B&O) Tax",
            "Gross receipts — no apportionment",
            "NOT an income tax. Applies to gross receipts (not net income). "
            "Multiple rate classifications based on activity type. "
            "No deduction for COGS. Evaluate economic nexus thresholds.",
        ),
        (
            "New York (NY)",
            "Corporate Franchise Tax (7.1%)",
            "Single sales factor, customer-based sourcing",
            "Receipts sourced to where the customer receives benefit. "
            "Business capital tax also applies (0.04% of capital, capped at $5M).",
        ),
    ]
    for state, tax_type, formula, notes in state_data:
        row = table.add_row().cells
        row[0].text = state
        row[1].text = tax_type
        row[2].text = formula
        row[3].text = notes

    doc.add_paragraph("")
    doc.add_heading("Key Considerations", level=2)

    doc.add_paragraph(
        "1. Zero vs. Missing Data: The state_factors.xlsx schedule may contain "
        "cells showing \"$0\" (zero presence — the company has confirmed no "
        "activity in that factor for the state) versus blank cells (data not "
        "provided — requires follow-up)."
    )
    doc.add_paragraph(
        "2. Non-Income Tax States: Texas (margin tax) and Washington (B&O tax) "
        "do not impose traditional income taxes. Their tax bases differ "
        "fundamentally from income-based states and require separate computations."
    )
    doc.add_paragraph(
        "3. Nexus Evaluation: For states with limited factor data (e.g., only "
        "a sales factor), consider whether economic nexus thresholds are met "
        "before computing a tax liability."
    )

    path = output_dir / _INPUT_DIR / "apportionment_rules.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, path)

    manifest.register(
        f"{_INPUT_DIR}/apportionment_rules.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Prompt & Expected Behavior ──────────────────────────────────────────────


def _write_prompt(output_dir: Path) -> None:
    """Write test_cases/TC-10/prompt.md per spec."""
    text = """\
Complete the multi-state apportionment and tax analysis for Cascade Industries.

1. Complete the apportionment schedule for all states.
2. Flag states with incomplete factor data and note what's missing.
3. Apply each state's apportionment formula to compute state taxable income.
4. For states with non-standard tax bases (TX margin tax, WA B&O tax), note that
   a different computation is needed and explain what information is required.
5. Flag any states where nexus may be questionable based on the factor data.
6. Summarize estimated state tax liability by jurisdiction.

Export as an Excel workbook.
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write test_cases/TC-10/expected_behavior.md per spec."""
    text = """\
# TC-10: Multi-State Apportionment — Expected Behavior

## Apportionment Schedule
- The agent should produce a complete apportionment schedule covering all 6 states:
  OR, TX, IL, CA, WA, NY.
- For OR, TX, IL: all three factors (sales, payroll, property) are available and
  should be used per each state's formula.
- For CA and WA: only the sales factor is provided. The agent should flag the
  missing payroll and property factors.
- For NY: all factor data is blank. The agent should flag this as incomplete and
  note that apportionment cannot be computed without the data.

## Zero vs. Missing Distinction
- The state_factors.xlsx schedule contains "$0" values (zero presence confirmed)
  and blank cells (data not provided). The agent MUST distinguish between these:
  - "$0" means the company has confirmed zero activity for that factor in that state.
  - Blank means the data has not been provided and requires follow-up.

## Non-Standard Tax States
- Texas (TX): The agent should recognize that Texas imposes a margin tax (franchise
  tax), not an income tax. It should NOT attempt to compute TX tax using pre-tax
  income as the base. Instead, it should explain that the margin tax base is gross
  revenue minus deductions (COGS, compensation, or 30% of revenue) and note what
  additional information is needed.
- Washington (WA): The agent should recognize that Washington imposes a B&O tax on
  gross receipts, not net income. It should NOT attempt to compute WA tax using
  apportioned income. It should explain the B&O tax structure and required data.

## Nexus Question
- WA has only a sales factor with no payroll or property data. The agent should
  raise a nexus question: without physical presence indicators, economic nexus
  thresholds should be evaluated before assuming tax liability.

## Income-Based States
- OR: Single sales factor. Apply sales factor to pre-tax income, tax at 6.6%.
- IL: Single sales factor. Apply sales factor to pre-tax income, tax at 9.5%.
- CA: Single sales factor (incomplete data). Can estimate based on available sales
  factor, but should note market-based sourcing requires customer location data.

## Output Quality
- Excel workbook with clear state-by-state layout.
- Distinct treatment of zero-presence vs. missing data.
- Estimated tax liability summary for computable states.
- Clear identification of what additional information is needed for TX, WA, and NY.
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold Standard ────────────────────────────────────────────────────────────


@register_gold("TC-10")
def _tc10_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """Build the TC-10 gold standard from the canonical model."""
    model: CascadeModel = model_kwargs["model"]
    is_stmt = build_income_statement(model.ledger, 2025)
    factors = _compute_state_factors(model)
    apportioned = _compute_apportioned_income(model, factors)

    # Build state-by-state expected outputs
    state_outputs: dict[str, dict[str, Any]] = {}
    for state in ["CA", "IL", "NY", "OR", "TX", "WA"]:
        sf = factors[state]
        ap = apportioned[state]

        state_info: dict[str, Any] = {
            "formula": ap["formula"],
            "tax_type": ap["tax_type"],
        }

        # Factors
        for factor_name in ["sales", "payroll", "property"]:
            val = sf[factor_name]
            if val is not None:
                state_info[f"{factor_name}_factor"] = float(val)
            else:
                state_info[f"{factor_name}_factor"] = "not_provided"

        # Apportioned income and tax
        if ap["apportioned_income"] is not None:
            state_info["apportioned_income"] = _fmt_dollars(ap["apportioned_income"])
        else:
            state_info["apportioned_income"] = "cannot_compute"

        if ap["estimated_tax"] is not None:
            state_info["estimated_tax"] = _fmt_dollars(ap["estimated_tax"])
        else:
            state_info["estimated_tax"] = "cannot_compute"

        if ap["notes"]:
            state_info["notes"] = ap["notes"]

        state_outputs[state] = state_info

    return GoldStandard(
        test_case="TC-10",
        expected_outputs={
            "file_type": "xlsx",
            "pre_tax_income": _fmt_dollars(is_stmt.pre_tax_income),
            "states": state_outputs,
            "flags": {
                "tx_margin_tax": "Texas uses margin tax, not income tax — different base required",
                "wa_bno_tax": "Washington uses B&O tax on gross receipts, not income tax",
                "wa_nexus_question": "WA has only sales data, no physical presence — evaluate economic nexus",
                "ca_incomplete_data": "California has sales factor only, missing payroll and property",
                "ny_no_data": "New York has no factor data provided",
                "zero_vs_missing": "Agent must distinguish '$0' (zero presence) from blank (not provided)",
            },
        },
        canary_verification={
            "read_consolidated_pl": canaries.canary_for("tc10_consolidated_pl"),
            "read_state_factors": canaries.canary_for("tc10_state_factors"),
            "read_apportionment_rules": canaries.canary_for("tc10_apportionment_rules"),
        },
        error_detection={
            "ERR-003": "IL sales factor has transposed digits",
            "ERR-023": "Total Operating Expenses has a rounding discrepancy",
        },
        scoring_hints={
            "correctness": (
                "Apportioned income correct for OR and IL (single sales factor); "
                "TX and WA flagged as non-income-tax states; CA and NY flagged "
                "for incomplete data; WA nexus question raised"
            ),
            "completeness": (
                "Complete schedule for all 6 states; incomplete data flagged; "
                "estimated tax for computable states; missing info listed for TX/WA/NY"
            ),
            "format_compliance": (
                "Valid xlsx; clear state-by-state layout; '$0' vs blank distinction preserved"
            ),
            "robustness": (
                "Distinguished '$0' from blank; applied correct formula per state; "
                "identified TX and WA as non-income taxes"
            ),
            "communication": (
                "Explained each state's formula; flagged TX margin tax and WA B&O "
                "with specific information needed; raised nexus question for WA"
            ),
        },
    )


# ── Public entry point ──────────────────────────────────────────────────────


def emit_tc10(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write all TC-10 files to *output_dir*."""
    _write_consolidated_pl(model, output_dir, canaries, errors, manifest)
    _write_state_factors(model, output_dir, canaries, errors, manifest)
    _write_apportionment_rules(output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
