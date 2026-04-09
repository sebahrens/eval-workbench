"""Formatter: TC-11 — Quality of Earnings (Advisory, Complex).

Emits:
- test_cases/TC-11/input_files/monthly_pl_fy2023_fy2024_fy2025.xlsx
  36 months of P&L data with line-item detail, targeting reported EBITDA ≈$28.4M
- test_cases/TC-11/input_files/management_adjustments.xlsx
  6 management-proposed EBITDA adjustments totalling $1.855M (2 traps)
- test_cases/TC-11/input_files/customer_contracts/CTR-001.pdf .. CTR-008.pdf
  8 customer contracts with terms, volumes, pricing, renewal dates
- test_cases/TC-11/input_files/management_interview_notes.docx
  Management Q&A notes covering business overview, growth, concentration, litigation
- test_cases/TC-11/prompt.md
- test_cases/TC-11/expected_behavior.md
- gold_standards/TC-11_gold.json

One planted error: ERR-011 classification_error — consulting fees appear in both
FY2024 and FY2025 P&L but management labels them "non-recurring" in adjustments.
The agent should recognise the recurrence and challenge the adjustment.

Uses the canonical model for revenue; constructs cost structure to hit
spec-mandated EBITDA ≈$28.4M (FY2025).
"""

from __future__ import annotations

import datetime
import io
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import docx
import openpyxl
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from generator.canaries import (
    CanaryRegistry,
    embed_canary_docx,
    embed_canary_xlsx,
)
from generator.errors import ErrorRegistry, PlantedError
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel
from generator.model.customers import (
    CONTRACTS,
    compute_customer_concentration,
    contracts_expiring_within,
    top_n_concentration,
)

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-11"
_INPUT_DIR = f"test_cases/{_TC}/input_files"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)

# Target EBITDA values per spec gold standard
_REPORTED_EBITDA_FY2025 = Decimal("28_400_000")

# Management adjustment amounts (from prompt.md TC-11)
_ADJUSTMENTS = [
    {
        "id": "ADJ-001",
        "description": "Owner compensation above-market",
        "amount": Decimal("180_000"),
        "category": "Compensation",
        "management_classification": "Non-recurring",
        "is_truly_nonrecurring": True,
        "notes": (
            "CEO Robert Cascade's compensation includes $180K above market "
            "rate for comparable role. Standard add-back in owner-operated "
            "businesses."
        ),
    },
    {
        "id": "ADJ-002",
        "description": "One-time legal settlement",
        "amount": Decimal("420_000"),
        "category": "Legal",
        "management_classification": "Non-recurring",
        "is_truly_nonrecurring": True,
        "notes": (
            "Settlement of Henderson v. Cascade Precision Components "
            "product liability claim. Paid in Q2 FY2025."
        ),
    },
    {
        "id": "ADJ-003",
        "description": "COVID-related PPP loan forgiveness",
        "amount": Decimal("250_000"),
        "category": "Government programs",
        "management_classification": "Non-recurring",
        "is_truly_nonrecurring": True,  # Non-recurring but stale (FY2023)
        "notes": (
            "PPP loan forgiveness recognised in FY2023. Management "
            "proposes adding back to normalise earnings."
        ),
    },
    {
        "id": "ADJ-004",
        "description": "Non-recurring consulting fees",
        "amount": Decimal("95_000"),
        "category": "Professional fees",
        "management_classification": "Non-recurring",
        "is_truly_nonrecurring": False,  # TRAP: appears in 2 of 3 years
        "notes": (
            "Strategic consulting engagement for operational improvement. "
            "Management classifies as non-recurring."
        ),
    },
    {
        "id": "ADJ-005",
        "description": "Run-rate adjustment for new Q4 customer",
        "amount": Decimal("600_000"),
        "category": "Revenue run-rate",
        "management_classification": "Pro forma",
        "is_truly_nonrecurring": False,  # TRAP: aggressive — 1 quarter only
        "notes": (
            "New customer won in Q4 FY2025 with $150K in initial orders. "
            "Management annualises to $600K as a pro-forma run-rate "
            "adjustment."
        ),
    },
    {
        "id": "ADJ-006",
        "description": "Facility relocation costs",
        "amount": Decimal("310_000"),
        "category": "Facilities",
        "management_classification": "Non-recurring",
        "is_truly_nonrecurring": True,
        "notes": (
            "One-time costs associated with Advanced Materials division "
            "moving to expanded facility in Q3 FY2025."
        ),
    },
]

_TOTAL_ADJUSTMENTS = sum(a["amount"] for a in _ADJUSTMENTS)
assert _TOTAL_ADJUSTMENTS == Decimal("1_855_000"), f"Expected $1.855M, got {_TOTAL_ADJUSTMENTS}"

# P&L line item structure — maps account name to approximate annual amounts
# Designed so FY2025 EBITDA ≈ $28.4M
_PL_LINE_ITEMS = [
    # Revenue lines (will be scaled to model revenue)
    ("Product Sales", "Revenue", Decimal("0.72")),  # fraction of total revenue
    ("Services Revenue", "Revenue", Decimal("0.23")),
    ("Other Revenue", "Revenue", Decimal("0.05")),
    # COGS lines (as fraction of revenue)
    ("Raw Materials", "COGS", Decimal("0.31")),
    ("Direct Labor", "COGS", Decimal("0.18")),
    ("Manufacturing Overhead", "COGS", Decimal("0.10")),
    ("Freight & Delivery", "COGS", Decimal("0.04")),
    # OpEx / SGA lines (absolute amounts for FY2025, scaled for other years)
    # These are tuned so: Revenue - COGS - OpEx = Operating Income
    # and Operating Income + D&A = EBITDA ≈ $28.4M
]

# FY2025 opex line items (absolute amounts, in dollars)
# Total opex = Revenue - COGS - Operating_Income
# COGS ratio = 0.63, so COGS ≈ $126M
# D&A ≈ $2.1M (from model), so Operating_Income = EBITDA - D&A = $26.3M
# Total opex = $200M - $126M - $26.3M = $47.7M
# Of that, D&A = $2.1M, cash opex = $45.6M
_OPEX_LINES_FY2025: list[tuple[str, Decimal]] = [
    ("Salaries & Benefits", Decimal("22_400_000")),
    ("Rent & Facilities", Decimal("4_800_000")),
    ("Professional Fees", Decimal("2_200_000")),
    ("Insurance", Decimal("1_850_000")),
    ("Marketing & Sales", Decimal("3_100_000")),
    ("Travel & Entertainment", Decimal("1_400_000")),
    ("Depreciation & Amortization", Decimal("2_100_000")),
    ("IT & Software", Decimal("2_600_000")),
    ("Office & General", Decimal("1_250_000")),
    ("Research & Development", Decimal("4_200_000")),
    ("Other Operating Expenses", Decimal("1_700_000")),
]

_COGS_RATIO = Decimal("0.63")


# ── Deterministic save helpers ──────────────────────────────────────────────


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)

    wb.properties.created = _FIXED_DATETIME
    wb.properties.modified = _FIXED_DATETIME

    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    fixed_date_time = (2025, 3, 15, 9, 0, 0)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=fixed_date_time)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)

    buf = io.BytesIO()
    doc.save(buf)

    fixed_date_time = (2025, 3, 15, 9, 0, 0)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=fixed_date_time)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _whole_dollars(d: Decimal) -> int:
    """Round a Decimal to whole dollars."""
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


# ── Monthly P&L computation ────────────────────────────────────────────────


def _build_monthly_pl(
    model: CascadeModel,
) -> list[dict[str, Any]]:
    """Build 36 months of P&L data from the canonical model.

    Revenue comes from model.revenue_records (FY2023–2025).
    COGS is derived as a fixed ratio of revenue.
    OpEx uses FY2025 absolute amounts, scaled proportionally for FY2023/FY2024.
    """
    # Aggregate monthly revenue from model
    monthly_revenue: dict[tuple[int, int], Decimal] = {}
    for rec in model.revenue_records:
        if rec.year in (2023, 2024, 2025):
            key = (rec.year, rec.month)
            monthly_revenue[key] = monthly_revenue.get(key, Decimal(0)) + rec.revenue

    # Annual revenue totals for scaling
    annual_revenue: dict[int, Decimal] = {}
    for (yr, _), rev in monthly_revenue.items():
        annual_revenue[yr] = annual_revenue.get(yr, Decimal(0)) + rev

    # Revenue mix fractions
    rev_fractions = [
        ("Product Sales", Decimal("0.72")),
        ("Services Revenue", Decimal("0.23")),
        ("Other Revenue", Decimal("0.05")),
    ]

    # COGS lines as fraction of monthly revenue
    cogs_fractions = [
        ("Raw Materials", Decimal("0.31")),
        ("Direct Labor", Decimal("0.18")),
        ("Manufacturing Overhead", Decimal("0.10")),
        ("Freight & Delivery", Decimal("0.04")),
    ]

    rows: list[dict[str, Any]] = []

    for year in (2023, 2024, 2025):
        # Scale factor for opex: proportional to revenue relative to FY2025
        opex_scale = annual_revenue[year] / annual_revenue[2025]

        for month in range(1, 13):
            total_rev = monthly_revenue.get((year, month), Decimal(0))
            if total_rev == 0:
                continue

            row: dict[str, Any] = {
                "year": year,
                "month": month,
            }

            # Revenue breakdown
            remaining_rev = total_rev
            for i, (name, frac) in enumerate(rev_fractions):
                if i == len(rev_fractions) - 1:
                    row[name] = remaining_rev  # plug last line
                else:
                    val = (total_rev * frac).quantize(
                        Decimal("0.01"), rounding=ROUND_HALF_UP,
                    )
                    row[name] = val
                    remaining_rev -= val

            row["Total Revenue"] = total_rev

            # COGS breakdown
            total_cogs = Decimal(0)
            for name, frac in cogs_fractions:
                val = (total_rev * frac).quantize(
                    Decimal("0.01"), rounding=ROUND_HALF_UP,
                )
                row[name] = val
                total_cogs += val

            row["Total COGS"] = total_cogs
            row["Gross Profit"] = total_rev - total_cogs

            # OpEx breakdown — monthly = annual / 12, scaled by year
            total_opex = Decimal(0)
            for name, annual_amt in _OPEX_LINES_FY2025:
                monthly_amt = (annual_amt * opex_scale / 12).quantize(
                    Decimal("0.01"), rounding=ROUND_HALF_UP,
                )
                # Inject the consulting fee trap: include $95K in FY2024 and
                # FY2025 to make it visibly recurring
                if name == "Professional Fees" and year in (2024, 2025):
                    # Bump professional fees slightly — the recurring consulting
                    # sits inside this line item
                    monthly_amt += Decimal("7917")  # ~$95K / 12
                row[name] = monthly_amt
                total_opex += monthly_amt

            row["Total Operating Expenses"] = total_opex
            row["Operating Income"] = row["Gross Profit"] - total_opex
            row["EBITDA"] = row["Operating Income"] + row.get(
                "Depreciation & Amortization", Decimal(0),
            )

            rows.append(row)

    return rows


def _compute_annual_ebitda(rows: list[dict[str, Any]], year: int) -> Decimal:
    """Sum EBITDA for a given year from monthly P&L rows."""
    return sum(r["EBITDA"] for r in rows if r["year"] == year)


# ── Monthly P&L XLSX ───────────────────────────────────────────────────────


def _write_monthly_pl(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> list[dict[str, Any]]:
    """Write monthly_pl_fy2023_fy2024_fy2025.xlsx and return the P&L rows."""
    file_key = "tc11_monthly_pl"
    canary_code = canaries.canary_for(file_key)

    wb = openpyxl.Workbook()
    loc = embed_canary_xlsx(wb, canary_code)

    pl_rows = _build_monthly_pl(model)

    ws = wb.active
    ws.title = "Monthly P&L"

    # Styles
    header_fill = PatternFill("solid", fgColor="1A3C6E")
    header_font = Font(bold=True, size=11, color="FFFFFF")
    bold_font = Font(bold=True, size=11)
    number_fmt = "#,##0"
    border = Border(bottom=Side(style="thin", color="CCCCCC"))

    # Title
    ws["A1"] = "Cascade Industries, Inc."
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "Profit & Loss Statement — Monthly Detail (FY2023–FY2025)"
    ws["A2"].font = Font(bold=True, size=12, color="666666")
    ws.merge_cells("A1:F1")
    ws.merge_cells("A2:F2")

    # Column headers — Year, Month, then each P&L line
    pl_columns = [
        "Product Sales", "Services Revenue", "Other Revenue", "Total Revenue",
        "Raw Materials", "Direct Labor", "Manufacturing Overhead",
        "Freight & Delivery", "Total COGS", "Gross Profit",
        "Salaries & Benefits", "Rent & Facilities", "Professional Fees",
        "Insurance", "Marketing & Sales", "Travel & Entertainment",
        "Depreciation & Amortization", "IT & Software", "Office & General",
        "Research & Development", "Other Operating Expenses",
        "Total Operating Expenses", "Operating Income", "EBITDA",
    ]
    all_headers = ["Year", "Month"] + pl_columns

    for col_idx, header in enumerate(all_headers, start=1):
        cell = ws.cell(row=4, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    # Data rows
    for row_idx, pl_row in enumerate(pl_rows, start=5):
        ws.cell(row=row_idx, column=1, value=pl_row["year"])
        ws.cell(row=row_idx, column=2, value=pl_row["month"])

        for col_idx, col_name in enumerate(pl_columns, start=3):
            val = pl_row.get(col_name, Decimal(0))
            cell = ws.cell(
                row=row_idx,
                column=col_idx,
                value=_whole_dollars(val),
            )
            cell.number_format = number_fmt
            cell.border = border

    # Subtotal rows for each year
    subtotal_row = 5 + len(pl_rows)
    for yr in (2023, 2024, 2025):
        yr_rows = [r for r in pl_rows if r["year"] == yr]
        ws.cell(row=subtotal_row, column=1, value=f"FY{yr} Total")
        ws.cell(row=subtotal_row, column=1).font = bold_font

        for col_idx, col_name in enumerate(pl_columns, start=3):
            total = sum(r.get(col_name, Decimal(0)) for r in yr_rows)
            cell = ws.cell(
                row=subtotal_row,
                column=col_idx,
                value=_whole_dollars(total),
            )
            cell.number_format = number_fmt
            cell.font = bold_font
            cell.border = Border(
                top=Side(style="thin"), bottom=Side(style="double"),
            )

        subtotal_row += 1

    # Column widths
    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 8
    for col_idx in range(3, len(all_headers) + 1):
        ws.column_dimensions[
            openpyxl.utils.get_column_letter(col_idx)
        ].width = 16

    # Register planted error ERR-011: consulting fees classified as
    # non-recurring even though they appear in FY2024 and FY2025.
    # The error is in the *management_adjustments* file where ADJ-004 labels
    # it non-recurring, but the P&L shows the fee in both years.
    # We register it here so the P&L file is recorded as the evidence file.
    errors.add(PlantedError(
        error_id="ERR-011",
        file=f"{_INPUT_DIR}/management_adjustments.xlsx",
        location="Sheet 'Adjustments', Row for ADJ-004 'Non-recurring consulting fees'",
        type="classification_error",
        description=(
            "Management classifies $95K consulting fees as non-recurring "
            "(ADJ-004), but the monthly P&L shows similar professional fees "
            "in both FY2024 and FY2025 (~$95K/year above baseline). "
            "The adjustment is actually recurring and should be rejected."
        ),
        severity="significant",
        which_test_cases_should_catch=["TC-11"],
    ))

    path = output_dir / f"{_INPUT_DIR}/monthly_pl_fy2023_fy2024_fy2025.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, path)

    canaries.set_location(
        file_key,
        f"{_INPUT_DIR}/monthly_pl_fy2023_fy2024_fy2025.xlsx",
        loc,
    )
    manifest.register(
        f"{_INPUT_DIR}/monthly_pl_fy2023_fy2024_fy2025.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=["TC-11"],
    )

    return pl_rows


# ── Management Adjustments XLSX ────────────────────────────────────────────


def _write_management_adjustments(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write management_adjustments.xlsx with 6 proposed adjustments."""
    file_key = "tc11_mgmt_adjustments"
    canary_code = canaries.canary_for(file_key)

    wb = openpyxl.Workbook()
    loc = embed_canary_xlsx(wb, canary_code)

    ws = wb.active
    ws.title = "Adjustments"

    # Styles
    header_fill = PatternFill("solid", fgColor="2E5090")
    header_font = Font(bold=True, size=11, color="FFFFFF")
    bold_font = Font(bold=True, size=11)
    number_fmt = "#,##0"

    # Title
    ws["A1"] = "Cascade Industries, Inc."
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "Management Proposed EBITDA Adjustments — FY2025"
    ws["A2"].font = Font(bold=True, size=12, color="666666")
    ws.merge_cells("A1:F1")
    ws.merge_cells("A2:F2")

    # Headers
    headers = [
        "Adjustment ID", "Description", "Amount ($)", "Category",
        "Classification", "Notes",
    ]
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    # Data rows
    for row_idx, adj in enumerate(_ADJUSTMENTS, start=5):
        ws.cell(row=row_idx, column=1, value=adj["id"])
        ws.cell(row=row_idx, column=2, value=adj["description"])
        cell = ws.cell(row=row_idx, column=3, value=_whole_dollars(adj["amount"]))
        cell.number_format = number_fmt
        ws.cell(row=row_idx, column=4, value=adj["category"])
        ws.cell(row=row_idx, column=5, value=adj["management_classification"])
        ws.cell(row=row_idx, column=6, value=adj["notes"])

    # Total row
    total_row = 5 + len(_ADJUSTMENTS)
    ws.cell(row=total_row, column=2, value="Total Proposed Adjustments")
    ws.cell(row=total_row, column=2).font = bold_font
    total_cell = ws.cell(
        row=total_row, column=3,
        value=_whole_dollars(_TOTAL_ADJUSTMENTS),
    )
    total_cell.number_format = number_fmt
    total_cell.font = bold_font
    total_cell.border = Border(
        top=Side(style="thin"), bottom=Side(style="double"),
    )

    # Column widths
    ws.column_dimensions["A"].width = 16
    ws.column_dimensions["B"].width = 36
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 20
    ws.column_dimensions["E"].width = 18
    ws.column_dimensions["F"].width = 50

    path = output_dir / f"{_INPUT_DIR}/management_adjustments.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, path)

    canaries.set_location(
        file_key,
        f"{_INPUT_DIR}/management_adjustments.xlsx",
        loc,
    )
    manifest.register(
        f"{_INPUT_DIR}/management_adjustments.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=["TC-11"],
    )


# ── Customer Contract PDFs ─────────────────────────────────────────────────


def _build_contract_elements(
    contract: Any,
    title_style: ParagraphStyle,
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
) -> list[Any]:
    """Build reportlab flowable elements for a single customer contract."""
    elements: list[Any] = []

    # Title
    elements.append(Paragraph("CUSTOMER AGREEMENT", title_style))
    elements.append(Spacer(1, 12))

    # Parties
    elements.append(Paragraph("PARTIES", heading_style))
    elements.append(Paragraph(
        f'This Agreement ("Agreement") is entered into between '
        f"<b>Cascade Industries, Inc.</b> through its subsidiary operating "
        f"under entity code <b>{contract.entity_code}</b> "
        f'("Supplier") and <b>{contract.customer_name}</b> '
        f'("Customer").',
        body_style,
    ))
    elements.append(Spacer(1, 6))

    # Contract details table
    elements.append(Paragraph("CONTRACT DETAILS", heading_style))

    coc_text = (
        "Yes — Customer may terminate within 90 days of ownership change"
        if contract.change_of_control_clause else "None"
    )
    auto_renew_text = (
        "Yes — automatic renewal for successive 12-month periods"
        if contract.auto_renew
        else "No — requires written renewal agreement"
    )

    detail_data = [
        ["Contract ID", contract.contract_id],
        ["Effective Date", contract.effective_date.strftime("%B %d, %Y")],
        ["Expiration Date", contract.expiration_date.strftime("%B %d, %Y")],
        ["Annual Volume", f"${contract.annual_volume:,.0f}"],
        ["Pricing Terms", contract.pricing_terms],
        ["Payment Terms", contract.payment_terms],
        ["Auto-Renewal", auto_renew_text],
        ["Change of Control", coc_text],
    ]

    detail_table = Table(
        detail_data,
        colWidths=[2 * inch, 4.5 * inch],
    )
    detail_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.Color(0.9, 0.9, 0.9)),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.7, 0.7, 0.7)),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
    ]))
    elements.append(detail_table)
    elements.append(Spacer(1, 12))

    # Terms and conditions
    elements.append(Paragraph("TERMS AND CONDITIONS", heading_style))

    terms = [
        (
            "1. SCOPE OF SUPPLY",
            f"Supplier agrees to manufacture and deliver products/services "
            f"to Customer in accordance with specifications mutually agreed "
            f"upon, at an approximate annual volume of "
            f"${contract.annual_volume:,.0f}.",
        ),
        (
            "2. PRICING",
            f"Pricing shall be determined on a {contract.pricing_terms.lower()} "
            f"basis. All prices are in US dollars.",
        ),
        (
            "3. PAYMENT",
            f"Customer shall pay all undisputed invoices within "
            f"{contract.payment_terms} of invoice date.",
        ),
        (
            "4. TERM AND RENEWAL",
            f"This Agreement is effective from "
            f"{contract.effective_date.strftime('%B %d, %Y')} through "
            f"{contract.expiration_date.strftime('%B %d, %Y')}. "
            f"{auto_renew_text}.",
        ),
        (
            "5. TERMINATION",
            "Either party may terminate this Agreement upon 90 days' "
            "written notice for material breach that remains uncured "
            "after 30 days' notice.",
        ),
    ]

    if contract.change_of_control_clause:
        terms.append((
            "6. CHANGE OF CONTROL",
            "In the event of a change of control of Supplier (defined as "
            "the acquisition of more than 50% of voting securities or "
            "substantially all assets), Customer may terminate this "
            "Agreement within 90 days of written notice of such change, "
            "without penalty or liability.",
        ))

    for heading, body in terms:
        elements.append(Paragraph(f"<b>{heading}</b>", body_style))
        elements.append(Paragraph(body, body_style))
        elements.append(Spacer(1, 4))

    # Notes if present
    if contract.notes:
        elements.append(Spacer(1, 8))
        elements.append(Paragraph("ADDITIONAL NOTES", heading_style))
        elements.append(Paragraph(contract.notes, body_style))

    # Signature block
    elements.append(Spacer(1, 24))
    sig_data = [
        ["FOR SUPPLIER:", "FOR CUSTOMER:"],
        ["", ""],
        ["_________________________", "_________________________"],
        ["Cascade Industries, Inc.", contract.customer_name],
        [f"Date: {contract.effective_date.strftime('%m/%d/%Y')}", "Date: ___/___/______"],
    ]
    sig_table = Table(sig_data, colWidths=[3.25 * inch, 3.25 * inch])
    sig_table.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
    ]))
    elements.append(sig_table)

    return elements


def _write_customer_contracts(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write 8 customer contract PDFs using reportlab."""
    contracts_dir = output_dir / f"{_INPUT_DIR}/customer_contracts"
    contracts_dir.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ContractTitle",
        parent=styles["Title"],
        fontSize=16,
        spaceAfter=20,
        alignment=1,  # center
    )
    heading_style = ParagraphStyle(
        "ContractHeading",
        parent=styles["Heading2"],
        fontSize=12,
        spaceBefore=12,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ContractBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )

    for i, contract in enumerate(CONTRACTS):
        file_key = f"tc11_contract_{i + 1:03d}"
        canary_code = canaries.canary_for(file_key)
        rel_path = f"{_INPUT_DIR}/customer_contracts/{contract.contract_id}.pdf"
        full_path = output_dir / rel_path

        elements = _build_contract_elements(
            contract, title_style, heading_style, body_style,
        )

        doc = SimpleDocTemplate(
            str(full_path),
            pagesize=letter,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
            title=f"Customer Agreement — {contract.contract_id}",
            author=f"CANARY: {canary_code}",
            creator="Cascade Industries Contract System",
            invariant=True,
        )
        doc.build(elements)

        canary_loc = "PDF metadata → Author"
        canaries.set_location(file_key, rel_path, canary_loc)
        manifest.register(rel_path, "pdf", canary=canary_code, test_cases=["TC-11"])


# ── Management Interview Notes DOCX ──────────────────────────────────────


def _write_interview_notes(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write management_interview_notes.docx."""
    file_key = "tc11_interview_notes"
    canary_code = canaries.canary_for(file_key)

    doc = docx.Document()
    embed_canary_docx(doc, canary_code)

    # Title
    title = doc.add_heading("Management Interview Notes", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Cascade Industries, Inc. — Quality of Earnings Due Diligence",
        style="Subtitle",
    )
    doc.add_paragraph(
        "Interview Date: March 10, 2026  |  Attendees: R. Cascade (CEO), "
        "M. Chen (CFO), D. Nakamura (CTO)",
    )
    doc.add_paragraph("")

    # Business Overview
    doc.add_heading("1. Business Overview", level=1)
    doc.add_paragraph(
        "Cascade Industries is a diversified industrial manufacturing company "
        "headquartered in Portland, Oregon, operating through three main "
        "subsidiaries: Cascade Precision Components (PC), Advanced Materials "
        "Corp (AM), and Distribution Services Inc (DS). The company was "
        "founded in 1987 and has grown to approximately $200 million in "
        "annual revenue."
    )
    doc.add_paragraph(
        "CEO Robert Cascade noted that the company has experienced steady "
        "growth over the past three years, with FY2025 representing the "
        "strongest performance in company history. He attributed the growth "
        "to investments in the Advanced Materials division and expansion of "
        "distribution capabilities."
    )

    # Growth Drivers
    doc.add_heading("2. Growth Drivers & Revenue Quality", level=1)
    doc.add_paragraph(
        "CFO Margaret Chen highlighted the following growth drivers:"
    )
    bullets = [
        "Precision Components (PC): Stable growth at ~5% annually, driven by "
        "long-term contracts with industrial OEMs. Acme Manufacturing remains "
        "the anchor customer.",
        "Advanced Materials (AM): Fastest growing division at ~18% YoY, driven "
        "by aerospace and defense demand. Recent government subcontracts "
        "through NextGen Composites have been significant.",
        "Distribution Services (DS): Moderate growth at ~6.5%, with expansion "
        "into third-party logistics services.",
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph(
        "Management noted a significant new customer win in Q4 FY2025 that "
        "they expect to contribute approximately $600,000 annually going "
        "forward. \"We signed them in October and already booked $150K in "
        "initial orders. Based on their projected needs, we're confident "
        "in the $600K annualized run-rate,\" said Mr. Cascade. He described "
        "the customer as \"a Fortune 500 company entering our space\" but "
        "acknowledged the relationship is still in early stages."
    )

    # Customer Concentration
    doc.add_heading("3. Customer Concentration", level=1)
    doc.add_paragraph(
        "When asked about customer concentration, CFO Chen acknowledged that "
        "Acme Manufacturing represents approximately 18% of consolidated "
        "revenue. \"They've been our largest customer for over a decade and "
        "the relationship is very stable,\" she said."
    )
    doc.add_paragraph(
        "However, she noted that the Acme contract expires at year-end "
        "(December 31, 2025) and does not auto-renew. \"We're in active "
        "renewal discussions. We've never had an issue renewing with Acme, "
        "but I understand the concern.\" The contract also contains a "
        "change-of-control termination provision."
    )
    doc.add_paragraph(
        "Management indicated that two other contracts are also expiring "
        "within the next 12 months: Columbia River Works (CTR-003, June 2025) "
        "and Quantum Materials (CTR-005, May 2025). Both are expected to "
        "renew, though CTR-005 requires renegotiation of pricing terms."
    )

    # Key Personnel
    doc.add_heading("4. Key Personnel & Compensation", level=1)
    doc.add_paragraph(
        "CEO Cascade's total compensation includes a salary that management "
        "acknowledges is approximately $180,000 above market rate for a "
        "comparable CEO role at a company of this size. \"It reflects my "
        "founder status and the fact that I've taken below-market comp in "
        "earlier years,\" he explained."
    )
    doc.add_paragraph(
        "CTO Nakamura is considered critical to the Advanced Materials "
        "division's technology roadmap. All three C-suite executives have "
        "employment agreements with change-of-control provisions."
    )

    # EBITDA Adjustments Discussion
    doc.add_heading("5. EBITDA Adjustments", level=1)
    doc.add_paragraph(
        "CFO Chen walked through the proposed EBITDA adjustments:"
    )
    adj_bullets = [
        "Owner compensation ($180K): \"Standard adjustment for an "
        "owner-operated business.\"",
        "Legal settlement ($420K): \"One-time settlement of the Henderson "
        "litigation. Fully paid in Q2 2025. We don't anticipate similar "
        "claims.\"",
        "PPP loan forgiveness ($250K): \"Recognised in FY2023. Not a "
        "recurring item.\"",
        "Consulting fees ($95K): \"We brought in McKinley & Associates for "
        "an operational improvement study. One-time engagement.\"",
        "Run-rate adjustment ($600K): \"Based on our new Q4 customer. We "
        "believe this represents the go-forward earnings power.\"",
        "Relocation costs ($310K): \"Advanced Materials moved to a larger "
        "facility in Q3 2025. One-time expense.\"",
    ]
    for bullet in adj_bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: When asked about consulting fees, CEO Cascade mentioned that "
        "\"we've worked with consultants periodically over the years — there "
        "was a similar engagement in FY2024 as well, though focused on a "
        "different area.\" This appears to contradict the 'non-recurring' "
        "classification."
    )

    # Pending Litigation
    doc.add_heading("6. Pending Litigation", level=1)
    doc.add_paragraph(
        "Management disclosed one pending litigation matter: Henderson v. "
        "Cascade Precision Components, a product liability claim with "
        "potential exposure of $2.5 million. The company has accrued $750,000 "
        "on the balance sheet. Outside counsel (Mitchell, Hartwell & "
        "Associates) believes the case will likely settle for less than the "
        "accrued amount. Depositions are scheduled for Q1 2026."
    )

    # Forward-Looking Items
    doc.add_heading("7. Forward-Looking Items", level=1)
    doc.add_paragraph(
        "Management identified the following items requiring further "
        "diligence:"
    )
    forward_bullets = [
        "Acme Manufacturing contract renewal status and pricing terms",
        "Advanced Materials division capacity constraints at the new facility",
        "Integration of new Q4 customer into standard production schedule",
        "Resolution timeline for Henderson litigation",
        "Impact of potential tariff changes on raw material costs",
    ]
    for bullet in forward_bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    rel_path = f"{_INPUT_DIR}/management_interview_notes.docx"
    path = output_dir / rel_path
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, path)

    canaries.set_location(
        file_key,
        rel_path,
        "Core properties → comments",
    )
    manifest.register(rel_path, "docx", canary=canary_code, test_cases=["TC-11"])


# ── Prompt ─────────────────────────────────────────────────────────────────


def _write_prompt(output_dir: Path) -> None:
    """Write the agent prompt for TC-11."""
    text = """\
# TC-11: Quality of Earnings Analysis

Perform a quality of earnings analysis for a potential acquisition of Cascade Industries.

## Input Files

- `monthly_pl_fy2023_fy2024_fy2025.xlsx` — 36 months of P&L data by account with line-item detail
- `management_adjustments.xlsx` — Management's proposed EBITDA adjustments
- `customer_contracts/` — 8 PDF contracts with top customers
- `management_interview_notes.docx` — Notes from management Q&A sessions

## Tasks

1. Compute reported EBITDA for each of the 36 months and annually.
2. Evaluate each of management's proposed adjustments:
   - Is it truly non-recurring?
   - Is the amount supportable?
   - Is it properly categorized (above/below the line)?
   - Challenge any adjustments that appear aggressive or recurring in nature.
3. Compute adjusted EBITDA after your accepted/modified adjustments.
4. Create a QofE bridge: Reported EBITDA → each adjustment → Adjusted EBITDA.
5. Analyze revenue quality:
   - Customer concentration (top 10 customers as % of revenue)
   - Contract renewal risk (any contracts expiring within 12 months)
   - Revenue trend sustainability
6. Draft the "Key Findings" section of the QofE report covering:
   - Adjusted EBITDA conclusion
   - Material adjustments and rationale
   - Revenue quality and risks
   - Items requiring further diligence

## Deliverables

Export:
- **Analysis workbook** (Excel) with EBITDA bridge, monthly detail, customer analysis
- **Key Findings memo** (Word document)
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Expected Behavior ──────────────────────────────────────────────────────


def _write_expected_behavior(output_dir: Path) -> None:
    """Write expected behavior notes for evaluators."""
    text = """\
# TC-11: Quality of Earnings — Expected Behavior

## Key Gold Standard Values

- **Reported EBITDA (FY2025)**: ~$28.4M
- **Management's total proposed adjustments**: $1,855,000
- **Properly adjusted EBITDA**: ~$29.3M

## Adjustment Evaluation

The agent should evaluate each adjustment individually:

| ID | Description | Amount | Expected Agent Response |
|----|-------------|--------|------------------------|
| ADJ-001 | Owner comp above-market | $180K | **Accept** — standard add-back for owner-operated business |
| ADJ-002 | Legal settlement | $420K | **Accept** — truly one-time (Henderson litigation) |
| ADJ-003 | PPP loan forgiveness | $250K | **Challenge/Reject** — stale; occurred in FY2023, not FY2025 |
| ADJ-004 | Consulting fees | $95K | **Reject** — recurring; similar fees appear in FY2024 P&L |
| ADJ-005 | Run-rate new customer | $600K | **Challenge** — aggressive; only 1 quarter of history ($150K actual) |
| ADJ-006 | Relocation costs | $310K | **Accept** — truly one-time facility move |

Accepted adjustments: $180K + $420K + $310K = **$910K**
Adjusted EBITDA: $28.4M + $0.91M ≈ **$29.31M** (~$29.3M)

## Error Detection

- **ERR-011**: Consulting fees (ADJ-004) are classified as non-recurring, but the
  monthly P&L shows professional fees bumped by ~$95K/year in both FY2024 and FY2025.
  The management interview notes also reveal CEO mentioned "a similar engagement
  in FY2024." The agent should identify the recurrence pattern and reject this
  adjustment.

## Data Challenges

- The PPP adjustment (ADJ-003) is for FY2023 but management includes it in an
  FY2025 adjustment schedule — the agent should note the temporal mismatch.
- The run-rate adjustment (ADJ-005) is based on only one quarter of actual orders.
  Management's Q4 interview enthusiasm should not override the thin data.
- Customer contracts show change-of-control clause for the top customer (Acme, 18%
  of revenue) — this is a significant risk in a potential acquisition.

## Revenue Quality

- Top customer (Acme Manufacturing): ~18% of consolidated revenue — flag
  concentration risk.
- 4 contracts expiring within 12 months of FY2025 year-end (CTR-001, CTR-002,
  CTR-007, CTR-008).
- Acme contract (CTR-001) has change-of-control termination provision — critical
  risk for acquisition context.

## Output Quality

- Analysis workbook should include: EBITDA bridge, monthly detail, customer
  analysis with concentration metrics.
- Key Findings memo should follow professional QofE report structure.
- Both files must be valid xlsx/docx that open without errors.
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold Standard ────────────────────────────────────────────────────────


@register_gold("TC-11")
def _tc11_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """Build the TC-11 gold standard from the canonical model."""
    model: CascadeModel = model_kwargs["model"]

    # Compute EBITDA from the P&L rows
    pl_rows = _build_monthly_pl(model)

    ebitda_fy2023 = _compute_annual_ebitda(pl_rows, 2023)
    ebitda_fy2024 = _compute_annual_ebitda(pl_rows, 2024)
    ebitda_fy2025 = _compute_annual_ebitda(pl_rows, 2025)

    # Annual revenue from model
    annual_rev: dict[int, Decimal] = {}
    for rec in model.revenue_records:
        if rec.year in (2023, 2024, 2025):
            annual_rev[rec.year] = annual_rev.get(rec.year, Decimal(0)) + rec.revenue

    # Customer concentration
    concentration = compute_customer_concentration(2025)
    top_customer = concentration[0]
    t10_pct = top_n_concentration(10, 2025)

    # Expiring contracts
    expiring = contracts_expiring_within(12)

    # Accepted adjustments (per spec gold standard)
    accepted = Decimal("180_000") + Decimal("420_000") + Decimal("310_000")
    adjusted_ebitda = ebitda_fy2025 + accepted

    return GoldStandard(
        test_case="TC-11",
        expected_outputs={
            "output_files": {
                "analysis_workbook": {
                    "type": "xlsx",
                    "required_sheets": [
                        "EBITDA Bridge",
                        "Monthly Detail",
                        "Customer Analysis",
                    ],
                },
                "key_findings_memo": {
                    "type": "docx",
                    "required_sections": [
                        "Adjusted EBITDA Conclusion",
                        "Material Adjustments",
                        "Revenue Quality",
                        "Items Requiring Further Diligence",
                    ],
                },
            },
            "financial_metrics": {
                "fy2023_revenue": _whole_dollars(annual_rev[2023]),
                "fy2024_revenue": _whole_dollars(annual_rev[2024]),
                "fy2025_revenue": _whole_dollars(annual_rev[2025]),
                "reported_ebitda_fy2023": _whole_dollars(ebitda_fy2023),
                "reported_ebitda_fy2024": _whole_dollars(ebitda_fy2024),
                "reported_ebitda_fy2025": _whole_dollars(ebitda_fy2025),
                "management_adjustments_total": _whole_dollars(_TOTAL_ADJUSTMENTS),
                "accepted_adjustments": _whole_dollars(accepted),
                "adjusted_ebitda_fy2025": _whole_dollars(adjusted_ebitda),
            },
            "adjustment_evaluation": {
                "ADJ-001": {"action": "accept", "amount": 180_000},
                "ADJ-002": {"action": "accept", "amount": 420_000},
                "ADJ-003": {"action": "reject", "reason": "stale — FY2023 item"},
                "ADJ-004": {"action": "reject", "reason": "recurring in FY2024 and FY2025"},
                "ADJ-005": {"action": "challenge", "reason": "aggressive — 1 quarter only"},
                "ADJ-006": {"action": "accept", "amount": 310_000},
            },
            "customer_concentration": {
                "top_customer_name": top_customer.customer_name,
                "top_customer_pct": float(top_customer.pct_of_consolidated),
                "top_10_pct": float(t10_pct),
            },
            "contract_renewal_risk": {
                "contracts_expiring_12m": len(expiring),
                "expiring_contract_ids": [c.contract_id for c in expiring],
                "change_of_control_contracts": ["CTR-001"],
            },
        },
        canary_verification={
            "read_monthly_pl": canaries.canary_for("tc11_monthly_pl"),
            "read_mgmt_adjustments": canaries.canary_for("tc11_mgmt_adjustments"),
            "read_interview_notes": canaries.canary_for("tc11_interview_notes"),
            "read_acme_contract": canaries.canary_for("tc11_contract_001"),
        },
        error_detection={
            "ERR-011": (
                "Consulting fees (ADJ-004) classified as non-recurring "
                "but appear in both FY2024 and FY2025 P&L"
            ),
        },
        scoring_hints={
            "correctness": (
                "Reported EBITDA within 1% of gold; accepted adjustments "
                "match gold; adjusted EBITDA within 1% of gold"
            ),
            "completeness": (
                "All 6 adjustments evaluated; EBITDA bridge present; "
                "customer concentration and renewal risk analysed"
            ),
            "judgment": (
                "ADJ-003 challenged as stale; ADJ-004 rejected as recurring; "
                "ADJ-005 challenged as aggressive; Acme concentration flagged"
            ),
            "communication": (
                "Professional QofE report language; clear rationale for each "
                "adjustment decision; forward-looking risk items identified"
            ),
        },
    )


# ── Public entry point ──────────────────────────────────────────────────────


def emit_tc11(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write all TC-11 files to *output_dir*."""
    _write_monthly_pl(model, output_dir, canaries, errors, manifest)
    _write_management_adjustments(output_dir, canaries, manifest)
    _write_customer_contracts(output_dir, canaries, manifest)
    _write_interview_notes(output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
