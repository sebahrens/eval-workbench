"""Formatter: TC-08-EU — European R&D Incentive Study (CIR + Forschungszulage).

Emits:
- test_cases/TC-08-EU/input_files/rd_employee_time_records_eu.csv
  ~2,184-row weekly time records for 42 R&D employees (30 CM + 12 CP)
- test_cases/TC-08-EU/input_files/rd_project_descriptions_eu/
  14 .docx project description files
- test_cases/TC-08-EU/input_files/payroll_data_eu_fy2025.xlsx
  Payroll register for CM and CP R&D employees
- test_cases/TC-08-EU/input_files/rd_supply_expenses_eu.xlsx
  Supply and materials expenses for CM and CP
- test_cases/TC-08-EU/input_files/rd_subcontractor_invoices_eu.xlsx
  Subcontracted R&D invoices for CM
- test_cases/TC-08-EU/input_files/prior_year_rd_data_eu.xlsx
  Prior year R&D spend (FY2023, FY2024) — contextual, NOT for averaging
- test_cases/TC-08-EU/prompt.md
- test_cases/TC-08-EU/expected_behavior.md
- gold_standards/TC-08-EU_gold.json

Planted error:
  ERR-EU-008: cost_category_misclassification — maintenance contract miscoded
  as R&D supply at CP (CP-RD-03)

Uses deterministic European R&D model — never hardcodes numbers that should
come from the model.
"""

from __future__ import annotations

import datetime
import io
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import docx
import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from generator.canaries import (
    CanaryRegistry,
    embed_canary_csv_comment,
    embed_canary_docx,
    embed_canary_xlsx,
)
from generator.errors import ErrorRegistry, PlantedError
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel
from generator.model.rd_eu import (
    _ERR_EU_008_EXPENSE,
    PRIOR_YEAR_RD_SPEND,
    RD_PROJECTS_EU,
    compute_consolidated_rd_benefit,
    generate_rd_employees_eu,
    generate_subcontractor_invoices_eu,
    generate_supply_expenses_eu,
    generate_time_records_eu,
)

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-08-EU"
_INPUT_DIR = f"test_cases/{_TC}/input_files"
_PROJ_DESC_DIR = f"{_INPUT_DIR}/rd_project_descriptions_eu"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)
_FIXED_ZIP_DT = (2025, 3, 15, 9, 0, 0)

# Canary keys
_KEY_TIME_RECORDS = "tc08eu_time_records"
_KEY_PAYROLL = "tc08eu_payroll"
_KEY_SUPPLY = "tc08eu_supply_expenses"
_KEY_SUBCONTRACTOR = "tc08eu_subcontractor_invoices"
_KEY_PRIOR_YEAR = "tc08eu_prior_year_rd"
_PROJECT_DOC_KEYS = [f"tc08eu_project_{i:03d}" for i in range(1, 15)]

# Styles
_HEADER_FILL = PatternFill("solid", fgColor="1A3C6E")
_HEADER_FONT = Font(bold=True, size=11, color="FFFFFF")
_NUM_FMT = "#,##0"
_NUM_FMT_2 = "#,##0.00"
_THIN_BORDER = Border(bottom=Side(style="thin", color="CCCCCC"))


# ── Deterministic save helpers ───────────────────────────────────────────────

def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)
    wb.properties.created = _FIXED_DATETIME
    wb.properties.modified = _FIXED_DATETIME

    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


# ── Time Records CSV ─────────────────────────────────────────────────────────

def _write_time_records_csv(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write rd_employee_time_records_eu.csv."""
    records = generate_time_records_eu()

    rel_path = f"{_INPUT_DIR}/rd_employee_time_records_eu.csv"
    abs_path = output_dir / rel_path
    abs_path.parent.mkdir(parents=True, exist_ok=True)

    canary = canaries.canary_for(_KEY_TIME_RECORDS)
    canary_line = embed_canary_csv_comment(canary)

    columns = [
        "entity_code",
        "employee_id",
        "employee_name",
        "week_ending",
        "project_code",
        "hours",
        "activity_description",
    ]

    lines: list[str] = [canary_line, ",".join(columns) + "\n"]
    for rec in records:
        desc = rec.activity_description
        if "," in desc or '"' in desc:
            desc = '"' + desc.replace('"', '""') + '"'
        row = [
            rec.entity_code,
            rec.employee_id,
            rec.employee_name,
            rec.week_ending,
            rec.project_code,
            str(rec.hours),
            desc,
        ]
        lines.append(",".join(row) + "\n")

    with open(abs_path, "w", newline="") as f:
        for line in lines:
            f.write(line)

    canaries.set_location(_KEY_TIME_RECORDS, rel_path, "Line 1 comment: # CANARY: ...")
    manifest.register(rel_path, "csv")


# ── Project Description DOCX files ───────────────────────────────────────────

def _write_project_descriptions(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write 14 project description .docx files."""
    projects = list(RD_PROJECTS_EU)
    desc_dir = output_dir / _PROJ_DESC_DIR
    desc_dir.mkdir(parents=True, exist_ok=True)

    for i, proj in enumerate(projects):
        file_key = _PROJECT_DOC_KEYS[i]
        canary = canaries.canary_for(file_key)

        doc = docx.Document()
        loc = embed_canary_docx(doc, canary)

        doc.add_heading(f"R&D Project Description: {proj.code}", level=1)
        doc.add_paragraph("")

        fields = [
            ("Project Code", proj.code),
            ("Entity", proj.entity_code),
            ("Project Name", proj.name),
            ("Status", proj.status),
        ]
        for label, value in fields:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(value)

        doc.add_heading("Objective", level=2)
        doc.add_paragraph(proj.objective)

        doc.add_heading("Technical Challenge", level=2)
        doc.add_paragraph(proj.technical_challenge)

        doc.add_heading("Methodology", level=2)
        doc.add_paragraph(proj.methodology)

        # Frascati criteria assessment
        doc.add_heading("Frascati Manual Criteria Assessment", level=2)
        criteria = [
            ("Novelty", proj.frascati_novelty),
            ("Creativity", proj.frascati_creativity),
            ("Uncertainty", proj.frascati_uncertainty),
            ("Systematicity", proj.frascati_systematicity),
            ("Transferability/Reproducibility", proj.frascati_transferability),
        ]
        for criterion, met in criteria:
            p = doc.add_paragraph()
            run = p.add_run(f"{criterion}: ")
            run.bold = True
            p.add_run("Yes" if met else "No")

        rel_path = f"{_PROJ_DESC_DIR}/{proj.code}.docx"
        abs_path = output_dir / rel_path
        _save_docx_deterministic(doc, abs_path)

        canaries.set_location(file_key, rel_path, f"{loc}")
        manifest.register(rel_path, "docx")


# ── Payroll XLSX ──────────────────────────────────────────────────────────────

def _write_payroll(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write payroll_data_eu_fy2025.xlsx for CM and CP R&D employees."""
    employees = generate_rd_employees_eu()
    canary = canaries.canary_for(_KEY_PAYROLL)

    wb = openpyxl.Workbook()
    loc = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Payroll Register FY2025"

    ws["A1"] = "Cascade Europe Group — R&D Payroll Register FY2025"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "Cascade Matériaux Avancés SAS (CM) / Cascade Präzisionsteile GmbH (CP)"
    ws["A2"].font = Font(bold=True, size=11, color="666666")
    ws.merge_cells("A1:G1")
    ws.merge_cells("A2:G2")

    headers = [
        "Entity Code",
        "Employee ID",
        "Employee Name",
        "Role",
        "Annual Gross Salary (EUR)",
        "Employer Social Charges (EUR)",
        "Total Employer Cost (EUR)",
    ]

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx, value=header)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    for row_idx, emp in enumerate(employees, start=5):
        ws.cell(row=row_idx, column=1, value=emp.entity_code)
        ws.cell(row=row_idx, column=2, value=emp.employee_id)
        ws.cell(row=row_idx, column=3, value=emp.name)
        ws.cell(row=row_idx, column=4, value=emp.role)

        for col_idx, val in enumerate(
            [emp.annual_gross_salary_eur, emp.employer_social_charges_eur,
             emp.total_employer_cost_eur],
            start=5,
        ):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.number_format = _NUM_FMT
            cell.border = _THIN_BORDER

    widths = [14, 14, 28, 24, 24, 26, 24]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    rel_path = f"{_INPUT_DIR}/payroll_data_eu_fy2025.xlsx"
    abs_path = output_dir / rel_path
    abs_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, abs_path)

    canaries.set_location(_KEY_PAYROLL, rel_path, loc)
    manifest.register(rel_path, "xlsx")


# ── Supply Expenses XLSX ──────────────────────────────────────────────────────

def _write_supply_expenses(
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write rd_supply_expenses_eu.xlsx with ERR-EU-008 planted."""
    expenses = generate_supply_expenses_eu()
    canary = canaries.canary_for(_KEY_SUPPLY)

    wb = openpyxl.Workbook()
    loc = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "R&D Supply Expenses"

    ws["A1"] = "Cascade Europe Group — R&D Supply & Materials Expenses FY2025"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:F1")

    headers = [
        "Entity Code",
        "Expense Date",
        "Description",
        "Amount (EUR)",
        "Cost Centre",
        "Project Code",
    ]

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=3, column=col_idx, value=header)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    err_row = None
    for row_idx, exp in enumerate(expenses, start=4):
        ws.cell(row=row_idx, column=1, value=exp.entity_code)
        ws.cell(row=row_idx, column=2, value=exp.expense_date)
        ws.cell(row=row_idx, column=3, value=exp.description)
        cell = ws.cell(row=row_idx, column=4, value=float(exp.amount_eur))
        cell.number_format = _NUM_FMT_2
        cell.border = _THIN_BORDER
        ws.cell(row=row_idx, column=5, value=exp.cost_center)
        ws.cell(row=row_idx, column=6, value=exp.project_code)

        # Track ERR-EU-008 row
        if (exp.entity_code == _ERR_EU_008_EXPENSE.entity_code
                and exp.description == _ERR_EU_008_EXPENSE.description
                and exp.amount_eur == _ERR_EU_008_EXPENSE.amount_eur):
            err_row = row_idx

    widths = [14, 14, 45, 16, 14, 14]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    rel_path = f"{_INPUT_DIR}/rd_supply_expenses_eu.xlsx"
    abs_path = output_dir / rel_path
    abs_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, abs_path)

    canaries.set_location(_KEY_SUPPLY, rel_path, loc)
    manifest.register(rel_path, "xlsx")

    # Register ERR-EU-008
    errors.add(PlantedError(
        error_id="ERR-EU-008",
        file=f"{_INPUT_DIR}/rd_supply_expenses_eu.xlsx",
        location=(
            f"Sheet 'R&D Supply Expenses', Row {err_row}, "
            "Entity CP, Project CP-RD-03"
        ),
        type="classification_error",
        description=(
            "Expense 'Calibration service — annual maintenance contract' (€12,400) "
            "at CP is coded to qualifying project CP-RD-03 but is a maintenance "
            "contract, not an R&D supply. For Forschungszulage, supplies don't "
            "qualify at all — the agent must recognise both the cost category error "
            "and the regime-specific ineligibility."
        ),
        severity="immaterial",
        which_test_cases_should_catch=[_TC],
    ))


# ── Subcontractor Invoices XLSX ───────────────────────────────────────────────

def _write_subcontractor_invoices(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write rd_subcontractor_invoices_eu.xlsx."""
    invoices = generate_subcontractor_invoices_eu()
    canary = canaries.canary_for(_KEY_SUBCONTRACTOR)

    wb = openpyxl.Workbook()
    loc = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Subcontracted R&D"

    ws["A1"] = "Cascade Matériaux Avancés SAS — Subcontracted R&D Invoices FY2025"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:F1")

    headers = [
        "Invoice ID",
        "Subcontractor Name",
        "Project Code",
        "Description",
        "Amount (EUR)",
        "Subcontractor Type",
    ]

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=3, column=col_idx, value=header)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    for row_idx, inv in enumerate(invoices, start=4):
        ws.cell(row=row_idx, column=1, value=inv.invoice_id)
        ws.cell(row=row_idx, column=2, value=inv.subcontractor_name)
        ws.cell(row=row_idx, column=3, value=inv.project_code)
        ws.cell(row=row_idx, column=4, value=inv.description)
        cell = ws.cell(row=row_idx, column=5, value=float(inv.amount_eur))
        cell.number_format = _NUM_FMT_2
        cell.border = _THIN_BORDER
        ws.cell(row=row_idx, column=6, value=inv.subcontractor_type)

    widths = [14, 55, 14, 55, 16, 20]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    rel_path = f"{_INPUT_DIR}/rd_subcontractor_invoices_eu.xlsx"
    abs_path = output_dir / rel_path
    abs_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, abs_path)

    canaries.set_location(_KEY_SUBCONTRACTOR, rel_path, loc)
    manifest.register(rel_path, "xlsx")


# ── Prior Year R&D Data XLSX ──────────────────────────────────────────────────

def _write_prior_year_data(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write prior_year_rd_data_eu.xlsx — contextual, NOT for averaging."""
    canary = canaries.canary_for(_KEY_PRIOR_YEAR)

    wb = openpyxl.Workbook()
    loc = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Prior Year R&D Data"

    ws["A1"] = "Cascade Europe Group — Prior Year R&D Expenditure"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:E1")
    ws["A2"] = "For contextual reference only — CIR and Forschungszulage use current year spend"
    ws["A2"].font = Font(italic=True, size=10, color="999999")
    ws.merge_cells("A2:E2")

    headers = [
        "Entity Code",
        "Entity Name",
        "Regime",
        "FY2023 Eligible R&D Spend (EUR)",
        "FY2024 Eligible R&D Spend (EUR)",
    ]

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx, value=header)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    # CM row
    ws.cell(row=5, column=1, value="CM")
    ws.cell(row=5, column=2, value="Cascade Matériaux Avancés SAS")
    ws.cell(row=5, column=3, value="France CIR")
    cell = ws.cell(row=5, column=4, value=int(PRIOR_YEAR_RD_SPEND["CM"][2023]))
    cell.number_format = _NUM_FMT
    cell = ws.cell(row=5, column=5, value=int(PRIOR_YEAR_RD_SPEND["CM"][2024]))
    cell.number_format = _NUM_FMT

    # CP row
    ws.cell(row=6, column=1, value="CP")
    ws.cell(row=6, column=2, value="Cascade Präzisionsteile GmbH")
    ws.cell(row=6, column=3, value="Germany Forschungszulage")
    cell = ws.cell(row=6, column=4, value=int(PRIOR_YEAR_RD_SPEND["CP"][2023]))
    cell.number_format = _NUM_FMT
    cell = ws.cell(row=6, column=5, value=int(PRIOR_YEAR_RD_SPEND["CP"][2024]))
    cell.number_format = _NUM_FMT

    widths = [14, 35, 24, 30, 30]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    rel_path = f"{_INPUT_DIR}/prior_year_rd_data_eu.xlsx"
    abs_path = output_dir / rel_path
    abs_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, abs_path)

    canaries.set_location(_KEY_PRIOR_YEAR, rel_path, loc)
    manifest.register(rel_path, "xlsx")


# ── Prompt ────────────────────────────────────────────────────────────────────

def _write_prompt(output_dir: Path) -> None:
    """Write test_cases/TC-08-EU/prompt.md."""
    text = """\
Perform an R&D incentive study for the Cascade Europe group for FY2025, covering
the French Crédit d'Impôt Recherche (CIR) for Cascade Matériaux Avancés SAS
and the German Forschungszulagengesetz (research allowance) for Cascade
Präzisionsteile GmbH.

For the French CIR (Cascade Matériaux Avancés SAS):

1. Review each CM project description and determine whether it meets the CIR
   eligibility criteria for qualifying R&D:
   - Activity must aim to advance scientific or technical knowledge
   - Must involve resolution of a technical or scientific uncertainty
   - Must be systematic (not routine adaptation, quality control, or market research)
2. For qualifying projects, compute eligible CIR expenses:
   - Researcher personnel costs: gross salary × forfaitaire rate of 43% for overhead
     (simplified method)
   - Consumables and supplies directly used in qualifying R&D
   - Subcontracted R&D: include at face value for private subcontractors;
     apply 2× multiplier for public research organisms
   - Exclude: market research, routine testing, general overhead not in forfaitaire
3. Compute the CIR credit: 30% × total eligible R&D expenditure
   (for spend ≤ €100M; rate drops to 5% above €100M)
4. Draft a summary memo for each qualifying project documenting the technical
   uncertainty and R&D activities performed.

For the German Forschungszulage (Cascade Präzisionsteile GmbH):

5. Review each CP project description and determine whether it qualifies under
   the Forschungszulagengesetz:
   - Must be fundamental research, industrial research, or experimental development
   - Must involve novelty, creativity, uncertainty, systematicity, and
     transferability/reproducibility (Frascati Manual criteria)
   - Routine quality testing and market research do not qualify
6. For qualifying projects, compute eligible personnel costs:
   - Only personnel costs of employees directly performing R&D qualify
   - Supplies, equipment, and subcontracted R&D do NOT qualify
   - Assessment basis capped at €2M per year
7. Compute the Forschungszulage: 25% × eligible personnel costs (max €500k)

For both entities:

8. Prepare a consolidated R&D incentive summary showing total benefit across
   both regimes.

Export:
- Project qualification analysis as Excel (project name, entity, regime,
  qualification determination, rationale)
- CIR computation as Excel (CM)
- Forschungszulage computation as Excel (CP)
- Consolidated incentive summary as Excel
- Documentation memos as a single Word document with one section per qualifying project
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Expected Behavior ─────────────────────────────────────────────────────────

def _write_expected_behavior(output_dir: Path) -> None:
    """Write test_cases/TC-08-EU/expected_behavior.md."""
    projects = list(RD_PROJECTS_EU)
    qual_lines: list[str] = []
    for p in projects:
        entity_label = "CM" if p.entity_code == "CM" else "CP"
        if p.qualifies == "yes":
            status = "Qualifies"
        elif p.qualifies == "borderline":
            status = "Borderline — flag for review"
        else:
            status = f"Does NOT qualify — {p.disqualification_reason}"
        qual_lines.append(f"- **{p.code} ({p.name})** [{entity_label}]: {status}")

    qual_table = "\n".join(qual_lines)

    result = compute_consolidated_rd_benefit()

    text = f"""\
# TC-08-EU: European R&D Incentive Study — Expected Behavior

## Project Qualification

The agent must evaluate each of the 14 R&D projects against the appropriate
regime criteria (CIR for CM, Forschungszulage for CP):

{qual_table}

### Key Traps

1. **CM-RD-10 (Market Analysis)**: Labelled as "R&D" but is pure market
   research. Does not advance scientific or technical knowledge. Must be
   excluded from CIR eligible base.

2. **CP-RD-04 (Quality R&D Programme)**: Uses established SPC methods —
   no novelty, creativity, or uncertainty. Fails Frascati criteria.

3. **CM-RD-08 and CM-RD-09 (Borderline)**: Should be flagged for review.
   CM-RD-08 is routine product adaptation; CM-RD-09 is compliance testing.

4. **No multi-year averaging**: Unlike the US ASC method, CIR and
   Forschungszulage both use current-year eligible spend. Prior year data
   is contextual only. If the agent applies a 3-year average formula,
   deduct points.

5. **Forschungszulage personnel-only rule**: CP eligible costs include
   ONLY personnel costs. If the agent includes CP supplies or subcontractor
   costs, deduct points.

6. **CIR forfaitaire**: Agent must apply 43% overhead on researcher
   salaries for CM. Omitting or using actual overhead is a deduction.

7. **Public research organism 2× multiplier**: The INSA Lyon subcontractor
   invoice (€80,000) should be counted as €160,000 in the CIR base.

## CIR Computation (CM)

- Qualifying researcher salaries: €{result.cir.qualifying_researcher_salary_eur:,}
- Forfaitaire overhead (43%): €{result.cir.forfaitaire_overhead_eur:,}
- Personnel total: €{result.cir.personnel_total_eur:,}
- Qualifying supplies: €{result.cir.qualifying_supplies_eur:,}
- Subcontracted R&D (private): €{result.cir.subcontracted_private_eur:,}
- Subcontracted R&D (public × 2): €{result.cir.subcontracted_public_doubled_eur:,}
- Total eligible base: €{result.cir.total_eligible_base_eur:,}
- **CIR credit (30%): €{result.cir.cir_credit_eur:,}**

## Forschungszulage Computation (CP)

- Qualifying personnel costs: €{result.forschungszulage.qualifying_personnel_cost_eur:,}
- Assessment basis (under €2M cap): €{result.forschungszulage.assessment_basis_eur:,}
- **Forschungszulage (25%): €{result.forschungszulage.benefit_eur:,}**

## Consolidated Group R&D Benefit

- CIR (CM): €{result.cir.cir_credit_eur:,}
- Forschungszulage (CP): €{result.forschungszulage.benefit_eur:,}
- **Total: €{result.total_benefit_eur:,}**

## Expected Deliverables

1. **Project Qualification Excel**: One row per project with columns for
   project code, entity, regime, Frascati/CIR criteria assessment,
   qualification determination, and rationale.
2. **CIR Computation Excel (CM)**: Personnel costs with forfaitaire,
   supplies, subcontracted R&D with public multiplier, total eligible base,
   and credit computation.
3. **Forschungszulage Computation Excel (CP)**: Personnel costs of qualifying
   employees, assessment basis, and benefit computation.
4. **Consolidated Summary Excel**: Group-level R&D benefit combining both
   regimes.
5. **Documentation Memos (Word)**: One section per qualifying project
   documenting the technical uncertainty and R&D activities.

## Data Challenges

- **42 employees × 52 weeks ≈ 2,184 time records**: Agent must aggregate
  hours by employee, project, and entity.
- **Dual-regime rules**: CIR allows supplies and subcontractors; Forschungszulage
  does not. Agent must apply different cost categories per entity.
- **ERR-EU-008**: Maintenance contract (€12,400) miscoded as R&D supply at CP.
  Agent should detect the cost category error.
- **European date format**: DD.MM.YYYY in time records and expense dates.
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold Standard ─────────────────────────────────────────────────────────────

@register_gold("TC-08-EU")
def _tc08_eu_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """TC-08-EU gold standard: European R&D incentive study."""
    result = compute_consolidated_rd_benefit()
    cir = result.cir
    fz = result.forschungszulage

    # Project qualification details
    project_details: dict[str, dict[str, Any]] = {}
    for p in RD_PROJECTS_EU:
        detail: dict[str, Any] = {
            "name": p.name,
            "entity_code": p.entity_code,
            "status": p.status,
            "qualifies": p.qualifies,
            "regime": "CIR" if p.entity_code == "CM" else "Forschungszulage",
        }
        if p.disqualification_reason:
            detail["disqualification_reason"] = p.disqualification_reason
        detail["frascati_criteria"] = {
            "novelty": p.frascati_novelty,
            "creativity": p.frascati_creativity,
            "uncertainty": p.frascati_uncertainty,
            "systematicity": p.frascati_systematicity,
            "transferability": p.frascati_transferability,
        }
        project_details[p.code] = detail

    expected_outputs: dict[str, Any] = {
        "projects_total": 14,
        "cm_projects": 10,
        "cp_projects": 4,
        "cm_qualifying": 7,
        "cm_borderline": 2,
        "cm_disqualified": 1,
        "cp_qualifying": 3,
        "cp_disqualified": 1,
        "cir": {
            "qualifying_researcher_salary_eur": str(cir.qualifying_researcher_salary_eur),
            "forfaitaire_overhead_eur": str(cir.forfaitaire_overhead_eur),
            "personnel_total_eur": str(cir.personnel_total_eur),
            "qualifying_supplies_eur": str(cir.qualifying_supplies_eur),
            "subcontracted_private_eur": str(cir.subcontracted_private_eur),
            "subcontracted_public_eur": str(cir.subcontracted_public_eur),
            "subcontracted_public_doubled_eur": str(cir.subcontracted_public_doubled_eur),
            "total_subcontracted_eur": str(cir.total_subcontracted_eur),
            "total_eligible_base_eur": str(cir.total_eligible_base_eur),
            "cir_credit_eur": str(cir.cir_credit_eur),
        },
        "forschungszulage": {
            "qualifying_personnel_cost_eur": str(fz.qualifying_personnel_cost_eur),
            "assessment_basis_eur": str(fz.assessment_basis_eur),
            "benefit_eur": str(fz.benefit_eur),
        },
        "consolidated_benefit_eur": str(result.total_benefit_eur),
        "project_details": project_details,
    }

    # Canary verification
    canary_verification: dict[str, str] = {
        "read_time_records": canaries.canary_for(_KEY_TIME_RECORDS),
        "read_payroll": canaries.canary_for(_KEY_PAYROLL),
        "read_supply_expenses": canaries.canary_for(_KEY_SUPPLY),
        "read_subcontractor_invoices": canaries.canary_for(_KEY_SUBCONTRACTOR),
        "read_prior_year_data": canaries.canary_for(_KEY_PRIOR_YEAR),
    }
    for i, key in enumerate(_PROJECT_DOC_KEYS):
        canary_verification[f"read_project_{RD_PROJECTS_EU[i].code}"] = (
            canaries.canary_for(key)
        )

    scoring_hints: dict[str, str] = {
        "correctness": (
            f"CIR credit: €{cir.cir_credit_eur:,}. "
            f"Forschungszulage: €{fz.benefit_eur:,}. "
            f"Consolidated: €{result.total_benefit_eur:,}. "
            "10 CM projects: 7 qualifying, 2 borderline, 1 disqualified. "
            "4 CP projects: 3 qualifying, 1 disqualified. "
            "Agent must NOT apply US ASC 3-year averaging formula."
        ),
        "completeness": (
            "5 deliverables: project qualification Excel, CIR computation "
            "Excel, Forschungszulage computation Excel, consolidated summary "
            "Excel, documentation memos Word."
        ),
        "dual_regime": (
            "CIR: personnel (forfaitaire 43%), supplies, subcontracted R&D "
            "(public 2×). Forschungszulage: personnel ONLY, no supplies, "
            "no subcontractors, €2M assessment cap. Mixing cost categories "
            "across regimes is a deduction."
        ),
        "terminology": (
            "Agent should use CIR, forfaitaire, Forschungszulage, Frascati "
            "Manual criteria — not US IRC Section 41 language."
        ),
        "communication": (
            "Qualification rationale documented per project using "
            "CIR criteria (CM) or Frascati criteria (CP). "
            "Borderline projects explicitly flagged for review."
        ),
    }

    return GoldStandard(
        test_case=_TC,
        expected_outputs=expected_outputs,
        canary_verification=canary_verification,
        error_detection={
            "ERR-EU-008": (
                "Maintenance contract (€12,400) miscoded as R&D supply at CP "
                "(CP-RD-03) — cost_category_misclassification in "
                "rd_supply_expenses_eu.xlsx"
            ),
        },
        scoring_hints=scoring_hints,
    )


# ── Public entry point ────────────────────────────────────────────────────────

def emit_tc08_eu(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
    **kwargs: object,
) -> None:
    """Write all TC-08-EU files to *output_dir*."""
    _write_time_records_csv(output_dir, canaries, manifest)
    _write_project_descriptions(output_dir, canaries, manifest)
    _write_payroll(output_dir, canaries, manifest)
    _write_supply_expenses(output_dir, canaries, errors, manifest)
    _write_subcontractor_invoices(output_dir, canaries, manifest)
    _write_prior_year_data(output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
