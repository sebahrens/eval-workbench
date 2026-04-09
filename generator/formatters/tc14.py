"""Formatter: TC-14 — 13-Week Cash Flow Forecast (Advisory, Routine).

Emits:
- test_cases/TC-14/input_files/balance_sheet_current.xlsx
  Balance sheet as of most recent week-end (Friday 2025-03-28).
  Includes cash balance, AR detail by aging bucket, AP detail by aging
  bucket, and current portion of debt.
- test_cases/TC-14/input_files/ap_aging_report.xlsx
  AP by vendor with due dates for the next 13 weeks.
- test_cases/TC-14/input_files/ar_aging_report.xlsx
  AR by customer with expected collection dates (DSO column).
- test_cases/TC-14/input_files/committed_discretionary_expenses.docx
  List of expenses: committed, semi-discretionary, discretionary with
  weekly/monthly amounts and payment timing.
- test_cases/TC-14/prompt.md
- test_cases/TC-14/expected_behavior.md
- gold_standards/TC-14_gold.json

Gold standard (from prompt.md):
  Cash trough in Week 8 at $1,340,000 (below $2M covenant).
  Defer marketing ($168,750/week) + training ($35,000/week) from Week 5 →
  trough $2,155,000.
  Sensitivity: collections slow 1 week → trough $780,000 even with
  deferrals.

Uses the canonical model for AR/AP aging and balance sheet data.
Planted errors:
  ERR-008 (mismatched_total) — AP aging total row omits one invoice ($1,800 delta).
  ERR-024 (rounding_discrepancy) — Balance-sheet Total Current Assets off by $1 vs. sum.
"""

from __future__ import annotations

import datetime
import io
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import openpyxl
from docx import Document
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from generator.canaries import CanaryRegistry, embed_canary_docx, embed_canary_xlsx
from generator.errors import (
    ErrorRegistry,
    PlantedError,
    mismatch_total,
    rounding_discrepancy,
)
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-14"
_INPUT_DIR = f"test_cases/{_TC}/input_files"

_FIXED_DATETIME = datetime.datetime(2025, 3, 28, 9, 0, 0)
_FIXED_ZIP_DT = (2025, 3, 28, 9, 0, 0)

# Week-end date for the "current" balance sheet (a Friday).
_BS_DATE = datetime.date(2025, 3, 28)

# 13 weeks starting from Week 1 = week ending 2025-04-04
_WEEK_START = datetime.date(2025, 3, 31)  # Monday of Week 1

# ── Expense schedule ────────────────────────────────────────────────────────
# Engineered so that the gold-standard numbers come out exactly right.
# Total weekly outflows: committed + semi-disc + discretionary.
#
# These are the weekly expense amounts that, combined with the AR inflows
# and AP outflows, produce exactly:
#   - Cash trough: Week 8 at $1,340,000
#   - Post-deferral trough: Week 8 at $2,155,000
#     (marketing $168,750 + training $35,000 = $203,750/week × 4 weeks = $815,000)
#   - Sensitivity trough: Week 8 at $780,000

# Committed expenses (cannot be deferred)
_COMMITTED_EXPENSES: list[tuple[str, str, int]] = [
    # (name, frequency, amount_dollars)
    ("Payroll", "Weekly", 580_000),
    ("Rent", "Monthly (Week 1 of each month)", 62_000),
    ("Debt Service — Interest", "Monthly (Week 2 of each month)", 45_000),
    ("Debt Service — Principal", "Monthly (Week 4 of each month)", 38_000),
    ("Insurance", "Monthly (Week 1 of each month)", 15_000),
    ("Utilities", "Monthly (Week 3 of each month)", 22_000),
]

# Semi-discretionary expenses (can be deferred 2–4 weeks)
_SEMI_DISC_EXPENSES: list[tuple[str, str, int]] = [
    ("Maintenance", "Bi-weekly (odd weeks)", 35_000),
    ("Professional Fees", "Monthly (Week 2 of each month)", 28_000),
]

# Discretionary expenses (can be cut)
_DISC_EXPENSES: list[tuple[str, str, int]] = [
    ("Marketing", "Weekly", 168_750),
    ("Travel", "Bi-weekly (even weeks)", 18_000),
    ("Training", "Weekly", 35_000),
    ("Bonuses", "Quarterly (Week 13 only)", 150_000),
]

# ── Weekly cash flow model ──────────────────────────────────────────────────
# These are the weekly AR collections and AP payments that, combined with
# the expense schedule above, produce the exact gold-standard values.
#
# Opening cash: $5,200,000 (derived from the model).
# The inflow/outflow patterns are designed to create the Week 8 trough.

_OPENING_CASH = 5_200_000

# Weekly AR collections — decline through Week 7, spike in Week 8
# (large quarterly client payment), then steady recovery.
_WEEKLY_AR_COLLECTIONS: list[int] = [
    900_000,    # Week 1
    840_000,    # Week 2
    780_000,    # Week 3
    722_000,    # Week 4
    660_000,    # Week 5
    600_000,    # Week 6
    570_000,    # Week 7
    1_375_000,  # Week 8 — quarterly client settlement
    1_400_000,  # Week 9
    1_550_000,  # Week 10
    1_650_000,  # Week 11
    1_700_000,  # Week 12
    1_800_000,  # Week 13
]

# Weekly AP payments (timing matches vendor due dates).
# Week 8 has a large vendor payment that coincides with the cash trough.
_WEEKLY_AP_PAYMENTS: list[int] = [
    380_000,   # Week 1
    350_000,   # Week 2
    420_000,   # Week 3
    310_000,   # Week 4
    390_000,   # Week 5
    440_000,   # Week 6
    360_000,   # Week 7
    755_000,   # Week 8 — large vendor payment
    300_000,   # Week 9
    370_000,   # Week 10
    400_000,   # Week 11
    320_000,   # Week 12
    280_000,   # Week 13
]


def _weekly_committed(week: int) -> int:
    """Compute committed expenses for a given week (1-indexed)."""
    total = _COMMITTED_EXPENSES[0][2]  # Payroll: every week
    # Rent: Week 1, 5, 9, 13 (first week of each 4-week "month")
    if week in (1, 5, 9, 13):
        total += _COMMITTED_EXPENSES[1][2]  # Rent
    # Debt service interest: Week 2, 6, 10
    if week in (2, 6, 10):
        total += _COMMITTED_EXPENSES[2][2]  # Interest
    # Debt service principal: Week 4, 8, 12
    if week in (4, 8, 12):
        total += _COMMITTED_EXPENSES[3][2]  # Principal
    # Insurance: Week 1, 5, 9, 13
    if week in (1, 5, 9, 13):
        total += _COMMITTED_EXPENSES[4][2]  # Insurance
    # Utilities: Week 3, 7, 11
    if week in (3, 7, 11):
        total += _COMMITTED_EXPENSES[5][2]  # Utilities
    return total


def _weekly_semi_disc(week: int) -> int:
    """Compute semi-discretionary expenses for a given week."""
    total = 0
    # Maintenance: odd weeks
    if week % 2 == 1:
        total += _SEMI_DISC_EXPENSES[0][2]
    # Professional fees: Week 2, 6, 10
    if week in (2, 6, 10):
        total += _SEMI_DISC_EXPENSES[1][2]
    return total


def _weekly_disc(week: int) -> int:
    """Compute discretionary expenses for a given week."""
    total = 0
    total += _DISC_EXPENSES[0][2]  # Marketing: every week
    # Travel: even weeks
    if week % 2 == 0:
        total += _DISC_EXPENSES[1][2]
    total += _DISC_EXPENSES[2][2]  # Training: every week
    # Bonuses: week 13 only
    if week == 13:
        total += _DISC_EXPENSES[3][2]
    return total


def _weekly_disc_deferred(week: int) -> int:
    """Compute discretionary expenses with marketing + training deferred from Week 5."""
    total = 0
    # Marketing: deferred starting Week 5
    if week < 5:
        total += _DISC_EXPENSES[0][2]
    # Travel: even weeks (not deferrable per spec)
    if week % 2 == 0:
        total += _DISC_EXPENSES[1][2]
    # Training: deferred starting Week 5
    if week < 5:
        total += _DISC_EXPENSES[2][2]
    # Bonuses: week 13 only
    if week == 13:
        total += _DISC_EXPENSES[3][2]
    return total


def _compute_weekly_cash_flows() -> list[dict[str, Any]]:
    """Compute the 13-week cash flow forecast (base scenario)."""
    balance = _OPENING_CASH
    weeks: list[dict[str, Any]] = []

    for w in range(1, 14):
        inflows = _WEEKLY_AR_COLLECTIONS[w - 1]
        ap = _WEEKLY_AP_PAYMENTS[w - 1]
        committed = _weekly_committed(w)
        semi_disc = _weekly_semi_disc(w)
        disc = _weekly_disc(w)
        total_outflows = ap + committed + semi_disc + disc
        net = inflows - total_outflows
        balance += net

        weeks.append({
            "week": w,
            "week_ending": str(_WEEK_START + datetime.timedelta(days=(w * 7 - 3))),
            "opening_balance": balance - net,
            "ar_collections": inflows,
            "ap_payments": ap,
            "committed_expenses": committed,
            "semi_disc_expenses": semi_disc,
            "disc_expenses": disc,
            "total_outflows": total_outflows,
            "net_cash_flow": net,
            "closing_balance": balance,
        })

    return weeks


def _compute_deferred_cash_flows() -> list[dict[str, Any]]:
    """Compute cash flows with marketing + training deferred from Week 5."""
    balance = _OPENING_CASH
    weeks: list[dict[str, Any]] = []

    for w in range(1, 14):
        inflows = _WEEKLY_AR_COLLECTIONS[w - 1]
        ap = _WEEKLY_AP_PAYMENTS[w - 1]
        committed = _weekly_committed(w)
        semi_disc = _weekly_semi_disc(w)
        disc = _weekly_disc_deferred(w)
        total_outflows = ap + committed + semi_disc + disc
        net = inflows - total_outflows
        balance += net

        weeks.append({
            "week": w,
            "closing_balance": balance,
            "net_cash_flow": net,
        })

    return weeks


def _compute_sensitivity_cash_flows() -> list[dict[str, Any]]:
    """Compute cash flows with 1-week collection delay + deferrals.

    Collections from Week N shift to Week N+1.  Week 1 gets zero
    collections (the prior week's collections haven't arrived yet).
    """
    balance = _OPENING_CASH
    weeks: list[dict[str, Any]] = []

    for w in range(1, 14):
        # Collections shifted by 1 week: Week 1 = 0, Week N = Week N-1's collections
        if w == 1:
            inflows = 0
        else:
            inflows = _WEEKLY_AR_COLLECTIONS[w - 2]

        ap = _WEEKLY_AP_PAYMENTS[w - 1]
        committed = _weekly_committed(w)
        semi_disc = _weekly_semi_disc(w)
        disc = _weekly_disc_deferred(w)  # deferrals still applied
        total_outflows = ap + committed + semi_disc + disc
        net = inflows - total_outflows
        balance += net

        weeks.append({
            "week": w,
            "closing_balance": balance,
            "net_cash_flow": net,
        })

    return weeks


# ── Styling helpers ─────────────────────────────────────────────────────────

_HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
_HEADER_FILL = PatternFill("solid", fgColor="4472C4")
_HEADER_ALIGN = Alignment(horizontal="center", wrap_text=True)
_THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)
_MONEY_FMT = '$#,##0'
_NUMBER_FMT = '#,##0'


def _style_header_row(ws: Any, row: int, max_col: int) -> None:
    """Apply header styling to a row."""
    for col in range(1, max_col + 1):
        cell = ws.cell(row=row, column=col)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = _HEADER_ALIGN
        cell.border = _THIN_BORDER


def _style_data_cell(cell: Any, fmt: str = _MONEY_FMT) -> None:
    """Apply data styling to a cell."""
    cell.number_format = fmt
    cell.border = _THIN_BORDER


# ── Deterministic save helpers ──────────────────────────────────────────────


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)
    wb.properties.modified = _FIXED_DATETIME

    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _pin_xlsx_dates(wb: openpyxl.Workbook) -> None:
    wb.properties.created = _FIXED_DATETIME
    wb.properties.modified = _FIXED_DATETIME


def _whole_dollars(d: Decimal) -> int:
    """Round a Decimal to whole dollars."""
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


def _fmt_dollars(v: int | Decimal) -> str:
    """Format as dollar string for gold standard."""
    if isinstance(v, Decimal):
        v = _whole_dollars(v)
    return f"${v:,}"


# ── AR aging report data ───────────────────────────────────────────────────
# Build customer-level AR with expected collection dates based on DSO.

def _build_ar_rows(model: CascadeModel) -> list[dict[str, Any]]:
    """Build AR aging rows with collection-week projection from customer DSO."""
    from generator.model.ar import CUSTOMERS, generate_ar_aging

    aging = generate_ar_aging(model.revenue_records, year=2025)
    rows: list[dict[str, Any]] = []

    for entry in aging:
        # Find customer DSO
        cust = next(c for c in CUSTOMERS if c.id == entry.customer_id)
        # Expected collection week = DSO / 7 (roughly)
        collection_weeks = max(1, cust.dso // 7)

        rows.append({
            "customer_id": entry.customer_id,
            "customer_name": entry.customer_name,
            "entity": entry.entity_code,
            "current": _whole_dollars(entry.current),
            "days_30": _whole_dollars(entry.days_30),
            "days_60": _whole_dollars(entry.days_60),
            "days_90": _whole_dollars(entry.days_90),
            "days_120_plus": _whole_dollars(entry.days_120_plus),
            "total": _whole_dollars(entry.total),
            "dso": cust.dso,
            "expected_collection_weeks": collection_weeks,
        })

    return rows


# ── AP aging report data ───────────────────────────────────────────────────

def _build_ap_rows(model: CascadeModel) -> list[dict[str, Any]]:
    """Build AP aging rows with due-date week projection."""
    from generator.model.ap import VENDORS, generate_ap_aging

    aging = generate_ap_aging(model.revenue_records, year=2025)
    rows: list[dict[str, Any]] = []

    for entry in aging:
        vendor = next(v for v in VENDORS if v.id == entry.vendor_id)
        # Due in N weeks based on payment terms
        due_weeks = max(1, vendor.payment_terms // 7)

        rows.append({
            "vendor_id": entry.vendor_id,
            "vendor_name": entry.vendor_name,
            "entity": entry.entity_code,
            "current": _whole_dollars(entry.current),
            "days_30": _whole_dollars(entry.days_30),
            "days_60": _whole_dollars(entry.days_60),
            "days_90_plus": _whole_dollars(entry.days_90_plus),
            "total": _whole_dollars(entry.total),
            "payment_terms": vendor.payment_terms,
            "due_week": due_weeks,
        })

    return rows


# ── File writers ────────────────────────────────────────────────────────────


def _write_balance_sheet(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write balance_sheet_current.xlsx — the starting point for cash flow."""
    from generator.model.views import build_balance_sheet

    bs = build_balance_sheet(model.ledger, _BS_DATE)
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    ws = wb.active
    ws.title = "Balance Sheet"

    # Header
    ws["A1"] = "Cascade Industries — Consolidated Balance Sheet"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = f"As of {_BS_DATE.strftime('%B %d, %Y')}"
    ws["A2"].font = Font(italic=True, size=11)

    # Assets section
    row = 4
    ws.cell(row=row, column=1, value="ASSETS").font = Font(bold=True, size=12)
    row += 1

    # Cash — use the engineered opening cash for the 13-week model
    ws.cell(row=row, column=1, value="Cash and Cash Equivalents")
    ws.cell(row=row, column=2, value=_OPENING_CASH)
    _style_data_cell(ws.cell(row=row, column=2))
    row += 1

    # AR — from the model
    ar_rows = _build_ar_rows(model)
    total_ar = sum(r["total"] for r in ar_rows)
    ws.cell(row=row, column=1, value="Accounts Receivable, net")
    ws.cell(row=row, column=2, value=total_ar)
    _style_data_cell(ws.cell(row=row, column=2))
    row += 1

    # Other current assets from model BS
    for acct_num in sorted(bs.assets.keys()):
        if acct_num in ("1010", "1100", "1150"):
            continue  # Already handled cash and AR
        acct_val = _whole_dollars(bs.assets[acct_num])
        if acct_val == 0:
            continue
        from generator.model.coa import ACCOUNTS_BY_NUMBER
        acct_info = ACCOUNTS_BY_NUMBER.get(acct_num)
        label = acct_info.name if acct_info else f"Account {acct_num}"
        ws.cell(row=row, column=1, value=label)
        ws.cell(row=row, column=2, value=acct_val)
        _style_data_cell(ws.cell(row=row, column=2))
        row += 1

    other_current = sum(
        _whole_dollars(v) for k, v in bs.assets.items()
        if k not in ("1010", "1100", "1150")
    )
    correct_total_current = _OPENING_CASH + total_ar + other_current

    # ERR-024: rounding discrepancy on Total Current Assets
    wrong_total_current = int(rounding_discrepancy(
        float(correct_total_current) + 0.6,  # force a $1 rounding difference
        decimal_places=0,
        direction="up",
    ))

    row += 1
    total_current_row = row
    ws.cell(row=row, column=1, value="Total Current Assets").font = Font(bold=True)
    ws.cell(row=row, column=2, value=wrong_total_current)
    _style_data_cell(ws.cell(row=row, column=2))

    errors.add(PlantedError(
        error_id="ERR-024",
        file=f"{_INPUT_DIR}/balance_sheet_current.xlsx",
        location=f"Sheet 'Balance Sheet', Row {total_current_row}, Column B (Total Current Assets)",
        type="rounding_discrepancy",
        description=(
            f"Total Current Assets shows ${wrong_total_current:,} "
            f"instead of ${correct_total_current:,}"
        ),
        severity="immaterial",
        which_test_cases_should_catch=["TC-14"],
    ))

    total_assets = correct_total_current
    row += 1
    ws.cell(row=row, column=1, value="TOTAL ASSETS").font = Font(bold=True)
    ws.cell(row=row, column=2, value=total_assets)
    _style_data_cell(ws.cell(row=row, column=2))

    # Liabilities section
    row += 2
    ws.cell(row=row, column=1, value="LIABILITIES").font = Font(bold=True, size=12)
    row += 1

    # AP from model
    ap_rows = _build_ap_rows(model)
    total_ap = sum(r["total"] for r in ap_rows)
    ws.cell(row=row, column=1, value="Accounts Payable")
    ws.cell(row=row, column=2, value=total_ap)
    _style_data_cell(ws.cell(row=row, column=2))
    row += 1

    # Current portion of debt
    debt_current = 2_000_000  # Fixed for model consistency
    ws.cell(row=row, column=1, value="Current Portion of Long-Term Debt")
    ws.cell(row=row, column=2, value=debt_current)
    _style_data_cell(ws.cell(row=row, column=2))
    row += 1

    for acct_num in sorted(bs.liabilities.keys()):
        if acct_num == "2010":
            continue  # AP already shown
        acct_val = _whole_dollars(bs.liabilities[acct_num])
        if acct_val == 0:
            continue
        from generator.model.coa import ACCOUNTS_BY_NUMBER
        acct_info = ACCOUNTS_BY_NUMBER.get(acct_num)
        label = acct_info.name if acct_info else f"Account {acct_num}"
        ws.cell(row=row, column=1, value=label)
        ws.cell(row=row, column=2, value=acct_val)
        _style_data_cell(ws.cell(row=row, column=2))
        row += 1

    # Minimum liquidity covenant note
    row += 1
    ws.cell(row=row, column=1, value="Note: Minimum liquidity covenant = $2,000,000")
    ws.cell(row=row, column=1).font = Font(italic=True, color="FF0000")

    ws.column_dimensions["A"].width = 40
    ws.column_dimensions["B"].width = 18

    # Canary
    canary = canaries.canary_for("tc14_balance_sheet_current")
    loc = embed_canary_xlsx(wb, canary)
    rel_path = f"{_INPUT_DIR}/balance_sheet_current.xlsx"
    canaries.set_location("tc14_balance_sheet_current", rel_path, loc)
    manifest.register(rel_path, "xlsx", canary=canary)

    abs_path = output_dir / rel_path
    abs_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, abs_path)


def _write_ar_aging(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write ar_aging_report.xlsx — AR by customer with DSO and collection weeks."""
    rows = _build_ar_rows(model)

    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    ws = wb.active
    ws.title = "AR Aging"

    headers = [
        "Customer ID", "Customer Name", "Entity", "Current (0-30)",
        "31-60 Days", "61-90 Days", "91-120 Days", "120+ Days",
        "Total AR", "DSO (days)", "Expected Collection (weeks)",
    ]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    for i, r in enumerate(rows, 2):
        ws.cell(row=i, column=1, value=r["customer_id"])
        ws.cell(row=i, column=2, value=r["customer_name"])
        ws.cell(row=i, column=3, value=r["entity"])
        ws.cell(row=i, column=4, value=r["current"])
        _style_data_cell(ws.cell(row=i, column=4))
        ws.cell(row=i, column=5, value=r["days_30"])
        _style_data_cell(ws.cell(row=i, column=5))
        ws.cell(row=i, column=6, value=r["days_60"])
        _style_data_cell(ws.cell(row=i, column=6))
        ws.cell(row=i, column=7, value=r["days_90"])
        _style_data_cell(ws.cell(row=i, column=7))
        ws.cell(row=i, column=8, value=r["days_120_plus"])
        _style_data_cell(ws.cell(row=i, column=8))
        ws.cell(row=i, column=9, value=r["total"])
        _style_data_cell(ws.cell(row=i, column=9))
        ws.cell(row=i, column=10, value=r["dso"])
        ws.cell(row=i, column=11, value=r["expected_collection_weeks"])

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 16
    ws.column_dimensions["B"].width = 28

    # Canary
    canary = canaries.canary_for("tc14_ar_aging_report")
    loc = embed_canary_xlsx(wb, canary)
    rel_path = f"{_INPUT_DIR}/ar_aging_report.xlsx"
    canaries.set_location("tc14_ar_aging_report", rel_path, loc)
    manifest.register(rel_path, "xlsx", canary=canary)

    abs_path = output_dir / rel_path
    abs_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, abs_path)


def _write_ap_aging(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write ap_aging_report.xlsx — AP by vendor with due-date weeks."""
    rows = _build_ap_rows(model)

    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    ws = wb.active
    ws.title = "AP Aging"

    headers = [
        "Vendor ID", "Vendor Name", "Entity", "Current (0-30)",
        "31-60 Days", "61-90 Days", "90+ Days",
        "Total AP", "Payment Terms", "Due (weeks)",
    ]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    for i, r in enumerate(rows, 2):
        ws.cell(row=i, column=1, value=r["vendor_id"])
        ws.cell(row=i, column=2, value=r["vendor_name"])
        ws.cell(row=i, column=3, value=r["entity"])
        ws.cell(row=i, column=4, value=r["current"])
        _style_data_cell(ws.cell(row=i, column=4))
        ws.cell(row=i, column=5, value=r["days_30"])
        _style_data_cell(ws.cell(row=i, column=5))
        ws.cell(row=i, column=6, value=r["days_60"])
        _style_data_cell(ws.cell(row=i, column=6))
        ws.cell(row=i, column=7, value=r["days_90_plus"])
        _style_data_cell(ws.cell(row=i, column=7))
        ws.cell(row=i, column=8, value=r["total"])
        _style_data_cell(ws.cell(row=i, column=8))
        ws.cell(row=i, column=9, value=f"Net {r['payment_terms']}")
        ws.cell(row=i, column=10, value=r["due_week"])

    # ERR-008: AP aging total row — omit one invoice from the total ($1,800 delta)
    totals_row = len(rows) + 2  # header row + data rows + 1
    correct_ap_total = sum(r["total"] for r in rows)
    _ERR008_DELTA = -1_800  # as if one small invoice was omitted
    wrong_ap_total = int(mismatch_total(correct_ap_total, _ERR008_DELTA))

    ws.cell(row=totals_row, column=1, value="TOTAL")
    ws.cell(row=totals_row, column=1).font = Font(bold=True)
    ws.cell(row=totals_row, column=8, value=wrong_ap_total)
    ws.cell(row=totals_row, column=8).font = Font(bold=True)
    _style_data_cell(ws.cell(row=totals_row, column=8))

    errors.add(PlantedError(
        error_id="ERR-008",
        file=f"{_INPUT_DIR}/ap_aging_report.xlsx",
        location=f"Sheet 'AP Aging', Row {totals_row}, Column H (Total AP)",
        type="mismatched_total",
        description=(
            f"AP aging total shows ${wrong_ap_total:,} "
            f"instead of ${correct_ap_total:,}"
        ),
        severity="material",
        which_test_cases_should_catch=["TC-14"],
    ))

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 16
    ws.column_dimensions["B"].width = 28

    # Canary
    canary = canaries.canary_for("tc14_ap_aging_report")
    loc = embed_canary_xlsx(wb, canary)
    rel_path = f"{_INPUT_DIR}/ap_aging_report.xlsx"
    canaries.set_location("tc14_ap_aging_report", rel_path, loc)
    manifest.register(rel_path, "xlsx", canary=canary)

    abs_path = output_dir / rel_path
    abs_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, abs_path)


def _write_expenses_docx(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write committed_discretionary_expenses.docx."""
    doc = Document()

    # Title
    doc.add_heading("Cascade Industries — Committed & Discretionary Expenses", level=1)
    doc.add_paragraph(
        "Weekly/monthly expense schedule for 13-week cash flow projection. "
        "Amounts are consolidated across all entities."
    )

    # Committed
    doc.add_heading("Committed Expenses (Cannot Be Deferred)", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Expense"
    hdr[1].text = "Frequency"
    hdr[2].text = "Amount"
    for name, freq, amt in _COMMITTED_EXPENSES:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = freq
        row[2].text = f"${amt:,}/week" if "Week" not in freq else f"${amt:,}/occurrence"

    doc.add_paragraph("")

    # Semi-discretionary
    doc.add_heading("Semi-Discretionary Expenses (Can Be Deferred 2-4 Weeks)", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Expense"
    hdr[1].text = "Frequency"
    hdr[2].text = "Amount"
    for name, freq, amt in _SEMI_DISC_EXPENSES:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = freq
        row[2].text = f"${amt:,}/occurrence"

    doc.add_paragraph("")

    # Discretionary
    doc.add_heading("Discretionary Expenses (Can Be Cut)", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Expense"
    hdr[1].text = "Frequency"
    hdr[2].text = "Amount"
    for name, freq, amt in _DISC_EXPENSES:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = freq
        row[2].text = f"${amt:,}/week" if freq == "Weekly" else f"${amt:,}/occurrence"

    doc.add_paragraph("")

    # Covenant note
    doc.add_heading("Liquidity Covenant", level=2)
    doc.add_paragraph(
        "Cascade Industries is subject to a minimum liquidity covenant of $2,000,000. "
        "Cash and cash equivalents must not fall below this threshold at any point."
    )

    # Canary
    canary = canaries.canary_for("tc14_committed_expenses")
    loc = embed_canary_docx(doc, canary)
    rel_path = f"{_INPUT_DIR}/committed_discretionary_expenses.docx"
    canaries.set_location("tc14_committed_expenses", rel_path, loc)
    manifest.register(rel_path, "docx", canary=canary)

    abs_path = output_dir / rel_path
    abs_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, abs_path)


def _write_prompt(output_dir: Path) -> None:
    """Write test_cases/TC-14/prompt.md."""
    text = """\
Build a 13-week cash flow forecast for Cascade Industries.

1. Start with the current cash balance.
2. Project weekly cash inflows based on the AR aging and historical collection patterns.
3. Project weekly cash outflows based on AP due dates and committed expense schedule.
4. Identify the projected cash trough (lowest cash balance and which week).
5. Determine if and when the company would breach its minimum liquidity
   covenant of $2,000,000.
6. If a breach is projected, identify which discretionary expenses could be
   deferred to avoid it, and show the revised forecast.

Export as an Excel workbook with:
- Weekly cash flow detail (inflows, outflows by category, net, cumulative balance)
- Summary dashboard showing the 13-week trend
- Sensitivity analysis: what happens if collections slow by 1 week?
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write test_cases/TC-14/expected_behavior.md."""
    text = """\
# TC-14: 13-Week Cash Flow Forecast — Expected Behavior

## Cash Flow Projection
- The agent should construct a week-by-week cash flow forecast starting from
  the current cash balance of $5,200,000.
- Weekly inflows should be derived from the AR aging report, projecting
  collections based on each customer's DSO.
- Weekly outflows should include AP payments (based on vendor due dates),
  committed expenses (payroll, rent, debt service, insurance, utilities),
  semi-discretionary expenses (maintenance, professional fees), and
  discretionary expenses (marketing, travel, training, bonuses).

## Covenant Breach Identification
- The agent MUST identify that the cash balance breaches the $2,000,000
  minimum liquidity covenant.
- The trough occurs in Week 8 at $1,340,000.
- The agent should flag this as a material covenant violation.

## Deferral Analysis
- The agent should identify that deferring marketing ($168,750/week) and
  training ($35,000/week) starting from Week 5 avoids the covenant breach.
- With these deferrals, the Week 8 trough improves to $2,155,000 —
  above the $2,000,000 covenant.
- The agent should NOT recommend deferring committed expenses (payroll,
  rent, debt service) as these are non-negotiable.

## Sensitivity Analysis
- The agent should model the impact of a 1-week delay in collections.
- Even with the marketing and training deferrals, a 1-week collection
  slowdown causes the trough to drop to $780,000 — well below the covenant.
- This demonstrates that the company's liquidity position is fragile and
  depends heavily on timely customer payments.

## Output Format
- Excel workbook with at least three sheets/sections:
  1. Weekly cash flow detail (inflows, outflows by category, net, cumulative)
  2. Summary dashboard / trend chart data
  3. Sensitivity analysis with revised projections
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold standard ──────────────────────────────────────────────────────────


@register_gold("TC-14")
def _tc14_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """Build the TC-14 gold standard from the engineered cash flow model."""
    base_flows = _compute_weekly_cash_flows()
    deferred_flows = _compute_deferred_cash_flows()
    sensitivity_flows = _compute_sensitivity_cash_flows()

    # Find trough in base scenario
    base_trough_week = min(base_flows, key=lambda w: w["closing_balance"])
    # Find trough in deferred scenario
    deferred_trough_week = min(deferred_flows, key=lambda w: w["closing_balance"])
    # Find trough in sensitivity scenario
    sensitivity_trough_week = min(sensitivity_flows, key=lambda w: w["closing_balance"])

    # Build weekly detail for gold standard
    weekly_detail: list[dict[str, Any]] = []
    for w in base_flows:
        weekly_detail.append({
            "week": w["week"],
            "week_ending": w["week_ending"],
            "ar_collections": _fmt_dollars(w["ar_collections"]),
            "ap_payments": _fmt_dollars(w["ap_payments"]),
            "committed_expenses": _fmt_dollars(w["committed_expenses"]),
            "semi_disc_expenses": _fmt_dollars(w["semi_disc_expenses"]),
            "disc_expenses": _fmt_dollars(w["disc_expenses"]),
            "total_outflows": _fmt_dollars(w["total_outflows"]),
            "net_cash_flow": _fmt_dollars(w["net_cash_flow"]),
            "closing_balance": _fmt_dollars(w["closing_balance"]),
        })

    return GoldStandard(
        test_case="TC-14",
        expected_outputs={
            "file_type": "xlsx",
            "opening_cash_balance": _fmt_dollars(_OPENING_CASH),
            "covenant_threshold": _fmt_dollars(2_000_000),
            "base_scenario": {
                "trough_week": base_trough_week["week"],
                "trough_balance": _fmt_dollars(base_trough_week["closing_balance"]),
                "covenant_breached": base_trough_week["closing_balance"] < 2_000_000,
                "weekly_detail": weekly_detail,
            },
            "deferral_scenario": {
                "deferred_expenses": ["Marketing ($168,750/week)", "Training ($35,000/week)"],
                "deferral_start_week": 5,
                "trough_week": deferred_trough_week["week"],
                "trough_balance": _fmt_dollars(deferred_trough_week["closing_balance"]),
                "covenant_breached": deferred_trough_week["closing_balance"] < 2_000_000,
            },
            "sensitivity_scenario": {
                "description": "Collections delayed by 1 week, with deferrals applied",
                "trough_week": sensitivity_trough_week["week"],
                "trough_balance": _fmt_dollars(sensitivity_trough_week["closing_balance"]),
                "covenant_breached": sensitivity_trough_week["closing_balance"] < 2_000_000,
            },
        },
        canary_verification={
            "read_balance_sheet": canaries.canary_for("tc14_balance_sheet_current"),
            "read_ar_aging": canaries.canary_for("tc14_ar_aging_report"),
            "read_ap_aging": canaries.canary_for("tc14_ap_aging_report"),
            "read_expense_schedule": canaries.canary_for("tc14_committed_expenses"),
        },
        error_detection={
            "ERR-008": "AP aging total row mismatched — one invoice omitted from the sum",
            "ERR-024": "Balance sheet Total Current Assets shows rounding discrepancy vs. component sum",
        },
        scoring_hints={
            "correctness": (
                "Cash trough at Week 8 = $1,340,000; deferral trough = $2,155,000 "
                "(marketing $168,750 + training $35,000 deferred from Week 5); "
                "sensitivity trough = $780,000; covenant = $2M"
            ),
            "completeness": (
                "All 13 weeks projected; all expense categories included; "
                "deferral analysis with specific expenses identified; "
                "sensitivity analysis for 1-week collection delay"
            ),
            "format_compliance": (
                "Valid xlsx; weekly detail sheet; summary/dashboard; "
                "sensitivity analysis section"
            ),
            "robustness": (
                "Correctly identifies covenant breach; recommends actionable deferrals; "
                "sensitivity shows fragility even with deferrals"
            ),
            "communication": (
                "Clear explanation of covenant risk; specific deferral recommendations; "
                "sensitivity analysis demonstrates liquidity fragility"
            ),
        },
    )


# ── Public entry point ──────────────────────────────────────────────────────


def emit_tc14(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write all TC-14 files to *output_dir*."""
    _write_balance_sheet(model, output_dir, canaries, errors, manifest)
    _write_ar_aging(model, output_dir, canaries, manifest)
    _write_ap_aging(model, output_dir, canaries, errors, manifest)
    _write_expenses_docx(output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
