"""Formatter: TC-16-EU — European Engagement Letter Generation.

Emits:
- test_cases/TC-16-EU/input_files/client_profile_eu.docx
  Cascade Europe Holdings B.V. profile with entity count, revenue tier,
  IFRS first-time adoption, jurisdictions, key contacts.
- test_cases/TC-16-EU/input_files/fee_schedule_eu.xlsx
  EUR fee matrix: 4 service lines (Statutory Audit, CIT, VAT, TP),
  European revenue tiers, complexity multipliers.
- test_cases/TC-16-EU/input_files/engagement_letter_template_eu.docx
  European variant template with ISA/IFRS/OECD language, NL law limitation.
- test_cases/TC-16-EU/prompt.md
- test_cases/TC-16-EU/expected_behavior.md
- gold_standards/TC-16-EU_gold.json

Gold standard (from design bead synth-data-eu.18):
  Statutory Audit: base EUR 210,000 x 1.20 IFRS adoption = EUR 252,000.
  CIT Compliance: EUR 72,000 + EUR 15,000 x 3 = EUR 117,000.
  VAT Compliance: EUR 35,000 + EUR 8,000 x 3 = EUR 59,000.
  Transfer Pricing: EUR 48,000 + EUR 12,000 x 3 = EUR 84,000.
  Total: EUR 512,000.

One planted error (ERR-EU-016): formula_error in fee_schedule_eu.xlsx
  — IFRS adoption multiplier omission in summary (shows EUR 210,000
    instead of EUR 252,000 for statutory audit total).
"""

from __future__ import annotations

import datetime
import io
from decimal import Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import openpyxl
from docx import Document
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from generator.canaries import CanaryRegistry, embed_canary_docx, embed_canary_xlsx
from generator.errors import ErrorRegistry, PlantedError, formula_error
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel

# -- Constants ----------------------------------------------------------------

_TC = "TC-16-EU"
_INPUT_DIR = f"test_cases/{_TC}/input_files"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)
_FIXED_ZIP_DT = (2025, 3, 15, 9, 0, 0)

# Fee schedule values from design bead synth-data-eu.18
_AUDIT_BASE_FEE = Decimal("210000")
_IFRS_COMPLEXITY_MULTIPLIER = Decimal("1.20")
_CIT_BASE_FEE = Decimal("72000")
_CIT_PER_ENTITY_ADDER = Decimal("15000")
_VAT_BASE_FEE = Decimal("35000")
_VAT_PER_ENTITY_ADDER = Decimal("8000")
_TP_BASE_FEE = Decimal("48000")
_TP_PER_ENTITY_ADDER = Decimal("12000")
_ADDITIONAL_ENTITIES = 3  # 3 subsidiaries beyond the parent

# Computed gold values
_AUDIT_FEE = _AUDIT_BASE_FEE * _IFRS_COMPLEXITY_MULTIPLIER  # EUR 252,000
_CIT_FEE = _CIT_BASE_FEE + _CIT_PER_ENTITY_ADDER * _ADDITIONAL_ENTITIES  # EUR 117,000
_VAT_FEE = _VAT_BASE_FEE + _VAT_PER_ENTITY_ADDER * _ADDITIONAL_ENTITIES  # EUR 59,000
_TP_FEE = _TP_BASE_FEE + _TP_PER_ENTITY_ADDER * _ADDITIONAL_ENTITIES  # EUR 84,000
_TOTAL_FEE = _AUDIT_FEE + _CIT_FEE + _VAT_FEE + _TP_FEE  # EUR 512,000

# Merge field values from design
_CLIENT_NAME = "Cascade Europe Holdings B.V."
_PARTNER_NAME = "Pieter de Jong"
_START_DATE = "15 March 2026"
_PAYMENT_TERMS = "Net 30"

# -- Deterministic file helpers -----------------------------------------------


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)

    buf = io.BytesIO()
    wb.properties.created = _FIXED_DATETIME
    wb.properties.modified = _FIXED_DATETIME
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


# -- Styling helpers ----------------------------------------------------------

_HEADER_FILL = PatternFill(start_color="1A3C6E", end_color="1A3C6E", fill_type="solid")
_HEADER_FONT = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
_DATA_FONT = Font(name="Calibri", size=11)
_BOLD_FONT = Font(name="Calibri", size=11, bold=True)
_CURRENCY_FMT = '#,##0'
_MULTIPLIER_FMT = '0.00"x"'
_THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)


# -- File generators ----------------------------------------------------------


def _write_client_profile(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Create test_cases/TC-16-EU/input_files/client_profile_eu.docx."""
    doc = Document()

    # Canary
    canary_code = canaries.canary_for("tc16eu_client_profile")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "tc16eu_client_profile",
        f"{_INPUT_DIR}/client_profile_eu.docx",
        location,
    )

    # Title
    doc.add_heading("Client Profile — Cascade Europe Holdings B.V.", level=1)

    # Company overview
    doc.add_heading("Company Overview", level=2)
    doc.add_paragraph(f"Company Name: {_CLIENT_NAME}")
    doc.add_paragraph("Entity Type: Dutch B.V. (Besloten Vennootschap) with 3 subsidiaries (GmbH, SAS, Ltd)")
    doc.add_paragraph(
        "Industry: Mid-market manufacturer (precision components, advanced materials, distribution)"
    )
    doc.add_paragraph("Registered Office: Amsterdam, Netherlands")
    doc.add_paragraph("Fiscal Year End: 31 December")

    # Revenue and size
    doc.add_heading("Revenue and Size", level=2)
    doc.add_paragraph("Consolidated Revenue (FY2025): ~\u20ac120,000,000")
    doc.add_paragraph("Revenue Tier: \u20ac75M\u2013\u20ac250M")
    doc.add_paragraph("Total Employees: 620")

    # Entity structure
    doc.add_heading("Entity Structure", level=2)
    doc.add_paragraph(
        "Cascade Europe Holdings B.V. operates through a parent company "
        "and three wholly-owned subsidiaries:"
    )

    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["Entity", "Location", "Type", "Revenue"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    entities = [
        (
            "Cascade Europe Holdings B.V. (Parent)",
            "Amsterdam, NL",
            "Parent \u2014 B.V.",
            "Consolidated: \u20ac120M",
        ),
        (
            "Cascade Pr\u00e4zisionsteile GmbH",
            "Munich, DE",
            "Core manufacturing",
            "\u20ac55M",
        ),
        (
            "Cascade Mat\u00e9riaux Avanc\u00e9s SAS",
            "Lyon, FR",
            "Specialty materials R&D",
            "\u20ac38M",
        ),
        (
            "Cascade Distribution Services Ltd",
            "Birmingham, UK",
            "Warehousing & logistics",
            "\u20ac27M",
        ),
    ]
    for row_idx, (name, loc, etype, rev) in enumerate(entities, start=1):
        table.rows[row_idx].cells[0].text = name
        table.rows[row_idx].cells[1].text = loc
        table.rows[row_idx].cells[2].text = etype
        table.rows[row_idx].cells[3].text = rev

    doc.add_paragraph("")  # spacer
    doc.add_paragraph(f"Total Entities: 4 (1 parent + {_ADDITIONAL_ENTITIES} subsidiaries)")

    # Complexity factors
    doc.add_heading("Complexity Factors", level=2)
    doc.add_paragraph(
        "The group is currently transitioning from Dutch GAAP to full IFRS "
        "for consolidated reporting. This is a first-time IFRS adoption, "
        "introducing additional complexity for the statutory audit engagement "
        "requiring IFRS 1 transition procedures and enhanced documentation."
    )
    doc.add_paragraph("IFRS First-Time Adoption: Yes \u2014 transitioning from Dutch GAAP to full IFRS")
    doc.add_paragraph(
        "Multi-Jurisdiction Operations: Yes \u2014 Netherlands, Germany, France, United Kingdom"
    )
    doc.add_paragraph(
        "Intercompany Transactions: Significant \u2014 cost-plus transfers, "
        "management fees, intercompany loans"
    )
    doc.add_paragraph("R&D Activity: Substantial \u2014 ~14% of Cascade Mat\u00e9riaux Avanc\u00e9s SAS revenue")

    # Key contacts
    doc.add_heading("Key Contacts", level=2)
    contacts = [
        ("Managing Director", "Hans van der Berg", "h.vanderberg@cascadeeurope.nl"),
        ("Group Finance Director", "Isabelle Moreau", "i.moreau@cascadeeurope.nl"),
        ("Group Controller", "Klaus Fischer", "k.fischer@cascadeeurope.nl"),
        ("External Counsel", "De Brauw Blackstone Westbroek (NL)", "engagement@debrauw.com"),
    ]
    for title, name, email in contacts:
        doc.add_paragraph(f"{title}: {name} ({email})")

    # Engagement history
    doc.add_heading("Engagement History", level=2)
    doc.add_paragraph(
        "Mitchell & Partners International LLP has served as Cascade Europe\u2019s "
        "external audit firm for the past 2 years. This engagement letter covers "
        "a combined statutory audit (IFRS), corporate income tax compliance, "
        "VAT compliance, and transfer pricing documentation engagement for FY2026."
    )

    # Save
    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME
    path = output_dir / _INPUT_DIR / "client_profile_eu.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, path)

    manifest.register(
        f"{_INPUT_DIR}/client_profile_eu.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


def _write_fee_schedule(
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Create test_cases/TC-16-EU/input_files/fee_schedule_eu.xlsx.

    Fee matrix with 4 European service lines and EUR revenue tiers.
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Fee Schedule"

    # Canary
    canary_code = canaries.canary_for("tc16eu_fee_schedule")
    location = embed_canary_xlsx(wb, canary_code)
    canaries.set_location(
        "tc16eu_fee_schedule",
        f"{_INPUT_DIR}/fee_schedule_eu.xlsx",
        location,
    )

    # Header row
    headers = [
        "Service Type",
        "Revenue Tier",
        "Entity Count Range",
        "Base Fee (\u20ac)",
        "Per-Entity Adder (\u20ac)",
        "Complexity Multiplier",
    ]
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = _THIN_BORDER

    # Fee data rows — 4 service lines with EUR revenue tiers
    fee_data = [
        # Statutory Audit (IFRS)
        ("Statutory Audit (IFRS)", "Under \u20ac25M", "1\u20133", 65000, 6000, 1.00),
        ("Statutory Audit (IFRS)", "Under \u20ac25M", "4\u20136", 85000, 8000, 1.00),
        ("Statutory Audit (IFRS)", "\u20ac25M\u2013\u20ac75M", "1\u20133", 120000, 10000, 1.00),
        ("Statutory Audit (IFRS)", "\u20ac25M\u2013\u20ac75M", "4\u20136", 155000, 14000, 1.00),
        ("Statutory Audit (IFRS)", "\u20ac75M\u2013\u20ac250M", "1\u20133", 175000, 15000, 1.00),
        ("Statutory Audit (IFRS)", "\u20ac75M\u2013\u20ac250M", "4\u20136", 210000, 18000, 1.00),  # Cascade
        ("Statutory Audit (IFRS)", "\u20ac75M\u2013\u20ac250M", "7\u201310", 260000, 20000, 1.00),
        ("Statutory Audit (IFRS)", "Over \u20ac250M", "1\u20133", 320000, 22000, 1.00),
        ("Statutory Audit (IFRS)", "Over \u20ac250M", "4\u20136", 380000, 25000, 1.00),
        ("Statutory Audit (IFRS)", "Over \u20ac250M", "7\u201310", 450000, 28000, 1.00),
        # CIT Compliance
        ("CIT Compliance", "Under \u20ac25M", "1\u20133", 35000, 8000, 1.00),
        ("CIT Compliance", "Under \u20ac25M", "4\u20136", 45000, 10000, 1.00),
        ("CIT Compliance", "\u20ac25M\u2013\u20ac75M", "1\u20133", 55000, 11000, 1.00),
        ("CIT Compliance", "\u20ac25M\u2013\u20ac75M", "4\u20136", 65000, 13000, 1.00),
        ("CIT Compliance", "\u20ac75M\u2013\u20ac250M", "1\u20133", 60000, 12000, 1.00),
        ("CIT Compliance", "\u20ac75M\u2013\u20ac250M", "4\u20136", 72000, 15000, 1.00),  # Cascade
        ("CIT Compliance", "\u20ac75M\u2013\u20ac250M", "7\u201310", 90000, 17000, 1.00),
        ("CIT Compliance", "Over \u20ac250M", "1\u20133", 110000, 18000, 1.00),
        ("CIT Compliance", "Over \u20ac250M", "4\u20136", 130000, 20000, 1.00),
        ("CIT Compliance", "Over \u20ac250M", "7\u201310", 155000, 22000, 1.00),
        # VAT Compliance
        ("VAT Compliance", "Under \u20ac25M", "1\u20133", 18000, 5000, 1.00),
        ("VAT Compliance", "Under \u20ac25M", "4\u20136", 24000, 6000, 1.00),
        ("VAT Compliance", "\u20ac25M\u2013\u20ac75M", "1\u20133", 28000, 6000, 1.00),
        ("VAT Compliance", "\u20ac25M\u2013\u20ac75M", "4\u20136", 32000, 7000, 1.00),
        ("VAT Compliance", "\u20ac75M\u2013\u20ac250M", "1\u20133", 30000, 7000, 1.00),
        ("VAT Compliance", "\u20ac75M\u2013\u20ac250M", "4\u20136", 35000, 8000, 1.00),  # Cascade
        ("VAT Compliance", "\u20ac75M\u2013\u20ac250M", "7\u201310", 42000, 9000, 1.00),
        ("VAT Compliance", "Over \u20ac250M", "1\u20133", 50000, 10000, 1.00),
        ("VAT Compliance", "Over \u20ac250M", "4\u20136", 60000, 12000, 1.00),
        ("VAT Compliance", "Over \u20ac250M", "7\u201310", 72000, 14000, 1.00),
        # Transfer Pricing
        ("Transfer Pricing", "Under \u20ac25M", "1\u20133", 25000, 8000, 1.00),
        ("Transfer Pricing", "Under \u20ac25M", "4\u20136", 32000, 10000, 1.00),
        ("Transfer Pricing", "\u20ac25M\u2013\u20ac75M", "1\u20133", 38000, 10000, 1.00),
        ("Transfer Pricing", "\u20ac25M\u2013\u20ac75M", "4\u20136", 44000, 11000, 1.00),
        ("Transfer Pricing", "\u20ac75M\u2013\u20ac250M", "1\u20133", 42000, 10000, 1.00),
        ("Transfer Pricing", "\u20ac75M\u2013\u20ac250M", "4\u20136", 48000, 12000, 1.00),  # Cascade
        ("Transfer Pricing", "\u20ac75M\u2013\u20ac250M", "7\u201310", 58000, 14000, 1.00),
        ("Transfer Pricing", "Over \u20ac250M", "1\u20133", 70000, 15000, 1.00),
        ("Transfer Pricing", "Over \u20ac250M", "4\u20136", 85000, 18000, 1.00),
        ("Transfer Pricing", "Over \u20ac250M", "7\u201310", 100000, 20000, 1.00),
    ]

    for row_idx, (stype, tier, ec, base, adder, mult) in enumerate(fee_data, start=2):
        ws.cell(row=row_idx, column=1, value=stype).font = _DATA_FONT
        ws.cell(row=row_idx, column=2, value=tier).font = _DATA_FONT
        ws.cell(row=row_idx, column=3, value=ec).font = _DATA_FONT

        c_base = ws.cell(row=row_idx, column=4, value=base)
        c_base.font = _DATA_FONT
        c_base.number_format = _CURRENCY_FMT

        c_adder = ws.cell(row=row_idx, column=5, value=adder)
        c_adder.font = _DATA_FONT
        c_adder.number_format = _CURRENCY_FMT

        c_mult = ws.cell(row=row_idx, column=6, value=mult)
        c_mult.font = _DATA_FONT
        c_mult.number_format = _MULTIPLIER_FMT

        for col in range(1, 7):
            ws.cell(row=row_idx, column=col).border = _THIN_BORDER
            ws.cell(row=row_idx, column=col).alignment = Alignment(
                horizontal="center" if col in (2, 3, 6) else "left",
                vertical="center",
            )

    # Complexity multiplier notes
    notes_start = len(fee_data) + 3
    ws.cell(row=notes_start, column=1, value="Complexity Multiplier Notes:").font = _BOLD_FONT
    notes = [
        (
            "IFRS First-Time Adoption",
            "1.20x",
            "Applied to statutory audit base fee when group is transitioning to IFRS",
        ),
        (
            "Cross-Border Restructuring",
            "1.15x",
            "Applied when group has active restructuring across jurisdictions",
        ),
        (
            "Pillar Two Readiness",
            "1.10x",
            "Applied for GloBE top-up tax preparation (Pillar Two)",
        ),
        (
            "UK Post-Brexit Customs",
            "1.10x",
            "Applied to VAT compliance when UK entity has import VAT obligations",
        ),
    ]
    for i, (factor, mult_val, desc) in enumerate(notes):
        row = notes_start + 1 + i
        ws.cell(row=row, column=1, value=factor).font = _DATA_FONT
        ws.cell(row=row, column=2, value=mult_val).font = _BOLD_FONT
        c_desc = ws.cell(row=row, column=3, value=desc)
        c_desc.font = _DATA_FONT
        c_desc.alignment = Alignment(wrap_text=True)

    # Per-entity adder notes
    adder_start = notes_start + len(notes) + 2
    ws.cell(row=adder_start, column=1, value="Per-Entity Adder Notes:").font = _BOLD_FONT
    c_adder_note = ws.cell(row=adder_start + 1, column=1, value=(
        "The per-entity adder applies to each entity beyond the first. "
        "For a group with 4 entities (1 parent + 3 subsidiaries), "
        "the adder is applied 3 times."
    ))
    c_adder_note.font = _DATA_FONT
    c_adder_note.alignment = Alignment(wrap_text=True)

    # Cascade Europe fee summary (pre-computed for convenience)
    # ERR-EU-016: The audit total omits the IFRS adoption complexity multiplier (1.20x).
    summary_start = adder_start + 3
    ws.cell(row=summary_start, column=1, value=(
        "Cascade Europe Holdings B.V. \u2014 Engagement Fee Summary"
    )).font = Font(name="Calibri", size=12, bold=True)

    summary_headers = ["Service Line", "Base Fee", "Adjustments", "Total Fee"]
    for col, h in enumerate(summary_headers, start=1):
        cell = ws.cell(row=summary_start + 1, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = _THIN_BORDER

    # Correct audit total = 210,000 * 1.20 = 252,000
    # Wrong audit total   = 210,000 * 1.00 = 210,000  (missing IFRS multiplier)
    correct_audit_total = int(_AUDIT_BASE_FEE * _IFRS_COMPLEXITY_MULTIPLIER)  # 252,000
    wrong_audit_total = int(formula_error(correct_audit_total, int(_AUDIT_BASE_FEE)))  # 210,000

    # Row: Statutory Audit (IFRS)
    audit_row = summary_start + 2
    ws.cell(row=audit_row, column=1, value="Statutory Audit (IFRS)").font = _DATA_FONT
    c = ws.cell(row=audit_row, column=2, value=int(_AUDIT_BASE_FEE))
    c.font = _DATA_FONT
    c.number_format = _CURRENCY_FMT
    ws.cell(row=audit_row, column=3, value="IFRS First-Time Adoption 1.20x").font = _DATA_FONT
    c = ws.cell(row=audit_row, column=4, value=wrong_audit_total)
    c.font = _BOLD_FONT
    c.number_format = _CURRENCY_FMT
    for col in range(1, 5):
        ws.cell(row=audit_row, column=col).border = _THIN_BORDER

    # Row: CIT Compliance
    cit_total = int(_CIT_FEE)  # 117,000
    cit_row = summary_start + 3
    ws.cell(row=cit_row, column=1, value="CIT Compliance").font = _DATA_FONT
    c = ws.cell(row=cit_row, column=2, value=int(_CIT_BASE_FEE))
    c.font = _DATA_FONT
    c.number_format = _CURRENCY_FMT
    ws.cell(
        row=cit_row, column=3,
        value=f"Per-entity adder \u00d7 {_ADDITIONAL_ENTITIES}",
    ).font = _DATA_FONT
    c = ws.cell(row=cit_row, column=4, value=cit_total)
    c.font = _BOLD_FONT
    c.number_format = _CURRENCY_FMT
    for col in range(1, 5):
        ws.cell(row=cit_row, column=col).border = _THIN_BORDER

    # Row: VAT Compliance
    vat_total = int(_VAT_FEE)  # 59,000
    vat_row = summary_start + 4
    ws.cell(row=vat_row, column=1, value="VAT Compliance").font = _DATA_FONT
    c = ws.cell(row=vat_row, column=2, value=int(_VAT_BASE_FEE))
    c.font = _DATA_FONT
    c.number_format = _CURRENCY_FMT
    ws.cell(
        row=vat_row, column=3,
        value=f"Per-entity adder \u00d7 {_ADDITIONAL_ENTITIES}",
    ).font = _DATA_FONT
    c = ws.cell(row=vat_row, column=4, value=vat_total)
    c.font = _BOLD_FONT
    c.number_format = _CURRENCY_FMT
    for col in range(1, 5):
        ws.cell(row=vat_row, column=col).border = _THIN_BORDER

    # Row: Transfer Pricing
    tp_total = int(_TP_FEE)  # 84,000
    tp_row = summary_start + 5
    ws.cell(row=tp_row, column=1, value="Transfer Pricing").font = _DATA_FONT
    c = ws.cell(row=tp_row, column=2, value=int(_TP_BASE_FEE))
    c.font = _DATA_FONT
    c.number_format = _CURRENCY_FMT
    ws.cell(
        row=tp_row, column=3,
        value=f"Per-entity adder \u00d7 {_ADDITIONAL_ENTITIES}",
    ).font = _DATA_FONT
    c = ws.cell(row=tp_row, column=4, value=tp_total)
    c.font = _BOLD_FONT
    c.number_format = _CURRENCY_FMT
    for col in range(1, 5):
        ws.cell(row=tp_row, column=col).border = _THIN_BORDER

    # Total row
    total_row = summary_start + 6
    wrong_total = wrong_audit_total + cit_total + vat_total + tp_total
    ws.cell(row=total_row, column=1, value="Total Engagement Fee").font = _BOLD_FONT
    c = ws.cell(row=total_row, column=4, value=wrong_total)
    c.font = Font(name="Calibri", size=11, bold=True, color="1A3C6E")
    c.number_format = _CURRENCY_FMT
    for col in range(1, 5):
        ws.cell(row=total_row, column=col).border = _THIN_BORDER

    # Register planted error
    errors.add(PlantedError(
        error_id="ERR-EU-016",
        file=f"{_INPUT_DIR}/fee_schedule_eu.xlsx",
        location=(
            "Sheet 'Fee Schedule', Cascade Europe Holdings B.V. Fee Summary, "
            "Statutory Audit Total Fee cell"
        ),
        type="formula_error",
        description=(
            f"Statutory audit total fee shows \u20ac{wrong_audit_total:,} instead of "
            f"\u20ac{correct_audit_total:,} \u2014 the IFRS first-time adoption "
            f"complexity multiplier (1.20x) was not applied to the base fee"
        ),
        severity="material",
        which_test_cases_should_catch=["TC-16-EU"],
    ))

    # Column widths
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 20
    ws.column_dimensions["C"].width = 22
    ws.column_dimensions["D"].width = 16
    ws.column_dimensions["E"].width = 20
    ws.column_dimensions["F"].width = 22

    # Page setup for PDF rendering
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    # Save
    path = output_dir / _INPUT_DIR / "fee_schedule_eu.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, path)

    manifest.register(
        f"{_INPUT_DIR}/fee_schedule_eu.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=[_TC],
    )


def _write_engagement_template(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Create test_cases/TC-16-EU/input_files/engagement_letter_template_eu.docx.

    European variant template with ISA/IFRS/OECD language and NL law
    limitation of liability (not a copy of the US template).
    """
    doc = Document()

    # Canary
    canary_code = canaries.canary_for("tc16eu_engagement_template")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "tc16eu_engagement_template",
        f"{_INPUT_DIR}/engagement_letter_template_eu.docx",
        location,
    )

    # Firm header
    doc.add_paragraph("Mitchell & Partners International LLP")
    doc.add_paragraph("Zuidas, Gustav Mahlerlaan 10  |  1082 PP Amsterdam  |  +31 20 555 0100")
    doc.add_paragraph("")

    # Date and addressee
    doc.add_paragraph("<<start_date>>")
    doc.add_paragraph("<<client_name>>")
    doc.add_paragraph("Cascade Europe Holdings B.V.")
    doc.add_paragraph("")

    doc.add_paragraph("Dear <<client_name>>,")
    doc.add_paragraph("")

    # Engagement Scope
    doc.add_heading("Engagement Scope", level=2)
    doc.add_paragraph(
        "We are pleased to confirm our understanding of the services we will provide "
        "to Cascade Europe Holdings B.V. and its subsidiaries for the fiscal year "
        "ending 31 December 2026."
    )
    doc.add_paragraph("<<engagement_scope>>")
    doc.add_paragraph(
        "Statutory audits will be conducted in accordance with International Standards "
        "on Auditing (ISA) as adopted in the respective jurisdictions."
    )
    doc.add_paragraph(
        "Corporate income tax compliance covers annual CIT returns in all four "
        "jurisdictions (Netherlands, Germany, France, United Kingdom)."
    )
    doc.add_paragraph(
        "Transfer pricing documentation will follow OECD BEPS Action 13 requirements "
        "(Master File and Local File)."
    )
    doc.add_paragraph(
        "VAT compliance covers periodic VAT return preparation for all EU-registered "
        "and UK-registered entities."
    )

    # Fees and Billing
    doc.add_heading("Fees and Billing", level=2)
    doc.add_paragraph(
        "Our fees for the services described above will be <<fee_amount>>. "
        "This fee estimate is based on the scope of services outlined herein "
        "and the anticipated level of complexity."
    )
    doc.add_paragraph(
        "Payment terms: <<payment_terms>>. Invoices will be submitted monthly "
        "as services are rendered. Any additional services outside the scope "
        "of this engagement will be billed separately at agreed-upon rates."
    )

    # Engagement Period
    doc.add_heading("Engagement Period", level=2)
    doc.add_paragraph(
        "This engagement will commence on <<start_date>> and is expected to conclude "
        "upon completion of the audit opinion, CIT filings, VAT returns, and transfer "
        "pricing documentation for the fiscal year ending 31 December 2026."
    )

    # Responsibilities
    doc.add_heading("Responsibilities", level=2)
    doc.add_paragraph("Our Responsibilities")
    doc.add_paragraph(
        "We will perform the engagement in accordance with professional standards "
        "applicable to the services described, including ISA as adopted in the "
        "relevant EU member states and the United Kingdom, and OECD Transfer "
        "Pricing Guidelines."
    )
    doc.add_paragraph("Management Responsibilities")
    doc.add_paragraph(
        "Management is responsible for providing us with access to all information "
        "and personnel necessary to complete the engagement, including the "
        "preparation of complete and accurate financial records in accordance "
        "with applicable accounting standards (IFRS)."
    )

    # Confidentiality
    doc.add_heading("Confidentiality", level=2)
    doc.add_paragraph(
        "We will maintain the confidentiality of all information obtained during "
        "the course of the engagement, in accordance with applicable professional "
        "standards and the EU General Data Protection Regulation (GDPR)."
    )

    # Limitation of Liability
    doc.add_heading("Limitation of Liability", level=2)
    doc.add_paragraph(
        "To the fullest extent permitted by the laws of the Netherlands, the "
        "aggregate liability of Mitchell & Partners International LLP arising "
        "from or in connection with this engagement shall not exceed the total "
        "fees paid under this engagement letter. This limitation applies to all "
        "claims whether in contract, tort (including negligence), or otherwise."
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "If the above terms are acceptable, please sign and return a copy of "
        "this letter. We appreciate the opportunity to serve Cascade Europe "
        "Holdings B.V."
    )
    doc.add_paragraph("")

    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("")
    doc.add_paragraph("<<partner_name>>")
    doc.add_paragraph("Engagement Partner")
    doc.add_paragraph("Mitchell & Partners International LLP")
    doc.add_paragraph("")
    doc.add_paragraph("")

    doc.add_paragraph("ACCEPTED AND AGREED")
    doc.add_paragraph("")
    doc.add_paragraph("Signature: _________________________________")
    doc.add_paragraph("")
    doc.add_paragraph("Name: _________________________________")
    doc.add_paragraph("")
    doc.add_paragraph("Title: _________________________________")
    doc.add_paragraph("")
    doc.add_paragraph("Date: _________________________________")

    # Save
    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME
    path = output_dir / _INPUT_DIR / "engagement_letter_template_eu.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, path)

    manifest.register(
        f"{_INPUT_DIR}/engagement_letter_template_eu.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# -- Prompt and expected behavior ---------------------------------------------


def _write_prompt(output_dir: Path) -> None:
    """Write test_cases/TC-16-EU/prompt.md."""
    text = """\
Generate a draft engagement letter for a combined audit, tax, and advisory
engagement for Cascade Europe Holdings B.V. using the template provided.

1. Look up the appropriate fees from the fee schedule based on the client's
   revenue tier and entity count.
2. The statutory audit fee should include the complexity multiplier for IFRS
   first-time adoption (1.20x) since the group is transitioning to IFRS.
3. The CIT compliance fee should cover annual corporate income tax returns
   for the parent and 3 subsidiaries across all four jurisdictions.
4. The VAT compliance fee should cover all EU and UK VAT-registered entities.
5. The transfer pricing fee should cover OECD master file and local files.
6. Populate the template with the correct values.
7. Set the start date to 15 March 2026 and payment terms to Net 30.
8. Use "Pieter de Jong" as the engagement partner.

Save as a Word document.
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write test_cases/TC-16-EU/expected_behavior.md."""
    text = """\
# TC-16-EU: European Engagement Letter Generation — Expected Behavior

## Fee Calculation
- The agent must identify Cascade Europe's revenue tier (\u20ac75M\u2013\u20ac250M)
  and entity count (4 entities \u2192 "4\u20136" range) from the client profile.
- Statutory Audit (IFRS) base fee: \u20ac210,000 (from fee schedule, \u20ac75M\u2013\u20ac250M tier, 4\u20136 entities).
- IFRS first-time adoption complexity multiplier: 1.20x (from complexity notes and prompt instruction).
- Statutory Audit fee: \u20ac210,000 \u00d7 1.20 = \u20ac252,000.
- CIT Compliance base fee: \u20ac72,000.
- CIT per-entity adder: \u20ac15,000 \u00d7 3 additional entities = \u20ac45,000.
- CIT Compliance fee: \u20ac72,000 + \u20ac45,000 = \u20ac117,000.
- VAT Compliance base fee: \u20ac35,000.
- VAT per-entity adder: \u20ac8,000 \u00d7 3 additional entities = \u20ac24,000.
- VAT Compliance fee: \u20ac35,000 + \u20ac24,000 = \u20ac59,000.
- Transfer Pricing base fee: \u20ac48,000.
- Transfer Pricing per-entity adder: \u20ac12,000 \u00d7 3 additional entities = \u20ac36,000.
- Transfer Pricing fee: \u20ac48,000 + \u20ac36,000 = \u20ac84,000.
- Total engagement fee: \u20ac252,000 + \u20ac117,000 + \u20ac59,000 + \u20ac84,000 = \u20ac512,000.

## Merge Field Population
- `<<client_name>>` \u2192 "Cascade Europe Holdings B.V."
- `<<engagement_scope>>` \u2192 Description covering combined statutory audit (IFRS),
  CIT compliance (NL/DE/FR/UK), VAT compliance, and transfer pricing documentation,
  including IFRS first-time adoption procedures.
- `<<fee_amount>>` \u2192 "\u20ac512,000"
  (or breakdown: Audit \u20ac252,000 + CIT \u20ac117,000 + VAT \u20ac59,000 + TP \u20ac84,000).
- `<<payment_terms>>` \u2192 "Net 30"
- `<<start_date>>` \u2192 "15 March 2026"
- `<<partner_name>>` \u2192 "Pieter de Jong"

## Template Preservation
- The agent must preserve the original formatting of the engagement letter template.
- All sections (Engagement Scope, Fees and Billing, Engagement Period,
  Responsibilities, Confidentiality, Limitation of Liability, Acceptance)
  should remain intact.
- The firm name "Mitchell & Partners International LLP" and Amsterdam address
  should be preserved.
- ISA/IFRS/OECD language in engagement scope section should be preserved.

## Output Format
- Word document (.docx) that looks like a professional engagement letter.
- All merge fields replaced with actual values \u2014 no unreplaced placeholders.

## Common Mistakes to Watch For
- Using US revenue tiers or USD amounts when EUR fee schedule is provided.
- Calculating CIT/VAT for parent only, ignoring the 3 subsidiaries across 4 jurisdictions.
- Forgetting to apply the 1.20x IFRS first-time adoption multiplier to the audit fee.
- Copying the summary total from the fee schedule (which contains ERR-EU-016)
  instead of computing from base fee \u00d7 multiplier.
- Populating engagement scope with US-style language (federal/state returns, GAAP)
  instead of European terminology (ISA, IFRS, CIT, VAT, OECD TP).
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# -- Gold standard ------------------------------------------------------------


def _fmt_euros(v: Decimal | int) -> str:
    """Format a value as whole-euro string with commas."""
    return f"\u20ac{int(v):,}"


@register_gold("TC-16-EU")
def _tc16_eu_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model: Any,
) -> GoldStandard:
    """Gold standard for TC-16-EU: European Engagement Letter Generation."""
    return GoldStandard(
        test_case="TC-16-EU",
        expected_outputs={
            "file_type": "docx",
            "fee_calculation": {
                "audit_base_fee": _fmt_euros(_AUDIT_BASE_FEE),
                "ifrs_complexity_multiplier": str(_IFRS_COMPLEXITY_MULTIPLIER),
                "audit_fee": _fmt_euros(_AUDIT_FEE),
                "cit_base_fee": _fmt_euros(_CIT_BASE_FEE),
                "cit_per_entity_adder": _fmt_euros(_CIT_PER_ENTITY_ADDER),
                "vat_base_fee": _fmt_euros(_VAT_BASE_FEE),
                "vat_per_entity_adder": _fmt_euros(_VAT_PER_ENTITY_ADDER),
                "tp_base_fee": _fmt_euros(_TP_BASE_FEE),
                "tp_per_entity_adder": _fmt_euros(_TP_PER_ENTITY_ADDER),
                "additional_entities": _ADDITIONAL_ENTITIES,
                "cit_fee": _fmt_euros(_CIT_FEE),
                "vat_fee": _fmt_euros(_VAT_FEE),
                "tp_fee": _fmt_euros(_TP_FEE),
                "total_engagement_fee": _fmt_euros(_TOTAL_FEE),
            },
            "merge_fields": {
                "client_name": _CLIENT_NAME,
                "engagement_scope": (
                    "Combined statutory audit (IFRS), corporate income tax compliance "
                    "(NL/DE/FR/UK), VAT compliance, and transfer pricing documentation "
                    "engagement, including IFRS first-time adoption procedures"
                ),
                "fee_amount": _fmt_euros(_TOTAL_FEE),
                "payment_terms": _PAYMENT_TERMS,
                "start_date": _START_DATE,
                "partner_name": _PARTNER_NAME,
            },
            "template_sections_preserved": [
                "Engagement Scope",
                "Fees and Billing",
                "Engagement Period",
                "Responsibilities",
                "Confidentiality",
                "Limitation of Liability",
                "ACCEPTED AND AGREED",
            ],
        },
        canary_verification={
            "read_client_profile": canaries.canary_for("tc16eu_client_profile"),
            "read_fee_schedule": canaries.canary_for("tc16eu_fee_schedule"),
            "read_engagement_template": canaries.canary_for("tc16eu_engagement_template"),
        },
        error_detection={
            "ERR-EU-016": (
                "Fee schedule statutory audit total shows \u20ac210,000 instead of \u20ac252,000 \u2014 "
                "IFRS first-time adoption complexity multiplier (1.20x) not applied"
            ),
        },
        scoring_hints={
            "correctness": (
                "Statutory Audit = \u20ac210,000 \u00d7 1.20 = \u20ac252,000; "
                "CIT = \u20ac72,000 + (\u20ac15,000 \u00d7 3) = \u20ac117,000; "
                "VAT = \u20ac35,000 + (\u20ac8,000 \u00d7 3) = \u20ac59,000; "
                "TP = \u20ac48,000 + (\u20ac12,000 \u00d7 3) = \u20ac84,000; "
                "Total = \u20ac512,000"
            ),
            "completeness": (
                "All 6 merge fields populated; template sections preserved; "
                "fee breakdown visible in the letter; all 4 service lines included"
            ),
            "format_compliance": (
                "Valid docx; professional engagement letter appearance; "
                "no unreplaced merge field placeholders; European firm name and address"
            ),
            "robustness": (
                "Correctly identifies EUR revenue tier and entity count from "
                "client profile; applies IFRS complexity multiplier only to audit fee; "
                "uses European terminology (ISA, IFRS, CIT, OECD TP)"
            ),
            "communication": (
                "Engagement scope clearly describes combined audit + CIT + VAT + TP services; "
                "IFRS first-time adoption mentioned; fee amount clearly stated in EUR"
            ),
        },
    )


# -- Public entry point -------------------------------------------------------


def emit_tc16_eu(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
    **kwargs: object,
) -> None:
    """Write all TC-16-EU files to *output_dir*."""
    _write_client_profile(output_dir, canaries, manifest)
    _write_fee_schedule(output_dir, canaries, errors, manifest)
    _write_engagement_template(output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
