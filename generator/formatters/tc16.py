"""Formatter: TC-16 — Engagement Letter Generation (Cross-service, Routine).

Emits:
- test_cases/TC-16/input_files/client_profile.docx
  Cascade Industries profile with entity count, revenue tier, complexity
  factors, key contacts.
- test_cases/TC-16/input_files/fee_schedule.xlsx
  Fee matrix: Service Type, Revenue Tier, Entity Count, Base Fee,
  Per-Entity Adder, Complexity Multiplier.
- test_cases/TC-16/input_files/engagement_letter_template.docx
  Copy of templates/engagement_letter_template.docx with merge fields.
- test_cases/TC-16/prompt.md
- test_cases/TC-16/expected_behavior.md
- gold_standards/TC-16_gold.json

Gold standard (from prompt.md):
  Audit base fee $285,000 × 1.15 IPO complexity = $327,750.
  Tax base fee $95,000 + $12,000 per additional entity × 3 = $131,000.
  Total engagement = $458,750.
  All merge fields correctly populated. Template formatting preserved.

One planted error (ERR-013): formula_error in fee_schedule.xlsx.
"""

from __future__ import annotations

import datetime
import io
from decimal import Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import openpyxl
from docx import Document
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from generator.canaries import CanaryRegistry, embed_canary_docx, embed_canary_xlsx
from generator.errors import ErrorRegistry, PlantedError, formula_error
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-16"
_INPUT_DIR = f"test_cases/{_TC}/input_files"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)
_FIXED_ZIP_DT = (2025, 3, 15, 9, 0, 0)

# Fee schedule values from prompt.md gold standard
_AUDIT_BASE_FEE = Decimal("285000")
_IPO_COMPLEXITY_MULTIPLIER = Decimal("1.15")
_TAX_BASE_FEE = Decimal("95000")
_PER_ENTITY_TAX_ADDER = Decimal("12000")
_ADDITIONAL_ENTITIES = 3  # 3 subsidiaries beyond the parent

# Computed gold values
_AUDIT_FEE = _AUDIT_BASE_FEE * _IPO_COMPLEXITY_MULTIPLIER  # $327,750
_TAX_FEE = _TAX_BASE_FEE + _PER_ENTITY_TAX_ADDER * _ADDITIONAL_ENTITIES  # $131,000
_TOTAL_FEE = _AUDIT_FEE + _TAX_FEE  # $458,750

# Merge field values from prompt.md
_CLIENT_NAME = "Cascade Industries, Inc."
_PARTNER_NAME = "Sarah Chen"
_START_DATE = "March 15, 2026"
_PAYMENT_TERMS = "Net 30"

# ── Deterministic file helpers ────────────────────────────────────────────────


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)

    buf = io.BytesIO()
    wb.properties.created = _FIXED_DATETIME
    wb.properties.modified = _FIXED_DATETIME
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _copy_docx_deterministic(src_path: Path, dst_path: Path, canary_code: str) -> str:
    """Copy a docx, embed canary, and repack with fixed timestamps.

    Returns the canary location description.
    """
    doc = Document(str(src_path))
    location = embed_canary_docx(doc, canary_code)
    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME
    dst_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, dst_path)
    return location


# ── Styling helpers ──────────────────────────────────────────────────────────

_HEADER_FILL = PatternFill(start_color="1A3C6E", end_color="1A3C6E", fill_type="solid")
_HEADER_FONT = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
_DATA_FONT = Font(name="Calibri", size=11)
_BOLD_FONT = Font(name="Calibri", size=11, bold=True)
_CURRENCY_FMT = '#,##0'
_MULTIPLIER_FMT = '0.00"x"'
_THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)


# ── File generators ──────────────────────────────────────────────────────────


def _write_client_profile(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Create test_cases/TC-16/input_files/client_profile.docx."""
    doc = Document()

    # -- Canary
    canary_code = canaries.canary_for("tc16_client_profile")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "tc16_client_profile",
        f"{_INPUT_DIR}/client_profile.docx",
        location,
    )

    # -- Title
    doc.add_heading("Client Profile", level=1)

    # -- Company overview
    doc.add_heading("Company Overview", level=2)
    doc.add_paragraph(f"Company Name: {_CLIENT_NAME}")
    doc.add_paragraph("Entity Type: US C-Corporation")
    doc.add_paragraph("Industry: Mid-market manufacturer (industrial parts, specialty materials, distribution)")
    doc.add_paragraph("Headquarters: Portland, Oregon")
    doc.add_paragraph("Fiscal Year End: December 31")

    # -- Revenue and size
    doc.add_heading("Revenue and Size", level=2)
    doc.add_paragraph("Consolidated Revenue (FY2025): ~$200,000,000")
    doc.add_paragraph("Revenue Tier: $100M–$500M")
    doc.add_paragraph("Total Employees: 850")

    # -- Entity structure
    doc.add_heading("Entity Structure", level=2)
    doc.add_paragraph(
        "Cascade Industries operates through a parent company and three wholly-owned subsidiaries:"
    )

    # Build entity table
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["Entity", "Location", "Type", "Revenue"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    entities = [
        ("Cascade Industries, Inc. (Parent)", "Portland, OR", "Parent — C-Corp", "Consolidated: $200M"),
        ("Cascade Precision Components LLC", "Portland, OR", "Core manufacturing", "$95M"),
        ("Cascade Advanced Materials, Inc.", "Austin, TX", "Specialty materials R&D", "$65M"),
        ("Cascade Distribution Services LLC", "Chicago, IL", "Warehousing & logistics", "$40M"),
    ]
    for row_idx, (name, loc, etype, rev) in enumerate(entities, start=1):
        table.rows[row_idx].cells[0].text = name
        table.rows[row_idx].cells[1].text = loc
        table.rows[row_idx].cells[2].text = etype
        table.rows[row_idx].cells[3].text = rev

    doc.add_paragraph("")  # spacer

    doc.add_paragraph(f"Total Entities: 4 (1 parent + {_ADDITIONAL_ENTITIES} subsidiaries)")

    # -- Complexity factors
    doc.add_heading("Complexity Factors", level=2)
    doc.add_paragraph(
        "The company is currently evaluating a potential IPO in the next 18–24 months. "
        "This introduces additional complexity for audit engagements, requiring "
        "public-readiness procedures and enhanced documentation."
    )
    doc.add_paragraph("IPO Readiness: Yes — considering IPO within 18–24 months")
    doc.add_paragraph("Multi-State Operations: Yes — OR, TX, IL, plus remote employees in CA, WA, NY")
    doc.add_paragraph(
        "Intercompany Transactions: Significant — cost-plus transfers, "
        "management fees, intercompany loans"
    )
    doc.add_paragraph("R&D Activity: Substantial — ~12% of Advanced Materials revenue")

    # -- Key contacts
    doc.add_heading("Key Contacts", level=2)
    contacts = [
        ("CEO", "Robert Chen", "robert.chen@cascadeind.com"),
        ("CFO", "Maria Santos", "maria.santos@cascadeind.com"),
        ("VP Finance", "David Kim", "david.kim@cascadeind.com"),
        ("External Counsel", "Patterson & Associates LLP", "jpatterson@pattersonlaw.com"),
    ]
    for title, name, email in contacts:
        doc.add_paragraph(f"{title}: {name} ({email})")

    # -- Engagement history
    doc.add_heading("Engagement History", level=2)
    doc.add_paragraph(
        "Mitchell & Partners LLP has served as Cascade Industries' external audit "
        "firm for the past 3 years. This engagement letter covers a combined "
        "audit and tax engagement for FY2026."
    )

    # -- Save
    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME
    path = output_dir / _INPUT_DIR / "client_profile.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, path)

    manifest.register(
        f"{_INPUT_DIR}/client_profile.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


def _write_fee_schedule(
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Create test_cases/TC-16/input_files/fee_schedule.xlsx.

    Fee matrix with columns:
    Service Type | Revenue Tier | Entity Count | Base Fee | Per-Entity Adder | Complexity Multiplier
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Fee Schedule"

    # -- Canary
    canary_code = canaries.canary_for("tc16_fee_schedule")
    location = embed_canary_xlsx(wb, canary_code)
    canaries.set_location(
        "tc16_fee_schedule",
        f"{_INPUT_DIR}/fee_schedule.xlsx",
        location,
    )

    # -- Header row
    headers = [
        "Service Type",
        "Revenue Tier",
        "Entity Count Range",
        "Base Fee",
        "Per-Entity Adder",
        "Complexity Multiplier",
    ]
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = _THIN_BORDER

    # -- Fee data rows
    # The agent needs to match Cascade's revenue tier ($200M → $100M–$500M)
    # and entity count (4 entities → 4–6 range) to find the right row.
    fee_data = [
        # Audit fees
        ("Audit", "Under $50M", "1–3", 85000, 8000, 1.00),
        ("Audit", "Under $50M", "4–6", 110000, 10000, 1.00),
        ("Audit", "$50M–$100M", "1–3", 165000, 10000, 1.00),
        ("Audit", "$50M–$100M", "4–6", 195000, 12000, 1.00),
        ("Audit", "$100M–$500M", "1–3", 245000, 12000, 1.00),
        ("Audit", "$100M–$500M", "4–6", 285000, 15000, 1.00),  # ← Cascade matches here
        ("Audit", "$100M–$500M", "7–10", 340000, 18000, 1.00),
        ("Audit", "Over $500M", "1–3", 420000, 20000, 1.00),
        ("Audit", "Over $500M", "4–6", 485000, 22000, 1.00),
        ("Audit", "Over $500M", "7–10", 560000, 25000, 1.00),
        # Tax fees
        ("Tax — Federal & State", "Under $50M", "1–3", 45000, 8000, 1.00),
        ("Tax — Federal & State", "Under $50M", "4–6", 55000, 10000, 1.00),
        ("Tax — Federal & State", "$50M–$100M", "1–3", 72000, 10000, 1.00),
        ("Tax — Federal & State", "$50M–$100M", "4–6", 85000, 11000, 1.00),
        ("Tax — Federal & State", "$100M–$500M", "1–3", 80000, 10000, 1.00),
        ("Tax — Federal & State", "$100M–$500M", "4–6", 95000, 12000, 1.00),  # ← Cascade matches here
        ("Tax — Federal & State", "$100M–$500M", "7–10", 115000, 14000, 1.00),
        ("Tax — Federal & State", "Over $500M", "1–3", 140000, 15000, 1.00),
        ("Tax — Federal & State", "Over $500M", "4–6", 165000, 18000, 1.00),
        ("Tax — Federal & State", "Over $500M", "7–10", 195000, 20000, 1.00),
        # Advisory fees (for completeness — not used in this engagement)
        ("Advisory — Financial", "$50M–$100M", "Any", 120000, 0, 1.00),
        ("Advisory — Financial", "$100M–$500M", "Any", 175000, 0, 1.00),
        ("Advisory — Financial", "Over $500M", "Any", 250000, 0, 1.00),
    ]

    for row_idx, (stype, tier, ec, base, adder, mult) in enumerate(fee_data, start=2):
        ws.cell(row=row_idx, column=1, value=stype).font = _DATA_FONT
        ws.cell(row=row_idx, column=2, value=tier).font = _DATA_FONT
        ws.cell(row=row_idx, column=3, value=ec).font = _DATA_FONT

        c_base = ws.cell(row=row_idx, column=4, value=base)
        c_base.font = _DATA_FONT
        c_base.number_format = _CURRENCY_FMT

        c_adder = ws.cell(row=row_idx, column=5, value=adder)
        c_adder.font = _DATA_FONT
        c_adder.number_format = _CURRENCY_FMT

        c_mult = ws.cell(row=row_idx, column=6, value=mult)
        c_mult.font = _DATA_FONT
        c_mult.number_format = _MULTIPLIER_FMT

        for col in range(1, 7):
            ws.cell(row=row_idx, column=col).border = _THIN_BORDER
            ws.cell(row=row_idx, column=col).alignment = Alignment(
                horizontal="center" if col in (2, 3, 6) else "left",
                vertical="center",
            )

    # -- Complexity multiplier notes
    notes_start = len(fee_data) + 3
    ws.cell(row=notes_start, column=1, value="Complexity Multiplier Notes:").font = _BOLD_FONT
    notes = [
        ("IPO Readiness / Public-Readiness", "1.15x", "Applied to audit base fee when client is preparing for IPO"),
        ("International Operations", "1.20x", "Applied when client has foreign subsidiaries or operations"),
        ("SEC Reporting", "1.25x", "Applied for SEC-registered entities"),
        ("Complex Revenue Recognition", "1.10x", "Applied when ASC 606 analysis is significant"),
    ]
    for i, (factor, mult_val, desc) in enumerate(notes):
        row = notes_start + 1 + i
        ws.cell(row=row, column=1, value=factor).font = _DATA_FONT
        ws.cell(row=row, column=2, value=mult_val).font = _BOLD_FONT
        c_desc = ws.cell(row=row, column=3, value=desc)
        c_desc.font = _DATA_FONT
        c_desc.alignment = Alignment(wrap_text=True)

    # -- Per-Entity Adder notes
    adder_start = notes_start + len(notes) + 2
    ws.cell(row=adder_start, column=1, value="Per-Entity Adder Notes:").font = _BOLD_FONT
    c_adder_note = ws.cell(row=adder_start + 1, column=1, value=(
        "The per-entity adder applies to each entity beyond the first. "
        "For a group with 4 entities (1 parent + 3 subsidiaries), "
        "the adder is applied 3 times."
    ))
    c_adder_note.font = _DATA_FONT
    c_adder_note.alignment = Alignment(wrap_text=True)

    # -- Cascade Industries fee summary (pre-computed for convenience)
    # ERR-013: The audit total omits the IPO complexity multiplier (1.15x).
    summary_start = adder_start + 3
    ws.cell(row=summary_start, column=1, value="Cascade Industries — Engagement Fee Summary").font = Font(
        name="Calibri", size=12, bold=True,
    )

    summary_headers = ["Service Line", "Base Fee", "Adjustments", "Total Fee"]
    for col, h in enumerate(summary_headers, start=1):
        cell = ws.cell(row=summary_start + 1, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = _THIN_BORDER

    # Correct audit total = 285,000 * 1.15 = 327,750
    # Wrong audit total   = 285,000 * 1.00 = 285,000  (missing IPO multiplier)
    correct_audit_total = int(_AUDIT_BASE_FEE * _IPO_COMPLEXITY_MULTIPLIER)  # 327,750
    wrong_audit_total = int(formula_error(correct_audit_total, int(_AUDIT_BASE_FEE)))  # 285,000

    audit_row = summary_start + 2
    ws.cell(row=audit_row, column=1, value="Audit").font = _DATA_FONT
    c = ws.cell(row=audit_row, column=2, value=int(_AUDIT_BASE_FEE))
    c.font = _DATA_FONT
    c.number_format = _CURRENCY_FMT
    ws.cell(row=audit_row, column=3, value="IPO Readiness 1.15x").font = _DATA_FONT
    c = ws.cell(row=audit_row, column=4, value=wrong_audit_total)
    c.font = _BOLD_FONT
    c.number_format = _CURRENCY_FMT
    for col in range(1, 5):
        ws.cell(row=audit_row, column=col).border = _THIN_BORDER

    tax_total = int(_TAX_FEE)  # 131,000 — correct
    tax_row = summary_start + 3
    ws.cell(row=tax_row, column=1, value="Tax — Federal & State").font = _DATA_FONT
    c = ws.cell(row=tax_row, column=2, value=int(_TAX_BASE_FEE))
    c.font = _DATA_FONT
    c.number_format = _CURRENCY_FMT
    ws.cell(row=tax_row, column=3, value=f"Per-entity adder × {_ADDITIONAL_ENTITIES}").font = _DATA_FONT
    c = ws.cell(row=tax_row, column=4, value=tax_total)
    c.font = _BOLD_FONT
    c.number_format = _CURRENCY_FMT
    for col in range(1, 5):
        ws.cell(row=tax_row, column=col).border = _THIN_BORDER

    total_row = summary_start + 4
    wrong_total = wrong_audit_total + tax_total  # 285,000 + 131,000 = 416,000
    _ = correct_audit_total + tax_total  # 327,750 + 131,000 = 458,750 (correct, not written)
    ws.cell(row=total_row, column=1, value="Total Engagement Fee").font = _BOLD_FONT
    c = ws.cell(row=total_row, column=4, value=wrong_total)
    c.font = Font(name="Calibri", size=11, bold=True, color="1A3C6E")
    c.number_format = _CURRENCY_FMT
    for col in range(1, 5):
        ws.cell(row=total_row, column=col).border = _THIN_BORDER

    # Register planted error
    errors.add(PlantedError(
        error_id="ERR-013",
        file=f"{_INPUT_DIR}/fee_schedule.xlsx",
        location="Sheet 'Fee Schedule', Cascade Industries Fee Summary, Audit Total Fee cell",
        type="formula_error",
        description=(
            f"Audit total fee shows ${wrong_audit_total:,} instead of "
            f"${correct_audit_total:,} — the IPO complexity multiplier "
            f"(1.15x) was not applied to the base fee"
        ),
        severity="material",
        which_test_cases_should_catch=["TC-16"],
    ))

    # -- Column widths
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 22
    ws.column_dimensions["D"].width = 15
    ws.column_dimensions["E"].width = 18
    ws.column_dimensions["F"].width = 22

    # -- Page setup for PDF rendering
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    # -- Save
    path = output_dir / _INPUT_DIR / "fee_schedule.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, path)

    manifest.register(
        f"{_INPUT_DIR}/fee_schedule.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=[_TC],
    )


def _write_engagement_template(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Copy templates/engagement_letter_template.docx into TC-16 input_files.

    The template already exists as a project-level template; we copy it to the
    TC-16 input directory and embed a canary for traceability.
    """
    repo_root = Path(__file__).resolve().parent.parent.parent
    src = repo_root / "templates" / "engagement_letter_template.docx"
    dst = output_dir / _INPUT_DIR / "engagement_letter_template.docx"
    dst.parent.mkdir(parents=True, exist_ok=True)

    canary_code = canaries.canary_for("tc16_engagement_template")
    location = _copy_docx_deterministic(src, dst, canary_code)
    canaries.set_location(
        "tc16_engagement_template",
        f"{_INPUT_DIR}/engagement_letter_template.docx",
        location,
    )

    manifest.register(
        f"{_INPUT_DIR}/engagement_letter_template.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Prompt and expected behavior ─────────────────────────────────────────────


def _write_prompt(output_dir: Path) -> None:
    """Write test_cases/TC-16/prompt.md."""
    text = """\
Generate a draft engagement letter for a combined audit and tax engagement
for Cascade Industries using the template provided.

1. Look up the appropriate fees from the fee schedule based on the client's
   revenue tier and entity count.
2. The audit fee should include the complexity multiplier for public-readiness
   (1.15x) since the company is considering an IPO.
3. The tax fee should cover federal and all state returns for the parent and
   3 subsidiaries.
4. Populate the template with the correct values.
5. Set the start date to March 15, 2026 and payment terms to Net 30.
6. Use "Sarah Chen" as the engagement partner.

Save as a Word document.
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write test_cases/TC-16/expected_behavior.md."""
    text = """\
# TC-16: Engagement Letter Generation — Expected Behavior

## Fee Calculation
- The agent must identify Cascade Industries' revenue tier ($100M–$500M)
  and entity count (4 entities → "4–6" range) from the client profile.
- Audit base fee: $285,000 (from fee schedule, $100M–$500M tier, 4–6 entities).
- IPO complexity multiplier: 1.15x (from complexity notes and prompt instruction).
- Audit fee: $285,000 × 1.15 = $327,750.
- Tax base fee: $95,000 (from fee schedule, $100M–$500M tier, 4–6 entities).
- Per-entity adder: $12,000 × 3 additional entities = $36,000.
- Tax fee: $95,000 + $36,000 = $131,000.
- Total engagement fee: $327,750 + $131,000 = $458,750.

## Merge Field Population
- `<<client_name>>` → "Cascade Industries, Inc."
- `<<engagement_scope>>` → Description covering combined audit and tax engagement,
  including IPO readiness procedures and multi-state tax returns.
- `<<fee_amount>>` → "$458,750" (or broken down as audit $327,750 + tax $131,000).
- `<<payment_terms>>` → "Net 30"
- `<<start_date>>` → "March 15, 2026"
- `<<partner_name>>` → "Sarah Chen"

## Template Preservation
- The agent must preserve the original formatting of the engagement letter template.
- All sections (Engagement Scope, Fees and Billing, Engagement Period,
  Responsibilities, Confidentiality, Limitation of Liability, Acceptance)
  should remain intact.
- The firm name "Mitchell & Partners LLP" and address should be preserved.

## Output Format
- Word document (.docx) that looks like a professional engagement letter.
- All merge fields replaced with actual values — no unreplaced placeholders.

## Common Mistakes to Watch For
- Using the wrong revenue tier or entity count range from the fee schedule.
- Forgetting to apply the 1.15x IPO complexity multiplier to the audit fee.
- Applying the per-entity adder incorrectly (should be × 3, not × 4).
- Not replacing all instances of merge fields (e.g., <<start_date>> appears twice).
- Altering the template structure or removing sections.
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold standard ────────────────────────────────────────────────────────────


def _fmt_dollars(v: Decimal | int) -> str:
    """Format a value as whole-dollar string with commas."""
    return f"${int(v):,}"


@register_gold("TC-16")
def _tc16_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model: Any,
) -> GoldStandard:
    """Gold standard for TC-16: Engagement Letter Generation."""
    return GoldStandard(
        test_case="TC-16",
        expected_outputs={
            "file_type": "docx",
            "fee_calculation": {
                "audit_base_fee": _fmt_dollars(_AUDIT_BASE_FEE),
                "ipo_complexity_multiplier": str(_IPO_COMPLEXITY_MULTIPLIER),
                "audit_fee": _fmt_dollars(_AUDIT_FEE),
                "tax_base_fee": _fmt_dollars(_TAX_BASE_FEE),
                "per_entity_adder": _fmt_dollars(_PER_ENTITY_TAX_ADDER),
                "additional_entities": _ADDITIONAL_ENTITIES,
                "tax_fee": _fmt_dollars(_TAX_FEE),
                "total_engagement_fee": _fmt_dollars(_TOTAL_FEE),
            },
            "merge_fields": {
                "client_name": _CLIENT_NAME,
                "engagement_scope": "Combined audit and tax engagement with IPO readiness procedures",
                "fee_amount": _fmt_dollars(_TOTAL_FEE),
                "payment_terms": _PAYMENT_TERMS,
                "start_date": _START_DATE,
                "partner_name": _PARTNER_NAME,
            },
            "template_sections_preserved": [
                "Engagement Scope",
                "Fees and Billing",
                "Engagement Period",
                "Responsibilities",
                "Confidentiality",
                "Limitation of Liability",
                "Acceptance",
            ],
        },
        canary_verification={
            "read_client_profile": canaries.canary_for("tc16_client_profile"),
            "read_fee_schedule": canaries.canary_for("tc16_fee_schedule"),
            "read_engagement_template": canaries.canary_for("tc16_engagement_template"),
        },
        error_detection={
            "ERR-013": (
                "Fee schedule audit total shows $285,000 instead of $327,750 — "
                "IPO complexity multiplier (1.15x) not applied"
            ),
        },
        scoring_hints={
            "correctness": (
                "Audit fee = $285,000 × 1.15 = $327,750; "
                "Tax fee = $95,000 + ($12,000 × 3) = $131,000; "
                "Total = $458,750"
            ),
            "completeness": (
                "All 6 merge fields populated; template sections preserved; "
                "fee breakdown visible in the letter"
            ),
            "format_compliance": (
                "Valid docx; professional engagement letter appearance; "
                "no unreplaced merge field placeholders"
            ),
            "robustness": (
                "Correctly identifies revenue tier and entity count from "
                "client profile; applies complexity multiplier only to audit fee"
            ),
            "communication": (
                "Engagement scope clearly describes combined audit + tax services; "
                "IPO readiness mentioned; fee amount clearly stated"
            ),
        },
    )


# ── Public entry point ──────────────────────────────────────────────────────


def emit_tc16(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
    **kwargs: object,
) -> None:
    """Write all TC-16 files to *output_dir*."""
    _write_client_profile(model, output_dir, canaries, manifest)
    _write_fee_schedule(output_dir, canaries, errors, manifest)
    _write_engagement_template(output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
