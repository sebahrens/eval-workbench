"""Formatter: TC-18 — Prior Year Workpaper Rollforward (Cross-service, Adversarial).

Emits:
- test_cases/TC-18/input_files/prior_year_workpapers/
    6 xlsx workpapers: revenue, expenses, balance_sheet, cash, fixed_assets, leases
    4 docx memos: planning, risk_assessment, summary, management_letter
    Each contains FY2024 data
- test_cases/TC-18/input_files/current_year_data/
    trial_balance_fy2025.csv  (format change from xlsx!)
    bank_statements_fy2025.csv
    lease_schedule_fy2025.xlsx  (with 2 new leases added)
    management_projections_fy2025.docx  (format change from xlsx!)
    goodwill_impairment_analysis.xlsx  (NEW — not in prior year)
- test_cases/TC-18/prompt.md
- test_cases/TC-18/expected_behavior.md
- gold_standards/TC-18_gold.json

No planted errors — adversarial signal comes from format changes,
new files, and judgment-call requirements.

Uses the canonical model — never hardcodes numbers.
"""

from __future__ import annotations

import csv
import datetime
import io
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import docx
import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.writer.excel import ExcelWriter

from generator.canaries import (
    CanaryRegistry,
    embed_canary_csv_comment,
    embed_canary_docx,
    embed_canary_xlsx,
)
from generator.errors import ErrorRegistry
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel
from generator.model.coa import ACCOUNTS_BY_NUMBER, AccountType
from generator.model.views import (
    build_balance_sheet,
    build_income_statement,
    consolidated_trial_balance_eliminated,
)

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-18"
_INPUT_DIR = f"test_cases/{_TC}/input_files"
_PY_DIR = f"{_INPUT_DIR}/prior_year_workpapers"
_CY_DIR = f"{_INPUT_DIR}/current_year_data"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)
_EOY_2024 = datetime.date(2024, 12, 31)
_EOY_2025 = datetime.date(2025, 12, 31)

# Canary file keys — 6 prior-year xlsx + 4 prior-year docx + 5 current-year files
_CANARY_KEYS: list[str] = sorted([
    "tc18_wp_revenue_fy2024",
    "tc18_wp_expenses_fy2024",
    "tc18_wp_balance_sheet_fy2024",
    "tc18_wp_cash_fy2024",
    "tc18_wp_fixed_assets_fy2024",
    "tc18_wp_leases_fy2024",
    "tc18_memo_planning_fy2024",
    "tc18_memo_risk_assessment_fy2024",
    "tc18_memo_summary_fy2024",
    "tc18_memo_management_letter_fy2024",
    "tc18_cy_trial_balance_fy2025",
    "tc18_cy_bank_statements_fy2025",
    "tc18_cy_lease_schedule_fy2025",
    "tc18_cy_mgmt_projections_fy2025",
    "tc18_cy_goodwill_impairment_fy2025",
])


# ── Helpers ──────────────────────────────────────────────────────────────────

def _pin_xlsx_dates(wb: openpyxl.Workbook) -> None:
    """Pin created timestamp for determinism."""
    wb.properties.created = _FIXED_DATETIME


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    path = Path(path)
    buf = io.BytesIO()
    wb.properties.modified = _FIXED_DATETIME
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    fixed_date_time = (2025, 3, 15, 9, 0, 0)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=fixed_date_time)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)

    fixed_date_time = (2025, 3, 15, 9, 0, 0)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=fixed_date_time)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _whole_dollars(d: Decimal | int) -> int:
    """Round a Decimal to whole dollars."""
    if isinstance(d, int):
        return d
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


_HEADER_FILL = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
_HEADER_FONT = Font(bold=True, color="FFFFFF", size=10)
_BOLD_FONT = Font(bold=True, size=10)
_NORMAL_FONT = Font(size=10)
_THIN_BORDER = Border(
    bottom=Side(style="thin"),
    top=Side(style="thin"),
    left=Side(style="thin"),
    right=Side(style="thin"),
)
_NUMBER_FMT = '#,##0'
_ACCT_FMT = Alignment(horizontal="left")


# ── Prior Year Workpapers (xlsx) ─────────────────────────────────────────────

def _write_wp_revenue(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> dict[str, Any]:
    """Write wp_revenue_fy2024.xlsx — revenue workpaper with FY2024 data."""
    tb = consolidated_trial_balance_eliminated(model.ledger, _EOY_2024)
    is_stmt = build_income_statement(model.ledger, 2024)

    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18_wp_revenue_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Revenue Analysis"

    # Header row
    headers = ["Account", "Account Name", "FY2024 Balance", "Notes"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.border = _THIN_BORDER

    # Revenue accounts from TB
    row = 2
    rev_data: dict[str, int] = {}
    for acct_num in sorted(tb.keys()):
        acct_info = ACCOUNTS_BY_NUMBER.get(acct_num)
        if acct_info and acct_info.account_type == AccountType.REVENUE:
            bal = _whole_dollars(-tb[acct_num])  # flip credit-normal to positive
            rev_data[acct_num] = bal
            ws.cell(row=row, column=1, value=acct_num).font = _NORMAL_FONT
            ws.cell(row=row, column=2, value=acct_info.name).font = _NORMAL_FONT
            ws.cell(row=row, column=3, value=bal).number_format = _NUMBER_FMT
            ws.cell(row=row, column=4, value="Per client TB").font = _NORMAL_FONT
            row += 1

    # Total row
    ws.cell(row=row, column=2, value="Total Revenue").font = _BOLD_FONT
    ws.cell(row=row, column=3, value=_whole_dollars(-is_stmt.total_revenue)).number_format = _NUMBER_FMT
    ws.cell(row=row, column=3).font = _BOLD_FONT

    # Commentary
    row += 2
    ws.cell(row=row, column=1, value="Prepared by: Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT
    conclusion = "Conclusion: Revenue agreed to TB and financial statements."
    ws.cell(row=row + 2, column=1, value=conclusion).font = _NORMAL_FONT

    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 25

    rel_path = f"{_PY_DIR}/wp_revenue_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])
    return rev_data


def _write_wp_expenses(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> dict[str, int]:
    """Write wp_expenses_fy2024.xlsx — expense analysis workpaper."""
    tb = consolidated_trial_balance_eliminated(model.ledger, _EOY_2024)
    is_stmt = build_income_statement(model.ledger, 2024)

    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18_wp_expenses_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Expense Analysis"

    headers = ["Account", "Account Name", "FY2024 Balance", "Variance Notes"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.border = _THIN_BORDER

    row = 2
    exp_data: dict[str, int] = {}
    for acct_num in sorted(tb.keys()):
        acct_info = ACCOUNTS_BY_NUMBER.get(acct_num)
        if acct_info and acct_info.account_type in (
            AccountType.COGS, AccountType.OPEX,
        ):
            bal = _whole_dollars(tb[acct_num])  # debit-normal, already positive
            exp_data[acct_num] = bal
            ws.cell(row=row, column=1, value=acct_num).font = _NORMAL_FONT
            ws.cell(row=row, column=2, value=acct_info.name).font = _NORMAL_FONT
            ws.cell(row=row, column=3, value=bal).number_format = _NUMBER_FMT
            ws.cell(row=row, column=4, value="").font = _NORMAL_FONT
            row += 1

    # Subtotals
    cogs_total = _whole_dollars(is_stmt.total_cogs)
    opex_total = _whole_dollars(is_stmt.total_opex)
    ws.cell(row=row, column=2, value="Total COGS").font = _BOLD_FONT
    ws.cell(row=row, column=3, value=cogs_total).number_format = _NUMBER_FMT
    row += 1
    ws.cell(row=row, column=2, value="Total OpEx").font = _BOLD_FONT
    ws.cell(row=row, column=3, value=opex_total).number_format = _NUMBER_FMT

    row += 2
    ws.cell(row=row, column=1, value="Prepared by: Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT

    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 30

    rel_path = f"{_PY_DIR}/wp_expenses_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])
    return exp_data


def _write_wp_balance_sheet(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> dict[str, int]:
    """Write wp_balance_sheet_fy2024.xlsx — balance sheet workpaper."""
    tb = consolidated_trial_balance_eliminated(model.ledger, _EOY_2024)
    bs = build_balance_sheet(model.ledger, _EOY_2024)

    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18_wp_balance_sheet_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Balance Sheet"

    headers = ["Account", "Account Name", "FY2024 Balance", "Classification"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.border = _THIN_BORDER

    row = 2
    bs_data: dict[str, int] = {}
    for acct_num in sorted(tb.keys()):
        acct_info = ACCOUNTS_BY_NUMBER.get(acct_num)
        if acct_info and acct_info.account_type in (
            AccountType.ASSET, AccountType.LIABILITY, AccountType.EQUITY,
        ):
            bal = _whole_dollars(tb[acct_num])
            bs_data[acct_num] = bal
            ws.cell(row=row, column=1, value=acct_num).font = _NORMAL_FONT
            ws.cell(row=row, column=2, value=acct_info.name).font = _NORMAL_FONT
            ws.cell(row=row, column=3, value=bal).number_format = _NUMBER_FMT
            ws.cell(row=row, column=4, value=acct_info.account_type.value).font = _NORMAL_FONT
            row += 1

    # Summary
    ws.cell(row=row, column=2, value="Total Assets").font = _BOLD_FONT
    ws.cell(row=row, column=3, value=_whole_dollars(bs.total_assets)).number_format = _NUMBER_FMT
    row += 1
    ws.cell(row=row, column=2, value="Total Liabilities + Equity").font = _BOLD_FONT
    ws.cell(row=row, column=3, value=_whole_dollars(bs.total_liabilities + bs.total_equity)).number_format = _NUMBER_FMT

    row += 2
    ws.cell(row=row, column=1, value="Prepared by: Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT

    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 20

    rel_path = f"{_PY_DIR}/wp_balance_sheet_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])
    return bs_data


def _write_wp_cash(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> dict[str, int]:
    """Write wp_cash_fy2024.xlsx — cash & bank reconciliation workpaper."""
    tb = consolidated_trial_balance_eliminated(model.ledger, _EOY_2024)

    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18_wp_cash_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Cash Reconciliation"

    headers = ["Account", "Account Name", "FY2024 Balance", "Bank Confirmed"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.border = _THIN_BORDER

    row = 2
    cash_data: dict[str, int] = {}
    # Cash accounts are 1000-1099
    for acct_num in sorted(tb.keys()):
        if acct_num.startswith("10"):
            acct_info = ACCOUNTS_BY_NUMBER.get(acct_num)
            if acct_info:
                bal = _whole_dollars(tb[acct_num])
                cash_data[acct_num] = bal
                ws.cell(row=row, column=1, value=acct_num).font = _NORMAL_FONT
                ws.cell(row=row, column=2, value=acct_info.name).font = _NORMAL_FONT
                ws.cell(row=row, column=3, value=bal).number_format = _NUMBER_FMT
                ws.cell(row=row, column=4, value="Yes").font = _NORMAL_FONT
                row += 1

    row += 1
    ws.cell(row=row, column=1, value="Prepared by: Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT
    conclusion = "Conclusion: Cash balances confirmed with bank statements."
    ws.cell(row=row + 2, column=1, value=conclusion).font = _NORMAL_FONT

    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 18

    rel_path = f"{_PY_DIR}/wp_cash_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])
    return cash_data


def _write_wp_fixed_assets(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> dict[str, int]:
    """Write wp_fixed_assets_fy2024.xlsx — fixed assets/depreciation workpaper."""
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18_wp_fixed_assets_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Fixed Assets"

    headers = ["Asset ID", "Description", "Entity", "Cost", "Accum Depr (FY2024)", "NBV"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.border = _THIN_BORDER

    row = 2
    asset_data: list[dict[str, Any]] = []
    for asset in sorted(model.assets, key=lambda a: a.asset_id):
        # Compute accumulated depreciation through FY2024
        accum_dec = sum(
            asset.book_depr_for_year(y)
            for y in range(asset.acquisition_date.year, 2025)
        )
        accum = _whole_dollars(accum_dec)
        cost = _whole_dollars(asset.cost)
        nbv = cost - accum

        ws.cell(row=row, column=1, value=asset.asset_id).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=asset.description).font = _NORMAL_FONT
        ws.cell(row=row, column=3, value=asset.entity_code).font = _NORMAL_FONT
        ws.cell(row=row, column=4, value=cost).number_format = _NUMBER_FMT
        ws.cell(row=row, column=5, value=accum).number_format = _NUMBER_FMT
        ws.cell(row=row, column=6, value=nbv).number_format = _NUMBER_FMT
        asset_data.append({
            "asset_id": asset.asset_id,
            "cost": cost,
            "accum_depr": accum,
            "nbv": nbv,
        })
        row += 1

    row += 1
    ws.cell(row=row, column=1, value="Prepared by: Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT

    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 10
    ws.column_dimensions["D"].width = 15
    ws.column_dimensions["E"].width = 18
    ws.column_dimensions["F"].width = 15

    rel_path = f"{_PY_DIR}/wp_fixed_assets_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])
    return {a["asset_id"]: a["nbv"] for a in asset_data}


def _write_wp_leases(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> list[str]:
    """Write wp_leases_fy2024.xlsx — lease schedule workpaper.

    Returns list of lease IDs present in FY2024.
    """
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18_wp_leases_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Lease Schedule"

    headers = ["Lease ID", "Description", "Entity", "Type", "Monthly Payment",
               "Start Date", "Term (months)", "ROU Asset", "Lease Liability"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.border = _THIN_BORDER

    row = 2
    # Only include leases that started on or before 2024-12-31
    fy2024_leases: list[str] = []
    fy2024_schedules = [s for s in model.lease_schedules if s.year == 2024]
    schedule_by_lease = {s.lease_id: s for s in fy2024_schedules}

    for lease in sorted(model.leases, key=lambda le: le.lease_id):
        if lease.commencement_date.year > 2024:
            continue
        fy2024_leases.append(lease.lease_id)
        sched = schedule_by_lease.get(lease.lease_id)

        ws.cell(row=row, column=1, value=lease.lease_id).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=lease.description).font = _NORMAL_FONT
        ws.cell(row=row, column=3, value=lease.entity_code).font = _NORMAL_FONT
        ws.cell(row=row, column=4, value=lease.lease_type.value).font = _NORMAL_FONT
        ws.cell(row=row, column=5, value=_whole_dollars(lease.monthly_base_rent)).number_format = _NUMBER_FMT
        ws.cell(row=row, column=6, value=str(lease.commencement_date)).font = _NORMAL_FONT
        ws.cell(row=row, column=7, value=lease.term_months).font = _NORMAL_FONT
        if sched:
            ws.cell(row=row, column=8, value=_whole_dollars(sched.rou_asset_end)).number_format = _NUMBER_FMT
            ws.cell(row=row, column=9, value=_whole_dollars(sched.lease_liability_end)).number_format = _NUMBER_FMT
        row += 1

    row += 1
    ws.cell(row=row, column=1, value="Prepared by: Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT

    for i, w in enumerate([12, 30, 10, 12, 15, 14, 14, 15, 15], 1):
        ws.column_dimensions[chr(64 + i)].width = w

    rel_path = f"{_PY_DIR}/wp_leases_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])
    return fy2024_leases


# ── Prior Year Workpapers (docx memos) ──────────────────────────────────────

def _write_memo_planning(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write memo_planning_fy2024.docx — audit planning memo.

    This is one of the 2 files that should be FLAGGED for manager attention
    (requires substantive rewriting, not just data updates).
    """
    is_stmt = build_income_statement(model.ledger, 2024)
    total_rev = _whole_dollars(-is_stmt.total_revenue)

    doc = docx.Document()
    key = "tc18_memo_planning_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_docx(doc, canary)

    doc.add_heading("Cascade Industries, Inc.", level=1)
    doc.add_heading("Audit Planning Memorandum — FY2024", level=2)

    doc.add_paragraph(
        "This memorandum documents the planned audit approach for the fiscal year "
        "ended December 31, 2024. The audit team has assessed the key risk areas "
        "and determined the scope of substantive testing required."
    )

    doc.add_heading("1. Client Overview", level=3)
    doc.add_paragraph(
        f"Cascade Industries is a mid-market manufacturer headquartered in Portland, OR "
        f"with consolidated revenue of approximately ${total_rev:,} for FY2024. "
        f"The company operates through three subsidiaries: Precision Components (Portland, OR), "
        f"Advanced Materials (Austin, TX), and Distribution Services (Chicago, IL)."
    )

    doc.add_heading("2. Key Risk Areas", level=3)
    doc.add_paragraph("• Revenue recognition — multi-element arrangements with custom machining")
    doc.add_paragraph("• Intercompany eliminations — material IC transactions between subsidiaries")
    doc.add_paragraph("• Lease accounting (ASC 842) — portfolio of operating and finance leases")
    doc.add_paragraph("• Inventory valuation — raw materials and WIP at Precision Components")

    doc.add_heading("3. Materiality", level=3)
    materiality = _whole_dollars(Decimal(str(total_rev)) * Decimal("0.01"))
    doc.add_paragraph(
        f"Planning materiality set at 1% of revenue = ${materiality:,}. "
        f"Performance materiality at 75% = ${_whole_dollars(Decimal(str(materiality)) * Decimal('0.75')):,}."
    )

    doc.add_heading("4. Audit Timeline", level=3)
    doc.add_paragraph("• Interim procedures: October 2024")
    doc.add_paragraph("• Year-end fieldwork: February–March 2025")
    doc.add_paragraph("• Report issuance: April 2025")

    doc.add_heading("5. Team Assignment", level=3)
    doc.add_paragraph("• Engagement Partner: J. Thompson")
    doc.add_paragraph("• Senior Manager: R. Patel")
    doc.add_paragraph("• Staff: To be assigned")

    p = doc.add_paragraph()
    p.add_run("\nPrepared: March 2025").italic = True

    rel_path = f"{_PY_DIR}/memo_planning_fy2024.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, full_path)

    canaries.set_location(key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary, test_cases=[_TC])


def _write_memo_risk_assessment(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write memo_risk_assessment_fy2024.docx — risk assessment memo."""
    doc = docx.Document()
    key = "tc18_memo_risk_assessment_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_docx(doc, canary)

    doc.add_heading("Cascade Industries, Inc.", level=1)
    doc.add_heading("Risk Assessment — FY2024 Audit", level=2)

    doc.add_heading("1. Fraud Risk Assessment", level=3)
    doc.add_paragraph(
        "Management override of controls assessed as present in all engagements. "
        "No specific fraud risk indicators identified beyond routine presumption."
    )

    doc.add_heading("2. Significant Risks", level=3)
    doc.add_paragraph("• Revenue recognition: Risk that revenue is recorded in the wrong period")
    doc.add_paragraph("• Related party transactions: Intercompany transactions require careful review")
    doc.add_paragraph("• Management estimates: Allowance for doubtful accounts, depreciation useful lives")

    doc.add_heading("3. Industry Risk Factors", level=3)
    doc.add_paragraph(
        "Manufacturing sector faces supply chain constraints and raw material price volatility. "
        "Customer concentration risk moderate — top 10 customers represent ~45% of revenue."
    )

    doc.add_heading("4. IT General Controls", level=3)
    doc.add_paragraph(
        "ERP system: SAP S/4HANA. Last IT audit performed in Q2 2024 with no significant findings. "
        "Logical access controls and change management procedures in place."
    )

    p = doc.add_paragraph()
    p.add_run("\nPrepared: March 2025").italic = True

    rel_path = f"{_PY_DIR}/memo_risk_assessment_fy2024.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, full_path)

    canaries.set_location(key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary, test_cases=[_TC])


def _write_memo_summary(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write memo_summary_fy2024.docx — audit summary memo."""
    is_stmt = build_income_statement(model.ledger, 2024)

    doc = docx.Document()
    key = "tc18_memo_summary_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_docx(doc, canary)

    doc.add_heading("Cascade Industries, Inc.", level=1)
    doc.add_heading("Audit Summary Memorandum — FY2024", level=2)

    doc.add_paragraph(
        "This memorandum summarizes the results of audit procedures performed for "
        "the fiscal year ended December 31, 2024."
    )

    doc.add_heading("1. Financial Highlights", level=3)
    doc.add_paragraph(f"• Revenue: ${_whole_dollars(-is_stmt.total_revenue):,}")
    doc.add_paragraph(f"• Net Income: ${_whole_dollars(is_stmt.net_income):,}")
    doc.add_paragraph(f"• Total Assets: ${_whole_dollars(build_balance_sheet(model.ledger, _EOY_2024).total_assets):,}")

    doc.add_heading("2. Audit Findings", level=3)
    doc.add_paragraph(
        "No material misstatements identified. One immaterial adjustment proposed "
        "for reclassification of prepaid expenses ($12,500). Management accepted all adjustments."
    )

    doc.add_heading("3. Going Concern Assessment", level=3)
    doc.add_paragraph(
        "No indicators of going concern identified. The company maintains adequate liquidity "
        "and positive operating cash flows."
    )

    doc.add_heading("4. Subsequent Events", level=3)
    doc.add_paragraph(
        "No significant subsequent events identified through the date of the auditor's report."
    )

    p = doc.add_paragraph()
    p.add_run("\nPrepared: March 2025").italic = True

    rel_path = f"{_PY_DIR}/memo_summary_fy2024.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, full_path)

    canaries.set_location(key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary, test_cases=[_TC])


def _write_memo_management_letter(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write memo_management_letter_fy2024.docx — management letter points."""
    doc = docx.Document()
    key = "tc18_memo_management_letter_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_docx(doc, canary)

    doc.add_heading("Cascade Industries, Inc.", level=1)
    doc.add_heading("Management Letter — FY2024", level=2)

    doc.add_paragraph(
        "During our audit of the financial statements for the year ended December 31, 2024, "
        "we noted certain matters involving internal control that we wish to bring to your attention."
    )

    doc.add_heading("Finding 1: Segregation of Duties — AP", level=3)
    doc.add_paragraph(
        "The accounts payable clerk has the ability to both enter invoices and approve payments "
        "up to $10,000. We recommend implementing dual authorization for all payments regardless of amount."
    )
    doc.add_paragraph("Management Response: Will implement in Q2 2025.", style="Intense Quote")

    doc.add_heading("Finding 2: Journal Entry Review", level=3)
    doc.add_paragraph(
        "Manual journal entries above $50,000 do not consistently require a second approval. "
        "We recommend a mandatory review process for all manual entries above this threshold."
    )
    doc.add_paragraph("Management Response: Agreed. Policy update in progress.", style="Intense Quote")

    doc.add_heading("Finding 3: Fixed Asset Physical Inventory", level=3)
    doc.add_paragraph(
        "Physical inventory of fixed assets has not been performed since 2022. "
        "We recommend conducting a physical count at least every two years to verify existence."
    )
    doc.add_paragraph("Management Response: Physical count scheduled for Q3 2025.", style="Intense Quote")

    p = doc.add_paragraph()
    p.add_run("\nPrepared: March 2025").italic = True

    rel_path = f"{_PY_DIR}/memo_management_letter_fy2024.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, full_path)

    canaries.set_location(key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary, test_cases=[_TC])


# ── Current Year Data ────────────────────────────────────────────────────────

def _write_cy_trial_balance(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> dict[str, int]:
    """Write trial_balance_fy2025.csv — FORMAT CHANGE from xlsx.

    This is the key adversarial element: the prior year TB was an xlsx workpaper,
    but the client now provides the current year TB as CSV.
    """
    tb = consolidated_trial_balance_eliminated(model.ledger, _EOY_2025)
    key = "tc18_cy_trial_balance_fy2025"
    canary = canaries.canary_for(key)

    rel_path = f"{_CY_DIR}/trial_balance_fy2025.csv"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    tb_data: dict[str, int] = {}
    with open(full_path, "w", newline="") as f:
        f.write(embed_canary_csv_comment(canary))
        writer = csv.writer(f)
        writer.writerow(["Account", "Account Name", "Debit", "Credit"])

        for acct_num in sorted(tb.keys()):
            acct_info = ACCOUNTS_BY_NUMBER.get(acct_num)
            if acct_info is None:
                continue
            bal = tb[acct_num]
            val = _whole_dollars(bal)
            tb_data[acct_num] = val
            if bal >= 0:
                writer.writerow([acct_num, acct_info.name, val, ""])
            else:
                writer.writerow([acct_num, acct_info.name, "", abs(val)])

    canaries.set_location(key, rel_path, "CSV comment line 1")
    manifest.register(rel_path, "csv", canary=canary, test_cases=[_TC])
    return tb_data


def _write_cy_bank_statements(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write bank_statements_fy2025.csv — current year bank statement data."""
    key = "tc18_cy_bank_statements_fy2025"
    canary = canaries.canary_for(key)

    rel_path = f"{_CY_DIR}/bank_statements_fy2025.csv"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    # Pull cash account balances from 2025 TB
    tb = consolidated_trial_balance_eliminated(model.ledger, _EOY_2025)

    with open(full_path, "w", newline="") as f:
        f.write(embed_canary_csv_comment(canary))
        writer = csv.writer(f)
        writer.writerow(["Account", "Account Name", "Statement Balance", "Date"])

        for acct_num in sorted(tb.keys()):
            if acct_num.startswith("10"):
                acct_info = ACCOUNTS_BY_NUMBER.get(acct_num)
                if acct_info:
                    writer.writerow([
                        acct_num,
                        acct_info.name,
                        _whole_dollars(tb[acct_num]),
                        "2025-12-31",
                    ])

    canaries.set_location(key, rel_path, "CSV comment line 1")
    manifest.register(rel_path, "csv", canary=canary, test_cases=[_TC])


def _write_cy_lease_schedule(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
    fy2024_lease_ids: list[str],
) -> list[str]:
    """Write lease_schedule_fy2025.xlsx — with 2 new leases added.

    Returns list of new lease IDs (leases in FY2025 but not in FY2024 workpaper).
    """
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18_cy_lease_schedule_fy2025"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Lease Schedule FY2025"

    headers = ["Lease ID", "Description", "Entity", "Type", "Monthly Payment",
               "Start Date", "Term (months)", "ROU Asset", "Lease Liability"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.border = _THIN_BORDER

    row = 2
    fy2025_schedules = [s for s in model.lease_schedules if s.year == 2025]
    schedule_by_lease = {s.lease_id: s for s in fy2025_schedules}
    new_leases: list[str] = []

    for lease in sorted(model.leases, key=lambda le: le.lease_id):
        # Include all leases active in FY2025
        sched = schedule_by_lease.get(lease.lease_id)
        if sched is None:
            continue

        if lease.lease_id not in fy2024_lease_ids:
            new_leases.append(lease.lease_id)

        ws.cell(row=row, column=1, value=lease.lease_id).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=lease.description).font = _NORMAL_FONT
        ws.cell(row=row, column=3, value=lease.entity_code).font = _NORMAL_FONT
        ws.cell(row=row, column=4, value=lease.lease_type.value).font = _NORMAL_FONT
        ws.cell(row=row, column=5, value=_whole_dollars(lease.monthly_base_rent)).number_format = _NUMBER_FMT
        ws.cell(row=row, column=6, value=str(lease.commencement_date)).font = _NORMAL_FONT
        ws.cell(row=row, column=7, value=lease.term_months).font = _NORMAL_FONT
        ws.cell(row=row, column=8, value=_whole_dollars(sched.rou_asset_end)).number_format = _NUMBER_FMT
        ws.cell(row=row, column=9, value=_whole_dollars(sched.lease_liability_end)).number_format = _NUMBER_FMT
        row += 1

    for i, w in enumerate([12, 30, 10, 12, 15, 14, 14, 15, 15], 1):
        ws.column_dimensions[chr(64 + i)].width = w

    rel_path = f"{_CY_DIR}/lease_schedule_fy2025.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])
    return new_leases


def _write_cy_mgmt_projections(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write management_projections_fy2025.docx — FORMAT CHANGE from xlsx.

    This is the second adversarial format change: prior year projections were in
    xlsx, but management now provides them as a docx narrative. This is one of
    the 2 files the agent should FLAG for manager attention.
    """
    is_2025 = build_income_statement(model.ledger, 2025)

    doc = docx.Document()
    key = "tc18_cy_mgmt_projections_fy2025"
    canary = canaries.canary_for(key)
    embed_canary_docx(doc, canary)

    doc.add_heading("Cascade Industries, Inc.", level=1)
    doc.add_heading("Management Financial Projections — FY2026", level=2)

    doc.add_paragraph(
        "The following projections reflect management's expectations for the fiscal year "
        "ending December 31, 2026, based on current market conditions and strategic initiatives."
    )

    total_rev = _whole_dollars(-is_2025.total_revenue)
    projected_rev = _whole_dollars(Decimal(str(total_rev)) * Decimal("1.08"))

    doc.add_heading("Revenue Projections", level=3)
    doc.add_paragraph(
        f"FY2025 Actual Revenue: ${total_rev:,}\n"
        f"FY2026 Projected Revenue: ${projected_rev:,} (8% growth)\n\n"
        "Growth drivers:\n"
        "• Advanced Composites: Expected 12% growth from new aerospace contracts\n"
        "• Custom Machining: Stable at 5% growth\n"
        "• Warehousing Services: Projected 15% growth from new Chicago facility expansion"
    )

    doc.add_heading("Expense Projections", level=3)
    doc.add_paragraph(
        "Operating expenses expected to increase 6% due to:\n"
        "• Headcount additions (50 new hires across all subsidiaries)\n"
        "• Raw material cost inflation of ~3%\n"
        "• New lease commitments for expanded warehouse space"
    )

    doc.add_heading("Capital Expenditures", level=3)
    doc.add_paragraph(
        "Planned CapEx of $8.5M:\n"
        "• CNC machinery upgrade at Precision Components: $4.2M\n"
        "• Lab equipment at Advanced Materials: $2.8M\n"
        "• Warehouse automation at Distribution Services: $1.5M"
    )

    doc.add_heading("Key Assumptions", level=3)
    doc.add_paragraph(
        "• No significant changes in customer base\n"
        "• Interest rates stable at current levels\n"
        "• No major acquisitions planned\n"
        "• Goodwill impairment testing scheduled for Q4 2026"
    )

    p = doc.add_paragraph()
    p.add_run("\nPrepared by: CFO Office, January 2026").italic = True

    rel_path = f"{_CY_DIR}/management_projections_fy2025.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, full_path)

    canaries.set_location(key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary, test_cases=[_TC])


def _write_cy_goodwill_impairment(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> dict[str, Any]:
    """Write goodwill_impairment_analysis.xlsx — NEW file not in prior year.

    The agent should flag this as a significant new audit area.
    """
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18_cy_goodwill_impairment_fy2025"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Goodwill Impairment"

    headers = ["Reporting Unit", "Carrying Amount", "Fair Value Estimate",
               "Impairment Indicated", "Notes"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.border = _THIN_BORDER

    # Goodwill from a prior acquisition of Advanced Materials
    goodwill_data = [
        {
            "unit": "Cascade Advanced Materials",
            "carrying": 12500000,
            "fair_value": 14200000,
            "impaired": "No",
            "notes": "Fair value exceeds carrying by 13.6%",
        },
        {
            "unit": "Cascade Distribution Services",
            "carrying": 3800000,
            "fair_value": 3950000,
            "impaired": "No",
            "notes": "Fair value exceeds carrying by 3.9% — close to threshold",
        },
    ]

    for i, item in enumerate(goodwill_data, 2):
        ws.cell(row=i, column=1, value=item["unit"]).font = _NORMAL_FONT
        ws.cell(row=i, column=2, value=item["carrying"]).number_format = _NUMBER_FMT
        ws.cell(row=i, column=3, value=item["fair_value"]).number_format = _NUMBER_FMT
        ws.cell(row=i, column=4, value=item["impaired"]).font = _NORMAL_FONT
        ws.cell(row=i, column=5, value=item["notes"]).font = _NORMAL_FONT

    row = len(goodwill_data) + 3
    ws.cell(row=row, column=1, value="Total Goodwill").font = _BOLD_FONT
    ws.cell(row=row, column=2, value=sum(d["carrying"] for d in goodwill_data)).number_format = _NUMBER_FMT
    ws.cell(row=row, column=2).font = _BOLD_FONT

    row += 2
    ws.cell(row=row, column=1, value="Methodology: Discounted Cash Flow (DCF)").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Discount rate: 10.5% WACC").font = _NORMAL_FONT
    ws.cell(row=row + 2, column=1, value="Terminal growth rate: 2.5%").font = _NORMAL_FONT
    ws.cell(row=row + 3, column=1, value="Prepared by: Management, December 2025").font = _NORMAL_FONT

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 18
    ws.column_dimensions["E"].width = 40

    rel_path = f"{_CY_DIR}/goodwill_impairment_analysis.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])
    return {
        "reporting_units": goodwill_data,
        "total_goodwill": sum(d["carrying"] for d in goodwill_data),
    }


# ── Prompt & expected behavior ───────────────────────────────────────────────

def _write_prompt(output_dir: Path) -> None:
    """Write the TC-18 prompt.md."""
    text = """\
# TC-18: Prior Year Workpaper Rollforward

## Input Files

### Prior Year Workpapers (`prior_year_workpapers/`)
- `wp_revenue_fy2024.xlsx` — Revenue analysis workpaper
- `wp_expenses_fy2024.xlsx` — Expense analysis workpaper
- `wp_balance_sheet_fy2024.xlsx` — Balance sheet workpaper
- `wp_cash_fy2024.xlsx` — Cash & bank reconciliation workpaper
- `wp_fixed_assets_fy2024.xlsx` — Fixed assets/depreciation workpaper
- `wp_leases_fy2024.xlsx` — Lease schedule workpaper (ASC 842)
- `memo_planning_fy2024.docx` — Audit planning memorandum
- `memo_risk_assessment_fy2024.docx` — Risk assessment memorandum
- `memo_summary_fy2024.docx` — Audit summary memorandum
- `memo_management_letter_fy2024.docx` — Management letter points

### Current Year Data (`current_year_data/`)
- `trial_balance_fy2025.csv` — **Note: format changed from xlsx to CSV**
- `bank_statements_fy2025.csv` — Bank statement data
- `lease_schedule_fy2025.xlsx` — Updated lease schedule (includes 2 new leases)
- `management_projections_fy2025.docx` — **Note: format changed from xlsx to docx**
- `goodwill_impairment_analysis.xlsx` — **New file not present in prior year**

## Instructions

Roll forward the FY2024 audit workpapers to FY2025 using the current year data.

1. For each prior year workpaper, identify which data needs to be updated.
2. Map the current year data files to the corresponding prior year workpapers,
   noting any format changes (e.g., CSV vs. xlsx).
3. Update all numerical data with current year figures.
4. Preserve the workpaper structure and any formulas/commentary that are still relevant.
5. Flag the following for manager attention:
   - Any structural changes in the client's data (new accounts, format changes,
     renamed fields)
   - The new goodwill impairment analysis file (not present in prior year —
     suggests a new accounting issue)
   - Any areas where prior year commentary may no longer be applicable
6. Update the planning memo with current year scope considerations.

Export the rolled-forward workpapers to a new folder.
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write the TC-18 expected_behavior.md."""
    text = """\
# TC-18: Expected Behavior

## Key Evaluation Criteria

### Workpaper Update Success (8 of 10)
The agent should successfully update 8 of 10 prior year workpapers:
- **Revenue workpaper** — Update with FY2025 revenue data from CSV trial balance
- **Expense workpaper** — Update with FY2025 expense data from CSV trial balance
- **Balance sheet workpaper** — Update with FY2025 balances from CSV trial balance
- **Cash workpaper** — Update with FY2025 bank statement data
- **Fixed assets workpaper** — Update with FY2025 depreciation data
- **Lease workpaper** — Update with FY2025 lease schedule (including 2 new leases)
- **Risk assessment memo** — Mechanical updates (dates, figures)
- **Management letter memo** — Update status of prior findings

### Files Requiring Manager Judgment (2 of 10)
- **Management projections** — Format changed from xlsx to docx; the prior year
  workpaper references specific cells that no longer exist in the new format.
  Agent should FLAG this for manager attention rather than attempt a mechanical update.
- **Planning memo** — Requires substantive rewriting to reflect FY2025 scope
  (new goodwill impairment area, updated materiality, changed risk landscape).
  A simple date/number swap is insufficient. Agent should FLAG this.

### Format Change Handling
- The CSV trial balance should be handled transparently — the agent should
  parse CSV and map data to the xlsx workpaper structure.
- The docx projections format change should be explicitly flagged.

### New Audit Area Detection
- The goodwill impairment analysis file was not present in prior year.
- The agent must flag this as a significant new audit area requiring
  new workpapers and procedures.
- Distribution Services unit is close to impairment threshold (3.9% cushion).

### Audit Summary Memo
- FY2024 financial highlights should be updated to FY2025 figures.
- Going concern and subsequent events sections should note they need
  current-year evaluation.
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold Standard ────────────────────────────────────────────────────────────

@register_gold("TC-18")
def _tc18_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """Build the TC-18 gold standard from the canonical model."""
    model: CascadeModel = model_kwargs["model"]

    is_2024 = build_income_statement(model.ledger, 2024)
    is_2025 = build_income_statement(model.ledger, 2025)
    bs_2024 = build_balance_sheet(model.ledger, _EOY_2024)
    bs_2025 = build_balance_sheet(model.ledger, _EOY_2025)

    # Identify new leases (in FY2025 schedule but not FY2024)
    fy2024_schedules = {s.lease_id for s in model.lease_schedules if s.year == 2024}
    fy2025_schedules = {s.lease_id for s in model.lease_schedules if s.year == 2025}
    new_lease_ids = sorted(fy2025_schedules - fy2024_schedules)

    # Canary verification — one per file
    canary_verification: dict[str, str] = {}
    for key in _CANARY_KEYS:
        label = f"read_{key.replace('tc18_', '')}"
        canary_verification[label] = canaries.canary_for(key)

    return GoldStandard(
        test_case=_TC,
        expected_outputs={
            "files_to_update": {
                "mechanical_updates": [
                    "wp_revenue_fy2024.xlsx",
                    "wp_expenses_fy2024.xlsx",
                    "wp_balance_sheet_fy2024.xlsx",
                    "wp_cash_fy2024.xlsx",
                    "wp_fixed_assets_fy2024.xlsx",
                    "wp_leases_fy2024.xlsx",
                    "memo_risk_assessment_fy2024.docx",
                    "memo_management_letter_fy2024.docx",
                ],
                "requires_manager_judgment": [
                    "memo_planning_fy2024.docx",
                    "management_projections_fy2025.docx",
                ],
            },
            "fy2024_financials": {
                "revenue": _whole_dollars(-is_2024.total_revenue),
                "net_income": _whole_dollars(is_2024.net_income),
                "total_assets": _whole_dollars(bs_2024.total_assets),
            },
            "fy2025_financials": {
                "revenue": _whole_dollars(-is_2025.total_revenue),
                "net_income": _whole_dollars(is_2025.net_income),
                "total_assets": _whole_dollars(bs_2025.total_assets),
            },
            "format_changes": [
                {
                    "file": "trial_balance_fy2025.csv",
                    "change": "Format changed from xlsx to CSV",
                    "handling": "Parse CSV and map to workpaper structure",
                },
                {
                    "file": "management_projections_fy2025.docx",
                    "change": "Format changed from xlsx to docx narrative",
                    "handling": "Flag for manager — cell references no longer apply",
                },
            ],
            "new_audit_area": {
                "file": "goodwill_impairment_analysis.xlsx",
                "description": "Goodwill impairment analysis not present in prior year",
                "significance": "New audit area requiring new workpapers and procedures",
                "total_goodwill": 16300000,
                "concern": "Distribution Services unit has only 3.9% cushion above carrying amount",
            },
            "new_leases": new_lease_ids,
            "planning_memo_flag": (
                "Planning memo requires substantive rewriting for FY2025: "
                "new goodwill impairment area, updated materiality based on "
                "FY2025 revenue, and changed risk landscape. "
                "A simple date/number replacement is insufficient."
            ),
        },
        canary_verification=canary_verification,
        error_detection={},
        scoring_hints={
            "correctness": (
                "8 of 10 workpapers updated with correct FY2025 data; "
                "2 correctly flagged for manager judgment; CSV format "
                "change handled transparently"
            ),
            "completeness": (
                "All 10 workpapers addressed (8 updated + 2 flagged); "
                "goodwill impairment file acknowledged as new audit area; "
                "all structural changes documented"
            ),
            "format_compliance": (
                "Workpaper structure preserved; formulas intact where "
                "still relevant; commentary updated; file naming consistent"
            ),
            "robustness": (
                "CSV format change handled transparently; new lease entries "
                "incorporated; stale commentary identified; current year data "
                "mapped to prior year workpaper fields despite format differences"
            ),
            "communication": (
                "All structural changes flagged (format changes, new goodwill file); "
                "clear distinction between mechanical updates and judgment calls; "
                "goodwill impairment highlighted as significant new audit area"
            ),
        },
    )


# ── Public entry point ──────────────────────────────────────────────────────

def emit_tc18(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write all TC-18 files to *output_dir*."""
    # Prior year workpapers (xlsx)
    _write_wp_revenue(model, output_dir, canaries, manifest)
    _write_wp_expenses(model, output_dir, canaries, manifest)
    _write_wp_balance_sheet(model, output_dir, canaries, manifest)
    _write_wp_cash(model, output_dir, canaries, manifest)
    _write_wp_fixed_assets(model, output_dir, canaries, manifest)
    fy2024_lease_ids = _write_wp_leases(model, output_dir, canaries, manifest)

    # Prior year workpapers (docx memos)
    _write_memo_planning(model, output_dir, canaries, manifest)
    _write_memo_risk_assessment(model, output_dir, canaries, manifest)
    _write_memo_summary(model, output_dir, canaries, manifest)
    _write_memo_management_letter(output_dir, canaries, manifest)

    # Current year data
    _write_cy_trial_balance(model, output_dir, canaries, manifest)
    _write_cy_bank_statements(model, output_dir, canaries, manifest)
    _write_cy_lease_schedule(model, output_dir, canaries, manifest, fy2024_lease_ids)
    _write_cy_mgmt_projections(model, output_dir, canaries, manifest)
    _write_cy_goodwill_impairment(model, output_dir, canaries, manifest)

    # Prompt and expected behavior
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
