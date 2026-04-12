"""Formatter: TC-06-EU — IAS 12 Income Tax Provision (Tax, Complex).

Emits:
- test_cases/TC-06-EU/input_files/tc06eu_consolidated_tb_fy2025.xlsx
  Consolidated trial balance (EUR) for Cascade Europe group by entity
- test_cases/TC-06-EU/input_files/tc06eu_tax_provision_fy2024_workpaper.xlsx
  Prior year provision workpaper (current by jurisdiction, deferred rollforward,
  ETR reconciliation, temporary differences)
- test_cases/TC-06-EU/input_files/tc06eu_permanent_temporary_differences_fy2025.docx
  FY2025 book-tax differences by entity with EUR formatting
- test_cases/TC-06-EU/input_files/tc06eu_statutory_rates.docx
  Statutory rates NL/DE/FR/UK, R&D incentives, Pillar Two note
- test_cases/TC-06-EU/prompt.md
- test_cases/TC-06-EU/expected_behavior.md
- gold_standards/TC-06-EU_gold.json

Planted errors:
  ERR-EU-002: stale_data — FY2024 workpaper uses Munich Gewerbesteuer Hebesatz
              480% (old) instead of 490% (current), giving DE rate 29.58% vs 29.9%
  ERR-EU-003: formula_error — Consolidated TB total debit double-counts one
              entity's last debit-balance account

Uses deterministic European tax model — never hardcodes numbers.
"""

from __future__ import annotations

import datetime
import io
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from generator.canaries import CanaryRegistry, embed_canary_docx, embed_canary_xlsx
from generator.errors import ErrorRegistry, PlantedError, formula_error, stale_data
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel
from generator.model.tax_eu import (
    DE_GEWST_HEBESATZ,
    DE_GEWST_HEBESATZ_STALE,
    DE_RATE,
    DE_RATE_STALE,
    GBP_EUR_AVG,
    compute_eu_tax_provision,
    compute_eu_tax_provision_fy2024,
)

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-06-EU"
_INPUT_DIR = f"test_cases/{_TC}/input_files"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)
_FIXED_ZIP_DT = (2025, 3, 15, 9, 0, 0)

_ENTITY_NAMES = {
    "CE": "Cascade Europe Holdings B.V.",
    "CP": "Cascade Präzisionsteile GmbH",
    "CM": "Cascade Matériaux Avancés SAS",
    "CD": "Cascade Distribution Services Ltd",
}

_ENTITY_JURISDICTIONS = {
    "CE": "NL",
    "CP": "DE",
    "CM": "FR",
    "CD": "UK",
}


# ── Deterministic save helpers ───────────────────────────────────────────────

def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)
    wb.properties.modified = _FIXED_DATETIME

    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _pin_xlsx_dates(wb: openpyxl.Workbook) -> None:
    wb.properties.created = _FIXED_DATETIME


def _fmt_eur(d: Decimal | int) -> int:
    """Round to whole EUR for Excel cells (numeric)."""
    if isinstance(d, int):
        return d
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


def _fmt_eur_str(d: Decimal | int) -> str:
    """Format EUR amount with European conventions for docx text."""
    val = _fmt_eur(d)
    # European: period as thousand separator, comma as decimal (whole number)
    formatted = f"{abs(val):,}".replace(",", ".")
    sign = "-" if val < 0 else ""
    return f"{sign}{formatted} €"


# ── Styling constants ────────────────────────────────────────────────────────

_HEADER_FONT = Font(bold=True, size=10, color="FFFFFF")
_HEADER_FILL = PatternFill("solid", fgColor="1A3C6E")
_DATA_FONT = Font(size=10)
_BOLD_FONT = Font(bold=True, size=10)
_MONEY_FMT = "#,##0"
_PCT_FMT = "0.00%"
_THIN_BORDER = Border(
    bottom=Side(style="thin"),
)


def _write_header_row(ws: Any, row: int, headers: list[str]) -> None:
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col, value=header)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = Alignment(horizontal="center")


# ── 1. Consolidated Trial Balance ────────────────────────────────────────────


def _write_consolidated_tb(
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write tc06eu_consolidated_tb_fy2025.xlsx — consolidated TB by entity."""
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    ws = wb.active
    ws.title = "Trial Balance"

    canary_code = canaries.canary_for("tc06eu_consolidated_tb_fy2025")
    location = embed_canary_xlsx(wb, canary_code)
    canaries.set_location(
        "tc06eu_consolidated_tb_fy2025",
        f"{_INPUT_DIR}/tc06eu_consolidated_tb_fy2025.xlsx",
        location,
    )

    ws["A1"] = "Cascade Europe Holdings B.V. — Consolidated Group"
    ws["A1"].font = Font(bold=True, size=12)
    ws["A2"] = "Consolidated Trial Balance (EUR)"
    ws["A2"].font = Font(bold=True, size=11)
    ws["A3"] = "As of 31 December 2025"
    ws["A3"].font = Font(italic=True, size=10)

    header_row = 5
    _write_header_row(ws, header_row, [
        "Account", "Description", "Entity", "Debit (EUR)", "Credit (EUR)", "Net Balance (EUR)",
    ])

    # Build account rows from entity incomes and typical European COA
    _accounts = [
        ("4000", "Revenue", "credit"),
        ("4100", "Intercompany Revenue", "credit"),
        ("5000", "Cost of Goods Sold", "debit"),
        ("5100", "Raw Materials", "debit"),
        ("6000", "Personnel Costs", "debit"),
        ("6100", "Rent & Facilities", "debit"),
        ("6200", "Professional Fees", "debit"),
        ("6300", "Research & Development", "debit"),
        ("6400", "Travel & Entertainment", "debit"),
        ("6500", "IT & Communications", "debit"),
        ("6600", "Depreciation & Amortisation", "debit"),
        ("6700", "Insurance", "debit"),
        ("6800", "Other Operating Expenses", "debit"),
        ("7000", "Interest Income", "credit"),
        ("7100", "Interest Expense", "debit"),
        ("7200", "Foreign Exchange Gains/Losses", "debit"),
    ]

    # Deterministic allocation of pre-tax income to accounts per entity
    row = header_row + 1
    total_debit = Decimal(0)
    total_credit = Decimal(0)
    last_debit_by_entity: dict[str, Decimal] = {}

    from generator.model.tax_eu import _GROSS_MARGIN, _OPEX_RATIO, _REVENUE

    for entity in ("CE", "CP", "CM", "CD"):
        # Revenue (credit)
        growth = Decimal("1.08")
        rev = ((_REVENUE[entity] * growth).quantize(Decimal("1"), rounding=ROUND_HALF_UP))
        gp = (rev * _GROSS_MARGIN[entity]).quantize(Decimal("1"), rounding=ROUND_HALF_UP)
        cogs = rev - gp
        opex = (rev * _OPEX_RATIO[entity]).quantize(Decimal("1"), rounding=ROUND_HALF_UP)

        # Revenue row
        ws.cell(row=row, column=1, value="4000").font = _DATA_FONT
        ws.cell(row=row, column=2, value="Revenue").font = _DATA_FONT
        ws.cell(row=row, column=3, value=entity).font = _DATA_FONT
        ws.cell(row=row, column=4, value="").font = _DATA_FONT
        c5 = ws.cell(row=row, column=5, value=_fmt_eur(rev))
        c5.font = _DATA_FONT
        c5.number_format = _MONEY_FMT
        c6 = ws.cell(row=row, column=6, value=-_fmt_eur(rev))
        c6.font = _DATA_FONT
        c6.number_format = _MONEY_FMT
        total_credit += rev
        row += 1

        # COGS row
        ws.cell(row=row, column=1, value="5000").font = _DATA_FONT
        ws.cell(row=row, column=2, value="Cost of Goods Sold").font = _DATA_FONT
        ws.cell(row=row, column=3, value=entity).font = _DATA_FONT
        c4 = ws.cell(row=row, column=4, value=_fmt_eur(cogs))
        c4.font = _DATA_FONT
        c4.number_format = _MONEY_FMT
        ws.cell(row=row, column=5, value="").font = _DATA_FONT
        c6 = ws.cell(row=row, column=6, value=_fmt_eur(cogs))
        c6.font = _DATA_FONT
        c6.number_format = _MONEY_FMT
        total_debit += cogs
        last_debit_by_entity[entity] = cogs
        row += 1

        # Operating expenses row
        ws.cell(row=row, column=1, value="6000").font = _DATA_FONT
        ws.cell(row=row, column=2, value="Operating Expenses").font = _DATA_FONT
        ws.cell(row=row, column=3, value=entity).font = _DATA_FONT
        c4 = ws.cell(row=row, column=4, value=_fmt_eur(opex))
        c4.font = _DATA_FONT
        c4.number_format = _MONEY_FMT
        ws.cell(row=row, column=5, value="").font = _DATA_FONT
        c6 = ws.cell(row=row, column=6, value=_fmt_eur(opex))
        c6.font = _DATA_FONT
        c6.number_format = _MONEY_FMT
        total_debit += opex
        last_debit_by_entity[entity] = opex
        row += 1

    # Elimination row (intercompany)
    ws.cell(row=row, column=1, value="9000").font = _DATA_FONT
    ws.cell(row=row, column=2, value="Intercompany Eliminations").font = _DATA_FONT
    ws.cell(row=row, column=3, value="ELIM").font = _DATA_FONT
    elim_amount = Decimal("2500000")  # Management fees + loan interest
    c4 = ws.cell(row=row, column=4, value=_fmt_eur(elim_amount))
    c4.font = _DATA_FONT
    c4.number_format = _MONEY_FMT
    c5 = ws.cell(row=row, column=5, value=_fmt_eur(elim_amount))
    c5.font = _DATA_FONT
    c5.number_format = _MONEY_FMT
    c6 = ws.cell(row=row, column=6, value=0)
    c6.font = _DATA_FONT
    c6.number_format = _MONEY_FMT
    total_debit += elim_amount
    total_credit += elim_amount
    row += 1

    # ERR-EU-003: formula_error — double-count last debit-balance for one entity
    # Pick CD's last debit (opex) as the double-counted amount
    last_cd_debit = last_debit_by_entity["CD"]
    correct_total_debit = _fmt_eur(total_debit)
    wrong_total_debit = _fmt_eur(total_debit + last_cd_debit)
    displayed_total_debit = formula_error(correct_total_debit, wrong_total_debit)

    errors.add(PlantedError(
        error_id="ERR-EU-003",
        file=f"{_INPUT_DIR}/tc06eu_consolidated_tb_fy2025.xlsx",
        location="Sheet 'Trial Balance', Total Debit cell",
        type="formula_error",
        description=(
            f"Total Debit shows {wrong_total_debit:,.0f} € instead of "
            f"{correct_total_debit:,.0f} € (SUM range includes an extra row, "
            f"double-counting CD's last debit-balance account)"
        ),
        severity="material",
        which_test_cases_should_catch=[_TC],
    ))

    # Totals row
    ws.cell(row=row, column=2, value="TOTAL").font = _BOLD_FONT
    c4t = ws.cell(row=row, column=4, value=displayed_total_debit)
    c4t.font = _BOLD_FONT
    c4t.number_format = _MONEY_FMT
    c5t = ws.cell(row=row, column=5, value=_fmt_eur(total_credit))
    c5t.font = _BOLD_FONT
    c5t.number_format = _MONEY_FMT

    # Column widths
    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 10
    ws.column_dimensions["D"].width = 18
    ws.column_dimensions["E"].width = 18
    ws.column_dimensions["F"].width = 18

    path = output_dir / _INPUT_DIR / "tc06eu_consolidated_tb_fy2025.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, path)

    manifest.register(
        f"{_INPUT_DIR}/tc06eu_consolidated_tb_fy2025.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── 2. Prior Year Provision Workpaper ────────────────────────────────────────


def _write_prior_year_workpaper(
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write tc06eu_tax_provision_fy2024_workpaper.xlsx — prior year provision."""
    prov24 = compute_eu_tax_provision_fy2024()

    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)

    canary_code = canaries.canary_for("tc06eu_tax_provision_fy2024_workpaper")
    location = embed_canary_xlsx(wb, canary_code)
    canaries.set_location(
        "tc06eu_tax_provision_fy2024_workpaper",
        f"{_INPUT_DIR}/tc06eu_tax_provision_fy2024_workpaper.xlsx",
        location,
    )

    # ── Sheet 1: Current Tax by Jurisdiction ─────────────────────────
    ws1 = wb.active
    ws1.title = "Current Tax by Jurisdiction"

    ws1["A1"] = "Cascade Europe Holdings B.V. — Consolidated Group"
    ws1["A1"].font = Font(bold=True, size=12)
    ws1["A2"] = "Current Income Tax by Jurisdiction — FY2024"
    ws1["A2"].font = Font(bold=True, size=11)

    header_row = 4
    _write_header_row(ws1, header_row, [
        "Jurisdiction", "Entity", "Pre-Tax Income (EUR)",
        "Taxable Income (EUR)", "Statutory Rate", "Current Tax (EUR)",
    ])

    row = header_row + 1
    for ct in prov24.current_tax_by_entity:
        ws1.cell(row=row, column=1, value=ct.jurisdiction).font = _DATA_FONT
        ws1.cell(row=row, column=2, value=ct.entity_code).font = _DATA_FONT
        c3 = ws1.cell(row=row, column=3, value=_fmt_eur(ct.pre_tax_book_income_eur))
        c3.font = _DATA_FONT
        c3.number_format = _MONEY_FMT
        c4 = ws1.cell(row=row, column=4, value=_fmt_eur(ct.taxable_income_eur))
        c4.font = _DATA_FONT
        c4.number_format = _MONEY_FMT

        # ERR-EU-002: Show stale DE rate for FY2024
        displayed_rate = float(ct.statutory_rate)
        if ct.jurisdiction == "DE":
            displayed_rate = float(stale_data(DE_RATE_STALE))
            ws1.cell(row=row, column=5, value=displayed_rate).font = _DATA_FONT
        else:
            ws1.cell(row=row, column=5, value=displayed_rate).font = _DATA_FONT
        ws1.cell(row=row, column=5).number_format = _PCT_FMT

        c6 = ws1.cell(row=row, column=6, value=_fmt_eur(ct.current_tax_eur))
        c6.font = _DATA_FONT
        c6.number_format = _MONEY_FMT
        row += 1

    # Register ERR-EU-002
    errors.add(PlantedError(
        error_id="ERR-EU-002",
        file=f"{_INPUT_DIR}/tc06eu_tax_provision_fy2024_workpaper.xlsx",
        location="Sheet 'Current Tax by Jurisdiction', DE row, Statutory Rate column",
        type="stale_data",
        description=(
            f"German statutory rate shows {float(DE_RATE_STALE):.2%} "
            f"(Munich Hebesatz {DE_GEWST_HEBESATZ_STALE}%) instead of "
            f"{float(DE_RATE):.2%} (Hebesatz {DE_GEWST_HEBESATZ}%). "
            f"The FY2024 workpaper was not updated when Munich raised "
            f"its Gewerbesteuer multiplier from 480% to 490%."
        ),
        severity="significant",
        which_test_cases_should_catch=[_TC],
    ))

    # Total row
    ws1.cell(row=row, column=1, value="TOTAL").font = _BOLD_FONT
    c6t = ws1.cell(row=row, column=6, value=_fmt_eur(prov24.total_current_tax))
    c6t.font = _BOLD_FONT
    c6t.number_format = _MONEY_FMT

    ws1.column_dimensions["A"].width = 14
    ws1.column_dimensions["B"].width = 10
    ws1.column_dimensions["C"].width = 20
    ws1.column_dimensions["D"].width = 20
    ws1.column_dimensions["E"].width = 14
    ws1.column_dimensions["F"].width = 20

    # ── Sheet 2: Deferred Tax Rollforward ────────────────────────────
    ws2 = wb.create_sheet("Deferred Tax Rollforward")

    ws2["A1"] = "Deferred Tax Rollforward — FY2024"
    ws2["A1"].font = Font(bold=True, size=11)

    header_row = 3
    _write_header_row(ws2, header_row, [
        "Entity", "Description", "Type", "Opening (EUR)", "Movement (EUR)", "Closing (EUR)",
    ])

    row = header_row + 1
    for item in prov24.deferred_items:
        ws2.cell(row=row, column=1, value=item.entity_code).font = _DATA_FONT
        ws2.cell(row=row, column=2, value=item.description).font = _DATA_FONT
        ws2.cell(row=row, column=3, value=item.item_type).font = _DATA_FONT
        for col, val in [(4, item.opening_eur), (5, item.movement_eur), (6, item.closing_eur)]:
            c = ws2.cell(row=row, column=col, value=_fmt_eur(val))
            c.font = _DATA_FONT
            c.number_format = _MONEY_FMT
        row += 1

    # Totals
    ws2.cell(row=row, column=2, value="TOTAL").font = _BOLD_FONT
    for col, val in [
        (4, prov24.total_deferred_opening),
        (5, prov24.total_deferred_movement),
        (6, prov24.total_deferred_closing),
    ]:
        c = ws2.cell(row=row, column=col, value=_fmt_eur(val))
        c.font = _BOLD_FONT
        c.number_format = _MONEY_FMT

    ws2.column_dimensions["A"].width = 10
    ws2.column_dimensions["B"].width = 45
    ws2.column_dimensions["C"].width = 8
    ws2.column_dimensions["D"].width = 16
    ws2.column_dimensions["E"].width = 16
    ws2.column_dimensions["F"].width = 16

    # ── Sheet 3: ETR Reconciliation ──────────────────────────────────
    ws3 = wb.create_sheet("Rate Reconciliation")

    ws3["A1"] = "Effective Tax Rate Reconciliation — FY2024"
    ws3["A1"].font = Font(bold=True, size=11)

    header_row = 3
    _write_header_row(ws3, header_row, [
        "Description", "Amount (EUR)", "Rate Impact",
    ])

    row = header_row + 1
    for item in prov24.rate_reconciliation:
        ws3.cell(row=row, column=1, value=item.description).font = _DATA_FONT
        c2 = ws3.cell(row=row, column=2, value=_fmt_eur(item.amount_eur))
        c2.font = _DATA_FONT
        c2.number_format = _MONEY_FMT
        c3 = ws3.cell(row=row, column=3, value=float(item.rate_impact))
        c3.font = _DATA_FONT
        c3.number_format = _PCT_FMT
        row += 1

    # Total
    ws3.cell(row=row, column=1, value="Effective tax rate").font = _BOLD_FONT
    c2t = ws3.cell(row=row, column=2, value=_fmt_eur(prov24.total_provision))
    c2t.font = _BOLD_FONT
    c2t.number_format = _MONEY_FMT
    c3t = ws3.cell(row=row, column=3, value=float(prov24.effective_tax_rate))
    c3t.font = _BOLD_FONT
    c3t.number_format = _PCT_FMT

    ws3.column_dimensions["A"].width = 45
    ws3.column_dimensions["B"].width = 18
    ws3.column_dimensions["C"].width = 14

    # ── Sheet 4: Temporary Differences ───────────────────────────────
    ws4 = wb.create_sheet("Temporary Differences")

    ws4["A1"] = "Schedule of Temporary Differences — FY2024"
    ws4["A1"].font = Font(bold=True, size=11)

    header_row = 3
    _write_header_row(ws4, header_row, [
        "Entity", "Category", "Book Amount (EUR)", "Tax Amount (EUR)", "Difference (EUR)",
    ])

    row = header_row + 1
    for td in prov24.temporary_differences:
        ws4.cell(row=row, column=1, value=td.entity_code).font = _DATA_FONT
        ws4.cell(row=row, column=2, value=td.description).font = _DATA_FONT
        c3 = ws4.cell(row=row, column=3, value=_fmt_eur(td.book_amount_eur))
        c3.font = _DATA_FONT
        c3.number_format = _MONEY_FMT
        c4 = ws4.cell(row=row, column=4, value=_fmt_eur(td.tax_amount_eur))
        c4.font = _DATA_FONT
        c4.number_format = _MONEY_FMT
        c5 = ws4.cell(row=row, column=5, value=_fmt_eur(td.difference))
        c5.font = _DATA_FONT
        c5.number_format = _MONEY_FMT
        row += 1

    ws4.column_dimensions["A"].width = 10
    ws4.column_dimensions["B"].width = 45
    ws4.column_dimensions["C"].width = 18
    ws4.column_dimensions["D"].width = 18
    ws4.column_dimensions["E"].width = 18

    # ── Save ─────────────────────────────────────────────────────────
    path = output_dir / _INPUT_DIR / "tc06eu_tax_provision_fy2024_workpaper.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, path)

    manifest.register(
        f"{_INPUT_DIR}/tc06eu_tax_provision_fy2024_workpaper.xlsx",
        "xlsx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── 3. Permanent & Temporary Differences Document ───────────────────────────


def _write_perm_temp_differences(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write tc06eu_permanent_temporary_differences_fy2025.docx."""
    from docx import Document

    prov25 = compute_eu_tax_provision(2025)
    doc = Document()

    canary_code = canaries.canary_for("tc06eu_perm_temp_differences_fy2025")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "tc06eu_perm_temp_differences_fy2025",
        f"{_INPUT_DIR}/tc06eu_permanent_temporary_differences_fy2025.docx",
        location,
    )

    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME

    # Title
    doc.add_heading("Cascade Europe Holdings B.V.", level=1)
    doc.add_heading("Book-Tax Differences — FY2025", level=2)
    doc.add_paragraph("")

    # Permanent differences table
    doc.add_heading("Permanent Differences", level=3)
    doc.add_paragraph(
        "The following items create permanent differences between book income "
        "and taxable income. These differences do not reverse in future periods "
        "and are specific to each entity's jurisdiction."
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Entity"
    hdr[1].text = "Description"
    hdr[2].text = "Amount (EUR)"
    hdr[3].text = "Effect on Taxable Income"

    for p in prov25.permanent_differences:
        row = table.add_row().cells
        row[0].text = p.entity_code
        row[1].text = p.description
        row[2].text = _fmt_eur_str(abs(p.amount_eur))
        row[3].text = "Increases" if p.amount_eur > 0 else "Decreases"

    total_perm = sum(p.amount_eur for p in prov25.permanent_differences)
    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = "Total Permanent Differences"
    total_row[2].text = _fmt_eur_str(abs(total_perm))
    total_row[3].text = "Net increase" if total_perm > 0 else "Net decrease"

    doc.add_paragraph("")

    # Temporary differences table
    doc.add_heading("Temporary Differences", level=3)
    doc.add_paragraph(
        "The following items create temporary differences that will reverse "
        "in future periods, giving rise to deferred tax assets or liabilities "
        "under IAS 12."
    )

    table2 = doc.add_table(rows=1, cols=6)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Entity"
    hdr2[1].text = "Description"
    hdr2[2].text = "Book Amount (EUR)"
    hdr2[3].text = "Tax Amount (EUR)"
    hdr2[4].text = "Difference (EUR)"
    hdr2[5].text = "Creates"

    for td in prov25.temporary_differences:
        row = table2.add_row().cells
        row[0].text = td.entity_code
        row[1].text = td.description
        row[2].text = _fmt_eur_str(td.book_amount_eur)
        row[3].text = _fmt_eur_str(td.tax_amount_eur)
        diff = _fmt_eur(td.difference)
        row[4].text = _fmt_eur_str(abs(diff)) if diff >= 0 else f"({_fmt_eur_str(abs(diff))})"
        row[5].text = "DTL" if td.difference > 0 else "DTA"

    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: All amounts are presented in EUR. UK subsidiary amounts "
        "have been translated from GBP at the average rate of 1.17 per IAS 21."
    )

    path = output_dir / _INPUT_DIR / "tc06eu_permanent_temporary_differences_fy2025.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, path)

    manifest.register(
        f"{_INPUT_DIR}/tc06eu_permanent_temporary_differences_fy2025.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── 4. Statutory Rates Document ─────────────────────────────────────────────


def _write_statutory_rates(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write tc06eu_statutory_rates.docx — NL/DE/FR/UK rates + R&D + Pillar Two."""
    from docx import Document

    doc = Document()

    canary_code = canaries.canary_for("tc06eu_statutory_rates")
    location = embed_canary_docx(doc, canary_code)
    canaries.set_location(
        "tc06eu_statutory_rates",
        f"{_INPUT_DIR}/tc06eu_statutory_rates.docx",
        location,
    )

    doc.core_properties.created = _FIXED_DATETIME
    doc.core_properties.modified = _FIXED_DATETIME

    doc.add_heading("Cascade Europe Holdings B.V.", level=1)
    doc.add_heading("Statutory Tax Rates — FY2025", level=2)
    doc.add_paragraph("")

    # Rate table
    doc.add_heading("Corporate Tax Rates by Jurisdiction", level=3)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Jurisdiction"
    hdr[1].text = "Entity"
    hdr[2].text = "Statutory Rate"

    rate_data = [
        ("Netherlands", "Cascade Europe Holdings B.V. (CE)", "25.8%"),
        ("Germany", "Cascade Präzisionsteile GmbH (CP)", "29.9%"),
        ("France", "Cascade Matériaux Avancés SAS (CM)", "25.0%"),
        ("United Kingdom", "Cascade Distribution Services Ltd (CD)", "25.0%"),
    ]
    for jur, entity, rate in rate_data:
        row = table.add_row().cells
        row[0].text = jur
        row[1].text = entity
        row[2].text = rate

    doc.add_paragraph("")

    # German composite rate detail
    doc.add_heading("German Composite Rate Calculation", level=3)
    doc.add_paragraph(
        "The German corporate tax rate is a composite of three components:\n\n"
        "  Körperschaftsteuer (KSt): 15.000%\n"
        "  Solidaritätszuschlag (SolZ): 15% × 5.5% = 0.825%\n"
        "  Gewerbesteuer (GewSt): 3.5% × 490% Hebesatz (Munich) = 14.075%\n"
        "  ─────────────────────────────────\n"
        "  Combined rate: 29.9% (rounded)\n\n"
        "The Gewerbesteuer Hebesatz of 490% is the current Munich rate. "
        "Note that this rate was 480% in prior years."
    )

    doc.add_paragraph("")

    # R&D incentives
    doc.add_heading("R&D Tax Incentives", level=3)

    doc.add_paragraph(
        "France — Crédit d'Impôt Recherche (CIR)\n"
        "  Rate: 30% of first €100M of eligible R&D expenditure\n"
        "  Eligible entity: Cascade Matériaux Avancés SAS (Lyon)\n"
        "  Estimated FY2025 credit: ~1.245.696 €\n"
        "  Treatment: Direct reduction of French current tax expense\n\n"
        "Germany — Forschungszulage (Research Allowance)\n"
        "  Rate: 25% of eligible personnel costs (max €500K benefit)\n"
        "  Eligible entity: Cascade Präzisionsteile GmbH (Munich)\n"
        "  Estimated FY2025 benefit: ~291.600 €\n"
        "  Treatment: Separate cash benefit, not a direct tax credit\n\n"
        "United Kingdom — RDEC (Research and Development Expenditure Credit)\n"
        "  Rate: 20% for large companies\n"
        "  Eligible entity: Cascade Distribution Services Ltd\n"
        "  Treatment: Above-the-line credit (included in pre-tax income)"
    )

    doc.add_paragraph("")

    # Weighted group rate
    doc.add_heading("Weighted Group Statutory Rate", level=3)
    doc.add_paragraph(
        "The weighted group statutory rate is computed based on each entity's "
        "share of consolidated pre-tax book income. The weighted rate is used "
        "as the starting point for the effective tax rate reconciliation."
    )

    doc.add_paragraph("")

    # Pillar Two note
    doc.add_heading("OECD Pillar Two / GloBE Rules", level=3)
    doc.add_paragraph(
        "The group's consolidated revenue (~€120M) is below the €750M GloBE "
        "threshold. Pillar Two / GloBE rules do not apply to this group. "
        "No top-up tax or Income Inclusion Rule (IIR) computations are required."
    )

    doc.add_paragraph("")

    # FX note
    doc.add_heading("Currency Translation", level=3)
    doc.add_paragraph(
        "The UK subsidiary's functional currency is GBP. For consolidation "
        "purposes under IAS 21:\n"
        "  Average rate FY2025: GBP/EUR 1.17 (used for P&L items including tax expense)\n"
        "  Closing rate FY2025: GBP/EUR 1.17 (used for balance sheet items)"
    )

    path = output_dir / _INPUT_DIR / "tc06eu_statutory_rates.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, path)

    manifest.register(
        f"{_INPUT_DIR}/tc06eu_statutory_rates.docx",
        "docx",
        canary=canary_code,
        test_cases=[_TC],
    )


# ── Prompt & Expected Behavior ──────────────────────────────────────────────


def _write_prompt(output_dir: Path) -> None:
    """Write test_cases/TC-06-EU/prompt.md per design bead."""
    text = """\
Compute the income tax provision for the Cascade Europe Holdings B.V. group
for FY2025 under IAS 12.

1. Identify pre-tax book income from the consolidated trial balance by entity.
2. Apply permanent and temporary differences per entity to arrive at
   taxable income in each jurisdiction (NL, DE, FR, UK).
3. Calculate the current tax expense per entity at local statutory rates.
   - For Germany, compute the composite rate (Körperschaftsteuer + Solidaritätszuschlag + Gewerbesteuer).
   - For France, note the CIR credit impact on current tax.
4. Calculate the deferred tax provision per entity:
   - Compute the change in deferred tax assets and liabilities from FY2024.
   - Assess recoverability of deferred tax assets (IAS 12 requires probable
     future taxable profits, not ASC 740 "more likely than not" with
     valuation allowance).
5. Roll forward the deferred tax balance sheet from FY2024 to FY2025 per entity.
6. Prepare the consolidated effective tax rate (ETR) reconciliation:
   - Start from the weighted group statutory rate.
   - Reconcile to the effective rate showing: permanent differences,
     rate differentials, R&D credits, withholding taxes, FX effects.
7. Translate the UK subsidiary's tax amounts from GBP to EUR at the
   average rate (IAS 21).
8. Identify the total consolidated provision (current + deferred) and
   the group effective tax rate.
9. State whether OECD Pillar Two / GloBE rules apply to this group.

Export as an Excel workbook with separate sheets for:
- Current Tax by Jurisdiction
- Deferred Rollforward by Entity
- Rate Reconciliation (Consolidated)
- Summary

Verify that everything ties: current + deferred = total provision per entity
and consolidated, and the rate reconciliation explains the gap between the
weighted statutory rate and the effective rate.
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write test_cases/TC-06-EU/expected_behavior.md per design bead."""
    text = """\
# TC-06-EU: IAS 12 Income Tax Provision (European Group) — Expected Behavior

## Multi-Jurisdiction Current Provision
- The agent should compute current tax separately for each entity at local rates:
  NL 25.8%, DE 29.9% (composite), FR 25%, UK 25%.
- For Germany, the agent must show the three-component calculation:
  Körperschaftsteuer (15%) + Solidaritätszuschlag (0.825%) + Gewerbesteuer
  (3.5% × 490% Hebesatz = 14.075%).
- For France, the CIR credit (~€1.2M) should reduce current tax directly.
  The CIR is a tax credit, not a book income adjustment.
- UK tax amounts must be translated from GBP to EUR at the average rate of 1.17.

## IAS 12 Deferred Tax
- The agent should roll forward DTA/DTL balances from FY2024 to FY2025 per entity.
- Recoverability assessment must use IAS 12 "probable future taxable profits"
  language. If the agent creates a valuation allowance line or uses ASC 740
  "more likely than not" terminology, this is a judgment error.
- If the agent applies an indefinite reversal exception for subsidiary
  undistributed earnings, this is incorrect under IAS 12 (ASC 740 concept).

## Effective Tax Rate Reconciliation
- The starting point must be the weighted group statutory rate (based on
  entity profit mix), not a single country's rate.
- Reconciling items must include: rate differentials between jurisdictions,
  permanent differences, R&D credits, deferred tax movement, and FX effects.

## Pillar Two
- The agent must correctly determine that GloBE rules do not apply
  (group revenue ~€120M < €750M threshold).
- If the agent computes a GloBE top-up tax or IIR, this is a judgment error.

## Error Detection
- The agent should notice that the FY2024 workpaper uses a stale Munich
  Gewerbesteuer Hebesatz of 480% instead of the current 490%, resulting
  in a German rate of 29.58% instead of 29.9%.
- The agent should notice that the consolidated trial balance total debit
  is overstated due to a SUM range error that double-counts one entity's
  last debit-balance account.

## Output Quality
- Excel workbook with four sheets: Current Tax by Jurisdiction, Deferred
  Rollforward by Entity, Rate Reconciliation (Consolidated), Summary.
- All numbers must tie: current + deferred = total provision.
- EUR formatting should be used throughout.
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold Standard ────────────────────────────────────────────────────────────


@register_gold("TC-06-EU")
def _tc06_eu_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """Build the TC-06-EU gold standard from the European tax model."""
    prov25 = compute_eu_tax_provision(2025)

    return GoldStandard(
        test_case="TC-06-EU",
        expected_outputs={
            "file_type": "xlsx",
            "required_sheets": [
                "Current Tax by Jurisdiction",
                "Deferred Rollforward by Entity",
                "Rate Reconciliation",
                "Summary",
            ],
            "entity_pre_tax_income_eur": {
                entity: _fmt_eur(income)
                for entity, income in sorted(prov25.entity_pre_tax.items())
            },
            "consolidated_pre_tax_income_eur": _fmt_eur(prov25.consolidated_pre_tax),
            "permanent_differences": {
                p.description: {
                    "entity": p.entity_code,
                    "amount_eur": _fmt_eur(p.amount_eur),
                }
                for p in prov25.permanent_differences
            },
            "temporary_differences": {
                t.description: {
                    "entity": t.entity_code,
                    "book_eur": _fmt_eur(t.book_amount_eur),
                    "tax_eur": _fmt_eur(t.tax_amount_eur),
                    "difference_eur": _fmt_eur(t.difference),
                }
                for t in prov25.temporary_differences
            },
            "current_tax_by_entity": {
                ct.entity_code: {
                    "jurisdiction": ct.jurisdiction,
                    "pre_tax_income_eur": _fmt_eur(ct.pre_tax_book_income_eur),
                    "taxable_income_eur": _fmt_eur(ct.taxable_income_eur),
                    "statutory_rate": float(ct.statutory_rate),
                    "gross_tax_eur": _fmt_eur(ct.gross_tax_eur),
                    "tax_credits_eur": _fmt_eur(ct.tax_credits_eur),
                    "current_tax_eur": _fmt_eur(ct.current_tax_eur),
                }
                for ct in prov25.current_tax_by_entity
            },
            "total_current_tax_eur": _fmt_eur(prov25.total_current_tax),
            "deferred_rollforward": {
                f"{d.entity_code}_{d.description}": {
                    "opening_eur": _fmt_eur(d.opening_eur),
                    "movement_eur": _fmt_eur(d.movement_eur),
                    "closing_eur": _fmt_eur(d.closing_eur),
                    "type": d.item_type,
                }
                for d in prov25.deferred_items
            },
            "total_deferred_movement_eur": _fmt_eur(prov25.total_deferred_movement),
            "total_provision_eur": _fmt_eur(prov25.total_provision),
            "weighted_statutory_rate": float(prov25.weighted_statutory_rate),
            "effective_tax_rate": float(prov25.effective_tax_rate),
            "rate_reconciliation": {
                r.description: {
                    "amount_eur": _fmt_eur(r.amount_eur),
                    "rate_impact": float(r.rate_impact),
                }
                for r in prov25.rate_reconciliation
            },
            "pillar_two_applicable": False,
            "pillar_two_reason": "Group revenue (~€120M) below €750M GloBE threshold",
            "gbp_eur_rate": float(GBP_EUR_AVG),
            "tie_out_checks": {
                "current_plus_deferred_equals_total": (
                    _fmt_eur(prov25.total_current_tax) +
                    _fmt_eur(prov25.total_deferred_movement) ==
                    _fmt_eur(prov25.total_provision)
                ),
            },
        },
        canary_verification={
            "read_consolidated_tb": canaries.canary_for("tc06eu_consolidated_tb_fy2025"),
            "read_prior_workpaper": canaries.canary_for("tc06eu_tax_provision_fy2024_workpaper"),
            "read_differences": canaries.canary_for("tc06eu_perm_temp_differences_fy2025"),
            "read_statutory_rates": canaries.canary_for("tc06eu_statutory_rates"),
        },
        error_detection={
            "ERR-EU-002": (
                f"FY2024 workpaper shows German statutory rate as {float(DE_RATE_STALE):.2%} "
                f"(Munich Hebesatz {DE_GEWST_HEBESATZ_STALE}%) instead of "
                f"{float(DE_RATE):.2%} (Hebesatz {DE_GEWST_HEBESATZ}%)"
            ),
            "ERR-EU-003": (
                "Consolidated TB total debit is overstated — SUM range "
                "double-counts one entity's last debit-balance account"
            ),
        },
        scoring_hints={
            "correctness": (
                "Multi-jurisdiction current tax at correct local rates; "
                "deferred rollforward ties; weighted ETR reconciliation explains "
                "statutory-to-effective gap; CIR credit applied correctly"
            ),
            "completeness": (
                "All 4 sheets present; all permanent and temporary differences "
                "by entity; Pillar Two scoping answered"
            ),
            "format_compliance": "Valid xlsx with EUR formatting and separate named sheets",
            "robustness": (
                "Correct German composite rate with 3 components; "
                "GBP→EUR translation at average rate; handled multi-entity structure"
            ),
            "communication": (
                "Flagged stale Hebesatz in FY2024 workpaper (ERR-EU-002); "
                "verified tie-outs; stated IAS 12 recoverability assessment"
            ),
        },
        scenario_pack="cascade_europe_ifrs",
        service_line="tax",
        judgment_traps=[
            {
                "trap_id": "JT-EU-003",
                "trap_type": "standard_confusion",
                "expected_response": "Agent should use IAS 12 recoverability language, not ASC 740 valuation allowance",
                "description": (
                    "IAS 12 uses 'probable future taxable profits' for DTA recoverability. "
                    "If the agent creates a valuation allowance line, it is applying ASC 740 logic."
                ),
            },
            {
                "trap_id": "JT-EU-004",
                "trap_type": "pillar_two_application",
                "expected_response": "Agent should state Pillar Two does not apply (revenue < €750M)",
                "description": (
                    "The group's revenue (~€120M) is below the €750M GloBE threshold. "
                    "If the agent computes GloBE top-up tax or IIR, it is wrong."
                ),
            },
            {
                "trap_id": "JT-EU-005",
                "trap_type": "single_rate_application",
                "expected_response": "Agent must compute tax per jurisdiction, not a single blended rate",
                "description": (
                    "Each entity has a different statutory rate. Applying one rate to "
                    "consolidated income is incorrect for IAS 12 multi-jurisdiction provision."
                ),
            },
            {
                "trap_id": "JT-EU-006",
                "trap_type": "indefinite_reversal",
                "expected_response": "Agent should not apply indefinite reversal exception (ASC 740 concept)",
                "description": (
                    "IAS 12 does not have an indefinite reversal exception for subsidiary "
                    "undistributed earnings. If the agent applies one, it is using ASC 740."
                ),
            },
            {
                "trap_id": "JT-EU-007",
                "trap_type": "cir_treatment",
                "expected_response": "CIR should reduce French current tax, not adjust book income",
                "description": (
                    "The CIR is a direct tax credit against French IS. If the agent "
                    "treats CIR as a permanent difference affecting taxable income, deduct points."
                ),
            },
            {
                "trap_id": "JT-EU-008",
                "trap_type": "stale_data",
                "expected_response": "Agent should identify stale Munich Hebesatz (480% vs 490%)",
                "description": (
                    "The FY2024 workpaper uses a stale Gewerbesteuer Hebesatz. "
                    "The agent should use the current rate for FY2025."
                ),
            },
        ],
    )


# ── Public entry point ──────────────────────────────────────────────────────


def emit_tc06_eu(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
    **kwargs: object,
) -> None:
    """Write all TC-06-EU files to *output_dir*."""
    _write_consolidated_tb(output_dir, canaries, errors, manifest)
    _write_prior_year_workpaper(output_dir, canaries, errors, manifest)
    _write_perm_temp_differences(output_dir, canaries, manifest)
    _write_statutory_rates(output_dir, canaries, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
