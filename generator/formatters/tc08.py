"""Formatter: TC-08 — R&D Tax Credit Study, Section 41 (Tax, Complex).

Emits:
- test_cases/TC-08/input_files/rd_employee_time_records.csv
  ~2,340-row weekly time records for 45 R&D-eligible employees
- test_cases/TC-08/input_files/rd_project_descriptions/RD-001.docx .. RD-012.docx
  12 R&D project description documents
- test_cases/TC-08/input_files/payroll_data_fy2025.xlsx
  Full payroll register for Advanced Materials
- test_cases/TC-08/input_files/rd_supply_expenses.xlsx
  Supply and materials expenses coded to R&D cost centres
- test_cases/TC-08/prompt.md
- test_cases/TC-08/expected_behavior.md
- gold_standards/TC-08_gold.json

Planted errors:
- ERR-005 mismatched_total in payroll_data_fy2025.xlsx (material)
- ERR-019 classification_error in rd_supply_expenses.xlsx (immaterial)

Uses the canonical model — never hardcodes numbers.
"""

from __future__ import annotations

import dataclasses
import datetime
import io
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import docx
import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from generator.canaries import (
    CanaryRegistry,
    embed_canary_csv_comment,
    embed_canary_docx,
    embed_canary_xlsx,
)
from generator.errors import (
    ErrorRegistry,
    PlantedError,
    classification_error,
    mismatch_total,
)
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel
from generator.model.rd import (
    PRIOR_YEAR_QRES,
    RD_PROJECTS,
    TARGET_CREDIT,
    TimeRecord,
    generate_rd_projects,
)
from generator.noise import apply_csv_noise, apply_docx_noise, make_noise_rng
from generator.scenario_context import ScenarioContext

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-08"
_INPUT_DIR = f"test_cases/{_TC}/input_files"
_PROJ_DESC_DIR = f"{_INPUT_DIR}/rd_project_descriptions"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)

# Canary file keys — one per generated file.
_KEY_TIME_RECORDS = "tc08_rd_employee_time_records"
_KEY_PAYROLL = "tc08_payroll_data_fy2025"
_KEY_SUPPLY = "tc08_rd_supply_expenses"
# 12 project description docx files
_PROJECT_DOC_KEYS = [f"tc08_rd_project_{i:03d}" for i in range(1, 13)]


# ── Deterministic save helpers ───────────────────────────────────────────────


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)

    wb.properties.created = _FIXED_DATETIME
    wb.properties.modified = _FIXED_DATETIME

    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    fixed_date_time = (2025, 3, 15, 9, 0, 0)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=fixed_date_time)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)

    buf = io.BytesIO()
    doc.save(buf)

    fixed_date_time = (2025, 3, 15, 9, 0, 0)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=fixed_date_time)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _whole_dollars(d: Decimal) -> int:
    """Round a Decimal to whole dollars."""
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


# ── Time Records CSV ────────────────────────────────────────────────────────


def _write_time_records_csv(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
    *,
    ctx: ScenarioContext | None = None,
) -> list[TimeRecord]:
    """Write rd_employee_time_records.csv and return the records."""
    records = model.rd_time_records

    rel_path = f"{_INPUT_DIR}/rd_employee_time_records.csv"
    abs_path = output_dir / rel_path

    canary = canaries.canary_for(_KEY_TIME_RECORDS)
    canary_line = embed_canary_csv_comment(canary)

    columns = [
        "employee_id",
        "week_ending",
        "project_code",
        "hours",
        "activity_description",
    ]

    # Build lines in memory so we can apply noise before writing.
    lines: list[str] = [canary_line, ",".join(columns) + "\n"]
    for rec in records:
        desc = rec.activity_description
        if "," in desc or '"' in desc:
            desc = '"' + desc.replace('"', '""') + '"'
        row = [
            rec.employee_id,
            rec.week_ending.isoformat(),
            rec.project_code,
            str(rec.hours),
            desc,
        ]
        lines.append(",".join(row) + "\n")

    # ── Controlled noise (csv_stdlib family pilot) ───────────────────
    # Canary line (index 0) is auto-protected by apply_csv_noise.
    # Header line (index 1) is excluded to preserve column names.
    from generator.noise import ExclusionZone
    noise_ctx = ctx if ctx is not None else ScenarioContext(seed=42)
    noise_rng = make_noise_rng(noise_ctx, "TC-08", _KEY_TIME_RECORDS)
    excl = ExclusionZone(rows={1})  # protect header row
    lines = apply_csv_noise(lines, noise_rng, excl)

    abs_path.parent.mkdir(parents=True, exist_ok=True)
    with open(abs_path, "w", newline="") as f:
        for line in lines:
            f.write(line)

    canaries.set_location(
        _KEY_TIME_RECORDS,
        rel_path,
        "Line 1 comment: # CANARY: ...",
    )
    manifest.register(rel_path, "csv")

    return records


# ── Project Description DOCX files ──────────────────────────────────────────


def _write_project_descriptions(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write 12 project description .docx files, one per R&D project."""
    projects = generate_rd_projects()

    desc_dir = output_dir / _PROJ_DESC_DIR
    desc_dir.mkdir(parents=True, exist_ok=True)

    for i, proj in enumerate(projects):
        file_key = _PROJECT_DOC_KEYS[i]
        canary = canaries.canary_for(file_key)

        doc = docx.Document()
        loc = embed_canary_docx(doc, canary)

        # Title
        doc.add_heading(f"R&D Project Description: {proj.code}", level=1)

        doc.add_paragraph("")

        # Fields
        fields = [
            ("Project Code", proj.code),
            ("Project Name", proj.name),
            ("Status", proj.status),
        ]
        for label, value in fields:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(value)

        doc.add_heading("Objective", level=2)
        doc.add_paragraph(proj.objective)

        doc.add_heading("Technical Uncertainty Addressed", level=2)
        doc.add_paragraph(proj.uncertainty)

        doc.add_heading("Methodology", level=2)
        doc.add_paragraph(proj.methodology)

        # ── Controlled noise (docx_python_docx family pilot) ────────
        # Canary lives in core_properties.comments — apply_docx_noise
        # never touches that field.  No planted errors in project
        # description docs, so no paragraph exclusions needed.
        noise_rng = make_noise_rng(
            ScenarioContext(seed=42), "TC-08", file_key,
        )
        apply_docx_noise(doc, noise_rng)

        rel_path = f"{_PROJ_DESC_DIR}/{proj.code}.docx"
        abs_path = output_dir / rel_path

        _save_docx_deterministic(doc, abs_path)

        canaries.set_location(
            file_key,
            rel_path,
            f"{loc}",
        )
        manifest.register(rel_path, "docx")


# ── Payroll XLSX ────────────────────────────────────────────────────────────


def _write_payroll(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write payroll_data_fy2025.xlsx — full payroll register for AM."""
    canary = canaries.canary_for(_KEY_PAYROLL)

    wb = openpyxl.Workbook()
    loc = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Payroll Register FY2025"

    # Styles
    header_fill = PatternFill("solid", fgColor="1A3C6E")
    header_font = Font(bold=True, size=11, color="FFFFFF")
    number_fmt = "#,##0"
    border = Border(bottom=Side(style="thin", color="CCCCCC"))

    # Title rows
    ws["A1"] = "Cascade Advanced Materials, Inc."
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "Payroll Register — FY2025"
    ws["A2"].font = Font(bold=True, size=12, color="666666")
    ws.merge_cells("A1:H1")
    ws.merge_cells("A2:H2")

    headers = [
        "Employee ID",
        "Employee Name",
        "Department",
        "Title",
        "W-2 Wages (Annual Salary)",
        "Employer FICA (7.65%)",
        "Employer Benefits",
        "Total Compensation",
    ]

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    # Filter to AM employees active in FY2025
    fy_start = datetime.date(2025, 1, 1)
    am_employees = sorted(
        [
            e for e in model.employees
            if e.entity_code == "AM"
            and (e.termination_date is None or e.termination_date >= fy_start)
        ],
        key=lambda e: e.employee_id,
    )

    total_wages = Decimal(0)
    for row_idx, emp in enumerate(am_employees, start=5):
        salary = Decimal(emp.annual_salary)
        total_wages += salary
        fica = (salary * Decimal("0.0765")).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP,
        )
        benefits = (salary * Decimal("0.18")).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP,
        )
        total = salary + fica + benefits

        ws.cell(row=row_idx, column=1, value=emp.employee_id)
        ws.cell(row=row_idx, column=2, value=emp.name)
        ws.cell(row=row_idx, column=3, value=emp.department)
        ws.cell(row=row_idx, column=4, value=emp.title)

        for col_idx, val in enumerate(
            [_whole_dollars(salary), _whole_dollars(fica),
             _whole_dollars(benefits), _whole_dollars(total)],
            start=5,
        ):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.number_format = number_fmt
            cell.border = border

    # ── ERR-005: mismatched payroll total wages ─────────────────────────────
    # The total row omits one employee's wages (~$2,400 delta).
    # Pick the last employee's salary as the omitted amount.
    omitted_salary = Decimal(am_employees[-1].annual_salary)
    correct_total_wages = _whole_dollars(total_wages)
    wrong_total_wages = mismatch_total(correct_total_wages, -_whole_dollars(omitted_salary))

    totals_row = 5 + len(am_employees)
    ws.cell(row=totals_row, column=1, value="")
    ws.cell(row=totals_row, column=2, value="TOTAL")
    ws.cell(row=totals_row, column=2).font = Font(bold=True)
    cell = ws.cell(row=totals_row, column=5, value=wrong_total_wages)
    cell.number_format = number_fmt
    cell.font = Font(bold=True)
    cell.border = border

    errors.add(PlantedError(
        error_id="ERR-005",
        file=f"{_INPUT_DIR}/payroll_data_fy2025.xlsx",
        location=f"Sheet 'Payroll Register FY2025', Row {totals_row}, Column E (Total W-2 Wages)",
        type="mismatched_total",
        description=(
            f"Payroll total wages row shows ${wrong_total_wages:,} instead of "
            f"${correct_total_wages:,} — one employee's salary "
            f"(${_whole_dollars(omitted_salary):,}) was omitted from the total"
        ),
        severity="material",
        which_test_cases_should_catch=["TC-08"],
    ))

    # Column widths
    widths = [14, 24, 20, 28, 22, 18, 18, 20]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    rel_path = f"{_INPUT_DIR}/payroll_data_fy2025.xlsx"
    abs_path = output_dir / rel_path
    _save_xlsx_deterministic(wb, abs_path)

    canaries.set_location(_KEY_PAYROLL, rel_path, loc)
    manifest.register(rel_path, "xlsx")


# ── Supply Expenses XLSX ────────────────────────────────────────────────────


def _write_supply_expenses(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write rd_supply_expenses.xlsx."""
    expenses = list(model.rd_supply_expenses)

    # ── ERR-019: classification_error — one expense under wrong project ─────
    # Pick expense at index 7 and swap its project code to a non-qualifying
    # project (RD-011, which is market research and should be excluded).
    err_idx = 7
    err_exp = expenses[err_idx]
    correct_project = err_exp.project_code
    wrong_project = classification_error(correct_project, "RD-011")
    expenses[err_idx] = dataclasses.replace(err_exp, project_code=wrong_project)

    canary = canaries.canary_for(_KEY_SUPPLY)

    wb = openpyxl.Workbook()
    loc = embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "R&D Supply Expenses"

    # Styles
    header_fill = PatternFill("solid", fgColor="1A3C6E")
    header_font = Font(bold=True, size=11, color="FFFFFF")
    number_fmt = "#,##0.00"
    border = Border(bottom=Side(style="thin", color="CCCCCC"))

    ws["A1"] = "Cascade Advanced Materials, Inc."
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "R&D Supply & Materials Expenses — FY2025"
    ws["A2"].font = Font(bold=True, size=12, color="666666")
    ws.merge_cells("A1:F1")
    ws.merge_cells("A2:F2")

    headers = [
        "Date",
        "Vendor",
        "Description",
        "Amount",
        "Project Code",
        "Cost Center",
    ]

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    for row_idx, exp in enumerate(expenses, start=5):
        ws.cell(row=row_idx, column=1, value=exp.date.isoformat())
        ws.cell(row=row_idx, column=2, value=exp.vendor)
        ws.cell(row=row_idx, column=3, value=exp.description)
        cell = ws.cell(row=row_idx, column=4, value=float(exp.amount))
        cell.number_format = number_fmt
        cell.border = border
        ws.cell(row=row_idx, column=5, value=exp.project_code)
        ws.cell(row=row_idx, column=6, value=exp.cost_center)

    widths = [12, 30, 36, 14, 14, 14]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    rel_path = f"{_INPUT_DIR}/rd_supply_expenses.xlsx"
    abs_path = output_dir / rel_path
    _save_xlsx_deterministic(wb, abs_path)

    canaries.set_location(_KEY_SUPPLY, rel_path, loc)
    manifest.register(rel_path, "xlsx")

    # Register ERR-019
    err_row = 5 + err_idx
    errors.add(PlantedError(
        error_id="ERR-019",
        file=f"{_INPUT_DIR}/rd_supply_expenses.xlsx",
        location=f"Sheet 'R&D Supply Expenses', Row {err_row}, Column E (Project Code)",
        type="classification_error",
        description=(
            f"Supply expense '{err_exp.description}' classified under "
            f"project {wrong_project} (market research — non-qualifying) "
            f"instead of {correct_project}"
        ),
        severity="immaterial",
        which_test_cases_should_catch=["TC-08"],
    ))


# ── Prompt & Expected Behavior ──────────────────────────────────────────────


def _write_prompt(output_dir: Path) -> None:
    """Write test_cases/TC-08/prompt.md per spec."""
    text = """\
Perform an R&D tax credit study for Cascade Advanced Materials, Inc. for FY2025.

1. Review each project description and determine whether it meets the four-part
   test for qualified research:
   - Permitted purpose (new or improved function, performance, reliability, quality)
   - Technological in nature (relies on principles of physical/biological/computer science)
   - Technological uncertainty (capability, method, or design uncertainty)
   - Process of experimentation (systematic evaluation of alternatives)
2. For qualifying projects, compute qualified research expenses (QREs):
   - Wages: allocate based on time records (% of time on qualifying activities × W-2 wages)
   - Supplies: include supplies directly used in qualified research
   - Do not include overhead or general administrative expenses
3. Compute the credit using the Alternative Simplified Credit (ASC) method:
   - Average QREs for FY2023-FY2025 (provide FY2023 and FY2024 QREs as $3.1M and $3.4M)
   - Credit = 14% × (current year QREs − 50% × average of prior 3 years QREs)
4. Draft a contemporaneous documentation memo for each qualifying project that
   summarizes the technical uncertainty and experimentation.

Export:
- Project qualification analysis as Excel (project name, qualification determination, rationale)
- QRE computation as Excel
- Credit calculation as Excel
- Documentation memos as a single Word document with one section per project
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write test_cases/TC-08/expected_behavior.md per spec."""
    # Build a project summary table for the expected behavior doc.
    projects = generate_rd_projects()
    qual_lines: list[str] = []
    for p in projects:
        if p.qualifies == "yes":
            status = "Qualifies"
        elif p.qualifies == "borderline":
            status = "Borderline — flag for manager review"
        else:
            status = f"Does NOT qualify — {p.disqualification_reason}"
        qual_lines.append(f"- **{p.code} ({p.name})**: {status}")

    qual_table = "\n".join(qual_lines)

    text = f"""\
# TC-08: R&D Tax Credit Study (Section 41) — Expected Behavior

## Project Qualification (4-Part Test)

The agent must evaluate each of the 12 R&D projects against the IRC §41
four-part test and classify them correctly:

{qual_table}

### Key Traps

1. **RD-011 (Market Analysis)**: Labeled as "R&D" in the time records but is
   pure market research. Does not meet the permitted purpose or technological
   uncertainty tests. The agent must exclude it from QREs despite employees
   logging time to this project code.

2. **RD-012 (ERP System Migration)**: An IT implementation project using
   established software. No technological uncertainty or experimentation.
   The agent must exclude it.

3. **RD-009 and RD-010 (Borderline)**: These should be flagged for manager
   review. RD-009 uses commercial off-the-shelf technology; RD-010 tests
   known materials against existing specs. Both have arguable technological
   uncertainty but are not clear-cut.

## QRE Computation

- **Wage QREs**: For each employee, compute:
  (hours on qualifying/borderline R&D projects / total hours) × W-2 wages
- **Supply QREs**: Sum supply expenses coded to qualifying R&D projects
- **Exclude**: Time logged to GEN-xxx overhead codes, RD-011, and RD-012
- **Total QREs for FY2025**: ~$2,885,714

## ASC Credit Calculation

- Prior year QREs: FY2023 = $3,100,000; FY2024 = $3,400,000
- Average of 3 years = (FY2023 + FY2024 + FY2025) / 3
- Credit = 14% × (FY2025 QREs − 50% × 3-year average)
- **Expected credit: ~$185,000** (within $500)

## Expected Deliverables

1. **Project Qualification Excel**: One row per project with columns for
   project code, name, 4-part test evaluation, qualification determination,
   and rationale.
2. **QRE Computation Excel**: Per-project breakdown of wage QREs and supply
   QREs, with employee-level detail.
3. **Credit Calculation Excel**: ASC method calculation showing prior year
   QREs, 3-year average, and credit computation.
4. **Documentation Memos (Word)**: One section per qualifying project
   summarizing the technical uncertainty addressed and the process of
   experimentation conducted.

## Data Challenges

- **45 employees × 52 weeks = ~2,340 time records**: The agent must aggregate
  hours by employee and project, then cross-reference with payroll data.
- **Non-qualifying project time**: Employees log time to RD-011 and RD-012,
  which must be excluded. The agent must read the project descriptions to
  determine this.
- **Overhead exclusion**: Time logged to GEN-xxx codes is general overhead
  and must not be included in QREs.
- **Borderline judgment**: The agent should include borderline project QREs
  in the computation but explicitly flag them for manager review.
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold Standard ───────────────────────────────────────────────────────────


@register_gold("TC-08")
def _tc08_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """TC-08 gold standard: R&D tax credit study."""
    model: CascadeModel = model_kwargs["model"]

    qre_result = model.rd_qre_result

    # Project qualification details
    project_details: dict[str, dict[str, Any]] = {}
    for p in RD_PROJECTS:
        detail: dict[str, Any] = {
            "name": p.name,
            "status": p.status,
            "qualifies": p.qualifies,
        }
        if p.disqualification_reason:
            detail["disqualification_reason"] = p.disqualification_reason
        if p.code in qre_result.qres_by_project:
            detail["qre_amount"] = str(qre_result.qres_by_project[p.code])
        else:
            detail["qre_amount"] = "0"
        project_details[p.code] = detail

    expected_outputs: dict[str, Any] = {
        "projects_total": 12,
        "projects_qualifying": 8,
        "projects_borderline": 2,
        "projects_disqualified": 2,
        "wage_qres": str(qre_result.wage_qres),
        "supply_qres": str(qre_result.supply_qres),
        "total_qres_fy2025": str(qre_result.total_qres),
        "prior_year_qres": {
            "fy2023": str(PRIOR_YEAR_QRES[2023]),
            "fy2024": str(PRIOR_YEAR_QRES[2024]),
        },
        "asc_credit": str(qre_result.credit),
        "target_credit": str(TARGET_CREDIT),
        "qres_by_project": {
            k: str(v) for k, v in sorted(qre_result.qres_by_project.items())
        },
        "project_details": project_details,
    }

    # Canary verification — all files the agent must read
    canary_verification: dict[str, str] = {
        "read_time_records": canaries.canary_for(_KEY_TIME_RECORDS),
        "read_payroll": canaries.canary_for(_KEY_PAYROLL),
        "read_supply_expenses": canaries.canary_for(_KEY_SUPPLY),
    }
    # Add project description canaries
    for i, key in enumerate(_PROJECT_DOC_KEYS):
        canary_verification[f"read_project_{RD_PROJECTS[i].code}"] = (
            canaries.canary_for(key)
        )

    scoring_hints: dict[str, str] = {
        "correctness": (
            "8 of 12 projects qualified; 2 borderline flagged for review; "
            "2 disqualified (market research mislabeled as R&D, ERP migration). "
            f"Credit = {qre_result.credit} (must be within $500 of $185,000). "
            "QRE amounts must match per-project gold standard."
        ),
        "completeness": (
            "All 4 deliverables: project qualification Excel, QRE computation "
            "Excel, credit calculation Excel, documentation memos Word. "
            "Memos must have one section per qualifying project."
        ),
        "format_compliance": (
            "Valid Excel and Word files. Excel has clear project-by-project "
            "breakdowns. Memos address four-part test for each project."
        ),
        "robustness": (
            "Four-part test applied rigorously to each project. "
            "Wages allocated using time records × W-2. "
            "Overhead and admin excluded. "
            "ASC method used correctly with prior year QREs "
            "($3.1M FY2023, $3.4M FY2024)."
        ),
        "communication": (
            "Qualification rationale documented per project. "
            "Borderline projects explicitly flagged for manager review. "
            "ASC method and prior-year QRE basis explained."
        ),
    }

    return GoldStandard(
        test_case=_TC,
        expected_outputs=expected_outputs,
        canary_verification=canary_verification,
        error_detection={
            "ERR-005": (
                "Payroll total wages row omits one employee — "
                "mismatched_total in payroll_data_fy2025.xlsx"
            ),
            "ERR-019": (
                "Supply expense classified under wrong project code (RD-011) — "
                "classification_error in rd_supply_expenses.xlsx"
            ),
        },
        scoring_hints=scoring_hints,
    )


# ── Public API ──────────────────────────────────────────────────────────────


def emit_tc08(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
    *,
    ctx: ScenarioContext | None = None,
    **kwargs: object,
) -> None:
    """Emit all TC-08 files."""
    _write_time_records_csv(model, output_dir, canaries, manifest, ctx=ctx)
    _write_project_descriptions(output_dir, canaries, manifest)
    _write_payroll(model, output_dir, canaries, errors, manifest)
    _write_supply_expenses(model, output_dir, canaries, errors, manifest)
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
