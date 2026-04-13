"""Formatter: TC-18-EU — IFRS/ISA Prior-Year Workpaper Rollforward.

European variant of TC-18.  Replaces US GAAP/PCAOB framing with IFRS/ISA
for Cascade Europe Holdings B.V. group (CE, CP, CM, CD).

Emits:
- test_cases/TC-18-EU/input_files/prior_year_workpapers/
    6 xlsx workpapers: revenue, operating_expenses, balance_sheet, cash, fixed_assets, leases
    4 docx memos: planning, risk_assessment, summary, management_letter
    Each contains FY2024 data with IFRS terminology and ISA audit framework
- test_cases/TC-18-EU/input_files/current_year_data/
    trial_balance_fy2025.csv  (format change from xlsx!)
    bank_statements_fy2025.csv  (multi-currency: EUR + GBP)
    lease_schedule_fy2025.xlsx  (IFRS 16 single model, 2 new leases)
    management_projections_fy2025.docx  (format change from xlsx! + Pillar Two trap)
    goodwill_impairment_analysis_ifrs.xlsx  (IAS 36 CGU-based, NEW file)
- test_cases/TC-18-EU/prompt.md
- test_cases/TC-18-EU/expected_behavior.md
- gold_standards/TC-18-EU_gold.json

Planted errors:
  ERR-EU-018: stale_data — CP product revenue in wp_revenue_fy2024.xlsx shows
              FY2023 value (€38,200,000) instead of FY2024 value (€40,100,000)

Key differences from US TC-18:
  - IFRS terminology throughout (statement of financial position, right-of-use assets, etc.)
  - ISA audit framework (ISA 315, ISA 600, ISA 701 KAM)
  - 4-entity group with component auditors and intercompany eliminations
  - Multi-currency: EUR + GBP with IAS 21 translation
  - IAS 36 CGU-based impairment (not ASC 350 single-step)
  - IAS 16 revaluation model for CP equipment
  - IAS 7 interest in financing activities
  - Pillar Two threshold trap (€120M < €750M → not applicable)
"""

from __future__ import annotations

import csv
import datetime
import io
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import docx
import openpyxl
from openpyxl.styles import Border, Font, PatternFill, Side
from openpyxl.writer.excel import ExcelWriter

from generator.canaries import (
    CanaryRegistry,
    embed_canary_csv_comment,
    embed_canary_docx,
    embed_canary_xlsx,
)
from generator.errors import ErrorRegistry, PlantedError, stale_data
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.build import CascadeModel

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-18-EU"
_INPUT_DIR = f"test_cases/{_TC}/input_files"
_PY_DIR = f"{_INPUT_DIR}/prior_year_workpapers"
_CY_DIR = f"{_INPUT_DIR}/current_year_data"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)
_FIXED_ZIP_DT = (2025, 3, 15, 9, 0, 0)
_EOY_2024 = datetime.date(2024, 12, 31)
_EOY_2025 = datetime.date(2025, 12, 31)

# Canary file keys — 6 prior-year xlsx + 4 prior-year docx + 5 current-year files
_CANARY_KEYS: list[str] = sorted([
    "tc18eu_wp_revenue_fy2024",
    "tc18eu_wp_operating_expenses_fy2024",
    "tc18eu_wp_balance_sheet_fy2024",
    "tc18eu_wp_cash_fy2024",
    "tc18eu_wp_fixed_assets_fy2024",
    "tc18eu_wp_leases_fy2024",
    "tc18eu_memo_planning_fy2024",
    "tc18eu_memo_risk_assessment_fy2024",
    "tc18eu_memo_summary_fy2024",
    "tc18eu_memo_management_letter_fy2024",
    "tc18eu_cy_trial_balance_fy2025",
    "tc18eu_cy_bank_statements_fy2025",
    "tc18eu_cy_lease_schedule_fy2025",
    "tc18eu_cy_mgmt_projections_fy2025",
    "tc18eu_cy_goodwill_impairment_ifrs_fy2025",
])

# ── European entity data ────────────────────────────────────────────────────

_ENTITIES = {
    "CE": {
        "name": "Cascade Europe Holdings B.V.",
        "jurisdiction": "Netherlands",
        "city": "Amsterdam",
        "currency": "EUR",
        "role": "European holding company",
    },
    "CP": {
        "name": "Cascade Präzisionsteile GmbH",
        "jurisdiction": "Germany",
        "city": "Munich",
        "currency": "EUR",
        "role": "Licensed manufacturer",
    },
    "CM": {
        "name": "Cascade Matériaux Avancés SAS",
        "jurisdiction": "France",
        "city": "Lyon",
        "currency": "EUR",
        "role": "R&D centre, IP developer",
    },
    "CD": {
        "name": "Cascade Distribution Services Ltd",
        "jurisdiction": "United Kingdom",
        "city": "Birmingham",
        "currency": "GBP",
        "role": "Post-Brexit distributor",
    },
}

# FX rates for IAS 21 translation (GBP → EUR)
_FX_CLOSING_2024 = Decimal("1.1530")  # 31 Dec 2024
_FX_AVERAGE_2024 = Decimal("1.1620")  # FY2024 average
_FX_CLOSING_2025 = Decimal("1.1480")  # 31 Dec 2025
_FX_AVERAGE_2025 = Decimal("1.1550")  # FY2025 average

# ── Revenue data by entity (EUR thousands, hardcoded per design) ────────────
# CP product revenue FY2023 = 38,200,000 (stale), FY2024 = 40,100,000 (correct)

_REVENUE_FY2024 = [
    ("CE", "Management fee revenue", Decimal("4_800_000")),
    ("CP", "Product revenue — precision parts", Decimal("40_100_000")),
    ("CP", "Service revenue — calibration", Decimal("3_200_000")),
    ("CM", "Product revenue — advanced materials", Decimal("28_500_000")),
    ("CM", "R&D service revenue", Decimal("5_100_000")),
    ("CD", "Distribution revenue (GBP translated)", Decimal("22_400_000")),
    ("CD", "Warehousing revenue (GBP translated)", Decimal("4_900_000")),
]

_REVENUE_FY2023_CP_PRODUCT = Decimal("38_200_000")  # stale value for ERR-EU-018

# Intercompany eliminations
_IC_ELIMINATIONS_REVENUE = [
    ("CP→CM", "Intercompany parts supply", Decimal("-2_100_000")),
    ("CP→CD", "Intercompany distribution supply", Decimal("-3_400_000")),
    ("CE→all", "Management fee elimination", Decimal("-4_800_000")),
]

_TOTAL_CONSOLIDATED_REVENUE_2024 = sum(
    r[2] for r in _REVENUE_FY2024
) + sum(e[2] for e in _IC_ELIMINATIONS_REVENUE)

# FY2025 revenue (for current-year data)
_REVENUE_FY2025 = [
    ("CE", "Management fee revenue", Decimal("5_040_000")),
    ("CP", "Product revenue — precision parts", Decimal("42_100_000")),
    ("CP", "Service revenue — calibration", Decimal("3_500_000")),
    ("CM", "Product revenue — advanced materials", Decimal("30_200_000")),
    ("CM", "R&D service revenue", Decimal("5_400_000")),
    ("CD", "Distribution revenue (GBP translated)", Decimal("23_800_000")),
    ("CD", "Warehousing revenue (GBP translated)", Decimal("5_300_000")),
]

_IC_ELIMINATIONS_REVENUE_2025 = [
    ("CP→CM", "Intercompany parts supply", Decimal("-2_300_000")),
    ("CP→CD", "Intercompany distribution supply", Decimal("-3_600_000")),
    ("CE→all", "Management fee elimination", Decimal("-5_040_000")),
]

_TOTAL_CONSOLIDATED_REVENUE_2025 = sum(
    r[2] for r in _REVENUE_FY2025
) + sum(e[2] for e in _IC_ELIMINATIONS_REVENUE_2025)

# ── Operating expenses by nature (IAS 1) ────────────────────────────────────

_OPEX_BY_NATURE_FY2024 = [
    ("Raw materials and consumables", Decimal("32_800_000")),
    ("Employee benefits expense", Decimal("28_500_000")),
    ("Depreciation and amortisation", Decimal("6_200_000")),
    ("Impairment losses", Decimal("150_000")),
    ("Other operating expenses", Decimal("8_700_000")),
    ("Social charges (DE §1 SvEV)", Decimal("5_700_000")),
    ("Social charges (FR charges sociales)", Decimal("3_800_000")),
    ("Social charges (UK NIC)", Decimal("620_000")),
    ("Social charges (NL premies)", Decimal("500_000")),
]

_TOTAL_OPEX_2024 = sum(e[1] for e in _OPEX_BY_NATURE_FY2024)

_OPEX_BY_NATURE_FY2025 = [
    ("Raw materials and consumables", Decimal("34_600_000")),
    ("Employee benefits expense", Decimal("30_100_000")),
    ("Depreciation and amortisation", Decimal("6_800_000")),
    ("Impairment losses", Decimal("200_000")),
    ("Other operating expenses", Decimal("9_200_000")),
    ("Social charges (DE §1 SvEV)", Decimal("6_000_000")),
    ("Social charges (FR charges sociales)", Decimal("4_000_000")),
    ("Social charges (UK NIC)", Decimal("660_000")),
    ("Social charges (NL premies)", Decimal("530_000")),
]

# ── Balance sheet (statement of financial position) data ─────────────────────

_BS_FY2024 = {
    # Non-current assets
    "Property, plant and equipment": Decimal("42_300_000"),
    "Right-of-use assets (IFRS 16)": Decimal("8_400_000"),
    "Goodwill": Decimal("15_800_000"),
    "Intangible assets": Decimal("4_200_000"),
    "Deferred tax assets (IAS 12)": Decimal("1_800_000"),
    # Current assets
    "Inventories": Decimal("12_600_000"),
    "Trade and other receivables": Decimal("18_200_000"),
    "Cash and cash equivalents": Decimal("9_500_000"),
    # Equity
    "Share capital": Decimal("-1_000_000"),
    "Share premium": Decimal("-5_000_000"),
    "Retained earnings": Decimal("-45_200_000"),
    "Revaluation reserve (IAS 16)": Decimal("-3_200_000"),
    "Translation reserve (IAS 21)": Decimal("-800_000"),
    # Non-current liabilities
    "Long-term borrowings": Decimal("-18_000_000"),
    "Lease liabilities — non-current (IFRS 16)": Decimal("-6_100_000"),
    "Deferred tax liabilities (IAS 12)": Decimal("-2_400_000"),
    "Provisions": Decimal("-1_500_000"),
    # Current liabilities
    "Trade and other payables": Decimal("-14_800_000"),
    "Contract liabilities (IFRS 15)": Decimal("-3_200_000"),
    "Lease liabilities — current (IFRS 16)": Decimal("-2_300_000"),
    "Current tax liabilities": Decimal("-2_200_000"),
    "Accrued expenses": Decimal("-5_100_000"),
}

_TOTAL_ASSETS_2024 = sum(
    v for v in _BS_FY2024.values() if v > 0
)
_TOTAL_EQUITY_2024 = sum(
    v for k, v in _BS_FY2024.items()
    if k in ("Share capital", "Share premium", "Retained earnings",
             "Revaluation reserve (IAS 16)", "Translation reserve (IAS 21)")
)

_BS_FY2025 = {
    "Property, plant and equipment": Decimal("44_100_000"),
    "Right-of-use assets (IFRS 16)": Decimal("9_200_000"),
    "Goodwill": Decimal("15_800_000"),
    "Intangible assets": Decimal("3_900_000"),
    "Deferred tax assets (IAS 12)": Decimal("1_900_000"),
    "Inventories": Decimal("13_400_000"),
    "Trade and other receivables": Decimal("19_500_000"),
    "Cash and cash equivalents": Decimal("10_800_000"),
    "Share capital": Decimal("-1_000_000"),
    "Share premium": Decimal("-5_000_000"),
    "Retained earnings": Decimal("-49_800_000"),
    "Revaluation reserve (IAS 16)": Decimal("-3_400_000"),
    "Translation reserve (IAS 21)": Decimal("-700_000"),
    "Long-term borrowings": Decimal("-16_500_000"),
    "Lease liabilities — non-current (IFRS 16)": Decimal("-7_100_000"),
    "Deferred tax liabilities (IAS 12)": Decimal("-2_600_000"),
    "Provisions": Decimal("-1_600_000"),
    "Trade and other payables": Decimal("-15_600_000"),
    "Contract liabilities (IFRS 15)": Decimal("-3_500_000"),
    "Lease liabilities — current (IFRS 16)": Decimal("-2_500_000"),
    "Current tax liabilities": Decimal("-2_400_000"),
    "Accrued expenses": Decimal("-5_500_000"),
}

_TOTAL_ASSETS_2025 = sum(v for v in _BS_FY2025.values() if v > 0)

# ── Cash flow data (IAS 7 — interest in financing) ──────────────────────────

_CASH_ACCOUNTS_FY2024 = [
    ("CE", "ING Bank N.V. — EUR current", Decimal("2_800_000"), "EUR"),
    ("CP", "Commerzbank AG — EUR operating", Decimal("3_200_000"), "EUR"),
    ("CM", "BNP Paribas — EUR current", Decimal("1_900_000"), "EUR"),
    ("CD", "Barclays — GBP sterling", Decimal("1_600_000"), "GBP"),
]

_CASH_ACCOUNTS_FY2025 = [
    ("CE", "ING Bank N.V. — EUR current", Decimal("3_100_000"), "EUR"),
    ("CP", "Commerzbank AG — EUR operating", Decimal("3_600_000"), "EUR"),
    ("CM", "BNP Paribas — EUR current", Decimal("2_200_000"), "EUR"),
    ("CD", "Barclays — GBP sterling", Decimal("1_900_000"), "GBP"),
]

# ── Fixed assets by entity (IAS 16 with component depreciation) ─────────────

_FIXED_ASSETS_FY2024 = [
    {
        "asset_id": "CE-FA-001",
        "description": "Office fit-out Amsterdam HQ",
        "entity": "CE",
        "cost": Decimal("1_200_000"),
        "accum_depr": Decimal("480_000"),
        "revaluation": Decimal("0"),
        "nbv": Decimal("720_000"),
    },
    {
        "asset_id": "CP-FA-001",
        "description": "CNC machining centre — line A",
        "entity": "CP",
        "cost": Decimal("8_500_000"),
        "accum_depr": Decimal("3_400_000"),
        "revaluation": Decimal("1_800_000"),
        "nbv": Decimal("6_900_000"),
        "note": "IAS 16 revaluation model applied",
    },
    {
        "asset_id": "CP-FA-002",
        "description": "CNC machining centre — line B",
        "entity": "CP",
        "cost": Decimal("6_200_000"),
        "accum_depr": Decimal("2_480_000"),
        "revaluation": Decimal("1_400_000"),
        "nbv": Decimal("5_120_000"),
        "note": "IAS 16 revaluation model applied",
    },
    {
        "asset_id": "CP-FA-003",
        "description": "Quality testing laboratory",
        "entity": "CP",
        "cost": Decimal("2_800_000"),
        "accum_depr": Decimal("1_120_000"),
        "revaluation": Decimal("0"),
        "nbv": Decimal("1_680_000"),
    },
    {
        "asset_id": "CM-FA-001",
        "description": "R&D pilot plant Lyon",
        "entity": "CM",
        "cost": Decimal("5_400_000"),
        "accum_depr": Decimal("2_160_000"),
        "revaluation": Decimal("0"),
        "nbv": Decimal("3_240_000"),
    },
    {
        "asset_id": "CM-FA-002",
        "description": "Laboratory equipment (component depreciation — IAS 16)",
        "entity": "CM",
        "cost": Decimal("3_100_000"),
        "accum_depr": Decimal("1_550_000"),
        "revaluation": Decimal("0"),
        "nbv": Decimal("1_550_000"),
    },
    {
        "asset_id": "CD-FA-001",
        "description": "Warehouse racking system Birmingham",
        "entity": "CD",
        "cost": Decimal("1_800_000"),
        "accum_depr": Decimal("540_000"),
        "revaluation": Decimal("0"),
        "nbv": Decimal("1_260_000"),
    },
]

# ── IFRS 16 lease data (single lessee model) ────────────────────────────────

_LEASES_FY2024 = [
    {
        "lease_id": "EU-LS-001",
        "description": "CE Amsterdam office — Zuidas",
        "entity": "CE",
        "monthly_payment": Decimal("35_000"),
        "currency": "EUR",
        "start_date": "2021-01-01",
        "term_months": 120,
        "ibr": "EURIBOR + 1.5%",
        "rou_asset": Decimal("3_200_000"),
        "lease_liability": Decimal("2_900_000"),
    },
    {
        "lease_id": "EU-LS-002",
        "description": "CP Munich factory — Freiham",
        "entity": "CP",
        "monthly_payment": Decimal("85_000"),
        "currency": "EUR",
        "start_date": "2020-07-01",
        "term_months": 180,
        "ibr": "EURIBOR + 2.0%",
        "rou_asset": Decimal("10_200_000"),
        "lease_liability": Decimal("9_400_000"),
    },
    {
        "lease_id": "EU-LS-003",
        "description": "CM Lyon R&D centre — Gerland",
        "entity": "CM",
        "monthly_payment": Decimal("28_000"),
        "currency": "EUR",
        "start_date": "2022-04-01",
        "term_months": 96,
        "ibr": "EURIBOR + 1.8%",
        "rou_asset": Decimal("2_100_000"),
        "lease_liability": Decimal("1_900_000"),
    },
    {
        "lease_id": "EU-LS-004",
        "description": "CD Birmingham warehouse — Erdington",
        "entity": "CD",
        "monthly_payment": Decimal("18_000"),
        "currency": "GBP",
        "start_date": "2023-01-01",
        "term_months": 60,
        "ibr": "SONIA + 2.2%",
        "rou_asset": Decimal("950_000"),
        "lease_liability": Decimal("850_000"),
    },
]

# Two new leases for FY2025
_NEW_LEASES_FY2025 = [
    {
        "lease_id": "EU-LS-005",
        "description": "CP factory extension — Freiham (new wing)",
        "entity": "CP",
        "monthly_payment": Decimal("42_000"),
        "currency": "EUR",
        "start_date": "2025-03-01",
        "term_months": 120,
        "ibr": "EURIBOR + 1.8%",
        "rou_asset": Decimal("4_200_000"),
        "lease_liability": Decimal("4_200_000"),
    },
    {
        "lease_id": "EU-LS-006",
        "description": "CD additional warehouse — Solihull",
        "entity": "CD",
        "monthly_payment": Decimal("12_000"),
        "currency": "GBP",
        "start_date": "2025-06-01",
        "term_months": 72,
        "ibr": "SONIA + 2.0%",
        "rou_asset": Decimal("780_000"),
        "lease_liability": Decimal("780_000"),
    },
]

# ── IAS 36 CGU-based goodwill impairment data ──────────────────────────────

_GOODWILL_CGUS = [
    {
        "cgu": "Precision Manufacturing (CP)",
        "carrying_amount": Decimal("8_200_000"),
        "recoverable_amount": Decimal("9_800_000"),
        "headroom_pct": "19.5%",
        "impaired": "No",
        "discount_rate": "9.8% pre-tax",
        "notes": "Value-in-use based on 5-year management projections",
    },
    {
        "cgu": "Advanced Materials R&D (CM)",
        "carrying_amount": Decimal("5_400_000"),
        "recoverable_amount": Decimal("6_100_000"),
        "headroom_pct": "13.0%",
        "impaired": "No",
        "discount_rate": "10.2% pre-tax",
        "notes": "Value-in-use; IP pipeline supports future cash flows",
    },
    {
        "cgu": "Distribution Services (CD)",
        "carrying_amount": Decimal("2_200_000"),
        "recoverable_amount": Decimal("2_350_000"),
        "headroom_pct": "6.8%",
        "impaired": "No",
        "discount_rate": "11.0% pre-tax",
        "notes": "Close to threshold — Brexit-related uncertainty in cash flow projections",
    },
]

_TOTAL_GOODWILL = sum(c["carrying_amount"] for c in _GOODWILL_CGUS)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _whole_euros(d: Decimal | int) -> int:
    """Round a Decimal to whole euros."""
    if isinstance(d, int):
        return d
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


def _pin_xlsx_dates(wb: openpyxl.Workbook) -> None:
    """Pin created timestamp for determinism."""
    wb.properties.created = _FIXED_DATETIME


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    path = Path(path)
    buf = io.BytesIO()
    wb.properties.modified = _FIXED_DATETIME
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)
    buf = io.BytesIO()
    doc.save(buf)

    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=_FIXED_ZIP_DT)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


_HEADER_FILL = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
_HEADER_FONT = Font(bold=True, color="FFFFFF", size=10)
_BOLD_FONT = Font(bold=True, size=10)
_NORMAL_FONT = Font(size=10)
_THIN_BORDER = Border(
    bottom=Side(style="thin"),
    top=Side(style="thin"),
    left=Side(style="thin"),
    right=Side(style="thin"),
)
_NUMBER_FMT = '#,##0'
_EUR_FMT = '€#,##0'


def _write_xlsx_header(ws: Any, headers: list[str]) -> None:
    """Write header row with EU-style formatting."""
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.border = _THIN_BORDER


# ── Prior Year Workpapers (xlsx) ─────────────────────────────────────────────

def _write_wp_revenue(
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write wp_revenue_fy2024.xlsx — IFRS 15 revenue by entity with IC eliminations.

    Contains ERR-EU-018: CP product revenue shows FY2023 value (stale).
    """
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18eu_wp_revenue_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    # -- Revenue Summary sheet --
    ws = wb.active
    ws.title = "Revenue Summary"

    _write_xlsx_header(ws, [
        "Entity", "Revenue Stream", "FY2024 (EUR)", "IFRS 15 Category", "Notes",
    ])

    row = 2
    err_planted = False
    for entity, stream, amount in _REVENUE_FY2024:
        # ERR-EU-018: CP product revenue shows stale FY2023 value
        if not err_planted and entity == "CP" and "Product revenue" in stream:
            stale_amount = _whole_euros(stale_data(_REVENUE_FY2023_CP_PRODUCT))
            correct_amount = _whole_euros(amount)
            errors.add(PlantedError(
                error_id="ERR-EU-018",
                file=f"{_PY_DIR}/wp_revenue_fy2024.xlsx",
                location="Sheet 'Revenue Summary', CP (Germany) Product Revenue row, Column C",
                type="stale_data",
                description=(
                    f"CP product revenue shows €{stale_amount:,} "
                    f"(FY2023 value) instead of €{correct_amount:,} (FY2024 value)"
                ),
                severity="material",
                which_test_cases_should_catch=[_TC],
            ))
            ws.cell(row=row, column=3, value=stale_amount).number_format = _EUR_FMT
            err_planted = True
        else:
            ws.cell(row=row, column=3, value=_whole_euros(amount)).number_format = _EUR_FMT

        ws.cell(row=row, column=1, value=entity).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=stream).font = _NORMAL_FONT
        ws.cell(row=row, column=4, value="Revenue from contracts with customers").font = _NORMAL_FONT
        ws.cell(row=row, column=5, value="Per client TB").font = _NORMAL_FONT
        row += 1

    # Intercompany eliminations
    row += 1
    ws.cell(row=row, column=1, value="Intercompany Eliminations").font = _BOLD_FONT
    row += 1
    for ic_ref, desc, amount in _IC_ELIMINATIONS_REVENUE:
        ws.cell(row=row, column=1, value=ic_ref).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=desc).font = _NORMAL_FONT
        ws.cell(row=row, column=3, value=_whole_euros(amount)).number_format = _EUR_FMT
        ws.cell(row=row, column=4, value="Elimination").font = _NORMAL_FONT
        row += 1

    # Consolidated total
    row += 1
    ws.cell(row=row, column=2, value="Consolidated Revenue (IFRS 15)").font = _BOLD_FONT
    ws.cell(row=row, column=3, value=_whole_euros(_TOTAL_CONSOLIDATED_REVENUE_2024)).number_format = _EUR_FMT
    ws.cell(row=row, column=3).font = _BOLD_FONT

    row += 2
    ws.cell(row=row, column=1, value="Prepared by: Group Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT
    ws.cell(row=row + 2, column=1, value=(
        "Note: European number formatting convention — comma as decimal "
        "separator in source ledgers. All amounts converted to standard EUR presentation."
    )).font = _NORMAL_FONT

    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 38
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 38
    ws.column_dimensions["E"].width = 20

    rel_path = f"{_PY_DIR}/wp_revenue_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])


def _write_wp_operating_expenses(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write wp_operating_expenses_fy2024.xlsx — expenses by nature (IAS 1)."""
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18eu_wp_operating_expenses_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Operating Expenses by Nature"

    _write_xlsx_header(ws, [
        "Expense Category (IAS 1 — by nature)", "FY2024 (EUR)", "% of Total", "Notes",
    ])

    row = 2
    for category, amount in _OPEX_BY_NATURE_FY2024:
        pct = (amount / _TOTAL_OPEX_2024 * 100).quantize(Decimal("0.1"))
        ws.cell(row=row, column=1, value=category).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=_whole_euros(amount)).number_format = _EUR_FMT
        ws.cell(row=row, column=3, value=f"{pct}%").font = _NORMAL_FONT
        ws.cell(row=row, column=4, value="").font = _NORMAL_FONT
        row += 1

    # Total
    ws.cell(row=row, column=1, value="Total Operating Expenses").font = _BOLD_FONT
    ws.cell(row=row, column=2, value=_whole_euros(_TOTAL_OPEX_2024)).number_format = _EUR_FMT
    ws.cell(row=row, column=2).font = _BOLD_FONT

    row += 2
    ws.cell(row=row, column=1, value=(
        "Classification: IAS 1 by nature (not by function as under US GAAP). "
        "Social charges split by jurisdiction per local statutory requirements."
    )).font = _NORMAL_FONT

    row += 2
    ws.cell(row=row, column=1, value="Prepared by: Group Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT

    ws.column_dimensions["A"].width = 45
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 12
    ws.column_dimensions["D"].width = 30

    rel_path = f"{_PY_DIR}/wp_operating_expenses_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])


def _write_wp_balance_sheet(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write wp_balance_sheet_fy2024.xlsx — statement of financial position (IFRS)."""
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18eu_wp_balance_sheet_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Statement of Financial Position"

    _write_xlsx_header(ws, [
        "Line Item", "FY2024 (EUR)", "Classification", "IFRS Reference",
    ])

    row = 2
    # Group by classification
    categories = [
        ("Non-current assets", [
            "Property, plant and equipment",
            "Right-of-use assets (IFRS 16)",
            "Goodwill",
            "Intangible assets",
            "Deferred tax assets (IAS 12)",
        ]),
        ("Current assets", [
            "Inventories",
            "Trade and other receivables",
            "Cash and cash equivalents",
        ]),
        ("Equity", [
            "Share capital",
            "Share premium",
            "Retained earnings",
            "Revaluation reserve (IAS 16)",
            "Translation reserve (IAS 21)",
        ]),
        ("Non-current liabilities", [
            "Long-term borrowings",
            "Lease liabilities — non-current (IFRS 16)",
            "Deferred tax liabilities (IAS 12)",
            "Provisions",
        ]),
        ("Current liabilities", [
            "Trade and other payables",
            "Contract liabilities (IFRS 15)",
            "Lease liabilities — current (IFRS 16)",
            "Current tax liabilities",
            "Accrued expenses",
        ]),
    ]

    ifrs_refs = {
        "Property, plant and equipment": "IAS 16",
        "Right-of-use assets (IFRS 16)": "IFRS 16",
        "Goodwill": "IFRS 3 / IAS 36",
        "Intangible assets": "IAS 38",
        "Deferred tax assets (IAS 12)": "IAS 12",
        "Inventories": "IAS 2",
        "Trade and other receivables": "IFRS 9",
        "Cash and cash equivalents": "IAS 7",
        "Lease liabilities — non-current (IFRS 16)": "IFRS 16",
        "Lease liabilities — current (IFRS 16)": "IFRS 16",
        "Contract liabilities (IFRS 15)": "IFRS 15",
        "Deferred tax liabilities (IAS 12)": "IAS 12",
        "Provisions": "IAS 37",
        "Revaluation reserve (IAS 16)": "IAS 16.39",
        "Translation reserve (IAS 21)": "IAS 21.39",
    }

    for cat_name, items in categories:
        ws.cell(row=row, column=1, value=cat_name).font = _BOLD_FONT
        row += 1
        subtotal = Decimal("0")
        for item in items:
            bal = _BS_FY2024[item]
            subtotal += bal
            ws.cell(row=row, column=1, value=f"  {item}").font = _NORMAL_FONT
            ws.cell(row=row, column=2, value=_whole_euros(abs(bal))).number_format = _EUR_FMT
            ws.cell(row=row, column=3, value=cat_name).font = _NORMAL_FONT
            ws.cell(row=row, column=4, value=ifrs_refs.get(item, "")).font = _NORMAL_FONT
            row += 1
        ws.cell(row=row, column=1, value=f"Total {cat_name}").font = _BOLD_FONT
        ws.cell(row=row, column=2, value=_whole_euros(abs(subtotal))).number_format = _EUR_FMT
        ws.cell(row=row, column=2).font = _BOLD_FONT
        row += 1

    row += 1
    ws.cell(row=row, column=1, value="Total Assets").font = _BOLD_FONT
    ws.cell(row=row, column=2, value=_whole_euros(_TOTAL_ASSETS_2024)).number_format = _EUR_FMT
    ws.cell(row=row, column=2).font = _BOLD_FONT

    row += 2
    ws.cell(row=row, column=1, value=(
        "Note: 'Statement of financial position' per IFRS terminology "
        "(not 'balance sheet' as under US GAAP). CD amounts translated "
        f"at closing rate {_FX_CLOSING_2024} GBP/EUR per IAS 21."
    )).font = _NORMAL_FONT

    row += 2
    ws.cell(row=row, column=1, value="Prepared by: Group Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT

    ws.column_dimensions["A"].width = 45
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 25
    ws.column_dimensions["D"].width = 20

    rel_path = f"{_PY_DIR}/wp_balance_sheet_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])


def _write_wp_cash(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write wp_cash_fy2024.xlsx — cash reconciliation with multi-currency (IAS 7)."""
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18eu_wp_cash_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Cash Reconciliation"

    _write_xlsx_header(ws, [
        "Entity", "Bank Account", "FY2024 Balance (Local)", "Currency",
        "FX Rate", "FY2024 Balance (EUR)", "Bank Confirmed",
    ])

    row = 2
    total_eur = Decimal("0")
    for entity, account, balance, ccy in _CASH_ACCOUNTS_FY2024:
        ws.cell(row=row, column=1, value=entity).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=account).font = _NORMAL_FONT
        ws.cell(row=row, column=3, value=_whole_euros(balance)).number_format = _NUMBER_FMT
        ws.cell(row=row, column=4, value=ccy).font = _NORMAL_FONT

        if ccy == "GBP":
            eur_amount = (balance * _FX_CLOSING_2024).quantize(Decimal("1"), rounding=ROUND_HALF_UP)
            ws.cell(row=row, column=5, value=str(_FX_CLOSING_2024)).font = _NORMAL_FONT
        else:
            eur_amount = balance
            ws.cell(row=row, column=5, value="1.0000").font = _NORMAL_FONT

        ws.cell(row=row, column=6, value=_whole_euros(eur_amount)).number_format = _EUR_FMT
        ws.cell(row=row, column=7, value="Yes").font = _NORMAL_FONT
        total_eur += eur_amount
        row += 1

    row += 1
    ws.cell(row=row, column=5, value="Total (EUR)").font = _BOLD_FONT
    ws.cell(row=row, column=6, value=_whole_euros(total_eur)).number_format = _EUR_FMT
    ws.cell(row=row, column=6).font = _BOLD_FONT

    row += 2
    ws.cell(row=row, column=1, value=(
        "IAS 7 policy choice: Interest paid classified under financing activities "
        "(not operating as typical under US GAAP)."
    )).font = _NORMAL_FONT

    row += 2
    ws.cell(row=row, column=1, value="Prepared by: Group Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT

    for i, w in enumerate([8, 35, 20, 10, 10, 20, 15], 1):
        ws.column_dimensions[chr(64 + i)].width = w

    rel_path = f"{_PY_DIR}/wp_cash_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])


def _write_wp_fixed_assets(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write wp_fixed_assets_fy2024.xlsx — PP&E with IAS 16 revaluation model."""
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18eu_wp_fixed_assets_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "Property, Plant and Equipment"

    _write_xlsx_header(ws, [
        "Asset ID", "Description", "Entity", "Cost (EUR)",
        "Accum Depr (EUR)", "Revaluation (EUR)", "NBV (EUR)", "Notes",
    ])

    row = 2
    for asset in _FIXED_ASSETS_FY2024:
        ws.cell(row=row, column=1, value=asset["asset_id"]).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=asset["description"]).font = _NORMAL_FONT
        ws.cell(row=row, column=3, value=asset["entity"]).font = _NORMAL_FONT
        ws.cell(row=row, column=4, value=_whole_euros(asset["cost"])).number_format = _EUR_FMT
        ws.cell(row=row, column=5, value=_whole_euros(asset["accum_depr"])).number_format = _EUR_FMT
        ws.cell(row=row, column=6, value=_whole_euros(asset["revaluation"])).number_format = _EUR_FMT
        ws.cell(row=row, column=7, value=_whole_euros(asset["nbv"])).number_format = _EUR_FMT
        ws.cell(row=row, column=8, value=asset.get("note", "")).font = _NORMAL_FONT
        row += 1

    # Totals
    total_cost = sum(a["cost"] for a in _FIXED_ASSETS_FY2024)
    total_depr = sum(a["accum_depr"] for a in _FIXED_ASSETS_FY2024)
    total_reval = sum(a["revaluation"] for a in _FIXED_ASSETS_FY2024)
    total_nbv = sum(a["nbv"] for a in _FIXED_ASSETS_FY2024)
    row += 1
    ws.cell(row=row, column=2, value="Total PP&E").font = _BOLD_FONT
    ws.cell(row=row, column=4, value=_whole_euros(total_cost)).number_format = _EUR_FMT
    ws.cell(row=row, column=5, value=_whole_euros(total_depr)).number_format = _EUR_FMT
    ws.cell(row=row, column=6, value=_whole_euros(total_reval)).number_format = _EUR_FMT
    ws.cell(row=row, column=7, value=_whole_euros(total_nbv)).number_format = _EUR_FMT

    row += 2
    ws.cell(row=row, column=1, value=(
        "Component depreciation applied per IAS 16. CP manufacturing equipment "
        "uses the revaluation model (IAS 16.31) — not permitted under US GAAP."
    )).font = _NORMAL_FONT

    row += 2
    ws.cell(row=row, column=1, value="Prepared by: Group Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT

    for i, w in enumerate([12, 42, 8, 14, 16, 16, 14, 35], 1):
        ws.column_dimensions[chr(64 + i)].width = w

    rel_path = f"{_PY_DIR}/wp_fixed_assets_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])


def _write_wp_leases(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write wp_leases_fy2024.xlsx — IFRS 16 lease schedule (single model)."""
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18eu_wp_leases_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "IFRS 16 Lease Schedule"

    _write_xlsx_header(ws, [
        "Lease ID", "Description", "Entity", "Monthly Payment",
        "Currency", "Start Date", "Term (months)", "IBR",
        "ROU Asset (EUR)", "Lease Liability (EUR)",
    ])

    row = 2
    for lease in _LEASES_FY2024:
        ws.cell(row=row, column=1, value=lease["lease_id"]).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=lease["description"]).font = _NORMAL_FONT
        ws.cell(row=row, column=3, value=lease["entity"]).font = _NORMAL_FONT
        ws.cell(row=row, column=4, value=_whole_euros(lease["monthly_payment"])).number_format = _NUMBER_FMT
        ws.cell(row=row, column=5, value=lease["currency"]).font = _NORMAL_FONT
        ws.cell(row=row, column=6, value=lease["start_date"]).font = _NORMAL_FONT
        ws.cell(row=row, column=7, value=lease["term_months"]).font = _NORMAL_FONT
        ws.cell(row=row, column=8, value=lease["ibr"]).font = _NORMAL_FONT
        ws.cell(row=row, column=9, value=_whole_euros(lease["rou_asset"])).number_format = _EUR_FMT
        ws.cell(row=row, column=10, value=_whole_euros(lease["lease_liability"])).number_format = _EUR_FMT
        row += 1

    row += 1
    ws.cell(row=row, column=1, value=(
        "IFRS 16 single lessee model — no operating/finance lease distinction "
        "(unlike ASC 842). Low-value and short-term exemptions applied where applicable."
    )).font = _NORMAL_FONT

    row += 2
    ws.cell(row=row, column=1, value="Prepared by: Group Audit Staff").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value="Date: March 2025").font = _NORMAL_FONT

    for i, w in enumerate([12, 38, 8, 16, 10, 12, 14, 16, 16, 18], 1):
        ws.column_dimensions[chr(64 + i)].width = w

    rel_path = f"{_PY_DIR}/wp_leases_fy2024.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])


# ── Prior Year Workpapers (docx memos) ──────────────────────────────────────

def _write_memo_planning(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write memo_planning_fy2024.docx — ISA 300 audit planning memorandum.

    One of 2 files requiring MANAGER JUDGMENT (substantive rewriting needed).
    """
    doc = docx.Document()
    key = "tc18eu_memo_planning_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_docx(doc, canary)

    doc.add_heading("Cascade Europe Holdings B.V. — Group", level=1)
    doc.add_heading("ISA 300 Audit Planning Memorandum — FY2024", level=2)

    doc.add_paragraph(
        "This memorandum documents the planned audit approach for the consolidated "
        "financial statements of Cascade Europe Holdings B.V. group for the fiscal year "
        "ended 31 December 2024, prepared in accordance with IFRS as adopted by the EU."
    )

    doc.add_heading("1. Group Overview", level=3)
    doc.add_paragraph(
        f"Cascade Europe Holdings B.V. is the European holding company of the Cascade "
        f"group, with consolidated revenue of approximately €{_whole_euros(_TOTAL_CONSOLIDATED_REVENUE_2024):,} "
        f"for FY2024. The group comprises four entities across four jurisdictions:"
    )
    for code, info in sorted(_ENTITIES.items()):
        doc.add_paragraph(
            f"• {info['name']} ({code}) — {info['city']}, {info['jurisdiction']} — {info['role']}",
        )

    doc.add_heading("2. ISA 315 Risk Assessment", level=3)
    doc.add_paragraph(
        "Significant risks identified per ISA 315 (Revised 2019):"
    )
    doc.add_paragraph("• Revenue recognition — IFRS 15 disaggregation across multi-entity group")
    doc.add_paragraph("• Intercompany eliminations — material IC transactions (CP→CM, CP→CD, CE management fees)")
    doc.add_paragraph("• Lease accounting (IFRS 16) — single lessee model across 4 jurisdictions")
    doc.add_paragraph("• IAS 21 FX translation — CD operates in GBP, consolidated in EUR")
    doc.add_paragraph("• Management estimates — IAS 36 impairment, useful lives, provisions (IAS 37)")

    doc.add_heading("3. ISA 600 Group Audit Considerations", level=3)
    doc.add_paragraph(
        "Component auditors engaged for the following subsidiaries:"
    )
    doc.add_paragraph("• CP (Germany) — BDO München, full-scope audit")
    doc.add_paragraph("• CM (France) — Mazars Lyon, full-scope audit")
    doc.add_paragraph("• CD (United Kingdom) — RSM UK, full-scope audit")
    doc.add_paragraph(
        "Group engagement team (Amsterdam) performs the CE audit and consolidation "
        "procedures. Component auditor instructions issued per ISA 600.40."
    )

    doc.add_heading("4. Materiality", level=3)
    materiality = _whole_euros(_TOTAL_CONSOLIDATED_REVENUE_2024 * Decimal("0.01"))
    perf_mat = _whole_euros(Decimal(str(materiality)) * Decimal("0.75"))
    doc.add_paragraph(
        f"Planning materiality set at 1% of consolidated revenue = €{materiality:,}. "
        f"Performance materiality at 75% = €{perf_mat:,}. "
        f"Component materiality allocated based on entity significance."
    )

    doc.add_heading("5. ISA 330 Audit Responses", level=3)
    doc.add_paragraph("• Substantive analytical procedures on revenue by entity and segment")
    doc.add_paragraph("• Tests of controls over intercompany elimination process")
    doc.add_paragraph("• IFRS 16 lease recalculation for new and modified leases")
    doc.add_paragraph("• IAS 21 translation testing: closing rate for balance sheet, average for P&L")
    doc.add_paragraph("• Management representation letters from all component management teams")

    doc.add_heading("6. Audit Timeline", level=3)
    doc.add_paragraph("• Component interim procedures: October 2024")
    doc.add_paragraph("• Group year-end fieldwork: February–March 2025")
    doc.add_paragraph("• Component reporting packages due: 15 March 2025")
    doc.add_paragraph("• Group audit completion: April 2025")

    doc.add_heading("7. Team Assignment", level=3)
    doc.add_paragraph("• Group Engagement Partner: M. van der Berg")
    doc.add_paragraph("• Group Senior Manager: A. Müller")
    doc.add_paragraph("• Component liaison: K. Lefèvre (France), J. Williams (UK)")

    p = doc.add_paragraph()
    p.add_run("\nPrepared: March 2025").italic = True

    rel_path = f"{_PY_DIR}/memo_planning_fy2024.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, full_path)

    canaries.set_location(key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary, test_cases=[_TC])


def _write_memo_risk_assessment(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write memo_risk_assessment_fy2024.docx — ISA 315 risk assessment."""
    doc = docx.Document()
    key = "tc18eu_memo_risk_assessment_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_docx(doc, canary)

    doc.add_heading("Cascade Europe Holdings B.V. — Group", level=1)
    doc.add_heading("ISA 315 Risk Assessment — FY2024 Audit", level=2)

    doc.add_heading("1. Fraud Risk Assessment", level=3)
    doc.add_paragraph(
        "Management override of controls assessed as present in all group entities "
        "per ISA 240. Revenue recognition presumed fraud risk per ISA 240.26. "
        "No entity-specific fraud indicators identified."
    )

    doc.add_heading("2. Significant Risks (ISA 315)", level=3)
    doc.add_paragraph("• Revenue recognition: IFRS 15 principal vs agent for intercompany sales")
    doc.add_paragraph("• Related party / intercompany: IC eliminations and transfer pricing (OECD)")
    doc.add_paragraph("• Management estimates: IAS 36 goodwill impairment — CGU allocation judgment")
    doc.add_paragraph("• IFRS 16 lease modifications: New leases and modification accounting across entities")
    doc.add_paragraph("• IAS 12 deferred tax: Temporary differences on intercompany unrealised profits")
    doc.add_paragraph("• IAS 21 FX translation: GBP→EUR translation for CD subsidiary")

    doc.add_heading("3. Industry and Jurisdiction Risk Factors", level=3)
    doc.add_paragraph(
        "Manufacturing sector in EU — subject to energy cost volatility and supply chain "
        "constraints. Customer concentration moderate (top 10 represent ~40% of group revenue). "
        "Brexit trade friction affects CD margins. German labor regulations (BetrVG) apply to CP. "
        "French social charges significantly impact CM cost base."
    )

    doc.add_heading("4. IT General Controls", level=3)
    doc.add_paragraph(
        "Group ERP: SAP S/4HANA (CE, CP). CM uses Sage X3 (France). CD uses Xero (UK). "
        "IT audit performed in Q2 2024 — no significant findings across group systems. "
        "Data consolidation performed centrally in Amsterdam using SAP BPC."
    )

    p = doc.add_paragraph()
    p.add_run("\nPrepared: March 2025").italic = True

    rel_path = f"{_PY_DIR}/memo_risk_assessment_fy2024.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, full_path)

    canaries.set_location(key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary, test_cases=[_TC])


def _write_memo_summary(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write memo_summary_fy2024.docx — ISA 220 audit completion summary."""
    doc = docx.Document()
    key = "tc18eu_memo_summary_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_docx(doc, canary)

    doc.add_heading("Cascade Europe Holdings B.V. — Group", level=1)
    doc.add_heading("Audit Completion Summary (ISA 220) — FY2024", level=2)

    doc.add_paragraph(
        "This memorandum summarises the results of audit procedures performed for "
        "the consolidated financial statements for the fiscal year ended 31 December 2024."
    )

    doc.add_heading("1. Financial Highlights", level=3)
    doc.add_paragraph(f"• Consolidated Revenue: €{_whole_euros(_TOTAL_CONSOLIDATED_REVENUE_2024):,}")
    net_income_2024 = _whole_euros(
        _TOTAL_CONSOLIDATED_REVENUE_2024 - _TOTAL_OPEX_2024
    )
    doc.add_paragraph(f"• Operating Result: €{net_income_2024:,}")
    doc.add_paragraph(f"• Total Assets: €{_whole_euros(_TOTAL_ASSETS_2024):,}")

    doc.add_heading("2. Audit Findings", level=3)
    doc.add_paragraph(
        "No material misstatements identified. Two immaterial adjustments proposed: "
        "(1) reclassification of €18,500 prepaid expense at CP; "
        "(2) FX translation adjustment of €7,200 at CD. Both accepted by management."
    )

    doc.add_heading("3. Key Audit Matters (ISA 701)", level=3)
    doc.add_paragraph(
        "KAM 1: Revenue recognition across multi-entity group — IFRS 15 disaggregation "
        "and intercompany elimination required significant auditor judgment."
    )
    doc.add_paragraph(
        "KAM 2: Goodwill allocated to cash-generating units — IAS 36 impairment testing "
        "of €15.8M goodwill required evaluation of management's assumptions, particularly "
        "for the CD CGU where headroom is limited."
    )

    doc.add_heading("4. Going Concern Assessment", level=3)
    doc.add_paragraph(
        "No indicators of going concern identified across group entities. All entities "
        "maintain adequate liquidity. Group banking facilities renewed through 2027."
    )

    doc.add_heading("5. Component Auditor Communications (ISA 600)", level=3)
    doc.add_paragraph(
        "All component auditor reporting packages received timely. No significant "
        "matters raised by BDO München (CP), Mazars Lyon (CM), or RSM UK (CD). "
        "Group engagement partner reviewed all component working papers."
    )

    doc.add_heading("6. Subsequent Events", level=3)
    doc.add_paragraph(
        "No significant subsequent events identified through 15 April 2025."
    )

    p = doc.add_paragraph()
    p.add_run("\nPrepared: March 2025").italic = True

    rel_path = f"{_PY_DIR}/memo_summary_fy2024.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, full_path)

    canaries.set_location(key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary, test_cases=[_TC])


def _write_memo_management_letter(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write memo_management_letter_fy2024.docx — management letter to B.V. board."""
    doc = docx.Document()
    key = "tc18eu_memo_management_letter_fy2024"
    canary = canaries.canary_for(key)
    embed_canary_docx(doc, canary)

    doc.add_heading("Cascade Europe Holdings B.V.", level=1)
    doc.add_heading("Management Letter — FY2024", level=2)

    doc.add_paragraph(
        "During our audit of the consolidated financial statements for the year ended "
        "31 December 2024, we noted certain matters involving internal control and IFRS "
        "compliance that we wish to bring to the attention of the board."
    )

    doc.add_heading("Finding 1: Transfer Pricing Documentation", level=3)
    doc.add_paragraph(
        "The OECD Master File documentation for FY2024 has not been completed. "
        "Under Dutch APA/ATR regulations and German §1 AStG, contemporaneous "
        "documentation is required. We recommend finalising the Master File before "
        "the filing deadline."
    )
    doc.add_paragraph("Management Response: Will complete by Q2 2025.", style="Intense Quote")

    doc.add_heading("Finding 2: IFRS 16 Lease Modifications", level=3)
    doc.add_paragraph(
        "Two lease modifications at CP were processed as new leases rather than "
        "modifications under IFRS 16.44. The impact is immaterial (€32,000 ROU asset "
        "difference) but the policy should be clarified to prevent future errors."
    )
    doc.add_paragraph("Management Response: Accounting policy updated.", style="Intense Quote")

    doc.add_heading("Finding 3: Intercompany Reconciliation Timing", level=3)
    doc.add_paragraph(
        "Monthly intercompany reconciliations between CP and CD show timing differences "
        "of up to 15 days due to different cut-off practices. We recommend synchronising "
        "month-end close procedures across entities."
    )
    doc.add_paragraph(
        "Management Response: Group finance will implement synchronised close calendar "
        "from Q1 2025.",
        style="Intense Quote",
    )

    doc.add_heading("Finding 4: GDPR Data Processing Agreements", level=3)
    doc.add_paragraph(
        "Two data processors used by CP do not have executed Art. 28 GDPR Data Processing "
        "Agreements. While this is primarily a compliance matter rather than a financial "
        "reporting issue, we flag it for board attention given potential regulatory fines."
    )
    doc.add_paragraph("Management Response: DPAs being negotiated.", style="Intense Quote")

    p = doc.add_paragraph()
    p.add_run("\nPrepared: March 2025").italic = True

    rel_path = f"{_PY_DIR}/memo_management_letter_fy2024.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, full_path)

    canaries.set_location(key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary, test_cases=[_TC])


# ── Current Year Data ────────────────────────────────────────────────────────

def _write_cy_trial_balance(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write trial_balance_fy2025.csv — FORMAT CHANGE from xlsx.

    IFRS account descriptions used (right-of-use assets, contract liabilities, etc.).
    Multi-entity consolidated with IC eliminations.
    """
    key = "tc18eu_cy_trial_balance_fy2025"
    canary = canaries.canary_for(key)

    rel_path = f"{_CY_DIR}/trial_balance_fy2025.csv"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    with open(full_path, "w", newline="") as f:
        f.write(embed_canary_csv_comment(canary))
        writer = csv.writer(f)
        writer.writerow([
            "Account", "Account Description (IFRS)", "Entity", "Debit (EUR)", "Credit (EUR)",
        ])

        # Balance sheet items (FY2025)
        for item, bal in sorted(_BS_FY2025.items()):
            val = _whole_euros(abs(bal))
            if bal >= 0:
                writer.writerow(["", item, "Group", val, ""])
            else:
                writer.writerow(["", item, "Group", "", val])

        # Revenue items (FY2025) — credit balances
        for entity, stream, amount in sorted(_REVENUE_FY2025, key=lambda r: (r[0], r[1])):
            writer.writerow(["", stream, entity, "", _whole_euros(amount)])

        # IC eliminations
        for ic_ref, desc, amount in sorted(_IC_ELIMINATIONS_REVENUE_2025, key=lambda r: r[0]):
            val = _whole_euros(abs(amount))
            writer.writerow(["", f"IC elimination: {desc}", ic_ref, val, ""])

        # Operating expenses (FY2025) — debit balances
        for category, amount in sorted(_OPEX_BY_NATURE_FY2025, key=lambda r: r[0]):
            writer.writerow(["", category, "Group", _whole_euros(amount), ""])

    canaries.set_location(key, rel_path, "CSV comment line 1")
    manifest.register(rel_path, "csv", canary=canary, test_cases=[_TC])


def _write_cy_bank_statements(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write bank_statements_fy2025.csv — multi-currency (EUR + GBP)."""
    key = "tc18eu_cy_bank_statements_fy2025"
    canary = canaries.canary_for(key)

    rel_path = f"{_CY_DIR}/bank_statements_fy2025.csv"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    with open(full_path, "w", newline="") as f:
        f.write(embed_canary_csv_comment(canary))
        writer = csv.writer(f)
        writer.writerow([
            "Entity", "Bank Account", "Statement Balance", "Currency", "Date",
        ])

        for entity, account, balance, ccy in sorted(_CASH_ACCOUNTS_FY2025, key=lambda r: r[0]):
            writer.writerow([
                entity, account, _whole_euros(balance), ccy, "2025-12-31",
            ])

    canaries.set_location(key, rel_path, "CSV comment line 1")
    manifest.register(rel_path, "csv", canary=canary, test_cases=[_TC])


def _write_cy_lease_schedule(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write lease_schedule_fy2025.xlsx — IFRS 16 with 2 new leases added."""
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18eu_cy_lease_schedule_fy2025"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "IFRS 16 Lease Schedule FY2025"

    _write_xlsx_header(ws, [
        "Lease ID", "Description", "Entity", "Monthly Payment",
        "Currency", "Start Date", "Term (months)", "IBR",
        "ROU Asset (EUR)", "Lease Liability (EUR)", "Status",
    ])

    row = 2
    all_leases = _LEASES_FY2024 + _NEW_LEASES_FY2025
    for lease in sorted(all_leases, key=lambda le: le["lease_id"]):
        ws.cell(row=row, column=1, value=lease["lease_id"]).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=lease["description"]).font = _NORMAL_FONT
        ws.cell(row=row, column=3, value=lease["entity"]).font = _NORMAL_FONT
        ws.cell(row=row, column=4, value=_whole_euros(lease["monthly_payment"])).number_format = _NUMBER_FMT
        ws.cell(row=row, column=5, value=lease["currency"]).font = _NORMAL_FONT
        ws.cell(row=row, column=6, value=lease["start_date"]).font = _NORMAL_FONT
        ws.cell(row=row, column=7, value=lease["term_months"]).font = _NORMAL_FONT
        ws.cell(row=row, column=8, value=lease["ibr"]).font = _NORMAL_FONT
        ws.cell(row=row, column=9, value=_whole_euros(lease["rou_asset"])).number_format = _EUR_FMT
        ws.cell(row=row, column=10, value=_whole_euros(lease["lease_liability"])).number_format = _EUR_FMT
        is_new = lease["lease_id"] in ("EU-LS-005", "EU-LS-006")
        ws.cell(row=row, column=11, value="NEW" if is_new else "Existing").font = _NORMAL_FONT
        row += 1

    for i, w in enumerate([12, 38, 8, 16, 10, 12, 14, 16, 16, 18, 10], 1):
        ws.column_dimensions[chr(64 + i)].width = w

    rel_path = f"{_CY_DIR}/lease_schedule_fy2025.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])


def _write_cy_mgmt_projections(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write management_projections_fy2025.docx — FORMAT CHANGE from xlsx.

    Includes Pillar Two impact assessment section (judgment trap: group revenue
    ~€120M is below €750M threshold, so Pillar Two does NOT apply).
    One of 2 files the agent should FLAG for manager attention.
    """
    doc = docx.Document()
    key = "tc18eu_cy_mgmt_projections_fy2025"
    canary = canaries.canary_for(key)
    embed_canary_docx(doc, canary)

    doc.add_heading("Cascade Europe Holdings B.V.", level=1)
    doc.add_heading("Management Financial Projections — FY2026", level=2)

    doc.add_paragraph(
        "The following projections reflect management's expectations for the fiscal year "
        "ending 31 December 2026, based on current market conditions and strategic initiatives "
        "across the Cascade Europe group."
    )

    total_rev_2025 = _whole_euros(_TOTAL_CONSOLIDATED_REVENUE_2025)
    projected_rev = _whole_euros(Decimal(str(total_rev_2025)) * Decimal("1.07"))

    doc.add_heading("Revenue Projections", level=3)
    doc.add_paragraph(
        f"FY2025 Actual Consolidated Revenue: €{total_rev_2025:,}\n"
        f"FY2026 Projected Consolidated Revenue: €{projected_rev:,} (7% growth)\n\n"
        "Growth drivers by entity:\n"
        "• CP (Germany): Expected 6% growth from new automotive OEM contracts\n"
        "• CM (France): 8% growth driven by aerospace materials R&D pipeline\n"
        "• CD (UK): 5% growth — cautious outlook due to sterling volatility\n"
        "• CE (Netherlands): Management fees aligned to subsidiary revenue growth"
    )

    doc.add_heading("Expense Projections", level=3)
    doc.add_paragraph(
        "Operating expenses expected to increase 5% due to:\n"
        "• Headcount additions: 35 new hires across group (15 CP, 10 CM, 5 CD, 5 CE)\n"
        "• German collective bargaining increase: ~3.2% (IG Metall)\n"
        "• French social charges inflation: ~2.5%\n"
        "• Raw material cost pressure: ~3% (steel, composites)\n"
        "• New IFRS 16 lease commitments (CP factory extension, CD warehouse)"
    )

    doc.add_heading("Capital Expenditures", level=3)
    doc.add_paragraph(
        "Planned CapEx of €6.2M:\n"
        "• CP: CNC upgrade line C — €3.1M\n"
        "• CM: Laboratory expansion — €2.0M\n"
        "• CD: Warehouse automation — £0.9M (~€1.0M at projected rates)\n"
        "• CE: IT infrastructure refresh — €0.1M"
    )

    doc.add_heading("OECD Pillar Two Impact Assessment", level=3)
    doc.add_paragraph(
        f"Group consolidated revenue for FY2025: €{total_rev_2025:,}.\n\n"
        "The OECD Pillar Two / GloBE rules apply to multinational groups with "
        "consolidated revenue exceeding €750 million in at least two of the four "
        "preceding fiscal years. Cascade Europe group revenue is significantly below "
        "this threshold.\n\n"
        "Assessment: Pillar Two does NOT currently apply to the group. No top-up tax "
        "calculations or Qualified Domestic Minimum Top-up Tax (QDMTT) analysis required. "
        "Management will continue to monitor revenue trajectory for future threshold assessment."
    )

    doc.add_heading("Key Assumptions", level=3)
    doc.add_paragraph(
        "• No significant changes in customer base across jurisdictions\n"
        "• EUR/GBP exchange rate assumed at 1.15 (stable)\n"
        "• No major acquisitions planned in FY2026\n"
        "• IAS 36 goodwill impairment testing scheduled for Q4 2026 (annual)\n"
        "• Interest rates: ECB deposit rate stable at current levels; BOE rate declining 25bps"
    )

    p = doc.add_paragraph()
    p.add_run("\nPrepared by: Group CFO Office, January 2026").italic = True

    rel_path = f"{_CY_DIR}/management_projections_fy2025.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_deterministic(doc, full_path)

    canaries.set_location(key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary, test_cases=[_TC])


def _write_cy_goodwill_impairment(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    """Write goodwill_impairment_analysis_ifrs.xlsx — IAS 36 CGU-based.

    NEW file not present in prior year. Agent should flag as significant new audit area.
    Uses pre-tax discount rates per IAS 36 (vs post-tax common in practice).
    """
    wb = openpyxl.Workbook()
    _pin_xlsx_dates(wb)
    key = "tc18eu_cy_goodwill_impairment_ifrs_fy2025"
    canary = canaries.canary_for(key)
    embed_canary_xlsx(wb, canary)

    ws = wb.active
    ws.title = "IAS 36 Goodwill Impairment"

    _write_xlsx_header(ws, [
        "Cash-Generating Unit (CGU)", "Carrying Amount (EUR)",
        "Recoverable Amount (EUR)", "Headroom %",
        "Impairment Indicated", "Discount Rate", "Notes",
    ])

    row = 2
    for cgu in _GOODWILL_CGUS:
        ws.cell(row=row, column=1, value=cgu["cgu"]).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=_whole_euros(cgu["carrying_amount"])).number_format = _EUR_FMT
        ws.cell(row=row, column=3, value=_whole_euros(cgu["recoverable_amount"])).number_format = _EUR_FMT
        ws.cell(row=row, column=4, value=cgu["headroom_pct"]).font = _NORMAL_FONT
        ws.cell(row=row, column=5, value=cgu["impaired"]).font = _NORMAL_FONT
        ws.cell(row=row, column=6, value=cgu["discount_rate"]).font = _NORMAL_FONT
        ws.cell(row=row, column=7, value=cgu["notes"]).font = _NORMAL_FONT
        row += 1

    # Total
    row += 1
    ws.cell(row=row, column=1, value="Total Goodwill").font = _BOLD_FONT
    ws.cell(row=row, column=2, value=_whole_euros(_TOTAL_GOODWILL)).number_format = _EUR_FMT
    ws.cell(row=row, column=2).font = _BOLD_FONT

    row += 2
    ws.cell(row=row, column=1, value="Methodology: Value-in-use (IAS 36.30)").font = _NORMAL_FONT
    ws.cell(row=row + 1, column=1, value=(
        "Discount rates: Pre-tax rates per IAS 36.BCZ85 "
        "(post-tax discount rates are common practice but IAS 36 requires pre-tax)"
    )).font = _NORMAL_FONT
    ws.cell(
        row=row + 2, column=1,
        value="Terminal growth rate: 2.0% (aligned with long-term EU GDP forecast)",
    ).font = _NORMAL_FONT
    ws.cell(
        row=row + 3, column=1,
        value="Projection period: 5 years based on management budgets and forecasts",
    ).font = _NORMAL_FONT
    ws.cell(row=row + 4, column=1, value="Prepared by: Group Management, December 2025").font = _NORMAL_FONT

    ws.column_dimensions["A"].width = 35
    ws.column_dimensions["B"].width = 22
    ws.column_dimensions["C"].width = 22
    ws.column_dimensions["D"].width = 14
    ws.column_dimensions["E"].width = 18
    ws.column_dimensions["F"].width = 18
    ws.column_dimensions["G"].width = 55

    rel_path = f"{_CY_DIR}/goodwill_impairment_analysis_ifrs.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    _save_xlsx_deterministic(wb, full_path)

    canaries.set_location(key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary, test_cases=[_TC])


# ── Prompt & expected behavior ───────────────────────────────────────────────

def _write_prompt(output_dir: Path) -> None:
    """Write the TC-18-EU prompt.md."""
    text = """\
# TC-18-EU: IFRS/ISA Prior-Year Workpaper Rollforward

## Input Files

### Prior Year Workpapers (`prior_year_workpapers/`)
- `wp_revenue_fy2024.xlsx` — IFRS 15 revenue analysis by entity (with IC eliminations)
- `wp_operating_expenses_fy2024.xlsx` — Operating expenses by nature (IAS 1)
- `wp_balance_sheet_fy2024.xlsx` — Consolidated statement of financial position
- `wp_cash_fy2024.xlsx` — Cash reconciliation with multi-currency (IAS 7)
- `wp_fixed_assets_fy2024.xlsx` — PP&E with IAS 16 revaluation and component depreciation
- `wp_leases_fy2024.xlsx` — IFRS 16 lease schedule (single lessee model)
- `memo_planning_fy2024.docx` — ISA 300 audit planning memorandum
- `memo_risk_assessment_fy2024.docx` — ISA 315 risk assessment
- `memo_summary_fy2024.docx` — ISA 220 audit completion summary
- `memo_management_letter_fy2024.docx` — Management letter to B.V. board

### Current Year Data (`current_year_data/`)
- `trial_balance_fy2025.csv` — **Note: format changed from xlsx to CSV**
- `bank_statements_fy2025.csv` — Multi-currency bank statement data (EUR + GBP)
- `lease_schedule_fy2025.xlsx` — Updated IFRS 16 schedule (includes 2 new leases)
- `management_projections_fy2025.docx` — **Note: format changed from xlsx to docx**
- `goodwill_impairment_analysis_ifrs.xlsx` — **New file not present in prior year** (IAS 36)

## Instructions

Roll forward the FY2024 IFRS audit workpapers to FY2025 using the current year
data for Cascade Europe Holdings B.V. group.

1. For each prior year workpaper, identify which data needs to be updated with
   FY2025 figures.
2. Map the current year data files to the corresponding prior year workpapers,
   noting any format changes (e.g., CSV vs. xlsx).
3. Update all numerical data with current year figures, ensuring IFRS
   presentation is preserved (statement of financial position, right-of-use
   assets, contract liabilities, etc.).
4. Preserve the workpaper structure and any ISA-referenced commentary that is
   still relevant.
5. Flag the following for the group engagement partner's attention:
   - Any structural changes in the client's data (new accounts, format changes,
     renamed fields)
   - The new IAS 36 goodwill impairment analysis file (not present in prior
     year — indicates a new significant risk area for the group audit)
   - The Pillar Two impact assessment in management projections (assess
     whether the group meets the €750M revenue threshold)
   - Any areas where prior year ISA 315 risk assessment commentary may no
     longer be applicable given current year developments
   - FX translation impacts from the UK subsidiary's GBP→EUR conversion
6. Update the ISA 300 planning memo with current year scope considerations,
   including any changes to component auditor arrangements.

Export the rolled-forward workpapers to a new folder.
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write the TC-18-EU expected_behavior.md."""
    text = """\
# TC-18-EU: Expected Behavior

## Key Evaluation Criteria

### Workpaper Update Success (8 of 10)
The agent should successfully update 8 of 10 prior year workpapers:
- **Revenue workpaper** — Update with FY2025 revenue data from CSV trial balance;
  preserve IFRS 15 disaggregation and intercompany elimination structure
- **Operating expenses workpaper** — Update with FY2025 data; maintain IAS 1
  classification by nature (not by function)
- **Balance sheet workpaper** — Update statement of financial position with FY2025
  figures; maintain IFRS terminology throughout
- **Cash workpaper** — Update with FY2025 bank statement data; handle multi-currency
  reconciliation (EUR + GBP)
- **Fixed assets workpaper** — Update with FY2025 PP&E data; preserve IAS 16
  revaluation reserve for CP equipment
- **Lease workpaper** — Update with FY2025 IFRS 16 schedule including 2 new leases
  (CP factory extension, CD warehouse)
- **Risk assessment memo** — Mechanical updates (dates, figures); preserve ISA
  references
- **Management letter memo** — Update status of prior findings

### Files Requiring Manager Judgment (2 of 10)
- **Management projections** — Format changed from xlsx to docx; includes Pillar Two
  section that requires threshold assessment (group revenue ~€120M < €750M threshold
  → Pillar Two does NOT apply). Agent should FLAG this for manager attention.
- **Planning memo (ISA 300)** — Requires substantive rewriting for FY2025 scope:
  new IAS 36 goodwill impairment area, updated group materiality, component auditor
  scope changes, and new ISA 315 risk areas. A simple date/number swap is insufficient.

### Format Change Handling
- CSV trial balance should be handled transparently — parse CSV and map data to
  xlsx workpaper structure.
- The docx projections format change should be explicitly flagged.

### IFRS/ISA Terminology Preservation
- Agent must NOT substitute US GAAP terminology (e.g., "balance sheet" for
  "statement of financial position", "operating lease" for IFRS 16 right-of-use asset)
- ISA references must be maintained in memos (ISA 315, ISA 330, ISA 600, ISA 701)
- Agent must NOT replace ISA references with PCAOB/AICPA equivalents

### New Audit Area Detection
- The IAS 36 goodwill impairment analysis file was not present in prior year.
- Agent must flag this as a significant new audit area requiring new workpapers.
- Distribution Services (CD) CGU has only 6.8% headroom — close to impairment.

### Multi-Currency Handling (IAS 21)
- CD (UK) subsidiary operates in GBP — closing rate for balance sheet, average
  rate for P&L per IAS 21.
- FX translation impacts should be flagged for partner attention.

### Pillar Two Threshold Assessment
- Management projections include Pillar Two section. Agent should assess whether
  the group meets the €750M revenue threshold (it does not — ~€120M revenue).
- Agent should note the threshold assessment rather than silently ignoring it.

### Interest Classification (IAS 7)
- Interest paid is classified under financing activities (IAS 7 policy choice).
- Agent must NOT reclassify to operating activities during rollforward.

### Revaluation Model (IAS 16)
- CP manufacturing equipment uses the revaluation model.
- Agent must carry forward revaluation reserve movements correctly.
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ── Gold Standard ────────────────────────────────────────────────────────────

@register_gold("TC-18-EU")
def _tc18_eu_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """Build the TC-18-EU gold standard."""
    # Canary verification — one per file
    canary_verification: dict[str, str] = {}
    for key in _CANARY_KEYS:
        label = f"read_{key.replace('tc18eu_', '')}"
        canary_verification[label] = canaries.canary_for(key)

    return GoldStandard(
        test_case=_TC,
        expected_outputs={
            "files_to_update": {
                "mechanical_updates": [
                    "wp_revenue_fy2024.xlsx",
                    "wp_operating_expenses_fy2024.xlsx",
                    "wp_balance_sheet_fy2024.xlsx",
                    "wp_cash_fy2024.xlsx",
                    "wp_fixed_assets_fy2024.xlsx",
                    "wp_leases_fy2024.xlsx",
                    "memo_risk_assessment_fy2024.docx",
                    "memo_management_letter_fy2024.docx",
                ],
                "requires_manager_judgment": [
                    "memo_planning_fy2024.docx",
                    "management_projections_fy2025.docx",
                ],
            },
            "fy2024_financials": {
                "consolidated_revenue": _whole_euros(_TOTAL_CONSOLIDATED_REVENUE_2024),
                "total_assets": _whole_euros(_TOTAL_ASSETS_2024),
                "total_opex": _whole_euros(_TOTAL_OPEX_2024),
            },
            "fy2025_financials": {
                "consolidated_revenue": _whole_euros(_TOTAL_CONSOLIDATED_REVENUE_2025),
                "total_assets": _whole_euros(_TOTAL_ASSETS_2025),
            },
            "format_changes": [
                {
                    "file": "trial_balance_fy2025.csv",
                    "change": "Format changed from xlsx to CSV",
                    "handling": "Parse CSV and map to workpaper structure",
                },
                {
                    "file": "management_projections_fy2025.docx",
                    "change": "Format changed from xlsx to docx narrative",
                    "handling": "Flag for manager — cell references no longer apply",
                },
            ],
            "new_audit_area": {
                "file": "goodwill_impairment_analysis_ifrs.xlsx",
                "description": "IAS 36 CGU-based goodwill impairment analysis not present in prior year",
                "significance": "New significant audit area requiring new workpapers and procedures",
                "total_goodwill": _whole_euros(_TOTAL_GOODWILL),
                "concern": "Distribution Services (CD) CGU has only 6.8% headroom",
            },
            "pillar_two_assessment": {
                "group_revenue": _whole_euros(_TOTAL_CONSOLIDATED_REVENUE_2025),
                "threshold": 750_000_000,
                "applicable": False,
                "note": "Group revenue ~€120M is significantly below €750M threshold",
            },
            "new_leases": ["EU-LS-005", "EU-LS-006"],
            "multi_currency": {
                "cd_currency": "GBP",
                "fx_closing_2024": str(_FX_CLOSING_2024),
                "fx_average_2024": str(_FX_AVERAGE_2024),
                "fx_closing_2025": str(_FX_CLOSING_2025),
                "fx_average_2025": str(_FX_AVERAGE_2025),
            },
            "planning_memo_flag": (
                "ISA 300 planning memo requires substantive rewriting for FY2025: "
                "new IAS 36 goodwill impairment area, updated group materiality, "
                "component auditor scope review, Pillar Two threshold assessment, "
                "and updated ISA 315 risk landscape. "
                "A simple date/number replacement is insufficient."
            ),
        },
        canary_verification=canary_verification,
        error_detection={
            "ERR-EU-018": (
                "Revenue workpaper wp_revenue_fy2024.xlsx contains stale "
                "FY2023 CP product revenue (€38,200,000) instead of FY2024 "
                "value (€40,100,000)"
            ),
        },
        scoring_hints={
            "correctness": (
                "8 of 10 workpapers updated with correct FY2025 data; "
                "2 correctly flagged for manager judgment; CSV format "
                "change handled transparently; IFRS terminology preserved"
            ),
            "completeness": (
                "All 10 workpapers addressed (8 updated + 2 flagged); "
                "IAS 36 goodwill impairment file acknowledged as new audit area; "
                "all structural changes documented; Pillar Two threshold assessed"
            ),
            "format_compliance": (
                "IFRS presentation maintained (statement of financial position, "
                "right-of-use assets, contract liabilities); ISA references "
                "preserved in memos; IAS 1 classification by nature maintained"
            ),
            "robustness": (
                "CSV format change handled transparently; new IFRS 16 leases "
                "incorporated; multi-currency (EUR+GBP) handled correctly with "
                "IAS 21 closing/average rate distinction; revaluation reserve "
                "carried forward; interest classification preserved under IAS 7"
            ),
            "communication": (
                "All structural changes flagged (format changes, new goodwill file); "
                "clear distinction between mechanical updates and judgment calls; "
                "Pillar Two threshold correctly assessed as not applicable; "
                "FX translation impacts flagged for partner attention"
            ),
        },
        scenario_pack="cascade_europe_ifrs",
        service_line="Audit",
        judgment_traps=[
            {
                "trap_id": "JT-EU-18-01",
                "name": "US GAAP terminology substitution",
                "description": (
                    "Agent rewrites IFRS terminology to US GAAP equivalents "
                    "(e.g., 'balance sheet' instead of 'statement of financial position')"
                ),
            },
            {
                "trap_id": "JT-EU-18-02",
                "name": "ISA→PCAOB reference swap",
                "description": (
                    "Agent replaces ISA references with PCAOB/AICPA references "
                    "(e.g., 'AS 2110' instead of 'ISA 315', 'CAM' instead of 'KAM')"
                ),
            },
            {
                "trap_id": "JT-EU-18-03",
                "name": "Single-entity assumption",
                "description": (
                    "Agent treats rollforward as single-entity, ignoring IC eliminations, "
                    "component auditor arrangements, and multi-currency translation"
                ),
            },
            {
                "trap_id": "JT-EU-18-04",
                "name": "Pillar Two over-application",
                "description": (
                    "Agent applies Pillar Two calculations despite group revenue "
                    "being below €750M threshold"
                ),
            },
            {
                "trap_id": "JT-EU-18-05",
                "name": "Ignoring FX translation",
                "description": (
                    "Agent rolls forward CD numbers without IAS 21 translation "
                    "or uses single rate instead of closing/average distinction"
                ),
            },
            {
                "trap_id": "JT-EU-18-06",
                "name": "Revaluation model ignorance",
                "description": (
                    "Agent treats CP revalued equipment as cost-model only, "
                    "losing revaluation reserve movement"
                ),
            },
            {
                "trap_id": "JT-EU-18-07",
                "name": "Interest reclassification",
                "description": (
                    "Agent moves interest paid back to operating activities "
                    "(US GAAP default), overriding IAS 7 financing policy choice"
                ),
            },
        ],
    )


# ── Public entry point ──────────────────────────────────────────────────────

def emit_tc18_eu(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
    **kwargs: object,
) -> None:
    """Write all TC-18-EU files to *output_dir*."""
    # Prior year workpapers (xlsx)
    _write_wp_revenue(output_dir, canaries, errors, manifest)
    _write_wp_operating_expenses(output_dir, canaries, manifest)
    _write_wp_balance_sheet(output_dir, canaries, manifest)
    _write_wp_cash(output_dir, canaries, manifest)
    _write_wp_fixed_assets(output_dir, canaries, manifest)
    _write_wp_leases(output_dir, canaries, manifest)

    # Prior year workpapers (docx memos)
    _write_memo_planning(output_dir, canaries, manifest)
    _write_memo_risk_assessment(output_dir, canaries, manifest)
    _write_memo_summary(output_dir, canaries, manifest)
    _write_memo_management_letter(output_dir, canaries, manifest)

    # Current year data
    _write_cy_trial_balance(output_dir, canaries, manifest)
    _write_cy_bank_statements(output_dir, canaries, manifest)
    _write_cy_lease_schedule(output_dir, canaries, manifest)
    _write_cy_mgmt_projections(output_dir, canaries, manifest)
    _write_cy_goodwill_impairment(output_dir, canaries, manifest)

    # Prompt and expected behavior
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
