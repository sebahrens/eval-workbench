"""Formatter: TC-12 — Data Room Triage & Document Index (Advisory, Adversarial).

Emits a 32-file data room across 7 categories plus a 65-item DD checklist:

  test_cases/TC-12/input_files/
  ├── data_room/
  │   ├── 01_corporate/   (5 PDFs)
  │   ├── 02_financial/   (2 PDFs + 3 XLSX)
  │   ├── 03_legal/       (3 contract PDFs + 1 DOCX + 2 PDFs)
  │   ├── 04_hr/          (1 XLSX + 1 PDF + 3 key-employee PDFs + 1 XLSX)
  │   ├── 05_tax/         (2 PDFs + 1 XLSX + 1 PDF)
  │   ├── 06_operations/  (1 PDF + 3 XLSX)
  │   └── 07_technology/  (1 PDF + 1 XLSX + 1 DOCX)
  └── dd_checklist_standard.docx

Planted errors:
- ERR-021 (missing_data) in employee_census.xlsx — blank salary cell

Gold standard red flags:
- Pending litigation: Henderson v. Cascade PC, $2.5M exposure
- CEO golden parachute: 3× salary ($975K) upon change of control
- Acme change-of-control termination clause (~18% of revenue)
- 2 patents expiring within 18 months of FY2025 year-end
- Incomplete IP assignments from 2 founding employees
- Missing from data room: environmental assessments, regulatory permits,
  insurance claim history, real property surveys

Uses the canonical model for financials; customers.py for contracts,
key personnel, litigation, and red flags.
"""

from __future__ import annotations

import datetime
import io
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import docx
import openpyxl
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from generator.canaries import (
    CanaryRegistry,
    embed_canary_docx,
    embed_canary_xlsx,
)
from generator.errors import ErrorRegistry, PlantedError, missing_data
from generator.golds.framework import GoldStandard, register_gold
from generator.manifest import Manifest
from generator.model.ar import CUSTOMERS as AR_CUSTOMERS
from generator.model.build import CascadeModel
from generator.model.customers import (
    CONTRACTS,
    KEY_PERSONNEL,
    LITIGATION,
    contracts_with_change_of_control,
)
from generator.model.entities import ENTITIES, PARENT, SUBSIDIARIES

# ── Constants ────────────────────────────────────────────────────────────────

_TC = "TC-12"
_INPUT_DIR = f"test_cases/{_TC}/input_files"
_DR = f"{_INPUT_DIR}/data_room"

_FIXED_DATETIME = datetime.datetime(2025, 3, 15, 9, 0, 0)

# ── Patent portfolio data (not in the model — TC-12-specific) ────────────────

_PATENTS = [
    {
        "number": "US 8,234,567",
        "title": "High-Precision Bearing Assembly Method",
        "filing_date": datetime.date(2012, 3, 15),
        "issue_date": datetime.date(2014, 8, 12),
        "expiration_date": datetime.date(2032, 3, 15),
        "assignee": "Cascade Precision Components LLC",
        "inventors": "Robert J. Cascade, James M. Foster",
        "status": "Active",
    },
    {
        "number": "US 9,012,345",
        "title": "Composite Layup Process for Aerospace Components",
        "filing_date": datetime.date(2014, 7, 22),
        "issue_date": datetime.date(2017, 1, 10),
        "expiration_date": datetime.date(2034, 7, 22),
        "assignee": "Advanced Materials Corp",
        "inventors": "David R. Nakamura, Lisa K. Thornton",
        "status": "Active",
    },
    {
        "number": "US 9,456,789",
        "title": "Automated Surface Treatment System",
        "filing_date": datetime.date(2015, 11, 3),
        "issue_date": datetime.date(2018, 4, 17),
        "expiration_date": datetime.date(2035, 11, 3),
        "assignee": "Cascade Precision Components LLC",
        "inventors": "Robert J. Cascade",
        "status": "Active",
    },
    {
        # RED FLAG: Expiring within 18 months of FY2025 year-end
        "number": "US 7,890,123",
        "title": "Thermal Spray Coating for Industrial Bearings",
        "filing_date": datetime.date(2007, 6, 10),
        "issue_date": datetime.date(2010, 2, 28),
        "expiration_date": datetime.date(2027, 6, 10),
        "assignee": "Cascade Precision Components LLC",
        "inventors": "James M. Foster, Robert J. Cascade",
        "status": "Active — expiring within 18 months",
    },
    {
        # RED FLAG: Expiring within 18 months of FY2025 year-end
        "number": "US 8,012,456",
        "title": "Advanced Polymer Matrix Reinforcement Process",
        "filing_date": datetime.date(2008, 1, 20),
        "issue_date": datetime.date(2011, 9, 6),
        "expiration_date": datetime.date(2027, 1, 20),
        "assignee": "Advanced Materials Corp",
        "inventors": "David R. Nakamura",
        "status": "Active — expiring within 18 months",
    },
    {
        "number": "US 10,234,567",
        "title": "Precision Alignment Fixture for Multi-Axis Machining",
        "filing_date": datetime.date(2019, 4, 8),
        "issue_date": datetime.date(2022, 11, 15),
        "expiration_date": datetime.date(2039, 4, 8),
        "assignee": "Cascade Precision Components LLC",
        "inventors": "Thomas H. Reynolds",
        "status": "Active",
    },
]

# IP assignments — deliberately incomplete (missing 2 founding employees)
_IP_ASSIGNMENTS = [
    {
        "employee": "Thomas H. Reynolds",
        "title": "Senior Engineer",
        "date": datetime.date(2019, 4, 1),
        "scope": "All inventions related to precision machining and fixtures",
        "status": "Complete",
    },
    {
        "employee": "Lisa K. Thornton",
        "title": "Materials Scientist",
        "date": datetime.date(2014, 6, 15),
        "scope": "All inventions related to composite materials and processes",
        "status": "Complete",
    },
    {
        "employee": "James M. Foster",
        "title": "Engineering Director (Retired)",
        "date": datetime.date(2012, 2, 1),
        "scope": "All inventions related to bearing assembly and coating processes",
        "status": "Complete",
    },
]

# The 2 founding employees with MISSING IP assignments (red flag):
_MISSING_IP_ASSIGNMENTS = [
    "Robert J. Cascade (CEO / Co-Founder — named inventor on 3 patents)",
    "David R. Nakamura (CTO / Co-Founder — named inventor on 2 patents)",
]

# Globex and Initech — not in the model, TC-12 only
_GLOBEX_CONTRACT = {
    "name": "Globex Corporation",
    "type": "Customer Agreement",
    "effective": datetime.date(2023, 7, 1),
    "expiration": datetime.date(2028, 6, 30),
    "value": Decimal("8_500_000"),
    "terms": "Fixed pricing with annual escalator tied to PPI index",
    "payment": "Net 30",
}

_INITECH_CONTRACT = {
    "name": "Initech Suppliers Ltd",
    "type": "Supplier Agreement",
    "effective": datetime.date(2024, 1, 1),
    "expiration": datetime.date(2026, 12, 31),
    "value": Decimal("12_000_000"),
    "terms": "Fixed unit pricing, quarterly volume commitments",
    "payment": "Net 45",
}

# ── DD Checklist (65 items across 9 categories) ─────────────────────────────

_DD_CHECKLIST = [
    # Corporate (8 items)
    ("Corporate", "Articles of Incorporation / Certificate of Formation"),
    ("Corporate", "Bylaws or Operating Agreement"),
    ("Corporate", "Board of Directors meeting minutes (last 3 years)"),
    ("Corporate", "Shareholder meeting minutes (last 3 years)"),
    ("Corporate", "Organizational chart (legal entity structure)"),
    ("Corporate", "List of officers and directors"),
    ("Corporate", "Good standing certificates (all jurisdictions)"),
    ("Corporate", "Capitalization table / equity ownership records"),
    # Financial (9 items)
    ("Financial", "Audited financial statements (last 3 fiscal years)"),
    ("Financial", "Interim / management financial statements (current YTD)"),
    ("Financial", "Annual operating budget (current year)"),
    ("Financial", "Revenue backlog and pipeline reports"),
    ("Financial", "Accounts receivable aging schedule"),
    ("Financial", "Accounts payable aging schedule"),
    ("Financial", "Debt schedule (all outstanding obligations)"),
    ("Financial", "Capital expenditure plans and forecasts"),
    ("Financial", "Working capital analysis"),
    # Tax (7 items)
    ("Tax", "Federal income tax returns (last 3 years)"),
    ("Tax", "State and local tax returns (last 3 years)"),
    ("Tax", "Tax notices, assessments, or audit correspondence"),
    ("Tax", "Transfer pricing documentation"),
    ("Tax", "Sales and use tax compliance records"),
    ("Tax", "Property tax records and assessments"),
    ("Tax", "Tax sharing agreements (intercompany)"),
    # Legal (8 items)
    ("Legal", "Material contracts (customer agreements)"),
    ("Legal", "Material contracts (supplier agreements)"),
    ("Legal", "Pending or threatened litigation summary"),
    ("Legal", "Settlement agreements (last 5 years)"),
    ("Legal", "Intellectual property assignments and licenses"),
    ("Legal", "Insurance policies and coverage summaries"),
    ("Legal", "Consent decrees or regulatory orders"),
    ("Legal", "Insurance claim history (last 5 years)"),
    # HR (7 items)
    ("HR", "Employee census / headcount by department"),
    ("HR", "Employee benefits summary (health, retirement, equity)"),
    ("HR", "Key employee employment agreements"),
    ("HR", "Non-compete and non-solicitation agreements"),
    ("HR", "Detailed organizational chart (with reporting lines)"),
    ("HR", "Workers' compensation claims history"),
    ("HR", "OSHA citations or workplace safety reports"),
    # Operations (7 items)
    ("Operations", "Facility leases and real property agreements"),
    ("Operations", "Equipment and machinery list (fixed assets)"),
    ("Operations", "Customer list with revenue by customer"),
    ("Operations", "Vendor list (top suppliers by spend)"),
    ("Operations", "Quality certifications (ISO, AS9100, etc.)"),
    ("Operations", "Business continuity / disaster recovery plans"),
    ("Operations", "Supply chain concentration analysis"),
    # Technology (6 items)
    ("Technology", "Patent portfolio summary"),
    ("Technology", "Software license inventory"),
    ("Technology", "IT infrastructure overview and architecture"),
    ("Technology", "Cybersecurity audit or assessment reports"),
    ("Technology", "Data privacy compliance (GDPR, CCPA) documentation"),
    ("Technology", "Source code escrow agreements"),
    # Environmental (6 items) — ALL MISSING from data room
    ("Environmental", "Phase I / Phase II environmental site assessments"),
    ("Environmental", "Environmental permits and compliance records"),
    ("Environmental", "Hazardous waste disposal records"),
    ("Environmental", "Environmental remediation obligations"),
    ("Environmental", "Air and water quality permits"),
    ("Environmental", "Environmental insurance policies"),
    # Regulatory (7 items) — MOSTLY MISSING from data room
    ("Regulatory", "Government permits and licenses"),
    ("Regulatory", "Regulatory filings and compliance reports"),
    ("Regulatory", "Export control / ITAR compliance documentation"),
    ("Regulatory", "FDA / product safety certifications"),
    ("Regulatory", "Antitrust compliance documentation"),
    ("Regulatory", "Real property surveys and title reports"),
    ("Regulatory", "Zoning and land use permits"),
]

assert len(_DD_CHECKLIST) == 65, f"Expected 65 DD items, got {len(_DD_CHECKLIST)}"


# ── Deterministic save helpers ──────────────────────────────────────────────


def _save_xlsx_deterministic(wb: openpyxl.Workbook, path: str | Path) -> None:
    """Save workbook with pinned timestamps and fixed zip entry dates."""
    from openpyxl.writer.excel import ExcelWriter

    path = Path(path)

    wb.properties.created = _FIXED_DATETIME
    wb.properties.modified = _FIXED_DATETIME

    buf = io.BytesIO()
    archive = ZipFile(buf, "w", ZIP_DEFLATED, allowZip64=True)
    writer = ExcelWriter(wb, archive)
    writer.save()

    fixed_date_time = (2025, 3, 15, 9, 0, 0)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=fixed_date_time)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _save_docx_deterministic(doc: Any, path: str | Path) -> None:
    """Save a python-docx Document with fixed zip entry timestamps."""
    path = Path(path)

    buf = io.BytesIO()
    doc.save(buf)

    fixed_date_time = (2025, 3, 15, 9, 0, 0)
    buf.seek(0)
    with ZipFile(buf, "r") as src, ZipFile(str(path), "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            info = ZipInfo(item.filename, date_time=fixed_date_time)
            info.compress_type = item.compress_type
            dst.writestr(info, src.read(item.filename))


def _whole_dollars(d: Decimal) -> int:
    """Round a Decimal to whole dollars."""
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


# ── Shared reportlab styles ────────────────────────────────────────────────

def _pdf_styles() -> tuple[
    ParagraphStyle, ParagraphStyle, ParagraphStyle, ParagraphStyle,
]:
    """Return (title, heading, subheading, body) paragraph styles."""
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle", parent=styles["Title"],
        fontSize=16, spaceAfter=20, alignment=1,
    )
    heading_style = ParagraphStyle(
        "DocHeading", parent=styles["Heading2"],
        fontSize=12, spaceBefore=12, spaceAfter=6,
    )
    subheading_style = ParagraphStyle(
        "DocSubheading", parent=styles["Heading3"],
        fontSize=10, spaceBefore=8, spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "DocBody", parent=styles["Normal"],
        fontSize=10, leading=14, spaceAfter=6,
    )
    return title_style, heading_style, subheading_style, body_style


def _build_simple_pdf(
    full_path: Path,
    canary_code: str,
    title_text: str,
    elements: list,
) -> str:
    """Build a reportlab PDF with canary in author metadata. Returns canary location."""
    doc = SimpleDocTemplate(
        str(full_path),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=title_text,
        author=f"CANARY: {canary_code}",
        creator="Cascade Industries Document System",
        invariant=True,
    )
    doc.build(elements)
    return "PDF metadata → Author"


# ── XLSX style constants ────────────────────────────────────────────────────

_HEADER_FONT = Font(bold=True, size=11)
_HEADER_FILL = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
_HEADER_FONT_WHITE = Font(bold=True, size=11, color="FFFFFF")
_THIN_BORDER = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin"),
)
_MONEY_FMT = '#,##0'
_MONEY_FMT_2 = '#,##0.00'
_DATE_FMT = 'YYYY-MM-DD'


def _style_header_row(ws: Any, row: int, col_count: int) -> None:
    """Apply header styling to a row."""
    for col in range(1, col_count + 1):
        cell = ws.cell(row=row, column=col)
        cell.font = _HEADER_FONT_WHITE
        cell.fill = _HEADER_FILL
        cell.border = _THIN_BORDER
        cell.alignment = Alignment(horizontal="center")


# ═══════════════════════════════════════════════════════════════════════════
# Category 01: Corporate (5 PDFs)
# ═══════════════════════════════════════════════════════════════════════════


def _write_articles_of_incorporation(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_articles_of_incorporation"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/01_corporate/articles_of_incorporation.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()
    elements = [
        Paragraph("Articles of Incorporation", ts),
        Paragraph("State of Oregon — Secretary of State", hs),
        Spacer(1, 12),
        Paragraph(
            f"<b>Entity Name:</b> {PARENT.name}", bs,
        ),
        Paragraph("<b>Date of Incorporation:</b> September 12, 1987", bs),
        Paragraph("<b>State:</b> Oregon", bs),
        Paragraph("<b>Registered Agent:</b> Pacific Northwest Corporate Services, Inc.", bs),
        Paragraph(
            "<b>Principal Office:</b> 4200 NW Industrial Parkway, Portland, OR 97210", bs,
        ),
        Spacer(1, 12),
        Paragraph("Article I — Name", hs),
        Paragraph(
            f"The name of the corporation shall be {PARENT.name}.", bs,
        ),
        Paragraph("Article II — Purpose", hs),
        Paragraph(
            "The corporation is organized for the purpose of engaging in any lawful "
            "business activity, including the manufacture, sale, and distribution of "
            "precision industrial components, advanced materials, and related services.",
            bs,
        ),
        Paragraph("Article III — Authorized Shares", hs),
        Paragraph(
            "The total number of shares authorized is 10,000,000 shares of common stock, "
            "par value $0.01 per share, and 1,000,000 shares of preferred stock, par "
            "value $0.01 per share.",
            bs,
        ),
        Paragraph("Article IV — Duration", hs),
        Paragraph("The corporation shall have perpetual existence.", bs),
        Paragraph("Article V — Registered Agent", hs),
        Paragraph(
            "The registered agent is Pacific Northwest Corporate Services, Inc., "
            "900 SW Fifth Avenue, Suite 2000, Portland, Oregon 97204.",
            bs,
        ),
    ]

    canary_loc = _build_simple_pdf(full_path, canary_code, "Articles of Incorporation", elements)
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_bylaws(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_bylaws"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/01_corporate/bylaws.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()
    elements = [
        Paragraph(f"Bylaws of {PARENT.name}", ts),
        Paragraph("As Amended and Restated — January 15, 2023", hs),
        Spacer(1, 12),
        Paragraph("Article I — Offices", hs),
        Paragraph(
            "Section 1.1. The principal office of the Corporation shall be located "
            "at 4200 NW Industrial Parkway, Portland, Oregon 97210.", bs,
        ),
        Paragraph("Article II — Meetings of Shareholders", hs),
        Paragraph(
            "Section 2.1. The annual meeting of shareholders shall be held on the "
            "third Tuesday of April of each year.", bs,
        ),
        Paragraph("Article III — Board of Directors", hs),
        Paragraph(
            "Section 3.1. The Board shall consist of not fewer than five (5) and "
            "not more than nine (9) directors. As of the date of this amendment, "
            "the Board consists of seven (7) members.", bs,
        ),
        Paragraph("Article IV — Officers", hs),
        Paragraph(
            "Section 4.1. Officers shall include a Chief Executive Officer, Chief "
            "Financial Officer, Chief Technology Officer, and Secretary. Officers "
            "are appointed by the Board and serve at the Board's pleasure.", bs,
        ),
        Paragraph("Article V — Indemnification", hs),
        Paragraph(
            "Section 5.1. The Corporation shall indemnify its directors and officers "
            "to the fullest extent permitted by Oregon law (ORS 60.387 et seq.).", bs,
        ),
    ]

    canary_loc = _build_simple_pdf(full_path, canary_code, "Bylaws", elements)
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_board_minutes(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
    year: int, file_key: str,
) -> None:
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/01_corporate/board_minutes_{year}.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()

    meeting_date = f"March 18, {year}"
    elements = [
        Paragraph("Board of Directors Meeting Minutes", ts),
        Paragraph(f"{PARENT.name}", hs),
        Paragraph(f"<b>Date:</b> {meeting_date}", bs),
        Paragraph("<b>Location:</b> Corporate Headquarters, Portland, OR", bs),
        Paragraph(
            "<b>Attendees:</b> R. Cascade (Chair), M. Chen, D. Nakamura, "
            "J. Wilson, S. Park, A. Morales, K. Fitzgerald",
            bs,
        ),
        Spacer(1, 12),
        Paragraph("1. Call to Order", hs),
        Paragraph(
            "The meeting was called to order at 9:00 AM by Chairman R. Cascade.", bs,
        ),
        Paragraph("2. Financial Review", hs),
        Paragraph(
            f"CFO M. Chen presented the financial results for FY{year - 1} and "
            f"the FY{year} budget. The Board approved the annual operating budget.",
            bs,
        ),
    ]

    if year == 2024:
        # Board minutes 2024 reference the Henderson litigation
        elements.extend([
            Paragraph("3. Legal Matters", hs),
            Paragraph(
                "General Counsel reported on pending litigation matters. The Board "
                "was briefed on Henderson v. Cascade Precision Components LLC "
                "(product liability — filed June 2024). Outside counsel "
                "(Mitchell, Hartwell &amp; Associates LLP) estimates potential "
                "exposure of up to $2,500,000. A reserve of $750,000 has been "
                "accrued. Discovery is ongoing.",
                bs,
            ),
            Paragraph("4. Strategic Initiatives", hs),
            Paragraph(
                "CEO R. Cascade reviewed the strategic plan, including continued "
                "investment in the Advanced Materials division and expansion of "
                "distribution capabilities.",
                bs,
            ),
        ])
    else:  # 2025
        # Board minutes 2025 reference patent portfolio review and CEO golden parachute
        elements.extend([
            Paragraph("3. Executive Compensation", hs),
            Paragraph(
                "The Compensation Committee presented its annual review. The Board "
                "ratified the CEO employment agreement, including the change-of-control "
                "provision (3&times; base salary golden parachute). The Committee "
                "confirmed this provision remains appropriate given the company's "
                "strategic positioning and market practice for mid-market companies.",
                bs,
            ),
            Paragraph("4. Technology &amp; IP Review", hs),
            Paragraph(
                "CTO D. Nakamura presented the patent portfolio status. The Board "
                "noted that two patents (US 7,890,123 and US 8,012,456) are "
                "approaching expiration within the next 18 months. Management was "
                "directed to evaluate whether continuation or new filings are "
                "warranted to protect core manufacturing processes.",
                bs,
            ),
            Paragraph("5. Litigation Update", hs),
            Paragraph(
                "The Henderson v. Cascade Precision Components matter remains in "
                "discovery. Depositions are scheduled for Q1 2026. Exposure "
                "estimate unchanged at $2,500,000.",
                bs,
            ),
        ])

    elements.extend([
        Spacer(1, 12),
        Paragraph("Adjournment", hs),
        Paragraph(
            "There being no further business, the meeting was adjourned at 11:30 AM.",
            bs,
        ),
        Spacer(1, 24),
        Paragraph("______________________________", bs),
        Paragraph("K. Fitzgerald, Secretary", bs),
    ])

    canary_loc = _build_simple_pdf(full_path, canary_code, f"Board Minutes {year}", elements)
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_org_chart_pdf(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_org_chart"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/01_corporate/org_chart.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()
    elements = [
        Paragraph("Corporate Organizational Chart", ts),
        Paragraph(PARENT.name, hs),
        Spacer(1, 12),
    ]

    # Simple table representation of org structure
    data = [
        ["Entity", "Location", "Business Line", "Key Executive"],
        [PARENT.name, PARENT.location, "Holding Company (Parent)", "Robert J. Cascade, CEO"],
    ]
    for code, entity in sorted(SUBSIDIARIES.items()):
        exec_name = "Operations Lead"
        if code == "PC":
            exec_name = "VP Manufacturing"
        elif code == "AM":
            exec_name = "David R. Nakamura, CTO"
        elif code == "DS":
            exec_name = "VP Distribution"
        data.append([entity.name, entity.location, entity.revenue_mix, exec_name])

    tbl = Table(data, colWidths=[2.2 * inch, 1.3 * inch, 2.0 * inch, 1.5 * inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.append(tbl)

    canary_loc = _build_simple_pdf(full_path, canary_code, "Org Chart", elements)
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# Category 02: Financial (2 PDFs + 3 XLSX)
# ═══════════════════════════════════════════════════════════════════════════


def _write_audited_financials_pdf(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
    year: int,
    file_key: str,
) -> None:
    from generator.model.views import build_balance_sheet, build_income_statement

    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/02_financial/audited_financials_fy{year}.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()

    inc = build_income_statement(model.ledger, year)
    bal = build_balance_sheet(model.ledger, datetime.date(year, 12, 31))

    elements = [
        Paragraph(f"Audited Financial Statements — FY{year}", ts),
        Paragraph(PARENT.name, hs),
        Paragraph(
            "Prepared by: Whitfield &amp; Associates, CPAs (Independent Auditors)", bs,
        ),
        Spacer(1, 12),
        Paragraph("Consolidated Income Statement", hs),
    ]

    # Income statement table
    inc_data = [
        ["", f"FY{year}"],
        ["Revenue", f"${_whole_dollars(inc.total_revenue):,}"],
        ["Cost of Goods Sold", f"${_whole_dollars(inc.total_cogs):,}"],
        ["Gross Profit", f"${_whole_dollars(inc.gross_profit):,}"],
        ["Operating Expenses", f"${_whole_dollars(inc.total_opex):,}"],
        ["Operating Income", f"${_whole_dollars(inc.operating_income):,}"],
        ["Other Expense (Net)", f"${_whole_dollars(inc.total_other):,}"],
        ["Pre-Tax Income", f"${_whole_dollars(inc.pre_tax_income):,}"],
        ["Tax Provision", f"${_whole_dollars(inc.total_tax):,}"],
        ["Net Income", f"${_whole_dollars(inc.net_income):,}"],
    ]
    tbl = Table(inc_data, colWidths=[3.5 * inch, 2.0 * inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
        ("LINEABOVE", (0, -1), (-1, -1), 1.5, colors.black),
    ]))
    elements.extend([tbl, Spacer(1, 24)])

    # Balance sheet (summarized)
    elements.append(Paragraph("Consolidated Balance Sheet", hs))
    bal_data = [
        ["", f"FY{year}"],
        ["Total Assets", f"${_whole_dollars(bal.total_assets):,}"],
        ["Total Liabilities", f"${_whole_dollars(bal.total_liabilities):,}"],
        ["Total Equity", f"${_whole_dollars(bal.total_equity):,}"],
    ]
    tbl2 = Table(bal_data, colWidths=[3.5 * inch, 2.0 * inch])
    tbl2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
    ]))
    elements.extend([tbl2, Spacer(1, 12)])

    elements.append(Paragraph(
        "<i>Note: Complete financial statements with notes are available upon request.</i>",
        bs,
    ))

    canary_loc = _build_simple_pdf(
        full_path, canary_code, f"Audited Financials FY{year}", elements,
    )
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_management_financials_xlsx(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    from generator.model.views import build_income_statement

    file_key = "tc12_management_financials_fy2025"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/02_financial/management_financials_fy2025_ytd.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    embed_canary_xlsx(wb, canary_code)
    ws = wb.active
    ws.title = "Management Financials"

    inc = build_income_statement(model.ledger, 2025)

    headers = ["Line Item", "FY2025 YTD", "FY2024 Actual", "Budget FY2025", "Variance"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    inc_2024 = build_income_statement(model.ledger, 2024)

    rows = [
        ("Revenue", inc.total_revenue, inc_2024.total_revenue, inc.total_revenue * Decimal("0.95")),
        ("COGS", inc.total_cogs, inc_2024.total_cogs, inc.total_cogs * Decimal("0.96")),
        ("Gross Profit", inc.gross_profit, inc_2024.gross_profit,
         inc.total_revenue * Decimal("0.95") - inc.total_cogs * Decimal("0.96")),
        ("Operating Expenses", inc.total_opex, inc_2024.total_opex,
         inc.total_opex * Decimal("1.02")),
        ("Operating Income", inc.operating_income, inc_2024.operating_income, None),
        ("Net Income", inc.net_income, inc_2024.net_income, None),
    ]

    for r, (label, ytd, prior, budget) in enumerate(rows, 2):
        ws.cell(row=r, column=1, value=label)
        ws.cell(row=r, column=2, value=_whole_dollars(ytd)).number_format = _MONEY_FMT
        ws.cell(row=r, column=3, value=_whole_dollars(prior)).number_format = _MONEY_FMT
        if budget is not None:
            ws.cell(row=r, column=4, value=_whole_dollars(budget)).number_format = _MONEY_FMT
            variance = _whole_dollars(ytd - budget)
            ws.cell(row=r, column=5, value=variance).number_format = _MONEY_FMT

    for col in range(1, 6):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 22

    _save_xlsx_deterministic(wb, full_path)
    canaries.set_location(file_key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary_code, test_cases=[_TC])


def _write_budget_xlsx(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    from generator.model.views import build_income_statement

    file_key = "tc12_budget_fy2025"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/02_financial/budget_fy2025.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    embed_canary_xlsx(wb, canary_code)
    ws = wb.active
    ws.title = "FY2025 Budget"

    inc_2024 = build_income_statement(model.ledger, 2024)

    headers = ["Line Item", "FY2024 Actual", "FY2025 Budget", "Growth %"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    # Budget assumes 9% growth on revenue (matching config growth rate)
    growth = Decimal("1.09")
    budget_rows = [
        ("Revenue", inc_2024.total_revenue, inc_2024.total_revenue * growth),
        ("COGS", inc_2024.total_cogs, inc_2024.total_cogs * growth * Decimal("0.98")),
        ("Operating Expenses", inc_2024.total_opex, inc_2024.total_opex * Decimal("1.05")),
    ]

    for r, (label, actual, budget) in enumerate(budget_rows, 2):
        ws.cell(row=r, column=1, value=label)
        ws.cell(row=r, column=2, value=_whole_dollars(actual)).number_format = _MONEY_FMT
        ws.cell(row=r, column=3, value=_whole_dollars(budget)).number_format = _MONEY_FMT
        pct = float((budget / actual - 1) * 100) if actual else 0
        ws.cell(row=r, column=4, value=round(pct, 1))

    for col in range(1, 5):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 20

    _save_xlsx_deterministic(wb, full_path)
    canaries.set_location(file_key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary_code, test_cases=[_TC])


def _write_debt_schedule_xlsx(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_debt_schedule"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/02_financial/debt_schedule.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    embed_canary_xlsx(wb, canary_code)
    ws = wb.active
    ws.title = "Debt Schedule"

    headers = ["Facility", "Lender", "Type", "Original Amount",
               "Outstanding Balance", "Rate", "Maturity"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    debt_items = [
        ("Revolving Credit Facility", "Pacific Northwest Bank", "Revolver",
         25_000_000, 8_500_000, "SOFR + 2.25%", "2028-06-30"),
        ("Term Loan A", "Pacific Northwest Bank", "Term Loan",
         30_000_000, 22_500_000, "SOFR + 2.75%", "2029-12-31"),
        ("Equipment Financing", "Cascade Equipment Finance", "Asset-Based",
         10_000_000, 6_200_000, "5.50% Fixed", "2027-03-31"),
        ("Intercompany Note (CI → AM)", "Internal", "IC Note",
         5_000_000, 5_000_000, "5.00% Fixed", "2027-12-31"),
    ]

    for r, (facility, lender, dtype, orig, bal, rate, mat) in enumerate(debt_items, 2):
        ws.cell(row=r, column=1, value=facility)
        ws.cell(row=r, column=2, value=lender)
        ws.cell(row=r, column=3, value=dtype)
        ws.cell(row=r, column=4, value=orig).number_format = _MONEY_FMT
        ws.cell(row=r, column=5, value=bal).number_format = _MONEY_FMT
        ws.cell(row=r, column=6, value=rate)
        ws.cell(row=r, column=7, value=mat)

    # Total row
    total_row = len(debt_items) + 2
    ws.cell(row=total_row, column=1, value="Total").font = Font(bold=True)
    ws.cell(
        row=total_row, column=5,
        value=sum(d[4] for d in debt_items),
    ).number_format = _MONEY_FMT
    ws.cell(row=total_row, column=5).font = Font(bold=True)

    for col in range(1, 8):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 22

    _save_xlsx_deterministic(wb, full_path)
    canaries.set_location(file_key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# Category 03: Legal (3 contract PDFs + 1 DOCX + 2 PDFs)
# ═══════════════════════════════════════════════════════════════════════════


def _write_customer_agreement_acme(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """RED FLAG: Acme contract with change-of-control termination clause."""
    file_key = "tc12_customer_agreement_acme"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/03_legal/material_contracts/customer_agreement_acme.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    acme = CONTRACTS[0]  # CTR-001 — Acme Manufacturing
    ts, hs, shs, bs = _pdf_styles()

    elements = [
        Paragraph("Master Supply Agreement", ts),
        Paragraph(
            f"Between {PARENT.name} (through Cascade Precision Components LLC) "
            f"and {acme.customer_name}",
            hs,
        ),
        Spacer(1, 12),
        Paragraph(f"<b>Contract ID:</b> {acme.contract_id}", bs),
        Paragraph(f"<b>Effective Date:</b> {acme.effective_date.isoformat()}", bs),
        Paragraph(f"<b>Expiration Date:</b> {acme.expiration_date.isoformat()}", bs),
        Paragraph(f"<b>Annual Volume:</b> ~${acme.annual_volume:,.0f}", bs),
        Paragraph(f"<b>Pricing:</b> {acme.pricing_terms}", bs),
        Paragraph(f"<b>Payment Terms:</b> {acme.payment_terms}", bs),
        Spacer(1, 12),
        Paragraph("Section 4 — Term and Renewal", hs),
        Paragraph(
            "This Agreement shall remain in effect from the Effective Date through "
            f"the Expiration Date ({acme.expiration_date.isoformat()}). This Agreement "
            "does NOT automatically renew. Either party must provide written notice "
            "of intent to renew at least 90 days prior to expiration.",
            bs,
        ),
        Spacer(1, 12),
        Paragraph("Section 8 — Change of Control", hs),
        Paragraph(
            "<b>8.1 Termination Right.</b> In the event of a Change of Control of "
            "Supplier (defined as any transaction or series of transactions resulting "
            "in a change in more than 50% of the voting securities or substantially "
            "all assets of Supplier), Customer shall have the right to terminate this "
            "Agreement upon ninety (90) days' written notice, without penalty or "
            "further obligation, delivered within one hundred twenty (120) days "
            "following the closing of such transaction.",
            bs,
        ),
        Paragraph(
            "<b>8.2 Notification.</b> Supplier shall provide Customer with written "
            "notice of any pending Change of Control transaction no later than "
            "thirty (30) days prior to the anticipated closing date.",
            bs,
        ),
        Spacer(1, 12),
        Paragraph("Section 12 — Governing Law", hs),
        Paragraph(
            "This Agreement shall be governed by the laws of the State of Oregon.",
            bs,
        ),
    ]

    canary_loc = _build_simple_pdf(
        full_path, canary_code, "Customer Agreement — Acme", elements,
    )
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_customer_agreement_globex(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_customer_agreement_globex"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/03_legal/material_contracts/customer_agreement_globex.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    g = _GLOBEX_CONTRACT
    ts, hs, _, bs = _pdf_styles()
    elements = [
        Paragraph("Customer Agreement", ts),
        Paragraph(f"Between {PARENT.name} and {g['name']}", hs),
        Spacer(1, 12),
        Paragraph(f"<b>Effective Date:</b> {g['effective'].isoformat()}", bs),
        Paragraph(f"<b>Expiration Date:</b> {g['expiration'].isoformat()}", bs),
        Paragraph(f"<b>Annual Value:</b> ~${g['value']:,.0f}", bs),
        Paragraph(f"<b>Pricing:</b> {g['terms']}", bs),
        Paragraph(f"<b>Payment Terms:</b> {g['payment']}", bs),
        Spacer(1, 12),
        Paragraph("Section 4 — Term", hs),
        Paragraph(
            "This Agreement shall automatically renew for successive one-year "
            "periods unless either party provides 60 days' written notice.",
            bs,
        ),
        Paragraph("Section 8 — Change of Control", hs),
        Paragraph(
            "No change of control provision. Standard assignment clause applies.",
            bs,
        ),
    ]

    canary_loc = _build_simple_pdf(
        full_path, canary_code, "Customer Agreement — Globex", elements,
    )
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_supplier_agreement_initech(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_supplier_agreement_initech"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/03_legal/material_contracts/supplier_agreement_initech.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    s = _INITECH_CONTRACT
    ts, hs, _, bs = _pdf_styles()
    elements = [
        Paragraph("Supplier Agreement", ts),
        Paragraph(f"Between {PARENT.name} and {s['name']}", hs),
        Spacer(1, 12),
        Paragraph(f"<b>Effective Date:</b> {s['effective'].isoformat()}", bs),
        Paragraph(f"<b>Expiration Date:</b> {s['expiration'].isoformat()}", bs),
        Paragraph(f"<b>Annual Value:</b> ~${s['value']:,.0f}", bs),
        Paragraph(f"<b>Pricing:</b> {s['terms']}", bs),
        Paragraph(f"<b>Payment Terms:</b> {s['payment']}", bs),
        Spacer(1, 12),
        Paragraph("Section 3 — Scope", hs),
        Paragraph(
            "Supplier shall provide raw materials (specialty alloys and polymers) "
            "for Cascade's Precision Components and Advanced Materials divisions.",
            bs,
        ),
    ]

    canary_loc = _build_simple_pdf(
        full_path, canary_code, "Supplier Agreement — Initech", elements,
    )
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_pending_litigation_docx(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """RED FLAG: Henderson v. Cascade PC — $2.5M product liability."""
    file_key = "tc12_pending_litigation_summary"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/03_legal/pending_litigation_summary.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()
    embed_canary_docx(doc, canary_code)

    title = doc.add_heading("Pending Litigation Summary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"{PARENT.name} — Confidential", style="Subtitle",
    )
    doc.add_paragraph("Prepared: March 15, 2025")
    doc.add_paragraph("")

    matter = LITIGATION[0]
    doc.add_heading(f"1. {matter.title}", level=1)

    tbl = doc.add_table(rows=9, cols=2)
    tbl.style = "Table Grid"
    fields = [
        ("Matter ID", matter.matter_id),
        ("Case Type", matter.case_type),
        ("Filing Date", matter.filing_date.isoformat()),
        ("Court", matter.court),
        ("Plaintiff", matter.plaintiff),
        ("Defendant", matter.defendant),
        ("Potential Exposure", f"${matter.potential_exposure:,.0f}"),
        ("Accrued Liability", f"${matter.accrued_liability:,.0f}"),
        ("Status", matter.status),
    ]
    for i, (label, value) in enumerate(fields):
        tbl.cell(i, 0).text = label
        tbl.cell(i, 1).text = value
        tbl.cell(i, 0).paragraphs[0].runs[0].bold = True if tbl.cell(i, 0).paragraphs[0].runs else None

    doc.add_paragraph("")
    doc.add_heading("Description", level=2)
    doc.add_paragraph(matter.description)

    doc.add_heading("Outside Counsel Assessment", level=2)
    doc.add_paragraph(
        f"Outside counsel ({matter.outside_counsel}) has assessed the likelihood "
        f"of an adverse outcome as 'reasonably possible.' The company has accrued "
        f"${matter.accrued_liability:,.0f} based on the most likely outcome estimate, "
        f"with a range of potential exposure from ${matter.accrued_liability:,.0f} "
        f"to ${matter.potential_exposure:,.0f}."
    )

    doc.add_paragraph("")
    doc.add_heading("2. Other Matters", level=1)
    doc.add_paragraph(
        "No other material pending or threatened litigation matters as of the "
        "date of this summary."
    )

    _save_docx_deterministic(doc, full_path)
    canaries.set_location(file_key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary_code, test_cases=[_TC])


def _write_ip_assignments_pdf(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """RED FLAG: Missing IP assignments for 2 founding employees."""
    file_key = "tc12_ip_assignment_agreements"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/03_legal/ip_assignment_agreements.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()
    elements = [
        Paragraph("Intellectual Property Assignment Agreements", ts),
        Paragraph(PARENT.name, hs),
        Spacer(1, 12),
        Paragraph(
            "The following employees have executed IP assignment agreements "
            "assigning all work-product intellectual property to the Company:",
            bs,
        ),
        Spacer(1, 8),
    ]

    # Assignments table
    data = [["Employee", "Title", "Date Executed", "Scope", "Status"]]
    for a in _IP_ASSIGNMENTS:
        data.append([
            a["employee"], a["title"], a["date"].isoformat(),
            a["scope"], a["status"],
        ])

    tbl = Table(data, colWidths=[1.4 * inch, 1.2 * inch, 0.9 * inch, 1.8 * inch, 0.7 * inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.extend([tbl, Spacer(1, 18)])

    # Note about missing assignments — this is the red flag
    elements.extend([
        Paragraph("Outstanding IP Assignments", hs),
        Paragraph(
            "<b>NOTE:</b> The following individuals are named inventors on company "
            "patents but do <b>not</b> have IP assignment agreements on file:",
            bs,
        ),
    ])
    for name in _MISSING_IP_ASSIGNMENTS:
        elements.append(Paragraph(f"&bull; {name}", bs))

    elements.append(Spacer(1, 12))
    elements.append(Paragraph(
        "Legal counsel has recommended that the Company obtain executed IP "
        "assignment agreements from the above individuals as soon as practicable, "
        "particularly prior to any change-of-control transaction.",
        bs,
    ))

    canary_loc = _build_simple_pdf(
        full_path, canary_code, "IP Assignment Agreements", elements,
    )
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_insurance_policies_pdf(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_insurance_policies_summary"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/03_legal/insurance_policies_summary.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()
    elements = [
        Paragraph("Insurance Policies Summary", ts),
        Paragraph(PARENT.name, hs),
        Spacer(1, 12),
    ]

    policies = [
        ["Type", "Carrier", "Policy Number", "Coverage Limit", "Deductible", "Expiration"],
        ["General Liability", "Pacific Mutual", "GL-2025-4872", "$5,000,000", "$50,000", "2026-01-31"],
        ["Product Liability", "Pacific Mutual", "PL-2025-4873", "$10,000,000", "$100,000", "2026-01-31"],
        ["D&O Liability", "Westcoast Underwriters", "DO-2025-1234", "$5,000,000", "$75,000", "2025-12-31"],
        ["Property", "National Insurance Co", "PR-2025-5678", "$25,000,000", "$250,000", "2026-03-31"],
        ["Workers' Compensation", "State Fund", "WC-2025-9012", "Statutory", "$0", "2026-01-31"],
        ["Umbrella / Excess", "Pacific Mutual", "UE-2025-3456", "$15,000,000", "$0", "2026-01-31"],
    ]

    tbl = Table(policies, colWidths=[1.1 * inch, 1.1 * inch, 1.0 * inch, 1.0 * inch, 0.8 * inch, 0.9 * inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.extend([tbl, Spacer(1, 12)])

    elements.append(Paragraph(
        "<i>Note: Insurance claim history is maintained by the broker and has "
        "not been included in this data room. Please request separately.</i>",
        bs,
    ))

    canary_loc = _build_simple_pdf(
        full_path, canary_code, "Insurance Policies Summary", elements,
    )
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# Category 04: HR (1 XLSX + 1 PDF + 3 key-employee PDFs + 1 XLSX)
# ═══════════════════════════════════════════════════════════════════════════


def _write_employee_census_xlsx(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    file_key = "tc12_employee_census"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/04_hr/employee_census.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    embed_canary_xlsx(wb, canary_code)
    ws = wb.active
    ws.title = "Employee Census"

    headers = ["Employee ID", "Name", "Entity", "Department", "Title",
               "Hire Date", "Annual Salary", "State", "Status"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    # Sort employees deterministically
    employees = sorted(model.employees, key=lambda e: e.employee_id)
    err021_target_idx = 4  # 5th employee in sorted order
    for r, emp in enumerate(employees, 2):
        ws.cell(row=r, column=1, value=emp.employee_id)
        ws.cell(row=r, column=2, value=emp.name)
        ws.cell(row=r, column=3, value=emp.entity_code)
        ws.cell(row=r, column=4, value=emp.department)
        ws.cell(row=r, column=5, value=emp.title)
        ws.cell(row=r, column=6, value=emp.hire_date.isoformat())
        if (r - 2) == err021_target_idx:
            correct_salary = emp.annual_salary
            ws.cell(row=r, column=7, value=missing_data())  # blank cell
            errors.add(PlantedError(
                error_id="ERR-021",
                file=f"{_DR}/04_hr/employee_census.xlsx",
                location=(
                    f"Sheet 'Employee Census', Row {r}, Column G (Annual Salary) "
                    f"for employee {emp.name} ({emp.employee_id})"
                ),
                type="missing_data",
                description=(
                    f"Employee salary is blank for {emp.name} ({emp.employee_id}) "
                    f"instead of ${correct_salary:,}"
                ),
                severity="immaterial",
                which_test_cases_should_catch=["TC-12"],
            ))
        else:
            ws.cell(row=r, column=7, value=emp.annual_salary).number_format = _MONEY_FMT
        ws.cell(row=r, column=8, value=emp.state)
        status = "Terminated" if emp.termination_date else "Active"
        ws.cell(row=r, column=9, value=status)

    for col in range(1, 10):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 18

    _save_xlsx_deterministic(wb, full_path)
    canaries.set_location(file_key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary_code, test_cases=[_TC])


def _write_benefits_summary_pdf(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_benefits_summary"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/04_hr/benefits_summary.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()
    elements = [
        Paragraph("Employee Benefits Summary", ts),
        Paragraph(f"{PARENT.name} — FY2025", hs),
        Spacer(1, 12),
        Paragraph("Health Insurance", hs),
        Paragraph(
            "All full-time employees are eligible for health, dental, and vision "
            "coverage through Pacific Health Group. The company pays 80% of premiums "
            "for employee-only coverage and 65% for family coverage.", bs,
        ),
        Paragraph("Retirement Plans", hs),
        Paragraph(
            "401(k) plan with company match of 4% of base salary (100% match on "
            "first 3%, 50% match on next 2%). Immediate vesting on employee "
            "contributions; 3-year cliff vesting on employer match.", bs,
        ),
        Paragraph("Paid Time Off", hs),
        Paragraph(
            "Standard PTO: 15 days/year for 0-5 years tenure, 20 days/year for "
            "5-10 years, 25 days/year for 10+ years. Separate sick leave: 10 days/year.",
            bs,
        ),
        Paragraph("Life Insurance", hs),
        Paragraph("Company-paid life insurance of 1x annual salary, up to $500,000.", bs),
    ]

    canary_loc = _build_simple_pdf(full_path, canary_code, "Benefits Summary", elements)
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_key_employee_agreement_pdf(
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
    kp_index: int,
    file_key: str,
    filename: str,
) -> None:
    """Write a key employee employment agreement PDF.

    RED FLAG for CEO: golden parachute 3x salary.
    """
    kp = KEY_PERSONNEL[kp_index]
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/04_hr/key_employee_agreements/{filename}"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()
    payout = Decimal(kp.base_salary) * kp.change_of_control_multiplier

    elements = [
        Paragraph("Employment Agreement", ts),
        Paragraph(f"{kp.person_name} — {kp.title}", hs),
        Spacer(1, 12),
        Paragraph(f"<b>Entity:</b> {ENTITIES[kp.entity_code].name}", bs),
        Paragraph(f"<b>Effective Date:</b> {kp.effective_date.isoformat()}", bs),
        Paragraph(f"<b>Base Salary:</b> ${kp.base_salary:,}", bs),
        Paragraph(f"<b>Bonus Target:</b> {float(kp.bonus_target_pct) * 100:.0f}% of base salary", bs),
        Spacer(1, 12),
        Paragraph("Section 3 — Severance", hs),
        Paragraph(
            f"Upon termination without cause, Employee shall receive {kp.severance_months} "
            f"months of base salary continuation plus pro-rated bonus.",
            bs,
        ),
        Paragraph("Section 4 — Change of Control Provision", hs),
        Paragraph(
            f"In the event of a Change of Control (as defined in Section 4.1), "
            f"Employee shall be entitled to a lump-sum payment equal to "
            f"<b>{kp.change_of_control_multiplier}× base salary "
            f"(${payout:,.0f})</b>, payable within 30 days of the closing of "
            f"the Change of Control transaction.",
            bs,
        ),
        Paragraph("Section 5 — Non-Compete", hs),
        Paragraph(
            f"Employee agrees to a non-compete period of {kp.non_compete_months} months "
            f"following termination, within a 100-mile radius of any Company facility.",
            bs,
        ),
    ]

    if kp.notes:
        elements.extend([
            Spacer(1, 12),
            Paragraph("Additional Notes", hs),
            Paragraph(kp.notes, bs),
        ])

    canary_loc = _build_simple_pdf(
        full_path, canary_code, f"Employment Agreement — {kp.title}", elements,
    )
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_org_chart_detailed_xlsx(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    file_key = "tc12_org_chart_detailed"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/04_hr/org_chart_detailed.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    embed_canary_xlsx(wb, canary_code)
    ws = wb.active
    ws.title = "Headcount by Entity & Dept"

    headers = ["Entity", "Department", "Headcount", "Avg Salary"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    # Aggregate from model.employees
    active = [e for e in model.employees if not e.termination_date]
    dept_counts: dict[tuple[str, str], list[int]] = {}
    for emp in active:
        key = (emp.entity_code, emp.department)
        dept_counts.setdefault(key, []).append(emp.annual_salary)

    row = 2
    for (entity, dept) in sorted(dept_counts.keys()):
        salaries = dept_counts[(entity, dept)]
        ws.cell(row=row, column=1, value=entity)
        ws.cell(row=row, column=2, value=dept)
        ws.cell(row=row, column=3, value=len(salaries))
        ws.cell(row=row, column=4, value=round(sum(salaries) / len(salaries))).number_format = _MONEY_FMT
        row += 1

    # Total row
    ws.cell(row=row, column=1, value="TOTAL").font = Font(bold=True)
    ws.cell(row=row, column=3, value=len(active)).font = Font(bold=True)

    for col in range(1, 5):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 20

    _save_xlsx_deterministic(wb, full_path)
    canaries.set_location(file_key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# Category 05: Tax (2 PDFs + 1 XLSX + 1 PDF)
# ═══════════════════════════════════════════════════════════════════════════


def _write_federal_returns_pdf(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
    year: int,
    file_key: str,
) -> None:
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/05_tax/federal_returns_fy{year}.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()

    tp = model.tax_provisions.get(year)
    pretax = _whole_dollars(tp.pre_tax_book_income) if tp else 0
    federal_tax = _whole_dollars(tp.federal_current) if tp else 0
    state_tax = _whole_dollars(tp.state_current) if tp else 0

    elements = [
        Paragraph(f"Federal Income Tax Return Summary — FY{year}", ts),
        Paragraph(f"{PARENT.name} (EIN: 93-1234567)", hs),
        Spacer(1, 12),
        Paragraph("<b>Filing Status:</b> C-Corporation (Form 1120)", bs),
        Paragraph(f"<b>Fiscal Year:</b> January 1 – December 31, {year}", bs),
        Paragraph(f"<b>Pre-Tax Income:</b> ${pretax:,}", bs),
        Paragraph(f"<b>Federal Tax Liability:</b> ${federal_tax:,}", bs),
        Paragraph(f"<b>State Tax Liability:</b> ${state_tax:,}", bs),
        Spacer(1, 12),
        Paragraph("Key Schedules", hs),
        Paragraph("Schedule M-1: Book-tax differences reconciliation filed.", bs),
        Paragraph("Schedule L: Balance sheet per books filed.", bs),
        Paragraph("Form 1125-A: Cost of Goods Sold attached.", bs),
    ]

    canary_loc = _build_simple_pdf(
        full_path, canary_code, f"Federal Returns FY{year}", elements,
    )
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_state_returns_xlsx(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    file_key = "tc12_state_returns_summary"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/05_tax/state_returns_summary.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    embed_canary_xlsx(wb, canary_code)
    ws = wb.active
    ws.title = "State Returns Summary"

    headers = ["State", "Entity", "FY2023 Tax", "FY2024 Tax", "FY2025 Tax (Est.)", "Filing Status"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    # State tax data (deterministic)
    state_data = [
        ("Oregon", "CI (Consolidated)", 285_000, 310_000, 345_000, "Filed"),
        ("Oregon", "PC", 180_000, 195_000, 215_000, "Filed"),
        ("Texas", "DS", 0, 0, 0, "No income tax — franchise tax filed"),
        ("Illinois", "AM", 62_000, 71_000, 82_000, "Filed"),
    ]

    for r, (state, entity, fy23, fy24, fy25, status) in enumerate(state_data, 2):
        ws.cell(row=r, column=1, value=state)
        ws.cell(row=r, column=2, value=entity)
        ws.cell(row=r, column=3, value=fy23).number_format = _MONEY_FMT
        ws.cell(row=r, column=4, value=fy24).number_format = _MONEY_FMT
        ws.cell(row=r, column=5, value=fy25).number_format = _MONEY_FMT
        ws.cell(row=r, column=6, value=status)

    for col in range(1, 7):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 22

    _save_xlsx_deterministic(wb, full_path)
    canaries.set_location(file_key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary_code, test_cases=[_TC])


def _write_tax_notices_pdf(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_tax_notices"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/05_tax/tax_notices.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()
    elements = [
        Paragraph("Tax Notices and Correspondence Summary", ts),
        Paragraph(PARENT.name, hs),
        Spacer(1, 12),
        Paragraph(
            "The Company has received no material tax notices, assessments, or "
            "audit correspondence from the IRS or any state taxing authority "
            "during the past three fiscal years (FY2023–FY2025).",
            bs,
        ),
        Spacer(1, 12),
        Paragraph("Minor Items", hs),
        Paragraph(
            "A routine CP2000 notice was received from the IRS in August 2024 "
            "regarding a minor Form 1099 discrepancy ($3,200). The matter was "
            "resolved with no additional tax owed.",
            bs,
        ),
    ]

    canary_loc = _build_simple_pdf(full_path, canary_code, "Tax Notices", elements)
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# Category 06: Operations (1 PDF + 3 XLSX)
# ═══════════════════════════════════════════════════════════════════════════


def _write_facility_leases_pdf(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    file_key = "tc12_facility_leases"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/06_operations/facility_leases.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()
    elements = [
        Paragraph("Facility Leases Summary", ts),
        Paragraph(PARENT.name, hs),
        Spacer(1, 12),
    ]

    data = [["Lease ID", "Lessee", "Lessor", "Commence", "Term (mo)", "Monthly Rent"]]
    for lease in sorted(model.leases, key=lambda ls: ls.lease_id):
        data.append([
            lease.lease_id,
            lease.lessee,
            lease.lessor,
            lease.commencement_date.isoformat(),
            str(lease.term_months),
            f"${_whole_dollars(lease.monthly_base_rent):,}",
        ])

    tbl = Table(data, colWidths=[0.7 * inch, 0.8 * inch, 1.4 * inch, 0.9 * inch, 0.7 * inch, 1.0 * inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.extend([tbl, Spacer(1, 12)])

    elements.append(Paragraph(
        f"<b>Total leases:</b> {len(model.leases)}", bs,
    ))

    canary_loc = _build_simple_pdf(full_path, canary_code, "Facility Leases", elements)
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_equipment_list_xlsx(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    manifest: Manifest,
) -> None:
    file_key = "tc12_equipment_list"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/06_operations/equipment_list.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    embed_canary_xlsx(wb, canary_code)
    ws = wb.active
    ws.title = "Equipment List"

    headers = ["Asset ID", "Description", "Entity", "Acquisition Date",
               "Cost", "Accum Depr", "Net Book Value"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    assets = sorted(model.assets, key=lambda a: a.asset_id)
    for r, asset in enumerate(assets, 2):
        # Compute accumulated depreciation through FY2025
        accum_dec = sum(
            asset.book_depr_for_year(y)
            for y in range(asset.acquisition_date.year, 2026)
        )
        cost = _whole_dollars(asset.cost)
        accum = _whole_dollars(accum_dec)
        nbv = cost - accum
        ws.cell(row=r, column=1, value=asset.asset_id)
        ws.cell(row=r, column=2, value=asset.description)
        ws.cell(row=r, column=3, value=asset.entity_code)
        ws.cell(row=r, column=4, value=asset.acquisition_date.isoformat())
        ws.cell(row=r, column=5, value=cost).number_format = _MONEY_FMT
        ws.cell(row=r, column=6, value=accum).number_format = _MONEY_FMT
        ws.cell(row=r, column=7, value=nbv).number_format = _MONEY_FMT

    for col in range(1, 8):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 18

    _save_xlsx_deterministic(wb, full_path)
    canaries.set_location(file_key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary_code, test_cases=[_TC])


def _write_customer_list_xlsx(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_customer_list_with_revenue"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/06_operations/customer_list_with_revenue.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    embed_canary_xlsx(wb, canary_code)
    ws = wb.active
    ws.title = "Customer Revenue"

    headers = ["Customer ID", "Customer Name", "Entity", "Revenue Share",
               "Est. Annual Revenue", "DSO"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    for r, cust in enumerate(AR_CUSTOMERS, 2):
        entity_target = Decimal(SUBSIDIARIES[cust.entity_code].revenue_target)
        annual_rev = (entity_target * cust.revenue_share).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP,
        )
        ws.cell(row=r, column=1, value=cust.id)
        ws.cell(row=r, column=2, value=cust.name)
        ws.cell(row=r, column=3, value=cust.entity_code)
        ws.cell(row=r, column=4, value=float(cust.revenue_share))
        ws.cell(row=r, column=5, value=_whole_dollars(annual_rev)).number_format = _MONEY_FMT
        ws.cell(row=r, column=6, value=cust.dso)

    for col in range(1, 7):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 22

    _save_xlsx_deterministic(wb, full_path)
    canaries.set_location(file_key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary_code, test_cases=[_TC])


def _write_vendor_list_xlsx(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_vendor_list"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/06_operations/vendor_list.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    embed_canary_xlsx(wb, canary_code)
    ws = wb.active
    ws.title = "Top Vendors"

    headers = ["Vendor Name", "Category", "Annual Spend (Est.)", "Payment Terms", "Entity"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    vendors = [
        ("Initech Suppliers Ltd", "Raw Materials", 12_000_000, "Net 45", "CI"),
        ("Pacific Steel Corp", "Raw Materials", 8_500_000, "Net 30", "PC"),
        ("Advanced Polymers Inc", "Raw Materials", 6_200_000, "Net 40", "AM"),
        ("Portland Power & Light", "Utilities", 2_800_000, "Net 30", "CI"),
        ("Northwest Logistics Co", "Freight & Shipping", 3_400_000, "Net 35", "DS"),
        ("TechServ Solutions", "IT Services", 1_200_000, "Net 30", "CI"),
        ("Pacific Equipment Leasing", "Equipment Leasing", 950_000, "Monthly", "PC"),
        ("Valley Packaging Supply", "Packaging Materials", 1_800_000, "Net 30", "DS"),
        ("Cascade Safety Products", "Safety & Compliance", 450_000, "Net 30", "CI"),
        ("Mitchell, Hartwell & Associates LLP", "Legal Services", 380_000, "Net 30", "CI"),
    ]

    for r, (name, cat, spend, terms, entity) in enumerate(vendors, 2):
        ws.cell(row=r, column=1, value=name)
        ws.cell(row=r, column=2, value=cat)
        ws.cell(row=r, column=3, value=spend).number_format = _MONEY_FMT
        ws.cell(row=r, column=4, value=terms)
        ws.cell(row=r, column=5, value=entity)

    for col in range(1, 6):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 24

    _save_xlsx_deterministic(wb, full_path)
    canaries.set_location(file_key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# Category 07: Technology (1 PDF + 1 XLSX + 1 DOCX)
# ═══════════════════════════════════════════════════════════════════════════


def _write_patent_portfolio_pdf(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    """RED FLAG: 2 patents expiring within 18 months of FY2025 year-end."""
    file_key = "tc12_patent_portfolio"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/07_technology/patent_portfolio.pdf"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    ts, hs, _, bs = _pdf_styles()
    elements = [
        Paragraph("Patent Portfolio Summary", ts),
        Paragraph(PARENT.name, hs),
        Spacer(1, 12),
    ]

    data = [["Patent No.", "Title", "Assignee", "Issue Date", "Expiration", "Status"]]
    for p in _PATENTS:
        data.append([
            p["number"],
            p["title"],
            p["assignee"],
            p["issue_date"].isoformat(),
            p["expiration_date"].isoformat(),
            p["status"],
        ])

    tbl = Table(data, colWidths=[0.9 * inch, 1.6 * inch, 1.4 * inch, 0.8 * inch, 0.8 * inch, 1.0 * inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.extend([tbl, Spacer(1, 16)])

    # Highlight expiring patents
    elements.append(Paragraph("Expiration Alert", hs))
    elements.append(Paragraph(
        "The following patents will expire within 18 months of FY2025 year-end "
        "(before June 30, 2027):",
        bs,
    ))
    for p in _PATENTS:
        if p["expiration_date"] <= datetime.date(2027, 6, 30):
            elements.append(Paragraph(
                f"&bull; <b>{p['number']}</b> — {p['title']} "
                f"(expires {p['expiration_date'].isoformat()})",
                bs,
            ))

    canary_loc = _build_simple_pdf(full_path, canary_code, "Patent Portfolio", elements)
    canaries.set_location(file_key, rel_path, canary_loc)
    manifest.register(rel_path, "pdf", canary=canary_code, test_cases=[_TC])


def _write_software_licenses_xlsx(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_software_licenses"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/07_technology/software_licenses.xlsx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    embed_canary_xlsx(wb, canary_code)
    ws = wb.active
    ws.title = "Software Licenses"

    headers = ["Software", "Vendor", "License Type", "Seats/Users",
               "Annual Cost", "Renewal Date"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header_row(ws, 1, len(headers))

    licenses = [
        ("Microsoft 365 E3", "Microsoft", "Subscription", 850, 306_000, "2026-01-01"),
        ("SAP Business One", "SAP", "Perpetual + Maintenance", 120, 185_000, "2026-03-15"),
        ("SolidWorks Professional", "Dassault Systèmes", "Subscription", 45, 135_000, "2025-12-31"),
        ("AutoCAD", "Autodesk", "Subscription", 30, 54_000, "2025-11-30"),
        ("Salesforce CRM", "Salesforce", "Subscription", 65, 117_000, "2026-06-30"),
        ("Oracle NetSuite", "Oracle", "Subscription", 40, 96_000, "2026-04-15"),
        ("Cisco Webex", "Cisco", "Subscription", 200, 24_000, "2025-09-30"),
        ("Fortinet Firewall", "Fortinet", "Hardware + License", 4, 32_000, "2026-02-28"),
        ("Veeam Backup", "Veeam", "Subscription", 1, 18_000, "2026-01-15"),
        ("Custom ERP Module", "In-house", "Proprietary", 850, 0, "N/A"),
    ]

    for r, (sw, vendor, ltype, seats, cost, renewal) in enumerate(licenses, 2):
        ws.cell(row=r, column=1, value=sw)
        ws.cell(row=r, column=2, value=vendor)
        ws.cell(row=r, column=3, value=ltype)
        ws.cell(row=r, column=4, value=seats)
        ws.cell(row=r, column=5, value=cost).number_format = _MONEY_FMT
        ws.cell(row=r, column=6, value=renewal)

    for col in range(1, 7):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 22

    _save_xlsx_deterministic(wb, full_path)
    canaries.set_location(file_key, rel_path, "Document properties → description")
    manifest.register(rel_path, "xlsx", canary=canary_code, test_cases=[_TC])


def _write_it_infrastructure_docx(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_it_infrastructure_overview"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_DR}/07_technology/it_infrastructure_overview.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()
    embed_canary_docx(doc, canary_code)

    title = doc.add_heading("IT Infrastructure Overview", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{PARENT.name}", style="Subtitle")
    doc.add_paragraph("Prepared: March 2025")
    doc.add_paragraph("")

    doc.add_heading("1. Network Architecture", level=1)
    doc.add_paragraph(
        "Cascade Industries operates a hub-and-spoke network with the primary "
        "data center at the Portland headquarters. Each subsidiary facility has "
        "a site-to-site VPN connection to the central data center."
    )

    doc.add_heading("2. Core Systems", level=1)
    doc.add_paragraph(
        "ERP: SAP Business One (on-premise) manages financials, inventory, and "
        "purchasing. A custom ERP module handles manufacturing scheduling for "
        "Precision Components. Oracle NetSuite is used by Distribution Services "
        "for order management and warehouse operations."
    )

    doc.add_heading("3. Security", level=1)
    doc.add_paragraph(
        "Perimeter security is managed by Fortinet firewalls with intrusion "
        "detection. Multi-factor authentication is enforced for all remote access. "
        "Annual penetration testing is performed by an external firm."
    )

    doc.add_heading("4. Disaster Recovery", level=1)
    doc.add_paragraph(
        "Full backups are performed nightly via Veeam Backup to an off-site "
        "location. RTO target: 4 hours. RPO target: 1 hour. DR tests are "
        "conducted semi-annually."
    )

    doc.add_heading("5. Headcount", level=1)
    doc.add_paragraph(
        "IT department: 12 FTEs plus 4 contractors. Managed by VP of Information "
        "Technology reporting to the CFO."
    )

    _save_docx_deterministic(doc, full_path)
    canaries.set_location(file_key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# DD Checklist (65 items, 9 categories)
# ═══════════════════════════════════════════════════════════════════════════


def _write_dd_checklist(
    output_dir: Path, canaries: CanaryRegistry, manifest: Manifest,
) -> None:
    file_key = "tc12_dd_checklist"
    canary_code = canaries.canary_for(file_key)
    rel_path = f"{_INPUT_DIR}/dd_checklist_standard.docx"
    full_path = output_dir / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()
    embed_canary_docx(doc, canary_code)

    title = doc.add_heading("Standard Due Diligence Checklist", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Acquisition of Cascade Industries, Inc.", style="Subtitle")
    doc.add_paragraph("65 Items | 9 Categories")
    doc.add_paragraph("")

    current_category = None
    item_num = 0
    for category, item_text in _DD_CHECKLIST:
        if category != current_category:
            doc.add_heading(category, level=1)
            current_category = category
        item_num += 1
        doc.add_paragraph(f"{item_num}. {item_text}")

    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: This checklist is a standard template. Not all items may have "
        "corresponding documents in the data room. Items without supporting "
        "documentation should be flagged as gaps for follow-up."
    )

    _save_docx_deterministic(doc, full_path)
    canaries.set_location(file_key, rel_path, "Core properties → comments")
    manifest.register(rel_path, "docx", canary=canary_code, test_cases=[_TC])


# ═══════════════════════════════════════════════════════════════════════════
# Prompt and Expected Behavior
# ═══════════════════════════════════════════════════════════════════════════


def _write_prompt(output_dir: Path) -> None:
    """Write the agent prompt for TC-12."""
    text = """\
# TC-12: Data Room Triage & Document Index

You have access to a deal data room for the potential acquisition of Cascade Industries.

## Input Files

- `data_room/` — 32 files organized into 7 categories:
  - `01_corporate/` — Articles, bylaws, board minutes, org chart
  - `02_financial/` — Audited financials, management financials, budget, debt schedule
  - `03_legal/` — Material contracts, pending litigation, IP assignments, insurance
  - `04_hr/` — Employee census, benefits, key employee agreements, org chart
  - `05_tax/` — Federal returns, state returns, tax notices
  - `06_operations/` — Facility leases, equipment list, customer/vendor lists
  - `07_technology/` — Patent portfolio, software licenses, IT overview
- `dd_checklist_standard.docx` — Standard due diligence checklist with 65 line items

## Tasks

1. Create a complete document index: for each file, provide:
   - File path
   - Document type/category
   - Date (if identifiable)
   - One-line summary of contents
   - Key data points or red flags noted
2. Cross-reference against the due diligence checklist:
   - Mark which checklist items have corresponding documents
   - Flag checklist items with NO corresponding document (gaps)
   - Note any documents in the data room not covered by the checklist
3. Identify red flags or items requiring immediate attention, such as:
   - Pending litigation
   - Contracts with change-of-control provisions
   - Key employee agreements with unusual terms
   - Missing critical documents
4. Prioritize the gaps: which missing documents are deal-critical vs. nice-to-have?

## Deliverables

Export:
- **Document index** as Excel (with columns: Path, Category, Date, Summary, Red Flags)
- **Gap analysis** as Excel (Checklist Item, Status, Priority, Notes)
- **Red flags summary** as Word memo
"""
    path = output_dir / f"test_cases/{_TC}/prompt.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _write_expected_behavior(output_dir: Path) -> None:
    """Write expected behavior notes for evaluators."""
    text = """\
# TC-12: Data Room Triage & Document Index — Expected Behavior

## Key Red Flags (agent must identify all 6)

| # | Red Flag | Location | Severity |
|---|----------|----------|----------|
| 1 | Pending product liability litigation — $2.5M exposure | `03_legal/pending_litigation_summary.docx` | High |
| 2 | CEO golden parachute — 3× salary ($975K) | `04_hr/key_employee_agreements/ceo_employment_agreement.pdf` | High |
| 3 | Acme change-of-control clause (~18% rev) | `03_legal/material_contracts/customer_agreement_acme.pdf` | High |
| 4 | 2 patents expiring within 18 months | `07_technology/patent_portfolio.pdf` | Medium |
| 5 | Incomplete IP assignments (2 founding employees) | `03_legal/ip_assignment_agreements.pdf` | High |
| 6 | Missing critical documents (see gap analysis below) | N/A — gaps in data room | High |

## Critical Document Gaps (missing from data room)

The agent should identify these as missing by cross-referencing the DD checklist:

| Category | Missing Item | Priority |
|----------|-------------|----------|
| Environmental | Phase I/II environmental site assessments | Deal-critical |
| Environmental | Environmental permits and compliance records | Deal-critical |
| Environmental | Hazardous waste disposal records | Deal-critical |
| Environmental | Environmental remediation obligations | Deal-critical |
| Environmental | Air and water quality permits | High |
| Environmental | Environmental insurance policies | Medium |
| Regulatory | Government permits and licenses | Deal-critical |
| Regulatory | Real property surveys and title reports | High |
| Legal | Insurance claim history (last 5 years) | High |

For a manufacturing company with chemical processes (Advanced Materials), the absence
of ALL environmental documentation is a critical deal issue.

## Additional Observations

- The agent should note that the CFO and CTO also have change-of-control provisions
  (2× salary each), though less unusual than the CEO's 3× provision.
- Board minutes 2025 explicitly references the expiring patents and CEO golden
  parachute ratification — cross-referencing these strengthens the analysis.
- The insurance policies summary notes that claim history is maintained elsewhere
  and not included — this should be flagged as a gap.

## Output Quality

- Document index should cover all 32 files with accurate summaries.
- Gap analysis should cover all 65 checklist items with clear status (Available /
  Partially Available / Not Available).
- Red flags memo should be professional, prioritized, and actionable.
- All output files must be valid xlsx/docx.
"""
    path = output_dir / f"test_cases/{_TC}/expected_behavior.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


# ═══════════════════════════════════════════════════════════════════════════
# Gold Standard
# ═══════════════════════════════════════════════════════════════════════════


@register_gold("TC-12")
def _tc12_gold(
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    **model_kwargs: Any,
) -> GoldStandard:
    """Build the TC-12 gold standard from the canonical model."""
    coc_contracts = contracts_with_change_of_control()

    # Patents expiring within 18 months of FY2025 year-end
    cutoff_18m = datetime.date(2027, 6, 30)
    expiring_patents = [p for p in _PATENTS if p["expiration_date"] <= cutoff_18m]

    return GoldStandard(
        test_case=_TC,
        expected_outputs={
            "output_files": {
                "document_index": {
                    "type": "xlsx",
                    "required_columns": [
                        "Path", "Category", "Date", "Summary", "Red Flags",
                    ],
                    "total_files_indexed": 32,
                },
                "gap_analysis": {
                    "type": "xlsx",
                    "required_columns": [
                        "Checklist Item", "Status", "Priority", "Notes",
                    ],
                    "total_checklist_items": 65,
                },
                "red_flags_memo": {
                    "type": "docx",
                    "required_sections": [
                        "Pending Litigation",
                        "Change-of-Control Provisions",
                        "Key Employee Agreements",
                        "IP and Patent Risks",
                        "Missing Critical Documents",
                    ],
                },
            },
            "red_flags": {
                "litigation": {
                    "title": LITIGATION[0].title,
                    "potential_exposure": _whole_dollars(LITIGATION[0].potential_exposure),
                    "accrued_liability": _whole_dollars(LITIGATION[0].accrued_liability),
                },
                "ceo_golden_parachute": {
                    "multiplier": float(KEY_PERSONNEL[0].change_of_control_multiplier),
                    "payout": _whole_dollars(
                        Decimal(KEY_PERSONNEL[0].base_salary)
                        * KEY_PERSONNEL[0].change_of_control_multiplier,
                    ),
                },
                "acme_change_of_control": {
                    "contract_id": coc_contracts[0].contract_id,
                    "annual_volume": _whole_dollars(coc_contracts[0].annual_volume),
                    "revenue_concentration": "~18%",
                },
                "expiring_patents": {
                    "count": len(expiring_patents),
                    "patent_numbers": [p["number"] for p in expiring_patents],
                },
                "incomplete_ip_assignments": {
                    "missing_persons": _MISSING_IP_ASSIGNMENTS,
                },
                "missing_documents": [
                    "Environmental site assessments",
                    "Environmental permits",
                    "Hazardous waste disposal records",
                    "Environmental remediation obligations",
                    "Government permits and licenses",
                    "Real property surveys and title reports",
                    "Insurance claim history",
                ],
            },
            "gap_analysis_summary": {
                "items_with_documents": 34,  # Roughly — checklist items with coverage
                "items_without_documents": 31,  # Remaining gaps
                "deal_critical_gaps": 7,
            },
        },
        canary_verification={
            "read_acme_contract": canaries.canary_for("tc12_customer_agreement_acme"),
            "read_litigation_summary": canaries.canary_for("tc12_pending_litigation_summary"),
            "read_ceo_agreement": canaries.canary_for("tc12_ceo_employment_agreement"),
            "read_patent_portfolio": canaries.canary_for("tc12_patent_portfolio"),
            "read_ip_assignments": canaries.canary_for("tc12_ip_assignment_agreements"),
            "read_dd_checklist": canaries.canary_for("tc12_dd_checklist"),
        },
        error_detection={},  # No planted ERR-xxx errors for TC-12
        scoring_hints={
            "correctness": (
                "All 32 files accurately indexed; all 65 checklist items "
                "correctly mapped to available/missing status"
            ),
            "completeness": (
                "All 6 red flags identified; all critical gaps flagged; "
                "cross-referencing between documents noted"
            ),
            "judgment": (
                "Missing environmental documentation flagged as deal-critical "
                "for a manufacturing company; gap prioritization is reasonable"
            ),
            "communication": (
                "Professional red flags memo; clear prioritization; "
                "actionable recommendations for each gap"
            ),
        },
    )


# ═══════════════════════════════════════════════════════════════════════════
# Public entry point
# ═══════════════════════════════════════════════════════════════════════════


def emit_tc12(
    model: CascadeModel,
    output_dir: Path,
    canaries: CanaryRegistry,
    errors: ErrorRegistry,
    manifest: Manifest,
) -> None:
    """Write all TC-12 files to *output_dir*."""
    # 01_corporate
    _write_articles_of_incorporation(output_dir, canaries, manifest)
    _write_bylaws(output_dir, canaries, manifest)
    _write_board_minutes(output_dir, canaries, manifest, 2024, "tc12_board_minutes_2024")
    _write_board_minutes(output_dir, canaries, manifest, 2025, "tc12_board_minutes_2025")
    _write_org_chart_pdf(output_dir, canaries, manifest)

    # 02_financial
    _write_audited_financials_pdf(
        model, output_dir, canaries, manifest, 2023, "tc12_audited_financials_fy2023",
    )
    _write_audited_financials_pdf(
        model, output_dir, canaries, manifest, 2024, "tc12_audited_financials_fy2024",
    )
    _write_management_financials_xlsx(model, output_dir, canaries, manifest)
    _write_budget_xlsx(model, output_dir, canaries, manifest)
    _write_debt_schedule_xlsx(output_dir, canaries, manifest)

    # 03_legal
    _write_customer_agreement_acme(output_dir, canaries, manifest)
    _write_customer_agreement_globex(output_dir, canaries, manifest)
    _write_supplier_agreement_initech(output_dir, canaries, manifest)
    _write_pending_litigation_docx(output_dir, canaries, manifest)
    _write_ip_assignments_pdf(output_dir, canaries, manifest)
    _write_insurance_policies_pdf(output_dir, canaries, manifest)

    # 04_hr
    _write_employee_census_xlsx(model, output_dir, canaries, errors, manifest)
    _write_benefits_summary_pdf(output_dir, canaries, manifest)
    _write_key_employee_agreement_pdf(
        output_dir, canaries, manifest, 0, "tc12_ceo_employment_agreement",
        "ceo_employment_agreement.pdf",
    )
    _write_key_employee_agreement_pdf(
        output_dir, canaries, manifest, 1, "tc12_cfo_employment_agreement",
        "cfo_employment_agreement.pdf",
    )
    _write_key_employee_agreement_pdf(
        output_dir, canaries, manifest, 2, "tc12_cto_employment_agreement",
        "cto_employment_agreement.pdf",
    )
    _write_org_chart_detailed_xlsx(model, output_dir, canaries, manifest)

    # 05_tax
    _write_federal_returns_pdf(
        model, output_dir, canaries, manifest, 2023, "tc12_federal_returns_fy2023",
    )
    _write_federal_returns_pdf(
        model, output_dir, canaries, manifest, 2024, "tc12_federal_returns_fy2024",
    )
    _write_state_returns_xlsx(model, output_dir, canaries, manifest)
    _write_tax_notices_pdf(output_dir, canaries, manifest)

    # 06_operations
    _write_facility_leases_pdf(model, output_dir, canaries, manifest)
    _write_equipment_list_xlsx(model, output_dir, canaries, manifest)
    _write_customer_list_xlsx(output_dir, canaries, manifest)
    _write_vendor_list_xlsx(output_dir, canaries, manifest)

    # 07_technology
    _write_patent_portfolio_pdf(output_dir, canaries, manifest)
    _write_software_licenses_xlsx(output_dir, canaries, manifest)
    _write_it_infrastructure_docx(output_dir, canaries, manifest)

    # DD checklist
    _write_dd_checklist(output_dir, canaries, manifest)

    # Prompt and expected behavior
    _write_prompt(output_dir)
    _write_expected_behavior(output_dir)
