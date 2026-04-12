"""Controlled document noise profiles for formatter outputs.

Adds deterministic, cosmetic noise to generated files — date format
variation, header perturbation, whitespace artifacts, encoding quirks —
without altering model facts, canary values, or planted errors.

All noise is seeded from ScenarioContext so it is byte-identical across
reruns.  Each file gets a dedicated sub-seed via the namespace
``noise:{tc_id}:{file_key}``.

Noise invariants (from docs/formatter-family-matrix.md):
1. Canary preservation — canaries remain findable after noise.
2. Model fact preservation — no semantic value changes.
3. Planted error preservation — errors remain detectable at registered locations.
4. Determinism — identical output across runs.
"""

from __future__ import annotations

import random
from dataclasses import dataclass, field
from typing import TYPE_CHECKING

from generator.scenario_context import ScenarioContext

if TYPE_CHECKING:
    from openpyxl import Workbook
    from openpyxl.worksheet.worksheet import Worksheet


# ---------------------------------------------------------------------------
# Exclusion zones
# ---------------------------------------------------------------------------

@dataclass
class ExclusionZone:
    """Cells/rows/paragraphs that noise must not touch.

    For XLSX: *cells* is a set of (sheet_title, row, col) tuples.
    For CSV: *rows* is a set of 0-based row indices.
    For DOCX: *paragraphs* is a set of 0-based paragraph indices.
    """

    cells: set[tuple[str, int, int]] = field(default_factory=set)
    rows: set[int] = field(default_factory=set)
    paragraphs: set[int] = field(default_factory=set)


# ---------------------------------------------------------------------------
# XLSX noise helpers
# ---------------------------------------------------------------------------

# Alternate date format strings for openpyxl cells that hold dates.
_DATE_FORMATS = [
    "MM/DD/YYYY",
    "M/D/YYYY",
    "YYYY-MM-DD",
    "DD-MMM-YY",
    "M/D/YY",
]

# Alternate header spellings keyed by normalised lowercase original.
_HEADER_VARIANTS: dict[str, list[str]] = {
    "account #": ["Acct #", "Account No.", "ACCOUNT #"],
    "account name": ["Acct Name", "account_name", "ACCOUNT NAME"],
    "account number": ["Acct No", "Account #", "ACCT NUMBER"],
    "debit": ["Dr", "DEBIT", "Debit "],
    "credit": ["Cr", "CREDIT", "Credit "],
    "net balance": ["Net Bal", "NET BALANCE", "Balance (Net)"],
    "description": ["Desc", "DESCRIPTION", "Desc."],
    "amount": ["Amt", "AMOUNT", "Amount "],
    "date": ["Dt", "DATE", "Date "],
    "reference": ["Ref", "REF", "Ref."],
    "period": ["Per", "PERIOD", "Period "],
    "entity": ["Entity", "ENTITY", "Co"],
    "vendor": ["Vendor Name", "VENDOR", "Vendor"],
    "invoice": ["Inv #", "INVOICE", "Invoice No."],
}


def apply_xlsx_noise(
    wb: Workbook,
    rng: random.Random,
    exclusions: ExclusionZone | None = None,
) -> None:
    """Apply controlled noise to an openpyxl Workbook **in place**.

    Noise dimensions (from formatter-family-matrix):
    - Header perturbation: column name casing/abbreviation drift
    - Trailing whitespace: pad some cell string values
    - Date format variation: change number_format on date-formatted cells

    The workbook's document properties (where canaries live) are never touched.

    Parameters
    ----------
    wb : the Workbook to mutate.
    rng : a seeded stdlib Random instance for determinism.
    exclusions : cells to skip (canary cells, planted-error cells).
    """
    if exclusions is None:
        exclusions = ExclusionZone()

    for ws in wb.worksheets:
        _apply_xlsx_header_noise(ws, rng, exclusions)
        _apply_xlsx_trailing_whitespace(ws, rng, exclusions)
        _apply_xlsx_date_format_noise(ws, rng, exclusions)


def _is_excluded(sheet_title: str, row: int, col: int, excl: ExclusionZone) -> bool:
    """Check whether a cell is in the exclusion zone."""
    return (sheet_title, row, col) in excl.cells


def _apply_xlsx_header_noise(
    ws: Worksheet, rng: random.Random, excl: ExclusionZone,
) -> None:
    """Perturb header cell text with casing/abbreviation variants.

    Scans the first 5 rows for cells whose lowercase value matches a known
    header.  Each matching cell has a 40% chance of being replaced with an
    alternate spelling.
    """
    for row in range(1, 6):
        for col in range(1, ws.max_column + 1):
            if _is_excluded(ws.title, row, col, excl):
                continue
            cell = ws.cell(row=row, column=col)
            if cell.value is None or not isinstance(cell.value, str):
                continue
            key = cell.value.strip().lower()
            if key in _HEADER_VARIANTS and rng.random() < 0.40:
                cell.value = rng.choice(_HEADER_VARIANTS[key])


def _apply_xlsx_trailing_whitespace(
    ws: Worksheet, rng: random.Random, excl: ExclusionZone,
) -> None:
    """Add trailing spaces to ~15% of non-empty string cells.

    Skips header rows (1–5), merged cells, and exclusion zones.
    """
    for row in range(6, ws.max_row + 1):
        for col in range(1, ws.max_column + 1):
            if _is_excluded(ws.title, row, col, excl):
                continue
            cell = ws.cell(row=row, column=col)
            if cell.value is None or not isinstance(cell.value, str):
                continue
            # Skip if cell is part of a merged range
            if _cell_is_merged(ws, row, col):
                continue
            if rng.random() < 0.15:
                spaces = rng.randint(1, 3)
                cell.value = cell.value + " " * spaces


def _apply_xlsx_date_format_noise(
    ws: Worksheet, rng: random.Random, excl: ExclusionZone,
) -> None:
    """Vary the number_format on date-formatted cells.

    Detects date cells by checking if the existing number_format contains
    common date tokens (mm, dd, yy).  Each matching cell has a 30% chance
    of getting a different date format.
    """
    _date_tokens = {"mm", "dd", "yy", "yyyy", "m/d", "d/m"}
    for row in range(1, ws.max_row + 1):
        for col in range(1, ws.max_column + 1):
            if _is_excluded(ws.title, row, col, excl):
                continue
            cell = ws.cell(row=row, column=col)
            fmt = (cell.number_format or "").lower()
            if any(tok in fmt for tok in _date_tokens):
                if rng.random() < 0.30:
                    cell.number_format = rng.choice(_DATE_FORMATS)


def _cell_is_merged(ws: Worksheet, row: int, col: int) -> bool:
    """Return True if (row, col) is part of a merged cell range."""
    for merged_range in ws.merged_cells.ranges:
        if (row, col) != (merged_range.min_row, merged_range.min_col):
            if merged_range.min_row <= row <= merged_range.max_row:
                if merged_range.min_col <= col <= merged_range.max_col:
                    return True
    return False


# ---------------------------------------------------------------------------
# CSV noise helpers
# ---------------------------------------------------------------------------

def apply_csv_noise(
    lines: list[str],
    rng: random.Random,
    exclusions: ExclusionZone | None = None,
    *,
    add_bom: bool | None = None,
) -> list[str]:
    """Apply controlled noise to CSV lines and return the modified list.

    Noise dimensions (from formatter-family-matrix):
    - Trailing whitespace in field values
    - Quoting variation (wrap some unquoted fields in quotes)
    - Optional UTF-8 BOM prefix

    The first line is preserved verbatim if it starts with ``# CANARY:``
    (canary comment line).

    Parameters
    ----------
    lines : list of CSV lines (including line terminators).
    rng : seeded stdlib Random.
    exclusions : row indices to skip.
    add_bom : if True, prepend BOM to first line.  If None, 25% chance.

    Returns
    -------
    Modified list of lines.
    """
    if exclusions is None:
        exclusions = ExclusionZone()

    result: list[str] = []

    for i, line in enumerate(lines):
        # Never touch canary comment line
        if i == 0 and line.startswith("# CANARY:"):
            result.append(line)
            continue

        if i in exclusions.rows:
            result.append(line)
            continue

        line = _csv_trailing_whitespace(line, rng)
        line = _csv_quoting_variation(line, rng)
        result.append(line)

    # BOM insertion
    if add_bom is None:
        add_bom = rng.random() < 0.25
    if add_bom and result:
        # Prepend BOM to first line content
        if not result[0].startswith("\ufeff"):
            result[0] = "\ufeff" + result[0]

    return result


def _csv_trailing_whitespace(line: str, rng: random.Random) -> str:
    """Add trailing spaces to ~10% of CSV field values."""
    if rng.random() >= 0.10:
        return line
    # Strip the line terminator, manipulate, then re-add
    stripped = line.rstrip("\r\n")
    terminator = line[len(stripped):]
    fields = stripped.split(",")
    if not fields:
        return line
    idx = rng.randrange(len(fields))
    # Don't add whitespace inside quoted fields
    if fields[idx].startswith('"'):
        return line
    fields[idx] = fields[idx] + " "
    return ",".join(fields) + terminator


def _csv_quoting_variation(line: str, rng: random.Random) -> str:
    """Wrap ~8% of unquoted string fields in unnecessary quotes."""
    if rng.random() >= 0.08:
        return line
    stripped = line.rstrip("\r\n")
    terminator = line[len(stripped):]
    fields = stripped.split(",")
    if not fields:
        return line
    # Pick a random field that looks like a string (not purely numeric)
    candidates = [
        j for j, f in enumerate(fields)
        if f and not f.startswith('"') and not _is_numeric(f.strip())
    ]
    if not candidates:
        return line
    idx = rng.choice(candidates)
    fields[idx] = f'"{fields[idx]}"'
    return ",".join(fields) + terminator


def _is_numeric(s: str) -> bool:
    """Return True if *s* looks like a number."""
    try:
        float(s.replace(",", ""))
        return True
    except ValueError:
        return False


# ---------------------------------------------------------------------------
# DOCX noise helpers
# ---------------------------------------------------------------------------

def apply_docx_noise(
    doc: "object",  # docx.Document — avoid import at module level
    rng: random.Random,
    exclusions: ExclusionZone | None = None,
) -> None:
    """Apply controlled noise to a python-docx Document **in place**.

    Noise dimensions (from formatter-family-matrix):
    - Font size jitter: vary font size ±0.5pt on body paragraph runs
    - Paragraph spacing: randomize before/after spacing slightly
    - Metadata clutter: set author/company to plausible values

    Core properties (where canaries live) are not touched — only the
    ``author`` and ``company`` fields in core_properties, which are distinct
    from the ``comments`` field used for canaries.

    Parameters
    ----------
    doc : a python-docx Document instance.
    rng : seeded stdlib Random.
    exclusions : paragraph indices to skip.
    """
    from docx.shared import Pt

    if exclusions is None:
        exclusions = ExclusionZone()

    _apply_docx_font_jitter(doc, rng, exclusions, Pt)
    _apply_docx_paragraph_spacing(doc, rng, exclusions, Pt)
    _apply_docx_metadata_clutter(doc, rng)


def _apply_docx_font_jitter(
    doc: "object", rng: random.Random, excl: ExclusionZone, Pt: type,
) -> None:
    """Vary font size ±0.5pt on ~20% of body paragraph runs.

    Operates on Run.font.size, not Run.text, so canary text is preserved.
    Skips heading paragraphs (style.name starts with 'Heading').
    """
    for i, para in enumerate(doc.paragraphs):  # type: ignore[attr-defined]
        if i in excl.paragraphs:
            continue
        style_name = (para.style.name or "") if para.style else ""
        if style_name.startswith("Heading"):
            continue
        for run in para.runs:
            if run.font.size is not None and rng.random() < 0.20:
                delta = rng.choice([-0.5, 0.5])
                new_size = max(6, run.font.size.pt + delta)
                run.font.size = Pt(new_size)


def _apply_docx_paragraph_spacing(
    doc: "object", rng: random.Random, excl: ExclusionZone, Pt: type,
) -> None:
    """Randomize before/after spacing on ~15% of paragraphs."""
    for i, para in enumerate(doc.paragraphs):  # type: ignore[attr-defined]
        if i in excl.paragraphs:
            continue
        if rng.random() < 0.15:
            pf = para.paragraph_format
            if pf.space_before is not None:
                delta = Pt(rng.choice([-1, 0.5, 1, 1.5]))
                new_val = max(0, pf.space_before + delta)
                pf.space_before = Pt(new_val / 12700)  # EMU to pt
            if pf.space_after is not None:
                delta = Pt(rng.choice([-1, 0.5, 1, 1.5]))
                new_val = max(0, pf.space_after + delta)
                pf.space_after = Pt(new_val / 12700)


_PLAUSIBLE_AUTHORS = [
    "John Smith", "S. Martinez", "Client Export", "scanner@cascade.local",
    "Admin", "J. Chen", "M. Williams", "Cascade Industries",
]

_PLAUSIBLE_COMPANIES = [
    "Cascade Industries, Inc.", "Cascade Industries",
    "CASCADE INDUSTRIES INC", "Cascade Ind.",
]


def _apply_docx_metadata_clutter(doc: "object", rng: random.Random) -> None:
    """Set author and company to plausible but varied values.

    Never touches the ``comments`` field — that's where canaries live.
    """
    props = doc.core_properties  # type: ignore[attr-defined]
    props.author = rng.choice(_PLAUSIBLE_AUTHORS)
    # company is not a standard core property in python-docx;
    # use category as a proxy for organisational metadata clutter
    props.category = rng.choice(_PLAUSIBLE_COMPANIES)


# ---------------------------------------------------------------------------
# Integration helper
# ---------------------------------------------------------------------------

def make_noise_rng(ctx: ScenarioContext, tc_id: str, file_key: str) -> random.Random:
    """Create a seeded stdlib Random for noise on a specific file.

    Uses the ScenarioContext namespace ``noise:{tc_id}:{file_key}`` so
    noise seeds are isolated from data-generation seeds.
    """
    seed = ctx.child_seed(f"noise:{tc_id}:{file_key}")
    return random.Random(seed)
