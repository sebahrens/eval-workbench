"""One-shot script to create templates/workpaper_memo_template.docx.

Run once, commit the .docx, then discard or keep the script for provenance.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def build_template() -> Document:
    doc = Document()

    # -- Page setup ----------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)

    # -- Define styles -------------------------------------------------------
    style = doc.styles

    # Normal text
    normal = style["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Heading 1 — section headers
    h1 = style["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)  # dark navy
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2 — sub-sections
    h2 = style["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)

    # -- Header --------------------------------------------------------------
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("CASCADE INDUSTRIES — CONFIDENTIAL")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.bold = True

    # Second header line with merge fields
    hp2 = header.add_paragraph()
    hp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = hp2.add_run("Engagement: {{ENGAGEMENT_NAME}}  |  Period: {{FISCAL_YEAR}}")
    run2.font.name = "Calibri"
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # -- Footer --------------------------------------------------------------
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = fp.add_run("{{FIRM_NAME}}  •  Work Product — Do Not Distribute  •  Page ")
    run_f.font.name = "Calibri"
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Add page number field
    fld_char1 = doc.element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    instr = doc.element.makeelement(qn("w:instrText"), {})
    instr.text = " PAGE "
    fld_char2 = doc.element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run_pg = fp.add_run()
    run_pg._r.append(fld_char1)
    run_pg._r.append(instr)
    run_pg._r.append(fld_char2)

    # -- Title block ---------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(4)
    run_t = title.add_run("AUDIT WORKPAPER MEMO")
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(20)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    run_s = subtitle.add_run("{{WORKPAPER_TITLE}}")
    run_s.font.name = "Calibri"
    run_s.font.size = Pt(14)
    run_s.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

    # Meta table
    meta = doc.add_table(rows=5, cols=2, style="Table Grid")
    meta.autofit = True
    labels = [
        ("Client:", "{{CLIENT_NAME}}"),
        ("Engagement:", "{{ENGAGEMENT_NAME}}"),
        ("Fiscal Year:", "{{FISCAL_YEAR}}"),
        ("Prepared By:", "{{PREPARED_BY}}"),
        ("Reviewed By:", "{{REVIEWED_BY}}"),
    ]
    for i, (label, value) in enumerate(labels):
        cell_l = meta.cell(i, 0)
        cell_l.text = label
        for p in cell_l.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(10)
        cell_v = meta.cell(i, 1)
        cell_v.text = value
        for p in cell_v.paragraphs:
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # -- Section: Objective --------------------------------------------------
    doc.add_heading("1. OBJECTIVE", level=1)
    doc.add_paragraph(
        "{{OBJECTIVE_TEXT — State the audit objective for this workpaper area. "
        "Example: To obtain sufficient appropriate audit evidence regarding the "
        "existence, completeness, valuation, and presentation of accounts receivable "
        "as of the balance sheet date.}}"
    )

    # -- Section: Scope ------------------------------------------------------
    doc.add_heading("2. SCOPE", level=1)
    doc.add_paragraph(
        "{{SCOPE_TEXT — Describe the scope of procedures, including the period under "
        "audit, the population tested, and any materiality thresholds applied.}}"
    )

    # -- Section: Procedures Performed ---------------------------------------
    doc.add_heading("3. PROCEDURES PERFORMED", level=1)
    doc.add_paragraph(
        "{{PROCEDURES_TEXT — Document each audit procedure performed. Use numbered "
        "sub-sections or a bulleted list. Reference workpaper exhibits where applicable.}}"
    )

    doc.add_heading("3.1 Procedure Detail", level=2)
    doc.add_paragraph(
        "{{PROCEDURE_DETAIL — For each major procedure, describe: (a) nature of the "
        "procedure, (b) extent of testing, (c) source documents examined, and "
        "(d) results of testing.}}"
    )

    # -- Section: Findings ---------------------------------------------------
    doc.add_heading("4. FINDINGS", level=1)
    doc.add_paragraph(
        "{{FINDINGS_TEXT — Summarize key findings from the procedures performed. "
        "Reference specific data points, exceptions noted, and their resolution. "
        "Include quantitative support where applicable.}}"
    )

    doc.add_heading("4.1 Exceptions Noted", level=2)
    doc.add_paragraph(
        "{{EXCEPTIONS_TEXT — Detail any exceptions or anomalies identified, "
        "including their magnitude, cause, and disposition.}}"
    )

    # -- Section: Conclusion -------------------------------------------------
    doc.add_heading("5. CONCLUSION", level=1)
    doc.add_paragraph(
        "{{CONCLUSION_TEXT — State the audit conclusion for this workpaper area. "
        "Example: Based on the procedures performed, we obtained sufficient "
        "appropriate audit evidence that accounts receivable is fairly stated "
        "in all material respects as of December 31, 2025.}}"
    )

    # -- Signature block -----------------------------------------------------
    doc.add_paragraph()  # spacer
    sig_table = doc.add_table(rows=2, cols=2, style="Table Grid")
    cells = [
        ("Prepared By:", "{{PREPARED_BY}}"),
        ("Date:", "{{PREPARATION_DATE}}"),
        ("Reviewed By:", "{{REVIEWED_BY}}"),
        ("Date:", "{{REVIEW_DATE}}"),
    ]
    for row_idx in range(2):
        for col_idx in range(2):
            cell = sig_table.cell(row_idx, col_idx)
            label, value = cells[row_idx * 2 + col_idx]
            p = cell.paragraphs[0]
            r_label = p.add_run(f"{label} ")
            r_label.font.bold = True
            r_label.font.name = "Calibri"
            r_label.font.size = Pt(10)
            r_val = p.add_run(value)
            r_val.font.name = "Calibri"
            r_val.font.size = Pt(10)

    return doc


if __name__ == "__main__":
    import pathlib

    out = pathlib.Path(__file__).resolve().parent.parent / "templates" / "workpaper_memo_template.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = build_template()
    doc.save(str(out))
    print(f"Wrote {out}")
