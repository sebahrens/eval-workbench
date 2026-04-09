"""One-shot script to create templates/engagement_letter_template.docx.

Run once, commit the .docx, then discard or keep the script for provenance.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def build_template() -> Document:
    doc = Document()

    # -- Page setup ----------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)

    # -- Define styles -------------------------------------------------------
    style = doc.styles

    normal = style["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1 = style["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    h2 = style["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)

    # -- Firm letterhead -----------------------------------------------------
    letterhead = doc.add_paragraph()
    letterhead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    letterhead.paragraph_format.space_before = Pt(12)
    letterhead.paragraph_format.space_after = Pt(0)
    run_firm = letterhead.add_run("Mitchell & Partners LLP")
    run_firm.font.name = "Calibri"
    run_firm.font.size = Pt(18)
    run_firm.font.bold = True
    run_firm.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

    addr = doc.add_paragraph()
    addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    addr.paragraph_format.space_before = Pt(0)
    addr.paragraph_format.space_after = Pt(2)
    run_addr = addr.add_run(
        "One Financial Plaza, Suite 4200  |  Chicago, IL 60601  |  (312) 555-0100"
    )
    run_addr.font.name = "Calibri"
    run_addr.font.size = Pt(9)
    run_addr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Horizontal rule (border-bottom on paragraph)
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(6)
    hr.paragraph_format.space_after = Pt(12)
    pPr = hr._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from lxml import etree

    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "002B5C")

    # -- Date and addressee --------------------------------------------------
    date_para = doc.add_paragraph("<<start_date>>")
    date_para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph("<<client_name>>")
    doc.add_paragraph("Cascade Industries, Inc.")
    doc.add_paragraph()  # blank line

    # -- Salutation ----------------------------------------------------------
    doc.add_paragraph("Dear <<client_name>>,")
    doc.add_paragraph()

    # -- Body: Engagement scope ----------------------------------------------
    doc.add_heading("Engagement Scope", level=1)
    doc.add_paragraph(
        "We are pleased to confirm our understanding of the services we will provide "
        "to Cascade Industries, Inc. and its subsidiaries (collectively, the \"Company\"). "
        "This letter confirms the terms of our engagement to provide the following "
        "professional services:"
    )
    doc.add_paragraph("<<engagement_scope>>")

    # -- Body: Fees and billing ----------------------------------------------
    doc.add_heading("Fees and Billing", level=1)
    doc.add_paragraph(
        "Our fees for the services described above will be <<fee_amount>>. "
        "This fee estimate is based on the anticipated cooperation of your personnel "
        "and the assumption that unexpected circumstances will not be encountered "
        "during the engagement."
    )
    doc.add_paragraph(
        "Payment terms: <<payment_terms>>. Invoices will be submitted monthly "
        "as services are rendered. Any additional services beyond the scope of this "
        "engagement will be billed at our standard hourly rates and are subject "
        "to prior written approval."
    )

    # -- Body: Engagement period ---------------------------------------------
    doc.add_heading("Engagement Period", level=1)
    doc.add_paragraph(
        "This engagement will commence on <<start_date>> and is expected to conclude "
        "upon completion of the agreed-upon services and delivery of all final "
        "reports and deliverables."
    )

    # -- Body: Responsibilities ----------------------------------------------
    doc.add_heading("Responsibilities", level=1)

    doc.add_heading("Our Responsibilities", level=2)
    doc.add_paragraph(
        "We will perform the engagement in accordance with professional standards "
        "applicable to the services described herein. We will provide timely "
        "communication of any matters that come to our attention during the "
        "engagement that may be significant to your management."
    )

    doc.add_heading("Management Responsibilities", level=2)
    doc.add_paragraph(
        "Management is responsible for providing us with access to all information "
        "and personnel necessary to perform the engagement, including all financial "
        "records, supporting documentation, and other information requested. "
        "Management is also responsible for the accuracy and completeness of the "
        "information provided."
    )

    # -- Body: Confidentiality -----------------------------------------------
    doc.add_heading("Confidentiality", level=1)
    doc.add_paragraph(
        "We will maintain the confidentiality of all information obtained during "
        "the course of the engagement, except as required by law, regulation, "
        "or professional standards. Our workpapers and files are the property of "
        "Mitchell & Partners LLP."
    )

    # -- Body: Limitation of liability ---------------------------------------
    doc.add_heading("Limitation of Liability", level=1)
    doc.add_paragraph(
        "To the fullest extent permitted by applicable law, the aggregate liability "
        "of Mitchell & Partners LLP and its partners, principals, and employees "
        "for any claims arising out of or related to this engagement shall not "
        "exceed the total fees paid under this engagement letter."
    )

    # -- Acceptance block ----------------------------------------------------
    doc.add_paragraph()
    doc.add_paragraph(
        "If the above terms are acceptable, please sign and return a copy of this "
        "letter. We appreciate the opportunity to serve Cascade Industries and "
        "look forward to working with you."
    )
    doc.add_paragraph()

    # Signature lines
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    sig = doc.add_paragraph("<<partner_name>>")
    sig.paragraph_format.space_after = Pt(0)
    doc.add_paragraph("Engagement Partner")
    doc.add_paragraph("Mitchell & Partners LLP")
    doc.add_paragraph()

    # Acceptance section
    hr2 = doc.add_paragraph()
    hr2.paragraph_format.space_before = Pt(18)
    hr2.paragraph_format.space_after = Pt(12)
    pPr2 = hr2._p.get_or_add_pPr()
    pBdr2 = etree.SubElement(pPr2, qn("w:pBdr"))
    bottom2 = etree.SubElement(pBdr2, qn("w:bottom"))
    bottom2.set(qn("w:val"), "single")
    bottom2.set(qn("w:sz"), "6")
    bottom2.set(qn("w:space"), "1")
    bottom2.set(qn("w:color"), "002B5C")

    accept = doc.add_paragraph()
    accept.paragraph_format.space_after = Pt(2)
    run_a = accept.add_run("ACCEPTED AND AGREED")
    run_a.font.bold = True
    run_a.font.size = Pt(12)
    run_a.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

    doc.add_paragraph()
    doc.add_paragraph("Signature: _________________________________")
    doc.add_paragraph()
    doc.add_paragraph("Name: _________________________________")
    doc.add_paragraph()
    doc.add_paragraph("Title: _________________________________")
    doc.add_paragraph()
    doc.add_paragraph("Date: _________________________________")

    return doc


if __name__ == "__main__":
    import pathlib

    out = (
        pathlib.Path(__file__).resolve().parent.parent
        / "templates"
        / "engagement_letter_template.docx"
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = build_template()
    doc.save(str(out))
    print(f"Wrote {out}")
